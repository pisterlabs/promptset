import sounddevice as sd
import soundfile as sf
import numpy as np
import openai
import whisper
import os
import requests
import re
import sys
from colorama import Fore, Style, init
from pydub import AudioSegment
from pydub.playback import play
from docx import Document

# import the required modules for text to speech conversion
from gtts import gTTS

init()

# Get the directory of the currently running script
script_directory = os.path.dirname(os.path.abspath(__file__))

def open_file(filepath):
    with open(filepath, 'r', encoding='utf-8') as infile:
        return infile.read()

# Construct the absolute paths to the files in the same directory as the script
api_key_path = os.path.join(script_directory, '_openai_apikey.txt')
chatbot_path = os.path.join(script_directory, '_chatbot.txt')

# Open the files using these paths
try:
    with open(api_key_path, 'r') as api_key_file:
        api_key = api_key_file.read().strip()
except FileNotFoundError:
    print("One or both files not found.")

if __name__ == "__main__":
    # Check if there are enough arguments provided
    if len(sys.argv) < 1:
        # Usage: voicegpt <docx_file_path>
        sys.exit(1)

docx_file_path = sys.argv[1]

conversation = []
with open(chatbot_path, 'r') as chatbot_file:
        chatbot = chatbot_file.read().strip()

def chatgpt(api_key, conversation, chatbot, user_input, temperature=0, frequency_penalty=0.2, presence_penalty=0):
    openai.api_key = api_key
    conversation.append({"role": "user","content": user_input})
    messages_input = conversation.copy()
    prompt = [{"role": "system", "content": chatbot}]
    messages_input.insert(0, prompt[0])
    completion = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=temperature,
        frequency_penalty=frequency_penalty,
        presence_penalty=presence_penalty,
        messages=messages_input)
    chat_response = completion['choices'][0]['message']['content']
    conversation.append({"role": "assistant", "content": chat_response})
    return chat_response

def text_to_speech(myText, voice_id):
    # passing the text and language to the engine
    # slow=False tells the module that the converted audio should have high speed
    audioFile = gTTS(text=myText, lang=voice_id, slow=False)

    # saving the converted audio to an mp3 file
    audioFile.save("output.mp3")

    # play the response as an audio file
    os.system("mpg321 output.mp3")

def print_colored(agent, text):
    agent_colors = {
        "Matilda:": Fore.YELLOW,
    }
    color = agent_colors.get(agent, "")
    print(color + f"{agent}: {text}" + Style.RESET_ALL, end="")

voice_id = 'EXAVITQu4vr4xnSDxMaL'

def record_and_transcribe(duration=5, fs=44100):
    print('Recording...')
    myrecording = sd.rec(int(duration * fs), samplerate=fs, channels=1)
    sd.wait()
    print('Recording complete.')
    filename = 'myrecording.wav'
    sf.write(filename, myrecording, fs)
    with open(filename, "rb") as file:
        openai.api_key = api_key
        result = openai.Audio.transcribe("whisper-1", file)
    transcription = result['text']
    return transcription

def record_and_transcript_local_whisper(duration=5, fs=44100):
    print('Recording...')
    myrecording = sd.rec(int(duration * fs), samplerate=fs, channels=1)
    sd.wait()
    print('Recording complete.')
    filename = 'myrecording.wav'
    sf.write(filename, myrecording, fs)

    model = whisper.load_model("base")
    result = model.transcribe(filename)
    
    return result["text"]

def save_as_docx(text, filename):
    if os.path.isfile(filename):
        # If the file exists, open it
        doc = Document(filename)
    else:
        # If the file doesn't exist, create a new document
        doc = Document()
    
    # Add the new paragraph with the provided text
    doc.add_paragraph(text)
    
    # Save the document
    doc.save(filename)

while True:
    print('running')
    # user_message = record_and_transcribe()
    user_message = record_and_transcript_local_whisper();
    response = chatgpt(api_key, conversation, chatbot, user_message)
    print_colored("Matilda:", f"{response}\n\n")
    
    # write it in a document
    save_as_docx(f"{response}\n\n", docx_file_path);
    
    user_message_without_generate_image = re.sub(r'(Response:|Narration:|Image: generate_image:.*|)', '', response).strip()
    text_to_speech(user_message_without_generate_image, "en")
    
