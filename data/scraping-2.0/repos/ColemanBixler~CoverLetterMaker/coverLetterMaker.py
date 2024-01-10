import os
import openai

from tkinter import *
from tkinter import messagebox

from gptGeneration import run_conversation

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.text.run import Font

from datetime import date

# Variables used in GPT's cover letter generation
current_date = date.today().strftime("%B %d, %Y")

user_name = ""
user_address = ""
user_phone = ""
user_email = ""
occupation_title = ""

recipient_name = ""
recipient_title = ""
company_name = ""
company_address = ""

past_work = ""
relevant_experiences = ""
excited_reasons = ""
relevant_passions = ""

cover_letter = ""

# Establish the gui window
root = Tk()
root.title("Coleman's CoverLetterMaker")
root.geometry('500x300')

# Create a frame to be altered throughout app usage
frame = Frame(root)
frame.grid()

# Function for clearing everything on screen
def clear_frame():
    for widget in frame.winfo_children():
        widget.destroy()

# Get applicant info and set to previously initialized variables
def section1():
    # Label the current section (Section 1)
    section1 = Label(frame, text = "Section 1: Applicant Information\n")
    section1.grid(column = "1", row = "0")

    # Get the user's name
    user_name_label = Label(frame, text = "Your Name: ")
    user_name_label.grid(column = "0", row = "1")
    user_name_entry = Entry(frame, bd = 5)
    user_name_entry.grid(column = "1", row = "1")

    # Get the user's home address
    user_address_label = Label(frame, text = "Your Home Address\n(include City, State, and ZIP Code): ")
    user_address_label.grid(column = "0", row = "2")
    user_address_entry = Entry(frame, bd = 5)
    user_address_entry.grid(column = "1", row = "2")

    # Get the user's phone number
    user_phone_label = Label(frame, text = "Your Phone Number: ")
    user_phone_label.grid(column = "0", row = "3")
    user_phone_entry = Entry(frame, bd = 5)
    user_phone_entry.grid(column = "1", row = "3")

    # Get the user's email
    user_email_label = Label(frame, text = "Your Email: ")
    user_email_label.grid(column = "0", row = "4")
    user_email_entry = Entry(frame, bd = 5)
    user_email_entry.grid(column = "1", row = "4")

    # Get the job the user is applying for
    occupation_title_label = Label(frame, text = "Role You're Applying For: ")
    occupation_title_label.grid(column = "0", row = "5")
    occupation_title_entry = Entry(frame, bd = 5)
    occupation_title_entry.grid(column = "1", row = "5")

    # End of Section 1 input
    def submit_1():
        # Assign submitted info to variables
        global user_name
        user_name = user_name_entry.get()
        
        global user_address
        user_address = user_address_entry.get()
        
        global user_phone
        user_phone = user_phone_entry.get()
        
        global user_email
        user_email = user_email_entry.get()
        
        global occupation_title
        occupation_title = occupation_title_entry.get()

        # Warn user about missing info
        if(user_name == "" or user_address == "" or 
           user_phone == "" or user_email == "" or
           occupation_title == ""):
            answer = messagebox.askokcancel("Missing Info", "Not all forms have been filled out. Do you still want to continue?")
            if answer:
                section2()
        else:
            section2()

    # Create Button for user to press once every thing is filled out
    btn_1 = Button(frame, text = "Submit", command = submit_1)
    btn_1.grid(column = "1", row = "6")

# Get recipient info and set to previously initialized variables
def section2():
    clear_frame()

    # Label the current section (Section 2)
    section2 = Label(frame, text = "Section 2: Recipient Information\n")
    section2.grid(column = "1", row = "0")

    # Get the email recipient's name
    recipient_name_label = Label(frame, text = "Name of the person you are emailing: ")
    recipient_name_label.grid(column = "0", row = "1")
    recipient_name_entry = Entry(frame, bd = 5)
    recipient_name_entry.grid(column = "1", row = "1")

    # Get the email recipient's title
    recipient_title_label = Label(frame, text = "Title of the person you are emailing: ")
    recipient_title_label.grid(column = "0", row = "2")
    recipient_title_entry = Entry(frame, bd = 5)
    recipient_title_entry.grid(column = "1", row = "2")

    # Get the name of the company the user is applying to
    company_name_label = Label(frame, text = "Name of the company you're applying to: ")
    company_name_label.grid(column = "0", row = "3")
    company_name_entry = Entry(frame, bd = 5)
    company_name_entry.grid(column = "1", row = "3")

    # Get the address of the company the user is applying to
    company_address_label = Label(frame, text = "The Company's Address\n(include City, State, and ZIP Code): ")
    company_address_label.grid(column = "0", row = "4")
    company_address_entry = Entry(frame, bd = 5)
    company_address_entry.grid(column = "1", row = "4")

    # End of Section 2 input
    def submit_2():
        # Assign submitted info to variables
        global recipient_name
        recipient_name = recipient_name_entry.get()
        
        global recipient_title
        recipient_title = recipient_title_entry.get()
        
        global company_name
        company_name = company_name_entry.get()
        
        global company_address
        company_address = company_address_entry.get()

        # Warn user about missing info
        if(recipient_name == "" or recipient_title == "" or 
           company_name == "" or company_address == ""):
            answer = messagebox.askokcancel("Missing Info", "Not all forms have been filled out. Do you still want to continue?")
            if answer: 
                section3()
        else:
            section3()

    # Create Button for user to press once every thing is filled out
    btn_2 = Button(frame, text = "Submit", command = submit_2)
    btn_2.grid(column = "1", row = "5")

# Get user's aptitude info and set to previously initialized variables
def section3():
    clear_frame()

    # Label the current section (Section 3)
    section3 = Label(frame, text = "Section 3: Aptitude Description\n")
    section3.grid(column = "1", row = "0")

    # Get info about user's past work
    past_work_label = Label(frame, text = "You have previously worked at... ")
    past_work_label.grid(column = "0", row = "1")
    past_work_entry = Entry(frame, bd = 5)
    past_work_entry.grid(column = "1", row = "1")

    # Get info about user's past relevant experiences
    relevant_experiences_label = Label(frame, text = "You have experience in... ")
    relevant_experiences_label.grid(column = "0", row = "2")
    relevant_experiences_entry = Entry(frame, bd = 5)
    relevant_experiences_entry.grid(column = "1", row = "2")

    # Get info about why the user wants to work there
    excited_reasons_label = Label(frame, text = "You are excited about this job because... ")
    excited_reasons_label.grid(column = "0", row = "3")
    excited_reasons_entry = Entry(frame, bd = 5)
    excited_reasons_entry.grid(column = "1", row = "3")

    # Get info about what the user is passionate about
    relevant_passions_label = Label(frame, text = "You are passionate about... ")
    relevant_passions_label.grid(column = "0", row = "4")
    relevant_passions_entry = Entry(frame, bd = 5)
    relevant_passions_entry.grid(column = "1", row = "4")

    # End of Section 3 input
    def submit_3():
        # Assign submitted info to variables
        global past_work
        past_work = past_work_entry.get()
        
        global relevant_experiences
        relevant_experiences = relevant_experiences_entry.get()
        
        global excited_reasons
        excited_reasons = excited_reasons_entry.get()
        
        global relevant_passions
        relevant_passions = relevant_passions_entry.get()

        # Warn user about missing info
        if(past_work == "" or relevant_experiences == "" or 
           excited_reasons == "" or relevant_passions == ""):
            answer = messagebox.askokcancel("Missing Info", "Not all forms have been filled out. Do you still want to continue?")
            if answer: 
                processing()
        else:    
            processing()

    # Create Button for user to press once every thing is filled out
    btn_3 = Button(frame, text = "Submit", command = submit_3)
    btn_3.grid(column = "1", row = "5")

def processing():
    clear_frame()

    # GPT Generation
    cover_letter = run_conversation(user_name, user_address, user_phone, user_email, occupation_title, 
                              recipient_name, recipient_title, company_name, company_address,
                              past_work, relevant_experiences, excited_reasons, relevant_passions,
                              current_date)

    # CoverLetter.docx Generation    
    docx_creation(cover_letter)

    # Completed label to signify the program has run its course
    program_complete = Label(frame, text = "Look at the 'CoverLetter.docx' file for the result. Thanks for using my program!")
    program_complete.grid(column = "1", row = "1")

def docx_creation(cover_letter):
    # turn letter into lines for iteration purposes
    cover_letter_lines = cover_letter.splitlines()
    
    # Create new document
    doc = Document()

    # Set the style of the document
    font = doc.styles['Normal'].font
    font.name = "Times New Roman"
    font.size = Pt(10)

    # First paragraph is personal info which should be centered
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Loop through lines of cover letter to create document
    for i, line in enumerate(cover_letter_lines):
        # First and Last lines will be bold; they are the sender's name
        if i == 0 or i == len(cover_letter_lines)-1:
            par.add_run(line + "\n").bold = True
        # The date should be on its own paragraph
        elif line == current_date:
            par = doc.add_paragraph(line)
        # Blank lines (signifiying line breaks) will start a new paragraph
        elif line == "":
            par = doc.add_paragraph()
        # The line before a line break or the date line should be added to the paragraph normally 
        elif cover_letter_lines[i+1] == "" or cover_letter_lines[i+1] == current_date:
            par.add_run(line)
        # Every other line will be added to the current paragraph, and then with a new line
        else:
            par.add_run(line + "\n")

    # Save the Final Product
    doc.save('CoverLetter.docx')

# Call to first fill out menu
section1()

# Let the program run
root.mainloop()

