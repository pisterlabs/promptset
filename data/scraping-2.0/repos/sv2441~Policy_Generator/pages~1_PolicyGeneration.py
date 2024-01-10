import streamlit as st
import pandas as pd
import base64
import csv
import math
import docx
import os
from langchain.output_parsers import OutputFixingParser
from langchain.llms import OpenAI
from langchain.chains import LLMChain, SimpleSequentialChain
from langchain.prompts import PromptTemplate, ChatPromptTemplate
from langchain.chat_models import ChatOpenAI
from langchain.output_parsers import ResponseSchema, StructuredOutputParser
from dotenv import load_dotenv
import pyperclip

load_dotenv()
os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
# os.environ["OPENAI_API_KEY"] = os.getenv('OPENAI_API_KEY')

# Initialize chat model
def dict_to_csv(data, filename, append=False):
    mode = 'a' if append else 'w'
    with open(filename, mode, newline='') as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=data.keys())
        if not append:
            writer.writeheader()
        writer.writerow(data)

# Function to Combine policy based on Topic
def combine_policy(df, topic_column):
    grouped_df = df.groupby(topic_column)['Policy'].apply(lambda x: ' '.join(x)).reset_index()
    grouped_df.rename(columns={'Policy': 'Policies'}, inplace=True)
    grouped_df.to_csv('grouped.csv', index=False)
    return grouped_df

# for combining all rows into a single variable
def combine_column_to_paragraph(df, column_name):
    column_data = df[column_name].tolist()
    paragraph = " ".join(str(item) for item in column_data)
    return paragraph

# Function to concatenate columns and download the modified CSV file
def process_csv(df):
    df['Combined'] = df['OP Title'].astype(str) + df['OP Description'].astype(str)
    modified_csv = df.to_csv(index=False)

    # Download the modified CSV file
    b64 = base64.b64encode(modified_csv.encode()).decode()
    href = f'<a href="data:file/csv;base64,{b64}" download="modified.csv">Download modified CSV file</a>'
    st.markdown(href, unsafe_allow_html=True)

    st.dataframe(df)
    df.to_csv('process_result.csv')

# Function to process the policy generation
def policy_generator(df, topic_column):
    Policy_schema = ResponseSchema(name="Policy", description="Policy Statement")
    response_schemas = [Policy_schema]
    output_parser = StructuredOutputParser.from_response_schemas(response_schemas)
    format_instructions = output_parser.get_format_instructions()
    new_parser = OutputFixingParser.from_llm(parser=output_parser, llm=ChatOpenAI())

    title_template = """ \ You are an AI Governance bot. Just Execute the set of steps one by one.
                Convert "{topic}" into policy statements in maximum 10 to 15 words.
                {format_instructions}
                """
    prompt = ChatPromptTemplate.from_template(template=title_template)
    df2 = df["Combined"]

    for i in range(len(df2)):
        messages = prompt.format_messages(topic=df2[i], format_instructions=format_instructions)
        response = chat_llm(messages)
        response_as_dict = new_parser.parse(response.content)
        data = response_as_dict
        dict_to_csv(data, 'policy.csv', append=True)

    result = pd.read_csv("policy.csv", names=['Policy'])
    final = pd.concat([df, result], axis=1)
    final.to_csv("policy_result.csv")
    # st.dataframe(final)

# Function to process the summary generation
def summary_generator(df, topic_column):
    combine_policy(df, topic_column)
    new_df = pd.read_csv('grouped.csv')

    Summary_schema = ResponseSchema(name="Summary", description="Summary as instructional policy paragraph")
    response_schemas = [Summary_schema]
    output_parser = StructuredOutputParser.from_response_schemas(response_schemas)
    format_instructions = output_parser.get_format_instructions()

    title_template = """ \ You are an AI Governance bot.  
                    Summarize "{topic}" as an instructional policy paragraph in {count} Words In Legal Language Style.
                     {format_instructions}
                """
    prompt = ChatPromptTemplate.from_template(template=title_template)

    for index, row in new_df.iterrows():
        topic = row['Policies']
        word_count_topic = len(topic.split())
        count = math.ceil(0.2 * word_count_topic)

        messages = prompt.format_messages(topic=topic, count=count, format_instructions=format_instructions)
        response = chat_llm(messages)
        content = str(response.content)
        response_as_dict = output_parser.parse(response.content)
        data = response_as_dict
        dict_to_csv(data, 'summary.csv', append=True)
    result = pd.read_csv("summary.csv", names=['Summary'])
    final = pd.concat([new_df, result], axis=1)
    final.to_csv("summary_result.csv")
    # st.dataframe(final)

def document_generator(df):
    contents = combine_column_to_paragraph(df, 'Summary')

    title_template = """ \ You are an AI Governance bot.  
                    Restructure {topic} based on topic into a policy document as a Paragraph in Legal Language Style.
                    
                """
    prompt = ChatPromptTemplate.from_template(template=title_template)

    if os.path.exists('Policy_Document.doc'):
        doc = docx.Document('Policy_Document.doc')
    else:
        doc = docx.Document()

    messages = prompt.format_messages(topic=contents)
    response = chat_llm(messages)
    content = str(response.content)
    pyperclip.copy(content)
    st.write(content)
    doc.add_paragraph(content)
    doc.save('Policy_Document.doc')
    

    # Add a button to download the policy document
    with open('Policy_Document.doc', 'rb') as f:
        doc_data = f.read()
    b64 = base64.b64encode(doc_data).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="Policy_Document.doc">Download Result</a>'
    st.markdown(href, unsafe_allow_html=True)

    # # Add a button to copy the text of the policy document
    # if st.button("Copy Policy Document Text"):
    #     st.write(content)

# Streamlit app
def main():
    st.title("Policy Prodago")

    # File upload
    uploaded_file = st.file_uploader("Upload CSV file", type=["csv"])

    if uploaded_file is not None:
        st.write("File uploaded successfully!")
        df = pd.read_csv(uploaded_file)

        st.subheader("CSV File Preview")
        st.dataframe(df)
        
        # Dropdown for selecting topic level
        topic_level = st.selectbox("Select Topic Level", ["topic", "L2 topic", "L3 topic", "L4 topic"])

        if st.button("Process the file"):
            process_csv(df)
            policy_generator(pd.read_csv("process_result.csv"), topic_level)
            st.write("Generating Summary...")
            summary_generator(pd.read_csv('policy_result.csv', usecols=['Policy', topic_level]), topic_level)
            st.write("Generating Policy...")
            document_generator(pd.read_csv("summary_result.csv"))

if __name__ == "__main__":
    main()
