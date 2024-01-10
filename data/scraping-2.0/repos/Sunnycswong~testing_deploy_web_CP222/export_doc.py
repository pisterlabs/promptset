#%%
## Import Library
import os
import copy
import json
import re
import logging
import datetime

#import openai
#========================================================================
#from langchain.chains import LLMChain, ConversationalRetrievalChain, RetrievalQA
#from langchain.llms import AzureOpenAI 
#from langchain.memory import ConversationBufferMemory, CosmosDBChatMessageHistory
#from langchain.prompts import PromptTemplate
#from langchain.chat_models import AzureChatOpenAI
#from langchain.embeddings.openai import OpenAIEmbeddings
#from langchain.vectorstores.azuresearch import AzureSearch
#from langchain.chains.question_answering import load_qa_chain
#from langchain.retrievers import AzureCognitiveSearchRetriever
#from langchain.schema import HumanMessage, Document

#========================================================================
from azure.storage.blob import BlobServiceClient
#from azure.core.exceptions import ResourceExistsError
#from azure.identity import AzureDeveloperCliCredential
#from azure.search.documents import SearchClient
#from azure.search.documents.indexes import SearchIndexClient
# from azure.search.documents.indexes.models import (
#     HnswParameters,
#     PrioritizedFields,
#     SearchableField,
#     SearchField,
#     SearchFieldDataType,
#     SearchIndex,
#     SemanticConfiguration,
#     SemanticField,
#     SemanticSettings,
#     SimpleField,
#     VectorSearch,
#     VectorSearchAlgorithmConfiguration,
# )
#===============================================================
from io import BytesIO
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

#from langdetect import detect
#from pypdf import PdfReader

# Blob Storage keys: hosted by Kenny
STORAGE_SERVICE = "creditproposal"
STORAGE_API_KEY = "hJ2qb//J1I1KmVeDHBpwEpnwluoJzm+b6puc5h7k+dnDSFQ0oxuh1qBz+qPB/ZT7gZvGufwRbUrN+ASto6JOCw=="
CONNECTION_STRING = f"DefaultEndpointsProtocol=https;AccountName={STORAGE_SERVICE};AccountKey={STORAGE_API_KEY}"
CONTAINER_NAME = "exportdocs"

# doc_intell Keys
DOC_INTELL_ENDPOINT = "https://doc-intelligence-test.cognitiveservices.azure.com/"
DOC_INTELL_KEY = "9fac3bb92b3c4ef292c20df9641c7374"


# set up openai environment - Jay
#os.environ["OPENAI_API_TYPE"] = "azure"
#os.environ["OPENAI_API_BASE"] = "https://pwcjay.openai.azure.com/"
#os.environ["OPENAI_API_VERSION"] = "2023-05-15"
#os.environ["OPENAI_API_KEY"] = "f282a661571f45a0bdfdcd295ac808e7"

# set up openai environment - Ethan
os.environ["OPENAI_API_TYPE"] = "azure"
os.environ["OPENAI_API_BASE"] = "https://lwyethan-azure-openai-test-01.openai.azure.com/"
os.environ["OPENAI_API_VERSION"] = "2023-05-15"
os.environ["OPENAI_API_KEY"] = "ff96d48045584cb9844fc70e5b802918"

# Setting up ACS -Jay
#os.environ["AZURE_COGNITIVE_SEARCH_SERVICE_NAME"] = search_service
#os.environ["AZURE_COGNITIVE_SEARCH_API_KEY"] = search_api_key
#os.environ["AZURE_INDEX_NAME"] = index_name

def create_docx(client_name, json_data):
    """
    Create a word document based on the latest generated text

    Parameters:
    ----------
    client_name: str
        The client name required
    json_data: json
        The input json according to our hierarchy format

    Return:
    -------
    blob_name: str
        The output filename
    document_bytes: io.BytesIO
        Output word object in io.BytesIO format
    """
    # Create a new Word document
    document = DocxDocument()

    title_text = "Credit Proposal for " + client_name
    title_size = 20 # Font size in points

    # Create a paragraph for the title
    title_paragraph = document.add_paragraph()

    # Add the title text to the paragraph
    title_run = title_paragraph.add_run(title_text)

    # Apply formatting to the title run
    title_run.bold = True
    title_run.font.size = Pt(title_size)

    # Set the alignment of the paragraph to the center
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Convert JSON values to section headers and paragraphs in the Word document
    for item in json_data['consolidated_text']:
        section = item['section']
        context = item['output']

        # Add the section header
        document.add_heading(section, level=1)

        # Split context into lines and check each line
        for line in context.split('\n'):
            # Create a new paragraph for each line
            paragraph = document.add_paragraph()

            # Search for the pattern [RM please ... ] using regex
            matches = re.findall(r'\[RM .*?\]', line)

            if matches:
                # If there's a match, split line into parts
                parts = re.split(r'(\[RM .*?\])', line)

                for part in parts:
                    run = paragraph.add_run(part)

                    if part in matches:
                        # This part should be colored red
                        run.font.color.rgb = RGBColor(255, 0, 0)  # RGB values for red
            else:
                # Normal text
                run = paragraph.add_run(line)
                
    current_time = datetime.datetime.now()
    time_string = current_time.strftime("_%Y_%m_%d_%H_%M_%S")
    blob_name = client_name + time_string + '_Word_proposal.docx'
    
    # Save the Word document to a BytesIO object
    document_bytes = BytesIO()
    document.save(document_bytes)
    document_bytes.seek(0)  # Reset the stream position to the beginning

    # Store the Word document in Azure Blob Storage

    # Temp disable the Blob Storage 
    #blob_service_client = BlobServiceClient.from_CONNECTION_STRING(CONNECTION_STRING)
    #container_client = blob_service_client.get_container_client(CONTAINER_NAME)
    #blob_client = container_client.get_blob_client(blob_name)
    #blob_client.upload_blob(document_bytes)

    # Make sure to reset the stream position again before returning
    document_bytes.seek(0) 
    
    return blob_name, document_bytes
