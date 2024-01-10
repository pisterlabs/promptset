import openai
from docx import Document
import argparse
from pydub import AudioSegment
from config import OPENAI_API_KEY

import tiktoken

encoding = tiktoken.get_encoding("cl100k_base")
# encoding = tiktoken.encoding_for_model("gpt4")

openai.api_key = OPENAI_API_KEY

# Create the parser
parser = argparse.ArgumentParser(description='Process meeting minutes from an audio file.')

# Add an argument
parser.add_argument('audio_file_path', type=str, help='The path to the audio file')

# Parse the arguments
args = parser.parse_args()

def transcribe_audio(audio_file_path, chunk_duration_ms=60000):  # Default chunk duration is 60 seconds
    audio = AudioSegment.from_wav(audio_file_path)
    chunks = make_chunks(audio, chunk_duration_ms)

    transcriptions = []
    for i, chunk in enumerate(chunks):
        chunk.export("./tmp/chunk{0}.wav".format(i), format="wav")
        with open("./tmp/chunk{0}.wav".format(i), 'rb') as audio_file:
            transcription = openai.Audio.transcribe("whisper-1", audio_file)
        transcriptions.append(transcription['text'])

    return ' '.join(transcriptions)

def make_chunks(audio, chunk_duration_ms):
    chunk_length = len(audio)
    chunks = []
    for i in range(0, chunk_length, chunk_duration_ms):
        chunks.append(audio[i:i+chunk_duration_ms])
    return chunks

def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }

def abstract_summary_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def key_points_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def action_item_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']

def sentiment_analysis(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)

def count_tokens(text):
    num_tokens = len(encoding.encode(text))
    return num_tokens

def chunk_transcription(transcription, max_tokens=8100):
    words = transcription.split(' ')
    chunks = []
    current_chunk = ''
    for word in words:
        if count_tokens(current_chunk + ' ' + word) > max_tokens:
            chunks.append(current_chunk)
            current_chunk = word
        else:
            current_chunk += ' ' + word
    chunks.append(current_chunk)
    return chunks


# Now you can use args.audio_file_path to get the audio file path
transcription = transcribe_audio(args.audio_file_path)
chunks = chunk_transcription(transcription)
minutes = [meeting_minutes(chunk) for chunk in chunks]

# Combine all the minutes into a single dictionary
combined_minutes = {
    'abstract_summary': '',
    'key_points': '',
    'action_items': '',
    'sentiment': ''
}
for minute in minutes:
    for key in combined_minutes:
        combined_minutes[key] += minute[key]

print(combined_minutes)

save_as_docx(combined_minutes, './output/meeting_minutes.docx')