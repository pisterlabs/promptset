import os

import openai
from docx import Document
import time
from dotenv import load_dotenv
import argparse
from glob import glob

load_dotenv()

openai.api_key = os.getenv('OPENAI_API_KEY')
model = "gpt-4"
temperature = 0
default_extension = 'm4a'


# ---------- Transcription --------------------

def transcribe_audio(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        transcription = openai.Audio.transcribe("whisper-1", audio_file)
    return transcription['text']


# ---------- Abstract Summary --------------------

def full_abstract_summary_extraction(transcription,
                                     meeting_description=None
                                     ):
    system_prompt = ("You are a highly skilled AI trained in language comprehension and summarization. " +
                     (f"This is a meeting about {meeting_description}. " if meeting_description else "") +
                     f"I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points.")

    response = openai.ChatCompletion.create(
        model=model,
        temperature=temperature,
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def chunked_abstract_summary_extraction(transcription,
                                        meeting_description=None
                                        ):
    max_tokens = 8000  # A bit less than 8192 to leave some room for the system message
    overlap = 1000  # Overlap size - tune this based on your use case

    transcript_parts = [transcription[i:i + max_tokens + overlap] for i in range(0, len(transcription), max_tokens)]

    final_summary = ""
    previous_summary = ""  # Initialize previous summary

    for part in transcript_parts:
        # Generate a summary of the chunk
        system_prompt = ("You are a highly skilled AI trained in language comprehension and summarization. " +
                         (f"This is a meeting about {meeting_description}. " if meeting_description else "") +
                         f"Previously, you summarized: '{previous_summary}'. Now, I would like you to read the following text and summarize it into a concise abstract paragraph, building upon your previous summary. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points.")

        response = openai.ChatCompletion.create(
            model=model,
            temperature=temperature,
            messages=[
                {
                    "role": "system",
                    "content": system_prompt
                },
                {
                    "role": "user",
                    "content": part
                }
            ]
        )

        previous_summary = response['choices'][0]['message']['content']  # Update previous summary
        final_summary += previous_summary + "\n"

    # Use GPT-4 to rephrase the final summary into a more cohesive paragraph
    response = openai.ChatCompletion.create(
        model=model,
        temperature=temperature,
        messages=[
            {
                "role": "system",
                "content": "As an AI trained in language comprehension and summarization, your task is to rephrase the following summaries into a more cohesive and concise paragraph. Please maintain the overall meaning and key details in your rephrasing."
            },
            {
                "role": "user",
                "content": final_summary
            }
        ]
    )

    final_summary = response['choices'][0]['message']['content']

    return final_summary


def abstract_summary_extraction(transcription,
                                meeting_description=None
                                ):
    try:
        # Try the original method first
        return full_abstract_summary_extraction(transcription, meeting_description)
    except openai.error.InvalidRequestError as e:
        # If the original method fails due to exceeding the maximum token limit, fall back to the chunking method
        if 'token' in str(e):
            print("Using chunking for abstract summary extraction.")
            # If the server returns a 502, wait 10 seconds then retry
            try:
                return chunked_abstract_summary_extraction(transcription, meeting_description)
            except openai.error.APIError as e:
                if e.http_status == 502:
                    print("API returned a 502 Bad Gateway error. Retrying in 10 seconds...")
                    time.sleep(10)
                    return chunked_abstract_summary_extraction(transcription, meeting_description)
                else:
                    # If the error is due to another reason, raise it
                    raise e
        else:
            # If the error is due to another reason, raise it
            raise e


# ---------- Key Points --------------------

def full_key_points_extraction(transcription,
                               meeting_description=None
                               ):
    system_prompt = ("You are a proficient AI with a specialty in distilling information into key points. " +
                     (f"This is a meeting about {meeting_description}. " if meeting_description else "") +
                     f"Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about.")

    response = openai.ChatCompletion.create(
        model=model,
        temperature=temperature,
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def chunked_key_points_extraction(transcription,
                                  meeting_description=None
                                  ):
    max_tokens = 8000  # A bit less than 8192 to leave some room for the system message
    overlap = 1000  # Overlap size - tune this based on your use case

    transcript_parts = [transcription[i:i + max_tokens + overlap] for i in range(0, len(transcription), max_tokens)]

    final_key_points = []
    previous_key_points = ""  # Initialize previous key points

    for part in transcript_parts:
        # Extract key points from the chunk
        system_prompt = ("You are a proficient AI with a specialty in distilling information into key points. " +
                         (f"This is a meeting about {meeting_description}. " if meeting_description else "") +
                         f"Previously, you identified: '{previous_key_points}'. Now, based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about.")

        response = openai.ChatCompletion.create(
            model=model,
            temperature=temperature,
            messages=[
                {
                    "role": "system",
                    "content": system_prompt
                },
                {
                    "role": "user",
                    "content": part
                }
            ]
        )

        previous_key_points = response['choices'][0]['message']['content']  # Update previous key points
        final_key_points.append(previous_key_points)

    # Combine all key points into a single string
    all_key_points = "\n".join(final_key_points)

    # Use GPT-4 to reformat and renumber the key points
    system_prompt = (
            "You are a proficient AI with a specialty in organizing and formatting information. " +
            f"Please take the following key points{f' (from a meeting about {meeting_description})' if meeting_description else ''} and reformat them into a coherent, numbered list. Ensure that the numbering is consistent, starts at number 1, and does not restart. Each key point should start on a new line."
    )

    response = openai.ChatCompletion.create(
        model=model,
        temperature=temperature,
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": all_key_points
            }
        ]
    )

    final_key_points = response['choices'][0]['message']['content']

    return final_key_points


def key_points_extraction(transcription,
                          meeting_description=None
                          ):
    try:
        # Try the original method first
        return full_key_points_extraction(transcription, meeting_description)
    except openai.error.InvalidRequestError as e:
        # If the original method fails due to exceeding the maximum token limit, fall back to the chunking method
        if 'token' in str(e):
            print("Using chunking for key points extraction.")
            # If the server returns a 502, wait 10 seconds then retry
            try:
                return chunked_key_points_extraction(transcription, meeting_description)
            except openai.error.APIError as e:
                if e.http_status == 502:
                    print("API returned a 502 Bad Gateway error. Retrying in 10 seconds...")
                    time.sleep(10)
                    return chunked_key_points_extraction(transcription, meeting_description)
                else:
                    # If the error is due to another reason, raise it
                    raise e
        else:
            # If the error is due to another reason, raise it
            raise e


# ---------- Action Items --------------------

def full_action_item_extraction(transcription,
                                meeting_description=None
                                ):
    system_prompt = ("You are an AI expert in analyzing conversations and extracting action items. " +
                     (f"This is a meeting about {meeting_description}. " if meeting_description else "") +
                     f"Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely.")

    response = openai.ChatCompletion.create(
        model=model,
        temperature=temperature,
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def chunked_action_item_extraction(transcription,
                                   meeting_description=None
                                   ):
    max_tokens = 8000  # A bit less than 8192 to leave some room for the system message
    overlap = 1000  # Overlap size - tune this based on your use case

    transcript_parts = [transcription[i:i + max_tokens + overlap] for i in range(0, len(transcription), max_tokens)]

    final_action_items = ""
    previous_action_items = ""  # Initialize previous action items

    for part in transcript_parts:
        # Extract action items from the chunk
        system_prompt = ("You are an AI expert in analyzing conversations and extracting action items. " +
                         (f"This is a meeting about {meeting_description}. " if meeting_description else "") +
                         f"Previously, you identified: '{previous_action_items}'. Now, please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done, building upon your previous list. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely.")

        response = openai.ChatCompletion.create(
            model=model,
            temperature=temperature,
            messages=[
                {
                    "role": "system",
                    "content": system_prompt
                },
                {
                    "role": "user",
                    "content": part
                }
            ]
        )

        previous_action_items = response['choices'][0]['message']['content']  # Update previous action items
        final_action_items += previous_action_items + "\n"

    # Use GPT-4 to consolidate the action items into a single, coherent list
    response = openai.ChatCompletion.create(
        model=model,
        temperature=temperature,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in synthesizing information, your task is to consolidate the following action items into a single, concise, and coherent list. Ensure the list is organized in a clear and concise manner. Do not overwhelm the reader with too many action items."
            },
            {
                "role": "user",
                "content": final_action_items
            }
        ]
    )

    final_action_items = response['choices'][0]['message']['content']

    return final_action_items


def action_item_extraction(transcription,
                           meeting_description=None
                           ):
    try:
        # Try the original method first
        return full_action_item_extraction(transcription, meeting_description)
    except openai.error.InvalidRequestError as e:
        # If the original method fails due to exceeding the maximum token limit, fall back to the chunking method
        if 'token' in str(e):
            print("Using chunking for action item extraction.")
            # If the server returns a 502, wait 10 seconds then retry
            try:
                return chunked_action_item_extraction(transcription, meeting_description)
            except openai.error.APIError as e:
                if e.http_status == 502:
                    print("API returned a 502 Bad Gateway error. Retrying in 10 seconds...")
                    time.sleep(10)
                    return chunked_action_item_extraction(transcription, meeting_description)
                else:
                    # If the error is due to another reason, raise it
                    raise e
        else:
            # If the error is due to another reason, raise it
            raise e


# ---------- Sentiment Analysis --------------------

def full_sentiment_analysis(transcription,
                            meeting_description=None
                            ):
    system_prompt = ("As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. " +
                     (f"This is a meeting about {meeting_description}. " if meeting_description else "") +
                     f"Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible.")

    response = openai.ChatCompletion.create(
        model=model,
        temperature=temperature,
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def chunked_sentiment_analysis(transcription,
                               meeting_description=None
                               ):
    max_tokens = 8000  # A bit less than 8192 to leave some room for the system message
    overlap = 1000  # Overlap size - tune this based on your use case

    transcript_parts = [transcription[i:i + max_tokens + overlap] for i in range(0, len(transcription), max_tokens)]

    final_sentiment = ""
    previous_sentiment = ""  # Initialize previous sentiment

    for part in transcript_parts:
        # Analyze the sentiment of the chunk
        system_prompt = (
                "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. " +
                (f"This is a meeting about {meeting_description}. " if meeting_description else "") +
                f"Previously, you analyzed: '{previous_sentiment}'. Now, please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible.")

        response = openai.ChatCompletion.create(
            model=model,
            temperature=temperature,
            messages=[
                {
                    "role": "system",
                    "content": system_prompt
                },
                {
                    "role": "user",
                    "content": part
                }
            ]
        )

        previous_sentiment = response['choices'][0]['message']['content']  # Update previous sentiment
        final_sentiment += previous_sentiment + "\n"

    # Use GPT-4 to rephrase the final sentiment analysis into a more cohesive paragraph
    response = openai.ChatCompletion.create(
        model=model,
        temperature=temperature,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to rephrase the following sentiment analysis into a more cohesive and concise paragraph. Please maintain the overall sentiment and key details in your rephrasing."
            },
            {
                "role": "user",
                "content": final_sentiment
            }
        ]
    )

    final_sentiment = response['choices'][0]['message']['content']

    return final_sentiment


def sentiment_analysis(transcription,
                       meeting_description=None
                       ):
    try:
        # Try the original method first
        return full_sentiment_analysis(transcription, meeting_description)
    except openai.error.InvalidRequestError as e:
        # If the original method fails due to exceeding the maximum token limit, fall back to the chunking method
        if 'token' in str(e):
            print("Using chunking for sentiment analysis.")
            # If the server returns a 502, wait 10 seconds then retry
            try:
                return chunked_sentiment_analysis(transcription, meeting_description)
            except openai.error.APIError as e:
                if e.http_status == 502:
                    print("API returned a 502 Bad Gateway error. Retrying in 10 seconds...")
                    time.sleep(10)
                    return chunked_sentiment_analysis(transcription, meeting_description)
                else:
                    # If the error is due to another reason, raise it
                    raise e
        else:
            # If the error is due to another reason, raise it
            raise e


# ---------- Main Functions --------------------

def save_as_docx(minutes,
                 filename
                 ):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)


def meeting_minutes(transcription,
                    meeting_description=None
                    ):
    abstract_summary = abstract_summary_extraction(transcription, meeting_description)
    key_points = key_points_extraction(transcription, meeting_description)
    action_items = action_item_extraction(transcription, meeting_description)
    sentiment = sentiment_analysis(transcription, meeting_description)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }


def cli():
    parser = argparse.ArgumentParser(prog="meeting-minutes.py", description="Generate meeting minutes from audio file")
    parser.add_argument("input_dir",
                        help="directory containing audio files.  Dir must be under 'audio' directory")
    parser.add_argument("-e", "--extension",
                        help="audio file extension",
                        default=default_extension)
    parser.add_argument("-r", "--review",
                        help="review the transcription before generating meeting minutes",
                        action="store_true")

    args = parser.parse_args()

    return {
        'input_dir': args.input_dir,
        'extension': args.extension,
        'review': args.review,
    }


if __name__ == '__main__':
    args = cli()

    # audio file extension
    if args['extension'].strip() == "":
        args['extension'] = default_extension
    extension = "*." + args['extension'].strip().replace('.', '')
    print(f"extension: '{extension}'")

    # input directory containing audio files
    input_dir = args['input_dir']
    if not os.path.exists(input_dir):
        print(f"input_dir does not exist: '{input_dir}'")
        exit(1)
    print(f"input dir: '{input_dir}'")
    print()

    audio_files = glob(os.path.join(input_dir, extension))
    print(f"found {len(audio_files)} audio files: {audio_files}")
    print()

    # Ask the user for an optional meeting description
    meeting_description = input('Complete the sentence: "This is a meeting about..." (or press Enter to skip): ')
    print()
    # If the user didn't provide a description, set meeting_description to None
    if meeting_description.strip() == "":
        meeting_description = None

    print("Transcribing audio files...")
    full_transcription = ""
    last_dir = os.path.basename(input_dir)

    for audio_file in audio_files:
        full_transcription += transcribe_audio(audio_file)

    transcription_file = f"{input_dir}/{last_dir}_transcription.txt"
    with open(transcription_file, 'w') as f:
        f.write(full_transcription)

    print(f"transcription files written to: '{transcription_file}'")
    print()

    if args['review']:
        print("Review and edit transcription as needed.")
        x = input("Press Enter to continue")
        with open(transcription_file, 'r') as f:
            full_transcription = f.read()
        print()

    print("Generating meeting minutes...")
    summary_text = meeting_minutes(full_transcription, meeting_description)
    summary_file = f"{input_dir}/{last_dir}_summary.docx"

    save_as_docx(summary_text, summary_file)
    print(f"Meeting summary written to: '{summary_file}'")
