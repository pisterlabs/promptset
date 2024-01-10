import streamlit as st
from llama_index import VectorStoreIndex, SimpleDirectoryReader, ServiceContext
from llama_index.llms import OpenAI
from PyPDF2 import PdfReader
import io
from dotenv import load_dotenv
import tempfile
from llama_index import Document
from llama_index.query_engine.retriever_query_engine import RetrieverQueryEngine
from llama_index.indices.vector_store.retrievers import VectorIndexRetriever
import os
import openai
import pyrebase
import markdown
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from io import BytesIO

import time
from streamlit_lottie import st_lottie
from streamlit_lottie import st_lottie_spinner
import requests


def app():
    # Retrieve OpenAI API key from secrets
    openai.api_key = st.secrets["OPENAI_API_KEY"]

    input_column, response_column = st.columns([3,2])
    input_column.image("onepager-logo.png", use_column_width="auto")


    def load_lottieurl(url: str):
        r = requests.get(url)
        if r.status_code != 200:
            return None
        return r.json()
    lottie_url_document = "https://lottie.host/6f7bd064-3488-46b0-a2c1-e74162937a76/kaFAwA0KFR.json"
    lottie_doc = load_lottieurl(lottie_url_document)


    # Add inputs for sender, recipient, and context
    sender_column1, sender_column2 = input_column.columns([1,2])
    sender_column1.markdown('&nbsp;')
    sender_column1.markdown(r"$\textsf{\large Who Sends the OnePager?}$")
    sender = sender_column2.text_input('', key='sender', placeholder="e.g. Assistant to the Board of Management at a medium-sized automotive supplier")

    recipient_column1, recipient_column2 = input_column.columns([1,2])
    recipient_column1.markdown('&nbsp;')
    recipient_column1.markdown(r"$\textsf{\large Who recieves the OnePager?}$")
    recipient_column1.markdown('&nbsp;')
    recipient = recipient_column2.text_input('', key='recipient', placeholder="e.g. Marketing director in our/another company")

    context = input_column.text_input(r"$\textsf{\large What is the context of the OnePager?}$", key='context', placeholder=" e.g. Proposal for cooperation with software company xy for MVP development")

    # Add dropdown for document structure
    # New code
    doc_structure = input_column.radio(r"$\textsf{\large How should the OnePager be structured?}$", ['Decision Paper','Speaking Points','Report'], horizontal=True)# Add sliders for tone, technicality, and length

    formality_labels = {1: 'Casual', 2: 'Somewhat Casual', 3: 'Neutral', 4: 'Somewhat Formal', 5: 'Formal'}
    formality_options = {'Casual': 1, 'Somewhat Casual': 2, 'Neutral': 3, 'Somewhat Formal': 4, 'Formal': 5}
    formality_label = input_column.select_slider(r"$\textsf{\large Formality}$", options=list(formality_options.keys()))
    tone_value = formality_options[formality_label]
    tone = formality_labels[tone_value]



    technicality_options = {'Non-technical': 1, 'Somewhat non-technical': 2, 'Neutral': 3, 'Somewhat technical': 4, 'Technical': 5}
    technicality_label = input_column.select_slider(r"$\textsf{\large Technicality}$", options=list(technicality_options.keys()))
    technicality = technicality_options[technicality_label]

    # Add slider for max tokens
    length_options = {'Short': 350, 'Medium': 500, 'Long': 650}
    length_label = input_column.select_slider(r"$\textsf{\large Length}$", options=list(length_options.keys()))
    max_tokens = length_options[length_label]

    # Add file uploader for background information
    uploaded_file = input_column.file_uploader(r"$\textsf{\large Upload a PDF with background information.}$", type="pdf")

    # Add inputs for source description, call to action, and additional info
    source_description = input_column.text_input(r"$\textsf{\large What kind of document is this? Why is it relevant?}$", key="source", placeholder="e.g. Relevant Case-Study to xy, which supports the argumentation")
    call_to_action = input_column.text_input(r"$\textsf{\large What is the recommendation for action?}$", key="rfa", placeholder="e.g. We should cooperate with software company xy")
    
    action_tone_labels = {1: 'Indirect', 2: 'Somewhat Indirect', 3: 'Moderate', 4: 'Somewhat Direct', 5: 'Direct'}
    action_tone_options = {'Indirect': 1, 'Somewhat Indirect': 2, 'Moderate': 3, 'Somewhat Direct': 4, 'Direct': 5}
    action_tone_label = input_column.select_slider(r"$\textsf{\large Action tone}$", options=list(action_tone_options.keys()))
    action_tone_value = action_tone_options[action_tone_label]
    action_tone = action_tone_labels[action_tone_value]


    additional_info = input_column.text_input(r"$\textsf{\large What additional information belongs in the OnePager?}$", placeholder="e.g. We have already had a successful workshop with the partner")

    # New code
    deadline_column1, deadline_column2 = input_column.columns(2)
    deadline_type = deadline_column1.radio(r"$\textsf{\large Is there a deadline?}$", ['Yes', 'No'])
    if deadline_type == 'Yes':
        deadline_date = deadline_column2.date_input(r"$\textsf{\large Select a date}$")



    if uploaded_file is not None:
        # Create the ServiceContext with the user-selected temperature
        service_context = ServiceContext.from_defaults(llm=OpenAI(temperature=0.2, model="gpt-4", max_tokens=max_tokens))
        status = st.empty()


        if input_column.button('Generate', key='generate-button'):
            with st.spinner('Reading PDF...'):
                status.text('Processing...')
                pdf = PdfReader(io.BytesIO(uploaded_file.getvalue()))
                text = " ".join(page.extract_text() for page in pdf.pages)
                documents = [Document(text=text)]
                index = VectorStoreIndex.from_documents(documents, service_context=service_context)

            # Determine formality phrase
            if tone_value == 1:
                formality = "In a casual and conversational style, "
            elif tone_value == 2:
                formality = "In a somewhat casual style, "
            elif tone_value == 3:
                formality = "In a neutral style, "
            elif tone_value == 4:
                formality = "In a somewhat formal style, "
            else:
                formality = "In a highly formal and academic style, "

            # Determine action tone phrase
            if action_tone == 1:
                action_formality = "In a casual and conversational style, "
            elif action_tone == 2:
                action_formality = "In a somewhat casual style, "
            elif action_tone == 3:
                action_formality = "In a neutral style, "
            elif action_tone == 4:
                action_formality = "In a somewhat formal style, "
            else:
                action_formality = "In a highly formal and academic style, "

            # Add user context and structure to the query
            if doc_structure == 'Decision Paper':
                query = f"As a {sender}, you will provide {recipient} with a decision paper. The context of the document is: {context}. To create the document, you can refer to the following three elements: source, additional information and call to action. The source is described as {source_description} and is delimited by triple backticks: ```source```. Additional information: {additional_info}. Call to action: {call_to_action}. Deadline: {deadline_date}. Your response will start with firstly “bold(Recommendation:  ) ”, secondly “bold(Deadline:  )  {deadline_date}” and then the three main text sections, “Background”, “Problem” and “Solution”. The writing style of the document should be {tone} with a {technicality} level of technicality. Please provide the response in markdown format with appropriate features."
            else:
                query = f"As {sender}, I need a {doc_structure} of the document for {recipient} that is {length_label} in length and {technicality} in technicality. My goal is {context}. I want the response in English. Please provide the response in markdown format with appropriate features. {tone}"
            # If source_description is provided, add it to the query
            if source_description and (doc_structure != "Decision Paper"):
                query += f" The source document is: {source_description}."


            with response_column:
                # Generate the response
                with st_lottie_spinner(lottie_doc):
                    retriever = VectorIndexRetriever(index=index)
                    query_engine = RetrieverQueryEngine(retriever=retriever)
                    response = query_engine.query(query)
                    # Store the response text in the session state
                    st.session_state['response'] = response.response
                    # Set the flag to True to indicate that the response is ready
                    st.session_state['response_ready'] = True
                status.text('Done processing.')

    # Display the response stored in the session state
    if 'response' in st.session_state:
        response_text = st.session_state['response']
        response_column.markdown(response_text)
        # Convert markdown to HTML
        html = markdown.markdown(response_text)

        # Parse HTML
        soup = BeautifulSoup(html, 'html.parser')

        # Create a new Document
        doc = DocxDocument()

        # Add each paragraph to the document
        for element in soup:
            if element.name == 'p':
                doc.add_paragraph(element.text)
            elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                doc.add_heading(element.text, level=int(element.name[1]))
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='ListBullet')
            elif element.name == 'ol':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='ListNumber')
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        # Add a download button for the response
        response_column.download_button("Download OnePager", buf.getvalue(), file_name="OnePager.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")