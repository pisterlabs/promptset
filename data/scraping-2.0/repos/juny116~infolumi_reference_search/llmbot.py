from langchain.prompts.chat import (
    ChatPromptTemplate,
    SystemMessagePromptTemplate,
    AIMessagePromptTemplate,
    HumanMessagePromptTemplate,
)
from langchain.schema import (
    messages_from_dict,
    messages_to_dict,
)

import discord
import requests
from discord.utils import escape_mentions, remove_markdown
import json
from models import load_model
from utils import (
    remove_mention,
    convert_id_to_name,
    convert_name_to_id,
    add_author,
    remove_author,
)
import hydra
from omegaconf import DictConfig
from time import sleep, time
import asyncio
import docx
import json
import requests
import xmltodict
import tempfile


def GetMedlinePage(start, end):
    if start == end:
        return start

    start_with_zero = (len(end) - len(start)) * "0" + start
    same_cnt = 0
    for e, s in zip(end, start_with_zero):
        if e == s:
            same_cnt += 1
        else:
            break
    medline_page = f":{start}-{end[same_cnt:]}"
    return medline_page


class LLMBot(discord.Client):
    def __init__(self, config, intents):
        super().__init__(heartbeat_timeout=60, intents=intents)
        self.config = config
        self.model = load_model(self.config["model"])

    async def on_ready(self):
        print("Logged on as {0}!".format(self.user))
        await self.change_presence(
            status=discord.Status.online, activity=discord.Game("대기중")
        )

    async def on_message(self, message):
        # Do not respond to ourselves
        if message.author == self.user:
            return
        # Do not respond to other system messages
        if message.content.startswith("***"):
            return
        # Only respond to the message if it is sent to the bot by mentioning
        if self.user.mentioned_in(message):
            # Check if the maximum number of turns has been reached
            thread = await message.create_thread(name="reply", auto_archive_duration=60)
            await thread.send("waiting for file to download...")
            for x in message.attachments:
                file_name = x.url.split("/")[-1]
                print(x.url)
                print(f"{self.config['temp_path']}/{file_name}")
                file_name = f"{self.config['temp_path']}/{file_name}"
                await x.save(file_name)
                self.loop.create_task(self.test(file_name, thread))
                break
            await thread.send("STARTED")

    async def test(self, fname, thread):
        # check the total time
        start_time = time()
        channel = self.get_channel(self.config["discord"]["channel_id"])
        system_message_prompt = SystemMessagePromptTemplate.from_template(
            self.config["template"]["system"]
        )
        parse_template = self.config["template"]["parse"]
        human_message_prompt = HumanMessagePromptTemplate.from_template(parse_template)
        chat_prompt = ChatPromptTemplate.from_messages(
            [system_message_prompt, human_message_prompt]
        )
        ref_list = []
        ref_size = 5
        templ = docx.Document(fname)
        for x, paragraph in enumerate(templ.paragraphs):
            ref_list.append(paragraph.text)
        tasks = []
        # TODO: limit the max task per single run
        max_turn = int(len(ref_list) / ref_size) + 1
        for i in range(max_turn):
            sub_list = ref_list[i * ref_size : (i + 1) * ref_size]
            temp = chat_prompt.format_prompt(
                references="\n".join(sub_list)
            ).to_messages()
            tasks.append(self.parse_reference(temp, channel, i, max_turn, thread))
            break
        results_list = await asyncio.gather(*tasks)
        dummy = {
            "authors": ["dummy"],
            "title": "dummy",
            "journal": "dummy",
            "year": 2023,
            "month": None,
            "day": None,
            "volume": None,
            "issue": None,
            "start_page": None,
            "end_page": None,
        }
        # try to json load the results if fail then add dummy
        for i, results in enumerate(results_list):
            try:
                results = json.loads(results)
            except:
                results = [dummy for i in range(ref_size)]

        results_list = [json.loads(results) for results in results_list]
        results_list = [item for sublist in results_list for item in sublist]
        uid_list = []
        for i, reference in enumerate(results_list):
            if reference == dummy:
                uid_list.append(None)
            else:
                uid_list.append(self.search_pubmed(reference))
                await asyncio.sleep(0.3)

        uid_without_none = [x for x in uid_list if x is not None]
        uid_to_article = self.fetch_pubmed(uid_without_none)
        results = self.revise_reference(uid_list, uid_to_article)

        fp = tempfile.NamedTemporaryFile(mode="w", delete=False, suffix=".docx")
        fname = fp.name
        fp.close()
        doc = docx.Document()
        for r in results:
            doc.add_paragraph(r)
        doc.save(fname)

        await thread.send(
            f"Done in {(time()-start_time):.2} seconds",
            file=discord.File(fname),
        )

    async def parse_reference(self, messages, channel, index, max_turn, thread):
        results = await self.model.agenerate(messages)
        await thread.send(f"parsing references {index+1}/{max_turn} done")
        return results

    def search_pubmed(self, reference):
        params = {
            "method": "auto",
            "authors": reference["authors"][0],
            "title": reference["title"],
            "pdat": str(reference["year"]),
            "volume": reference["volume"],
            "journal": reference["journal"],
        }
        try:
            res = requests.get(
                "https://pubmed.ncbi.nlm.nih.gov/api/citmatch/?", params=params
            )
            result = res.json().get("result")
            uids = result.get("uids")
            if not uids:
                return None
            # return the first uid
            for uid in uids:
                for k, v in uid.items():
                    if k == "pubmed":
                        return v
        except:
            return None

        return None

    def fetch_pubmed(self, uid_list):
        uid_string = ",".join(uid_list)
        data = {"db": "pubmed", "id": uid_string, "retmode": "xml"}
        res = requests.post(
            "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?", data=data
        )
        results = xmltodict.parse(res.text)

        fetch_dict = {}
        for k, uid in enumerate(uid_list):
            fetch_dict[uid] = results["PubmedArticleSet"]["PubmedArticle"][k][
                "MedlineCitation"
            ]["Article"]

        return fetch_dict

    def revise_reference(self, uid_list, uid_to_article):
        revised_refs = []
        for i, uid in enumerate(uid_list):
            max_authors = False
            revised = f"{i+1}. "
            if uid is None:
                revised_refs.append(revised + "Not Found")
            else:
                authors = uid_to_article[uid]["AuthorList"]["Author"]
                title = uid_to_article[uid]["ArticleTitle"]
                journal = uid_to_article[uid]["Journal"]
                page = uid_to_article[uid]["Pagination"]

                if len(authors) > 6:
                    authors = authors[:3]
                    max_authors = True

                if type(authors) == list:
                    authors = [
                        f"{author['LastName']} {author['Initials']}"
                        for author in authors
                    ]
                elif type(authors) == dict:
                    authors = [f"{authors['LastName']} {authors['Initials']}"]
                revised += ", ".join(authors)
                if max_authors:
                    revised += ", et al"

                revised += f". {title} "
                revised += f"{journal['ISOAbbreviation']}. "
                revised += f"{journal['JournalIssue']['PubDate']['Year']}"
                if journal["JournalIssue"].get("Volume"):
                    revised += f";{journal['JournalIssue']['Volume']}"
                if journal["JournalIssue"].get("Issue"):
                    if journal["JournalIssue"].get("Volume") is None:
                        revised += f";({journal['JournalIssue']['Issue']})"
                    else:
                        revised += f"({journal['JournalIssue']['Issue']})"

                if page.get("EndPage"):
                    revised += GetMedlinePage(
                        page.get("StartPage"), page.get("EndPage")
                    )
                else:
                    revised += f":{page.get('StartPage')}"
                revised += "."
                revised_refs.append(revised)
        return revised_refs


@hydra.main(version_base=None, config_path="conf", config_name="config")
def main(config: DictConfig) -> None:
    intents = discord.Intents.all()
    client = LLMBot(config=config, intents=intents)
    client.run(config["discord"]["token"])


if __name__ == "__main__":
    main()
