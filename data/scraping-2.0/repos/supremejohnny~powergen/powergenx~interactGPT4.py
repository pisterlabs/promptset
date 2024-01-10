import os
from docx import Document
import openai
from IPython.display import display, Markdown
import fitz  # PyMuPDF
import re

def read_word_document(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


def read_pdf_document(pdf_path):

    pdf_document = fitz.open(pdf_path)
    text_content = []

    # Iterate over each page in the PDF
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        # Extract text from the page
        text_content.append(page.get_text())

    pdf_document.close()
    return '\n'.join(text_content)

def user_input_configurator():
    print("Welcome to the PowerPoint Presentation Configurator")

    # Number of slides
    slide_count = int(input("Enter the number of slides (1-10): "))

    # Text style
    text_style_options = ['A) concise', 'B) medium', 'C) detailed']
    print("Choose text style (A/B/C):", ', '.join(text_style_options))
    text_style_choice = input().upper()
    text_style_map = {'A': 'concise', 'B': 'medium', 'C': 'detailed'}
    text_style = text_style_map.get(text_style_choice, 'concise')

    # Duration
    duration = int(input("Enter the presentation duration in minutes (1-10): "))

    # Theme
    theme_options = ['A) default', 'B) modern', 'C) classic', 'D) colorful']
    print("Choose a theme (A/B/C/D):", ', '.join(theme_options))
    theme_choice = input().upper()
    theme_map = {'A': 'default', 'B': 'modern', 'C': 'classic', 'D': 'colorful'}
    theme = theme_map.get(theme_choice, 'default')

    # Layout
    layout_options = ['A) standard', 'B) creative', 'C) professional', 'D) minimalistic']
    print("Choose a layout (A/B/C/D):", ', '.join(layout_options))
    layout_choice = input().upper()
    layout_map = {'A': 'standard', 'B': 'creative', 'C': 'professional', 'D': 'minimalistic'}
    layout = layout_map.get(layout_choice, 'standard')

    # Generate content based on user inputs
    presentation_code = generate_presentation_content(
        proposal_text,
        slide_count=slide_count,
        text_style=text_style,
        duration=duration,
        theme=theme,
        layout=layout
    )

    return presentation_code

import requests

def search_image(search_query, api_key, file_path):
    base_url = "https://api.pexels.com/v1/search"
    headers = {"Authorization": api_key}
    params = {"query": search_query, "per_page": 1}

    response = requests.get(base_url, headers=headers, params=params)

    if response.status_code == 200:
        results = response.json()
        if results["photos"]:
            image_url = results["photos"][0]["src"]["original"]
            # Download the image
            img_data = requests.get(image_url).content
            with open(file_path, 'wb') as handler:
                handler.write(img_data)
            return file_path
        else:
            return "No images found."
    else:
        return "Error in API request."

from openai import OpenAI

client = OpenAI(
    api_key="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"
)


def generate_presentation_content(text, slide_count, text_style, duration, theme, layout, model="gpt-4"):
    content_prompt = (f"Create a {text_style}, engaging PowerPoint presentation outline for a proposal with the following guidelines:\n\n"
                      f"{text}\n\n"
                      "Guidelines:\n"
                      "- Simplify and limit words on each screen. Use key phrases and include only essential information.\n"
                      "- Limit punctuation and avoid all-capital letters. Use empty space to enhance readability.\n"
                      "- Use contrasting colors for text and background, preferably light text on a dark background.\n"
                      "- Clearly label each screen. Use a larger font (35-45 points) or different color for the title.\n"
                      "- For bullet points, adhere to the 6 x 6 Rule: One thought per line, no more than 6 words per line, and no more than 6 lines per slide.\n"
                      "- Avoid abbreviations and acronyms.\n\n"
                      f"The presentation should contain {slide_count} slides, including a title slide. Each slide should be clearly defined. The presentation time is {duration} minutes. Theme: {theme}, Slide Layout: {layout}. Provide a title and content for each slide.")

    reference_code = """
                      a = Poweregen()
                      a.add_title(title = "Self Intro", subtitle="Yuquan Gan\nlevel iii of Computer Science")
                      a.add_content_bullet("Myself", "Yuquan Gan\nJohnny\n3rd year", font_size=24)
                      a.add_content_subbullet("Education", ['School', 'Hey'], [['No.3', 'International'], ['test1', 'test2']], double_spacing=True, font_size=24)
                      a.add_content_bullet("Hobbies", "Video Games!", font_size=24)
                      a.add_content_bullet("Hobbies cont", "The hitbox in Street Fighter\nThe ragdoll system in video games", font_size=24)
                      a.add_comparison_slide('Video games', left_title='FPS', left_content='fast pace\nStrong hit\nStrong feedback', right_title='RPG', right_content='Low pace\nGrow up\nExperience\nStoryline', font_size=24)
                      a.add_content_bullet("Special Thanks", "Enlightment\nParents\nMy audience", double_spacing=True, font_size=24)
                      a.save("Selfintro")
                      """

    code_prompt = (f"Now, generate only Python code (without text) using the 'Poweregen' class to create this PowerPoint presentation based on the above specifications and the reference code:\n\n{reference_code}\n\n"
                   "Adjust the code to fit the presentation outline generated previously. You should restrictly follow the rules of reference code.\n\n"
                   "You can't include any list in the add_content_bullet. You will need to use '\n' to divide each bullet point for add_content_bullet (subbullet doesn't work like this). Your generation should at least have one comparison slide and subbullet content slide.\n\n"
                   "You should adjust the font_size to make all the text fit well in the silde.")

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": code_prompt},
            {"role": "user", "content": content_prompt}
        ],
        temperature=0.7
    )
    return response.choices[0].message.content


def generate_image_ideas_and_complete_code(existing_code, model="gpt-4"):
    updated_code = existing_code
    updated_code_lines = []
    image_paths = []
    slide_number = 0

    code_lines = existing_code.split('\n')

    # Iterate through the lines to find slide-creating commands
    for line in code_lines:
        line = line.strip()
        if line.startswith('a.add_title') or line.startswith('a.add_content_bullet') or line.startswith(
                'a.add_content_subbullet') or line.startswith('a.add_comparison_slide'):
            slide_number += 1

        updated_code_lines.append(line)

        # Generate and append image addition code after specific slide-creating lines
        if slide_number and line.startswith('a.add_content_bullet'):
            image_idea_prompt = f"Based on the content of slide {slide_number} from the presentation, suggest an idea for a related image. Your idea should be between 2-7 words. And try to be creative."
            image_idea_response = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": image_idea_prompt}],
                temperature=0.7
            )
            image_idea = image_idea_response.choices[0].message.content

            # Assume search_image function returns a path to the downloaded image
            file_path = f"slide_{slide_number}.jpg"
            downloaded_image_path = search_image(image_idea, api_key, file_path)
            image_paths.append(downloaded_image_path)

            # Append the code to add an image slide at the correct position
            updated_code_lines.append(
                f"a.add_picture_slide({slide_number}, '/content/{downloaded_image_path}', 'right lower corner/mid lower/ left lower corner')")

    # Join the updated code lines into a single string
    updated_code = '\n'.join(updated_code_lines)
    # Prompt GPT-4 to generate the complete code with image integrations
    complete_code_prompt = (
        f"Revise and output the final Python code and output only python code to create a PowerPoint presentation using the 'Poweregen' class, integrating the downloaded images appropriately based on this existing code:\n\n{updated_code}\n\n"
        "You should make sure that the picture will not influence the text that already on the slide. So you can adjust the font_size of those silde, and also adjust the third parameter of add_picture:right lower corner/mid lower/ lift lower corner.\n\n"
        "You should apply all the picture on the corresponding slide that have provided to you.\n\n"
        "The first parameter of add_picture_slide is the number of slide you will add picture, so the file number in the second parameter is also the silde number. \n\n"
        "The a.save() should be at the last line of the entire code \n\n"
        "Output only the Python code without any text \n\n"
        "add_picture_slide should apply after the conclusion slide but before the a.save()")
    complete_code_response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": complete_code_prompt}
        ],
        temperature=0.7
    )
    complete_code = complete_code_response.choices[0].message.content

    return complete_code, image_paths


def run_PowerGenX_code(code_block):
    # Use a regular expression to find the code starting with 'a = Poweregen()'
    # and remove Markdown backticks
    code_to_run = re.sub(r'^```python\n', '', code_block, flags=re.MULTILINE)
    code_to_run = re.sub(r'```$', '', code_to_run, flags=re.MULTILINE)

    # Make sure we start executing from the instantiation of Poweregen
    match = re.search(r'a = Poweregen\(\)', code_to_run)
    if match:
        code_to_run = code_to_run[match.start():]  # Extract from 'a = Poweregen()' to the end
        try:
            # Execute the extracted code
            exec(code_to_run, globals())
            print("Code executed successfully.")
        except Exception as e:
            print(f"Error executing the code: {e}")
    else:
        print("Could not find the starting point of Poweregen instantiation.")
