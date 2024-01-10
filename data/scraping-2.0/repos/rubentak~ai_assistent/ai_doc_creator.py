
# --- MANUAL TEMPLATE ---
# Import libraries
from docx import Document
from docx.shared import Inches
import openai
import credentials
import os

script_dir = os.path.dirname(os.path.abspath(__file__))
startfile = os.path.join(script_dir, "template.docx")

#%%
# --- CHATGPT PROMPT  ---
header = input("\nTitle of the document: ")
openai.api_key = credentials.api_key
number_paragraphs = int(input("Number of Paragraphs?: "))

docu_list = []

for i in range(number_paragraphs):
    prompt = input(f"Question {i+1}/{number_paragraphs}: ")
    docu_list.append(prompt)

def gpt_docu(prompt):
    try:
      print(f"----- {prompt} -----")
      completion = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
          {"role": "user", "content": prompt}
        ]
      )

      answer = completion.choices[0].message.content
      for i in range(2):
          try:
              if answer[0] == "\n":
                  answer = answer[1:]
          except:
              pass


      return answer
    except:
      print("Something went wrong. Please try again.")


# append all answers to a list
answer_list = []
try:
    for i in range(len(docu_list)):
        print(f"\n----- Generating answer for: Question {i+1}/{len(docu_list)} -----")
        answer = gpt_docu(docu_list[i])
        answer_list.append(answer)
        print(f"----- Answer : {answer} -----\n")
except:
    print("Something went wrong. Please try again.")

#%%
# --- CREATE DOCUMENT ---
# Open template document `hello_world.docx`
document = Document(startfile)
# Clear document
document._body.clear_content()

# Add title
document.add_heading(header, 1)

for i in range(len(answer_list)):
    # Add header first
    document.add_heading(docu_list[i], 2)
    # Add paragraph
    p = document.add_paragraph(answer_list[i])

# Save document with the first 4 words from the header if the exist.

try:
    document.save('/Users/erictak/Desktop/' + header.split()[0] +' '+  header.split()[1] +' '+ header.split()[2] + header.split()[3] + '.docx')
except:
    try:
        document.save('/Users/erictak/Desktop/' + header.split()[0] +' '+  header.split()[1] +' '+ header.split()[2] + '.docx')
    except:
        try:
            document.save('/Users/erictak/Desktop/' + header.split()[0] +' '+  header.split()[1] + '.docx')
        except:
            try:
                document.save('/Users/erictak/Desktop/' + header.split()[0] + '.docx')
            except:
                document.save('/Users/erictak/Desktop/Untitled.docx')

