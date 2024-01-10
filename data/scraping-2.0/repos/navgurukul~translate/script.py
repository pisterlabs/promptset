import sys
import openai
from docx import Document
from dotenv import dotenv_values
import csv
import os
import time
import random
from docx.shared import RGBColor

# Imports the Google Cloud Translation library
from google.cloud import translate

# Load the environment variables from the .env file
env_vars = dotenv_values('.env')
openai.api_key = env_vars.get('OPENAI_API_KEY')
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = env_vars.get('GOOGLE_APPLICATION_CREDENTIALS')

pre_replacements_file = 'pre-phrases.csv'
post_replacements_file = 'post-phrases.csv'

# Read the pre and post replacements from the csv files
with open(pre_replacements_file, 'r', encoding='utf-8') as csvfile:
    reader = csv.reader(csvfile)
    pre_replacements = {row[0]: row[1] for row in reader}

with open(post_replacements_file, 'r', encoding='utf-8') as csvfile:
    reader = csv.reader(csvfile)
    post_replacements = {row[0]: row[1] for row in reader}

# Read the secret key from the .env file
def read_secret_key():
    # Load the environment variables from the .env file
    env_vars = dotenv_values('.env')
    # Access the SECRET_KEY variable
    secret_key = env_vars.get('SECRET_KEY')
    return secret_key

# Replace the phrases in the text with the pre-defined replacements before translation
def pre_replace_phrases(text):
    for phrase, replacement in pre_replacements.items():
        text = text.replace(phrase, replacement)

    return text

# Replace the phrases in the text with the replacements post automatic translation
def post_replace_phrases(text):
    for phrase, replacement in post_replacements.items():
        text = text.replace(phrase, replacement)

    return text

# Initialize Translation client
def translate_paragraph_google(paragraph, target_language='hindi') -> translate.TranslationServiceClient:
    """Translating Text."""

    project_id = "chanakya-259818"

    if target_language == 'hindi':
        target_language = 'hi'

    client = translate.TranslationServiceClient()
    location = "global"
    parent = f"projects/{project_id}/locations/{location}"

    # Translate text from source_language to target_language
    # Detail on supported types can be found here:
    # https://cloud.google.com/translate/docs/supported-formats
    response = client.translate_text(
        request={
            "parent": parent,
            "contents": [paragraph],
            "mime_type": "text/plain",  # mime types: text/plain, text/html
            "source_language_code": "en-US",
            "target_language_code": target_language,
        }
    )

    return response.translations[0].translated_text

def translate_paragraph_gpt(paragraph, target_language='hindi'):
    # Prepare the system message
    # system_message = {
    #     "role": "system",
    #     "content": "
    # }

    if paragraph.strip() == '':
        return ''
    
    PROMPT = env_vars.get('GPT_PROMPT').replace('TARGET_LANGUAGE', target_language)

    # Prepare the user message with the paragraph to translate
    user_message = {
        "role": "user",
        "content": PROMPT + paragraph
    }

    # Generate translation using ChatGPT
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[user_message],
            max_tokens=3900,
            temperature=0.7,
            n=1,
            stop=None,
        )
    except Ellipsis as e:
        print("An error occured: ", str(e))
        sleep_duration = random.randint(5, 10)
        time.sleep(sleep_duration)
        return translate_paragraph_gpt(paragraph, target_language='hindi')
    # Retrieve the translated text from the API response

    translated_text = response.choices[0].text.strip()
    return translated_text

def break_paragraph_into_subparagraphs(paragraph, max_length=2000):
    # Break the paragraph into sub-paragraphs of maximum length
    sub_paragraphs = []
    while len(paragraph) > max_length:
        # Find the last period within the maximum length
        last_period_index = paragraph.rfind(".", 0, max_length)
        if last_period_index == -1:
            # If no period is found within the maximum length, just split at the max length
            sub_paragraph = paragraph[:max_length]
            paragraph = paragraph[max_length:]
        else:
            # Split at the last period found
            sub_paragraph = paragraph[:last_period_index + 1]
            paragraph = paragraph[last_period_index + 1:]
        sub_paragraphs.append(sub_paragraph)

    # Append the remaining part of the paragraph, which is now less than the max_length
    if paragraph:
        sub_paragraphs.append(paragraph)

    return sub_paragraphs

def translate_paragraph(model,sub_paragraph, target_language):
    # Translate the sub-paragraph based on the model
    if model=='gpt':
        return translate_paragraph_gpt(sub_paragraph, target_language)
    elif model=='google':
        return translate_paragraph_google(sub_paragraph, target_language)
    else:
        raise Exception('Invalid model')

model_colors = {
    'gpt': RGBColor(25, 25, 112),
    'google': RGBColor(0, 128, 0)
}

def output_translation(input_path, output_path, target_language='en'):
    # Load the input document
    doc = Document(input_path)

    # Create a new document for the output
    output_doc = Document()

    # Iterate over paragraphs in the input document
    for paragraph in doc.paragraphs:
        text=paragraph.text.strip()
        if text=='':
            continue

        # Replace phrases in the paragraph before translation
        replaced_paragraph = pre_replace_phrases(text)
        sub_paragraphs = break_paragraph_into_subparagraphs(replaced_paragraph, max_length=2000)

        # Translate each sub-paragraph
        for sub_paragraph in sub_paragraphs:
            # Add the original sub-paragraph to the output document
            output_doc.add_paragraph(sub_paragraph)

            # Translate the sub-paragraph using each model
            for model in model_colors.keys():
                # Add the translated sub-paragraph to the output document
                translated_sub_paragraph = translate_paragraph(model,sub_paragraph, target_language)
                translated_sub_paragraph = post_replace_phrases(translated_sub_paragraph).replace('\n', '')
                run = output_doc.add_paragraph().add_run()
                run.text = translated_sub_paragraph
                font = run.font
                font.color.rgb = model_colors[model]

    # Save the output document
    output_doc.save(output_path)

if __name__ == '__main__':
    # Check if the correct number of command line arguments is provided
    if len(sys.argv) < 3:
        print("Usage: python script.py input_file output_file [target_language]")
        sys.exit(1)

    # Extract the input and output file names from command line arguments
    input_file = sys.argv[1]
    output_file = sys.argv[2]

    # Extract the target language if provided, default to 'en' if not specified
    target_language = sys.argv[3] if len(sys.argv) > 3 else 'hindi'

    # Call the function to output the translation
    output_translation(input_file, output_file, target_language)