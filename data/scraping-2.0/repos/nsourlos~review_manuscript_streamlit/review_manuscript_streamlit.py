#Import dependencies
from IPython.display import display, Markdown
import os
import streamlit as st
import pdfplumber
import docx
import io

#Used only in env file provided - not here
from dotenv import load_dotenv
load_dotenv()
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY', 'YourAPIKeyIfNotSet')

from langchain import OpenAI
from langchain.text_splitter import TokenTextSplitter

# Start Of Streamlit page
st.set_page_config(page_title="Paper review assistant", page_icon=":robot:")

# Start Top Information - List of emojis in #https://streamlit-emoji-shortcodes-streamlit-app-gwckff.streamlit.app/
st.header(":pencil2: LLM Assisted Paper Review")
col1, col2 = st.columns(2)

with col1:
    st.markdown("Have a manuscript to review? Want to Q&A with it? This tool is meant to help you with these things! \
                \n\nThere is a limit of 16k tokens (with response). \
                Takes only pdf files as input. Click 'Generate Output' after uploading file. \
                If you want to only know the price, just upload the document (no key needed) and click the above button. \
                \n\nThis tool is made  by [Nikos Sourlos](https://linkedin.com/in/nsourlos). \
                \n\nSource Code on [Github](https://github.com/nsourlos/review_manuscript_streamlit/blob/main/review_manuscript_streamlit.py)")

with col2:
      st.image( #Idea taken from https://discuss.streamlit.io/t/how-to-centre-and-resize-gifs-in-streamlit/18333/4
            "https://raw.githubusercontent.com/nsourlos/review_manuscript_streamlit/main/paper_review.gif", 
            width=300, #Change width of GIF manually - using the column-width parameter would make it weirdly big. 
        ) 
# End Top Information

st.markdown("## :muscle: Upload PDF document")

uploaded_file = st.file_uploader("Choose a pdf file", type="pdf") #Upload file based on https://discuss.streamlit.io/t/how-to-upload-a-pdf-file-in-streamlit/2428/2

#Show a box to be filled with the OpenAI API Key by the user
OPENAI_API_KEY = st.text_input(label="OpenAI API Key (or set it as .env variable)",  placeholder="Ex: sk-2twmA8tfCb8un4...", key="YourAPIKeyIfNotSet")

#Click one of options
prompt_text = st.radio(
    "Task:",
    ('Generate Review Questions', 'Q&A'))

if prompt_text == 'Generate Review Questions':

    #Insert placeholder text for prompt and format it properly
    prompt = """You are a experienced reviewer of scientific manuscripts. You provide concise feedback on the
                manuscript as well as specific suggestions for things that should be modified based on the content of it. Provide at least 10 suggestions
                tailored to the content of the specific manuscript. Avoid general remarks and give specific recommendations on what should change.
                Explain why what is already written is not sufficient and expand each point raised by providing ways to improve.
                The scientific manuscript is: """
    prompt=prompt.replace('                ','').replace('\n',' ')

elif prompt_text == 'Q&A':

    prompt=""" You are a helpful chatbot that helps the user answer a question with information obtained from a scientific manuscript. The question is: """

    #Insert placeholder text for prompt and format it properly
    placeholder_text = """How to mitigate biases in AI algorithms?"""
    
    #Show a box to be filled in with prompt to be sent to OpenAI
    question = st.text_area('Enter Question Below:', 
                            height=300, 
                            placeholder=placeholder_text) 


button_ind = st.button("*Generate Output*", type='secondary', help="Click to generate answer")
# st.write(':heavy_exclamation_mark: Please refresh page everytime you click that button')

#Select options from dropdown box
# download_option = st.selectbox("Download questions as docx?",('Yes', 'No'),index=1) #https://stackoverflow.com/questions/65026852/set-default-value-for-selectbox

if button_ind: #When button is clicked
    
    if uploaded_file is None: #If no file is uploaded raise an error and stop execution
        st.warning('Please provide a PDF file', icon="⚠️")
        st.stop()

    #Information on how to load uploaded file in https://docs.streamlit.io/library/api-reference/widgets/st.file_uploader
    #For pdfs https://discuss.streamlit.io/t/how-to-upload-a-pdf-file-in-streamlit/2428
    st.write("Loading PDF...")
    
    #https://github.com/jsvine/pdfplumber/issues/147
    pdf_file=pdfplumber.open(uploaded_file) #No other way to load PDF since it's only loader in buffer memory and not saved in disk
       
    docs = []
    for page in pdf_file.pages: #Add all documents to one
        docs.append(page.extract_text(x_tolerance=1)) #Tolerance to have proper spacing as stated in https://github.com/jsvine/pdfplumber/issues/334
    pdf_file.close()
    paper=''.join(docs) #Combine all documents
    
    #Calculate token usage and price based on paper length - Same as to send to OpenAI but free
    st.write("Calculating price...")
    text_splitter = TokenTextSplitter(chunk_size=1, chunk_overlap=0,model_name='gpt-3.5-turbo-16k')
    ind_tokens=text_splitter.split_text(paper)
    st.write("Total tokens:",len(ind_tokens))
    st.write("Price for them:",round(len(ind_tokens)*0.003/1000,4),"$")
    st.write("Loading LLM output...")

    if not OPENAI_API_KEY: #If no OpenAI API key is set raise an error and stop execution
        st.warning('Please insert OpenAI API Key. Instructions [here](https://help.openai.com/en/articles/4936850-where-do-i-find-my-secret-api-key)', icon="⚠️")
        st.stop()

    if prompt_text=='Generate Review Questions': 
        review_prompt=prompt
    elif prompt_text=='Q&A': 
        review_prompt=prompt+question+' The scientific manuscript is: '

    llm=OpenAI(openai_api_key=OPENAI_API_KEY,temperature=0,model_name='gpt-3.5-turbo-16k') #Initialize LLM - 16k context length to fit the whole paper
    st.write('Prompt given to LLM:',review_prompt)
    questions_final=llm.predict(review_prompt+paper) #Predict response using LLM 
    st.markdown(f"#### LLM Output:")
    st.write(questions_final)

    #Create a docx to be saved if user clicks button below
    document = docx.Document() #Create word document
    document.add_heading('https://review-paper.streamlit.app/', level=1) #Add title
    document.add_heading("Prompt Used:",level=3) #Add text
    document.add_paragraph(review_prompt) #Add text
    document.add_heading("LLM Output:",level=3) #Add text
    document.add_paragraph(questions_final) #Add text
    
    #https://discuss.streamlit.io/t/downloading-a-ms-word-document/28850/3
    bio = io.BytesIO() #Create BytesIO to save docx from memory
    document.save(bio) #Save document from memory
    if document: #If document is created
        st.download_button( #Create download button
            label="Download as 'docx'",
            data=bio.getvalue(),
            file_name="LLM_output.docx",
            mime="docx"
        )

#Notes: Streamlit gives 1CPU, 1GB of RAM and 1GB of disk space (https://discuss.streamlit.io/t/problem-on-resources-limit/12605)
#Streamlit implementation adapted from https://github.com/gkamradt/llm-interview-research-assistant/blob/main/main.py
#Deployment based on https://docs.streamlit.io/streamlit-community-cloud/get-started/deploy-an-app