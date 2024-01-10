# Importing modules 
import streamlit as st
from streamlit import session_state as ss
import re
from helfunc import extract_data , parse_string , disGen , file_upload_check , load_api_key , save_api_key , parse_C_string
import os
# Importing Python Docx Reader
from docx import Document
# Importing PyPanDoc
import pypandoc
# Importing Tempfile
import tempfile
# Importing Beautiful Soup
from bs4 import BeautifulSoup
# importing langchain modules
from langchain.llms import OpenAI
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain

# Setting page config to wide mode
st.set_page_config(layout="wide")

# Load API key from the configuration file
saved_api_key = load_api_key()

# Global variable to store the API key
api_key = st.text_input("Enter your OpenAI API key:", type="password", key="api_key", value=saved_api_key)

# If the user provided an API key, save it to the configuration file
if api_key and api_key != saved_api_key:
    save_api_key(api_key)
    os.environ['OPENAI_API_KEY'] = api_key

# Title for the web app with the OpenAI API key
st.title(f'MCQ Generator')

# Uploading the file
uploaded_file = st.file_uploader("Upload Your Files",type=['docx'])


# Creating a prompt template
genQtemplate= """I have questions in specific format and you have to generate and return new innovative practice question and removing the year from question section from same sub-topic with appropriate content even if it's not there in what i have sent to you for students in json format , keys format should be strictly same , keep sub-topic same : {question}"""
# template= """I have questions in specific format and you have to generate new innovative practice question from same sub-topic with appropriate content even if it's not there in what i have sent to you for students in json format , keys format should be strictly same : {question}"""
# genCtemplate= """I have questions in specific format and you have to correct the content and return the corrected question in json format only , improve the hint or explanation if you feel it's not correct , keys format should be strictly same , keep sub-topic same , don't send additional data like description of what you have done : {question}"""
genCtemplate= """I have questions in specific format and you have to correct the hint and explanation only and return the corrected question in json format only , rephrase the hint or explanation if you feel it's not correct , keys format should be strictly same , keep sub-topic same , don't send additional data like description of what you have done : {question}"""
genAdtemplate= """I have questions in specific format and you have to generate and return new innovative practice question and removing the year from question section from same sub-topic but it should involve some advanced concepts or more tough with appropriate content even if it's not there in what i have sent to you for students in json format , keys format should be strictly same , keep sub-topic same : {question}"""


llm = ChatOpenAI(model="gpt-4-1106-preview",openai_api_key=ss.api_key)  # os.environ['OPENAI_API_KEY'])
# prompt = PromptTemplate(template=genQtemplate, input_variables=["question"])
# llm_chain = LLMChain(prompt=prompt, llm=llm)

# Using ss to store and retrieve generated content
if 'generatedContent' not in ss:
    ss.generatedContent = None

# Checking if the file is uploaded or not
if uploaded_file:

    if file_upload_check(uploaded_file):
        # By Default creating a Document
        if 'doc' not in ss:
            ss.doc=Document()
            ss.genQs=-1 
    else:
        del ss.doc
        ss.doc=Document()
        ss.genQs=-1
        
    # Creating a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        # Write the uploaded file's content to the temporary file
        tmp_file.write(uploaded_file.getvalue())
        tmp_file_path = tmp_file.name

    # Checking data in text format
    # opening_file=Document(uploaded_file)
    opening_file=pypandoc.convert_file(tmp_file_path, 'html',extra_args=['--quiet'])
    # Converting into Soup HTML format
    soup = BeautifulSoup(opening_file, 'html.parser')
    # Checking the number of tables in the document
    
    configs=soup.find_all('table')[0]
    questions=soup.find_all('table')[1:]
    no_questions=len(questions)
    st.info(f"No of questions in the document are {no_questions}")


    # Creating json file for the questions
    if 'question_set' not in ss:
        question_set={}
        format=["Question","Option A","Option B","Option C","Option D","Correct Answer","Hint","Explanation","Sub-topic"]
        rows = configs.find_all('tr')
        # Creating a dictionary for the configs
        question_set["Config"]={cols[0].text: cols[1].text for row in rows for cols in [row.find_all(['th', 'td'])] if cols}
        # st.info(question_set["Config"])
        for key,values in question_set["Config"].items():
            # Checking in console
            # print(f"{key} : {values}")
            st.write(f"{key} : {values}" )
        # Creating a dictionary for the questions
        # question_set["Questions"]=[dict(zip(format,[(values.text[:-6].replace('\n','') if '(20' in values.text else values.text.replace('\n','')) for values in questions[i].columns[0].cells])) for i in range(no_questions)]    
        question_set["Questions"]=[extract_data(question) for question in questions] 

    # Creating Columns
    ReadQuestion,GenerateQuestion=st.columns([4,5]) 

    # Reading the question and setting up Column 1
    with ReadQuestion:
        q_no = st.number_input(f"Select Question No : Max {no_questions} ", 1, no_questions, 1) 

        # st.write(question_set["Questions"][q_no - 1])
        # print(questions[q_no-1])
        qs=question_set["Questions"][q_no - 1]
        if questions[q_no - 1].find('img') and ' ⁉📷' not in qs['Question']:
            # + f" {'Image 📷 Missing' if question.find('img') else ''}"
            qs['Question'] += ' ⁉📷'
        disGen(qs)
        st.write(f"\n")
        st.write(f"\n")

        match=re.search(r'\(\b20\d{2}\b\)',str(questions[q_no - 1]))
        if match:
            year=f" \t {match.group()} "
        else:
            year=None
        
        # Dividing the Button into five Columns
        gen_similar,gen_corrected,copyit,gen_advanced,addit,download=st.columns(6)

        # Creating Buttons side by side
        with gen_similar:
            similar=st.button("⚙️",help="Generate Similar Question")
            
        with gen_corrected:
            corrected=st.button("✅",help="Generate Corrected Question")
        
        with copyit:
            copied=st.button("📋",help="Copy the same Question into File")
            
        with gen_advanced:
            advanced=st.button("↗️", help="Generate Advanced Question")
        
        with addit:
            addIt=st.button(" ➕ " , help="Add the Generated Question into File")
        
        with download:
            downld=st.button("⏬" , help="Save and Download the File")


    with GenerateQuestion:
        print(f"Gen Qs : {similar} , Correct It : {corrected} , Copied : {copied} Gen Advanced : {advanced} , addIt : {addIt} , downld : {downld}" )
        if ' ⁉📷' in qs['Question']:
            qs['Question'] = qs['Question'].replace(' ⁉📷', '')

        # print(year)
        
        if similar :
            prompt = PromptTemplate(template=genQtemplate, input_variables=["question"])
            llm_chain = LLMChain(prompt=prompt, llm=llm)
            ss.generatedContent = llm_chain.run(str(qs))
            # For logging purpose
            print(f" Similar : {ss.generatedContent}")
            st.write(f"Generated Question : ") 
            json_str=parse_string(ss.generatedContent)
            disGen(json_str) 

        elif corrected:
            prompt = PromptTemplate(template=genCtemplate, input_variables=["question"])
            llm_chain = LLMChain(prompt=prompt, llm=llm)
            ss.generatedContent = llm_chain.run(str(qs))
            print(f" Corrected : {ss.generatedContent}")
            st.write(f"Generated Question : ")
            json_str=parse_C_string(ss.generatedContent)
            disGen(json_str) 

        elif copied:
            ss.genQs+=1
            json_str=qs
            ss.doc.add_table(rows=9, cols=1, style='Table Grid')
            for j, items in enumerate(json_str.items()):
                match=re.search(r'\(\b20\d{2}\b\)', items[1])
                if match:
                    ss.doc.tables[ss.genQs].rows[j].cells[0].text = items[1].replace(match.group(), f'\n{match.group()}')
                else:
                    ss.doc.tables[ss.genQs].rows[j].cells[0].text = items[1]

            ss.doc.add_paragraph()
            with copyit:
                st.write(" Qn 📋ed")
            print(f"Table Added")

        elif advanced:
            prompt = PromptTemplate(template=genAdtemplate, input_variables=["question"])
            llm_chain = LLMChain(prompt=prompt, llm=llm)
            ss.generatedContent = llm_chain.run(str(qs))
            print(f" Advanced : {ss.generatedContent}")
            st.write(f"Generated Question : ")
            json_str=parse_string(ss.generatedContent)
            disGen(json_str) 

        
        if addIt:
            ss.genQs+=1
            try:
                print(f" Generated Question : {ss.genQs+1} ")
                json_str=parse_string(ss.generatedContent)
                ss.doc.add_table(rows=9, cols=1, style='Table Grid')
                for j, items in enumerate(json_str.items()):
                    match=re.search(r'\(\b20\d{2}\b\)', items[1])
                    if match:
                        ss.doc.tables[ss.genQs].rows[j].cells[0].text = items[1].replace(match.group(), f'\n{match.group()}')
                    else:
                        ss.doc.tables[ss.genQs].rows[j].cells[0].text = items[1]
                ss.doc.add_paragraph()
                # ss.doc.save('generated.docx')
                with addit:
                    st.write("Qs ➕ed ")
                print(f"Table Added")
            except:
                st.warning("Please Generate the Question First")
                ss.genQs-=1
    
        with download:
            if downld:
                st.write("File Saved")
                ss.doc.save('generated.docx')
                st.download_button(label=" ⬇️ ",data=open('generated.docx','rb').read(),file_name='generated.docx',mime='application/octet-stream',)
