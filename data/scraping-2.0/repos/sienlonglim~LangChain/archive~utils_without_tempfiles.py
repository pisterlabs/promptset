import re
from io import StringIO
from PyPDF2 import PdfReader
from docx import Document as docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

def remove_delimiters(document_chunks, config : dict):
    delimiters_to_remove = config['splitter_options']['delimiters_to_remove']
    for chunk in document_chunks:
        for delimiter in delimiters_to_remove:
            chunk.page_content = re.sub(delimiter, ' ', chunk.page_content)
    return document_chunks

def remove_pages(reader : object, config : dict):
    front_pages_to_remove = config['splitter_options']['front_pages_to_remove']
    last_pages_to_remove = config['splitter_options']['last_pages_to_remove']
    # Remove pages
    for _ in range(front_pages_to_remove):
        del reader.pages[0]
    for _ in range(last_pages_to_remove):
        reader.pages.pop()
        print(f'\tNumber of pages after skipping: {len(reader.pages)}')
    return reader

def get_pdf(file : bytes, pdf_title : str, config : dict, splitter : object):
    remove_pages_before = config['splitter_options']['remove_pages']
      
    reader = PdfReader(file)
    for key, value in reader.metadata.items():
        if 'title' in key.lower():
            pdf_title = value
    print(f'\t\tOriginal no. of pages: {len(reader.pages)}')
    if remove_pages_before:
        reader = remove_pages(reader, config)

    # Each file will be split into LangChain Documents and kept in a list
    document_chunks = [] 
    # Loop through each page and extract the text and write in metadata
    if splitter:
        for page_number, page in enumerate(reader.pages):
            page_chunks = splitter.split_text(page.extract_text()) # splits into a list of strings (chunks)
            for chunk in page_chunks: # Write into Document format with metadata
                document_chunks.append(Document(page_content=chunk, metadata={'source': pdf_title , 'page':str(page_number+1)}))
    else:  # This will split purely by page
        for page_number, page in enumerate(reader.pages):
            document_chunks.append(Document(page_content=page.extract_text(),  metadata={'source': pdf_title , 'page':str(page_number+1)}))
            i += 1
    print(f'\t\tExtracted no. of documents: {len(document_chunks)}')

    return pdf_title, document_chunks

def get_txt(file : bytes, title : str, config : dict, splitter : object):
    # file.getvalue() returns bytes which is decoded using utf-8, to instantiate a StringIO. Then we getvalue of the StringIO
    data = StringIO(file.getvalue().decode("utf-8"))
    text = data.getvalue()

    if splitter:
        splits = splitter.split_text(text) 
        document_chunks = []
        for split in splits:
            document_chunks.append((Document(page_content=split,  metadata={'source': title , 'page':'N/A'})))
    else:
        document_chunks = [(Document(page_content=text,  metadata={'source': title , 'page':'N/A'}))]

    print(f'\t\tExtracted no. of documents: {len(document_chunks)}')
    return title, document_chunks

def get_docx(file : bytes, title : str, config : dict, splitter : object):
    word_doc = docx(file)
    text = ''
    for paragraph in word_doc.paragraphs:
        text += paragraph.text

    if splitter:
        splits = splitter.split_text(text) 
        document_chunks = []
        for split in splits:
            document_chunks.append((Document(page_content=split,  metadata={'source': title , 'page':'na'})))
    else:
        document_chunks = [(Document(page_content=text,  metadata={'source': title , 'page':'na'}))]
    print(f'\t\tExtracted no. of documents: {len(document_chunks)}')
    
    return title, document_chunks

def get_chunks(file_input : list, config : dict):
    '''
    Main function to split or load all the documents into chunks
    '''
    remove_leftover_delimiters = config['splitter_options']['remove_leftover_delimiters']
    use_splitter = config['splitter_options']['use_splitter']
    if use_splitter:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=config['splitter_options']['chunk_size'],
            chunk_overlap=config['splitter_options']['chunk_overlap'],
            separators = config['splitter_options']['chunk_separators']
            )
    else:
        splitter = None

    # Main list of all LangChain Documents
    document_chunks_full = []
    document_names = []
    print(f'Splitting documents: total of {len(file_input)}')

    # Handle file by file
    for file_index, file in enumerate(file_input):

        # Get the file type and file name
        file_type = file.type.split('/')[1]
        print(f'\tSplitting file {file_index+1} : {file.name}')
        file_name = file.name.split('.')[:-1]
        file_name = ''.join(file_name)

        # Handle different file types
        if file_type =='pdf':
            title, document_chunks = get_pdf(file, file_name, config, splitter)
        elif file_type == 'txt' or file_type == 'plain':
            title, document_chunks = get_txt(file, file_name, config, splitter)
        elif file_type == 'vnd.openxmlformats-officedocument.wordprocessingml.document':
            title, document_chunks = get_docx(file, file_name, config, splitter)

        # Remove all left over the delimiters and extra spaces
        if remove_leftover_delimiters:
            document_chunks = remove_delimiters(document_chunks, config)

        document_names.append(title)
        document_chunks_full.extend(document_chunks)
             
    print(f'\tNumber of document chunks extracted in total: {len(document_chunks_full)}')
    return document_chunks_full, document_names
