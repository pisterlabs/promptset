"""
this file will act as the main text generator, collecting all the necessary text.
Then a post-processing pipeline will be used, such as error correction and text optimization.
In the final stage, the data will be translated and placed in a user-friendly file such as word (docx).
"""
# base libraries
import requests
import json
import deepl
import random
import os
import openai
import math
import cv2
import numpy
import re
from docx import Document
import threading

# external files
import stencil


# global functions
def translateTo(text: str, target_lang="CS") -> str:
    with open("key.txt", "r") as key:
        for line in key.readlines():
            if "deepl: " in line:
                deepl_key = line.replace("deepl: ", "").replace("\n", "")

    translator = deepl.Translator(deepl_key)
    result = translator.translate_text(text, target_lang=target_lang)
    translated_funfact = result.text

    return translated_funfact


def openai_response(prompt: str, token_lenght=150, ) -> str:
    with open("key.txt", "r") as key:
        for line in key.readlines():
            if "openai: " in line:
                openai.api_key = line.replace("openai: ", "").replace("\n", "")

    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=prompt,
        temperature=0.7,
        max_tokens=token_lenght,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0
    )

    resp =  response['choices'][0]['text']
    # for removing unvanted spaces in front of a response
    for char in resp:
        if char == " " or char == "\n":
            resp = resp[1:]
        else:
            return resp


def text_only(text: str) -> str:
    return re.sub(r'[^a-zA-Z]', '', text)



class bibliographic_data:
    def __init__(self, LN) -> None:
        self.LN = LN
        self.max_words_person_name = 4 # insurance for checking if the name responses are valid
        self.max_words_publisher_name = 5 # insurance for checking if the publisher response is valid
        self.on_error_repeats = 3 # insurance for checking if the publisher response is valid
        self.category = translateTo(stencil.categories[0], language)
        self.output = {
            "Author": "",
            "Illustrator": "",
            "Translator": "",
            "Publisher": "",
            "Page_count": "",
        }

    def Author(self):
        for _ in range(self.on_error_repeats):
            data = openai_response(f"who is the author of {book}, respond with one name, nothing more.")
            if data.count(" ") < self.max_words_person_name:
                self.output["Author"] = data
                return 1
        
        self.output["Author"] = "Unknown"
        return -1

    def Illustrator(self):
        for _ in range(self.on_error_repeats):
            data = openai_response(f"who is the illustrator of {book}, respond with one name, nothing more.")
            if data.count(" ") < self.max_words_person_name:
                self.output["Illustrator"] = data
                return 1
        
        self.output["Illustrator"] = "Unknown"
        return -1

    def Translator(self):
        if not "EN" in self.LN:
            for _ in range(self.on_error_repeats):
                data = openai_response(f"who is the translator of {book} to {self.LN}, respond with one name, nothing more.")
                if data.count(" ") < self.max_words_person_name and data != self.output["Author"]:
                    self.output["Translator"] = data
                    return 1
            
            self.output["Translator"] = "Unknown"
            return -1
        else:
            self.output["Translator"] = self.output["Author"]
            return 1


    def Publisher(self):
        for _ in range(self.on_error_repeats):
            data = openai_response(f"name one book publishing company that published {book} to {self.LN}, respond with one name only.")
            if data.count(" ") < self.max_words_person_name:
                self.output["Publisher"] = data
                return 1
        
        self.output["Publisher"] = "Unknown"
        return -1

    def Page_Count(self):
        for _ in range(self.on_error_repeats):
            data = openai_response(f"give me number of pages of a {self.LN} {book} book, give me one number only, try to find books with lower number.")
            try:
                self.output["Page_count"] = int(data)
                return 1
            except ValueError:
                continue
        
        self.output["Page_count"] = "Unknown"
        return -1

    def complete(self):
        error_rate = [
            self.Author(),
            self.Illustrator(),
            self.Translator(),
            self.Publisher(),
            self.Page_Count(),
        ]
        # print(error_rate)
        error_rate = sum(error_rate)
        return error_rate
        

class author_info:
    def __init__(self, author) -> None:
        self.author = author
        self.output = ""
        self.tokens = 250
        self.category = translateTo(stencil.categories[1], language)

    def get_data(self, result, index):
        data = openai_response(
            f"give me some info on {self.author} using the text provided as reference, make your text a bit longer than the text provided and try to touch the same points as in the reference text.\n'{stencil.stencil_filled[1]}'",
            self.tokens
        )
        self.output = translateTo(data, language)
        result[index] = 1
        return 1


class style_info:
    def __init__(self, author) -> None:
        self.author = author
        self.output = ""
        self.tokens = 350
        self.style = ""
        self.on_error_repeats = 3
        self.category = translateTo(stencil.categories[2], language)

    def get_style_name(self):
        for _ in range(self.on_error_repeats):
            style = openai_response(f"what artistic style does {self.author} belong to? Use single word only. Choose from these: {' '.join(stencil.artistic_styles)}")
            if style.lower() in stencil.artistic_styles:
                self.style = style
                return 1
        self.style = "Unknown"
        return -1

    def get_data(self, result, index):
        if self.get_style_name() == 1:
            data = openai_response(
                f"write me a text about {self.style}, focus on these points:\n{' '.join(stencil.style_points)}",
                self.tokens
            )
            self.output = translateTo(data, language)
            result[index] = 1
            return 1
        else:
            self.output = "Unknown"
            result[index] = -1
            return -1


class content:
    def __init__(self) -> None:
        self.output = ""
        self.tokens = 350
        self.category = translateTo(stencil.categories[3], language)

    def get_data(self, result, index):
        data = openai_response(
            f"write me a text about the story of the book {book}. write about the story only, no other information.",
            self.tokens,
        )
        self.output = translateTo(data, language)
        result[index] = 1
        return 1

class characteristics:
    def __init__(self) -> None:
        self.max_character_count = 3
        self.output = ""
        self.tokens = 300
        self.on_error_repeats = 3
        self.category = translateTo(stencil.categories[4], language)

    def get_data(self, result, index):
        for _ in range(self.on_error_repeats):
            data = openai_response(
                f"Describe main characters in {book}. Max {self.max_character_count}. Let the each paragraph start with number.",
                self.tokens
            )
            if (data.count("\n") - 1) <= self.max_character_count:
                self.output = translateTo(data, language)
                result[index] = 1
                return 1
        self.output = "Unknown"
        result[index] = -1
        return -1


class type_genre_classification:
    def __init__(self) -> None:
        self.on_error_repeats = 10
        self.literar_class_index: int
        self.category = translateTo(stencil.categories[5], language)
        self.output = {
            "literar_class": "",
            "genre": "",
        }

    def get_literar_class(self):
        for _ in range(self.on_error_repeats):
            data = openai_response(
                f"what literar class is the book {book}, choose from these options and type your chosen answer only\n{' '.join(stencil.literar_classes)}."
            )
            data = text_only(data)
            try:
                if data.lower() in stencil.literar_classes:
                    self.literar_class_index = stencil.literar_classes.index(data.lower())
                    self.output["literar_class"] = data
                    return 1
            except:
                continue

        
        self.output["literar_class"] = f"({translateTo(data, language)})"
        return -1

    def complete(self, result, index):
        if self.get_literar_class() == 1:
            self.output["literar_class"] = translateTo(self.output["literar_class"], language)
            for _ in range(self.on_error_repeats):
                options = stencil.literar_classes_options[self.literar_class_index]
                data = openai_response(
                    f"what genre is the book {book}, choose from these options, type the your chosen option, nothing more\n{' '.join(options)}."
                )
                data = text_only(data)
                if data.lower() in options:
                    self.output["genre"] = translateTo(data, language)
                    result[index] = 1
                    return 1
            self.output["genre"] = "Unknown"
            result[index] = -1
            return -1
        else:
            data = openai_response(
                    f"what genre is the book {book}, single word response only."
                )
            data = text_only(data)
            self.output["genre"] = f"({translateTo(data, language)})"
            result[index] = -1
            return -1

class space_time:
    def __init__(self) -> None:
        self.on_error_repeats = 3
        self.category = translateTo(stencil.categories[6], language)
        self.output = {
            "place": "",
            "time": "",
        }

    def get_data(self, result, index):
        data = openai_response(
            f"where the plot of the book {book} takes place. answer as straight as possible.",
        )
        self.output["place"] = translateTo(data, language)
        data = openai_response(
            f"at what time the plot of the book {book} happened. answer with the time period only.",
        )
        self.output["time"] = translateTo(data, language)
        result[index] = 1
        return 1


class circumstances:
    def __init__(self) -> None:
        self.output = ""
        self.tokens = 350
        self.category = translateTo(stencil.categories[7], language)

    def get_data(self, result, index):
        data = openai_response(
            f"write me a text about the circumstances of writing the book {book}, do not write about the content of the book",
            self.tokens,
        )
        self.output = translateTo(data, language)
        result[index] = 1
        return 1

class philosophy:
    def __init__(self) -> None:
        self.output = ""
        self.tokens = 350
        self.category = translateTo(stencil.categories[8], language)

    def get_data(self, result, index):
        data = openai_response(
            f"write me about the thought of the book {book}, write about the idea, that the book is introducing. do not add any introduction to the text.",
            self.tokens,
        )
        self.output = translateTo(data, language)
        result[index] = 1
        return 1
    
class opinion:
    def __init__(self) -> None:
        self.output = ""
        self.tokens = 350
        self.category = translateTo(stencil.categories[9], language)

    def get_data(self, result, index):
        data = openai_response(
            f"write me a text where you rate the book {book} and talk about your opinion, from {'male' if opinion_guide['is_man'] else 'female'} perspective, the evaluation of the text is expressed by the following value: {opinion_guide['book_rating']}, 1 is worst and 10 best, some additional info: {opinion_guide['additional_info']}",
            self.tokens,
        )
        self.output = translateTo(data, language)
        result[index] = 1
        return 1



def main(book, language):

    bib = bibliographic_data(language)
    bib.complete()
    Author = bib.output["Author"]

    aut = author_info(Author)
    sty = style_info(Author)
    con = content()
    cha = characteristics()
    typ = type_genre_classification()
    spa = space_time()
    cir = circumstances()
    phi = philosophy()
    opi = opinion()

    text_classes = [bib, aut, sty, con, cha, typ, spa, cir, phi, opi]
    results = [None] * 9

    t1 = threading.Thread(target=aut.get_data, args=(results, 0))
    t2 = threading.Thread(target=sty.get_data, args=(results, 1))
    t3 = threading.Thread(target=con.get_data, args=(results, 2))
    t4 = threading.Thread(target=cha.get_data, args=(results, 3))
    t5 = threading.Thread(target=typ.complete, args=(results, 4))
    t6 = threading.Thread(target=spa.get_data, args=(results, 5))
    t7 = threading.Thread(target=cir.get_data, args=(results, 6))
    t8 = threading.Thread(target=phi.get_data, args=(results, 7))
    t9 = threading.Thread(target=opi.get_data, args=(results, 8))

    threads = [t1, t2, t3, t4, t5, t6, t7, t8, t9]


    for t in threads:
        t.start()

    for t in threads:
        t.join()

    for i, value in enumerate(results):
        if value == None:
            results[i] = -1

    error_rate = (len(threads) - sum(results)) / (len(threads) * 2)

    print(results)
    print(error_rate)

    #create word document

    document = Document()

    document.add_heading(translateTo("Čtenářský deník", language))

    for text_class in text_classes:
        document.add_heading(text_class.category, 2)

        if type(text_class.output) == type("lol"):
            document.add_paragraph(str(text_class.output))
        
        else:
            keys = list(text_class.output.keys())
            for key in keys:
                document.add_heading(key, 3)
                document.add_paragraph(str(text_class.output[key]))


    document.save("result.docx")



book = "Démon (Lermontov)"
language = "CS"
opinion_guide = {
    "book_rating": 6,      # integer [1, 10], how much is the opinion author enjoying the book (1 worst, 10 best)
    "is_man": True,            # boolean, for determining the sex of opinion author
    "additional_info": "Student in high school"   # string, ads some more info for specifing the opinion author
}

if __name__ == "__main__":
    main(book, language)