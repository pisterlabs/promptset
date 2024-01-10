#Lebab - Word docx translator, from one language to another - Copyright © 2023 Iwan van der Kleijn - See LICENSE.txt for conditions
import sys, os
from typing import Iterator, List
from docx import Document #type: ignore
from openai import OpenAI
from typing import List, Iterator
from itertools import islice


ContentItems = Iterator[List[str]]

llm = OpenAI()

SEPARATOR = "\n⸻⸻⸻⸻⸻\n"

def translate_content_items(content_elements: List[str], source_lang: str, target_lang: str) -> ContentItems:
    content_length = 0
    joined_content = []
    for element in content_elements:
        joined_content.append(element)
        if content_length + len(element) + len(SEPARATOR) > 4000:
            # Call translate_content with the joined content
            translated_content = translate_content(SEPARATOR.join(joined_content), source_lang, target_lang)
            # Split the translated content on the separator
            translated_content_list = translated_content.split(SEPARATOR)
            # Yield each element of the translated content list
            yield from translated_content_list
            # Reset the joined content and content length
            joined_content = []
            content_length = 0
        # Join the element with the separator
        
        content_length += len(element) + len(SEPARATOR)
    
    # Handle the remaining joined content
    if joined_content:
        translated_content = translate_content(SEPARATOR.join(joined_content), source_lang, target_lang)
        translated_content_list = translated_content.split(SEPARATOR)
        yield from translated_content_list
    
def translate_content(content: str, source_lang: str, target_lang:str)-> str:
        
    # Constructing the prompt for translation
    translation_prompt = f"""Translate the following text from {source_lang} to {target_lang}: 

The text  consists of text elements seperated by the following seperator: {SEPARATOR}

Leave the seperator in place, and translate the text elements in between the seperators.
Don't change the seperator itself. Dont'a add anything to the text elements, or remove anything from them.

Translate all of the text below until the (END OF TEXT) marker.):

{content}

(END OF TEXT)
""" 
    
    completion = llm.chat.completions.create(
    model="gpt-4",
    #model="gpt-4-1106-preview",
    
    messages=[
        {"role": "system", "content": "You are a profesional translator of many different languages. Your skill is the ability to strike a good ballance between semantic and communicative translation"},
        {"role": "user", "content": translation_prompt}
    ])
    
    translated_content = completion.choices[0].message.content
    if translated_content is None:
        raise Exception("No translation returned")
    return translated_content 

def get_content_elements(doc: Document) -> List[str]:
    content_items = []
    for paragraph in doc.paragraphs:
        if text := paragraph.text.strip():
            content_items.append(text)
            
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if text := paragraph.text.strip():
                        content_items.append(text)
    
    return content_items

def set_content(doc: Document, content_items:ContentItems):
    try:
        # overwrite the text in the paragraphs
        
        for paragraph in doc.paragraphs:
            if text := paragraph.text.strip():
                paragraph.text = next(content_items)
                
        # overwrite the text in the tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if text := paragraph.text.strip():
                            paragraph.text = next(content_items)
                           
    except StopIteration:
        print("Warning: less transalated content items then expected")
        return
            
def lebab(file_path, source_lang, target_lang):
    # Copy the file to a new file with the specified format
    new_file_path = f"{os.path.splitext(file_path)[0]}_{target_lang}.docx"
    doc = Document(file_path)
   
    doc.save(new_file_path)

    # Access the new file
    new_doc = Document(new_file_path)

    content_elements: List[str] = get_content_elements(new_doc)

    #print(content_elements)
    #write content to a text file
    #with open(f"{os.path.splitext(file_path)[0]}_{source_lang}.txt", "w") as text_file:
    #    text_file.write(content)
    
    translated_content_items: ContentItems = translate_content_items(content_elements, source_lang, target_lang)
    # with open(f"{os.path.splitext(file_path)[0]}_{target_lang}.txt", "r") as text_file:
    #     translated_content = text_file.read()
    
    
    #write translated_content to a text file
    #with open(f"{os.path.splitext(file_path)[0]}_{target_lang}.txt", "w") as text_file:
    #    text_file.write(translated_content)
      
    set_content(new_doc, translated_content_items)
    
    # Save the new file
    new_doc.save(new_file_path)

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: lebab document.docx source_language target_language")
        sys.exit(1)

    file_path = sys.argv[1]
    source_lang = sys.argv[2]
    target_lang = sys.argv[3]

    print(f"Translating {file_path} from {source_lang} to {target_lang}")
    lebab(file_path, source_lang, target_lang)
    print("Done")
