from config import Prompt_Limitation

from dataclasses import dataclass, field
from typing import List, Any
import copy
import json

from tools.llm.api_client import *
import openai
openai.api_base = "http://116.62.63.204:8001/v1"

import fitz
from fitz import TextPage

import docx
from docx import Document   # api.Document
from docx.oxml.ns import qn


import docx  # 导入python-docx库
# from docx.document import Document  # document.Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

# llm = LLM_Qwen()

from enum import Enum, auto

# ============================关于角色提示============================
# 一、你希望llm了解哪些信息：
# 1) where are you based?
# 2) what do you do for work?
# 3) what are your hobbies and interests?
# 4) what subject can you talk about for hours?
# 5) what are some goals you have?
# 二、你希望llm怎样回复：
# 1) how formal or casual should llm be?
# 2) how long or short should responses generally be?
# 3) how do you want to be addressed?
# 4) should llm have opinions on topics or remain neutral?
# ============================关于角色提示============================

LLM_Doc_DEBUG = False
def dprint(*args, **kwargs):
    if LLM_Doc_DEBUG:
        print(*args, **kwargs)

def is_win():
    import platform
    sys_platform = platform.platform().lower()
    if "windows" in sys_platform:
        return True
    else:
        return False

# =========================================管理doc的层次化递归数据结构=================================================
from utils.Hierarchy_Node import Hierarchy_Node
# node 数据中的image数据
@dataclass
class Image_Data():
    name: str
    data: str
    width: int
    height: int

# node 数据中的table数据
@dataclass
class Table_Data():
    # table标题
    index:      str = ''  # '2.1.6.3.1-1'
    head:       str = ''  # '南麂微电网2015年各支路潮流及各配变负载率计算结果'
    unit:       str = ''  # '万元'
    annotate:   str = ''  # '注: ...'
    # table内容
    text:       str = ''  # 表格内容文本
    obj:        Any = None  # 表格的Table对象

# node 数据
@dataclass
class Doc_Node_Data():
    type:str = ''   # 'text' 'table' 'image'
    text:str = ''
    table:Table_Data = None    # Table_Data 对象
    image:Image_Data = None

# node 基本信息
@dataclass
class Doc_Node_Content():
    level: int      # 必需属性
    name: str       # 必需属性, 如: '1.1.3'
    heading: str    # 必需属性, 如: '建设必要性'
    data_list:List[Doc_Node_Data] = field(default_factory=list)   # 元素包括(text:str, image:Image_Part, table:str)
    # text: str       # 如: '本项目建设是必要的...'
    # image: Image_Part

# 用于控制prompt长度的参数
# @dataclass
# class Prompt_Limitation():
#     toc_max_len:int = 4096          # 返回目录(toc)字符串的最大长度
#     toc_nonsense_min_len:int = 300  # 返回目录(toc)内容太短从而无法进行总结的长度
#     context_max_len:int = 8192      # 返回文本(content)字符串的最大长度

# =========================================管理doc的层次化递归数据结构=================================================
# docx: 采用python-docx解析文档，采用win32com解决页码问题
# pdf:  采用pitz解析文档
class LLM_Doc():
    def __init__(self, in_file_name, in_llm=None):
        self.doc_name = in_file_name    # 文件名: 如c:/xxx.docx
        self.doc_base_name = os.path.splitext(os.path.basename(self.doc_name))[0] # 获取 c:/xxx.docx的xxx

        # 用于控制prompt长度的参数
        self.prompt_limit = Prompt_Limitation()

        # doc文档的目录标题是否含有1.1.3这样的索引数字(有可能docx中，1.1.3仅出现在域中)，大部分情况下标题是有索引数字的
        self.toc_heading_has_index = True

        self.win32_doc = None
        self.win32_doc_app = None
        self.win32_constant = None

        self.doc = None                 # Document对象
        self.doc_root = None            # 存放doc层次化数据
        self.doc_root_parsed = False    # 是否已经解析为层次化数据

        # for pdf
        self.pdf_doc = None # pdf的pitz对象
        self.pdf_toc = None # pdf的现成目录

        self.llm = in_llm
        print(f'------------------------------------------------------------------------------------------')
        print(f'LLM_Doc(): self.llm={self.llm}')
        print(f'------------------------------------------------------------------------------------------')
        if self.llm is None:

            self.llm = LLM_Client(
                history=False,
                # history_max_turns=50,
                # history_clear_method='pop',
                temperature=0,
                url='http://127.0.0.1:8001/v1',
                need_print=False,
            )
            print(f'------------------------------------------------------------------------------------------')
            print(f'LLM_Doc(): temperature={self.llm.temperature}')
            print(f'------------------------------------------------------------------------------------------')
            self.llm.set_role_prompt('你是文档问答助手，不管问你什么问题，都不做任何解释，直接按要求回复指定格式的内容')

        self.question_types = [
            '"关于文档总体的提问"',
            '"关于文档细节的提问"',
            '"关于文档指定章节的问题"',
            '"关于文档表格的提问"',
            '"与文档无关的问题"',
        ]

        try:
            if 'docx' in self.doc_name:
                self.doc = Document(self.doc_name)
            elif 'pdf' in self.doc_name:
                self.pdf_doc = fitz.open(self.doc_name)
        except Exception as e:
            print(f'打开文件"{self.doc_name}"失败。')

    # llm对user的提问进行分类，返回问题类型的index
    def llm_classify_question(self, in_question):
        question = f'用户正在对文档进行提问，问题是："{in_question}"，请问该问题属于哪种类型的提问，请从以下类型中选择一种：[{",".join(self.question_types)}]'
        print(f'llm_classify_question() question is : {question}')
        result = self.llm.ask_prepare(question).get_answer_and_sync_print()
        print(f'llm_classify_question result is: {result}')
        for i in range(len(self.question_types)):
            if self.question_types[i].replace('"', '') in result :
                return i
        return -1

    # 调用tools
    def call_tools(self, in_tool_index, in_question, in_toc, in_tables=None):
        content = ''
        answer = ''
        print(f'----------------------------------call_tools()输入----------------------------------------')
        print(f'输入问题为: "{in_question}"')
        print(f'所选择的工具编号为: {in_tool_index}')
        print(f'获取的目录信息长度为: {len(in_toc)}')
        if in_tables is not None:
            print(f'已输入table信息')
        else:
            print(f'未输入table信息')
        print(f'----------------------------------call_tools()输出----------------------------------------')

        match in_tool_index:
            case 0: #关于文档总体的提问
                # question = f'{in_toc}. 以上是一个文档的目录结构，请问该文档的总体内容描述应该在这个目录中的哪个章节中，请返回唯一的章节标题，返回内容仅为"章节号 章节标题"这样的字符串，不能返回其他任何解释、前缀或多余字符，而且，如果该文档目录为英文，则返回的章节标题也必须为英文'
                question = f'{in_toc}. 以上是一个文档的目录结构，用户针对这个文档提出了问题"{in_question}"，请问所提问题涉及的内容最可能出现在这个目录的哪个章节中，请返回唯一的章节标题，返回内容仅为"章节号 章节标题"这样的字符串，不能返回其他任何解释、前缀或多余字符，而且，如果该文档目录为英文，则返回的章节标题也必须为英文'
                print(f'question: {question}')
                chapter = self.llm.ask_prepare(question).get_answer_and_sync_print()
                print(f'call_tools[0] 选择chapter raw: "{chapter}"')
                # chapter = re.search(r'\d+(.\d+)*', chapter)
                # if chapter is not None:
                #     chapter = chapter.group(0)
                # print(f'call_tools[0] 选择chapter: {chapter}')

                print(f'---------------------------------定位的chapter为: -------------------------\n{chapter}')
                content = self.get_text_from_doc_node(in_node_heading=chapter, in_if_similar_search=False)
                print(f'---------------------------------返回内容content为: -------------------------\n{content}')
                content = self.long_content_summary(content)
                question = f'{content}. 以上是从文档中获取的具体内容，用户针对这块内容提出了问题"{in_question}"，请根据这块内容用中文回答问题，回复格式要层次清晰、便于理解，该换行的地方要换行，该编序号和缩进的地方要编制序号和缩进'
                print(f'call_tools[0] 最终问题:\n{question}')
                answer = self.llm.ask_prepare(question).get_answer_generator()

            case 1: # 关于文档细节的提问
                question = f'{in_toc}. 以上是一个文档的目录结构，用户针对这个文档提出了问题"{in_question}"，请问所提问题涉及的内容最可能出现在这个目录的哪个章节中，请返回唯一的章节标题，返回内容仅为"章节号 章节标题"这样的字符串，不能返回其他任何解释、前缀或多余字符，而且，如果该文档目录为英文，则返回的章节标题也必须为英文'
                chapter = self.llm.ask_prepare(question).get_answer_and_sync_print()

                print(f'---------------------------------定位的chapter为: -------------------------\n{chapter}')
                content = self.get_text_from_doc_node(in_node_heading=chapter, in_if_similar_search=True)
                print(f'---------------------------------返回内容content为: -------------------------\n{content}')
                # content = self.get_text_from_doc_node(in_node_heading=chapter)
                content = self.long_content_summary(content)
                question = f'{content}. 以上是从文档中获取的具体内容，用户针对这块内容提出了问题"{in_question}"，请根据这块内容用中文回答问题，回复格式要层次清晰、便于理解，该换行的地方要换行，该编序号和缩进的地方要编制序号和缩进'
                print(f'call_tools[0] 最终问题:\n{question}')
                answer = self.llm.ask_prepare(question).get_answer_generator()
            case 2: # 关于文档指定章节的问题
                chapter = in_question
                print(f'call_tools[0] 选择chapter raw: {chapter}')
                chapter = re.search(r'\d+(.\d+)*', chapter)
                if chapter is not None:
                    chapter = chapter.group(0)
                print(f'call_tools[0] 选择chapter: {chapter}')

                content = self.get_text_from_doc_node(in_node_heading=chapter)
                content = self.long_content_summary(content)
                question = f'{content}. 以上是从文档中获取的具体内容，用户针对这块内容提出了问题"{in_question}"，请根据这块内容用中文回答问题，回复格式要层次清晰、便于理解，该换行的地方要换行，该编序号和缩进的地方要编制序号和缩进'
                print(f'call_tools[0] 最终问题:\n{question}')
                answer = self.llm.ask_prepare(question).get_answer_generator()
            case 3: # 关于文档表格的提问
                if in_tables is not None:
                    table_names = [f'"{table.head}"' for table in in_tables]
                    question = f'[{", ".join(table_names)}]. 以上是从文档中所有表格名称的清单，用户的提问是"{in_question}"，不要做任何解释，请直接返回提问相关的表格名称。'
                    table_name = self.llm.ask_prepare(question).get_answer_and_sync_print()
                    print(f'call_tools[0] 选择table_name raw: {table_name}')
                    table_name = re.search(r'".+"', table_name)
                    if table_name is not None:
                        table_name = table_name.group(0)
                    print(f'call_tools[0] 选择table_name: {table_name}')

                    content = self.get_table_content_by_head(table_name)
                    content = self.long_content_summary(content)
                    question = f'{content}. 以上是从文档中获取的表格内容，用户针对这块内容提出了问题"{in_question}"，请根据表格内容回答问题，回复格式要层次清晰、便于理解，该换行的地方要换行，该编序号和缩进的地方要编制序号和缩进'
                    print(f'call_tools[0] 最终问题:\n{question}')
                    answer = self.llm.ask_prepare(question).get_answer_generator()
                else:
                    # pdf暂时无法解析table
                    question = f'文档无法解析出表格内容，用户针对这块内容提出了问题"{in_question}"，请回答该问题'
                    print(f'call_tools[0] 最终问题:\n{question}')
                    answer = self.llm.ask_prepare(question).get_answer_generator()
            case 4: # 与文档无关的问题
                answer = self.llm.ask_prepare(in_question).get_answer_generator()
            case -1:
                print('call_tools(): 未匹配到tool')

        print(f'----------------------------------call_tools()输出----------------------------------------')
        return answer

    # 对超长文本进行精简
    def long_content_summary(self, in_content):
        print(f'==========long_content_summary(长度{len(in_content)}) =============')
        content = in_content
        if len(content) <= self.prompt_limit.context_max_len:
            # 如果文本长度没有超过Prompt_Limitation.context_max_len(4096)，则直接返回
            return content

        # 如果文本长度超过Prompt_Limitation.context_max_len(4096)，进行分段精简并汇编
        paras = content.split('\n')
        para_len = 0
        content_list_to_summary = []
        one_content = ''
        for para in paras:
            one_content += para + '\n'
            para_len += len(para)
            if para_len >= self.prompt_limit.context_max_len:
                content_list_to_summary.append(one_content)
                one_content = ''
                para_len = 0

        answer_list = []
        for content in content_list_to_summary:
            print(f'==========需要总结的文本(长度{len(content)})为: =============\n{content}')
            question = f'"{content}", 请对这些文字进行总结，总结一定要简明扼要、要抓住重点、字数要少于100字，不要进行解释，直接返回总结后的文字'
            gen = self.llm.ask_prepare(question).get_answer_generator()
            print(f'--------------------------------该文本的总结结果--------------------------------\n')
            answer = ''
            for chunk in gen:
                print(chunk, end='', flush=True)
                answer += chunk
            print()
            print(f'-----------------------------------------------------------------------------\n')

            answer_list.append(answer)

        final_answer =''
        answers = '\n'.join(answer_list)
        question = f'"{answers}", 请对这些文字进行总结，总结一定要简明扼要、要抓住重点、字数要少于2000字，不要进行解释，直接返回总结后的文字'
        print(f'==========需要总结的文本(长度{len(answers)})为: =============\n{answers}')
        print(f'--------------------------------该文本的总结结果--------------------------------\n')
        gen = self.llm.ask_prepare(question).get_answer_generator()
        final_answer = ''
        for chunk in gen:
            print(chunk, end='', flush=True)
            final_answer += chunk
        print(f'-----------------------------------------------------------------------------\n')

        return final_answer

    def ask_docx(self, in_query, in_max_level=3):
        file = self.doc_name
        doc = self
        doc.parse_all_docx()
        toc = doc.get_toc_list_json_string(in_max_level=in_max_level)

        # -------------------------------找到query内容所在章节---------------------------------------
        prompt = '''
        这是文档的目录结构"{toc}",
        请问这个问题"{query}"涉及的内容应该在具体的哪个章节中，不解释，请直接以"章节编号"形式返回。
        '''
        prompt = prompt.format(toc=toc, query=in_query)
        self.llm.need_print = False         # 关闭print输出
        res = self.llm.ask_prepare(prompt).get_answer_and_sync_print()

        # --------------将'1.1.3 some章节'转换为'1.1.3'----------------------
        re_result = re.search(r"\d+(.\d+)*", res).group(0)

        # --------------获取'1.1.3'对应章节下的text_got----------------------
        node = doc.find_doc_node(re_result)
        inout_text = []
        doc.get_text_from_doc_node(inout_text, node)
        text_got = '\n'.join(inout_text)

        # --------------对text_got进行限定范围的query----------------------
        prompt2 = '''
        请根据材料"{text_got}"中的内容, 回答问题"{query}"。
        '''

        prompt2 = prompt2.format(text_got=text_got, query=in_query)
        print(prompt2)
        print(f'材料长度为: {len(text_got)}')
        # self.llm.need_print = True          # 打开print输出
        gen = self.llm.ask_prepare(prompt2).get_answer_generator()
        return gen

    def win32com_init(self):
        print(f'win32尝试打开文件"{self.doc_name}"')
        if is_win():
            import win32com.client as win32
            from win32com.client import constants
            self.win32_constant = constants
        else:
            print(f"未在windows系统下运行，无法获取word文档页码.")
            return

        import os

        self.win32_doc_app = win32.gencache.EnsureDispatch('Word.Application')  # 打开word应用程序
        self.win32_doc_app.Visible = False
        # doc_app.Visible = True
        curr_path = os.getcwd()
        self.win32_doc = self.win32_doc_app.Documents.Open(self.doc_name)
        print(f'win32打开了文件"{self.doc_name}"')

    def win32_close_file(self):
        if self.win32_doc:
            try:
                self.win32_doc_app.Documents.Close(self.doc_name)
            except Exception as e:
                print(f'关闭文件"{self.doc_name}"出错: {e}')

    # 遍历python-docx数据某个节点下所有block里的paragraph内容或table内容
    def iter_block_items(self, parent):
        """
        Yield each paragraph and table child within *parent*, in document order.
        Each returned value is an instance of either Table or Paragraph. *parent*
        would most commonly be a reference to a main Document object, but
        also works for a _Cell object, which itself can contain paragraphs and tables.
        """
        if isinstance(parent, docx.document.Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    # 获取所有node中所有table组成的清单
    def get_all_tables(self):
        table_names = []
        if self.doc_root is None:
            return []

        def _get_all_tables(in_node):
            for child in in_node.children:
                for item in child.node_data.data_list:
                    if item.type=='table':
                        table_names.append(item.table)
                _get_all_tables(child)

        root = self.find_doc_node('root')
        _get_all_tables(root)

        return table_names

    # 从table清单中按照表明(head)查找一个表
    def get_table_content_by_head(self, in_head):
        table_head = in_head.replace('"', '').replace('“', '')
        print(f'get_table_content_by_head() table_head: {table_head}')
        for table in self.get_all_tables():
            if table_head in table.head :
                return table.text

    # 获取完整目录(table of content)的md格式
    def legacy_get_toc_md_for_tool(
            self,
            in_max_level='auto',    # 'auto' | 1 | 2 | 3 | ...
            in_if_render=False
    ):
        toc = []
        if self.doc_root is None:
            return []

        max_level = 0
        if in_max_level == 'auto':
            max_level = 3
            toc_max_len = self.prompt_limit.toc_max_len
            self.doc_root.get_toc_md_for_tool(toc, self.doc_root, in_max_level=max_level, in_if_render=in_if_render)
            if len('\n'.join(toc)) > toc_max_len:
                toc = []
                max_level = 2
                self.doc_root.get_toc_md_for_tool(toc, self.doc_root, in_max_level=max_level, in_if_render=in_if_render)
        else:
            max_level = in_max_level
            self.doc_root.get_toc_md_for_tool(toc, self.doc_root, max_level, in_if_render=in_if_render)

        # 统计toc标题中是否有1.1.3这样的数字
        num_head_has_index = 0
        for item in toc:
           if re.search(r"\d+(.\d+)*", item) is not None:
               num_head_has_index += 1

        if num_head_has_index/len(toc) > 0.7:
            self.toc_heading_has_index = True
            # 表明该文档的目录标题中还有1.1.3这样的数字
            print(f'文档"{self.doc_name}"的目录中有{num_head_has_index}个标题中包含数字，占比为{num_head_has_index/len(toc)*100:.2f}%')
            print(f'判定文档"{self.doc_name}"的目录标题包含索引数字')
        else:
            self.toc_heading_has_index = False
            # 表明该文档的目录标题中没有1.1.3这样的数字(如1.1.3在自动增长的域中)，此时需要设置in_if_head_has_index=False，生成专门的1.1.3
            print(f'文档"{self.doc_name}"的目录中有{num_head_has_index}个标题中包含数字，占比为{num_head_has_index/len(toc)*100:.2f}%')
            print(f'判定文档"{self.doc_name}"的目录标题不包含索引数字')
            toc = []
            # max_level直接用上面的结果
            self.doc_root.get_toc_md_for_tool(toc, self.doc_root, max_level, in_if_head_has_index=False, in_if_render=in_if_render)


        return '\n'.join(toc)

    # 获取某节点下的目录(table of content)的md格式
    def get_toc_md_for_tool_by_node(
            self,
            in_node,
            in_max_level='auto',    # 'auto' | 1 | 2 | 3 | ...
            in_if_render=False,
            in_if_md_head=True,
    ):
        toc = []
        if in_node is None:
            return []

        max_level = 0
        if in_max_level == 'auto':
            max_level = 3
            toc_max_len = self.prompt_limit.toc_max_len
            self.doc_root.get_toc_md_for_tool(toc, in_node, in_max_level=max_level, in_if_render=in_if_render, in_if_md_head=in_if_md_head)
            if len('\n'.join(toc)) > toc_max_len:
                toc = []
                max_level = 2
                self.doc_root.get_toc_md_for_tool(toc, in_node, in_max_level=max_level, in_if_render=in_if_render, in_if_md_head=in_if_md_head)
        else:
            max_level = in_max_level
            self.doc_root.get_toc_md_for_tool(toc, in_node, max_level, in_if_render=in_if_render, in_if_md_head=in_if_md_head)

        # 统计toc标题中是否有1.1.3这样的数字
        num_head_has_index = 0
        for item in toc:
           if re.search(r"\d+(.\d+)*", item) is not None:
               num_head_has_index += 1

        if num_head_has_index/len(toc) > 0.7:
            self.toc_heading_has_index = True
            # 表明该文档的目录标题中还有1.1.3这样的数字
            print(f'文档"{self.doc_name}"的目录中有{num_head_has_index}个标题中包含数字，占比为{num_head_has_index/len(toc)*100:.2f}%')
            print(f'判定文档"{self.doc_name}"的目录标题包含索引数字')
        else:
            self.toc_heading_has_index = False
            # 表明该文档的目录标题中没有1.1.3这样的数字(如1.1.3在自动增长的域中)，此时需要设置in_if_head_has_index=False，生成专门的1.1.3
            print(f'文档"{self.doc_name}"的目录中有{num_head_has_index}个标题中包含数字，占比为{num_head_has_index/len(toc)*100:.2f}%')
            print(f'判定文档"{self.doc_name}"的目录标题不包含索引数字')
            toc = []
            # max_level直接用上面的结果
            self.doc_root.get_toc_md_for_tool(toc, in_node, max_level, in_if_head_has_index=False, in_if_render=in_if_render, in_if_md_head=in_if_md_head)


        return '\n'.join(toc)

    # 获取目录(table of content)的json格式, list形式，节省字符串长度
    def get_toc_list_json_string(self, in_max_level=3):
        import json
        toc = []
        if self.doc_root is None:
            return json.dumps([], indent=2, ensure_ascii=False)

        self.doc_root.get_toc_list_json(toc, self.doc_root, in_max_level)

        return json.dumps(toc, indent=2, ensure_ascii=False)

    # 获取下目录(table of content)的json格式, dict形式，比较占用字符串长度
    def get_toc_dict_json_string(self, in_max_level=3):
        import json
        toc = {}
        if self.doc_root is None:
            return json.dumps({}, indent=2, ensure_ascii=False)

        self.doc_root.get_toc_dict_json(toc, self.doc_root, in_max_level)

        return json.dumps(toc, indent=2, ensure_ascii=False)

    # 用'1.1.3'这样的字符串查找node name
    def find_doc_node(self, in_node_s):
        if self.doc_root is None:
            return None

        return self.doc_root.find(in_node_s)

    # 用'1.1.3'这样的字符串查找node head
    def find_doc_node_by_head(self, in_node_s):
        if self.doc_root is None:
            return None

        return self.doc_root.find_by_head(in_node_s)

    # 遍历打印node下的所有内容
    def print_from_doc_node(self, in_node='root'):   # in_node为'1.1.3'或者Hierarchy_Node对象
        # 如果输入'1.1.3'这样的字符串
        if type(in_node)==str:
            node_s = in_node
            in_node = self.doc_root.find(node_s)
            if in_node is None:
                print(f'节点"{node_s}"未找到.')
                return

        # 如果输入Hierarchy_Node对象
        if in_node is None:
            return
        else:
            node = in_node

        node_level = node.node_data.level
        node_name = node.node_data.name
        node_heading = node.node_data.heading

        # 获取node文本
        node_content = self.get_node_data_callback(node)

        print(f'{"-"*node_level}{node_name}-{node_heading}{"-"*(80-node_level-len(node_name)-len(node_heading))}')
        print(f'{node_content}')

        for child in node.children:
            self.print_from_doc_node(child)

    # 遍历输出node下的所有内容
    def get_text_from_doc_node(
            self,
            in_node_heading='root',
            in_if_similar_search=False,
    ):
        def _get_text_from_doc_node(
            inout_text_list,
            in_node='root'
        ):   # in_node为'1.1.3'或者Hierarchy_Node对象
            # 如果输入'1.1.3'这样的字符串
            if type(in_node) == str:
                if in_if_similar_search==False:
                    # ====================对输入的node进行精确查找====================
                    if self.toc_heading_has_index==True:
                        print(f'====================_get_text_from_doc_node() self.toc_heading_has_index={self.toc_heading_has_index}================================')
                        # 搜索heading '1.1.3 建设必要性'，因为重新编码导致搜索name'1.19.2'是错的，而heading为'1.1.3 建设必要性'
                        node_s = in_node
                        in_node = self.doc_root.find_by_head(node_s)
                        if in_node is None:
                            dprint(f'节点"{node_s}"未找到.')
                            return
                    else:
                        print(f'====================_get_text_from_doc_node() self.toc_heading_has_index={self.toc_heading_has_index}================================')
                        # 搜索name '1.1.3' (heading='建设必要性')，因为没有重新编码将heading'1.1.3 建设必要性'编为name'1.19.2'
                        node_s = in_node
                        # in_node = self.doc_root.find(node_s)
                        # 先通过node name查找1.1.3
                        in_node = self.doc_root.find(node_s)
                        if in_node is None:
                            # 然后通过node heading查找1.1.3
                            in_node = self.doc_root.find_by_head(node_s)
                            if in_node is None:
                                dprint(f'节点"{node_s}"未找到.')
                                return
                else:
                    # ====================对输入的node字符串进行模糊查找====================
                    # 将"一次性的模糊查找结果"，改为"返回模糊相关度大于0的所有目录标题"，然后再问llm选取chapter
                        node_s = in_node
                        in_node = self.doc_root.find_similar_by_head(self.toc_heading_has_index, node_s)
                        if in_node is None:
                            dprint(f'节点"{in_node}"未找到.')
                            return

            # 如果输入Hierarchy_Node对象
            if in_node is None:
                return
            else:
                node = in_node

            node_level = node.node_data.level
            node_name = node.node_data.name
            node_heading = node.node_data.heading

            # 获取node文本
            node_content = self.get_node_data_callback(node)

            inout_text_list.append(f'{"-"*node_level}{node_name}-{node_heading}{"-"*(80-node_level-len(node_name)-len(node_heading))}')
            inout_text_list.append(f'{node_content}')

            for child in node.children:
                _get_text_from_doc_node(inout_text_list, child)
                # print('$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$')

        # print(f'------------------------------开始模糊查找，node: "{in_node}"-----------------------------')
        inout_text_list = []

        _get_text_from_doc_node(inout_text_list, in_node_heading)
        # -------------------------如果返回的content长度大于Prompt_Limitation.context_max_len(4096)-------------------------
        length = len('\n'.join(inout_text_list))
        # if length > self.prompt_limit.context_max_len:
        #     if in_if_similar_search==True:
        #         toc_node = self.doc_root.find_similar_by_head(self.toc_heading_has_index, in_node_heading)
        #     else:
        #         toc_node = self.doc_root.find_by_head( in_node_heading)
        #
        #     toc = self.get_toc_md_for_tool_by_node(toc_node, 100)  # 100表示搜索所有子目录
        #     print(f'warning: 返回content的长度为: {length} ,超过限制: Prompt_Limitation.context_max_len({self.prompt_limit.context_max_len})')
        #     print(f'warning: 改为返回对应目录内容, 目录内容长度为: {len(toc)}')
        #
        #     if len(toc) < self.prompt_limit.toc_nonsense_min_len:
        #         # -------------------------如果返回的content长度超限，但返回的目录toc又太短-------------------------
        #         # ------则仍然返回content，并且后续进行逐段的总结，然后汇编总结------
        #         return '\n'.join(inout_text_list)
        #
        #     return toc

        # print(f'------------------------------结束模糊查找------------------------------------------------')
        return '\n'.join(inout_text_list)

    # 为table增加annotate'注: ...'
    def get_table_annotate_from_text(self, in_text, inout_table_obj):
        match_annotate = re.search(r'\s*(?<=注[:：]).*', in_text)
        match_annotate = match_annotate.group(0).strip() if match_annotate else ''
        inout_table_obj.annotate = match_annotate
        return inout_table_obj

    # 从文本中解析出table的标题
    def get_table_head_from_text(self, in_text, inout_table_obj):
        '\n附表2.1.6.3.1 -1   南麂岛10kV线路总概算表   单位：万元'
        match_index = re.search(r'^((\s*)(表|附表))[\d\.-]*', in_text)  # 开头可能有'\n'这种字符，2.1.6.3.1和-1中间可能有空格' ', 不管这个空格
        # match_index = re.search(r'^((\s*)(表|附表))[\d\.\s-]*', in_text)  # 开头可能有'\n'这种字符，2.1.6.3.1和-1中间可能有空格' '
        match_index = match_index.group(0).strip() if match_index else ''

        match_unit = re.search(r'(?<=单位[：|:])\s*[\u4e00-\u9fa5]+', in_text)     # (?<=exp)匹配exp后面，(?=exp)匹配exp前面，(?!exp)匹配后面跟的不是exp，(?<!exp)匹配前面不是exp
        match_unit = match_unit.group(0).strip() if match_unit else ''

        match_unit_all = re.search(r'单位[：|:]\s*[\u4e00-\u9fa5]+', in_text)     # (?<=exp)匹配exp后面，(?=exp)匹配exp前面，(?!exp)匹配后面跟的不是exp，(?<!exp)匹配前面不是exp
        match_unit_all = match_unit_all.group(0).strip() if match_unit_all else ''

        match_head = in_text.replace(match_index, '').replace(match_unit_all, '').strip()

        inout_table_obj.index = match_index     # '附表21'
        inout_table_obj.head = match_head       # '南麂岛10kV线路总概算表'
        inout_table_obj.unit = match_unit       # '万元'
        # inout_table_obj.annotate= ''        # '注: ...'

        return inout_table_obj

    # 替代para.text，因为para.text解析word xml的文字时，会漏掉'w:smartTag'或'w:ins'下的'w:r'，导致如'表2.1.6.3.4'变为'表.3.4'的问题
    def _patch_get_text(self, in_para):
        text = []
        for elem in in_para._element:
            if elem.tag == qn('w:r'):
                text.append(elem.text)
            elif elem.tag == qn('w:ins') or elem.tag == qn('w:smartTag'):
                for sub_elem in elem:
                    if sub_elem.tag == qn('w:r'):
                        text.append(sub_elem.text)
        return ''.join(text)

    # 将doc解析为层次结构，每一级标题（容器）下都有text、table、image等对象
    def parse_all_docx(self):
        # 获取node上一级即parent的name，如'1.1.3'的上一级为'1.1', '1'的上一级为'root'
        def find_parent_node(in_node_name):
            dprint(f'================查找节点"{in_node_name}"的父节点', end='')
            if len(in_node_name.split('.')) == 1:
                parent_name = 'root'
                dprint(f'"{parent_name}"========')
                return self.doc_root.find(parent_name)
            else:
                # name_list = in_node_name.split('.')
                # name_list.pop()
                # parent_name = '.'.join(name_list)
                # print(f'"{parent_name}"========')
                # return self.doc_root.find(parent_name)

                name_list = in_node_name.split('.')
                dprint(f'========in_node_name: "{in_node_name}"========', end='')

                while True:
                    # 循环pop()，直到找到parent_node，例如3.4后面突然出现3.4.1.1，这时候的parent_node就3.4而不是3.4.1
                    if not name_list:   # 有可能 1.1之上没有1，要直接返回root
                        return self.doc_root

                    name_list.pop()
                    parent_name = '.'.join(name_list)
                    dprint(f'"parent_name: "{parent_name}" ========')
                    node = self.doc_root.find(parent_name)
                    if node:
                        return node

        # 处理root
        print(f'===================doc_base_name : {self.doc_base_name}===================')
        self.doc_root = Hierarchy_Node(Doc_Node_Content(0, 'root', self.doc_base_name))

        current_node = self.doc_root
        current_node_name = 'root'
        current_level = 0

        # 上一个文本，主要用于查找table标题
        last_para_text = ''
        # 用于判断'注：...'前面是否为表格
        last_is_table = False
        last_table_obj = None

        # 递归遍历doc
        # for para in self.doc.paragraphs:
        for block in self.iter_block_items(self.doc):
            if block.style.name == 'Normal Table':
                #--------------------------------------------找到表格-----------------------------------------------------
                table = block
                inout_table_obj = Table_Data(obj=table)
                self.get_table_head_from_text(last_para_text, inout_table_obj)  # 解析table的标题
                # table_obj = Table_Data(head=self.get_table_head_from_text(last_para_text), obj=table)
                # 添加内容(text、table、image等元素)
                self.parse_node_data_callback(
                    in_node_data_list_ref=current_node.node_data.data_list,
                    in_data=Doc_Node_Data(type='table', table=inout_table_obj)
                )  # 这里para.text是文本内容 Doc_Node_Data.data_list中的text

                last_is_table = True
                last_table_obj = inout_table_obj
            else:
                #--------------------------------------------找到标题head、文本text、图image等------------------------------
                para = block
                if last_is_table:
                    self.get_table_annotate_from_text(self._patch_get_text(para), last_table_obj)
                    # *******************测试para.text的bug************************
                    # if '注：由于各配电变压器' in para.text:
                    #     print(f'注：由于各配电变压器: [{self._patch_get_text(para)}]')
                    #     print(f'注：由于各配电变压器: [{para.text}]')
                        # print(f'注：由于各配电变压器: [{para._element.xml}]')
                        # print(f'annotate: {last_table_obj.annotate}')
                    # *******************测试para.text的bug************************

                last_para_text = self._patch_get_text(para)
                style = para.style.name                     # "Heading 1"
                style_name  = style.split(' ')[0]           # 标题 "Heading

                if style_name=='Heading':
                    # ----------------------------------找到标题heading，处理Hierarchy层次结构-------------------------------
                    new_level = int(style.split(' ')[1])    # 标题级别 1

                    # 计算current_node_name, 如：'1.1.1'或'1.2'
                    if current_node_name=='root':
                        current_node_name = '.'.join(['1']*new_level)   # 1级为'1', 如果直接为2级就是'1.1'
                    else:
                        if new_level == current_level:
                            dprint(f'----------------------------current_node_name: {current_node_name}----------------------------')
                            dprint(f'new_level: {new_level}')
                            dprint(f'current_level: {current_level}')
                            # ‘1.1.1’变为'1.1.2'
                            new_node_list = current_node_name.split('.')  # ['1', '1', '1']
                            last_num = int(new_node_list[-1]) + 1  # 2
                            new_node_list.pop()  # ['1', '1']
                            current_node_name = '.'.join(new_node_list) + '.' + str(last_num)  # '1.1.2'
                            # current_node_name = current_node_name[:-1] + str(int(current_node_name[-1])+1)
                        elif new_level > current_level:
                            dprint(f'----------------------------current_node_name: {current_node_name}----------------------------')
                            dprint(f'new_level: {new_level}')
                            dprint(f'current_level: {current_level}')
                            # ‘1.1.1’变为'1.1.1.1.1'
                            current_node_name += '.' + '.'.join(['1']*(new_level-current_level))
                        elif new_level < current_level:
                            dprint(f'----------------------------current_node_name: {current_node_name}----------------------------')
                            dprint(f'new_level: {new_level}')
                            dprint(f'current_level: {current_level}')
                            # ‘1.1.1’变为'1.2' 或 ‘1.1.1’变为'2'
                            new_node_list = current_node_name.split('.')    # ['1', '1', '1']
                            for i in range(current_level-new_level):
                                new_node_list.pop()                         # ['1', '1']
                            last_num = int(new_node_list[-1]) +1                      # 2
                            new_node_list.pop()                             # ['1']
                            if len(new_node_list)>0:
                                current_node_name = '.'.join(new_node_list) + '.' + str(last_num)   # '1.2'
                            else:
                                current_node_name = str(last_num)
                            # current_node_name = current_node_name[:-1-2*(current_level-new_level)] + str(int(current_node_name[-1-2*(current_level-new_level)])+1)
                    current_level = new_level

                    # 找到parent节点，并添加new_node
                    new_node = Hierarchy_Node(Doc_Node_Content(current_level, current_node_name, self._patch_get_text(para))) # 这里para.text是标题内容 Doc_Node_Data.heading
                    parent_node = find_parent_node(current_node_name)
                    parent_node.add_child(new_node)

                    # 刷新current状态
                    current_node = new_node
                else:
                    # ------------------------------------------找到内容：text、image等------------------------------------
                    self.parse_node_data_callback(
                        in_node_data_list_ref=current_node.node_data.data_list,
                        in_data=Doc_Node_Data(type='text', text=self._patch_get_text(para))
                    )    # 这里para.text是文本内容 Doc_Node_Data.data_list中的text

                last_is_table = False
                last_table_obj = None

        self.doc_root_parsed = True

    # 解析node数据的callback
    def parse_node_data_callback(self, in_node_data_list_ref, in_data):
        # 处理node中table的text内容
        node_content = ''
        node_data = in_data
        if node_data.type=='table':
            # 表格
            # table头
            tbl_index = node_data.table.index
            tbl_head = node_data.table.head
            tbl_unit = node_data.table.unit
            tbl_annotate = node_data.table.annotate

            node_content += '表格索引: ' + tbl_index + '\n'
            node_content += '表格名称: ' + tbl_head + '\n'
            node_content += '表格数值的单位: ' + tbl_unit + '\n'

            # table内容
            for row in node_data.table.obj.rows:
                for cell in row.cells:
                    node_content += cell.text + '\t'
                node_content += '\n'
            # table注解
            if tbl_annotate:
                node_content += '表格的注解: ' + tbl_annotate + '\n'

            node_data.table.text = node_content     # 解析过程中，将表格数据存入text

        # 将处理后的node_data添加到node_data_list
        in_node_data_list_ref.append(node_data)

    # 读取node数据的callback
    def get_node_data_callback(self, in_node):
        node_content = ''

        for node_data in in_node.node_data.data_list:
            if node_data.type=='text':
                # 普通文本
                node_content += node_data.text + '\n'
            elif node_data.type=='table':
                # 表格
                # table头
                tbl_index = node_data.table.index
                tbl_head = node_data.table.head
                tbl_unit = node_data.table.unit
                tbl_annotate = node_data.table.annotate
                half1 = (80-len(tbl_index)-len(tbl_head)-len(tbl_unit)-len('表格[::]'))//2
                half2 = (80-len(tbl_annotate)-len('[注: ]'))//2

                node_content += '-'*half1 + f'表格[{tbl_index}:{tbl_head}:{tbl_unit}]' + '-'*half1 + '\n'

                # table内容
                for row in node_data.table.obj.rows:
                    for cell in row.cells:
                        node_content += cell.text + '\t'
                    node_content += '\n'
                # table注解
                if tbl_annotate:
                    node_content += '-'*half2 + f'[注: {tbl_annotate}]' + '-'*half2 + '\n'
                else:
                    node_content += '-'*80

                node_data.table.text = node_content     # 解析过程中，将表格数据存入text
            else:
                pass

        return node_content

    # 将pdf解析为层次结构，每一级标题（容器）下都有text、image等对象
    def parse_all_pdf(self):
        # 获取node上一级即parent的name，如'1.1.3'的上一级为'1.1', '1'的上一级为'root'
        def find_parent_node(in_node_name):
            dprint(f'================查找节点"{in_node_name}"的父节点', end='')
            if len(in_node_name.split('.')) == 1:
                parent_name = 'root'
                dprint(f'"{parent_name}"========')
                return self.doc_root.find(parent_name)
            else:
                # name_list = in_node_name.split('.')
                # name_list.pop()
                # parent_name = '.'.join(name_list)
                # print(f'"{parent_name}"========')
                # return self.doc_root.find(parent_name)

                name_list = in_node_name.split('.')
                dprint(f'========in_node_name: "{in_node_name}"========', end='')

                while True:
                    # 循环pop()，直到找到parent_node，例如3.4后面突然出现3.4.1.1，这时候的parent_node就3.4而不是3.4.1
                    if not name_list:   # 有可能 1.1之上没有1，要直接返回root
                        return self.doc_root

                    name_list.pop()
                    parent_name = '.'.join(name_list)
                    dprint(f'"parent_name: "{parent_name}" ========')
                    node = self.doc_root.find(parent_name)
                    if node:
                        return node

        # 处理root
        doc_name = os.path.splitext(os.path.basename(self.doc_name))[0] # 获取c:/xxx.pdf的xxx
        print(f'===================doc_name : {doc_name}===================')
        self.doc_root = Hierarchy_Node(Doc_Node_Content(0, 'root', doc_name))

        current_node = self.doc_root
        current_node_name = 'root'
        current_level = 0

        # 上一个文本，主要用于查找table标题
        last_para_text = ''
        # 用于判断'注：...'前面是否为表格
        last_is_table = False
        last_table_obj = None

        self.pdf_toc = self.pdf_doc.get_toc()

        # 递归遍历toc
        for i in range(len(self.pdf_toc)):
            # 当前为章节如1.1
            item = self.pdf_toc[i]
            next_item = self.pdf_toc[i+1 if i<len(self.pdf_toc)-1 else len(self.pdf_toc)-1] # 如果一共10个，i为9时，则i+1也取9
            level = item[0]
            head = item[1]
            page_num = item[2]
            next_item_page_num = next_item[2]
            dprint(f'目录级别：{level} ', end='')
            dprint(f'目录标题：{head} ', end='')
            dprint(f'目录页码：{page_num}')

            # ----------------------------------找到标题heading，处理Hierarchy层次结构-------------------------------
            new_level = int(level)   # 目前节点的级别，如1、2、3

            # 计算current_node_name, 如：'1.1.1'还是'1.2'
            if current_node_name == 'root':
                current_node_name = '.'.join(['1'] * new_level)  # 1级为'1', 如果直接为2级就是'1.1'
            else:
                if new_level == current_level:
                    # 相同层级
                    dprint(f'----------------------------current_node_name: {current_node_name}----------------------------')
                    dprint(f'new_level: {new_level}')
                    dprint(f'current_level: {current_level}')
                    # ‘1.1.1’变为'1.1.2'
                    new_node_list = current_node_name.split('.')  # ['1', '1', '1']
                    last_num = int(new_node_list[-1]) + 1  # 2
                    new_node_list.pop()  # ['1', '1']
                    current_node_name = '.'.join(new_node_list) + '.' + str(last_num)  # '1.1.2'
                    # current_node_name = current_node_name[:-1] + str(int(current_node_name[-1])+1)
                elif new_level > current_level:
                    # 上级
                    dprint(f'----------------------------current_node_name: {current_node_name}----------------------------')
                    dprint(f'new_level: {new_level}')
                    dprint(f'current_level: {current_level}')
                    # ‘1.1.1’变为'1.1.1.1.1'
                    current_node_name += '.' + '.'.join(['1'] * (new_level - current_level))
                elif new_level < current_level:
                    # 下级
                    dprint(f'----------------------------current_node_name: {current_node_name}----------------------------')
                    dprint(f'new_level: {new_level}')
                    dprint(f'current_level: {current_level}')
                    # ‘1.1.1’变为'1.2' 或 ‘1.1.1’变为'2'
                    new_node_list = current_node_name.split('.')  # ['1', '1', '1']
                    for i in range(current_level - new_level):
                        new_node_list.pop()  # ['1', '1']
                    last_num = int(new_node_list[-1]) + 1  # 2
                    new_node_list.pop()  # ['1']
                    if len(new_node_list) > 0:
                        current_node_name = '.'.join(new_node_list) + '.' + str(last_num)  # '1.2'
                    else:
                        current_node_name = str(last_num)
                    # current_node_name = current_node_name[:-1-2*(current_level-new_level)] + str(int(current_node_name[-1-2*(current_level-new_level)])+1)

            # （1）添加节点（容器）：找到parent节点，并添加new_node
            new_node = Hierarchy_Node(Doc_Node_Content(
                level=level,                # 3
                name=current_node_name,     # 1.1.3
                heading=head                # 建设必要性
            ))
            # 在parent节点的位置添加new_node
            parent_node = find_parent_node(current_node_name)
            parent_node.add_child(new_node)

            # 刷新current节点状态
            current_level = new_level
            current_node = new_node

            # （2）添加节点内容：text、image等
            text_to_add = []
            for page in self.pdf_doc.pages(page_num-1,next_item_page_num-1):
                # =====================整块获取text========================
                text = page.get_text('text')
                # =====================按line获取text========================
                # for block in page.get_text('dict')['blocks']:
                #     lines = block.get('lines')
                #     if lines:
                #         for line in lines:
                #             spans = line['spans']
                #             for span in spans:
                #                 text = span['text']
                #                 print(f'【line】: {text}')
                text_to_add.append(text)
            text_to_add = '\n'.join(text_to_add)
            self.parse_node_data_callback(
                in_node_data_list_ref=current_node.node_data.data_list,
                in_data=Doc_Node_Data(type='text', text=text_to_add)
            )


        self.doc_root_parsed = True

    def get_para_inline_images(self, in_para):
        image = {
            'name':'',
            'data':None,
            'width':0,
            'height':0,
        }
        images_list = []

        # 打印ImagePart的成员函数
        # for item in dir(docx.parts.image.ImagePart):
        #     print(item)

        # 这一段由python-docx库的issue中网友mfripp提供：https://github.com/python-openxml/python-docx/issues/249
        for run in in_para.runs:
            for inline in run._r.xpath("w:drawing/wp:inline"):
                width = float(inline.extent.cx)  # in EMUs https://startbigthinksmall.wordpress.com/2010/01/04/points-inches-and-emus-measuring-units-in-office-open-xml/
                height = float(inline.extent.cy)
                if inline.graphic.graphicData.pic is not None:
                    rId = inline.graphic.graphicData.pic.blipFill.blip.embed
                    image_part = self.doc.part.related_parts[rId]
                    filename = image_part.filename      # 文件名称(其实是类型), 如"image.wmf"
                    bytes_of_image = image_part.blob    # 文件数据(bytes)
                    image['name'] = filename
                    image['data'] = bytes_of_image
                    image['width'] = width
                    image['height'] = height
                    images_list.append(copy.deepcopy(image))

                    # print(f'image_part:{image_part}, width:{width}, height:{height}, rId: {rId}, image:{image}, filename:{filename}')
                    # with open('a.wmf', 'wb') as f:  # make a copy in the local dir
                    #     f.write(bytes_of_image)
        return images_list


    def get_all_inline_images(self):
        images_list = []

        for para in self.doc.paragraphs:
            images_list += self.get_para_inline_images(para)
        return images_list

    def get_paras(self):
        for para in self.doc.paragraphs:
            yield para.text

    # 检查某doc文档所有para的错别词
    def check_wrong_written_in_docx_file(self, in_llm):
        try:
            doc = Document(self.doc_name)
        except Exception as e:
            print(f'文件"{self.doc_name}" 未找到。')
            return

        ii = 0
        result_list = []
        for para in self.get_paras():
            ii += 1
            # print(f"正在分析第{ii}个段落...")

            # para result登记
            print("*"*30+f"正在分析第{ii}个段落"+"*"*30)
            check_prompt = f'现在帮我在这段文字中找错别字："{para}"'
            print(f'本段落内容："{para}"')
            # in_llm.print_history()
            # print(f'check prompt 为：{check_prompt}')
            result_list.append(in_llm.ask_prepare(check_prompt).get_answer_and_sync_print())

        return result_list

def main1():
    llm = LLM_Client(
        url='http://116.62.63.204:8001/v1',
        history=False,
        history_max_turns=50,
        history_clear_method='pop',
        temperature=0.9,
    )
    role_prompt = '你是一位汉字专家。你需要找出user提供给你的文本中的所有错别字，并给出修改意见。'
    example_prompts = [
        '例如，user发送给你的文字中有单个字的笔误："你是我的好彭友，我们明天粗去玩吧？"，你要指出"彭"应为"朋"、"粗"应为"出"。',
        '例如，user发送给你的文字中有涉及语义的笔误："我们已对全社会效益已经财务效益进行了全面分析。"，你要指出"已经"应为"以及"。',
    ]
    # nagetive_example_prompt = '需要注意的是，一个词语不完整或者多余，并不属于错别字，例如"社会效益最大"应为"社会效益最大化"、"电影院"应为"电影"就不属于错别字，不要将这种情况误判为错别字。'
    nagetive_example_prompt = ''
    style_prompt = '你的错别字修改意见要以json格式返回，具体的json格式要求是，有错别字时像这样：{"result":"有错别字", [{"原有词语":"彭友", "修改意见":"朋友"}, {"原有词语":"粗去", "修改意见":"出去"}]}，没用错别字时像这样：{"result":"无错别字"}。'
    other_requirement = '直接返回json意见，不作任何解释。一步一步想清楚。'
    llm.set_role_prompt(role_prompt+''.join(example_prompts)+nagetive_example_prompt+style_prompt+other_requirement)

    # llm.print_history()
    # llm.ask_prepare('现在帮我在这段文字中找错别字："报告所提投资优化分析适用于新型电力系统、综合能源项目、微电网项目以及传统电力系统项目，具体支持冷热电气各类机组和设备模型，负荷类型支持城市类型、工业类型等，优化目标支持社会效益最大和财务效益最佳等。计算中已经内置了8760h负荷特性、新能源出力特性已经分时电价等信息。"').get_answer_and_sync_print()

    doc = LLM_Doc('d:/server/life-agent/tools/doc/错别字案例.docx')
    doc.win32com_init()
    res_list = doc.check_wrong_written_in_docx_file(llm)
    doc.win32_close_file()

    print("*"*30+"纠错结果汇编"+"*"*30)
    jj=0
    for item in res_list:
        jj += 1
        print(f'第{jj}段检查结果：\n {item}')

def main_toc():
    llm = LLM_Client(
        url='http://116.62.63.204:8001/v1',
        history=False,
        history_max_turns=50,
        history_clear_method='pop',
        temperature=0.9,
    )
    file = 'd:/server/life-agent/tools/doc/南麂岛离网型微网示范工程-总报告.docx'
    import docx
    doc = docx.Document(file)
    for para in doc.paragraphs:
        if para.style.name=="Heading 1":
            print(para.text)
        if para.style.name=="Heading 2":
            print('\t'+para.text)
        if para.style.name=="Heading 3":
            print('\t\t'+para.text)
        # if para.style.name=="Heading 4":
        #     print('\t\t\t'+para.text)
    # doc = LLM_Doc(file)
    # doc.win32com_init()
    # for para in doc.get_paras():
    #     print(para)
    # doc.win32_close_file()

def main_image():
    file = 'd:/server/life-agent/tools/doc/南麂岛离网型微网示范工程-总报告.docx'
    doc = LLM_Doc(file)
    # for para in doc.get_paras():
    #     print(para)
    # for image in doc.doc.inline_shapes:
    #     print(f'image: {image.type}')
        # print(f'image: {docx.enum.shape.WD_INLINE_SHAPE(3)}')
        # print(docx.enum.shape.WD_INLINE_SHAPE.PICTURE)

    # for image in doc.get_all_inline_images():
    #     print(', '.join([
    #         f'image name: {image["name"]}',
    #         f'image size: {len(image["data"])}',
    #         f'image width: {image["width"]}',
    #         f'image height: {image["height"]}'
    #     ]))

    doc.parse_all_docx()
    # doc.print_doc_root()
    # doc.print_doc_root('2.1.7')
    # node = doc.find_doc_root('9')
    node = doc.find_doc_node('2.1.6.3')
    # node = doc.find_doc_root('2.1.7')
    text = []
    doc.get_text_from_doc_node(text, node)
    print(''.join(text))
    # doc.print_from_doc_node(node)

    # text_list = []
    # doc.get_text_from_doc_node(text_list, node)
    # print('\n'.join(text_list))


    print(doc.get_toc_list_json_string(in_max_level=2))
    # print(doc.get_toc_json_string(in_max_level=3))

def ask_docx(in_filename='d:/server/life-agent/tools/doc/南麂岛离网型微网示范工程-总报告.docx'):
    llm = LLM_Client(
        history=False,
        # history_max_turns=50,
        # history_clear_method='pop',
        temperature=0.7,
        url='http://127.0.0.1:8001/v1',
        need_print=False,
    )

    file = in_filename
    doc = LLM_Doc(file, llm)
    doc.parse_all_docx()
    toc = doc.get_toc_list_json_string(in_max_level=3)



    while True:
        query = input('User: ')

        json_example = '{"head":"1.1.3"}'
        prompt = '''
        这是文档的目录结构"{toc}",
        请问这个问题"{query}"涉及的内容应该在具体的哪个章节中，不解释，请直接以"章节编号"形式返回。
        '''

        prompt = prompt.format(toc=toc, query=query, json_example=json_example)
        # print(f'--------发给LLM的prompt----------')
        # print(prompt)
        # print(f'--------发给LLM的prompt----------')
        res = llm.ask_prepare(prompt).get_answer_and_sync_print()
        print(f'Bot: {res}')
        re_result = re.search(r"\d+(.\d+)*",res).group(0)
        print(f'RE: {re_result}')

        node = doc.find_doc_node(re_result)
        # text_got = node.node_data.text
        inout_text = []
        doc.get_text_from_doc_node(inout_text, node)
        text_got = '\n'.join(inout_text)
        print(f'text_got: {text_got}')


        prompt2 = '''
        请根据材料"{text_got}"中的内容, 回答问题"{query}"。
        '''

        prompt2 = prompt2.format(text_got=text_got, query=query)
        llm.need_print = True
        res = llm.ask_prepare(prompt2).get_answer_and_sync_print()

# Color枚举类
class Color(Enum):
    red=auto()
    green=auto()
    blue=auto()

def main_table():
    # doc1 = Document('d:/server/life-agent/tools/doc/南麂岛离网型微网示范工程-总报告.docx')
    doc = LLM_Doc('d:/server/life-agent/tools/doc/南麂岛离网型微网示范工程-总报告.docx')

    # for table in doc1.tables:
    #     print('-------------------输出table-------------------')
    #     for row in table.rows:
    #         for cell in row.cells:
    #             print(cell.text, end='\t', flush=True)
    #         print()

    for block in doc.iter_block_items(doc.doc):
        # block.style.name可以直接返回：heading 1、normal、normal table
        if block.style.name == 'Normal Table':
            table = block
            print('-------------------输出table-------------------')
            for row in table.rows:
                for cell in row.cells:
                    print(cell.text, end='\t', flush=True)
                print()

    # doc.parse_all_docx()
    # # doc.print_doc_root()
    # # doc.print_doc_root('2.1.7')
    # node = doc.find_doc_root('2.1.3.2')
    # doc.print_from_doc_node(node)

def main():


    doc = LLM_Doc(in_file_name='d:/server/life-agent/tools/doc/南麂岛离网型微网示范工程-总报告.docx')
    while True:
        query = input("User: ")
        print('Assistant: ')
        gen = doc.ask_docx(query)
        for chunk in gen:
            print(chunk, end='', flush=True)
        print()

def main_llm_pdf():
    # doc = LLM_Doc(in_file_name='d:/server/life-agent/tools/doc/南麂岛离网型微网示范工程-总报告.docx')
    # doc = LLM_Doc(in_file_name='d:/server/life-agent/tools/doc/WorldEnergyOutlook2023.docx')
    # doc.parse_all_docx()
    # toc = doc.get_toc_md_string(2, in_show_md=False)
    # print(f'root: {doc.doc_root.node_data.heading}')
    # print(toc)

    doc = LLM_Doc(in_file_name='d:/server/life-agent/tools/doc/WorldEnergyOutlook2023.pdf')
    doc.parse_all_pdf()
    tables = doc.get_all_tables()

    # node = doc.find_doc_node_by_head('5.6')
    # print(f'==============node found: {node} ============')
    # inout_text = []
    # doc.get_text_from_doc_node(inout_text, node)
    # for item in inout_text:
    #     print(item)

    # question = '报告讲了什么？'
    # question = '报告2.2.3讲了什么？'
    # question = '负荷预测表返回给我'
    question = '今天天气如何？'

    # toc = doc.get_toc_md_for_tool(4)
    toc = doc.get_toc_md_for_tool_by_node(doc.doc_root, 3, in_if_render=False)
    print(toc)
    print(f'toc 长度: {len(toc)}')

    # print(f'user: {question}')
    # tool = doc.llm_classify_question(question)
    # print(f'选择工具: {tool}')
    # answer = doc.call_tools(tool, question, toc, in_tables=None)
    # for chunk in answer:
    #     print(chunk, end='', flush=True)
    # print()


    # doc = fitz.open("D:/server/life-agent/tools/doc/WorldEnergyOutlook2023.pdf")
    # 获取Document 文档对象的属性和方法
    # 1、获取pdf 页数
    # pageCount = doc.page_count
    # print("pdf 页数", pageCount)

    # 2、获取pdf 元数据
    # metaData = doc.metadata
    # print("pdf 元数据:", metaData)

    # pix = doc[3].get_pixmap()
    # pix._writeIMG(f'page{3}.png')
    print('==============================获取页面图片============================================')
    # image_list =  doc[0].get_images()
    # for img in image_list:
    #     print(f'img: {img}')
    #     xref = img[0]
    #     pix = fitz.Pixmap(doc, xref)
    #     if str(fitz.csRGB) == str(pix.colorspace):
    #         img_path = f'D:/server/life-agent/tools/doc/image{189}_{xref}.png'
    #         pix.save(img_path)
    #         print(f'D:/server/life-agent/tools/doc/image{189}_{xref}.png')
    print('===============================页面转图片===========================================')
    # 设置缩放和旋转系数,zoom_x, zoom_y取相同值，表示等比例缩放
    # page_index = 1
    # zoom_x = zoom_y = 3
    # rotation_angle = 0
    # trans = fitz.Matrix(zoom_x, zoom_y).prerotate(rotation_angle)
    # pm = doc[page_index].get_pixmap(matrix=trans, alpha=False)
    # 开始写图像
    # pm.save(f'D:/server/life-agent/tools/doc/gen.png')  # 第1张图片名：1.png，以此类推
    # print('D:/server/life-agent/tools/doc/gen.png saved')  # 第1张图片名：1.png，以此类推
    print('================================表格处理==========================================')
    # doc_tab_page = 289
    # tabs = doc[doc_tab_page].find_tables()
    # i = 0
    # for tab in tabs:
    #     i += 1
    #     print(f'---------------------table-{i} found.------------------------')
    #     df = tab.to_pandas()
    #     print(f'df: {df}')
    #     df.to_excel(f'{doc_tab_page}-{i}.xlsx')
    #     print(f'tab-{i} on page[{doc_tab_page}] saved.')

    print('================================================================================')


    # 3、获取pdf 目录信息
    # llm = LLM_Qwen(
    #     history=False,
    #     # history_max_turns=50,
    #     # history_clear_method='pop',
    #     temperature=0.7,
    #     url='http://127.0.0.1:8001/v1',
    #     need_print=False,
    # )

    # toc = doc.get_toc()
    # for item in toc:
    #     print(f"pdf 目录：{item}")


    # 4、遍历para
    # for page in doc.pages(0, 1):
    #     print(f'--------{page}---------')
    #     print(page.get_text('text'))
    #     # print(f'--------{page}---------')
    #     # print(page.get_text('blocks'))
    #     print(f'--------{page}---------')
    #     # print(json.dumps(page.get_text('dict'), indent=2))
    #     text_dict = page.get_text('dict')
    #     blocks = text_dict['blocks']
    #     for block in blocks:
    #         # lines = block['lines']
    #         lines = block.get('lines')
    #         if lines:
    #             for line in lines:
    #                 spans = line['spans']
    #                 for span in spans:
    #                     text = span['text']
    #                     print(f'【line】: {text}')

        # print(f'--------{page}---------')
        # print(page.get_text().encode('utf8'))
        # print(f'--------{page}---------')
        # print(page.get_text())

def main_llm():
    def call_llm():
        tables = doc.get_all_tables()
        # for table in tables:
        #     print(f'table: {table.text}')

        # question = '负荷预测水平是多少？'
        question = '主要结论是多少？'
        # question = '投资估算是多少？'
        # question = '报告讲了什么？'
        # question = '报告2.2.3讲了什么？'
        # question = '负荷预测表返回给我'
        # question = '今天天气如何？'

        print(f'user: {question}')
        tool = doc.llm_classify_question(question)
        print(f'选择工具: {tool}')
        answer = doc.call_tools(tool, question, toc, tables)
        for chunk in answer:
            print(chunk, end='', flush=True)

    doc = LLM_Doc(in_file_name='d:/server/life-agent/tools/doc/南麂岛离网型微网示范工程-总报告.docx')
    # doc = LLM_Doc(in_file_name='d:/server/life-agent/tools/doc/WorldEnergyOutlook2023.docx')
    doc.llm.need_print = False

    doc.parse_all_docx()

    toc = doc.get_toc_md_for_tool_by_node(
        doc.doc_root,
        # in_max_level='auto',
        in_max_level=2,
        in_if_md_head=False
    )
    # toc = doc.get_toc_md_for_tool(3, in_if_render=False)
    print(toc)
    print(f'toc 长度: {len(toc)}')

    # node_to_find = '2.1.2 负荷预测'
    # doc.doc_root.find_similar_node_by_head_and_ask_llm(doc.toc_heading_has_index, node_to_find)

    # call_llm()

if __name__ == "__main__":
    # main_llm_pdf()
    main_llm()
    # main_table()

    # (? <= \s)\d + (?=\s)
    # main_image()
    # match_unit = re.search(r'\s*(?<=注[:：]).*', ' 注：由于各配电变压器均只有一条631或632线路支线作为进线，各支线的潮流与流入配电变压器的功率相等，简化起见，在表中有功、无功和视在功率统一表示。')
    # match_unit = re.search(r'^(\s(表|附表))\d+(.\d+)*', '\n附表16                                 南麂岛智能用电系统安装工程费用表                                 单位：元')
    # match_unit = re.search(r'(?<=单位[：|:])\s*[\u4e00-\u9fa5]+', '附表21   南麂岛10kV线路总概算表   单位： 万元')
    # match_unit = match_unit.group(0) if match_unit else ''
    # print(match_unit)

    # doc = fitz.open("D:/server/life-agent/WorldEnergyOutlook2023.pdf")
    # # 获取Document 文档对象的属性和方法
    # # 1、获取pdf 页数
    # pageCount = doc.page_count
    # print("pdf 页数", pageCount)
    #
    # # 2、获取pdf 元数据
    # metaData = doc.metadata
    # print("pdf 元数据:", metaData)

    # 3、获取pdf 目录信息
    # llm = LLM_Qwen(
    #     history=False,
    #     # history_max_turns=50,
    #     # history_clear_method='pop',
    #     temperature=0.7,
    #     url='http://127.0.0.1:8001/v1',
    #     need_print=False,
    # )

    # toc = doc.get_toc()
    # print("pdf 目录：")

    # 4、遍历para
    # for page in doc.pages(100, 102):
    #     print(f'--------{page}---------')
    #     print(page.get_text('text'))
    #     # print(f'--------{page}---------')
    #     # print(page.get_text('blocks'))
    #     print(f'--------{page}---------')
    #     # print(json.dumps(page.get_text('dict'), indent=2))
    #     text_dict = page.get_text('dict')
    #     blocks = text_dict['blocks']
    #     for block in blocks:
    #         lines = block['lines']
    #         for line in lines:
    #             spans = line['spans']
    #             for span in spans:
    #                 text = span['text']
    #                 print(f'【line】: {text}')

        # print(f'--------{page}---------')
        # print(page.get_text().encode('utf8'))
        # print(f'--------{page}---------')
        # print(page.get_text())

    # for item in toc:
    #     level = item[0]
    #     head = item[1]
    #     print(f'{"-"*level}{head}')
    #
    # # prompt = f'"{toc_json}"为一本书的目录结构列表，注意列表中每一个元素的数据结构为[level, toc_head, page]，请只把level为1和2的目录标题翻译为中文后返回给我'
    #
    #     prompt = f'把{head}翻译为中文'
    #     res = llm.ask_prepare(prompt, in_max_tokens=4096).get_answer_and_sync_print()
    #     print(res)




    # print(f'color: {Color(1)}')
    # print(f'color: {Color.blue}')


# 若win32com打开word文件报错：AttributeError: module 'win32com.gen_py.00020905-0000-0000-C000-000000000046x0x8x7' has no attribute 'CLSIDToClassMap'
# 则删除目录C:\Users\tutu\AppData\Local\Temp\gen_py\3.10中的对应缓存文件夹00020905-0000-0000-C000-000000000046x0x8x7即可

# doc.win32_close_file()若报错：pywintypes.com_error: (-2147352567, '发生意外。', (0, 'Microsoft Word', '类型不匹配', 'wdmain11.chm', 36986, -2146824070), None)
# 很可能是和wps有关，据说卸载word，win32.gencache.EnsureDispatch('Word.Application')会成功调用wps