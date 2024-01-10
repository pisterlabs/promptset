import docx
import sys
import sqlite3
import html
import yaml
import re
import openai
import uuid
import pandas as pd

import sentiment as SE
import meetingsummarydb as DB
import summarize as SU
import mod_utilities as MU
import terms as TE
import meetingreport as MR

openai.api_key = MU.get_textfromfile('marmot.txt')


def create_tight_summary(prompt):
    '''Use the OpenAI endpoint to create a summary.'''

    try:

        response = openai.Completion.create(
            model="text-davinci-003",
            prompt="Summarize this for a fifth-grade student: {}".format(prompt),
            temperature=0.7,
            max_tokens=2000,
            top_p=1.0,
            frequency_penalty=0.0,
            presence_penalty=0.0
            ) 
        return response["choices"][0]["text"]

    except Exception as e:
        print(e)
        return "Error researching server."


def add_zero(instring):
    if len(instring) == 1:
        outstring = "0" + instring
    else:
        outstring = instring
    return outstring


def clean_timestamp(intimestring):
    slots = intimestring.split(":")
    newstring = add_zero(slots[0]) + ":" + add_zero(slots[1]) + ":" +  add_zero(slots[2])
    return newstring


def load_table_line(row, dbpath):
    '''Load the document table with a processDocument object.'''
    try:
        conn = sqlite3.connect(dbpath)
        cur = conn.cursor()
        cur.execute('INSERT INTO line (ID, TStamp, Speaker, \
            Verbatim, SENT_POS, SENT_NEU, SENT_NEG, Sentiment) VALUES \
            ( ?, ?, ?, ?, ?, ?, ?, ?)', \
            ( row[0], row[1],  row[2],  row[3],  row[4],  row[5],  row[6],  row[7]) )
        conn.commit()
        cur.close()
    except Exception as e:
        print("An error occurred in loadDocument/document {}".format(e))


def table_transcript_word(filein, dbpath, size=0):
    '''With a filepath to a MS Word transcript, and an optional line index, return a table.'''

    doc = docx.Document(filein)
    if size == 0:
        size = len(doc.paragraphs)
    
    parsed = ""

    for indx, i in enumerate(doc.paragraphs):
        if indx > 3 and indx < size:
            try:
                if re.search("[0-9][0-9]:[0-9][0-9]:[0-9][0-9]", i.text):
                    parsed = i.text.split(" ")
                else:
                    print("Line: {}".format(size-indx))
                    esctext = html.escape(i.text)
                    sent = SE.get_sentiment(i.text)
                    row = [indx, parsed[0], parsed[1] + " " + parsed[2], esctext, sent['neg'], sent['neu'], sent['pos'], sent['compound']]
                    load_table_line(row, dbpath)
            except Exception as e:
                print(e)


def table_transcript_teams(filein, dbpath, size=0):
    '''With a filepath to a Teams transcript, and an optional line index, return a table.'''

    doc = docx.Document(filein)
    if size == 0:
        size = len(doc.paragraphs)
    
    parsed = ""

    for indx, i in enumerate(doc.paragraphs):
        if indx > 0 and indx < size:
            try:
                print("Line: {}".format(size-indx))
                entry = i.text.split("\n")
                tempstamp = entry[0].split("-->")[1].split(".")[0]
                timestamp = clean_timestamp(tempstamp.strip())
                speaker = entry[1]
                esctext = html.escape(entry[2])
                sent = SE.get_sentiment(esctext)
                row = [indx, timestamp, speaker, esctext, sent['neg'], sent['neu'], sent['pos'], sent['compound']]
                load_table_line(row, dbpath)
            except Exception as e:
                print(e)


def get_verbatims(dbpath):
    '''With a loaded DB, collect varbitms.'''
    conn = sqlite3.connect(dbpath)
    cur = conn.cursor()
    speaker_list = list(cur.execute('SELECT * From speaker_list;'))
    verbatims = list(cur.execute('SELECT * From verbatims;'))
    cur.close()

    # get speakers
    verbatims_speaks = {}
    for i in speaker_list:
        verbatims_speaks[i[0]] = ""
    
    # prep long summaries
    summaries = {}
    for i in speaker_list:
        summaries[i[0]] = ""
    
    # prep short summaries
    shorts = {}
    for i in speaker_list:
        shorts[i[0]] = ""
    
    # get verbatims
    for i in speaker_list:
        for j in verbatims:
            if j[0] == i[0]:
                extract = str(j[1]) + " "
                verbatims_speaks[i[0]] += extract
    
    # get long summary of verbatims
    for i in speaker_list:
        summary = SU.summarize_text(html.unescape(verbatims_speaks[i[0]]))
        summaries[i[0]] = summary
        short = create_tight_summary(summary).strip()
        shorts[i[0]] = short
    
    # get short summary and store all summary fields
    for i in speaker_list:
        conn = sqlite3.connect(dbpath)
        cur = conn.cursor()
        cur.execute('INSERT INTO summary (Speaker, ShortSummary, LongSummary, Allsaid) \
            VALUES ( ?, ?, ?, ?)', \
            ( i[0], shorts[i[0]], summaries[i[0]], verbatims_speaks[i[0]]))
        conn.commit()
        cur.close()

def create_and_load_db(config):

    dbpath = config["reportpath"] + config["stem"] + ".db"
    DB.create_db(dbpath)
    if config["source"] == "word":
        table_transcript_word(config["transcript"], dbpath)
    elif config["source"] == "teams":
        table_transcript_teams(config["transcript"], dbpath)
    else:
        print("Need to set the source attribute.")
        exit
    get_verbatims(dbpath)

    # Load meeting data

    conn = sqlite3.connect(dbpath)
    cur = conn.cursor()
    m_id = config["stem"]
    m_title  = config["title"]
    m_meetingdate  = config["date"]
    m_attendees  = config["attendees"]
    m_agenda = config["agenda"]
    m_actionitems  = config["actions"]
    m_notes  = config["notes"]
    m_transcript = config["transcript"]
    m_recording  = config["recording"]

    try:
        conn = sqlite3.connect(dbpath)
        cur = conn.cursor()
        cur.execute('INSERT INTO meeting (ID, Title, MeetingDate, Attendees, \
            Agenda, ActionItems, Notes, Transcript, Recording) VALUES \
            (?, ?, ?, ?, ?, ?, ?, ?, ?)', ( m_id, m_title, \
            m_meetingdate, m_attendees, m_agenda, m_actionitems, \
            m_notes, m_transcript, m_recording) )
        conn.commit()
        cur.close()
    except Exception as e:
        print(e)

    # Extract entities from the verbatims.

    conn = sqlite3.connect(dbpath)
    cur = conn.cursor()
    speaker_list = list(cur.execute('SELECT * From speaker_list;'))
    cur.close()
    speakers = []
    for i in speaker_list:
        speakers.append(i[0])
    for s in speakers:
        conn = sqlite3.connect(dbpath)
        cur = conn.cursor()
        summaries = pd.read_sql("SELECT * from summary Where Speaker='{}'".format(s), con=conn)
        conn.close()
        fulltext = html.unescape(summaries["Allsaid"].values[0])
        entities = TE.get_top_fifty(fulltext, "{}".format(s))
        keys = entities.keys()
        for k in keys:
            entity_item = entities[k]["keyword"].lower()
            alternate = ""
            if entity_item[-1:] == "s":
                alternate = entity_item
                entity_item = entity_item[:-1]
            try:
                conn = sqlite3.connect(dbpath)
                cur = conn.cursor()
                cur.execute('INSERT INTO entity VALUES ("{}", "{}");'.format(entity_item, alternate))
                conn.commit()
                cur.close()
            except Exception as e:
                print(e)

    # match entities to lines create the KWIC

    conn = sqlite3.connect(dbpath)
    cur = conn.cursor()
    entity_list = list(cur.execute('SELECT Entity From entity;'))
    lines = list(cur.execute('SELECT ID, Verbatim From line;'))
    cur.close()
    conn.close()

    for i in entity_list:
        for j in lines:
            line = j[1].lower()
            if line.find(i[0]) > 0:
                conn = sqlite3.connect(dbpath)
                cur = conn.cursor()
                cur.execute('INSERT INTO occurance VALUES ("{}", "{}", "{}");'.format(str(uuid.uuid4()), i[0], j[0]))
                conn.commit()
                cur.close()
    
    # Rank the keywords

    conn = sqlite3.connect(dbpath)
    cur = conn.cursor()
    keyword_return = list(cur.execute('SELECT Distinct Entity From KWIC;'))
    cur.close()
    conn.close()

    keywords = []
    for i in keyword_return:
        keywords.append(i[0])

    for i in keywords:
        conn = sqlite3.connect(dbpath)
        cur = conn.cursor()
        rank = list(cur.execute('SELECT COUNT (Entity) FROM KWIC WHERE Entity="{}";'.format(i)))
        cur.close()
        conn.close()

        try:
            conn = sqlite3.connect(dbpath)
            cur = conn.cursor()
            cur.execute('INSERT INTO ranks VALUES ("{}", "{}");'.format(i, rank[0][0]))
            conn.commit()
            cur.close()
        except Exception as e:
            print(e)
    
    return dbpath

def main():
    '''Useage: python ./sumeeting.py "<pathtoconfig>" '''

    # load config

    listofitems = list(sys.argv)
    configpath = listofitems[1]

    try:
        with open (configpath, "r") as stream:
            config = yaml.load(stream, Loader=yaml.CLoader) 
    except Exception as e:
        print("Please provide a path to a job file (meeting.yml).")
        print(e)
        exit()

    # Create and load DB
    dbpath = create_and_load_db(config)

    # get template and generate markdown report
    
    template_string = MU.get_textfromfile("reporttemplate1.md")
    MR.create_report(template_string, config["reportpath"], dbpath, config["stem"])

    print("{} - Done".format(config["title"]))

if __name__ == "__main__":
    main()