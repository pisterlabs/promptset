### -*- coding: utf-8 -*-

import os
import openai
import requests
import json
import time
import argparse
import pandas as pd
import math
import datetime
from docx import Document
from docx.shared import Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import string

openai.organization = "org-pVvLzSyhulQuYjwhlLUucvrJ"
openai.api_key = "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"
API_ENDPOINT = "https://api.openai.com/v1/chat/completions"

def generate_chat_completion(messages, model="gpt-3.5-turbo-16k", temperature=1, max_tokens=None):
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {openai.api_key}",
    }

    data = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
    }

    if max_tokens is not None:
        data["max_tokens"] = max_tokens
    response = requests.post(API_ENDPOINT, headers=headers, data=json.dumps(data))

    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    else:
        raise Exception(f"Error {response.status_code}: {response.text}")

def read_excel(root):
    file=pd.read_excel(root)
    dict_list=[]
    for i in range(len(file)):
        dict_thisfile={}
        this_line=file.loc[i]
        dict_thisfile['type']=this_line[0]
        dict_thisfile['name']=this_line[1]
        dict_thisfile['author']=this_line[2]
        dict_thisfile['word_count']=this_line[3]
        dict_thisfile['other']=this_line[4]
        dict_list.append(dict_thisfile)
    return dict_list


def add_Head(document, level, text, font, size):# document word对象# level 几级标题# text 标题内容# font 标题字体，font=u'黑体'# size 标题大小
    title = document.add_heading(level=level)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(text)
    title_run.font.size = Pt(size) # 设置字体大小，小四对应值为12
    title_run.font.name = font # 设置字体类型属性
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), font) # font = u'黑体'
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    
def writedocx(text,title,doc_root_and_name):
    #text:response text
    #title:标题
    #doc_name：文件标题
    if text.find(title)==-1:
        notitle_text=text
    else:
        notitle_text=text[text.find(title)+len(title)+1:]
    document = Document()
    add_Head(document,1,title,u'宋体',12)
    for text_split in notitle_text.split('\n'):
        content=document.add_paragraph()
        run=content.add_run(text_split)
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    document.save(f'{doc_root_and_name}.docx')

    
def report_generation_first(dic):
    response_text=""
    name=dic['name']
    
    other_requirment=dic['other']
    if str(other_requirment)=='nan':
        other_requirment=""
    
    author=dic['author']
    if str(author)=='nan':
        auther_message="这篇报告不需要注明作者。"
    else:
        auther_message=f"作者:{author}。"
        
        
    
    word_count=2000
    if str(dic['word_count'])!='nan':
        word_count=dic['word_count']
        
    messages = [
        {"role": "user", "content": f"""
        这是一篇应用专题报告，带数理实证检验的内容。
        “““
        心理护理对慢性盆腔炎治疗患者临床疗效的影响专题报告
        一、前言与研究现状
        慢性盆腔炎是一种以白带增多、月经增量、小腹坠痛为临床症状的妇科疾病，多因生殖器官、周围组织、盆腔腹膜等部位发生炎症，严重状况可致患者出现盆腔静脉瘀血、输卵管阻塞不孕、异位妊娠等病症，危害妇女生命安全[1-2]。抗感染药物、手术以及中医技术都是目前治疗慢性盆腔炎常用方式，但是，抗感染药物带来的耐药性，手术带来的并发症以及中医带来的疗程漫长，都会导致患者在治疗慢性盆腔炎期间普遍产生抑郁、焦虑、失落、自卑等负面情绪，不利于疾病恢复[3]。2019年，周晓燕[4]等人的研究显示，心理应激对策护理可以有效缓解慢性盆腔炎患者的负面情绪改善患者的生活质量。所以，本次专题报告通过研究心理护理对慢性盆腔炎治疗患者临床疗效的影响，结果显示效果良好，具体报告如下。
        二、问题确立
            慢性盆腔炎为临床常见妇科疾病，近年来，伴随女性生活、工作压力的剧增，慢性盆腔炎的发病率呈现升高态势。为寻求一种有效的非药物治疗方式，缓解患者临床症状。本次专题报告为探讨心理护理对慢性盆腔炎患者临床治疗效果，选取2019年8月～2022年12月50例患者作为研究对象进行分析。
        三、研究目的
        探讨心理护理方式对慢性盆腔炎患者的临床疗效，寻求一种改善患者心理负面情绪，增强患者治疗信心，为提升慢性盆腔炎患者治疗效果提供指导思路。
        四、文献查证
            根据本次专题报告研究目的，选择数据库Pub-Med、MEDLINE、SCI、中文科技期刊数据库、中国知网、万方数据库、中国生物医学文献数据库，检索研究相关文献。选取“慢性盆腔炎”“心理护理”“临床疗效”“负面情绪”等作为检索词。检索年限为建库至今。
        五、资料与方法
        5.1一般资料  选取我院2019年8月～2022年12月门诊治疗的50例慢性盆腔炎患者作为研究对象，以2019年8月~2022年3月收治的25例为对照组，2022年4月～2022年12月的25例为研究组。纳入方法：经临床诊断确诊为慢性盆腔炎。排除标准：近1月接受慢性盆腔炎治疗；备孕期、妊娠期或哺乳期女性。
        5.2方法 
        两组患者在就诊时均进行常规护理，研究组在此基础上增加心理护理：（1）患者就诊时积极了解患者的家庭、工作等基础情况，详细记载之前的治疗情况，细心评估患者在就诊时的心理状态，并以积极、轻松的语气回应患者的问题；（2）帮助患者全面认识到慢性盆腔炎疾病的发病原因、治疗方式以及治疗效果等，告知患者该疾病的治疗需要较长的时间周期，帮助患者以积极的心态面对疾病，并且帮助建立治愈的信心；（3）为患者提供舒适、放松的治疗环境，多种供患者转移注意力的方式，比如看电视、听音乐、专家讲座等；（4）提供饮食指导，建议患者的饮食以易消化、高蛋白、高热量为主，要每天补充维生素，适当增加锻炼，提高身体免疫力，促进疾病疗效。
        5.3观察指标  （1）比较两组患者护理后的临床效果；痊愈：临床症状消失，血液指标水平正常；显效：临床症状大部分改善，血液指标水平正常；有效：临床症状部分改善，血液指标较治疗前趋于正常水平；无效：临床症状改善不明显，血液指标较治疗前没有明显变化，临床效果=痊愈+显效+有效；（2）比较两组患者护理前后的心理焦虑抑郁变化，通过焦虑自评量表（SAS）和抑郁自评量表（SDS）来评估，评分越高表明焦虑或者抑郁情况越严重。
        5.4  统计学方法  采用Spss18.0系统软件进行处理分析，计量资料以均数±标准差（x±s）表示，实施t检验;计数资料采用卡方检验，P＜0.05 为具有统计学意义。
        六、实证结果
        6.1  比较两组患者护理后的临床效果比较  结果显示，护理后，研究组的临床疗效明显优于对照组（P<0.05）。
        表1  两组患者护理后的临床效果比较 （n %）
        组别	痊愈	显效	有效	无效	临床疗效
        对照组（n=25）	7（28.00）	7（28.00）	5（20.00）	8（32.00）	17（68.00）
        研究组（n=25）	10（40.00）	8（32.00）	5（20.00）	2（8.00）	23（92.00）
        χ2					4.50
        P					P<0.05
        6.2 两组患者护理前后的心理焦虑抑郁变化  研究显示，护理前，两组SAS和SDS评分没有明显区别（P>0.05），护理后，研究组的两个评分明显较低（P<0.05）。
        表2  两组患者护理前后的心理焦虑抑郁变化 （分）
        组别	SAS评分	SDS评分
            护理前	护理后	护理前	护理后
        对照组（n=25）	59.37±4.58	48.57±3.03	58.26±3.51	46.28±3.88
        研究组（n=25）	59.38±3.99	38.46±3.47	58.31±3.48	37.15±4.01
        t	0.0082	10.97	0.051	8.18
        P	0.99	P<0.05	0.96	P<0.05
        七、讨论
        慢性盆腔炎不仅危害患者身体健康，降低生活质量，还容易引起并发症，产生严重后果。为改善患者病情，对疾病实施有效的方法治疗和调理，对改善临床症状，促进预后有重要意义。在治疗慢性盆腔炎方面，除了采取常规的中医、西医治疗，还可以通过心理护理的方式。本次研究报告表明，心理护理不仅可以改善慢性盆腔炎患者的临床疗效，还可以有效改善患者的焦虑抑郁情况，值得临床推广。
        参考文献
        [1]	邓艳春.心理护理对慢性盆腔炎治疗的效果影响分析[J].中文科技期刊数据库(文摘版)医药卫生, 2022(5):3.
        [2]	彭姬源.慢性盆腔炎治疗现状及进展[J].中国科技期刊数据库医药, 2021(4):2.

        ”””
        我想让你帮我仿照这篇范文生成一篇应用专题报告，题目是：《{name}》
        要求：
        尽量详细地去写，字数越多越好。起码要写到{int(word_count)+500}到{int(word_count)+1000}个中文字符（这个是硬性要求，字数必须满足要求）。
        文章要看起来专业、规范；
        也需要生成一个类似范文中的数理实证内容，包括t检验和卡方检验与对应p值，置于表格中。请尽量提供更多细节，看起来更加详实；
        文章中请至少包括：前言与研究现状、问题确立、研究目的、文献查证、资料与方法、实证、实证结果分析、讨论、参考文献等模块；
        二级标题下要有三级标题，最好实证内容看起来规范可靠；
        可以在文章最后生成几篇看上去很真实的中文参考文献，参考文献和作者必须看上去非常规范、专业、可靠；
        请只输出这篇报告内容，请不要输出其他的内容；
        {other_requirment}
        你可以帮我运用你的专业知识写这篇文章吗?

        """}
    ]
    while len(response_text)<(word_count/3):
        response_text = generate_chat_completion(messages,model='gpt-3.5-turbo-16k')
    return response_text

def report_generation_second(first_text,dic):
    response_text=""
    name=dic['name']
    word_count=2000
    if str(dic['word_count'])!='nan':
        word_count=dic['word_count']
    messages = [
        {"role": "user", "content": f"""
        这是一篇应用专题报告，带数理实证检验的内容。
        “““    
        {first_text}
        ”””

        这篇报告目前的问题是字数达不到要求。要求的字数为{int(word_count)+500}到{int(word_count)+1000}个中文字符。我想让你帮我把这篇应用专题报告扩充一下生成内容并直接加到原文中，并将满足字数要求的、扩充后的报告全文生成给我，题目还是：《{name}》
        要求：
        你要做的事情就是对原文进行改写，改写后输出的字数越多越好。改写扩充后的全文需要达到{int(word_count)+500}到{int(word_count)+2000}个中文字符，这个是硬性要求，必须在这个字符数量范围内或比这个字符数量更多。
        你需要将原文中的每一句话都尝试变成相同意思但是字数更多的两三个句子来扩充全文的字数，此外，你也可以在原文中插入可以补充原文逻辑的句子，比如现象阐述的后面你可以加入更加细致的解释；
        请一定注意，绝对不要续写，也不要超过改写原文所要表达的含义。请不要将扩充内容作为单独的一部分续写在文章后；
        最好实证内容看起来规范可靠；
        添加的内容必须符合原文的行文排版逻辑；
        请将所有参考文献附在文章最后面，最终输出的中间段落不要出现参考文献部分；
        请只输出这篇报告内容，请不要输出其他的内容。
        你可以帮我运用你的专业知识写这篇文章吗?

        """}
    ]
    while len(response_text)<(word_count/3):
        response_text = generate_chat_completion(messages,model='gpt-3.5-turbo-16k')    
    return response_text


##普刊

def pukan_generation_first(dic):
    response_text=""
    name=dic['name']
    
    other_requirment=dic['other']
    if str(other_requirment)=='nan':
        other_requirment=""
    
    author=dic['author']
    if str(author)=='nan':
        auther_message="这篇报告不需要注明作者。"
    else:
        auther_message=f"作者:{author}。"
        
        
    
    word_count=2000
    if str(dic['word_count'])!='nan':
        word_count=dic['word_count']
        
    messages = [
        {"role": "user", "content": f"""
        这是一篇普刊格式的文章，带数理实证检验的内容。
        “““
                凝血酶联合西咪替丁治疗新生儿上消化道出血的临床疗效
        一、	摘要
        目的：对于患有上消化道出血的新生儿，为其治疗当中，凝凝血酶、西咪替丁的治疗效果。方法：上消化道出血新生儿共纳入82例（病例选取时间：2022.3~2023.12），随机分成常规组、对照组（常规疗法，41例）、实验组（凝血酶+西咪替丁，41例），观察两组不良反应发生情况、止血时间、临床疗效。结果：实验组治疗效果比对照组好，止血时间较对照组短，不良反应发生率低于对照组，（P＜0.05）。结论 ：对于上消化道出血新生儿，联合使用凝血酶与西咪替丁治疗，疗效显著，能够提升止血效果，减少不良反应。
        关键词：新生儿上消化道出血；西咪替丁；凝血酶

        二、引言
        新生儿消化道出血会对患儿消化道的任何位置造成影响，其主要分成上下消化道出血[1]。有较多的因素可使新生儿出现消化道出血的症状，例如：颅内出血等等[2]。患儿会有便血及呕血等等症状，病情严重的患者会有失血性的休克，有着极高的死亡率。对于胃酸过度分泌进行抑制，能够使得胃黏膜受损程度明显减轻，属于避免患儿死亡的主要措施，所以应为患者选择科学的治疗方案。有鉴于此，本次纳入82例上消化道出血新生儿共，探讨凝血酶+西咪替丁的疗效，如下。三、资料与方法
        3.1一般资料
        上消化道出血新生儿共纳入82例（病例选取时间：2022.3~2023.12），随机分成常规组、对照组（常规疗法，41例）、实验组（凝血酶+西咪替丁，41例），对照组：日龄：3~7天 , 平均日龄：（5.02±0.22）天，性别：男21例，女20例。研究组：日龄：3~6天 , 平均日龄：（5.03±0.23）天，性别：男20例，女21例，两组患儿一般资料，无显著差异，（P＞0.05）。
        3.2方法
        对照组：常规治疗法：在保障血糖水平与血压水平稳定时，适当的为患儿补充维生素K1；若是患儿存在感染情况为其抗感染。洗胃用1%碳酸氢钠进行，再为患儿注射125 mg 的酚磺乙胺注射液（企业：万邦德制药集团有限公司；批准文号：国药准字H13022163），进行止血，静脉滴注30 毫升10%谷氨酰胺合成酶，每天1次，持续治疗2到7天。 
        实验组：西咪替丁+凝血酶治疗：通过胃管于胃内予以患儿0.5 kU的矛头腹蛇血凝酶（批准文号：国药准字H20041418；生产企业：蓬莱诺康药业有限公司），且经胃管为其注射5 毫升的生理盐水，之后予以其0.5kU的血凝酶，持续一周。且予以其西咪替丁注射液（生产企业：无锡济煜山禾药业股份有限公司批准文号：国药准字H32026476），滴速为14 mg/（kg·h），每隔8 个小时为患者滴注一次，治疗5天到7天。
        3.3观察指标
        （1）不良反应发生：呕吐、腹泻等，发生率=发生数/总数*100%。（2）止血时间：统计止血时间。临床疗效：治疗48 h 内止血，胃引流液检查未见新血液渗出，无咖啡样液体，治疗72 小时以内，便血情况阴。（3）疗效评估标准：显效：止血时间在治疗后48小时以内，新血液未渗出，且有咖啡样的液体，治疗72小时之内便血已经转阴；有效：止血时间在治疗3天后，胃引流液血液等明显减少，症状改善明显；治疗3 天依然出血，可观察到新血液，便血检验结果为阳性。有效率与显效率和。
        3.4统计学方法
        计量资料以（ ±s）展示，检验t，计数资料以%表示，2检验；P＜0.05则数据差异显著。 统计学软件：SPSS21.0，
        四、结果  
        4.1两组不良反应发生情况比较
        不良反应发生率：和对照组比较，实验组低，（P＜0.05），见下表。
        表1比较不良反应发生情况[n，（%）]
        组别	例数	呕吐	恶心	腹泻	总发生率
        对照组	41	2(4.88)	2(4.88)	2(4.88)	6(14.64)
        实验组	41	1(2.44)	0（0.00）	0（0.00）	1(2.44)
        2	-	-	-	-	3.905 
        P	-	-	-	-	0.048 
        4.2两组止血时间及两组治疗效果比较
        两组治疗效果比较，和对照组比较，实验组优，两组止血时间比较，实验组短（P＜0.05），见表3。
        表2比较治疗效果、止血时间[n，（%）， ±s]
        组别	例数	显效	有效	无效	总有效率	住院时间（天）
        对照组	41	12(29.27)	18(43.90)	11(26.83)	30(73.17)	2.06±0.11
        实验组	41	17(41.46)	23(56.10)	1(2.44)	40(97.56)	1.55±0.21
        2	-	-	-	-	9.762	13.775
        P	-	-	-	-	0.002	＜0.001
        五、讨论
        临床当前为新生儿消化道出血常规治疗方法为对其生命体征监测同时实施针对性止血疗效，主要包含禁食、减轻胃肠刺激、采用维生素K1等等常规疗效，造必要时为患儿输血。但是对于有消化道出血症状的新生儿，常规治疗方案，在改善其止血效果方面均需要进一步提升[4]。
        凝血酶在与出血的病灶接触之后，会使索状凝固膜形成，可以增强平滑肌收缩功能，可以促使出血创面收敛，能够使毛细血管通透性下降，并且可以使得局部炎症反应与局部水肿，能明显的减轻，有着较好的止血效果。凝血酶也能够集聚血小板，使得溶胶状态纤维蛋白原逐渐转化成纤维蛋白，进而可以堵塞上消化道出血点。为患者治疗前为其进行洗胃，能够使中和胃内的氢离子。西咪替丁有着组胺以及增强免疫效果，能对创面修复，使得胃黏膜可再生。西咪替丁可抑制因为刺激因素导致的胃酸分泌，使胃内的pH酸度可以降低，能够将应激性的胃溃疡症状得以改善，止血效果显著。实验组止血时间比对照组更短，治疗效果比对照组好，不良反应发生率比对照组低。说明联合使用凝血酶与西咪替丁止血效果好，安全性高。
        综上：针对上消化道出血新生儿将凝血酶和西咪替丁联合使用为患儿治疗，效果显著，对提升止血效果，减少不良反应具有重要意义。
        参考文献：
        [1]郭修平.奥曲肽联合凝血酶治疗肝硬化合并上消化道出血的疗效[J].黑龙江医药科学,2023,46(01):192-193+196.
        [2]叶士芬,姜立飞.注射用白眉蛇毒凝血酶联合奥美拉唑治疗急性消化道出血的临床效果[J].临床合理用药,2023,16(01):73-75.



        ”””
        我想让你帮我仿照这篇范文生成一篇普刊格式的文章，题目是：《{name}》
        要求：
        尽量详细地去写，字数越多越好。起码要写到{int(word_count)+500}到{int(word_count)+1000}个中文字符（这个是硬性要求，字数必须满足要求）。
        文章要看起来专业、规范；
        也需要生成一个类似范文中的数理实证内容，包括t检验和卡方检验与对应p值，置于表格中。请尽量提供更多细节，看起来更加详实；
        文章中请至少包括：摘要、（目的、方法、结果和结论）、关键词、引言、资料与方法（一般资料、方法、观察指标、统计学方法）、结果、讨论与结论、参考文献等模块；
        二级标题下要有三级标题，最好实证内容看起来规范可靠；
        可以在文章最后生成几篇看上去很真实的中文参考文献，参考文献和作者必须看上去非常规范、专业、可靠；
        请只输出这篇文章内容，请不要输出其他的内容；
        {other_requirment}
        你可以帮我运用你的专业知识写这篇文章吗?

        """}
    ]
    while len(response_text)<(word_count/3):
        response_text = generate_chat_completion(messages,model='gpt-3.5-turbo-16k')
    return response_text

def pukan_generation_second(first_text,dic):
    response_text=""
    name=dic['name']
    word_count=2000
    if str(dic['word_count'])!='nan':
        word_count=dic['word_count']
    messages = [
        {"role": "user", "content": f"""
        这是一篇普刊格式的文章，带数理实证检验的内容。
        “““    
        {first_text}
        ”””

        这篇文章目前的问题是字数达不到要求。要求的字数为{int(word_count)+500}到{int(word_count)+1000}个中文字符。我想让你帮我把这篇文章扩充一下生成内容并直接加到原文中，并将满足字数要求的、扩充后的报告全文生成给我，题目还是：《{name}》
        要求：
        你要做的事情就是对原文进行改写，改写后输出的字数越多越好。改写扩充后的全文需要达到{int(word_count)+500}到{int(word_count)+2000}个中文字符，这个是硬性要求，必须在这个字符数量范围内或比这个字符数量更多。
        你需要将原文中的每一句话都尝试变成相同意思但是字数更多的两三个句子来扩充全文的字数，此外，你也可以在原文中插入可以补充原文逻辑的句子，比如现象阐述的后面你可以加入更加细致的解释；
        请一定注意，绝对不要续写，也不要超过改写原文所要表达的含义。请不要将扩充内容作为单独的一部分续写在文章后；
        最好实证内容看起来规范可靠；
        添加的内容必须符合原文的行文排版逻辑；
        请将所有参考文献附在文章最后面，最终输出的中间段落不要出现参考文献部分；
        请只输出这篇报告内容，请不要输出其他的内容。
        你可以帮我运用你的专业知识写这篇文章吗?

        """}
    ]
    while len(response_text)<(word_count/3):
        response_text = generate_chat_completion(messages,model='gpt-3.5-turbo-16k')    
    return response_text


def yewu_generation_first(dic):
    response_text=""
    name=dic['name']
    
    other_requirment=dic['other']
    if str(other_requirment)=='nan':
        other_requirment=""
    
    author=dic['author']
    if str(author)=='nan':
        auther_message="这篇报告不需要注明作者。"
    else:
        auther_message=f"作者:{author}。"
        
        
    
    word_count=2000
    if str(dic['word_count'])!='nan':
        word_count=dic['word_count']
        
    messages = [
        {"role": "user", "content": f"""
        这是一篇业务报告格式的文章，带数理实证检验的内容。
        “““
                超敏C反应蛋白和肺部感染的方面业务报告
        一、摘要
        目的：探讨超敏C反应蛋白（hs-CRP）在肺部感染诊断中的应用价值。方法：选取2021年1月至2022年9月在某基层医院住院治疗的肺部感染患者120例（其中有重复住院者），按照感染类型分为细菌性肺部感染组（60例）和病毒性肺部感染组（60例），同时选取同期健康体检者60例作为对照组。检测各组血清hs-CRP水平，并分析其与感染类型、严重程度、预后等因素的关系。结果：细菌性肺部感染组的hs-CRP水平明显高于病毒性肺部感染组和对照组，差异有统计学意义（P<0.05）。hs-CRP水平与肺部感染患者的白细胞计数、中性粒细胞比例、C反应蛋白（CRP）水平、降钙素原（PCT）水平呈正相关（r>0.5，P<0.05），与住院时间、住院费用、死亡率呈正相关（r>0.5，P<0.05）。hs-CRP水平在治疗前后有显著下降（P<0.05），且下降幅度与治疗效果呈正相关（r>0.5，P<0.05）。结论：hs-CRP是一种敏感、特异的肺部感染诊断指标，可反映感染类型、严重程度、预后及治疗效果，对于基层医院的临床诊治具有重要意义。
        [关键词] 超敏C反应蛋白；肺部感染；诊断；预后

        二、引言
        肺部感染是一种常见的呼吸系统疾病，主要由细菌、病毒或真菌等微生物引起，可导致肺泡或支气管内膜的炎症反应，影响气体交换功能[1]。肺部感染的临床表现多样，包括发热、咳嗽、咳痰、胸闷、呼吸困难等。肺部感染的诊断主要依据临床表现、实验室检查和影像学检查[2]。然而，在基层医院，由于条件限制，常规的实验室检查和影像学检查往往不能及时、准确地进行，导致肺部感染的诊断困难和误诊率高[3]。因此，寻找一种简便、快速、敏感、特异的肺部感染诊断指标，对于基层医院的临床诊治具有重要意义。
        超敏C反应蛋白（hypersensitive C-reactive protein，hs-CRP）是一种由肝脏合成的急性期蛋白，是机体发生急性或慢性炎症时产生的非特异性标志物[4]。hs-CRP的检测方法比常规的C反应蛋白（CRP）检测方法更敏感，能检测到极低浓度的CRP，反映低水平的炎症状态。近年来，越来越多的研究表明，hs-CRP在心血管疾病、糖尿病、肿瘤等多种疾病的诊断、预后和治疗中具有重要的应用价值[5]。然而，hs-CRP在肺部感染诊断中的应用价值尚不明确，尤其是在基层医院的实际情况下，hs-CRP是否能作为一种有效的肺部感染诊断指标，值得进一步探讨。
        本研究选取2021年1月至2022年9月在某基层医院住院治疗的肺部感染患者120例（其中有重复住院者），按照感染类型分为细菌性肺部感染组（60例）和病毒性肺部感染组（60例），同时选取同期健康体检者60例作为对照组。检测各组血清hs-CRP水平，并分析其与感染类型、严重程度、预后等因素的关系，旨在探讨hs-CRP在肺部感染诊断中的应用价值。

        三、资料及方法
        3.1 一般资料
        本研究选取2021年1月至2022年9月在某基层医院住院治疗的肺部感染患者120例（其中有重复住院者），按照感染类型分为细菌性肺部感染组（60例）和病毒性肺部感染组（60例）。细菌性肺部感染组包括肺炎链球菌肺炎、金黄色葡萄球菌肺炎、铜绿假单胞菌肺炎等；病毒性肺部感染组包括流感病毒肺炎、呼吸道合胞体病毒肺炎、新型冠状病毒肺炎等。感染类型的判断依据为临床表现、实验室检查和影像学检查。同时选取同期健康体检者60例作为对照组，无呼吸系统相关的临床表现和异常检查结果。三组间年龄、性别分布无统计学差异（P>0.05），具体见表1。
        表1 三组患者一般资料比较
        组别	例数	年龄（岁）	性别（男/女）
        细菌性肺部感染组	60	61.23±15.37	32/28
        病毒性肺部感染组	60	59.67±16.42	31/29
        对照组	60	58.53±14.29	30/30
        P值	-	0.670	0.980
        注：数据均以均数±标准差表示
        3.2 方法
        所有患者在入院后24小时内采集静脉血5ml，离心后分离血清，采用免疫比浊法检测血清hs-CRP水平，采用酶联免疫吸附法（ELISA）检测血清CRP水平，采用荧光免疫法检测血清降钙素原（PCT）水平，采用自动血球分析仪检测白细胞计数（WBC）和中性粒细胞比例（NEUT%）。所有检测方法均按照生物试剂盒的说明书操作。对于治疗后出院的患者，记录其住院时间、住院费用、治疗效果等情况。治疗效果分为显效、有效和无效，判断标准为：显效：临床症状和体征消失，实验室检查和影像学检查恢复正常；有效：临床症状和体征明显改善，实验室检查和影像学检查有所好转；无效：临床症状和体征无明显改善或加重，实验室检查和影像学检查无改善或恶化。对于治疗后死亡的患者，记录其死亡原因和死亡时间。对于重复住院的患者，只记录其最后一次住院的情况。
        3.3 观察指标
        主要观察指标为各组患者的血清hs-CRP水平；次要观察指标为各组患者的WBC、NEUT%、CRP、PCT水平，以及肺部感染患者的住院时间、住院费用、死亡率、治疗效果等情况。
        3.4 统计学方法
        数据的处理和分析采用SPSS 22.0软件。计量资料以均数±标准差表示，组间比较采用单因素方差分析或t检验，多重比较采用LSD-t法；计数资料以百分数表示，组间比较采用卡方检验或Fisher精确概率法。相关性分析采用Pearson相关系数或Spearman秩相关系数。结果中P<0.05表示差异有统计学意义。
        四、结果
        4.1 各组患者的hs-CRP水平比较
        细菌性肺部感染组的hs-CRP水平明显高于病毒性肺部感染组和对照组，差异有统计学意义（P<0.050）。病毒性肺部感染组的hs-CRP水平也高于对照组，但差异无统计学意义（P>0.050）。具体见表2。
        表2 各组患者的hs-CRP水平比较（mg/L）
        组别	hs-CRP
        细菌性肺部感染组	12.34±3.56
        病毒性肺部感染组	6.78±2.45
        对照组	5.67±1.23
        P值	<0.050
        注：数据均以均数±标准差表示
        4.2 hs-CRP水平与其他指标的相关性分析
        hs-CRP水平与肺部感染患者的WBC、NEUT%、CRP、PCT水平呈正相关（r>0.500，P<0.050），与住院时间、住院费用、死亡率呈正相关（r>0.500，P<0.050）。具体见表3。
        表3 hs-CRP水平与其他指标的相关性分析
        指标	r	P
        WBC	0.630	<0.050
        NEUT%	0.680	<0.050
        CRP	0.720	<0.050
        PCT	0.750	<0.050
        住院时间	0.540	<0.050
        住院费用	0.570	<0.050
        死亡率	0.590	<0.050
        注：数据均以Pearson相关系数表示
        4.3 hs-CRP水平在治疗前后的变化
        肺部感染患者在治疗前后的hs-CRP水平有显著下降（P<0.050），且下降幅度与治疗效果呈正相关（r>0.500，P<0.050）。具体见表4。
        表4 hs-CRP水平在治疗前后的变化（mg/L）
        治疗前后	hs-CRP
        治疗前	9.56±3.21
        治疗后	6.34±2.34
        P值	<0.050
        注：数据均以均数±标准差表示
        五、讨论
        hs-CRP是一种由肝脏合成的急性期蛋白，是机体发生急性或慢性炎症时产生的非特异性标志物。hs-CRP的检测方法比常规的CRP检测方法更敏感，能检测到极低浓度的CRP，反映低水平的炎症状态。近年来，越来越多的研究表明，hs-CRP在心血管疾病、糖尿病、肿瘤等多种疾病的诊断、预后和治疗中具有重要的应用价值。然而，hs-CRP在肺部感染诊断中的应用价值尚不明确，尤其是在基层医院的实际情况下，hs-CRP是否能作为一种有效的肺部感染诊断指标，值得进一步探讨。
        本研究发现，细菌性肺部感染组的hs-CRP水平明显高于病毒性肺部感染组和对照组，差异有统计学意义。这说明hs-CRP能够区分细菌性和病毒性肺部感染，对于感染类型的判断有一定的参考价值。这可能与细菌性肺部感染引起的炎症反应更强、更持久有关。hs-CRP作为一种急性期蛋白，能够敏感地反映机体的炎症状态，其水平与细菌数量和毒力成正比。而病毒性肺部感染引起的炎症反应相对较弱、较短暂，其水平与细菌数量和毒力成负比。因此，hs-CRP能够区分细菌性和病毒性肺部感染。然而，本研究也发现，病毒性肺部感染组的hs-CRP水平也高于对照组，但差异无统计学意义。这可能与病毒性肺部感染的病原体种类和检测方法的敏感性有关。病毒性肺部感染的病原体包括流感病毒、呼吸道合胞体病毒、新型冠状病毒等，其中一些病毒具有较强的变异能力，可能导致hs-CRP检测方法的敏感性降低。此外，病毒性肺部感染的诊断还需要结合其他的实验室检查和影像学检查，单纯依靠hs-CRP水平可能存在一定的局限性。因此，hs-CRP在区分细菌性和病毒性肺部感染时，仍需结合临床综合判断，不能作为唯一的诊断依据。
        本研究还发现，hs-CRP水平与肺部感染患者的WBC、NEUT%、CRP、PCT水平呈正相关，与住院时间、住院费用、死亡率呈正相关。这说明hs-CRP能够反映肺部感染患者的炎症程度、严重程度、预后及治疗效果。这可能与hs-CRP作为一种急性期蛋白，能够敏感地反映机体的炎症反应和免疫功能有关。hs-CRP能够激活补体系统，增强吞噬细胞的吞噬功能，促进白细胞和中性粒细胞的释放和迁移，从而增加WBC和NEUT%。hs-CRP还能够诱导内皮细胞表达黏附分子，促进白细胞和血小板的黏附和聚集，从而增加CRP和PCT。hs-CRP水平越高，说明机体的炎症反应越强，免疫功能越低，肺部感染越严重，住院时间越长，住院费用越高，死亡率越高。
        本研究还发现，肺部感染患者在治疗前后的hs-CRP水平有显著下降，且下降幅度与治疗效果呈正相关。这说明hs-CRP能够监测肺部感染患者的治疗反应和恢复情况。这可能与hs-CRP作为一种急性期蛋白，能够快速地响应机体的炎症变化有关。当机体受到感染刺激时，hs-CRP水平迅速升高；当机体消除感染刺激时，hs-CRP水平迅速下降。因此，hs-CRP水平能够及时地反映肺部感染患者的治愈程度，对于指导临床治疗具有重要意义。
        六、体会
        hs-CRP是一种敏感、特异的肺部感染诊断指标，可反映感染类型、严重程度、预后及治疗效果，对于基层医院的临床诊治具有重要意义。然而，本研究也存在一些局限性，如样本量较小、观察时间较短、未考虑其他影响因素等，需要进一步的研究来验证和完善。

        参考文献
        [1]施惠娟.全自动血液细胞仪在超敏C反应蛋白联合血常规检验诊断小儿细菌性感染性疾病中的价值[J].中国冶金工业医学杂志,2023,40(03):333-334.

        
        ”””
        我想让你帮我仿照这篇范文生成一篇业务报告格式的文章，题目是：《{name}》
        要求：
        尽量详细地去写，字数越多越好。起码要写到{int(word_count)+500}到{int(word_count)+1000}个中文字符（这个是硬性要求，字数必须满足要求）。
        文章要看起来专业、规范；
        也需要生成一个类似范文中的数理实证内容，包括t检验等统计量和对应p值，置于表格中。请尽量提供更多细节，看起来更加详实；
        文章中请至少包括：摘要、引言、资料与方法（一般资料、方法、观察指标、统计学方法）、结果、讨论、体会等模块；
        二级标题下要有三级标题，最好实证内容看起来规范可靠；
        体会部分需要详细去写；
        可以在文章最后生成几篇看上去很真实的中文参考文献，参考文献和作者必须看上去非常规范、专业、可靠；
        请只输出这篇文章内容，请不要输出其他的内容；
        {other_requirment}
        你可以帮我运用你的专业知识写这篇文章吗?

        """}
    ]
    while len(response_text)<(word_count/3):
        response_text = generate_chat_completion(messages,model='gpt-3.5-turbo-16k')
    return response_text

def yewu_generation_second(first_text,dic):
    response_text=""
    name=dic['name']
    word_count=2000
    if str(dic['word_count'])!='nan':
        word_count=dic['word_count']
    messages = [
        {"role": "user", "content": f"""
        这是一篇业务报告格式的文章，带数理实证检验的内容。
        “““    
        {first_text}
        ”””

        这篇文章目前的问题是字数达不到要求。要求的字数为{int(word_count)+500}到{int(word_count)+1000}个中文字符。我想让你帮我把这篇文章扩充一下生成内容并直接加到原文中，并将满足字数要求的、扩充后的报告全文生成给我，题目还是：《{name}》
        要求：
        你要做的事情就是对原文进行改写，改写后输出的字数越多越好。改写扩充后的全文需要达到{int(word_count)+500}到{int(word_count)+2000}个中文字符，这个是硬性要求，必须在这个字符数量范围内或比这个字符数量更多。
        你需要将原文中的每一句话都尝试变成相同意思但是字数更多的两三个句子来扩充全文的字数，此外，你也可以在原文中插入可以补充原文逻辑的句子，比如现象阐述的后面你可以加入更加细致的解释；
        请一定注意，绝对不要续写，也不要超过改写原文所要表达的含义。请不要将扩充内容作为单独的一部分续写在文章后；
        最好实证内容看起来规范可靠；
        添加的内容必须符合原文的行文排版逻辑；
        请将所有参考文献附在文章最后面，最终输出的中间段落不要出现参考文献部分；
        请只输出这篇报告内容，请不要输出其他的内容。
        你可以帮我运用你的专业知识写这篇文章吗?

        """}
    ]
    while len(response_text)<(word_count/3):
        response_text = generate_chat_completion(messages,model='gpt-3.5-turbo-16k')    
    return response_text

def kepu_first(dic):
    response_text=""
    name=dic['name']
    
    other_requirment=dic['other']
    if str(other_requirment)=='nan':
        other_requirment=""
    
    author=dic['author']
    if str(author)=='nan':
        auther_message="这篇报告不需要注明作者。"
    else:
        auther_message=f"作者:{author}。"
        
    word_count=1500
    if str(dic['word_count'])!='nan':
        word_count=dic['word_count']

    messages = [
        {"role": "user", "content": f"""
        我想让你帮我运用你的全面知识生成一篇科普文章，题目是：《{name}》
        要求：
        1. 文章的字数要求很高，因此请每一段中都详细去写。文章的总字数起码要写到{int(word_count)+500}到{int(word_count)+1500}个中文字符（这个是硬性要求，字数必须满足要求）。对于每一个环节你都可以使用很华丽和看起来非常专业的语言将你想表达的句子复杂化以扩充字数。要尽量扩充文章内容。
        2. 要求看起来专业、规范，分不同的模块。
        3. 请生成看起来内容充实、丰富，包含尽量多的对科普对象各种方面的文章。你可以根据这个题目去定义一些科普文章可能会有的不同模块，每一个环节都可以分不同的小点去详细介绍，每个小点都可以使用尽量多的词语去从更多的方面去描述，用尽量多的词汇去表达内容以扩充字数；
        4. 不需要参考文献；
        5. 请按照标题、文章的严格顺序进行输出，不要输出其他的内容；
        {other_requirment}
        你可以帮我运用你的专业知识写一下这篇文章吗？

        """}
    ]

    while len(response_text)<(word_count/3):
        response_text = generate_chat_completion(messages,model='gpt-4')
    return response_text



def kepu_second(first_text,dic):
    response_text=""
    name=dic['name']
    
    other_requirment=dic['other']
    if str(other_requirment)=='nan':
        other_requirment=""
    
    author=dic['author']
    if str(author)=='nan':
        auther_message="这篇报告不需要注明作者。"
    else:
        auther_message=f"作者:{author}。"
        
    word_count=1500
    if str(dic['word_count'])!='nan':
        word_count=dic['word_count']

    messages = [
        {"role": "user", "content": f"""
        这是一篇科普文章。
        “““    
        {first_text}
        ”””
        这篇文章目前问题是字数没有达到要求.我要求全文达到{int(word_count)+800}到{int(word_count)+2000}个中文字符。
        你需要对原文进行改写以最终满足字数要求。请对原文进行改写，将改写的内容替换掉原文内容，而不要直接在原文后添加扩充内容。并将改写后的全文作为输出，题目不变。
        要求：
        你要做的事情就是对原文进行改写，改写后输出的字数越多越好。改写扩充后的全文需要达到{int(word_count)+800}到{int(word_count)+2000}个中文字符，这个是硬性要求，必须在这个字符数量范围内或比这个字符数量更多。
        你需要将原文中的每一句话都尝试变成相同意思但是字数更多的两三个句子来扩充全文的字数，此外，你也可以在原文中插入可以补充原文逻辑的句子，比如现象阐述的后面你可以加入更加细致的解释；
        你可以将原文中的一些简单的话改写以增加字数，或者对一些分类的小点进行更加详细的解释，或者阅读全文寻找有没有可以添加的段落并进行扩充；
        请一定注意，绝对不要续写，也不要超过改写原文所要表达的含义。请不要将扩充内容作为单独的一部分续写在文章后；
        请保持原文的小标题和结构；
        扩写后的文章总结部分应当只有一小段，改写内容应主要集中在文章前中段；
        请不要添加参考文献；
        请不要输出除了文章本身之外的其他内容。
        你可以帮我运用你的专业知识写这篇文章吗?

        """}
    ]
    while len(response_text)<(word_count/3):
        response_text = generate_chat_completion(messages,model='gpt-3.5-turbo-16k')
    return response_text

def kepu_tpf(dic):
    response_text=""
    name=dic['name']
    
    other_requirment=dic['other']
    if str(other_requirment)=='nan':
        other_requirment=""
    
    author=dic['author']
    if str(author)=='nan':
        auther_message="这篇报告不需要注明作者。"
    else:
        auther_message=f"作者:{author}。"
        
    word_count=1500
    if str(dic['word_count'])!='nan':
        word_count=dic['word_count']

    messages = [
        {"role": "user", "content": f"""
        这是一篇科普报告范文：
        ~~~
        《磁共振与心脏疾病：揭示心脏内部的秘密》
        心脏是人体最重要的器官之一，承担着泵血、供氧和运输养分等重要功能。然而，心脏疾病却是全球范围内导致死亡的主要原因之一。为了更好地了解和诊断心脏疾病，科学家们开发了一项重要的技术——磁共振成像（Magnetic Resonance Imaging，MRI）。磁共振成像利用磁场和无害的无线电波来获取人体内部的详细图像，特别是在心脏领域的应用尤为重要。通过磁共振成像，我们可以揭示心脏内部的秘密，深入了解心脏的结构、功能和血流情况，从而实现更准确的心脏疾病诊断和治疗。本文将带您深入探索磁共振与心脏疾病的关系，解密心脏内部的奥秘。
        一、磁共振成像的原理和基本步骤
        1.1 原理
        磁共振成像利用磁场和无害的无线电波来获取人体内部的详细图像，通过利用水分子中的氢原子产生的共振信号，重建出人体内部的结构图像。人体的组织和器官中含有丰富的水分和氢原子，当置于强磁场中时，氢原子会发生共振现象，即吸收并释放特定频率的无线电波。通过探测和分析这些共振信号，可以生成高分辨率的图像，展示人体内部组织和器官的细节。这种原理使得磁共振成像成为一种非侵入性、无放射性的先进影像技术，广泛应用于医学诊断和研究。
        1.2 步骤
        进行心脏磁共振成像通常包括以下步骤：
        a) 患者准备：患者需要躺在磁共振扫描床上，并确保身体保持静止。在某些情况下，可能需要给患者注射对比剂以提高图像质量。
        b) 定位和校准：通过将患者放置在特定的位置和方向，并使用参考标记，确保图像的准确性和一致性。
        c) 图像获取：利用磁场和无线电波，扫描仪会对心脏进行多个方向的扫描，收集数据以生成图像。这通常需要数分钟到数十分钟的时间。
        d) 数据处理和分析：通过计算机软件对收集到的数据进行处理和分析，生成高质量的图像，供医生进行诊断和评估。
        二、磁共振成像在心脏疾病诊断中的应用
        2.1 结构与功能评估
        心脏磁共振成像可以提供高分辨率的心脏图像，帮助医生评估心脏的结构和功能。通过观察心脏的大小、形状、心室和心房的壁厚度、心瓣的功能等指标，医生可以了解心脏是否存在结构异常或功能异常，进而判断是否存在心脏疾病。例如，心肌梗死导致的心肌损伤可以在磁共振图像中清晰可见，有助于确定患者的病情和制定治疗方案。
        2.2 冠状动脉成像
        冠状动脉是供应心肌血液的重要血管，冠状动脉疾病是心脏疾病的主要原因之一。心脏磁共振成像可以通过注射对比剂和高分辨率的成像技术来评估冠状动脉的状况。医生可以观察冠状动脉的通畅程度、是否存在狭窄或堵塞等问题，以及评估心肌缺血的程度，有助于制定治疗策略。
        2.3 纤维化检测
        心脏疾病常常伴随着心肌纤维化的发生，这是心脏肌肉组织的异常增生和纤维组织沉积。磁共振成像可以使用特殊的扫描序列来检测和定量纤维化程度。通过观察图像中的纤维化区域，医生可以判断心肌损伤的范围和严重程度，并评估患者的预后情况。
        2.4 心脏运动与血流分析
        心脏磁共振成像可以通过动态图像来分析心脏的运动和血流情况。医生可以观察心脏收缩和舒张的过程，评估心肌的收缩功能和心脏瓣膜的功能情况。此外，磁共振成像还可以提供心脏血流的定量测量，了解心脏各区域的灌注情况，帮助诊断和评估心肌缺血和心肌梗死等疾病。
        三、心脏磁共振的优势和局限性
        3.1 优势
        a) 非侵入性：心脏磁共振成像是一种非侵入性的检查方法，不需要使用放射性物质或插入导管等。这降低了患者的风险和不适感，并减少了检查的并发症。
        b) 多维信息：心脏磁共振成像可以提供多维信息，包括心脏的结构、功能、血流动力学等。这使得医生可以全面评估心脏疾病，制定更准确的治疗方案。
        c) 高分辨率：心脏磁共振成像具有高分辨率的特点，能够清晰显示心脏的细微结构和病变区域。这使得医生能够更准确地定位和评估心脏病变，提高诊断的准确性。
        3.2 局限性
        a) 依赖患者合作度：心脏磁共振成像需要患者在狭小的扫描仪中保持静止，并耐受长时间的扫描过程。对于一些无法配合或有严重焦虑的患者，可能会造成困扰或不适。
        b) 对金属物质的敏感性：由于心脏磁共振成像使用强磁场，患者身上的金属物质（如心脏起搏器、人工心脏瓣膜等）可能对成像质量产生干扰。因此，具有金属植入物的患者可能不适合进行心脏磁共振检查。
        c) 相对较高的成本：与其他常规影像检查方法相比，心脏磁共振成像的设备和操作成本相对较高。这可能限制了它在某些地区或医疗机构的普及程度。
        四、结论
        综上所述，心脏磁共振成像是一项强大而有效的非侵入性检查技术，可以提供高分辨率的心脏图像，帮助医生评估心脏的结构、功能、血流动力学以及心肌纤维化情况。它在心脏疾病的诊断、治疗效果评估和疾病监测中发挥着重要作用。尽管心脏磁共振成像存在一些局限性，如患者合作度要求高、对金属物质敏感以及相对较高的成本，但其优势远大于不足之处。随着技术的不断进步和应用的推广，心脏磁共振成像将在心脏疾病领域展现出更大的潜力和发展空间，为患者的健康带来更多的利益。

        ~~~
        我想让你帮我运用你的全面知识仿照这篇范文生成一篇科普文章，题目是：《{name}》
        要求：
        1. 尽量详细地去写，字数越多越好。文章的总字数起码要写到{int(word_count)+500}到{int(word_count)+1500}个中文字符（这个是硬性要求，字数必须满足要求）。对于每一个环节你都可以使用很华丽和看起来非常专业的语言将你想表达的句子复杂化以扩充字数。要尽量扩充文章内容。
        2. 要求看起来专业、规范，分不同的小标题。一级标题使用汉字数字加‘、’识别（如：一、二、），二级标题使用阿拉伯数字和‘.’识别（如1.1 1.2）。
        3. 请生成看起来内容充实、丰富，包含尽量多的对科普对象各种方面的文章。你可以根据这个题目去定义一些科普文章可能会有的不同模块，每一个环节都可以分不同的小点去详细介绍，每个小点都可以使用尽量多的词语去从更多的方面去描述，用尽量多的词汇去表达内容以扩充字数；
        4. 文章的第一段必须是‘一、引言’，结尾不需要参考文献；
        5. 请按照标题、文章的严格顺序进行输出，不要输出其他的内容；
        {other_requirment}
        你可以帮我运用你的专业知识写一下这篇文章吗？

        """}
    ]

    while len(response_text)<(word_count/3):
        response_text = generate_chat_completion(messages,model='gpt-3.5-turbo-16k')
    return response_text

def yykepu_tpf(dic):
    response_text=""
    name=dic['name']
    
    other_requirment=dic['other']
    if str(other_requirment)=='nan':
        other_requirment=""
    
    author=dic['author']
    if str(author)=='nan':
        auther_message="这篇报告不需要注明作者。"
    else:
        auther_message=f"作者:{author}。"
        
    word_count=1500
    if str(dic['word_count'])!='nan':
        word_count=dic['word_count']

    messages = [
        {"role": "user", "content": f"""
        这是一篇科普报告范文：
        ~~~
                人体扫描之谜：X射线如何揭示身体的隐秘？
        一、引言
        X射线是一项重要的医学成像技术，被广泛应用于诊断和治疗领域。通过使用X射线，医生们可以揭示人体内部的隐秘，了解身体的结构、异常和疾病。本文将带您深入探索X射线在人体扫描中的原理、应用和局限性，解密如何利用X射线揭示身体的内部。
        二、X射线成像的原理
        X射线成像是基于X射线在物体中的吸收和散射特性，通过获取物体内部的图像信息。当X射线经过物体时，不同组织对其吸收能力不同，导致图像呈现出差异。骨骼和金属物质对X射线的吸收能力较强，呈现出明亮的区域，而软组织对X射线的吸收能力较弱，呈现出暗影区域。通过观察和分析这些影像，医生可以了解身体内部结构和异常情况。
        通过X射线成像可以获得人体内部的高分辨率图像，这对于诊断和治疗是非常重要的。X射线成像技术的诞生和发展为医疗领域带来了革命性的变化，使医生能够更准确地了解患者的病情，并根据这些信息做出相应的治疗方案。
        三、X射线在人体扫描中的应用
        当谈到人体扫描，X射线技术的应用在医学领域无疑是一项卓越的成就。通过在人体内部引入无形的X射线，医生们能够揭示身体内部的隐秘，从而更准确地诊断疾病，规划治疗方案，甚至挽救生命。以下将更详细地探讨X射线在不同方面的应用。
        3.1 骨骼成像
        3.1.1 骨折诊断和评估
        在遭受外伤或意外事故后，骨折成为常见问题。X射线在骨折诊断和评估中起到了关键作用。通过将X射线投射到受伤部位，医生能够获得关于骨骼状态的影像，从而准确判断骨折的类型（如横向骨折、纵向骨折、骨裂等）、位置以及骨折程度。这种能力让医生能够选择最合适的治疗方法，无论是保守治疗还是手术干预。
        3.1.2 骨质疏松检测
        随着年龄增长，骨质疏松风险逐渐升高。通过X射线吸收测量，医生可以评估骨骼的密度，以此来判断骨骼的健康状态。这项技术可以帮助发现骨骼中的钙含量减少，从而引发骨质疏松症。早期的诊断让患者能够尽早采取预防和治疗措施，减少骨折风险。
        3.2 胸部成像
        3.2.1 肺部疾病诊断
        胸部X射线是最常用的肺部疾病诊断方法之一。肺部疾病，如肺炎、结核、肿瘤等，可以通过X射线影像来识别。医生可以观察肺部纹理和阴影，确定异常区域的位置，从而进行深入的检查和治疗。此外，胸部X射线还有助于观察肋骨的状态，以检测是否存在异常。
        3.2.2 心脏疾病评估
        胸部X射线不仅用于评估肺部，还可以提供关于心脏健康的信息。医生可以观察心脏的形状、大小和位置，了解心脏是否扩大、是否存在瓣膜问题，甚至评估冠状动脉的状况。这对于诊断心脏疾病至关重要，因为心脏问题常常表现为胸部疼痛和不适。
        3.3 腹部成像
        3.3.1 腹部器官评估
        腹部X射线是评估腹部器官健康的方法之一。通过观察肝脏、胆囊、胃、肾脏、脾脏等器官的影像，医生可以评估它们的大小、形态和位置，以判断是否存在异常情况，如肿瘤、结石、囊肿等。这种非侵入性的成像方法可以在早期发现问题，为患者提供更早的治疗。
        3.3.2 消化道问题诊断
        腹部X射线还有助于检测消化道问题，如肠道梗阻、胃肠道穿孔等。医生可以通过观察消化道的影像，判断是否有肠道阻塞、积气等问题。这有助于及早诊断消化道疾病，为患者提供相应治疗方案。
        3.4 放射治疗
        3.4.1 癌症治疗
        放射治疗是在癌症治疗中广泛使用的方法之一。高能X射线能够杀灭癌细胞，阻止其生长，从而缩小肿瘤体积并减少其对周围组织的影响。通过精确的计算和定位，医生可以将X射线能量精确地照射到肿瘤部位，最大限度地降低健康组织受损的风险。
        3.4.2 定位和计划
        放射治疗需要精确的定位和计划。医生会利用不同角度的X射线影像，确定肿瘤的位置和周围解剖结构。然后，他们会制定详细的照射计划，确保X射线能够准确地照射到肿瘤，同时最小化对周围健康组织的损伤。这种精确性有助于提高放射治疗的效果，降低患者的不良影响。
        3.5 血管造影
        3.5.1 静脉造影
        静脉造影是一种利用X射线成像来观察静脉系统的检查方法。在静脉造影中，医生会将一种对X射线有强吸收能力的造影剂通过静脉注射进入患者的血液循环中。这种造影剂会随着血流进入不同的静脉，通过X射线成像设备可以清晰地显示出静脉的轮廓和分布情况。通过观察静脉造影图像，医生可以判断静脉的通畅性、是否存在血栓形成以及其他静脉系统的异常情况。
        3.5.2 动脉造影
        动脉造影是一种通过向动脉内注入造影剂，然后使用X射线成像设备来观察动脉系统的检查方法。在动脉造影中，医生会将造影剂注入特定的动脉，造影剂会随着血流进入动脉系统，通过X射线成像可以清晰地显示出动脉的轮廓、狭窄程度以及其他异常情况。动脉造影在心血管疾病的诊断和治疗中具有重要作用，可以帮助医生判断动脉粥样硬化、动脉狭窄、血栓形成等问题，从而指导后续的治疗方案
        四、结论
        综上所述，X射线成像是一项重要的医学成像技术，可以揭示人体内部的结构、异常和疾病。它在骨骼成像、胸部成像、腹部成像和放射治疗等方面都有广泛应用。尽管X射线成像具有易操作、快速和相对较低成本的优势，但也存在辐射暴露和对柔软组织限制的局限性。随着技术的不断进步和应用的推广，X射线成像将继续发展，并发挥更大的作用，为人体健康带来更多的利益。

        ~~~
        我想让你帮我运用你的全面知识仿照这篇范文生成一篇科普文章，题目是：《{name}》
        要求：
        1. 尽量详细地去写，字数越多越好。文章的总字数起码要写到{int(word_count)+500}到{int(word_count)+1500}个中文字符（这个是硬性要求，字数必须满足要求）。对于每一个环节你都可以使用很华丽和看起来非常专业的语言将你想表达的句子复杂化以扩充字数。要尽量扩充文章内容。
        2. 要求看起来专业、规范，分不同的小标题。一级标题使用汉字数字加‘、’识别（如：一、二、），二级标题使用阿拉伯数字和‘.’识别（如1.1 1.2）。
        3. 请生成看起来内容充实、丰富，包含尽量多的对科普对象各种方面的文章。你可以根据这个题目去定义一些科普文章可能会有的不同模块，每一个环节都可以分不同的小点去详细介绍，每个小点都可以使用尽量多的词语去从更多的方面去描述，用尽量多的词汇去表达内容以扩充字数；
        4. 文章的第一段必须是‘一、引言’，结尾不需要参考文献；
        5. 请按照标题、文章的严格顺序进行输出，不要输出其他的内容；
        {other_requirment}
        你可以帮我运用你的专业知识写一下这篇文章吗？

        """}
    ]

    while len(response_text)<(word_count/3):
        response_text = generate_chat_completion(messages,model='gpt-3.5-turbo-16k')
    return response_text

def find_paragraph(text):
    #return list, including each title, paragraph without title and according position
    title_pos_list=[]
    chinese_number_list=["一、","二、","三、","四、","五、","六、","七、","八、","九、","十、"]
    for title in chinese_number_list:
        if text.find(title)!=-1:
            if "\n" in text[text.find(title)-10:text.find(title)+20]:
                title_pos_list.append(text.find(title))
    final_list=[]
    for n_para in range(len(title_pos_list)-1):
        this_para=text[title_pos_list[n_para]:title_pos_list[n_para+1]]
        title_this_para=this_para[:this_para.find('\n')]
        this_para_without_title=this_para[len(title_this_para):]
        final_list.append([title_this_para,this_para_without_title,title_pos_list[n_para]])
    return final_list


#############################
########扩段-具体函数##############
###########################
def para_extend(text):
    response_text=''
    messages = [
        {"role": "user", "content": f"""
        这是一篇文章中的一段，这个段落的问题是字数不够。我希望你能对这个段落的原文进行改写扩充，改写后输出的全文字数越多越好。你需要将原文中的每一句话都尝试变成相同意思但是字数更多的两三个句子来扩充全文的字数，此外，你也可以在原文中插入可以补充原文逻辑的句子，比如现象阐述的后面你可以加入更加细致的解释。切记，是改写这个段落，不要续写后续内容，也请不要总结内容，不要超过这一段所要表达的含义，最终目标是扩充本段字数，相比原文字数越多越好。请你将改写之后的内容直接发给我。
    
        ~~~
        {text}
        ~~~

        """}
    ]
    model_kind='gpt-3.5-turbo'
    if len(text)>1500:
        model_kind='gpt-3.5-turbo-16k'
    while len(response_text)<200:
        response_text = generate_chat_completion(messages,model=model_kind)
    return response_text

#####################
##总扩段函数#########

def extend_with_num(text,number_of_para):
    para_list=find_paragraph(text)
    enlarged_text=para_extend(para_list[number_of_para][1])
    extended_text=text[:para_list[number_of_para][2]+len(para_list[number_of_para][0])]+'\n'+enlarged_text+'\n'+text[para_list[number_of_para][2]+len(para_list[number_of_para][0])+len(para_list[number_of_para][1]):]
    return extended_text

###################################################################
#####字数统计函数：输入文章，输出包括中文、数字和标点符号的数量###############
####################################################################
def count_zishu(input_text: str):
    count_chinese=0
    for s in input_text:
        if s in string.ascii_letters:
            count_chinese+=0.2
        elif s.isalpha() or s.isdigit():
            count_chinese+=1
    return int(count_chinese)



##########################################################
############多进程函数#####################################


import multiprocessing

def article_generation(line):
    today=datetime.date.today()
    author_message=""
    if str(line['author'])!='nan':
        author_message=str(line['author'])+"-"
    RESPONSE=""
    
    
    
    ###############
    
    ##生成
    if (f"{author_message}{line['name']}.docx" or f"not_enough_{author_message}{line['name']}.docx") not in os.listdir(f"output/submission_{today}/{line['type']}"):
        length=0
        if line['type']=='科普':    
            while (length<line['word_count']*2/5):
                RESPONSE=kepu_tpf(line)
                length=int(count_zishu(RESPONSE))
            print(line['name'],'---第一次生成，字数：',length)
        
        elif line['type']=='应用科普':    
            while (length<line['word_count']*2/5):
                RESPONSE=yykepu_tpf(line)
                length=int(count_zishu(RESPONSE))
            print(line['name'],'---第一次生成，字数：',length)

        elif line['type']=='报告':
            while (length<line['word_count']*3/5):
                RESPONSE=report_generation_first(line)
                length=int(count_zishu(RESPONSE))
            print(line['name'],'---第一次生成，字数：',length)

        elif line['type']=='普刊':
            while (length<line['word_count']*1/2):
                RESPONSE=pukan_generation_first(line)
                length=int(count_zishu(RESPONSE))
            print(line['name'],'---第一次生成，字数：',length)
            # if len(RESPONSE) <line['word_count']+0:
            #     RESPONSE=pukan_generation_second(RESPONSE,line)
            #     print('第二次生成，字数：',count_zishu(RESPONSE))
        
        elif line['type']=='业务报告':
            while (length<line['word_count']*1/2):
                RESPONSE=yewu_generation_first(line)
                length=int(count_zishu(RESPONSE))
            print(line['name'],'---第一次生成，字数：',length)
            # if len(RESPONSE) <line['word_count']+0:
            #     RESPONSE=yewu_generation_second(RESPONSE,line)
            #     print('第二次生成，字数：',count_zishu(RESPONSE))
        
        
        ##############################################################################
        ##扩段
        
        ##为扩段预留多余字数
        if line['type'] in ['报告','业务报告','普刊']:
                exceed_amount=0
        if line['type'] in ['科普','应用科普']:
                exceed_amount=0
                
        if len(find_paragraph(RESPONSE))>2:
            if count_zishu(RESPONSE)<line['word_count']+exceed_amount:
           
                ##v1 从文章的中后部分开始扩########################
                ##逻辑设置为40-80为重点修改部分。如果字数不足再向两边拓展################

                ##开始位置
                percentage=0.6

                para_list=range(len(find_paragraph(RESPONSE)))
                start_pos=int(len(find_paragraph(RESPONSE))*percentage)
                rearrange_result = []
                rearrange_result.append(para_list[start_pos])

                offset = 1  # 初始偏移量
                while start_pos - offset >= 0 or start_pos + offset < len(para_list):
                    if start_pos + offset < len(para_list):
                        rearrange_result.append(para_list[start_pos + offset])  # 处理中心位置+偏移量
                    if start_pos - offset >= 0:
                        rearrange_result.append(para_list[start_pos - offset])  # 处理中心位置-偏移量
                    offset += 1

                ###得到重排序后的段落列表：rearrange_result
                for i in rearrange_result:
                    if i!=0:#不扩第一段
                        if line['type'] in ['报告','业务报告','普刊']:
                            if ('实证' not in find_paragraph(RESPONSE)[i][0]) and ('结果' not in find_paragraph(RESPONSE)[i][0]):
                                RESPONSE=extend_with_num(RESPONSE,i)
                                print(line['name'],f'---扩第{i+1}段，字数：',int(count_zishu(RESPONSE)))
                        if line['type'] in ['科普','应用科普']:
                            RESPONSE=extend_with_num(RESPONSE,i)
                            print(line['name'],f'---扩第{i+1}段，字数：',int(count_zishu(RESPONSE)))
                        if count_zishu(RESPONSE)>line['word_count']+exceed_amount:
                            break
        ##############################################################################                    
        ##输出结果
        if count_zishu(RESPONSE) > line['word_count']+exceed_amount:
            writedocx(RESPONSE,line['name'],f"output/submission_{today}/{line['type']}/{author_message}"+line['name'])
            print(line['name']+'————完成', "包括数字的中文字符数:",int(count_zishu(RESPONSE)))
            
        else:
            writedocx(RESPONSE,line['name'],f"output/submission_{today}/{line['type']}/not_enough_{author_message}"+line['name'])
            print(line['name']+'————字数不足', "包括数字的中文字符数:",int(count_zishu(RESPONSE)))
            




########################################
##########主函数#########################


def main(file_name, num_tasks):
    today=datetime.date.today()
    try:
        os.mkdir(f"output/submission_{today}")
        os.mkdir(f"output/submission_{today}/报告")
        os.mkdir(f"output/submission_{today}/科普")
        os.mkdir(f"output/submission_{today}/应用科普")
        os.mkdir(f"output/submission_{today}/普刊")
        os.mkdir(f"output/submission_{today}/业务报告")
    except:
        pass
    

    files = read_excel(f'{file_name}.xlsx')  
    pool = multiprocessing.Pool(processes = num_tasks)


    for n in range(len(files)):
        file=files[n]
        pool.apply_async(article_generation, args=(file,))

    pool.close()
    pool.join()

    print(f"生成结束")


############################
##########运行主函数##########

if __name__ == '__main__':
    main(file_name='input/8.29',num_tasks=7)
