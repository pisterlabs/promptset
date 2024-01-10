"""
This file contains various helper functions for a Telegram bot.

Functionality provided by this program includes:
    data_parser: 
        Parses news data from a specified URL and returns a list of news articles.
    learning_resources_links: 
        Retrieves a list of learning resource links related to a given query.
    weather_classified: 
        Classifies weather-related text and extracts the city name mentioned in the text.
    AI_responce: 
        Generates an AI response based on a user's conversation history using the OpenAI GPT-3.5 Turbo model.
    Encryption: 
        Provides encryption and decryption functions.
    translate_text: 
        Translates text from one language to another using a translation service.
    generate_image: 
        Generates an image based on a specified query.
    imageClassification:
         Performs image classification using a pre-trained model.
    backGroundRemove: 
        Removes the background from an input image and saves the result.
    document_creating_for_speccial_topic: 
        Creates a document with a specified topic, group, your name, and instructor name.
    find_weekday: Finds the weekday for a given date.

"""

   

import json
import openai


import locale
import config
import easyocr

import requests
import datetime
import itertools

import numpy as np
import pytesseract
import pyshorteners
import urllib.request

from pyowm import OWM
import tensorflow as tf

from PIL import Image
from rembg import remove
from docx import Document
from docx.shared import Pt
from pymystem3 import Mystem

from langdetect import detect

from bs4 import BeautifulSoup
from googlesearch import search
from translate import Translator
...
from datetime import datetime, timedelta
from pyowm.commons.exceptions import NotFoundError


from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from tensorflow.keras.applications.resnet50 import preprocess_input, decode_predictions



    


"""                                                     """
"""    Data parsing and news info based on query        """ 
"""                                                     """



def data_parser(url):
    """
    Parses news data from the specified URL and returns a list of news articles.
    
    Args:
        url (str): The URL of the webpage containing the news data.
    
    Returns:
        list: A list of news articles extracted from the webpage.
    """
    
    news = []
    response = requests.get(url)
    soup = BeautifulSoup(response.content, "html.parser")

    for el in soup.select(".newsfeed-column"):
        title_links = el.select('a')
        for title in title_links:
            new = title.get_text()
            news.append(new)
    
    return news[:5]


def newsInfo(query):

    api_key = '4df3aad8b5804bab879012bab2d7b8cc'
    url = 'https://newsapi.org/v2/top-headlines'

    params = {
        'apiKey': api_key,
        'country': 'ua',
        'language': 'uk',
        'q': query
    }

    response = requests.get(url, params=params)
    data = response.json()

    if data['status'] == 'ok':
        articles = data['articles']
        for article in articles:
            title = article['title']
            description = article['description']
           
            return title, description
    else:
        print('Error fetching news.')




"""                                                   """
"""             Resources links to query              """
"""                                                   """
def learning_resources_links(query):
    """
    Retrieves a list of learning resource links related to the given query.
    
    Args:
        query (str): The query to search for learning resources.
    
    Returns:
        list: A list of learning resource URLs.
    """
    search_results = search(f"Информация про {query}", advanced=True)
    limited_results = itertools.islice(search_results, 3)  # Limit to 3 results


    url_list = []
    for result in limited_results:
        if not result.url.endswith('.ru'):  # Exclude .ru domain URLs
            url_list.append(result.url)

    return url_list


import requests
from bs4 import BeautifulSoup
import urllib.request


def search_pdf(query):
    """
    Searches the web for PDF files for a given query and downloads them.

    Args:
        query (str): The query to search.

    Returns:
        None
    """
    url = "https://www.google.com/search"
    params = {
        "q": query + " filetype:pdf",
        "num": 1
    }
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/80.0.3987.149 Safari/537.36"
    }
    
    response = requests.get(url, params=params, headers=headers)
    
    soup = BeautifulSoup(response.text, "html.parser")
    results = soup.select(".yuRUbf a")
    
    fixed_filename = "pdf_file.pdf"  # Фиксированное имя файла
    
    for result in results:
        link = result["href"]
        if link.endswith(".pdf"):
            try:
                response = requests.get(link, stream=True)
                with open(fixed_filename, "wb") as f:
                    for chunk in response.iter_content(chunk_size=1024):
                        f.write(chunk)
                print("Скачан файл:", fixed_filename)
            except Exception as e:
                print("Не удалось скачать файл:", fixed_filename)





"""                                                   """
"""    Weather classification and getting info        """
"""                                                   """
def weather_classified(text):
    """
    This function, weather_classified, is used to classify weather-related text and extract the city name mentioned in the text.
 
    Data:
        text (str): The input text containing weather-related information.

    Returns:
        city_result (str): The extracted city name from the input text.
        The function performs the following steps:

        Defines a set of stop words and weather-related keywords.
        Converts the input text to lowercase and splits it into tokens.
        Filters out the tokens that are weather-related keywords or stop words.
        Retrieves the last remaining token as the potential city name.
        Initializes the Mystem object from the pymystem3 library.
        Lemmatizes the potential city name using Mystem to normalize the word form.
        Returns the extracted city name as the final result.
    """
    
    stop_words = {'как', 'какая', 'какие', 'какой', 'мне', 'нужна', 'скажи', 'сейчас', 'там', 'что', 'в', 'город', 'города', 'городе', 'погода', 'погоде', 'поживает', 'як', 'яка', 'які', 'який', 'мені', 'потрібна', 'скажи', 'зараз', 'там', 'що', 'в', 'місто', 'міста', 'місті', 'погода', 'погоді', 'поживає'}
    weather_keywords = {'погода', 'город', "місто", }
    tokens = text.lower().split()
    filtered_tokens = [token for token in tokens if token not in weather_keywords and token not in stop_words]
    city_result = filtered_tokens[-1] if filtered_tokens else ''
    
    m = Mystem()
    
    city_name = city_result
    lemmas = m.lemmatize(city_name)
    city_result = lemmas[0]

    return city_result


def get_weather(text):
    """ 
    Get info about the weather in the city

    """ 
    classified_place = weather_classified(text)
    print(classified_place)
    

    # Determine the weather
    owm = OWM('6366e634aae1cdf800864fe8535f22b3')
    mgr = owm.weather_manager()

    try: 
        # Search for current weather in Kyiv  and get details
        observation = mgr.weather_at_place(classified_place)  # -> This uses the function from functions file
        w = observation.weather # all data 

        detailed_status = w.detailed_status  # 'clouds'
        wind = w.wind()['speed'] # {'speed': 4.6, 'deg': 330} -> example
        humidity = w.humidity    # 87   Humidity
        temperature = w.temperature('celsius')['temp'] 
        rain = w.rain # example ->  {}
        clouds = w.clouds # example -> 75

       
        
        return f"_У місті {classified_place.capitalize()} {round(temperature)} градусів за Цельсієм_",temperature
        
    except NotFoundError:
        return "Місто не знайдено, спробуйте ввести інше місто"




"""                                           """
"""    Init OpenAI and generate answer        """
"""                                           """


async def AI_response(user_input, chat_history, user_name, api_key):
    """
    This function, AI_response, is used to generate a response from an AI chatbot based on the user input and chat history.

    Parameters:
    - user_input (str): The user's input message.
    - chat_history (list): The chat history containing the previous user and AI messages.
    - user_name (str): The name of the user.

    Returns:
    - response_content (str): The generated response from the AI chatbot.
    - updated_chat_history (list): The updated chat history including the user input and AI response.

    The function performs the following steps:
    - Initializes the chat message log with the provided chat history.
    - Appends the user's input message to the message log.
    - Limits the message log to the last 4 messages, if its length exceeds 4.
    - Specifies the model engine, maximum tokens, and OpenAI key for the chatbot.
    - Generates a response from the AI chatbot using OpenAI's ChatCompletion API.
    - Appends the AI's response to the message log.
    - Returns the generated response from the AI chatbot and the updated chat history.
    """

    message_log = chat_history

    if len(message_log) == 0:
        message_log.append({'role': 'user', 'content': f"Привет, я - {user_name}, {user_input}"})
    else:
        message_log.append({'role': 'user', 'content': user_input})
        
        if len(message_log) > 4:
            message_log = message_log[-4:]


    model_engine = "gpt-3.5-turbo"
    max_tokens = 564  # default 1024
    OpenAIKey = api_key
    openai.api_key = OpenAIKey
    
    response = openai.ChatCompletion.create(
        model=model_engine,
        messages=message_log,
        max_tokens=max_tokens,
        temperature=0.7,
        top_p=1,
        stop=None
    )

    # Обновляем лог сообщений с ответом ИИ
    message_log.append({
        'role': 'assistant',
        'content': response.choices[0].message.content
    })
   
    # Возвращаем текст ответа ИИ и обновленный лог сообщений
    return response.choices[0].message.content, message_log



def AI_responce_without_log(message, language_target: bool = False):
    """
    Generates an AI response using the OpenAI GPT-3.5 Turbo model.
    
    Args:
        message: Message of the user for whom the AI response is generated.
    
    Returns:
        AI answer for function document_creating_for_speccial_topic
    """
    
    openai.api_key = config.OpenAIKey
    model_engine = "gpt-3.5-turbo"
    max_tokens = 1024  # default 1024
    language = "спілкуємося українскою" if language_target == False else ""

    messages = [
        {"role": "system", "content": "You: " + message},
        {"role": "user", "content": f"{message}, {language}"}
    ]
   
    response = openai.ChatCompletion.create(
        model=model_engine,
        messages=messages,
        max_tokens=max_tokens,
        temperature=0.7,
        top_p=1,
        stop=None
    )

    for choice in response.choices:
        if "text" in choice:
            return choice.text
        
    # If no response with text is found, return the first response's content (which may be empty)
    return response.choices[0].message.content


class Encryption():
    ...
    

language_names = {
    'af': 'африкаанс',
    'ar': 'арабська',
    'bg': 'болгарська',
    'bn': 'бенгальська',
    'ca': 'каталонська',
    'cs': 'чеська',
    'cy': 'валлійська',
    'da': 'данська',
    'de': 'німецька',
    'el': 'грецька',
    'en': 'англійська',
    'es': 'іспанська',
    'et': 'естонська',
    'fa': 'перська',
    'fi': 'фінська',
    'fr': 'французька',
    'gu': 'гуджараті',
    'he': 'іврит',
    'hi': 'гінді',
    'hr': 'хорватська',
    'hu': 'угорська',
    'id': 'індонезійська',
    'it': 'італійська',
    'ja': 'японська',
    'kn': 'каннада',
    'ko': 'корейська',
    'lt': 'литовська',
    'lv': 'латвійська',
    'mk': 'македонська',
    'ml': 'малаялам',
    'mr': 'маратхі',
    'ne': 'непальська',
    'nl': 'голландська',
    'no': 'норвезька',
    'pa': 'панджабі',
    'pl': 'польська',
    'pt': 'португальська',
    'ro': 'румунська',
    'ru': 'російська',
    'sk': 'словацька',
    'sl': 'словенська',
    'so': 'сомалі',
    'sq': 'албанська',
    'sv': 'шведська',
    'sw': 'суахілі',
    'ta': 'тамільська',
    'te': 'телугу',
    'th': 'тайська',
    'tl': 'тагальська',
    'tr': 'турецька',
    'uk': 'українська',
    'ur': 'урду',
    'vi': 'вʼєтнамська',
    'zh-cn': 'китайська'
}

def translate_text(text, source_lang, target_lang, detect_language=False):
    """
    Translates the given text from the source language to the target language.
    
    Args:
        text (str): The text to be translated.
        source_lang (str): The source language code.
        target_lang (str): The target language code.
    
    Returns:
        str: The translated text.
    """

    if detect_language:
        source_lang = detect(text)
        if source_lang in language_names:
            source_lang = language_names[source_lang]
        translator = Translator(from_lang=source_lang, to_lang="uk")
        translation = translator.translate(text)
        return translation, source_lang
    else:
        translator = Translator(from_lang=source_lang, to_lang=target_lang)
        translation = translator.translate(text)
        return translation



"""                                                                 """
"""  1. Generate Image, using Openai neural network DALL-E API,     """
"""  2. Image Classification, using a pre-trained ResNet50 model,   """
"""  3. BackGround removing                                         """
"""                                                                 """
    
def generate_image(description: str) -> str:
    """
    Generates an image based on the provided description using the OpenAI API.
    
    Args:
        description (str): The description or prompt for generating the image.
    
    Returns:
        str: The URL of the generated image.
    """

    url = "https://api.openai.com/v1/images/generations"
    payload = {
        "model": "image-alpha-001",
        "prompt": description,
        "num_images": 1,
        "size": "512x512",
        "response_format": "url"
    }
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {config.OpenAIKey}"
    }
    response = requests.post(url, headers=headers, data=json.dumps(payload))
    if response.status_code == 200:
        image_url = response.json()["data"][0]["url"]
        return image_url
    else:
        raise ValueError(f"Ошибка {response.status_code}: {response.text}")




def imageClassification(image_path):
    """
    Performs image classification on the given image using a pre-trained ResNet50 model.
    
    Args:
        image_path (str): The path to the image file.
    
    Returns:
        str: The classification result with the highest probability.
    """

    # Loading a pre-trained ResNet50 model
    model = tf.keras.applications.ResNet50(weights='imagenet')

    # Image loading and preprocessing
    image_path = image_path
    image = Image.open(image_path)
    image = image.resize((224, 224))
    image = image.convert('RGB')  # Convert image to RGB mode
    image_array = np.array(image)
    image_array = preprocess_input(image_array[np.newaxis, :])

    # Image classification
    predictions = model.predict(image_array)
    decoded_predictions = decode_predictions(predictions, top=3)[0]
    for _, label, probability in decoded_predictions:
        translation = translate_text(label, "en", "uk", False)
        translation = translation.replace("NAME OF TRANSLATORS", "")
        print(f'Це {translation} з вірогідністю {probability * 100:.2f}%')
        return f'Це {translation} з вірогідністю {probability * 100:.2f}%'


def backGroundRemove(input_path, output_path):
    """
     A function to remove the background from an image.
    
     Arguments:
     - input_path: string, path to the input image.
     - output_path: string, path to save the processed image.
    """
    input_data = Image.open(input_path)
    output = remove(input_data)
    output.save(output_path)



""" Image to text """


def image_to_text(image_path):
    # Set the path to the Tesseract OCR executable
    pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'

    # Open image
    image = Image.open(image_path)

    # Text recognising
    text = pytesseract.image_to_string(image)

    return text



def uk_image_to_text(image_path):
    # Создание экземпляра OCR
    reader = easyocr.Reader(['ru', 'uk'])

    # Открытие изображения
    image = Image.open(image_path)

    # Распознавание текста
    results = reader.readtext(np.array(image))

    # Извлечение распознанного текста
    text = ' '.join([result[1] for result in results])

    return text


"""  Create a new document"""
def document_creating(discipline, group, studentName, teacher):
    """
    Creates a new document with the provided details.
    
    Args:
        discipline (str): The name of the discipline.
        group (str): The name of the group.
        studentName (str): The name of the student.
        teacher (str): The name of the teacher.
    
    Returns:
        Document: The created document object.
    """

    doc = Document()

     # Добавление данных на титульный лист
    doc.add_paragraph("МІНІСТЕРСТВО НАУКИ ТА ОСВІТИ УКРАЇНИ")
    doc.add_paragraph("НАЦІОНАЛЬНИЙ ТЕХНІЧНИЙ УНІВЕРСИТЕТ")
    doc.add_paragraph("«ХАРКІВСЬКИЙ ПОЛІТЕХНІЧНИЙ ІНСТИТУТ»")
    doc.add_paragraph("КАФЕДРА ІНФОРМАЦІЙНІ СИСТЕМИ ТА ТЕХНОЛОГІЇ")

    # Align paragraphs to the center
    
    doc.add_paragraph("\n")
    doc.add_paragraph(f"ЛАБОРАТОРНА РОБОТА № ")
    doc.add_paragraph(f"з дисципліни «{discipline}»")
    doc.add_paragraph("\n")
    doc.add_paragraph("Виконав:")
    doc.add_paragraph(f"студент {group}")
    doc.add_paragraph(f"{studentName}")
    doc.add_paragraph("\n")
    doc.add_paragraph("Перевірив:")
    doc.add_paragraph(f"проф. каф. ІСТ")
    doc.add_paragraph(f"{teacher}")
    doc.add_paragraph("\n")
    doc.add_paragraph("Харків-2023")
    
   

    for paragraph in doc.paragraphs[:7]:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for paragraph in doc.paragraphs[8:]:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Adding a page break
    doc.add_page_break()

    # Create a "ZVIT" section on a new page
    doc.add_paragraph("\n")
    doc.add_paragraph("\n")
    doc.add_paragraph("\n")
    doc.add_paragraph("\n")           
    doc.add_paragraph("ЗМІСТ")
    
    for paragraph in doc.paragraphs[22:23]:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
 

    # Adding sections and subsections
    try:
        gpt_answer = AI_responce_without_log(f"напиши Зміст для теми {discipline} без пояснень, одразу к ділу, ограничься 5  пунктами", True)
        # Split the response by newlines to separate each item
        items = gpt_answer.split('\n')
        # Remove empty items
        items = [item.strip() for item in items if item.strip()]

        if len(items) == 5:
            print("Valid response:")
            for item in items:
                print(item)
        else:
            print("Invalid response. Expected exactly 5 items.")

        for i, item in enumerate(items, start=1):
            doc.add_heading(item, level=2).runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    except Exception as e:
        print("An error occurred:", str(e))
        doc.add_heading("1 Завдання", level=2)
        doc.add_heading("2 Код програми на мові С", level=2)
        doc.add_heading("3 Покрокове пояснення коду", level=2)
        doc.add_heading("4 Демонстрація роботи програми", level=2)
        doc.add_heading("5 Висновок", level=2)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = False

    doc.save("yourVariant.docx")
   



def document_creating_for_speccial_topic(discipline):
    """
    Creates a document for a special topic based on the provided discipline.
    
    Args:
        discipline (str): The name of the discipline.
    
    Returns:
        Document: The created document object containing the special topic.
    """

    doc = Document()

    # Adding data to the cover page
    doc.add_paragraph("МІНІСТЕРСТВО НАУКИ ТА ОСВІТИ УКРАЇНИ")
    doc.add_paragraph("НАЦІОНАЛЬНИЙ ТЕХНІЧНИЙ УНІВЕРСИТЕТ")
    doc.add_paragraph("«ХАРКІВСЬКИЙ ПОЛІТЕХНІЧНИЙ ІНСТИТУТ»")
    doc.add_paragraph("КАФЕДРА ІНФОРМАЦІЙНІ СИСТЕМИ ТА ТЕХНОЛОГІЇ")

    # Align paragraphs to the center
    
    doc.add_paragraph("\n")
    doc.add_paragraph(f"ЛАБОРАТОРНА РОБОТА № ")
    doc.add_paragraph(f"з дисципліни «...»")
    doc.add_paragraph("\n")
    doc.add_paragraph("Виконав:")
    doc.add_paragraph(f"студент ")
    doc.add_paragraph(f"...")
    doc.add_paragraph("\n")
    doc.add_paragraph("Перевірив:")
    doc.add_paragraph(f"проф. каф. ІСТ")
    doc.add_paragraph(f"...")
    doc.add_paragraph("\n")
    doc.add_paragraph("Харків-2023")
    
   

    for paragraph in doc.paragraphs[:7]:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for paragraph in doc.paragraphs[8:]:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Adding a page break
    doc.add_page_break()
    
    # Create a "ZVIT" section on a new page
    doc.add_paragraph("\n")
    doc.add_paragraph("\n")
    doc.add_paragraph("\n")
          
    doc.add_paragraph("ЗМІСТ")
    
    for paragraph in doc.paragraphs[22:23]:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # Adding sections and subsections
    try:
        gpt_answer = AI_responce_without_log(f"напиши Зміст для теми {discipline} українською мовою без пояснень, одразу к ділу, ограничься 5  пунктами", True)
        # Split the response by newlines to separate each item
        items = gpt_answer.split('\n')
        # Remove empty items
        items = [item.strip() for item in items if item.strip()]

        if len(items) == 5:
            print("Valid response:")
            for item in items:
                print(item)
        else:
            print("Invalid response. Expected exactly 5 items.")

        for i, item in enumerate(items, start=1):
            doc.add_heading(item, level=2).runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    except Exception as e:
        print("An error occurred:", str(e))
        doc.add_heading("1 Завдання", level=2)
        doc.add_heading("2 Код програми на мові С", level=2)
        doc.add_heading("3 Покрокове пояснення коду", level=2)
        doc.add_heading("4 Демонстрація роботи програми", level=2)
        doc.add_heading("5 Висновок", level=2)
    
    doc.add_page_break()
    doc.add_paragraph("\n")
    doc.add_paragraph("\n") 

    """ Address to chatgpt Function, to get data about topics """

    paragraphFirst_main = doc.add_paragraph(items[0])
    gpt_answer = AI_responce_without_log(f"Напиши текст 150 слів на тему {items[0]} українською мовою, без пояснень, одразу к тексту", True)
    paragraphFirst = doc.add_paragraph(f"\t{gpt_answer}")
    doc.add_paragraph("\n")

    paragraphSecond_main = doc.add_paragraph(items[1])
    gpt_answer = AI_responce_without_log(f"Напиши текст 150 слів на тему {items[1]} українською мовою, без пояснень, одразу к тексту", True)
    paragraphSecond = doc.add_paragraph(f"\t{gpt_answer}")
    doc.add_paragraph("\n")

    paragraphThird_main = doc.add_paragraph(items[2])
    gpt_answer = AI_responce_without_log(f"Напиши текст 150 слів на тему {items[2]} українською мовою, без пояснень, одразу к тексту", True)
    paragraphThird = doc.add_paragraph(f"\t{gpt_answer}")
    doc.add_paragraph("\n")

    paragraphFour_main = doc.add_paragraph(items[3])
    gpt_answer = AI_responce_without_log(f"Напиши текст 150 слів на тему {items[3]} українською мовою, без пояснень, одразу к тексту")
    paragraphFour = doc.add_paragraph(f"\t{gpt_answer}")
    doc.add_paragraph("\n")

    paragraphFive_main = doc.add_paragraph(items[4])
    gpt_answer = AI_responce_without_log(f"Напиши текст 150 слів на тему {items[4]} українською мовою, без пояснень, одразу к тексту", True)
    paragraphFive = doc.add_paragraph(f"\t{gpt_answer}")


    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = False
            paragraphFirst_main.runs[0].bold = True
            paragraphSecond_main.runs[0].bold = True
            paragraphThird_main.runs[0].bold = True
            paragraphFour_main.runs[0].bold = True
            paragraphFive_main.runs[0].bold = True


    doc.save("yourVariant.docx")





def find_weekday():
    locale.setlocale(locale.LC_TIME)
    
    now = datetime.now()
    weekday = now.strftime("%A")
    if weekday == 'Monday':
        weekday = 'Понеділок'
    elif weekday == 'Tuesday':
        weekday = 'Вівторок'
    elif weekday == 'Wednesday':
        weekday = 'Середа'
    elif weekday == 'Thursday':
        weekday = 'Четверг'
    elif weekday == 'Friday':
        weekday = "П'ятниця"
    current_time = now.strftime("%H:%M")
    return weekday



