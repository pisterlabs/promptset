import pandas as pd
import openai
from docx import Document

# Charger le fichier CSV
df = pd.read_csv('/Users/antoine/Documents/GitHub.nosync/SWD.COVID.CONF/Data/Database/SWD.asymptomatic.sentences.sample.csv')

# Lire le contenu du premier prompt depuis le fichier Word
doc = Document('/Users/antoine/Documents/GitHub.nosync/SWD.COVID.CONF/Code/Scripts/GPT_3.5/Prompt.docx')
prompt_text_frame = "\n".join([para.text for para in doc.paragraphs])

# Supposons que le deuxième prompt est dans un autre document ou une autre partie du même document
doc_check = Document('/Users/antoine/Documents/GitHub.nosync/SWD.COVID.CONF/Code/Scripts/GPT_3.5/verification_prompt.docx')
prompt_text_check = "\n".join([para.text for para in doc_check.paragraphs])

# Configurer l'API OpenAI
openai.api_key = 

# Fonction pour interroger GPT-3.5 avec le premier prompt
def query_gpt3_frame(sentences):
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=prompt_text_frame.format(sentences=sentences),
        max_tokens=11,
        temperature=0,
        top_p=1
    )
    return response.choices[0].text.strip()

# Fonction pour interroger GPT-3.5 avec le deuxième prompt
def query_gpt3_check(sentences):
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=prompt_text_check.format(sentences=sentences),
        max_tokens=3,
        temperature=0,
        top_p=1
    )
    return response.choices[0].text.strip()

# Analyser chaque extrait avec le premier prompt et obtenir le cadre
df['frame'] = df['sentences'].apply(query_gpt3_frame)

# Analyser chaque extrait avec le deuxième prompt et obtenir la vérification
df['check'] = df['sentences'].apply(query_gpt3_check)

# Sauvegarder les résultats dans un nouveau fichier CSV
df.to_csv('/Users/antoine/Documents/GitHub.nosync/SWD.COVID.CONF/Data/Database/SWD.frame.analysis.asymptomatic.sentences.csv', index=False)
