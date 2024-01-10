import pandas as pd
import openai
from docx import Document

# Charger le fichier CSV
df = pd.read_csv('/Users/antoine/Documents/GitHub.nosync/SWD.COVID.CONF/Data/Database/QC.schools.sentences.sample.csv')
print("Fichier CSV chargé.")

# Lire le contenu du premier prompt depuis le fichier Word
doc = Document('/Users/antoine/Documents/GitHub.nosync/SWD.COVID.CONF/Code/Scripts/GPT_3.5/Prompt.docx')
prompt_text_frame = "\n".join([para.text for para in doc.paragraphs])
print("Premier prompt lu depuis le fichier Word.")

# Supposons que le deuxième prompt est dans un autre document ou une autre partie du même document
doc_check = Document('/Users/antoine/Documents/GitHub.nosync/SWD.COVID.CONF/Code/Scripts/GPT_3.5/verification_prompt.docx')
prompt_text_check = "\n".join([para.text for para in doc_check.paragraphs])
print("Deuxième prompt lu depuis le fichier Word.")

# Configurer l'API OpenAI
openai.api_key = 
print("API OpenAI configurée.")

# Fonction pour interroger GPT-3.5 avec le premier prompt
def query_gpt3_frame(sentences):
    full_prompt = prompt_text_frame.format(sentences=sentences)
    print(f"Envoi à GPT-3 (cadre) : {full_prompt}")
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Vous êtes un assistant de recherche en science politique. Votre tâche est de réaliser une analyse de cadre."},
                {"role": "user", "content": full_prompt}
            ]
        )
        result = response.choices[0].message['content'].strip()
        print(f"Réponse de GPT-3 (cadre) : {result}")
    except Exception as e:
        print(f"Une erreur s'est produite lors de l'interrogation de GPT-3 (cadre) : {e}")
        result = None
    return result

# Fonction pour interroger GPT-3.5 avec le deuxième prompt
def query_gpt3_check(sentences):
    full_prompt = prompt_text_check.format(sentences=sentences)
    print(f"Envoi à GPT-3 (vérification) : {full_prompt}")
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Vous êtes un assistant de recherche en science politique. Votre tâche est de réaliser une analyse de cadre."},
                {"role": "user", "content": full_prompt}
            ]
        )
        result = response.choices[0].message['content'].strip()
        print(f"Réponse de GPT-3 (vérification) : {result}")
    except Exception as e:
        print(f"Une erreur s'est produite lors de l'interrogation de GPT-3 (vérification) : {e}")
        result = None
    return result


# Analyser chaque extrait avec le premier prompt et obtenir le cadre
print("Début de l'analyse des cadres...")
df['frame'] = df['sentences'].apply(query_gpt3_frame)

# Analyser chaque extrait avec le deuxième prompt et obtenir la vérification
print("Début de la vérification des extraits...")
df['check'] = df['sentences'].apply(query_gpt3_check)

# Sauvegarder les résultats dans un nouveau fichier CSV
df.to_csv('/Users/antoine/Documents/GitHub.nosync/SWD.COVID.CONF/Data/Database/QC.frame.analysis.schools.sentences.csv', index=False)
print("Résultats sauvegardés dans un fichier CSV.")
