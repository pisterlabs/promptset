import os 
import openai
import docx
import platform

openai.api_key = os.getenv("OPENAI_API_KEY")

def create_word_document(cover_letter):
    doc = docx.Document()
    doc.add_paragraph(cover_letter)
    doc.save("cover_letter.docx")

def convert_to_pdf(pdf_filename):
    if platform.system() == 'Windows':
        import win32com.client
        wdFormatPDF = 17
        word = win32com.client.Dispatch('Word.Application')
        doc_path = os.path.abspath("cover_letter.docx")

        # Save the PDF in the "exports" subfolder
        pdf_path = os.path.abspath(os.path.join("Cover Letters", pdf_filename + ".pdf"))

        doc = word.Documents.Open(doc_path)
        doc.SaveAs(pdf_path, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()
    else:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph

        doc_path = "cover_letter.docx"
        pdf_path = "cover_letter.pdf"

        # Load the cover letter from Word document
        doc = docx.Document(doc_path)
        cover_letter = ""
        for paragraph in doc.paragraphs:
            cover_letter += paragraph.text + "\n"

        # Create PDF using ReportLab
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        content = [Paragraph(cover_letter, styles["Normal"])]
        doc.build(content)

def revise(cover_letter, messages):
    # Add the cover letter to the messages list
    messages.append({"role": "assistant", "content": cover_letter})
    # Prompt the user to write a revision prompt and add it to the messages list
    second_prompt = input("Write a revision prompt for the cover letter: ")
    messages.append({"role": "user", "content": second_prompt})
    
    # Get the response from the assistant
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",  # Use the gpt 3.5 model
        messages=messages,
        temperature=0.7,
        max_tokens=1000
    )
    print("ChatGPT Response:\n", response)
    cover_letter = response['choices'][0]['message']['content'].strip()
    return cover_letter

def generate_cover_letter():
    # Read the job description from the file
    job_description_file = "job_description.txt"
    with open(job_description_file, "r", encoding="utf-8") as f:
        job_description = f.read()

    # Read the resume from the file
    resume_file = "resume.txt"
    with open(resume_file, "r", encoding="utf-8") as f:
        resume = f.read()
    
    prompt = f"Write a cover letter for the following job description:\n{job_description}\n\nHere is my resume:\n{resume}\n\n"
    # Feel free to personalize the system message to include a more knowledgeable assistant in your respective field
    system_message = "You are a professional recruiter that helps people write cover letters for job applications. You focus on expressing the proper skills and experience for the job in short but effective manner."

    # Starting the conversation
    messages = [
        {"role": "system", "content": system_message},
        {"role": "user", "content": prompt}
    ]

    # Initial response from the assistant
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",  # Use the gpt 3.5 model
        messages=messages,
        temperature=0.7,
        max_tokens=1000
    )
    print("ChatGPT Response:\n", response)
    cover_letter = response['choices'][0]['message']['content'].strip()

    # write a loop where the user can revise the cover letter
    while True:
        yesorno = input("Would you like to revise the cover letter? (y/n): ")
        if (yesorno == "n" or yesorno == "no"):
            # Export the cover letter to a Word document and then convert it to a PDF and break out of the loop
            create_word_document(cover_letter)
            pdf_filename = input("Enter the PDF filename (without extension): ")
            convert_to_pdf(pdf_filename)
            print(f"Cover letter exported to PDF as '{pdf_filename}.pdf'.")
            break
        elif (yesorno != "y" and yesorno != "yes" and yesorno != "n" and yesorno != "no"):
            # Just print an error message and continue the loop
            print("Invalid input. Please input either y or n.")
        else:
            # Revise the cover letter and continue the loop
            cover_letter = revise(cover_letter, messages)     
    

if __name__ == "__main__":
    generate_cover_letter()
