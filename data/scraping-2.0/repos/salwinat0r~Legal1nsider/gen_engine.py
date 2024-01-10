import re
import docx
from docx import Document
import openai


openai.api_key = open("key.txt", 'r').read().strip('\n')

def generate_response(prompt):
    completion = openai.ChatCompletion.create(model="gpt-3.5-turbo",
    messages=[
        {"role": "system", "content": "You are legal document assistant"},
        {"role": "user", "content": prompt},
    ])
    response = completion["choices"][0]["message"]['content']
    return response


def extract_clauses_from_document(doc_path):
    doc = Document(doc_path)
    content = ""
    for paragraph in doc.paragraphs:
        content += paragraph.text + "\n"

    pattern1 = r'“([^”]+)”\s+shall have the meaning ascribed to it in Clause (\d+\.\d+|\d+)'
    pattern2 = r'Clause\s+(\d+\.\d+|\d+)\s*\(([^)]+)\)'
    pattern3 = r'“([^”]+)”\s+shall have the meaning ascribed to it under Clause ([^ ]+)'
    pattern4 = r'“([^”]+)”\s+shall have the meaning ascribed to such term in Clause (\d+\.\d+|\d+)'

    matches = re.findall(pattern1 + '|' + pattern2 + '|' + pattern3 + '|' + pattern4, content)

    clauses = []
    for match in matches:
        if match[0] != '':
            clause_name = match[0]
            clause_number = match[1]
        elif match[2] != '':
            clause_number = match[2]
            clause_name = match[3]
        elif match[4] != '':
            clause_name = match[4]
            clause_number = match[5]
        else:
            clause_name = match[6]
            clause_number = match[7]
        clauses.append({"clause_number": clause_number, "clause_name": clause_name.strip()})

    return clauses

def extract_title(file_path):
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    lines = text.split('\n\n')

# Search for a line that contains typical title formatting
    for line in lines:
        # Remove leading and trailing whitespace from the line
        line = line.strip()

        # Check if the line meets the criteria for a title (e.g., length, capitalization, etc.)
        if len(line) > 0 and line.isupper():
            return line
            break
    return None