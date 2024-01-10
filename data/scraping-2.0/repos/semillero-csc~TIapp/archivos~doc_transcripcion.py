# importamos las librerias
import streamlit as st
import datetime
import openai 
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
from decouple import config
from dotenv import load_dotenv
from email.message import EmailMessage
import ssl
import smtplib




def transcripcion_doc(list_transcripciones:dict):
    # Agrega un párrafo al documento
    encabezado: str = '''--------------------------------------------------------------------------------------------------------------------
Esta transcripción es producto de los desarrollos del equipo de la Linea
de investigación aplicada en Territorios Inteligentes, que hace parte del
grupo de investigación: Redes y Actores Sociales (RAS) del departamento de
Sociología de la Facultad de Ciencias Sociales y Humanas de la Universidad
de Antioquia. Su efectividad está supeditado a la calidad del audio, téngalo
en cuenta al momento de revisarlo.\t\n
--------------------------------------------------------------------------------------------------------------------'''
    


    metadata: str = '''- Nombre de la transcripción: {nombre_archivo}
- Fecha y hora en la que se realizó la transcripción: {fecha}
- Numero de palabras transcritas: {numero_palabras}

- Texto:'''

    template = '''"{texto}"\t\n
--------------------------------------------------------------------------------------------------------------------'''

    # Defino el nombre del documento
    if len(list_transcripciones) > 0:
        nombre = list_transcripciones[0]['nombre_archivo'][:-7]
        nombre_archivo_docx = f"archivos/docs/transcripcion_{nombre}.docx"

    # Crea un nuevo documento
    doc = docx.Document()

    #doc.add_paragraph(encabezado)

    # Agregar un título
    titulo = doc.add_paragraph(encabezado, style= 'Body Text')
    titulo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Alineación justificada

    for dato in list_transcripciones:
        print(template.format(**dato))

    # Agregar metedatos

    metadatos_ = doc.add_paragraph(metadata.format(**dato), style= 'Body Text')
    metadatos_.bold = True

    # Agregar un párrafo y configurar la alineación como justificada
    paragraph = doc.add_paragraph(template.format(**dato), style= 'Body Text')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Alineación justificada

    # doc.add_paragraph(template.format(**dato), style= 'Body Text' )
    try:
    # Guarda el documento en un archivo .docx
        if nombre_archivo_docx :
            doc.save(nombre_archivo_docx)
            st.success(f'Ya se ha creado su archivo .docx. Ahora puede dercargarlo dando click en "Download docx"')

            # Agregamos el boon de descarga
            with open(nombre_archivo_docx, "rb") as file:
                btn = st.download_button(
                        label="Download docx",
                        data=file,
                        file_name=nombre_archivo_docx,
                    )
                
        
                # Esto es para enviar el archivo a un correo
                imail_emisor = config('CORREO_PERSONAL')
                imail_contraseña = config('GOOGLE_KEY')
                imail_receptor = config('CORREO_U')
                asunto = 'Archivo_traducción'
                cuerpo = 'Se adjunta archivo con la traducción'

                em = EmailMessage()
                em['From'] = imail_emisor
                em['To'] = imail_receptor
                em['Subject'] = asunto
                em.set_content(cuerpo)

                with open(nombre_archivo_docx, "rb") as f:
                    em.add_attachment(
                        f.read(),
                        filename=nombre_archivo_docx,
                        maintype="application",
                        subtype="docx"
                    )

                contexto = ssl.create_default_context()

                with smtplib.SMTP_SSL('smtp.gmail.com', 465, context=contexto) as smtp:
                    smtp.login(imail_emisor, imail_contraseña)
                    smtp.sendmail(imail_emisor, imail_receptor, em.as_string()) 
                    smtp.quit()

    except:
        print('error')
