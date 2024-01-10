import streamlit as st
import pandas as pd
import base64
import csv
import math
import docx
import os
from langchain.output_parsers import OutputFixingParser
from langchain.llms import OpenAI
from langchain.chains import LLMChain, SimpleSequentialChain
from langchain.prompts import PromptTemplate, ChatPromptTemplate
from langchain.chat_models import ChatOpenAI
from langchain.output_parsers import ResponseSchema, StructuredOutputParser
from dotenv import load_dotenv
from langchain.callbacks import get_openai_callback
from langchain.llms import OpenAI


load_dotenv()


os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
# os.environ["OPENAI_API_KEY"] = os.getenv('OPENAI_API_KEY')

chat_llm = ChatOpenAI(temperature=0.0)

# # Determine the execution environment (development or production)
is_dev = os.getenv("IS_DEV", "false").lower() == "true"
data_path = "data" if is_dev else "/data"


doc = docx.Document()

def read_docx(file_path):
    """
    Read the content of a docx file and return it as a text string.
    """
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

file_path = r'data/Questions to ask summary.docx'
document_text = read_docx(file_path)
questions_to_ask_summary= document_text

def generate_questions(use_case,questions_to_ask_summary):
    title_template = """
            Ask about the following in questions form questions based on the "{usecase}" provided.
            Each Queston should have its in deatil description related to the "{usecase}"Given . 
            Use "{questions_to_ask_summary}" document as a knowledge base for Generation of Question and in detail description related to "{usecase}". 
            
            1) Nature of Use Case: 
            2) Number of User Interactions:
            3) Purpose of Use Case:             
            1) Intended User Group:
            2) Sensitivity of Use Case and Data:   
            1) Nature of LLMs Used:
            2) Embedding Approach: 
            3) Vector Stores:
            4) Prompting Approach:
            5) Fine-Tuning Approach:
            6) Type of Evaluations: 
            7) Guardrails:
            8) Monitoring Approach:
            9) Deployment Model:
            10) Humans in the Loop:
            11) Logging and Feedback Mechanism:
                """
    prompt = ChatPromptTemplate.from_template(template=title_template)
    messages = prompt.format_messages(usecase=use_case,questions_to_ask_summary=questions_to_ask_summary)
    response = chat_llm(messages)
    doc.add_paragraph("Questions:")
    doc.add_paragraph(response.content)
    questions=response.content
    st.write(questions)
    answers = st.text_input('Write your use case based on question asked')
    if st.button('Save'):
        doc.add_paragraph("Answers:")
        doc.add_paragraph(answers)
        
        st.write("Go to Next page For generation of Risk and Actionables")
    
    return doc.save('questions_and_answers.docx')

# Streamlit UI layout
st.title('LLM Risk Assessment Engine :S1')
st.subheader('Specializes in providing preliminary risk indicators for any use case')

# Text input for use case
use_case = st.text_input('Write your use case')

# Submit button for the use case
if st.button('Generate Questions'):
    st.write("Generating Questions")
    questions=generate_questions(use_case,questions_to_ask_summary)
    
    st.write(questions)
    
