import streamlit as st
import pandas as pd
import base64
from datetime import datetime
import csv
import math
import docx
import os
from langchain.output_parsers import OutputFixingParser
from langchain.llms import OpenAI
from langchain.chains import LLMChain, SimpleSequentialChain
from langchain.prompts import PromptTemplate, ChatPromptTemplate
from langchain.chat_models import ChatOpenAI
from langchain.output_parsers import ResponseSchema, StructuredOutputParser
from dotenv import load_dotenv
from langchain.callbacks import get_openai_callback
from langchain.llms import OpenAI


load_dotenv()


os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
# os.environ["OPENAI_API_KEY"] = os.getenv('OPENAI_API_KEY')

chat_llm = ChatOpenAI(temperature=0.0)


def read_docx(file_path):
    """
    Read the content of a docx file and return it as a text string.
    """
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


def save_string_to_docx(text, filename):
    doc = docx.Document()
    doc.add_paragraph(text)

    try:
        doc.save(filename)
        print(f'Saved to {filename}')
    except Exception as e:
        print(f'Error: {str(e)}')
        
def generate_risk_rank():
    doc = docx.Document()

    file_path = r'questions_and_answers.docx'
    questions_and_answers = read_docx(file_path)
    doc.add_paragraph("Answers:")
    doc.add_paragraph(questions_and_answers)

    file_path = r'data/Documented nature of risks.docx'
    Documented_nature_of_risks= read_docx(file_path)
        
    title_template = """
                Identify the risks that apply to the "{questions_and_answers}". Use the information in "{Documented_nature_of_risks}" document to identify the applicable risks. 
            Provide atleast 1 or 2 examples of each risk using the use case brief and the user responses to the questions in "{questions_and_answers}"
                """
    prompt = ChatPromptTemplate.from_template(template=title_template)
    messages = prompt.format_messages(questions_and_answers=questions_and_answers,Documented_nature_of_risks=Documented_nature_of_risks)
    response2 = chat_llm(messages)
    risk_information=response2.content
    doc.add_paragraph("Key Risks:")
    doc.add_paragraph(risk_information)
    title_template = """
                Rank the "{risk_information}" in terms of priority and provide a criticality score as high/ medium/ low given for "{questions_and_answers}".
                It should have Criticality Score and Reason for the above "{risk_information}".
                
                """
    prompt = ChatPromptTemplate.from_template(template=title_template)
    messages = prompt.format_messages(risk_information=risk_information,questions_and_answers=questions_and_answers)
    response3 = chat_llm(messages)
    risk_ranking=response3.content
    doc.add_paragraph("Risk Ranking:")
    doc.add_paragraph(risk_ranking)
    
    title_template = """Provide Actionable steps for governance to address each identified risk for "{risk_ranking}".
                For each risk compile a set of actionables to address the "{risk_ranking}". These actionables shall be governance actionables.
                """
    prompt = ChatPromptTemplate.from_template(template=title_template)
    
    messages = prompt.format_messages(risk_ranking=risk_ranking)
    response4 = chat_llm(messages)
    Actionables=response4.content
    doc.add_paragraph("Actionables:")
    doc.add_paragraph(Actionables)
   
    doc.add_paragraph(Actionables)
   
    summary = ""
    for paragraph in doc.paragraphs:
        # Here you define the logic to decide which paragraphs to include in the summary
        summary += paragraph.text + "\n"
    
    title_template = """Compile All information in "{summary}" . and Sturcture in below Format
                The document shall contain the following information: 
                Section A: Brief about the use case . 
                Section B: List of high-level risks associated with the use case.
                Section C: Table containing key risks with their risk ranking along with the reasons for the risk ranking.
                Section D: List of actionables for each risk listed in Section C.
                                """
    prompt = ChatPromptTemplate.from_template(template=title_template)
    
    messages = prompt.format_messages(summary=summary)
    response2 = chat_llm(messages)
    final_document=response2.content
    st.write(final_document)
    filename = "Final_document.docx"
    result=save_string_to_docx(final_document, filename)
    
    with open('Final_document.docx', 'rb') as f:
        doc_data = f.read()
    b64 = base64.b64encode(doc_data).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="Ai_Risk_Document.docx">Download Result</a>'
    st.markdown(href, unsafe_allow_html=True)
    
    

# Streamlit UI layout
st.title('LLM Risk Assessment Engine :S2')
st.subheader('Specializes in providing preliminary risk indicators for any use case')




# Submit button for the use case
if st.button('Generate Document'):
    st.write("Generating Key Risk and Actionables.....")
    generate_risk_rank()
    st.write("Download Final Document.....")
    # generate_document()
    
