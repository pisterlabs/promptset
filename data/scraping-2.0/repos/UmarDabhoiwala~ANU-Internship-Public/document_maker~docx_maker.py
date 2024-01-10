import sys 
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

import config
import openai
from docx import Document 
import pypandoc
from halo import Halo

openai.api_key = config.OPENAI_API_KEY

spinner = Halo(text='Loading', spinner='dots')

def generate_text(paragraph_gen_msg, title = False):
    
    if title:
        text, paragraph_gen_msg = chat_gpt_completion(paragraph_gen_msg, append = True, usr_prompt= "Generate a sub title relating to the topic. make it very short and conscise. A single sentence is best")
    else:
        text, paragraph_gen_msg = chat_gpt_completion(paragraph_gen_msg, append = True, usr_prompt= "Generate another paragraph in the same topic please. Keep it short and conscise. Attempt to maintain flow with the rest of the text")
    
    
    print(text)
    
    return text, paragraph_gen_msg


def replace_paragraphs(docx_file, paragraph_gen_msg):
    # Load the document
    doc = Document(docx_file)
    
    #Replace the text in each paragraph
    for paragraph in doc.paragraphs:
        # Clear the current paragraph's text
        counter = 0
        for run in paragraph.runs:
            counter += len(run.text)
            run.text = ''
            
        

        #Add the generated text
        if counter > 100:
            print("para")
            text, paragraph_gen_msg = generate_text(paragraph_gen_msg)
            paragraph.add_run(text)
        elif counter > 60:
            print("title")
            text, paragraph_gen_msg = generate_text(paragraph_gen_msg)
            paragraph.add_run(text)
            
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("List Bullet"):
            # Clear the current paragraph's text
            for run in paragraph.runs:
                run.text = ""

    # Save the modified document
    output_file = os.path.splitext(docx_file)[0] + "_modified.docx"
    doc.save(output_file)
    print(f"Modified file saved as: {output_file}")
    
def get_full_text (docx_file):
    
    try: 
        text = pypandoc.convert_file(docx_file, 'plain', extra_args=['--wrap=none'])
    except OSError:
        text = ""
        print("Pandoc not found. Please install pandoc to use this function.")
        sys.exit(1)
    except Exception as e:
        text = ""
        print(e)
        
    return text 


def chat_gpt_completion(chat_message, append = False, usr_prompt = ""):
    
    spinner.start("text generating")
        
    completion = openai.ChatCompletion.create(
        model="gpt-3.5-turbo", 
        messages= chat_message
    )
    
    spinner.succeed("text generated")
    
    response = completion['choices'][0]["message"]["content"]
    
    if append:
        chat_message.append({"role": "system", "content": response})
        chat_message.append({"role": "user", "content": usr_prompt})
        
    
    return response, chat_message


def text_summary_prompt(text): 
    
    system_prompt = """You are a text summariser, The user will give you a large amount of text 
    and you are to summarise the text into a single very small paragraph."""
    
    initial_prompt = f"""The user has given you the following text to summarise: {text}"""
    
    
    chat_MSG = [{"role": "system", "content": system_prompt},
            {"role": "user", "content": initial_prompt}]     
    
    return chat_MSG
    

def paragraph_gen_prompt(text):
    
    system_prompt = """You are a paragraph generator, The user will give you a topic
    and you are to generate a single very small paragraph from relating to the topic.  
    Each paragraph should be a maximum of 150 tokens. 
    Only respond with the answer, do not include any prompts or instructions."""
    
    initial_prompt = f"""This is the topic: {text}"""
    
    
    chat_MSG = [{"role": "system", "content": system_prompt},
            {"role": "user", "content": initial_prompt}]     
    
    return chat_MSG
    
        
input_file = "sample_doc_2.docx"

text = get_full_text(input_file)

text_summary_msg = text_summary_prompt(text)   

text_summary, text_summary_msg = chat_gpt_completion(text_summary_msg)

print(text_summary)
    
paragraph_gen_msg = paragraph_gen_prompt(text_summary)

replace_paragraphs(input_file, paragraph_gen_msg)

