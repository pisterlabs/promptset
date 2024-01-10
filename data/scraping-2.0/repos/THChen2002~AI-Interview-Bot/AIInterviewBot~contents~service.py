import openai
from django.conf import settings
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.oxml.ns import qn
import os
from django.utils import timezone

class ContentsService:
    def __init__():
        openai.api_key = settings.OPENAI_API_KEY
    
    # 取得OPEN AI API對話訊息
    def get_messages(user_prompt):
        messages=[
            {
                "role": "system",
                "content": "#zh-tw You are an excellent interviewee and want to apply for work or school."
            },
            user_prompt
        ]
        return messages
    
    # 取得OPEN AI API回覆訊息
    def get_reply(messages, model="gpt-3.5-turbo", type="text"):
        try:
            response = openai.ChatCompletion.create(
                model=model,
                messages=messages,
                response_format={"type": type},
                # max_tokens=1024, 
            )
            reply = response["choices"][0]["message"]["content"]
        except openai.OpenAIError as err:
            reply = f"發生 {err.error.type} 錯誤\n{err.error.message}"
        return reply
    
    # 取得OPEN AI API回覆訊息(Streaming)
    def get_reply_s(messages):
        try:
            response = openai.ChatCompletion.create(
                model = "gpt-3.5-turbo",
                messages = messages,
                stream = True
            )
            for chunk in response:
                if 'content' in chunk['choices'][0]['delta']:
                    yield chunk["choices"][0]["delta"]["content"]
        except openai.OpenAIError as err:
            reply = f"發生 {err.error.type} 錯誤\n{err.error.message}"

    # 匯出簡歷
    def export_resume(resume, output_path):
        # 設定段落格式
        def set_run_font(run, ch_font_name, en_font_name, font_size, bold=False, italic=False, underline=False, font_color=None):
            run.font.name = en_font_name
            run.element.rPr.rFonts.set(qn('w:eastAsia'), ch_font_name)
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.italic = italic
            run.font.underline = underline
            if font_color:
                run.font.color.rgb = font_color

        # 填入簡約履歷       
        def fill_simple_resume(resume, doc):
            # 使用已傳入的Document對象的第一個表格
            table = doc.tables[0]
            
            # 填入數據
            table.cell(1, 2).text = resume['name']
            table.cell(2, 2).text = resume['gender']
            table.cell(3, 2).text = resume['birth_date']
            table.cell(5, 2).text = resume['email']
            table.cell(7, 1).text = resume['self_introduction']
            table.cell(9, 1).text = resume['education']
            table.cell(9, 3).text = resume['personal_experience']
            table.cell(11,1).text = resume['skill']       
            table.cell(13, 3).text = resume['interest']

        # 填入專業履歷
        def fill_professional_resume(resume, doc):
            table = doc.tables[0]
            table.cell(1, 3).text = resume['name']
            table.cell(3, 2).text = resume['birth_date']
            table.cell(5, 2).text = resume['gender']
            table.cell(5, 5).text = resume['email']
            table.cell(7, 0).text = resume['education']
            table.cell(7, 4).text = resume['self_introduction']
            table.cell(9, 0).text = resume['skill']
            table.cell(10, 0).text = resume['interest']
            table.cell(11, 4).text = resume['personal_experience']

        # 填入創意履歷
        def fill_creative_resume(resume, doc):
            table = doc.tables[0]
            table.cell(2, 5).text = resume['name']
            table.cell(3, 5).text = resume['self_introduction']
            table.cell(4, 1).text = resume['education']
            table.cell(6, 1).text = resume['personal_experience']
            table.cell(9, 1).text = resume['skill']
            table.cell(14, 1).text = resume['interest']
            table.cell(10, 6).text = resume['gender']
            table.cell(12, 6).text = resume['birth_date']
            table.cell(16, 6).text = resume['email']
                
        # 根據風格選擇模板
        template_file = {
            '1': 'SimpleResume.docx',
            '2': 'ProfessionalResume.docx',
            '3': 'CreativeResume.docx'
        }.get(resume['style'], 'DefaultResume.docx')

        template_path = os.path.join(settings.STATICFILES_DIRS[0], 'resume_templates', template_file)

        # 載入模板文檔
        doc = Document(template_path)

        # 根據風格調用對應的填充函式
        if resume['style'] == '1':
            fill_simple_resume(resume, doc)
        elif resume['style'] == '2':
            fill_professional_resume(resume, doc)
        elif resume['style'] == '3':
            fill_creative_resume(resume, doc)

        # 遍歷表格設定樣式
        # for row in table.rows:
        #     for cell in row.cells:
        #         for paragraph in cell.paragraphs:
        #             for run in paragraph.runs:
        #                 set_run_font(run, '標楷體', 'Times New Roman', 12)
        #                 # Set vertical alignment for cell text
        #                 cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # 保存文件
        doc.save(output_path)
    
    # 取得日期區間開始日期
    def get_start_date(date_filter):
        if date_filter == 'today':
            start_date = timezone.now().date()
        elif date_filter == 'week':
            start_date = timezone.now().date() - timezone.timedelta(days=6)
        elif date_filter == 'month':
            start_date = timezone.now().date() - timezone.timedelta(weeks=9)
        elif date_filter == 'year':
            start_date = timezone.now().date() - timezone.timedelta(days=365)
        return start_date
    
    # 取得折線圖資料
    def get_line_chart_data(avg_scores, labels, date_filter):
        return {
            'labels': labels,
            'datasets': [
                {
                    'label': '專業能力',
                    'data': [record.professional_score if date_filter == 'today' else record['professional_score'] for record in avg_scores],
                    'borderWidth': 1,
                }, 
                {
                    'label': '創意能力',
                    'data': [record.creative_score if date_filter == 'today' else record['creative_score'] for record in avg_scores],
                    'borderWidth': 1,
                },
                {
                    'label': '策略能力',
                    'data': [record.strategy_score if date_filter == 'today' else record['strategy_score'] for record in avg_scores],
                    'borderWidth': 1,
                },
                {
                    'label': '溝通能力',
                    'data': [record.communication_score if date_filter == 'today' else record['communication_score'] for record in avg_scores],
                    'borderWidth': 1,
                },
                {
                    'label': '自主學習能力',
                    'data': [record.self_learning_score if date_filter == 'today' else record['self_learning_score'] for record in avg_scores],
                    'borderWidth': 1,
                },
                {
                    'label': '綜合能力',
                    'data': [record.comprehensive_score if date_filter == 'today' else record['comprehensive_score'] for record in avg_scores],
                    'borderWidth': 1,
                }
            ]
        }
    
    # 取得雷達圖資料
    def get_radar_chart_data(dashboard):
        return {
            'labels': ['專業能力', '創意能力', '策略能力', '溝通能力', '自主學習能力'],
            'datasets': [{
                'label': '專業能力',
                'data': [dashboard.professional_score, dashboard.creative_score, dashboard.strategy_score, dashboard.communication_score, dashboard.self_learning_score],
            }]
        }