from __future__ import annotations

from pathlib import Path

import httpx
import stamina
from bs4 import BeautifulSoup
from django.conf import settings
from django.db import models
from django.db.models.query import QuerySet
from django.utils.translation import gettext_lazy as _
from docx import Document as DocxDocument
from langchain.document_loaders import PyPDFLoader
from langchain.document_loaders import UnstructuredURLLoader
from langchain.embeddings import FakeEmbeddings
from langchain.embeddings import OpenAIEmbeddings
from langchain.text_splitter import Document as LangchainTextDoc
from langchain.text_splitter import RecursiveCharacterTextSplitter
from llama_hub.youtube_transcript import YoutubeTranscriptReader
from model_utils.models import TimeStampedModel
from pgvector.django import L2Distance
from pgvector.django import VectorField

from leerming.users.models import User

EMBEDDING_SIZE = 1536

TEXT_CHUNK_SIZE = 1000
TEXT_CHUNK_OVERLAP = 100

if settings.DEBUG:
    openai_embeddings = FakeEmbeddings(size=EMBEDDING_SIZE)
else:
    openai_embeddings = OpenAIEmbeddings()

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=TEXT_CHUNK_SIZE, chunk_overlap=TEXT_CHUNK_OVERLAP
)


@stamina.retry(on=httpx.HTTPError, attempts=3)
def get_title_from(*, url: str):
    response = httpx.get(url, follow_redirects=True)
    response.raise_for_status()
    soup = BeautifulSoup(response.text, "html.parser")
    # Find the first available title (h1, h2, h3)
    first_title = (
        soup.find("title")
        or soup.find("h1")
        or soup.find("h2")
        or soup.find("h3")
        or soup.find("h4")
    )
    return first_title.text.strip()


class UploadedDocument(TimeStampedModel):
    chunks: QuerySet[DocumentChunk]

    class DocType(models.TextChoices):
        PDF_DOC = "PDF_DOC", _("Document PDF")
        DOCX_DOC = "DOCX_DOC", _("Document Word")
        WEB_DOC = "HTML_DOC", _("Page web")
        YOUTUBE_VIDEO = "YOUTUBE_VIDEO", _("Vidéo Youtube")
        RAW_TEXT = "RAW_TEXT", _("Texte brut")

    owner = models.ForeignKey(
        "users.User", on_delete=models.CASCADE, related_name="uploaded_documents"
    )
    title = models.CharField(max_length=255)
    url = models.URLField(blank=True)
    doc_type = models.CharField(max_length=255, choices=DocType.choices)

    class Meta:
        constraints = [
            models.UniqueConstraint(
                fields=["owner", "title"], name="unique_uploaded_document"
            )
        ]

    def __str__(self):
        return self.title

    def get_relevant_chunks_for(self, query: str) -> QuerySet[DocumentChunk]:
        embedded_query = openai_embeddings.embed_query(query)
        return self.chunks.order_by(L2Distance("embedding", embedded_query))[:2]  # noqa

    def get_relevant_text_for(self, query: str):
        chunks = self.get_relevant_chunks_for(query)
        return "\n".join(chunk.content for chunk in chunks)

    def create_chunks(self, documents: list[LangchainTextDoc]) -> None:
        texts = text_splitter.split_documents(documents)
        vectors = openai_embeddings.embed_documents([t.page_content for t in texts])

        chunks = [
            DocumentChunk(
                document=self,
                content=texts[index].page_content,
                embedding=vector,
            )
            for index, vector in enumerate(vectors)
        ]
        DocumentChunk.objects.bulk_create(chunks)

    @classmethod
    def create_from_pdf_file(
        cls, *, temp_file: Path, title: str, owner: User
    ) -> UploadedDocument:
        loader = PyPDFLoader(str(temp_file))
        documents = loader.load_and_split()
        temp_file.unlink(missing_ok=True)
        uploaded_document = cls.objects.create(
            title=title, url="", doc_type=cls.DocType.PDF_DOC, owner=owner
        )
        uploaded_document.create_chunks(documents=documents)
        return uploaded_document

    @classmethod
    def create_from_docx_file(
        cls, *, temp_file: Path, title: str, owner: User
    ) -> UploadedDocument:
        docx = DocxDocument(str(temp_file))
        documents = [LangchainTextDoc(page_content=p.text) for p in docx.paragraphs]
        temp_file.unlink(missing_ok=True)
        uploaded_document = cls.objects.create(
            title=title, url="", doc_type=cls.DocType.DOCX_DOC, owner=owner
        )
        uploaded_document.create_chunks(documents=documents)
        return uploaded_document

    @classmethod
    def create_from_web_page(cls, *, url: str, owner: User, title) -> UploadedDocument:
        loader = UnstructuredURLLoader(
            urls=[url], continue_on_failure=False, headers={"User-Agent": "value"}
        )
        documents = loader.load()
        uploaded_document = cls.objects.create(
            title=title, url=url, doc_type=cls.DocType.WEB_DOC, owner=owner
        )
        uploaded_document.create_chunks(documents=documents)
        return uploaded_document

    @classmethod
    def create_from_raw_text(
        cls, *, text: str, owner: User, title: str
    ) -> UploadedDocument:
        uploaded_document = cls.objects.create(
            title=title, url="", doc_type=cls.DocType.RAW_TEXT, owner=owner
        )
        uploaded_document.create_chunks(documents=[LangchainTextDoc(page_content=text)])
        return uploaded_document

    @classmethod
    def create_from_youtube(cls, *, url: str, owner: User, title) -> UploadedDocument:
        loader = YoutubeTranscriptReader()
        documents = loader.load_data(ytlinks=[url])
        documents = [d.to_langchain_format() for d in documents]
        uploaded_document = cls.objects.create(
            title=title, url=url, doc_type=cls.DocType.YOUTUBE_VIDEO, owner=owner
        )
        uploaded_document.create_chunks(documents=documents)
        return uploaded_document

    @classmethod
    def get_create_func(cls, doc_type: DocType) -> callable:
        return {
            cls.DocType.YOUTUBE_VIDEO: cls.create_from_youtube,
            cls.DocType.PDF_DOC: cls.create_from_pdf_file,
            cls.DocType.DOCX_DOC: cls.create_from_docx_file,
            cls.DocType.WEB_DOC: cls.create_from_web_page,
            cls.DocType.RAW_TEXT: cls.create_from_raw_text,
        }.get(doc_type)


class DocumentChunk(TimeStampedModel):
    document = models.ForeignKey(
        UploadedDocument, on_delete=models.CASCADE, related_name="chunks"
    )
    content = models.TextField()
    embedding = VectorField(dimensions=EMBEDDING_SIZE)
