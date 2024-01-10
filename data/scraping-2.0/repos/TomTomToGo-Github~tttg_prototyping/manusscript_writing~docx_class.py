"""
Ideas for this class:
    - Proposal writingg
    - Paper writingg
    - School preparations
    - Company internal documents
"""

## built-in modules
import re
import os
from pathlib import Path
from uuid import uuid5, NAMESPACE_DNS
## pip installed modules
from docx import Document
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
import xml.etree.ElementTree as ET
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from gpt4all import GPT4All, pyllmodel
from openai import OpenAI

## pip install tensorflow -> windows c++ build tools SDK windows 10 - were necessary: winget install Microsoft.VisualStudio.2022.BuildTools --force --override "--wait --passive --add Microsoft.VisualStudio.Component.VC.Tools.x86.x64 --add Microsoft.VisualStudio.Component.Windows10SDK"
## pip3 install torch torchvision torchaudio -> also needs windows c++ build tools SDK windows 10 (or 11)
import tensorflow as tf; print(tf.reduce_sum(tf.random.normal([1000, 1000])))
a = tf.constant([[1.0, 2.0], [3.0, 4.0]])
from torch import Tensor as Tensor_pt
b = Tensor_pt([[1.0, 2.0], [3.0, 4.0]])


load_dotenv()

class DocxJson():
    def __init__(self):
        self.document = Document()
        self.scanned_doc = None
        self.scanned_doc_raw = None
        self.element_object_mapping = None
        self.img_count_write = 0
    
    ## Methods related to reading doc
    def docx_to_json(self, file_path):
        doc = Document(file_path)
        self.scanned_doc_raw = doc
        element_object_mapping = self.link_elements_to_objects(doc)
        self.element_object_mapping = element_object_mapping
        full_doc = []
        element_count = 0
        for element in doc.element.body:
            element_count += 1
            if element.tag.endswith('p'):
                # print(f'docx_to_json: paragraph scanned')
                paragraph = element_object_mapping['paragraphs'][element]
                new_element = self.paragraph_to_json(para=paragraph, count=element_count)
            elif element.tag.endswith('tbl'):
                print(f'docx_to_json: Table scanned')
                table = element_object_mapping['tables'][element]
                new_element = self.table_to_json(table, element_count)
                # print(new_element)
            elif 'graphicData' in element.xml:  # ?? old ?? element.tag.endswith('drawing'):
                print(f'docx_to_json: Image scanned')
                image_path, formatting = element_object_mapping['images'][element]
                new_element = self.image_to_json(image_path, element_count, formatting=formatting, run=None)
            elif element.tag.endswith('sectPr'):
                print("docx_to_json: section properties scanned")
                new_element = None
                break
            else:
                print("docx_to_json: Something else scanned")
                element_object_mapping['others'][element] = "??"
                new_element = None
                break
            if new_element:
                full_doc.append(new_element)

        if element_object_mapping['images_floating']:
            full_doc.append({"type": "header", "content": "Floating Images", "level": 1})
            for _, (img_float_path, formatting) in element_object_mapping['images_floating'].items():
                element_count += 1
                new_element = self.image_to_json(img_float_path, element_count, formatting=formatting, run=None)
                if new_element:
                    full_doc.append(new_element)
            
        # print("Full doc:", full_doc)
        self.scanned_doc = full_doc
        return full_doc

    def link_elements_to_objects(self, doc):
        element_object_mapping = {'paragraphs': {}, 'tables': {}, 'images': {}, 'images_floating': {}, 'others': {}}
        paragraph_iter = iter(doc.paragraphs)
        table_iter = iter(doc.tables)
        image_iter = self.iter_and_save_images_from_docx(doc)
        for element in doc.element.body:
            if element.tag.endswith('p'):
                try:
                    paragraph = next(paragraph_iter)
                    element_object_mapping['paragraphs'][element] = paragraph
                    for run in paragraph.runs:
                        xml_str = run._r.xml
                        if 'graphicData' in xml_str:
                            print("Link elements: Found image in paragraph run")
                            image_local_path, ctr_img_read = next(image_iter)
                            # print(xml_str)
                            formatting = self.get_image_formatting_from_xml(xml_str)
                            element_object_mapping['images'][ctr_img_read] = (image_local_path, formatting)
                except StopIteration:
                    break
            elif element.tag.endswith('tbl'):
                try:
                    table = next(table_iter)
                    element_object_mapping['tables'][element] = table
                except StopIteration:
                    break
            # elif 'graphicData' in element.xml:  # images are always part of run or floating
            #     print("Link elements: Found image as stand - allone")
            #     try:
            #         image_local_path, ctr_img_read, image_formatting = next(image_iter)
            #         element_object_mapping['images'][ctr_img_read] = (image_local_path, image_formatting)
            #     except StopIteration:
            #         break
            elif element.tag.endswith('sectPr'):
                print("Link elements: section properties found")
                break
            else:
                print("Link elements: Something else found")
                element_object_mapping['others'][element] = "??"
                break

        for image_local_path, ctr_img_read in image_iter:
            print("Link elements: Found more images - probably floating")
            formatting = None
            element_object_mapping['images_floating'][ctr_img_read] = (image_local_path, formatting)

        return element_object_mapping

    def iter_and_save_images_from_docx(self, doc):
        doc_rels = doc.part.rels
        ctr_img_read = 0
        for rel in doc_rels.values():
            if "image" in rel.reltype:
                ctr_img_read += 1
                image_data = rel._target._blob
                image_path =  self.get_temp_img_name(ctr_img_read)
                with open(image_path, 'wb') as img_file:
                    img_file.write(image_data)
                yield image_path, ctr_img_read  # , image_formatting

    @staticmethod
    def get_image_formatting_from_xml(xml_str):
        xml_tree = ET.fromstring(xml_str)
        formatting = None
        try:
            ## Extract size information
            ns = {
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
            }
            for elem in xml_tree.iter(f"{{{ns['a']}}}ext"):
                width = int(elem.attrib.get('cx')) / 914400
                height = int(elem.attrib.get('cy')) / 914400
            formatting = {'width': width, 'height': height}
        except:
            print("Image size could not be extracted")
        return formatting

    @staticmethod
    def get_temp_img_name(num):
        return f'__temp__saved_image_{num}.png'

    def paragraph_to_json(self, para, count):
        new_paragraph = None
        par_style = para.style.name  # 'Normal' equals 'Default Paragraph Style'
        if par_style == 'Title':
            new_paragraph = {
                "type": "header",
                "level": 0,
                "content": para.text,
                # "body": []
            }
        elif par_style.startswith('Heading'):
            new_paragraph = {
                "type": "header",
                "level": int(par_style[-1]),
                "content": para.text,
                # "body": []
            }
        elif par_style in ['Normal', 'Intense Quote', 'List Bullet', 'List Number']:
            if len(para.runs) == 1 and par_style != 'Normal':
                if 'graphicData' in para.runs[0]._r.xml:
                    print(f'Par2json: Image found in single paragraph')
                    new_run = {
                        "type": "image",
                        "style": 'bold',
                        "content": "There was an image here 1"
                    }
                content = para.text
            elif len(para.runs) >= 1 and par_style == 'Normal':
                content = []
                for run in para.runs:
                    if 'graphicData' in run._r.xml:
                        print(f'Par2json: Image found in run')
                        new_run = self.image_to_json(None, count, formatting=None, run=True)
                    else:
                        new_run = {
                            "type": "paragraph",
                            "style": 'bold' if run.bold else 'italic' if run.italic else 'Normal',
                            "content": run.text
                        }
                    content.append(new_run)
            else:
                print("Unhandled paragraph case!!!")
                content = 'UNHANDLED CASE'
            new_paragraph = {
                "type": "paragraph",
                "style": par_style,
                "content": content,
            }
            new_paragraph['alignment'] = para.alignment
            new_paragraph['element_id'] = uuid5(NAMESPACE_DNS, str(new_paragraph['content']) + str(count)).hex
        return new_paragraph

    def table_to_json(self, table, count):
        content = []  # rowwise
        formatting = []  # rowwise
        for row in table.rows:
            row_content = []
            row_formats = []
            for cell in row.cells:
                row_content.append(cell.text)
                row_formats.append({
                    "width": cell.width,
                    "vertical_alignment": cell.vertical_alignment,
                    # "text_direction": cell.text_direction,
                })
            content.append(row_content)
            formatting.append(row_formats)
        new_table = {
            "type": "table",
            "style": table.style,
            "content": content,
            "formatting": formatting
        }
        return new_table

    def image_to_json(self, image_path, img_count, formatting=None, run=True):
        self.img_count_write += 1
        if not formatting:
            formatting = {}
        if not run:
            new_image = {
                "type": "image",
                "id": uuid5(NAMESPACE_DNS, str(image_path) + str(img_count)).hex,
                "path":  image_path,
                "width": formatting.get('width', 4.25),
                "height": formatting.get('height', 1.25)
            }
            new_image.update(formatting)
        else: 
            image_path, formatting = self.element_object_mapping['images'][self.img_count_write]
            new_image = {
                "type": "image",
                "id": uuid5(NAMESPACE_DNS, str(image_path) + str(img_count)).hex,
                "path":  image_path,
                "width": formatting.get('width', 4.25),
                "height": formatting.get('height', 1.25)
            }
            new_image.update(formatting)
        return new_image

    ## Methods related to writing doc
    def create_docx_from_json(self, data):
        if isinstance(data, list):  # recursion for nested elements (e.g. body of header is list)
            for item in data:
                # an item is a dictionary of with
                #   mandatory key 'type' allowing header, text, table, image
                #   optional keys, 'level', 'style', 'width', body depending on type
                self.create_docx_from_json(item)
        if isinstance(data, dict):
            try:
                if data['type'] == 'header':
                    print("Writing header")
                    header = data.get('content', 'Header x')
                    level = data.get('level', 1)
                    self.document.add_heading(header, level=level)
                    if 'body' in data:
                        self.create_docx_from_json(data['body'])
                elif 'paragraph' in data['type']:  # paragraph and paragraph_run
                    self.write_paragraph(data)
                elif data['type'] == 'table':
                    self.write_table(data)
                elif data['type'] == 'image':
                    self.add_image(data, in_run=False)
            except FileNotFoundError as e:
                print(e)

    def write_paragraph(self, data):
        content = data.get('content', 'RANDOM PARAGRAPH TEXT')
        par_style= data.get('style', 'Normal')
        alignment = data.get('alignment', 'RIGHT')  # Get the alignment from the data

        # Convert the alignment to a value from the WD_ALIGN_PARAGRAPH enumeration
        if alignment == 'LEFT':
            alignment_value = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == 'CENTER':
            alignment_value = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'RIGHT':
            alignment_value = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            alignment_value = WD_ALIGN_PARAGRAPH.LEFT  # Default to left alignment

        if isinstance(content, str):
            paragraph = self.document.add_paragraph(content, style=par_style.title())
            paragraph.alignment = alignment_value  # Set the alignment of the paragraph
        elif isinstance(content, list):
            self.current_paragraph = self.document.add_paragraph("")
            self.current_paragraph.alignment = alignment_value  # Set the alignment of the paragraph
            for item in content:
                if 'paragraph' in item['type']:  # paragraph and paragraph_runs
                    print("Writing paragraph")
                    self.add_paragraph_element(item['content'], font_style=item.get('style', 'Normal'))
                elif item['type'] == 'image':
                    self.add_image(item, in_run=True)

    def add_paragraph_element(self, text, font_style='Normal'):
        if font_style.lower() == 'normal':
            print("     -> Adding normal run")
            self.current_paragraph.add_run(text)
        elif font_style.lower() == 'bold':
            print("     -> Adding bold run")
            self.current_paragraph.add_run(text).bold = True
        elif font_style.lower() == 'italic':
            print("     -> Adding italic run")
            self.current_paragraph.add_run(text).italic = True

    def add_image(self, new_image, in_run=False):
        width = Inches(new_image.get('width', 4))
        height = Inches(new_image.get('height', 1.25))
        if in_run:
            print("     -> Adding image run")
            current_run = self.current_paragraph.add_run()
        else:
            print("Writing new paragraph\n  ->     Adding image run")
            current_run = self.document.add_paragraph().add_run()
        current_run.add_picture(new_image['path'], width=width, height=height)

    def write_table(self, data):
        content = data.get('content', 'DEFAULT TABLE TEXT')
        formatting = data.get('formatting', None)
        table = self.document.add_table(rows=len(content), cols=len(content[0]))
        table.style = data.get('style', 'Table Grid')
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell.text = content[row_idx][col_idx]
                # Set cell width
                if formatting:
                    cell.width = formatting[row_idx][col_idx].get('width', Pt(50))
                    cell.vertical_alignment = formatting[row_idx][col_idx].get('vertical_alignment', Pt(50))        

    def write_document(self, file_path):
        self.document.save(file_path)
        


class LlmBaseClass:
    def __init__(self, llm_source, model_name, model_path=None, project_name=None, device='cpu'):
        self._prompt_templates = {}
        self.models_loaded = {'openai': {}, 'gpt4all': {}}
        self.model_names = {'openai': [], 'gpt4all': []}
        if llm_source == 'openai':
            if project_name:
                ## OPEN AI CLIENT
                openai_key_project = f"OPENAI_API_KEY_{project_name}"
                if os.getenv(openai_key_project, None):
                    print("Loaded openAI key: ", openai_key_project)
                else:
                    print("Loaded openAI key: OPENAI_KEY_GENERAL")
                model = OpenAI(
                    api_key=os.getenv(openai_key_project, os.getenv("OPENAI_API_KEY_GENERAL"))
                )
            else:
                print("No project name given. -> Loading default OpenAI key.")
                if os.getenv("OPENAI_API_KEY"):
                    print("Loaded openAI key: OPENAI_API_KEY")
                    model = OpenAI(
                        api_key=os.getenv(openai_key_project, os.getenv("OPENAI_API_KEY"))
                    )
                else:
                    print("Unable to load OPENAI_API_KEY. Check .env No project name given. -> No OpenAI client loaded.")
        elif llm_source == 'gpt4all':
            model = GPT4All(model_name, model_path=model_path) # device='amd', device='intel'
            print(f"Loaded model: {model_name} using gpt4all")

        self.models_loaded[llm_source][model_name] = model
        self.model_names[llm_source].append(model_name)
        self.last_output = None

    def prompt_llm(self, llm_source='gpt4all', model_name=None, messages=None, **kwargs):
        if not model_name and self.model_names[llm_source]:
            model_name = self.model_names[llm_source][-1]
            print("Prompting with default model: ", model_name)
        ## Format parameters according to model
        prompt, prompt_params = self.get_prompt_params(llm_source, messages, **kwargs)
        print("Prompt: ", prompt)
        # print(prompt_params, prompt)
        ## Prompt the model
        print(f"Prompting the LLM: {model_name} from source {llm_source}")
        if llm_source == 'openai':
            self.last_output = self.models_loaded[llm_source][model_name].chat.completions.create(
                model=model_name,
                messages=prompt
                # **prompt_params
            )
            # print(self.last_output.choices[0].message.content)
            return self.last_output.choices[0].message.content
        elif llm_source == 'gpt4all':
            self.last_output =  self.models_loaded['gpt4all'][model_name].generate(prompt, **prompt_params)
            return self.last_output
    
    def get_prompt_params(self, llm_source, messages, **kwargs):
        prompt_params = {}
        if llm_source == 'openai':
            prompt = messages
            prompt_params = {
                "temperature": 0.7,
                "usage": {
                    "prompt_tokens": kwargs.get('prompt_tokens', 100),
                    "completion_tokens": kwargs.get('completion_tokens', 100),
                    "total_tokens": kwargs.get('total_tokens', kwargs.get('max_tokens', 200)),
                }
            }
        elif llm_source == 'gpt4all':
            prompt = " ".join(["{}: {}".format(msg.get('role', '').title(), msg.get('content', '')) for msg in messages])
            prompt_params = {
                'max_tokens': kwargs.get('max_tokens', 200),
                'temp': kwargs.get('temp', 0.3),
                'top_k': kwargs.get('top_k', 40),
                'top_p': kwargs.get('top_p', 0.4),
                'repeat_penalty': kwargs.get('repeat_penalty', 1.18),
                'repeat_last_n': kwargs.get('repeat_last_n', 64),
                'n_batch': kwargs.get('n_batch', 8),
                # 'n_predict': kwargs.get('', 200,  # same as max tokens -> just for backwards compatibility
                'streaming': kwargs.get('streaming', False),
                'callback': kwargs.get('callback', pyllmodel.empty_response_callback),
            }
        # print("Prompt params: ", prompt_params)
        return prompt, prompt_params

    def add_prompt_template(self, template_name, template):
        self._prompt_templates[template_name] = template

    def _prompt_template_to_text(self, template_name, inputs=None):
        template = self._prompt_templates[template_name]
        template = "Test {{{{INP}}}}, {{{{INP}}}}, {{{{INP}}}}"
        inputs = ["1", "2", "3"]
        if isinstance(inputs, list):
            nr_inputs = len(re.findall(r"{{{{\w*}}}}", template))
            if nr_inputs == len(inputs):
                for input in inputs:
                    template = template.replace("{{{{INP}}}}", input, 1)
            else:
                print("Number of inputs does not match number of input placeholders")
        elif isinstance(inputs, str):
            template = template.replace("{{{{INP}}}}", inputs)
        elif isinstance(inputs, dict):
            for key, value in inputs.items():
                template = template.replace(f"{{{{{{{{{key}}}}}}}}}", value)

    def _prompt_text_to_messages(self, prompt_text):
        # prompt_text = "# FILEPATH: /c:/Users/thoma/Desktop/Tommy/Programmieren/tttg_prototyping/manusscript_writing/docx_class.py\n# BEGIN: ed8c6549bwf9\n    def _prompt_text_to_messages(self, prompt_text):\n        \n        return messages\n# END: ed8c6549bwf9"
        split_input = re.split(r'#+', prompt_text)
        print(split_input)
        return messages

    def prompt_template_to_messages(self, template_name, inputs):
        prompt_text = self._prompt_template_to_text(self, template_name, inputs=None)
        return self.text_to_messages(prompt_text)

class DocxTextEdit():
    def __init__(self, file_path):
        self.document = None
        file_path = Path(file_path)
        if file_path.is_file() and file_path.name.endswith('.docx'):  # check if file exists
            self.load_doc(file_path)

    def load_doc(self, file_path):
        try:
            self.document = Document(file_path)
        except FileNotFoundError as e:
            print(e)

    def write_document(self, file_path):
        self.document.save(file_path)
        
    def extract_paragraph_text_to_edit(self, regexes=None):
        doc = self.document
        text_edit_dict = {}  # self.text_edit_dict
        for idx, para in enumerate(doc.paragraphs):
            text_before, text_to_replace, text_after = self.split_text(para.text, regexes=None)
            id = uuid5(NAMESPACE_DNS, str(para.text) + str(idx)).hex,
            text_edit_dict[id] = {
                'id': id,
                'original_location': idx,
                'current_location': idx,
                'text_before': text_before,
                'text_to_replace': text_to_replace,
                'text_after': text_after,
                'paragraph_object': para
            }
            text_edit_dict = {id: values for id, values in text_edit_dict.items() if values['text_to_replace']}
        return text_edit_dict

    @staticmethod
    def split_text(text, regexes=None):
        # text = "Some text that should stay $%$$%$$%$% Some text to replace §$%$%$%%§$%$ Some text that should stay"
        # text = "Some text that should stay $%$$%$$%$% Some text to replace No text that should stay"
        # text = "Some text that should stay No text to replace Some text that should stay"
        # text = "######WUHU IT WORKS######"
        # text = "WUHU IT  ###### WORKS ####"
        standard_delimiter = "[\W]{5,}"
        if not isinstance(regexes, dict):
            regexes = {
                "reg_start": standard_delimiter,
                "reg_end": standard_delimiter
            }
        reg_start = regexes.get('reg_start', standard_delimiter)
        reg_end = regexes.get('reg_end', standard_delimiter)
        text_len = len(text)
        text_before = text
        text_to_replace = ""
        text_after = ""
        delim_start = re.search(reg_start, text)
        if delim_start:
            span_start = delim_start.span()
            text_before = text[:span_start[0]]
            delim_end = re.search(reg_end, text[::-1])
            if delim_end and (span_start[1] < (text_len - delim_end.span()[1])):
                span_end = delim_end.span()
                text_to_replace = text[span_start[1]:-span_end[1]]
                text_after = text[text_len - span_end[0]:]
            else:
                text_to_replace = text[span_start[1]:]
        # print(text_before)
        # print(text_to_replace)
        # print(text_after)
        return text_before, text_to_replace, text_after

    def get_par_by_id(self, id):
        for para in self.document.paragraphs:
            if id == self.text_edit_dict['id']:
                return para        


class LlmEditDocx(LlmBaseClass, DocxTextEdit):
    def __init__(self, llm_source, **kwargs):
        LlmBaseClass.__init__(self, llm_source, **kwargs)
        self.system_message = "You are a writing assistant, skilled in writing academic proposals using professional language and tone."
        self.topic = "How to integrate AI into your business"
        self.prompt_template = f"Write at least one paragraph on the topic of '{self.topic}'. Include the following information: '{text_to_replace}'"
        # super().__init__(*args, **kwargs)

    def replace_marked_text_in_docx(self, regexes=None, **kwargs):
        text_edit_dict = self.extract_paragraph_text_to_edit(regexes=regexes)
        total_replacements = len(list(text_edit_dict.keys()))
        for idx, values_dict in enumerate(text_edit_dict.values()):
            print(f"Text replacement {idx + 1} / {total_replacements} in progress")
            para = values_dict.get('paragraph_object', None)
            if isinstance(para, Paragraph):
                text_before = values_dict.get('text_before', '')
                text_to_replace = values_dict.get('text_to_replace', '')
                text_after = values_dict.get('text_after', '')
                prompt_params = {
                    'max_tokens': 800,
                    'temp': 0.3,
                    'top_k': 40,
                    'top_p': 0.4,
                    'repeat_penalty': 1.18,
                    'repeat_last_n': 64,
                    'n_batch': 8,
                }
                messages = [
                    {"role": "system", "content": self.system_message},
                    {"role": "user", "content": self.prompt_template}
                ]

                ## Prompt model and get output
                # replacement_text = self.prompt_llm(llm_source='gpt4all', model_name=None, messages=messages, **prompt_params)
                replacement_text = self.prompt_llm(messages=messages, **kwargs)  ## kwargs to specify model options
                print("\n\nPrevious text in document", para.text)
                print("\n\nMessages: ", messages)
                print("Replacement_text: ", replacement_text)
                para.text = text_before + replacement_text + text_after
                print("\nEndresult: ", para.text)

# from langchain.chat_models import ChatOpenAI, ChatPromptTemplate
# from langchain.prompts import HumanMessagePromptTemplate
# from langchain.schema.messages import SystemMessage

# chat_template = ChatPromptTemplate.from_messages(
#     [
#         SystemMessage(
#             content=(
#                 "You are a helpful assistant that re-writes the user's text to "
#                 "sound more upbeat."
#             )
#         ),
#         HumanMessagePromptTemplate.from_template("Some form of text: {text}"),
#     ]
# )


if __name__ == '__main__':
    ## Write from predetermined json
    # -> make paragraf be one item with subitems!
    json_data = {
        "type": "header",
        "level": 0,
        "content": "Document Title",
        "body": [
            {
                "type": "paragraph",
                "style": "Normal",
                "content": "A plain paragraph having some normal text."
            },
            {
                "type": "paragraph",
                "style": "Normal",
                "content": [
                    {
                        "type": "paragraph",
                        "style": "Normal",
                        "content": "A plain paragraph having some normal text plus some "
                    },
                    {
                        "type": "paragraph",
                        "style": "bold",
                        "content": "bold"
                    },
                    {
                        "type": "paragraph",
                        "style": "Normal",
                        "content": " and some "
                    },
                    {
                        "type": "paragraph",
                        "style": "italic",
                        "content": "italic."
                    },
                    {
                        "type": "paragraph",
                        "style": "Normal",
                        "content": "."
                    }
                ]
            },
            {
                "type": "header",
                "content": "Heading, level 1",
                "level": 1
            },
            {
                "type": "paragraph",
                "content": "Intense quote",
                "style": "Intense Quote"
            },
            {
                "type": "paragraph",
                "content": "first item in unordered list",
                "style": "List Bullet"
            },
            {
                "type": "paragraph",
                "content": "first item in ordered list",
                "style": "List Number"
            },
            {
                "type": "header",
                "content": "Heading, level 1",
                "level": 1
            },
            {
                "type": "header",
                "content": "Heading, level 2",
                "level": 2
            },
            {
                "type": "header",
                "content": "Heading, level 3",
                "level": 3
            },
            {
                "type": "paragraph",
                "content": "Some random text to accompany the image",
                "style": "Normal"
            },
            {
                "type": "image",
                "path": "test.png",
                "width": 1.25
            },
            {
                "type": "header",
                "content": "Heading, level 1",
                "level": 1
            },
            {
                "type": "header",
                "content": "Heading, level 2",
                "level": 2
            },
            {
                "type": "paragraph",
                "content": "There should be a table underneath here:",
                "style": "Normal"
            },
            {
                "type": "table",
                "content": [
                    ["Qty", "Id", "Desc"],
                    ["3", "101", "Spam"],
                    ["7", "422", "Eggs"],
                    ["4", "631", "Spam, spam, eggs, and spam"]
                ]
            }
        ]
    }


    # docx_instance = DocxJson()
    # docx_instance.create_docx_from_json(json_data)
    # docx_instance.write_document('test.docx')
    
    ## Scan document and rewrite it
    # docx_instance = DocxJson()
    # scanned_json = docx_instance.docx_to_json('test.docx')
    # docx_instance.create_docx_from_json(scanned_json)
    # docx_instance.write_document('test_2.docx')
    
    ## Gpt4All - lcoal
    # Init 1: Local without licence and project specification
    llm_source = 'gpt4all'
    # model_name = "orca-mini-3b-gguf2-q4_0.gguf"
    # model_name = "mistral-7b-instruct-v0.1.Q4_0.gguf"
    model_name = "orca-2-7b.Q4_0.gguf"
    model_path = Path(r"C:/Users/thoma/Desktop/Tommy/Programmieren/gpt4all/models")
    project_name = None
    # Init 2: OpenAI API with project token specifications
    llm_source = 'openai'
    model_path = None
    model_name = "gpt-3.5-turbo-1106"
    project_name = "PROPOSAL"
    
    ## Initalize the class with LLM capabilities and load word document 
    docx_llm = LlmEditDocx(llm_source, model_name=model_name, model_path=model_path, project_name=project_name)
    docx_llm.load_doc('test_3.docx')
    ## Replace marked text in document
    # docx_llm.replace_marked_text_in_docx(llm_source=None, model_name=None, regexes=None)  # Defaults: 'gpt4all' (no costs), last model loaded of source, regex to identify 5 special characters
    docx_llm.replace_marked_text_in_docx(llm_source='openai', model_name=None, regexes=None)  # Defaults: last model loaded of source, regex to identify 5 special characters
    ## Save document with text replacement
    docx_llm.write_document('test_4.docx')
    
    
    
    topic = "How to integrate AI into your business"
    text_to_replace = "How can I convince my coworkers?"
    messages = [
        {"role": "system", "content": "You are a writing assistant, skilled in writing academic proposals using professional language and tone."},
        # {"role": "user", "content": f"Write at least one paragraph on the topic of '{topic}'. Include the following information: '{text_to_replace}'"}
    ]
    docx_llm.prompt_llm(messages=messages)
    
    
    

    
    # ## Define messages and params
    # llm_source = 'gpt4all'
    # model_name = "orca-mini-3b-gguf2-q4_0.gguf"
    # model_name = "orca-mini-3b.ggmlv3.q4_0.bin"
    # prompt_params = {
    #     'max_tokens': 2400,
    #     'temp': 0.3,
    #     'top_k': 40,
    #     'top_p': 0.4,
    #     'repeat_penalty': 1.18,
    #     'repeat_last_n': 64,
    #     'n_batch': 8,
    # }
    # messages = [
    #     {"role": "system", "content": "You are a poetic assistant, skilled in explaining complex programming concepts with creative flair."},
    #     {"role": "user", "content": "Compose a poem that explains the concept of recursion in programming. Respond in only 6 lines of poetry."}
    # ]
    
    # ## Prompt model and get output
    # output_str = docx_llm.prompt_llm(llm_source=llm_source, model_name=model_name, messages=messages, **prompt_params)

    # print(output_str)
