import os
import re
import glob
import langchain.schema
import random
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.document_loaders import DirectoryLoader
from langchain.embeddings import HuggingFaceInstructEmbeddings
from langchain.document_loaders.csv_loader import CSVLoader
from langchain.vectorstores.pgvector import PGVector
from dotenv import load_dotenv
from pdfminer.high_level import extract_text
from Crypto.Cipher import AES
from docx import Document
from tqdm import tqdm

load_dotenv()

embeddings = HuggingFaceInstructEmbeddings(model_name='thenlper/gte-large')
#embeddings = HuggingFaceInstructEmbeddings(model_name='BAAI/bge-base-en')
#embeddings = HuggingFaceInstructEmbeddings(model_name='BAAI/bge-small-en')
#embeddings = HuggingFaceInstructEmbeddings(model_name='BAAI/bge-large-en')

text_splitter = RecursiveCharacterTextSplitter(chunk_size=3000, chunk_overlap=300)

directory = r'/home/metodi/dev/documentqna/staging/' 

def contains_isolated_chars(text, threshold=2):
    lines = text.split('\n')
    isolated_char_count = sum(1 for line in lines if len(line.strip()) == 1)
    return isolated_char_count >= threshold

EXCLUDED_SECTIONS = [
    "References", 
    "Index", 
    "Contents", 
    "Contributors",
    "Acknowledgment",
    "Funding",
    "Conflict of Interest",
    "Ethics Approval and Consent to Participate"
]

def remove_excluded_sections(text, sections=EXCLUDED_SECTIONS):
    for section in sections:
        pattern = re.compile(rf'{section}.*', re.DOTALL)
        text = re.sub(pattern, '', text)
    return text

def clean_text(text):
    
    text = re.sub(r'\(cid:\d+\)', '', text)
    text = text.replace('\x0c', '')
    text = re.sub(r'(\n\w\n)+', '\n', text)
    text = re.sub(r'\n+', '\n', text)
    text = text.strip('\n')
    
    common_keywords_patterns = [
        r"Copyright:", 
        r"Licensee",
        r"@",
        r"Downloaded",
        r"Download",
        r"Figure",
        r"open access article",
        r"correspondence",
        r"Creative Commons",
        r"Citation:",
        r"Academic Editor:",
        r"Received:",
        r"Accepted:",
        r"Published:",
        r"Publisher's Note:",
        r"https://",
        r"www\.",
        r"\d{4}, \d+, \d+ \d+ of \d+"
    ]
    
    for pattern in common_keywords_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    
    return text

def create_document(text, file_path):
    text = text.replace('\0', '')
    doc = langchain.schema.Document(
        page_content=text, 
        metadata={
            "source": f'{file_path}', 
            "row": random.randint(999999, 9999999)
        }
    )
    return doc

def parse_documents(path):
    """
    Parse and process documents from a specified path.

    This function reads files from a given path (or directory) and processes them based on their file type.
    Currently, it supports PDF, DOCX, and CSV files. The function extracts text content, applies necessary 
    preprocessing, chunks the content if required, and then prepends the filename to each chunk.

    Parameters:
    - path (str): The path to the file or directory containing the files to be processed.

    Returns:
    - list: A list of processed documents with filenames prepended.

    Notes:
    - For PDF files, the content is extracted, optionally cleaned, and then chunked if required.
    - For DOCX files, the content from paragraphs and tables is combined, optionally cleaned, and then chunked.
    - For CSV files, the content is loaded using a custom CSVLoader. Each row is treated as a document,
      which is then chunked if required.
    - Any errors encountered during processing are printed to the console with the respective filename.
    """
    if os.path.isdir(path):
        file_paths = glob.glob(os.path.join(path, '*.*'))
    else:
        file_paths = [path]

    documents = []
    num_docs = 0
    
    for file_path in tqdm(file_paths):
        try:
            file_name = os.path.basename(file_path)

            if file_path.endswith(".pdf"):
                combined_text = extract_text(file_path)
                if combined_text:
                    cleaned_content = clean_text(combined_text)
                    cleaned_content = combined_text
                    section_removed_content = remove_excluded_sections(cleaned_content)
                    chunked_content = text_splitter.split_text(section_removed_content)
                    
                    for chunk in chunked_content:
                        if contains_isolated_chars(chunk):
                            continue
                        chunk_with_filename = f"{file_name}\n{chunk}"
                        doc = create_document(chunk_with_filename, file_path)
                        documents.append(doc)
                        num_docs += 1

            elif file_path.endswith(".docx"):
                docx_file = Document(file_path)
                combined_text = ''.join(paragraph.text for paragraph in docx_file.paragraphs)
                combined_text += ''.join(cell.text for table in docx_file.tables for row in table.rows for cell in row.cells)
                cleaned_content = clean_text(combined_text)
                cleaned_content = combined_text
                chunked_content = text_splitter.split_text(cleaned_content)
                
                for chunk in chunked_content:
                    chunk_with_filename = f"{file_name}\n{chunk}"
                    doc = create_document(chunk_with_filename, file_path)
                    documents.append(doc)
                    num_docs += 1

            if file_path.endswith(".csv"):
                loader = CSVLoader(file_path=file_path)
                loaded_documents = loader.load()
                chunked_content = text_splitter.split_documents(loaded_documents)
                
                for chunk in chunked_content:
                    chunk_with_filename = f"{file_name}\n{chunk}"
                    doc = create_document(chunk_with_filename, file_path)
                    documents.append(doc)
                    num_docs += 1

        except Exception as e:
            print(f"Error processing file {file_path}: {e}")

    print(f'Parsed {num_docs} documents.')
    return documents

documents = parse_documents(directory)

vectordb = PGVector.from_documents(
    documents=documents,
    embedding=embeddings,
    collection_name='papers',
    pre_delete_collection=True,
    connection_string='postgresql+psycopg2://pgadmin:pgadmin@localhost:5432/embeddings')