import os

try:
  import openai
  import dotenv
  import youtube_dl
  from pytube import YouTube

  from docx import Document
except ImportError as e:
  print("Missing dependencies detected. Run pip install -r requirements.txt to install...")

MODEL = "gpt-3.5-turbo-16k-0613"

dotenv.load_dotenv()

openai.api_key = os.getenv("API_KEY")

def get_video_description(video_url):
    if video_url.startswith("C:") or video_url.startswith("\\") or video_url.startswith("/"):
        return None
    elif not video_url.startswith("http") and not video_url.startswith("www") and not video_url.startswith("youtube"):
        video_url = "https://www.youtube.com/watch?v=" + video_url

    # with youtube_dl.YoutubeDL() as ydl:
    #     info_dict = ydl.extract_info(video_url, download=False)
    #     return info_dict.get('description', None)

    try:
        yt = YouTube(video_url)
        return yt.initial_data["engagementPanels"][1]["engagementPanelSectionListRenderer"]["content"]["structuredDescriptionContentRenderer"]["items"][1]["expandableVideoDescriptionBodyRenderer"]["attributedDescriptionBodyText"]["content"]
    except:
        return None

def chatbot(questions_path, transcript_path, save_path, youtube_url=None):
    try:
       doc = Document(questions_path)
    except:
       print("Error")

    questions = ""
    if questions_path:
        if questions_path.endswith(".docx") or questions_path.endswith(".doc"):
            for paragraph in doc.paragraphs:
                questions += paragraph.text
        elif questions_path.endswith(".txt"):
            with open(questions_path, encoding="utf-8") as txt:
                questions = txt.read()
    elif youtube_url:
        questions = get_video_description(youtube_url)
    else:
        print("No questions provided.")
        exit(1)

    if questions is None or questions == "":
        print("Failed to get questions from description.")
        exit(1) 

    transcript = ""
    with open(transcript_path, encoding="utf-8") as txt:
       transcript = txt.read()

    messages = [
        {"role": "system", "content": "Jsi AI asistent, který odpovídá na otázky z dějepisu, češtiny, nebo jiných školních předmětů. Odpověz na tyto otázky stručně a jednoduše. Je velmi důležité, že každá odpověď je správná. Informace získáš z přiloženého výkladu."},
    ]

    messages.append({"role": "user", "content": transcript})

    print(messages)

    print("Sending prompt with transcript...")

    response = openai.chat.completions.create(
        model=MODEL,
        messages=messages
    )
    
    print("Sending prompt with questions...")

    messages.append({"role": "user", "content": questions})
    print(messages)
    response = openai.chat.completions.create(
        model=MODEL,
        messages=messages
    )

    try:
        response_text = response.choices[0].message.content

        if save_path.endswith(".docx") or save_path.endswith(".doc"):
            doc = Document()
            doc.add_paragraph(response_text)
            doc.save(save_path)
        elif save_path.endswith(".txt"):
            with open(save_path, "w", encoding="utf-8") as txt:
                txt.write(response_text)
    except:
        print("Error while saving answers. Will print to console instead.")

        print(response)