import requests
from bs4 import BeautifulSoup
from gtts import gTTS
import openai
import json
import time
from datetime import datetime, timedelta
from docx import Document
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive

# Load the configuration file; you will need to have your openai and newsapi keys in this file
with open('config.json') as config_file:
    config = json.load(config_file)

# API keys setup
openai.api_key = config['openai_key']
newsapi_key = config['newsapi_key']

# Google Drive setup
# you will need to have a file client_secrets.json that has your Google api credentials
gauth = GoogleAuth()
gauth.LocalWebserverAuth()
drive = GoogleDrive(gauth)

# Function to upload file to Google Drive
def upload_file(filename):
    print(f"Uploading {filename} to Google Drive...")
    start = time.time()
    gfile = drive.CreateFile({'title': filename})
    gfile.SetContentFile(filename)
    gfile.Upload()
    print(f"Done uploading. Time taken: {time.time() - start} seconds.")

# Function to scrape the news article text
def scrape_news(url):
    print(f"Scraping article at {url}...")
    start = time.time()
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')
    paragraphs = soup.find_all('p')
    news_text = ' '.join([para.text for para in paragraphs])
    print(f"Done scraping. Time taken: {time.time() - start} seconds.")
    return news_text

# Function to generate a summary of the text using OpenAI's GPT-4
def summarize_text(text):
    print("Summarizing text...")
    start = time.time()

    prompt = "As an employee researcher, summarize the following article focusing on how AI could be used by employees or in the workplace:"
    input_text = f"{prompt}\n{text[:2048]}"

    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=input_text,
        temperature=0.3,
        max_tokens=150
    )
    print(f"Done summarizing. Time taken: {time.time() - start} seconds.")
    return response.choices[0].text.strip()

# Function to convert text to speech using Google Text-to-Speech
def text_to_speech(text, filename):
    print(f"Converting text to speech and saving as {filename}...")
    start = time.time()
    tts = gTTS(text=text, lang='en')
    tts.save(filename)
    print(f"Done converting to speech. Time taken: {time.time() - start} seconds.")
    upload_file(filename)

# Main function to scrape news, generate summary, convert summary to audio and upload to Google Drive
def main():
    print("Fetching news articles...")
    start = time.time()

    # Date five days ago from now in 'YYYY-MM-DD' format
    from_date = (datetime.now() - timedelta(5)).strftime('%Y-%m-%d')
    today_date = datetime.now().strftime('%Y-%m-%d')  # Today's date for MP3 filename and intro text

    # Fetch news related to "generative AI" from the last 5 days, top 10 articles
    newsapi_url = f"https://newsapi.org/v2/everything?q=generative%20ai%20workplace&from={from_date}&pageSize=10&apiKey={newsapi_key}"
    news_json = requests.get(newsapi_url).json()
    print(f"Done fetching articles. Time taken: {time.time() - start} seconds.")

    combined_summary = f"News summary for generative AI in the workplace on {today_date}. "
    doc = Document()  # Word document

    for article in news_json['articles']:
        text = scrape_news(article['url'])
        summary = summarize_text(text)
        print(f"Summary: {summary}")

        if len(summary.split()) >= 5:
            combined_summary += f"Article titled {article['title'].replace('/', '_')} has the following summary: {summary}. "
            doc.add_paragraph(f"Article: {article['title']}")
            doc.add_paragraph(f"URL: {article['url']}")
            doc.add_paragraph(f"Summary: {summary}\n")

    doc_filename = f"generative_ai_summary_{today_date}.docx"
    doc.save(doc_filename)
    upload_file(doc_filename)

    text_to_speech(combined_summary, f"generative_ai_summary_{today_date}.mp3")

# Run the main function
if __name__ == "__main__":
    main()
