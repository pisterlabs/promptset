from openai import OpenAI
import os
from dotenv import load_dotenv
from docx import Document

load_dotenv()
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')

client = OpenAI(
    # defaults to os.environ.get("OPENAI_API_KEY")
     api_key=OPENAI_API_KEY,
)

# def transcribe_audio(audio_file_path):
#     with open(audio_file_path, 'rb') as audio_file:
#         transcription = client.audio.transcriptions.create(
#             model="whisper-1",
#             file= audio_file,
#             response_format="json",
#             language="en"
#             )
#         print(transcription)
#         # This statement continually returns with the error 'Transcription' object is not subscriptable. 
#         # I just printed the transcription object to the console and used it from there.
#         return transcription['text']

def abstract_summary_extraction(transcription,completion_model):
    response = client.chat.completions.create(
        model=completion_model,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    # print(response)
    abstract = response.choices[0].message.content
    print(abstract)
    return abstract


def key_points_extraction(transcription,completion_model):
    response = client.chat.completions.create(
        model=completion_model,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    # print(response)
    key_points = response.choices[0].message.content
    print(key_points)
    return key_points


def action_item_extraction(transcription,completion_model):
    response = client.chat.completions.create(
        model=completion_model,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    # print(response)
    action_items = response.choices[0].message.content
    print(action_items)
    return action_items

def sentiment_analysis(transcription,completion_model):
    response = client.chat.completions.create(
        model=completion_model,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    # print(response)
    # print(response.choices[0].message.content)
    sentiment = response.choices[0].message.content
    # print(sentiment)
    return sentiment

def meeting_minutes(transcription,completion_model):
    abstract_summary = abstract_summary_extraction(transcription,completion_model)
    #write to file
    with open('abstract_summary.txt', 'w') as f:
        f.write(abstract_summary)
    key_points = key_points_extraction(transcription,completion_model)
    #write to file
    with open('key_points.txt', 'w') as f:
        f.write(key_points)
    action_items = action_item_extraction(transcription,completion_model)
    #write to file
    with open('action_items.txt', 'w') as f:
        f.write(action_items)
    sentiment = sentiment_analysis(transcription,completion_model)
    #write to file
    with open('sentiment.txt', 'w') as f:
        f.write(sentiment)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)

# basepath = os.path.dirname(__file__)
# audio_base_path = basepath+"/"+"Section-"
# audio_extension = ".mp3"
# transcript_base_path = basepath+"/"+"Transcript-"
# transcript_extension = ".txt"
# for i in range(3, 8):
#     transcription = transcribe_audio(audio_base_path+str(i)+audio_extension)
#     # Save transcription to a text file so that if the original part succeeeds and the later parts fail, we don't keep paying to transcribe. 
#     with open(transcript_base_path+str(i)+transcript_extension, 'w') as f:
#         f.write(transcription)

basepath = os.path.dirname(__file__)
audio_base_path = basepath+"/"+"Meeting-lossy"
audio_extension = ".mp3"
transcript_base_path = basepath+"/"+"Transcript-lossy"
transcript_extension = ".txt"
# transcription = transcribe_audio(audio_base_path+audio_extension)
# Save transcription to a text file so that if the original part succeeeds and the later parts fail, we don't keep paying to transcribe. 
# with open(transcript_base_path+transcript_extension, 'w') as f:
#     f.write(transcription)

# Load the transcription from the text file
with open(transcript_base_path+transcript_extension, 'r') as f:
    transcription = f.read()

# print(transcription)
# completion_model = "gpt-4-32k"
completion_model = "gpt-3.5-turbo-1106"
minutes = meeting_minutes(transcription,completion_model)
print(minutes)

save_as_docx(minutes, basepath+"/"+'meeting_minutes.docx')