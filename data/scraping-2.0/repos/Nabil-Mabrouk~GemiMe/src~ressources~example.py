from dotenv import load_dotenv
from langchain.vectorstores import Chroma
import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import CharacterTextSplitter
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.chains.summarize import load_summarize_chain
from langchain.llms import OpenAI
from langchain.callbacks import get_openai_callback
from docx import Document
from docx.table import _Cell
from streamlit_extras.add_vertical_space import add_vertical_space
import sys


def clear_history():
    if "history" in st.session_state:
        del st.session_state["history"]

def format_chat_history(chat_history):
    formatted_history = ""
    for entry in chat_history:
        question, answer = entry
        # Added an extra '\n' for the blank line
        formatted_history += f"Question: {question}\nAnswer: {answer}\n\n"
    return formatted_history


def extract_text_from_table(table):
    text = ""
    for row in table.rows:
        for cell in row.cells:
            if isinstance(cell, _Cell):
                text += cell.text + "\n"
    return text.strip()
#side bar contents

# Configure Streamlit page settings
st.set_page_config(page_title="PDFReader")
st.title("PDF & Word Reader âœ¨")
    

def main():
    if "model" not in st.session_state:
        st.session_state.model = "text-davinci-003"
    # brief summary
    with st.sidebar:
        st.title('ðŸ¤—ðŸ’¬ LLM PDFReader App')
        st.markdown("""
        ## About
        This app is an LLM-powered chatbot built using:
        - [Streamlit](https://streamlit.io/)
        - [Langchain](https://python.langchian.com/)
        - [OpenAI](https://platform.openai.com/docs/models) LLM model        
        """)
        st.radio(
        "Model ðŸ‘‰",
        key="model",
        options=["text-ada-001", "text-davinci-002", "text-davinci-003"],
        )
        add_vertical_space(5)




    llm = OpenAI(temperature=0.7, model=st.session_state.model)
    #llmchat = OpenAI(temperature=0.7, model_name='gpt-3.5-turbo')
    chain = load_summarize_chain(llm, chain_type="stuff")
    chain_large = load_summarize_chain(llm, chain_type="map_reduce")
    chain_qa = load_qa_chain(llm, chain_type="stuff")
    chain_large_qa = load_qa_chain(llm, chain_type="map_reduce")

   # Load environment variables 
    load_dotenv()


    # Upload file
    uploaded_file  = st.file_uploader("Upload your file", type=["pdf", "docx"])

    # Initialize session state
    if 'pdf_name' not in st.session_state:
        st.session_state.pdf_name = None

    # Extract the text
    if uploaded_file  is not None :
        file_type = uploaded_file.type

        # Clear summary if a new file is uploaded
        if 'summary' in st.session_state and st.session_state.file_name != uploaded_file.name:
            st.session_state.summary = None

        st.session_state.file_name = uploaded_file.name

        try:
            if file_type == "application/pdf":
                # Handle PDF files
                pdf_reader = PdfReader(uploaded_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()

            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # Handle Word documents
                doc = Document(uploaded_file)
                paragraphs = [p.text for p in doc.paragraphs]
                text = "\n".join(paragraphs)

                # Extract text from tables
                for table in doc.tables:
                    table_text = extract_text_from_table(table)
                    if table_text:
                        text += "\n" + table_text

            else:
                st.error("Unsupported file format. Please upload a PDF or DOCX file.")
                return

            # Split text into chunks, use this if you only use this app for small documents.
            # text_splitter = CharacterTextSplitter(
            #     separator="\n",
            #     chunk_size=1000,
            #     chunk_overlap=200,
            #     length_function=len
            # )

            # Split text into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len
            )
            chunks = text_splitter.split_text(text)


            # Create embeddings
            embeddings = OpenAIEmbeddings(disallowed_special=())
            knowledge_base = FAISS.from_texts(chunks, embeddings)


            st.header("Here's a brief summary of your file:")
            pdf_summary = "Give me a concise summary, use the language that the file is in. "

            docs = knowledge_base.similarity_search(pdf_summary)
            
            
            if 'summary' not in st.session_state or st.session_state.summary is None:
              with st.spinner('Wait for it...'):
                    with get_openai_callback() as scb:
                        try:
                            st.session_state.summary = chain.run(input_documents=docs, question=pdf_summary)    
                        except Exception as maxtoken_error:
                            # Fallback to the larger model if the context length is exceeded
                            print(maxtoken_error)
                            st.session_state.summary = chain_large.run(input_documents=docs, question=pdf_summary)
                        print(scb)    
                            
            st.write(st.session_state.summary)


            # User input for questions
            user_question = st.text_input("Ask a question about your file:")
            if user_question:
                docs = knowledge_base.similarity_search(user_question)
                with st.spinner('Wait for it...'):
                  with get_openai_callback() as cb:
                    try:
                        response = chain_qa.run(input_documents=docs, question=user_question)
                    except Exception as maxtoken_error:
                        print(maxtoken_error)
                        response = chain_large_qa.run(input_documents=docs, question=user_question) 
                    print(cb)
                    # Show/hide section using st.beta_expander
                    #with st.expander("Used Tokens", expanded=False):
                       #st.write(cb)
                st.write(response)
                
        except IndexError:
            #st.caption("Well, Seems like your PDF doesn't contain any text, try another one.ðŸ†–")
            st.error("Please upload another PDF. It seems like this PDF doesn't contain any text.")
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")



if __name__ == '__main__':
    main()