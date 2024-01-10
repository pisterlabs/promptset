import streamlit as st
import os
import openai

import io
# import docx
# import docx2pdf
from docx import Document
from langchain.document_loaders import PyPDFLoader
import tempfile
from langchain.chains.summarize import load_summarize_chain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.llms import OpenAI
from langchain import PromptTemplate
from langchain.utilities import WikipediaAPIWrapper
global doc_name, pdf_name


st.set_page_config(page_title="Paper Generator 🤖")
st.header("Paper Generator 📚🧠🤔")
accepted_file_types = ["pdf"]

lang = st.radio("Choose a Language for Paper Generator:", ("Urdu", "Arabic"))

method = st.radio("Choose a  Paper Generator:", ("Through Topic", "Through Pdf"))



openai.api_key = "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"
os.environ['OPENAI_API_KEY'] = "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"



def create_mcqs(content,lang,method,topic):
    if method== "Through Topic":
        
        message=[
            {
                "role": "system",
                "content": f"Please generate 5 multiple-choice questions (MCQs) with four answer options for each question based on the content. Ensure that the questions cover the content's main topics and key points. Create questions that collectively address all the topics from these content. Each question should have one correct answer and three plausible distractors. The questions should be clear, concise, and designed to test the reader's understanding of the content.Please note that the language should be : {lang} . The vocabulary of the should be simple "
            },
            {
                "role": "user",
                "content": content
            }
        ]
    else:
       
        message=[
            {
                "role": "system",
                "content": f"Please generate 5 multiple-choice questions (MCQs) with four answer options for each question based on the topic of the content provided. Ensure that the questions covers the topic main content and key points not whole content. Create questions that address only that topic from these content. Each question should have one correct answer and three plausible distractors. The questions should be clear, concise, and designed to test the reader's understanding of the topic.Please note that the language should be : {lang} . The vocabulary should be simple and the topic is {topic} "
            },
            {
                "role": "user",
                "content": content
            }
        ]
        
        
    
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0.5,
        messages=message
    )
  
    return response.choices[0].message.content

def short_question(content,lang,method,topic):
    if method== "Through Topic":
        message=[
            {
                "role": "system",
                "content": f"You are a highly skilled AI trained in creating quizzes for students. Generate 5 short unique questions only, based on the content. Ensure that the questions cover the content's main topics and key points. Create questions that collectively address all the topics from these content. The questions should be clear, concise, and designed to test the reader's understanding of the content. Only provide questions nothing else .Please note that the language should be : {lang} . The vocabulary should be simple "
            },
            {
                "role": "user",
                "content": content
            }
        ]
    else:
      
        message=[
            {
                "role": "system",
                "content": f"You are a highly skilled AI trained in creating quizzes for students. Generate 5 short unique questions only, based on the topic provided from the content. Ensure that the questions cover the topic main concepts and key points. Create questions that address topic only from the content. The questions should be clear, concise, and designed to test the reader's understanding of the content. Only provide questions nothing else .Please note that the language should be : {lang} . The vocabulary should be simple and the topic is: {topic} "
            },
            {
                "role": "user",
                "content": content
            }
        ]
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=message
    )
    return response.choices[0].message.content


def long_question(content, short_questions,lang,method,topic):
    if method== "Through Topic":
        message=[
            {
                "role": "system",
                "content": f"You are a highly skilled AI trained in creating quizes for students. Generate 3 long section questions,based on the content. Ensure that the questions cover the content's main topics and key points. Create questions that collectively address all the topics from the content. The questions should be clear, concise, and designed to test the reader's understanding of the content. The questions should not include these questions in any way: {short_questions}. Be unique.Ask different questions please try from different topics, Only provide questions nothing else. Please note that the language should be : {lang} . The vocabulary should be simple "
            },
            {
                "role": "user",
                "content": content
            }
        ]
    else:
        message=[
            {
                "role": "system",
                "content":f"You are a highly skilled AI trained in creating quizes for students. Generate 3 long section questions,based on the topic that is provided from the content. Ensure that the questions cover the topic only and topic'skey points. Create questions that collectively address all the concepts from the topic. The questions should be clear, concise, and designed to test the reader's understanding of the topic. The questions should not include these questions in any way: {short_questions}. Be unique.Ask different questions, Only provide questions nothing else. Please note that the language should be : {lang} . The vocabulary should be simple and the topic is : {topic}"
            },
            {
                "role": "user",
                "content": content
            }
        ]
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=message
    )
    
    return response.choices[0].message.content

def func(content,lang,method,topic):
    mcqs = create_mcqs(content,lang,method,topic)
    short = short_question(content,lang,method,topic)
    long = long_question(content, short,lang,method,topic)

    return {
        '1)_MCQS_Questions': mcqs,
        '2)_Short_Questions': short,
        '3)_Long_Questions': long,
    }

def save_as_docx(name, filename):
    doc = Document()
    for key, value in name.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)
    
@st.cache_data
def extract_text_from_pdf(uploaded_file,lang,topic):
    st.write(topic)
    # text = []
    llm = OpenAI(model_name="text-davinci-003")    

    loader = PyPDFLoader(uploaded_file)
    documents = loader.load()
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    text = text_splitter.split_documents(documents)
    map_custom_prompt='''
    Summarize the following text in a clear and concise way:
    TEXT:`{text}`
    Brief Summary:
    '''
    combine_custom_prompt='''
    Generate a summary of the following text that includes the following elements:

    * A title that accurately reflects the content of the text.
    * An introduction paragraph that provides an overview of the topic.
    * Bullet points that list the key points of the text.
    * A conclusion paragraph that summarizes the main points of the text.

    Text:`{text}`
'''
    map_prompt_template = PromptTemplate (
    input_variables=['text'],
    template=map_custom_prompt
    )
    
    combine_prompt_template = PromptTemplate(
    template=combine_custom_prompt, 
    input_variables=['text']
    )
    
    chain = load_summarize_chain(llm, chain_type="map_reduce", verbose=True ,map_prompt=map_prompt_template,combine_prompt=combine_prompt_template)
    response = chain.run(text)
    
    if response:
        with open(f"{lang}.txt", "w") as f:
            f.write(response)
            
    return True



def generate_paper(docs, lang, method,topic):
    if method == "Through Pdf":
        extracted_text = []
        for doc in docs:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                temp_file.write(doc.read())
                temp_file_path = temp_file.name

            extract_text_from_pdf(temp_file_path, lang,topic)
            os.remove(temp_file_path)
            with open(f"{lang}.txt", "r") as f:
                file = f.read()
                items = func(file, lang,method,topic)
                save_as_docx(items, f"{lang}.docx")
                st.success(f"Paper has been generated and saved as '{lang}.docx'")

            
    else:
        text = docs

        if text:
            
            file = text

            items = func(file, lang,method,topic)
            save_as_docx(items, f"{lang}.docx")
            st.success(f"Paper has been generated and saved as '{lang}.docx'")


   
# def generate_paper(docs,lang,method):
#         if method == "Through Pdf":
#             extracted_text = []
#             for doc in docs:
#                 with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
#                     temp_file.write(doc.read())
#                     temp_file_path = temp_file.name
                    
#                 text = extract_text_from_pdf(temp_file_path,lang)
#                 os.remove(temp_file_path)
#         else:
#             text = docs
              
#             if text:
#                  if method == "Through Pdf":
#                     with open(f"{lang}.txt","r") as f:
#                         file = f.read()
#                  else:
#                         file = text
#         items = func(file,lang)
#         save_as_docx(items, f"{lang}.docx")
#         st.success(f"Paper has been generated and saved as '{lang}.docx'")
    
                    
                    
def main(method,lang):
    if method == "Through Topic":
        # summary = Extractive(text)
        user_input = st.text_input("Enter Topic:")
        button_clicked = st.button('Generate!')
        if user_input and button_clicked:
            wikipedia = WikipediaAPIWrapper()
            response= wikipedia.run(user_input)
            
            with st.spinner("Runnings 🏃💨"):
                generate_paper(response,lang,method,None)
        
    elif method == "Through Pdf":
     
        docs = st.file_uploader("Upload Your Notes", type=accepted_file_types, accept_multiple_files=True)
        topic = st.text_input("Enter Topic From Pdf: ")
        button_clicked = st.button('Generate!')
        
        
        
        if docs and button_clicked:
            with st.spinner("Runnings 🏃💨"):
                generate_paper(docs,lang,method,topic)
                

   
main(method,lang)

