import streamlit as st
from pdfminer.high_level import extract_text
from pdfminer.high_level import extract_text_to_fp
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import json
import aiohttp
import asyncio
import time
import openai
import os
import re
import pandas as pd
from dotenv import load_dotenv


'''
Code copied from https://github.com/ChintaKrishnaMourya/resume_builder

'''
load_dotenv()  # take environment variables from .env

api_key = os.getenv("OPENAI_API_KEY")
API_KEY = api_key
openai.api_key = api_key



def convert_files_to_text(uploaded_file):
    try:
        file_name = uploaded_file.name
        if file_name.endswith('.pdf'):
            text = convert_pdf_to_text2(uploaded_file)
        elif file_name.endswith('.docx'):
            text = convert_docx_to_text(uploaded_file)
        elif file_name.endswith('.txt'):
            text = convert_txt_to_text(uploaded_file)
        else:
            return "Not a Resume"
        return text
    except Exception as e:
        print(f"Error converting file {file_name} to text: {e}")
        st.error(f"Error converting file {file_name} to text: {e}")
        return ""

def convert_pdf_to_text2(uploaded_file):
    try:
        output_string = BytesIO()
        extract_text_to_fp(BytesIO(uploaded_file.read()), output_string)
        return output_string.getvalue().decode()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""

def convert_docx_to_text(uploaded_file):
    try:
        doc = Document(BytesIO(uploaded_file.read()))
        text = ''
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\\n'
        return text
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

def convert_txt_to_text(uploaded_file):
    try:
        content = uploaded_file.getvalue().decode('utf-8')
        return content
    except Exception as e:
        return f"An error occurred: {e}"


def truncate_text_by_words(text, max_words=4000):
    """
    Truncates the text to a specified number of words.
    """
    words = text.split()
    if len(words) <= max_words:
        return text
    return " ".join(words[:max_words])


system1='''You are an excellent NLP engineer, skilled talent recruiter and data scientist and your task is to analyse and parse candidate resumes into meaningful structured JSON format.
        You will be provided with candidate resume text.
        The system instruction is:
        Step-1:
        Analyse and parse the following information from the candidate's resume, do not just extract the data, rephrase it meaningfully;
        return the meaningful parsed data in a sturctured JSON format with key and corresponding value format as follows-
        'name': string,
        'gmail': string,
        'phone number' : string,
        'address' : string
        'social media links': list of string,
        If value of a key is missing in the resume then value should be null.
        If not a resume then all the key's value should be null.
        Step-2:
        Only return the parsed JSON format resume, nothing else.'''

system2='''You are an excellent NLP engineer, skilled talent recruiter and data scientist and your task is to analyse and parse candidate resumes into meaningful structured JSON format.
        You will be provided with candidate resume text.
        The system instruction is:
        Step-1:
        Analyse and parse the following information from the candidate's resume, do not just extract the data, rephrase it meaningfully;
        return the meaningful parsed data in a sturctured JSON format with key and corresponding value format as follows-
        'about': string under 200 tokens
        'skillset and expertise': list of string,
        'certifications': list of string,
        'Explanation of projects': list of string under 200 tokens,
        'Explanation of position of responsibilities': list of string under 200 tokens,
        If value of a key is missing in the resume then value should be null.
        If not a resume then all the key's value should be null.
        Step-2:
        Only return the parsed JSON format resume, nothing else.'''

system3='''You are an excellent NLP engineer, skilled talent recruiter and data scientist and your task is to analyse and parse candidate resumes into meaningful structured JSON format.
        You will be provided with candidate resume text.
        The system instruction is:
        Step-1:
        Analyse and parse the following information from the candidate's resume, do not just extract the data, rephrase it meaningfully;
        return the meaningful parsed data in a sturctured JSON format with key and corresponding value format as follows-
        'years of experience': string,
        'Previous work experience description': list of string under 200 tokens,
        'educational qualification': list of string,
        'extracurriculars': list of string,
        'awards and achievements': list of string,
        'previous job title': list of string
        If value of a key is missing in the resume then value should be null.
        If not a resume then all the key's value should be null.
        Step-2:
        Only return the parsed JSON format resume, nothing else. '''

systems = [system1, system2, system3]


async def async_openai_request(session, resumetext, system):
    url = "https://api.openai.com/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    prompt = f"""
            Only return the structured parsed json format of the resume of candidate.
            Information about the candidate's resume is given inside text delimited by triple backticks.

            Candidate's Resume :```{resumetext}```

            """
    data = {
        "model": "gpt-3.5-turbo-16k",
        "messages": [{"role": "system", "content": system}, {"role": "user", "content": prompt}],
        "temperature": 0
    }
    async with session.post(url, json=data, headers=headers) as response:
        st.write(system)
        return await response.json()


async def fetch_responses(resumetext, system):
    st.write('Start fetching!')
    async with aiohttp.ClientSession() as session:
        results = await asyncio.gather(*(async_openai_request(session, resumetext, system) for system in systems))
    st.write('Ending fetching!')
    return results

def process_responses(responses):
    st.write('Starting to process fetched reponses!')
    output_list = [json.loads(resp['choices'][0]['message']['content']) for resp in responses]
    # Merging dictionaries
    combined_dict = {k: v for response in output_list for k, v in response.items()}
    st.write('Finished to process fetched reponses!')
    return combined_dict


## OpenAI api call
def get_choice_text_from_prompt(messages):
    try:
        
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo-16k",
            messages=messages,
            temperature=0,
            max_tokens=10000
        )
        choice_text = response.choices[0]["message"]["content"]
        return choice_text
    except Exception as e:
        print("Error in get_choice_text_from_prompt:", str(e))
        st.error(f"Error communicating with OpenAI: {e}")
        return ""


## Parsing resume
def parse_resume(resumetext):
    try:
        # Ensure resume text does not exceed 4000 tokens
        # resumetext = truncate_text_by_words(resumetext, 4000)
        resumetext = truncate_text_by_words(resumetext, 4000)
        system = """
        You are an excellent NLP engineer, skilled talent recruiter and data scientist and your task is to analyse and parse candidate resumes into meaningful structured JSON format.
        You will be provided with candidate resume text.

        The system instruction is:

        Step-1:
        Analyse and parse the following information from the candidate's resume, do not just extract the data, rephrase it meaningfully:
        name, gmail, phone number, address, social media links, about, skillset and expertise, certifications, Explanation of projects,
        Explanation of position of responsibilities, years of work experience in years and its explanation mentioning the job roles,
        Previous work experience description, educational qualification,
        extracurriculars,awards and achievements, previous job title.
        If value of a key is missing in the resume then value should be null.
        If not a resume then all the key's value should be null

        Step-2:
        Return the meaningful parsed data in a sturctured JSON format with key and corresponding value format as follows-
        'name': string,
        'gmail': string,
        'phone number' : string,
        'address' : string
        'social media links': list of string,
        'about': string under 200 tokens
        'skillset and expertise': list of string,
        'certifications': list of string,
        'Explanation of projects': list of string under 200 tokens,
        'Explanation of position of responsibilities': list of string under 200 tokens,
        'years of experience': string,
        'Previous work experience description': list of string under 200 tokens,
        'educational qualification': list of string,
        'extracurriculars': list of string,
        'awards and achievements': list of string,
        'previous job title': list of string
        If not a resume then all the key's value should be null.

         Step-3:
        Only return the parsed JSON format resume, nothing else.
        """
        prompt = f"""
        Only return the structured parsed json format of the resume of candidate.
        Information about the candidate's resume is given inside text delimited by triple backticks.

        Candidate's Resume :```{resumetext}```

        """

        messages =  [
        {'role':'system', 'content':system},
        {'role':'user', 'content': prompt}]

        # input_string = system + prompt
        # print('Resume Parsing input!')
        # st.write("Resume Parsing input!")
        # gpt_token_counts(input_string)

        start_time = time.time()
        parsed_resume = get_choice_text_from_prompt(messages)
        end_time = time.time()
        parsing_time = round(end_time - start_time, 1)
        st.sidebar.write('Time taken to parse: ',parsing_time,' seconds')
        print(f"Time taken to parse: {parsing_time} seconds")

        # print('Resume Parsing output!')
        # st.write("Resume Parsing output!")
        # gpt_token_counts(parsed_resume)

        return parsed_resume
    except Exception as e:
        print(f"Error parsing resume: {e}")
        st.error(f"Error parsing resume: {e}")
        return ""



# Function to set space after a paragraph to zero
def set_space_after(paragraph, space):
    p_spacing = OxmlElement('w:spacing')
    p_spacing.set(qn('w:after'), str(space))
    paragraph._element.get_or_add_pPr().append(p_spacing)

# Function to set cell background color
def set_cell_background(cell, fill):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Function for creating a document from a JSON-like structure
# def create_doc_from_json_template1(json_data, filename):

#     doc = Document()
    
#     # Set the page margins
#     section = doc.sections[0]
#     section.top_margin = Inches(0.5)
#     section.bottom_margin = Inches(0.5)
#     section.left_margin = Inches(0.5)
#     section.right_margin = Inches(0.5)

#     # Add a table for the header
#     header_table = doc.add_table(rows=1, cols=2)
#     header_table.autofit = False
#     header_table.columns[0].width = Inches(4.25)
#     header_table.columns[1].width = Inches(2.25)
    
#     # Populate header table
#     name_cell = header_table.cell(0, 0) 
#     name_cell.text = f"{json_data['name']}\n{json_data['social media links'][0]}"
#     set_cell_background(name_cell, 'ADD8E6')  # Light gray background
#     contact_cell = header_table.cell(0, 1)
#     contact_info = f"Contact No: {json_data['phone number']}\nEmail: {json_data['gmail']}\n"
#     contact_cell.text = contact_info
#     set_cell_background(contact_cell, '87CEEB')  # Slightly darker gray background

#     # Add content sections with titles
#     titles = ['educational qualification','skillset and expertise', 'Previous work experience description','certifications','awards and achievements','Explanation of position of responsibilities', 'Explanation of projects', 'years of experience', 'previous job title']
#     for title in titles:
#         table = doc.add_table(rows=2, cols=1)
#         table.autofit = False
#         table.columns[0].width = Inches(6.5)

#         title_cell = table.cell(0, 0)
#         # Apply bold to the title and set a light blue background
#         run = title_cell.paragraphs[0].add_run(title.title())
#         run.bold = True  # Convert title to title case
#         set_cell_background(title_cell, 'ADD8E6')  # Dark gray background for title

#         content_cell = table.cell(1, 0)
#         # Retrieve content from json_data based on title
#         content = json_data.get(title, 'Content not provided')
#         print(title,content)
#         # Special formatting for 'previous work experience description'
#         if title == 'previous work experience description' and isinstance(content, list):
#             for item in content:
#             # Create a paragraph for each experience item
#                 p = content_cell.add_paragraph(style='ListBullet')
#                 # Split the item into subtitle and description
#                 job_title, _, description = item.partition(': ')
#                 # Add the job title as bold
#                 p.add_run(job_title + ': ').bold = True
#                 # Continue with the description
#                 p.add_run(description)
#         elif title == 'skillset and expertise' and isinstance(content, list):
#             # Join the skills with a comma and a space for the 'skillset and expertise' section
#             content_cell.text = ', '.join(content)
#         elif isinstance(content, list):
#             # Add content as bullet points for list-type contents
#             for item in content:
#                 content_cell.add_paragraph(item, style='ListBullet')
#         elif content is None:  # Ensure content is not None
#             content = []
#         else:
#             content_cell.text = content  # Directly add content if not a list
        
#         print('/n \n')

#         # Set the space after each table to zero
#         set_space_after(table.rows[0].cells[0].paragraphs[0], 0)
#         set_space_after(table.rows[1].cells[0].paragraphs[0], 0)


#     # Save the document
#     doc.save(filename)
#     # Read the saved file into BytesIO object
#     with open(filename, "rb") as file:
#         doc_bytes = BytesIO(file.read())
#     return doc_bytes

def create_doc_from_json_template1(json_data, filename):
    doc = Document()
    
    # Set the page margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Add a table for the header
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Inches(4.25)
    header_table.columns[1].width = Inches(2.25)
    
    # Populate header table
    name_cell = header_table.cell(0, 0)
    name = json_data.get('name', 'Name not provided')
    social_media_links = json_data.get('social media links', [])
    social_media_link = social_media_links[0] if social_media_links else "Not provided"
    name_cell.text = f"{name}\n{social_media_link}"
    set_cell_background(name_cell, 'ADD8E6')  # Light blue background

    contact_cell = header_table.cell(0, 1)
    phone_number = json_data.get('phone number', 'Phone number not provided')
    gmail = json_data.get('gmail', 'Email not provided')
    contact_info = f"Contact No: {phone_number}\nEmail: {gmail}\n"
    contact_cell.text = contact_info
    set_cell_background(contact_cell, '87CEEB')  # Slightly darker blue background

    # Add content sections with titles
    titles = ['educational qualification', 'skillset and expertise', 'Previous work experience description', 
              'certifications', 'awards and achievements', 'Explanation of position of responsibilities', 
              'Explanation of projects', 'years of experience', 'previous job title']
    
    for title in titles:
        table = doc.add_table(rows=2, cols=1)
        table.autofit = False
        table.columns[0].width = Inches(6.5)

        title_cell = table.cell(0, 0)
        run = title_cell.paragraphs[0].add_run(title.title())
        run.bold = True  # Convert title to title case
        set_cell_background(title_cell, 'ADD8E6')  # Dark blue background for title

        content_cell = table.cell(1, 0)
        content = json_data.get(title)
        if content is None:
            content_cell.text = "Content not provided"
        elif isinstance(content, list):
            for item in content:
                content_cell.add_paragraph(item, style='ListBullet')
        else:
            content_cell.text = str(content)  # Ensure content is a string
            
        set_space_after(table.rows[0].cells[0].paragraphs[0], 0)
        set_space_after(table.rows[1].cells[0].paragraphs[0], 0)


    # Save the document
    doc.save(filename)
    # Read the saved file into BytesIO object
    with open(filename, "rb") as file:
        doc_bytes = BytesIO(file.read())
    return doc_bytes


# Function to set space after a paragraph to zero
def set_space_after2(paragraph, space):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Inches(space)

# Function for creating a document from a JSON-like structure
def create_doc_from_json_template2(json_data, filename):
    # Create a new Document
    doc = Document()
    
    # Set the page margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Header with name and contact information
    # Header with name and contact information
    doc.add_heading(json_data.get('name', 'Name Not Provided'), level=1)
    contact_paragraph = doc.add_paragraph()

    # Add email and phone number
    contact_paragraph.add_run(f"Email: {json_data.get('gmail', 'Email Not Provided')} | ")
    contact_paragraph.add_run(f"Phone: {json_data.get('phone number', 'Phone Not Provided')} | ")

    # Check and add LinkedIn link if available
    social_media_links = json_data.get('social media links', [])
    if social_media_links:
        contact_paragraph.add_run(f"LinkedIn: {social_media_links[0]}\n")
    else:
        contact_paragraph.add_run("LinkedIn: Not Provided\n")

    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_space_after2(contact_paragraph, 0.1)

    # # Skillset and Expertise
    # doc.add_heading('Skillset and Expertise', level=2)
    # skills = ', '.join(json_data['skillset and expertise'])
    # doc.add_paragraph(skills)
    
    # # Certifications
    # doc.add_heading('Certifications', level=2)
    # for certification in json_data['certifications']:
    #     doc.add_paragraph(certification, style='ListBullet')
    
    # # Projects
    # doc.add_heading('Projects', level=2)
    # for project in json_data['Explanation of projects']:
    #     doc.add_paragraph(project, style='ListBullet')
    
    # # Positions of Responsibility
    # doc.add_heading('Positions of Responsibility', level=2)
    # for position in json_data['Explanation of position of responsibilities']:
    #     doc.add_paragraph(position, style='ListBullet')

    # # Experience
    # doc.add_heading('Experience', level=2)
    # for experience in json_data['Previous work experience description']:
    #     doc.add_paragraph(experience, style='ListBullet')
    
    # # Education
    # doc.add_heading('Education', level=2)
    # for education in json_data['educational qualification']:
    #     doc.add_paragraph(education, style='ListBullet')
    
    # # Extracurriculars
    # doc.add_heading('Extracurriculars', level=2)
    # for activity in json_data['extracurriculars']:
    #     doc.add_paragraph(activity, style='ListBullet')
    
    # # Awards and Achievements
    # doc.add_heading('Awards and Achievements', level=2)
    # for award in json_data['awards and achievements']:
    #     doc.add_paragraph(award, style='ListBullet')
    
    # # Save the document
    # doc.save(filename)

    # Define the keys and their corresponding headings
    sections = {
        'skillset and expertise': 'Skillset and Expertise',
        'certifications': 'Certifications',
        'Explanation of projects': 'Projects',
        'Explanation of position of responsibilities': 'Positions of Responsibility',
        'Previous work experience description': 'Experience',
        'educational qualification': 'Education',
        'extracurriculars': 'Extracurriculars',
        'awards and achievements': 'Awards and Achievements'
    }

    # Iterate through each section
    for key, title in sections.items():
        # Check if the key exists and is not None
        if key in json_data and json_data[key] is not None:
            # Add heading
            doc.add_heading(title, level=2)
            # Iterate through items and add them to the document
            for item in json_data[key]:
                doc.add_paragraph(item, style='ListBullet')
    
    # Save the document
    doc.save(filename)
    # Read the saved file into BytesIO object
    with open(filename, "rb") as file:
        doc_bytes = BytesIO(file.read())
    return doc_bytes


# Function to add a cell with a colored background
def set_cell_background(cell, fill):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)

## Function to create a two-column table for layout
def create_two_column_table(doc, json_data):
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Cm(5.5)  # Set the width of the left column
    table.columns[1].width = Cm(11.5) # Set the width of the right column

    # Left column content
    left_cell = table.cell(0, 0)
    set_cell_background(left_cell, "ADD8E6")  # Light blue background
    left_paragraph = left_cell.paragraphs[0]
    
    # Process each section
    process_section(left_paragraph, json_data, 'name', "Name", font_size=14, bold=True)
    process_section(left_paragraph, json_data, 'gmail', "Email")
    process_section(left_paragraph, json_data, 'phone number', "Phone")
    process_section(left_paragraph, json_data, 'social media links', "LinkedIn", is_list=True)
    process_section(left_paragraph, json_data, 'educational qualification', "Educational Qualification", is_list=True)
    process_section(left_paragraph, json_data, 'skillset and expertise', "Skillset and Expertise", is_list=True)
    process_section(left_paragraph, json_data, 'previous job title', "Previous Job Roles", is_list=True)
    process_section(left_paragraph, json_data, 'years of experience', "Years of Experience")
    process_section(left_paragraph, json_data, 'certifications', "Certifications", is_list=True)

    # Right column content
    right_cell = table.cell(0, 1)
    right_paragraph = right_cell.paragraphs[0]

    process_section(right_paragraph, json_data, 'Previous work experience description', "Previous Work Experience Description", is_list=True)
    process_section(right_paragraph, json_data, 'Explanation of projects', "Projects", is_list=True)
    process_section(right_paragraph, json_data, 'awards and achievements', "Awards and Achievements", is_list=True)
    process_section(right_paragraph, json_data, 'Explanation of position of responsibilities', "Explanation of Position of Responsibilities", is_list=True)
    process_section(right_paragraph, json_data, 'extracurriculars', "Extracurricular Activities", is_list=True)

def process_section(paragraph, json_data, key, title, is_list=False, font_size=None, bold=False):
    content = json_data.get(key)
    if content is None:
        return  # Skip this section if content is None

    if bold:
        run = paragraph.add_run(title + "\n")
        run.bold = True
        if font_size:
            run.font.size = Pt(font_size)
    else:
        paragraph.add_run(title + ":\n").bold = True

    if is_list and isinstance(content, list):
        for item in content:
            paragraph.add_run(f"• {item}\n")
    else:
        paragraph.add_run(f"{content}\n")


# Function for creating a document from a JSON-like structure
def create_doc_from_json_template3(json_data, filename):
    doc = Document()
    # Set the page margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Create the two-column layout
    create_two_column_table(doc, json_data)

    # Save the document
    doc.save(filename)
    # Read the saved file into BytesIO object
    with open(filename, "rb") as file:
        doc_bytes = BytesIO(file.read())
    return doc_bytes

def download_docx(doc_bytes, filename):
    st.download_button(
        label='Download Resume',
        data=doc_bytes,
        file_name=filename,
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

def run_async_code(systems, resumetext):
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    st.write('Getting fetching reponses!')
    responses = loop.run_until_complete(fetch_responses(systems, resumetext))
    return process_responses(responses)


