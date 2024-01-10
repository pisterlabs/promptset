import os
from flask import Flask, request, render_template
from flask_wtf import FlaskForm
from wtforms import FileField, SelectMultipleField, SubmitField
from werkzeug.utils import secure_filename
from docx import Document
import openai
from flask_wtf.file import FileRequired
from flask_bootstrap import Bootstrap
from docx.shared import RGBColor

app = Flask(__name__)
app.config['SECRET_KEY'] = 'ef1ed9ef3d738bb2c12bdc436828ac6b'
Bootstrap(app)

openai.api_key = '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'

# Define form
class UploadForm(FlaskForm):
    file = FileField('Upload File', validators=[FileRequired()])
    language = SelectMultipleField('Select Languages', choices=[
        ('Mandarin Chinese', 'Mandarin Chinese'), 
        ('Spanish', 'Spanish'),
        # ... add all the other languages here ...
        ('Simplified Chinese', 'Simplified Chinese'), 
        ('Traditional Chinese', 'Traditional Chinese')
    ])
    submit = SubmitField('Translate')

@app.route('/', methods=['GET', 'POST'])
def index():
    form = UploadForm()
    if form.validate_on_submit():
        filename = secure_filename(form.file.data.filename)
        upload_folder = 'uploads'
        if not os.path.exists(upload_folder):
            os.makedirs(upload_folder)
        form.file.data.save(os.path.join(upload_folder, filename))
        selected_languages = form.language.data  # get selected languages
        return process_file(os.path.join(upload_folder, filename), selected_languages)
    return render_template('index.html', form=form)

def process_file(filepath, languages):
    document = Document(filepath)
    translation_doc = Document()

    for language in languages:
        translation_doc.add_heading(language, level=1)

        for paragraph in document.paragraphs:
            # Skip paragraph if it is empty or only contains whitespace
            if not paragraph.text.strip():
                continue

            # Add the original English text in red
            original_paragraph = translation_doc.add_paragraph()
            original_paragraph.add_run(paragraph.text).font.color.rgb = RGBColor(255, 0, 0)  # RGB values for red

            translation = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "user", "content": 'Translate the following English text to {}: "{}"'.format(language, paragraph.text)}
                ]
            )
            translation_doc.add_paragraph(translation['choices'][0]['message']['content'])

            translation_doc.add_paragraph()

        translation_doc.save('translated_doc.docx')
        
   

    return 'Translation completed'

if __name__ == "__main__":
    app.run(debug=True)
