from docx import Document
from docx.shared import RGBColor
from langchain import PromptTemplate
from chatgpt_client_chapter2 import ChatGPTBase, ChatGPTGoogleSearch, ChatGPTBingSearch
from myutils import read_yaml, read_text_file


class ReportGenerator:
    """
    This is the main class for generating the reports related to the project.
    """

    def __init__(self, template_config_path: str):
        self.chat = ChatGPTBase()
        self.chat_google = ChatGPTGoogleSearch()
        self.chat_bing = ChatGPTBingSearch()
        self.project_report = Document()
        self.prompt_eng_report = Document()
        self.template_config_path = template_config_path
        self.template = self.load_template_config()
        self.sections = self.template["content"]
        self.textfiles_dir = self.template["textfiles_dir"]
        self.responses = {} # a dictionary to store the responses of the chatgpt to the prompts in case we need them later


    def generate_content_menu(self, sections):
        for section in sections:
            indent = (section['level'] - 1) * 4
            self.project_report.add_heading((indent * ' ' + section['title']), level=section['level'])
            if len(section['subsections']):
                self.generate_content_menu(section['subsections'])
            else:
                continue

    def write_the_chapter(self, sections):
        for section in sections:

            # write the title of the section
            print("writing section title: " + section['title'] + " ...")
            self.project_report.add_heading(section['title'], level=section['level'])
            self.prompt_eng_report.add_heading(section['title'], level=section['level'])

            # write the text file to the section
            if len(section['text_files']):
                for text_file in section['text_files']:
                    print("writing text file: " + text_file + " ...")
                    text = read_text_file(self.textfiles_dir + text_file)
                    self.project_report.add_paragraph(text)
                    self.prompt_eng_report.add_paragraph(text)

            # write the prompt to the prompt_eng section
            if len(section['prompt_files']):
                for prompt_file in section['prompt_files']:
                    print("reading prompt file: " + prompt_file + " ...")
                    prompt = read_text_file(self.textfiles_dir + prompt_file)

                    # create a prompt template object from the prompt to check if it has input variables
                    lang_chain_prompt = PromptTemplate.from_template(prompt)

                    # check if the prompt has input variables. some prompts need the answer
                    # of the previous prompt which is replicated in those prompts as {prompt_file_name}
                    # so we need to replace those with the actual answer of the previous prompt. there 
                    # should be just one input variable in the prompt in this design.
                    if len(lang_chain_prompt.input_variables):
                        print("prompt has input variables ...")
                        # format the prompt with the input variables
                        print("formatting prompt with input variables ...")
                        input_variable_name = lang_chain_prompt.input_variables[0]
                        prompt = prompt.format(**{input_variable_name:self.responses[input_variable_name]})
                    
                    # write the prompt to the section in prompt eng report
                    p1 = self.prompt_eng_report.add_paragraph()
                    p1.add_run(prompt).font.color.rgb = RGBColor(26, 93, 26)  # green

                    # send the prompt to the chatgpt
                    print("sending prompt to chatgpt ...")
                    response_to_prompt_from_file = self.chat.llm.predict(prompt)
                    self.responses[prompt_file[:-4]] = response_to_prompt_from_file
                    # write the response to the section
                    print("writing response to prompt to prompt eng report...")
                    p2 = self.prompt_eng_report.add_paragraph()
                    p2.add_run(response_to_prompt_from_file).font.color.rgb = RGBColor(128, 0, 128)  # purple

                    # write the response to the project report section
                    print("writing answer to project report section...")
                    p2_report = self.project_report.add_paragraph()
                    p2_report.add_run(response_to_prompt_from_file).font.color.rgb = RGBColor(128, 0, 128) #purple

                    # send the prompt to chatgpt with google search
                    print("sending prompt to chatgpt with google search tool ...")
                    response_with_google_search = self.chat_google.agent.run(prompt)
                    self.responses[prompt_file[:-4]] += ("\n" + response_with_google_search) # [:-4] to remove the .txt from the prompt file name
                    # write the response to the section
                    print("writing response of chatgpt with google search to prompt eng report...")
                    self.prompt_eng_report.add_paragraph("chatgpt answer with google search: ")
                    p3 = self.prompt_eng_report.add_paragraph()
                    p3.add_run(response_with_google_search).font.color.rgb = RGBColor(128, 0, 128)  # purple

                    # write the response to the project report section
                    print("writing response of chatgpt with google search to project report...")
                    self.project_report.add_paragraph("chatgpt answer with google search: ")
                    p3_report = self.project_report.add_paragraph()
                    p3_report.add_run(response_with_google_search).font.color.rgb = RGBColor(128, 0, 128) #purple

                    # send the prompt to chatgpt with bing search
                    print("sending prompt to chatgpt with bing search tool ...")
                    response_with_bing_search = self.chat_bing.agent.run(prompt)
                    self.responses[prompt_file[:-4]] += ("\n" + response_with_bing_search)
                    # write the response to the section
                    print("writing response of chatgpt with bing search to prompt eng report...")
                    self.prompt_eng_report.add_paragraph("chatgpt answer with bing search: ")
                    p4 = self.prompt_eng_report.add_paragraph()
                    p4.add_run(response_with_bing_search).font.color.rgb = RGBColor(128, 0, 128)  # purple
                    # write the final answer to the project report section
                    print("writing final answer to project report section...")
                    self.project_report.add_paragraph("chatgpt answer with bing search: ")
                    p4_report = self.project_report.add_paragraph()
                    p4_report.add_run(response_with_bing_search).font.color.rgb = RGBColor(128, 0, 128)  # purple

            if len(section['subsections']):
                self.write_the_chapter(section['subsections'])

    def load_template_config(self):
        return read_yaml(self.template_config_path)


    def save_document(self):
        self.project_report.save('output/project_report_chapter2.docx')
        self.prompt_eng_report.save('output/prompt_eng_report_chapter2.docx')

