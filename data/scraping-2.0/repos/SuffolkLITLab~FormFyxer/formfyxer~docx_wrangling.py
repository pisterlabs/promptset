import docx
import sys
import os
from openai import OpenAI

import tiktoken
import json
from docx.oxml import OxmlElement
import re

from typing import List, Tuple, Optional, Union

__all__ = [
    "get_labeled_docx_runs",
    "update_docx",
    "modify_docx_with_openai_guesses",
    "get_docx_repr",
    "get_modified_docx_runs",
    "make_docx_plain_language",
]


def add_paragraph_after(paragraph, text):
    p = OxmlElement("w:p")
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text

    r.append(t)
    p.append(r)
    paragraph._element.addnext(p)


def add_paragraph_before(paragraph, text):
    p = OxmlElement("w:p")
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text

    r.append(t)
    p.append(r)
    paragraph._element.addprevious(p)

def add_run_after(run, text):
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text

    r.append(t)
    run._element.addnext(r)


def update_docx(
    document: Union[docx.Document, str], modified_runs: List[Tuple[int, int, str, int]]
) -> docx.Document:
    """Update the document with the modified runs.

    Note: OpenAI is probabilistic, so the modified run indices may not be correct. 
    When the index of a run or paragraph is out of range, a new paragraph 
    will be inserted at the end of the document or a new run at the end of the
    paragraph's runs.

    Take a careful look at the output document to make sure it is still correct.

    Args:
        document: the docx.Document object, or the path to the DOCX file
        modified_runs: a tuple of paragraph number, run number, the modified text, a question (not used), and whether a new paragraph should be inserted (for conditional text)

    Returns:
        The modified document.
    """
    ## Sort modified_runs in reverse order so inserted paragraphs are in the correct order
    modified_runs.sort(key=lambda x: x[0], reverse=True)

    if isinstance(document, str):
        document = docx.Document(document)

    for item in modified_runs:
        if len(item) > 4:
            continue
        paragraph_number, run_number, modified_text, new_paragraph = item
        if paragraph_number >= len(document.paragraphs):
            add_paragraph_after(document.paragraphs[-1], modified_text)
            continue
        paragraph = document.paragraphs[paragraph_number]
        if run_number >= len(paragraph.runs):
            add_run_after(paragraph.runs[-1], modified_text)
            continue
        run = paragraph.runs[run_number]
        if new_paragraph == 1:
           add_paragraph_after(paragraph, modified_text)
        elif new_paragraph == -1:
           add_paragraph_before(paragraph, modified_text)
        else:
            run.text = modified_text
    return document

def get_docx_repr(docx_path: str, paragraph_start:int=0, paragraph_end:Optional[int]=None):
    """Return a JSON representation of the paragraphs and runs in the DOCX file.

    Args:
        docx_path: path to the DOCX file
    
    Returns:
        A JSON representation of the paragraphs and runs in the DOCX file.
    """
    items = []
    paragraphs = docx.Document(docx_path).paragraphs[paragraph_start:paragraph_end]
    for pnum, paragraph in enumerate(paragraphs):
        for rnum, run in enumerate(paragraph.runs):
            items.append(
                [
                    pnum,
                    rnum,
                    run.text,
                ]
            )
    return repr(items)

def get_labeled_docx_runs(
    docx_path: Optional[str] = None,
    docx_repr = Optional[str],
    custom_people_names: Optional[Tuple[str, str]] = None,
    openai_client: Optional[OpenAI] = None,
    api_key: Optional[str] = None,
) -> List[Tuple[int, int, str, int]]:
    """Scan the DOCX and return a list of modified text with Jinja2 variable names inserted.

    Args:
        docx_path: path to the DOCX file
        docx_repr: a string representation of the paragraphs and runs in the DOCX file, if docx_path is not provided. This might be useful if you want
        custom_people_names: a tuple of custom names and descriptions to use in addition to the default ones. Like: ("clients", "the person benefiting from the form")

    Returns:
        A list of tuples, each containing a paragraph number, run number, and the modified text of the run.
    """

    custom_name_text = ""
    if custom_people_names:
        assert isinstance(custom_people_names, list)
        for name, description in custom_people_names:
            custom_name_text += f"    {name} ({description}), \n"

    custom_example = """Example input, with paragraph and run numbers indicated:
    [
        [0, 1, "Dear John Smith:"],
        [1, 0, "This sentence can stay as is in the output and will not be in the reply."],
        [2, 0, "[Optional: if you are a tenant, include this paragraph]"],
    ]"""

    instructions = """The purpose of the resulting document is to be used as a template within a Docassemble interview, with Jinja2 markup.
    Steps:
    1. Analyze the document. Identify placeholder text and repeated _____ that should be replaced with a variable name.
    2. Insert jinja2 tags around a new variable name that represents the placeholder text.
    3. Mark optional paragraphs with conditional Jinja2 tags.
    4. Text intended for verbatim output in the final document will remain unchanged.

    Example reply, indicating paragraph, run, the new text, and a number indicating if this changes the 
    current paragraph, adds one before, or adds one after (-1, 0, 1):

    {"results":
    [
        [0, 1, "Dear {{ other_parties[0] }}:", 0],
        [2, 0, "{%p if is_tenant %}", -1],
        [3, 0, "{%p endif %}", 1],
    ]
    }
    """

    instructions += f"""
    Rules for variable names:
        1. Variables usually refer to people or their attributes.
        2. People are stored in lists.
        3. We use Docassemble objects and conventions.
        4. Use variable names and patterns from the list below. Invent new variable names when it is appropriate.

    List names for people:
        {custom_people_names}
        users (for the person benefiting from the form, especially when for a pro se filer)
        other_parties (the opposing party in a lawsuit or transactional party)
        plaintiffs
        defendants
        petitioners
        respondents
        children
        spouses
        parents
        caregivers
        attorneys
        translators
        debt_collectors
        creditors
        witnesses
        guardians_ad_litem
        guardians
        decedents
        interested_parties

        Name Forms:
            users (full name of all users)
            users[0] (full name of first user)
            users[0].name.full() (Alternate full name of first user)
            users[0].name.first (First name only)
            users[0].name.middle (Middle name only)
            users[0].name.middle_initial() (First letter of middle name)
            users[0].name.last (Last name only)
            users[0].name.suffix (Suffix of user's name only)

    Attribute names (replace `users` with the appropriate list name):
        Demographic Data:
            users[0].birthdate (Birthdate)
            users[0].age_in_years() (Calculated age based on birthdate)
            users[0].gender (Gender)
            users[0].gender_female (User is female, for checkbox field)
            users[0].gender_male (User is male, for checkbox field)
            users[0].gender_other (User is not male or female, for checkbox field)
            users[0].gender_nonbinary (User identifies as nonbinary, for checkbox field)
            users[0].gender_undisclosed (User chose not to disclose gender, for checkbox field)
            users[0].gender_self_described (User chose to self-describe gender, for checkbox field)
            user_needs_interpreter (User needs an interpreter, for checkbox field)
            user_preferred_language (User's preferred language)

        Addresses:
            users[0].address.block() (Full address, on multiple lines)
            users[0].address.on_one_line() (Full address on one line)
            users[0].address.line_one() (Line one of the address, including unit or apartment number)
            users[0].address.line_two() (Line two of the address, usually city, state, and Zip/postal code)
            users[0].address.address (Street address)
            users[0].address.unit (Apartment, unit, or suite)
            users[0].address.city (City or town)
            users[0].address.state (State, province, or sub-locality)
            users[0].address.zip (Zip or postal code)
            users[0].address.county (County or parish)
            users[0].address.country (Country)

        Other Contact Information:
            users[0].phone_number (Phone number)
            users[0].mobile_number (A phone number explicitly labeled as the "mobile" number)
            users[0].phone_numbers() (A list of both mobile and other phone numbers)
            users[0].email (Email)

        Signatures:
            users[0].signature (Signature)
            signature_date (Date the form is completed)

        Information about Court and Court Processes:
            trial_court (Court's full name)
            trial_court.address.county (County where court is located)
            trial_court.division (Division of court)
            trial_court.department (Department of court)
            docket_number (Case or docket number)
            docket_numbers (A comma-separated list of docket numbers)
            
    When No Existing Variable Name Exists:
        1. Craft short, readable variable names in python snake_case.
        2. Represent people with lists, even if only one person.
        3. Use valid Python variable names within complete Jinja2 tags, like: {{ new_variable_name }}.

        Special endings:
            Suffix _date for date values.
            Suffix _value or _amount for currency values.

        Examples: 
        "(State the reason for eviction)" transforms into `{{ eviction_reason }}`.
    """
    return get_modified_docx_runs(
        docx_path = docx_path,
        docx_repr = docx_repr,
        custom_example=custom_example,
        instructions=instructions,
        openai_client=openai_client,
        api_key=api_key,
    )

def get_modified_docx_runs(
        docx_path: Optional[str] = None,
        docx_repr: Optional[str] = None,
        custom_example:str = "",
        instructions:str = "",
        openai_client: Optional[OpenAI] = None, 
        api_key:Optional[str]=None,
        temperature=0.5,
) -> List[Tuple[int, int, str, int]]:
    """Use GPT to rewrite the contents of a DOCX file paragraph by paragraph. Does not handle tables, footers, or
    other structures yet.

    This is a light wrapper that provides the structure of DOCX paragraphs and runs to your prompt
    to OpenAI to facilitate the rewriting of the document without disrupting formatting.

    For example, this could be used to:
    * Remove any passive voice
    * Replace placeholder text with variable names
    * Rewrite to a 6th grade reading level
    * Do an advanced search and replace, without requiring you to use a regex

    By default, the example prompt includes a sample like this:

    [
        [0, 0, "Dear "],
        [0, 1, "John Smith:"],
        [1, 0, "I hope this letter finds you well."],
    ]

    Your custom instructions should include an example of how the sample will be modified, like the one below: 
    
    Example reply, indicating paragraph, run, the new text, and a number indicating if this changes the 
    current paragraph, adds one before, or adds one after (-1, 0, 1):

    {"results":
        [
            [0, 1, "Dear {{ other_parties[0] }}:", 0],
            [2, 0, "{%p if is_tenant %}", -1],
            [3, 0, "{%p endif %}", 1],
        ]
    }

    You may also want to customize the input example to better match your use case.

    Args:
        docx_path (str): path to the DOCX file
        docx_repr (str): a string representation of the paragraphs and runs in the DOCX file, if docx_path is not provided.
        custom_example (Optional[str]): a string containing the purpose and overview of the task
        instructions (str) a string containing specific instructions for the task
        openai_client (Optional[OpenAI]): an OpenAI client object. If not provided a new one will be created.
        api_key (Optional[str]): an OpenAI API key. If not provided, it will be obtained from the environment
        temperature (float): the temperature to use when generating text. Lower temperatures are more conservative.

    Returns:
        A list of tuples, each containing a paragraph number, run number, and the modified text of the run.
    """
    if docx_path:
        docx_repr = get_docx_repr(docx_path)
    elif not docx_repr:
        raise Exception("Either docx_path or docx_repr must be provided.")

    assert isinstance(docx_repr, str)

    if not openai_client:
        openai_client = OpenAI(
            api_key = api_key or os.environ.get("OPENAI_API_KEY")
        )

    if not custom_example:
        custom_example = """[
        [0, 0, "Dear"],
        [0, 1, "John Smith:"],
        [1, 0, "I hope this letter finds you well."],
    ]"""

    if not "[" in instructions: # Make sure we have at least a minimal example of the output
        instructions += """The result will look like this:

    {"results":
        [
            [0, 1, "modified run", 0],
            [1, 0, "another modified run, skipping the run that should be left alone", 0],
        ]
    }
    """
        
    role_description = f"""
    You will process a DOCX document and return a JSON structure that transforms the DOCX file
    based on the following guidelines and examples. The DOCX will be provided as an annotated series of
    paragraphs and runs in JSON structure, like this:

    { custom_example }

    The result will be a JSON structure that includes a list of modified runs, each run represented as a list with exactly 4 items:
    1. The paragraph number
    2. The run number
    3. The modified text of the run
    4. A number indicating if this changes the current paragraph, adds one before, or adds one after (-1, 0, 1)
    
    {instructions}

    The reply ONLY contains the runs that have modified text.
    """

    encoding = tiktoken.encoding_for_model("gpt-4")

    encoding = tiktoken.encoding_for_model("gpt-4")
    token_count = len(encoding.encode(role_description + docx_repr))

    if token_count > 128000:
        raise Exception(
            f"Input to OpenAI is too long ({token_count} tokens). Maximum is 128000 tokens."
        )

    moderation_response = openai_client.moderations.create(input=role_description + docx_repr)
    if moderation_response.results[0].flagged:
        raise Exception(
            f"OpenAI moderation error: {moderation_response.results[0]}"
        )

    response = openai_client.chat.completions.create(
        model="gpt-4-1106-preview",
        messages=[
            {"role": "system", "content": role_description},
            {"role": "user", "content": docx_repr},
        ],
        response_format={"type": "json_object"},
        temperature=temperature,
        max_tokens=4096,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0,
    )

    assert isinstance(response.choices[0].message.content, str)

    # check finish reason
    if response.choices[0].finish_reason != "stop":
        raise Exception(
            f"OpenAI did not finish processing the document. Finish reason: {response.choices[0].finish_reason}"
        )
    guesses = json.loads(response.choices[0].message.content)["results"]
    return guesses

def make_docx_plain_language(docx_path: str) -> docx.Document:
    """
    Convert a DOCX file to plain language with the help of OpenAI.
    """
    guesses = get_modified_docx_runs(
        docx_path,
        custom_example="""[
        [0, 0, "If the location of the land is in a state other than the state in which the tribe’s reservation is located, the tribe’s justification of anticipated benefits from the acquisition will be subject to greater scrutiny."],
        [1, 0, "When the process of freeing a vehicle that has been stuck results in ruts or holes, the operator will fill the rut or hole created by such activity before removing the vehicle from the immediate area."],
    ]""",
        instructions="""        
        You are a plain language expert whose goal is to rewrite the document at a 6th grade reading level, without changing the meaning of the document.
        You will rewrite passive voice sentences in the active voice. You will use simple vocabulary words to replace complex ones. You will use short sentences and short paragraphs.

        The result will look like this:

    {"results":
        [
            [0, 0, "If the land is in a different State than the tribe’s reservation, we will scrutinize the tribe’s justification of anticipated benefits more thoroughly.", 0],
            [1, 0, "If you make a hole while freeing a stuck vehicle, you must fill the hole before you drive away.", 0],
        ]
    }
    """,
    
    )
    return update_docx(docx.Document(docx_path), guesses)

def modify_docx_with_openai_guesses(docx_path: str) -> docx.Document:
    """Uses OpenAI to guess the variable names for a document and then modifies the document with the guesses.

    Args:
        docx_path (str): Path to the DOCX file to modify.

    Returns:
        docx.Document: The modified document, ready to be saved to the same or a new path
    """
    guesses = get_labeled_docx_runs(docx_path)

    return update_docx(docx.Document(docx_path), guesses)


if __name__ == "__main__":
    new_doc = modify_docx_with_openai_guesses(sys.argv[1])
    new_doc.save(sys.argv[1] + ".output.docx")