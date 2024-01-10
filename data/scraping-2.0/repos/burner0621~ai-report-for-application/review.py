import json
import openai
import re
import enchant
import zipcodes
import pycountry
from uszipcode import SearchEngine
from geopy.geocoders import Nominatim
from geopy.exc import GeocoderTimedOut
from datetime import datetime, timedelta
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, RGBColor

openai.api_key = "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"

STYLE_BUNDLE = {
    'TITLE': {
        'fontStyle': 'Arial',
        'fontSize': 48  
    },
    'HEADER_ONE' : {
        'fontStyle': "Times",
        'fontSize': 32
    },
    'NORMAL': {
        'fontStyle': "Times",
        'fontSize': 12
    },
    'ITEM' : {
        'fontStyle': "Arial",
        'fontSize': 16
    }
}

class Review():
    
    data = {}
    documentTitle = 'Report For Application'
    pdf = None
    fileName = 'Report'
    width = 210
    height = 297
    marginW = marginH = 5
    
    def __init__(self, inputFilePath): # inputFileNam = 'input.json'
        super().__init__()
        with open(inputFilePath, 'r') as f:
            self.data = json.load(f)

        self.patternPinfoName = r"^[A-za-z]+(\s|,)+[A-za-z]+(\s|,)+[A-za-z]*$"
        self.patternPinfoMaiden = r"^(\s*N/A\s*) | ([A-za-z]+(\s|,)+[A-za-z]+(\s|,)+[A-za-z]*)$"
        self.patternPhoneNumber = r"^\d\d\d-\d\d\d-\d\d\d\d$"
        self.patternEmail = r"^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$"
        self.patternAddress = r"^\d+\s+\w+\s+\w+\s+\w+$"
        self.questionPinfoName = "This is someone's First Name and Last Name.\n" + self.data['Personal Information']['Name'] + "\n Is this valid format?"
        self.questionPinfoMaiden = "This is someone's Maiden Name.\n" + self.data['Personal Information']['Name'] + "\n Is this valid format for common maiden name?"
        self.questionPinfoPassport = "This is " + self.data['Personal Information']['Country of Issuance'] + "' passport - " + self.data['Personal Information']['Passport Number'] + " Is this format correct?"
        self.doc = Document()
        
        self.titles ()

    def titles(self):
        self.doc.add_heading(self.documentTitle, 0)
        self.doc.add_paragraph ("This report can help you correct your applications.")

    def getAnswerFromAI(self, question):
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                    {"role": "system", "content": "You are a chatbot"},
                    {"role": "user", "content": question},
                ]
        )

        result = ''
        for choice in response.choices:
            result += choice.message.content

        return result

    def checkFormat(self, pattern, str):
        match = re.match(pattern, str)
        if(match == None):
            return False
        else:
            return True
        
    def spellCheck(self, text):
        d = enchant.Dict("en_US")
        words = text.split()
        misspelled = [word for word in words if word.isalpha() and not d.check(word)]
        return misspelled

    def isValidAddress(self, address):
        geolocator = Nominatim(user_agent='review-app')
        try:
            location = geolocator.geocode(address)
            return True if location else False
        except GeocoderTimedOut:
            return False
    
    def validateCityStateZip(self, city, state, zip_code):
        search = SearchEngine()
        result = search.by_city_and_state(city, state)
        if not result:
            return False
        zip_info = zipcodes.matching(zip_code)
        if not zip_info:
            return False
        zip_info = zip_info[0]
        if zip_info['city'] != city or zip_info['state'] != state:
            return False
        return True

    def validatePassport(self, passport_number):
        pattern = r'^[A-Z]{2}\d{7}$'
        if re.match(pattern, passport_number):
            return True
        else:
            return False
    
    def validateCountry(self, country_name):
        try:
            # Attempt to lookup country by name
            country = pycountry.countries.get(name=country_name)
            return True
        except KeyError:
            try:
                # Attempt to lookup country by ISO code
                country = pycountry.countries.get(alpha_2=country_name.upper())
                return True
            except KeyError:
                return False

    ############################################
    ###         Personal Information         ###
    ############################################
    def validatePersonalInformation(self):
        result = ''
        query = ''
        pinfo = self.data['Personal Information']
        query = 'Name string is ' + pinfo['Name'] + ". Do you know if this name format is correct? If not correct, write only reason. If correct, write 'correct' string"
        query = """
                Full Name string is \"""" + pinfo['Name'] + """\" . Do you know if format of this name is correct? 
                If not correct, write only reason. If  correct, print just only string 'correct' simply. No need any description. Name format is \"First name, 
                [Middle name],Last name\". Middle name can be contained or not. Check only format, Don't care about spelling."""
        
        query += """Next item.
                The value of someone's maiden name is \"""" + pinfo['Maiden'] + """\". If this value is empty or "", not correct. Should be "N/A".Do you know if 
                format of this value is correct? Check only format. Don't care about spelling. If correct, print just only string "correct" simply.
                 No need any description. If not correct, write only reason."""
        
        query += """Next item
                Address value is '""" + pinfo['Address'] + """', City value is '""" + pinfo['City'] + """', State value is '""" + pinfo['State'] + """' and then 
                zipcode is '""" + pinfo['Zip'] + """'. Do you know if these values are correct? If correct, print just only string 'correct' simply. No need any 
                description. But if not correct, print just only reason.
                """
        
        query += """Next item

                """
        
        answer_string = self.getAnswerFromAI(query)
        print (answer_string)
        pams = answer_string.split ("\n")
        self.doc.add_heading ('Personal Information', level = 1)
        i = 0
        for pam in pams:
            if pam.strip() == "": continue
            if i == 0:
                p = self.doc.add_paragraph()
                p.add_run("Name: ").bold = True
                p.add_run(pam)
            if i == 1:
                p = self.doc.add_paragraph()
                p.add_run("Maiden: ").bold = True
                p.add_run(pam)
            if i == 2:
                p = self.doc.add_paragraph()
                p.add_run("Address, City, State and Zipcode: ").bold = True
                p.add_run(pam)

                p = self.doc.add_paragraph()
                p.add_run("Country of Issuance: ").bold = True
                if self.validateCountry(pinfo['Country of Issuance']) == False: p.add_run("This value is incorrect.")
                else: p.add_run ("correct.")

                p = self.doc.add_paragraph()
                p.add_run("Country of Legal Residence: ").bold = True
                if self.validateCountry(pinfo['Country of Issuance']) == False: p.add_run("This value is incorrect.")
                else: p.add_run ("correct.")
                
                p = self.doc.add_paragraph()
                p.add_run("Passport: ").bold = True
                if self.validatePassport(pinfo['Passport Number']) == False: p.add_run("Above country name is incorrect.")
                else: p.add_run ("correct.")

                p = self.doc.add_paragraph()
                p.add_run("Aliases: ").bold = True
                if pinfo['Aliases'].strip() == "": p.add_run("This value should be 'N/A', not empty.")
                else: p.add_run ("correct.")

                p = self.doc.add_paragraph()
                p.add_run("Home Phone: ").bold = True
                if self.checkFormat(self.patternPhoneNumber, pinfo['Home Phone']) == False: p.add_run("Phone number format is not correct, Please input again.")
                else: p.add_run ("correct.")
                
                p = self.doc.add_paragraph()
                p.add_run("Email Address: ").bold = True
                if self.checkFormat(self.patternEmail, pinfo['Email Address']) == False: p.add_run("Phone number format is not correct, Please input again.")
                else: p.add_run ("correct.")
            i = i + 1

    def validateAddressHistory(self):
        query = ""
        ah = self.data['Address History']
        self.doc.add_heading ('Address History', level = 1)
        
        permanent = ah['Permanent']
        p = self.doc.add_paragraph ("Permanent", style='List Bullet')
        if(self.checkFormat(self.patternAddress, permanent['Address']) == False):
            self.doc.add_paragraph ("\t" + permanent['Address'] + ":" + "Permanent address format is incorrect. Please have a look on that and fix if it's possible.")
        else:
            self.doc.add_paragraph ("\t" + permanent['Address'] + " format:" + "Correct.")

        if(self.isValidAddress(permanent['Address']) == False):
            self.doc.add_paragraph ("\t" + permanent['Address'] + " reality:" + "Permanent address you input is invalid. Please try again to input it.")
        else:
            self.doc.add_paragraph ("Correct.")

        if(self.validateCityStateZip(permanent['City'], permanent['State'], permanent['Zip']) == False):
            self.doc.add_paragraph ("\t" + permanent['City'] + "," + permanent['State'] + "," + permanent['Zip'] + ": " + 
                        "Not matched. Please have a look on them.")
        else:
            self.doc.add_paragraph ("\t" + permanent['City'] + "," + permanent['State'] + "," + permanent['Zip'] + ": " +"Correct.")
        
        # # History
        history = ah['History']
        total_days = timedelta(days = 0)

        for item in history:
            if(item['To'] != 'Present'):
                to_date = datetime.strptime(item['To'], '%m/%d/%Y')
            else:
                to_date = datetime.today()

            from_date = datetime.strptime(item['From'], '%m/%d/%Y')

            total_days += to_date - from_date

            p = self.doc.add_paragraph (str(from_date) + " To " + str(to_date), style='List Bullet')
            if(self.isValidAddress(item['Address']) == False):
                p = self.doc.add_paragraph (item['Address'] + ": " + "Invalid.")
                if len(self.spellCheck(item['Address'])) > 0:
                    p.add_run ("Incorrect spelling - " + ", ".join(self.spellCheck(item['Address'])))
            else:
                self.doc.add_paragraph (item['Address'] + ": " + "Correct.")

            if(self.validateCityStateZip(item['City'], item['State'], item['Zip']) == False):
                self.doc.add_paragraph (item['City'] + "," + item['State'] + "," + item['Zip'] + ": " + "Not matched. Please have a look on them.")
            else:
                self.doc.add_paragraph (item['City'] + "," + item['State'] + "," + item['Zip'] + ": " + "Correct.")
        
        # # History
        to_date = datetime.today()
        from_date = datetime.strptime(history[len(history) - 1]['From'], '%m/%d/%Y')
        if total_days.days != (to_date-from_date).days:
            p = self.doc.add_paragraph()
            p.add_run ("There is overlaps or gaps in Address History. Please take a look on it and fix it.").bold = True

    def validateEducationHistory(self):
        self.doc.add_heading ('Education History', level = 1)
        educaton_history = self.data["Education History"]
        edu_summary = educaton_history["Summary"]
        edu_history = educaton_history["History"]
        edu_achievements = educaton_history["Achievements"]

        total_days = timedelta(days = 0)
        for item in edu_history:
            # print("\n\n",item,"\n\n")
            if item['type'].find("University")>=0 or item['type'].find("College")>=0:
                from_date = datetime.strptime(item['From'], '%m/%d/%Y')
                to_date = datetime.strptime(item['To'], '%m/%d/%Y')
                delta = to_date - from_date
                total_days+=delta
        if total_days.days < int(edu_summary["YearOfCollege"])*365:
            p = self.doc.add_paragraph ()
            p.add_run ("Years of College: ").bold = True
            p.add_run ("Total years calculated of College and high education is incorrect.")
        if len(self.spellCheck(edu_summary['Degree'])) > 0:
            p = self.doc.add_paragraph ()
            p.add_run ("Degree: ").bold = True
            p.add_run ("Spelling incorrect. Incorrect spelling word - " + ", ".join(self.spellCheck(edu_summary['Degree'])))
        if not edu_summary["FluentInEnglish"].capitalize() in ["Yes", "No"]:
            p = self.doc.add_paragraph ()
            p.add_run ("Fluent in English: ").bold = True
            p.add_run ("Incorrect. Should be 'Yes' or 'No'")
        if edu_summary["OtherLanguages"].strip() == "":
            p = self.doc.add_paragraph ()
            p.add_run ("Other Languages: ").bold = True
            p.add_run ("Incorrect. Should be N/A if you don't know other languages.")
        if edu_summary["OtherLanguages"].find("English"):
            p = self.doc.add_paragraph ()
            p.add_run ("If you know only English, Other Langauges field should be set as N/A.")

        highschool_count = 0
        i  =  0
        for item in edu_history:
            self.doc.add_heading (item['type'] + ' ', level = 3)
            i += 1
            text=""
            if len(self.spellCheck(item["School"])) > 0:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("School: " + item["School"] + " ").bold = True
                p.add_run ("is incorrect in spelling. Incorrect spelling word - " + ", ".join(self.spellCheck(item["School"])))
            if len (self.spellCheck(item["Program"])) > 0:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Program: " + item["Program"] + " ").bold = True
                p.add_run ("is incorrect in spelling. Incorrect spelling word - " + ", ".join(self.spellCheck(item["Program"])))
            full_addr = item['Address'] + "," + item['City'] + "," + item['State']
            if not self.isValidAddress(full_addr):
                p = self.doc.add_paragraph ("\t")
                p.add_run (full_addr + ": ").bold = True
                p.add_run ("is invalid.")
            if item["Graduate"].capitalize() not in ["Yes", "No"]:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Graduate: ").bold = True
                p.add_run ("Should be 'Yes' or 'No'.")
            if item['type'].find("High School")==-1:
                if item["GPA"].strip()=="":
                    p = self.doc.add_paragraph ("\t")
                    p.add_run (item["GPA"] + ": ").bold = True
                    p.add_run ("GPA field shoulb be actual number like 3.5 or N/A.")
                elif item["GPA"]!="N/A":
                    try:
                        float(item["GPA"])
                    except ValueError:
                        p = self.doc.add_paragraph ("\t")
                        p.add_run (item["GPA"] + ": ").bold = True
                        p.add_run ("GPA field shoulb be actual number like 3.5 or N/A.")
            if item['type'].find("High School")>=0:
                highschool_count+=1
                if item['Program'].strip()!="High School Diploma":
                    p = self.doc.add_paragraph ("\t")
                    p.add_run (item["Program"] + ": ").bold = True
                    p.add_run ("We recommend to set Program field of High School as High School Diploma.")
                if item['Graduate'].strip().capitalize() not in ["Yes", "No"]:
                    p = self.doc.add_paragraph ("\t")
                    p.add_run (item["Graduate"] + ": ").bold = True
                    p.add_run ("Should be 'Yes' or 'No'.")
                try:
                    float(item["GPA"])
                except ValueError:
                    p = self.doc.add_paragraph ("\t")
                    p.add_run (item["GPA"] + ": ").bold = True
                    p.add_run ("GPA field shoulb be actual number like 3.5 or N/A.")
            
        if highschool_count==0:
            p = self.doc.add_paragraph ()
            p.add_run ("All applications must have High School.").bold = True

    def validateDriversRecord(self):
        self.doc.add_heading ('Drivers Record', level = 1)

        drivers_record = self.data["Drivers Record"]
        dr_summary = drivers_record["Summary"]
        dr_violations = drivers_record["Violations"]

        if dr_summary["License"]=="":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("License: ").bold = True
            p.add_run ("Should not blank.")
        if dr_summary["State"] =="":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("State: ").bold = True
            p.add_run ("Should not blank.")
        from_date = datetime.today()
        to_date = datetime.strptime(dr_summary["Expires"], '%m/%d/%Y')
        total_days = timedelta(days = 0)
        total_days += to_date - from_date
        if total_days.days < 90:
            p = self.doc.add_paragraph ()
            p.add_run ("Caution! Driver License will be expired within 90 days.").bold = True
        for item in dr_violations:
            err = False
            r_text = ""
            if len(self.spellCheck(item["Violation"])) > 0:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Violation: " + item["Violation"] + " ").bold = True
                p.add_run ("is incorrect in spelling. Incorrect spelling word - " + ", ".join(item["Violation"]))
            if (datetime.strptime(item["Date"], '%m/%d/%Y')-datetime.now()).total_seconds()>0:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Violation date is not past date").bold = True
            if len(self.spellCheck(item["City"])) > 0:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("City: " + item["City"] + " ").bold = True
                p.add_run ("is incorrect in spelling. Incorrect spelling word - " + ", ".join(item["City"]))
            if len(self.spellCheck(item["County"])) > 0:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("County: " + item["County"] + " ").bold = True
                p.add_run ("is incorrect in spelling. Incorrect spelling word - " + ", ".join(item["County"]))
            if len(self.spellCheck(item["Disposition"])) > 0:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Disposition: " + item["Disposition"] + " ").bold = True
                p.add_run ("is incorrect in spelling. Incorrect spelling word - " + ", ".join(item["Disposition"]))

    def validateCriminalRecord(self):
        self.doc.add_heading ('Criminal Record', level = 1)
        cr = self.data['Criminal Record']
        
        if cr["Driving while Impaired"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Driving While Impaired: " + cr["Driving while Impaired"] + " -> ").bold = True
            p.add_run ("Should be 'Yes' or 'No'.")
        
        if cr["Under the Influence"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Under the Influence: " + cr["Under the Influence"] + " -> ").bold = True
            p.add_run ("Should be 'Yes' or 'No'.")
        
        if cr["Driving While Intoxicated"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Driving While Intoxicated: " + cr["Driving While Intoxicated"] + " -> ").bold = True
            p.add_run ("Should be 'Yes' or 'No'.")
                
        if cr["License Suspended"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("License Suspended: " + cr["License Suspended"] + " -> ").bold = True
            p.add_run ("Should be 'Yes' or 'No'.")

        if cr["License Revoked"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("License Revoked: " + cr["License Revoked"] + " -> ").bold = True
            p.add_run ("Should be 'Yes' or 'No'.")

        if cr["Additional Details"].strip() == "":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Additional Details: ").bold = True
            p.add_run ("If you have no additinal details information, this field value should be 'N/A', not empty")

        if cr["Past 10years Criminal"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Past 10years Criminal: " + cr["Past 10years Criminal"] + " -> ").bold = True
            p.add_run ("Should be 'Yes' or 'No'.")

        if cr["IfYes"].strip() == "":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("In the past ten (10) years, have you ever been convicted of a crime involving a felony, misdemeanor, infraction, or violation of any law? Please exclude all traffic violations.: ").bold = True
            p.add_run ("If you have no details information, this field value should be 'N/A', not empty")
    
    def validateEmploymentGeneral(self):
        self.doc.add_heading ('Employment General', level = 1)
        cr = {}
        try:
            cr = self.data['Employment General']
        except:
            return
        
        legal_to_work = cr["LegalToWork"]
        able_to_relocate = cr["AbleToRelocate"]
        contact_present = cr["ContactPresent"]
        contact_previous = cr["ContactPrevious"]
        ever_discharged = cr["EverDischarged"]
        details = cr["Details"]

        if legal_to_work.strip().capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Legal to work in U.S.: " + cr["LegalToWork"] + " -> ").bold = True
            p.add_run ("Should be 'Yes' or 'No'.")
        if able_to_relocate.strip().capitalize() == "YES" or able_to_relocate.strip().capitalize() == "NO":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Able to Relocate: " + cr["AbleToRelocate"] + " -> ").bold = True
            p.add_run ("Should be 'Yes' or 'No'.")
        if contact_present.strip().capitalize() == "YES" or contact_present.strip().capitalize() == "NO":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Contact present employer: " + cr["ContactPresent"] + " -> ").bold = True
            p.add_run ("Should be 'Yes' or 'No'.")
        if contact_previous.strip().capitalize() == "YES" or contact_previous.strip().capitalize() == "NO":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Contact previous Employer: " + cr["ContactPrevious"] + " -> ").bold = True
            p.add_run ("Should be 'Yes' or 'No'.")
        if details.strip() == "":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Details: ").bold = True
            p.add_run ("If you have no details information, this field value should be 'N/A', not empty")
        

    def validateEmploymentPresent(self):
        self.doc.add_heading ('Employment Present', level = 1)

        employment_present = self.data["Employment Present"]
        if employment_present["To"].strip() != "Present":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("To: ").bold = True
            p.add_run ("To date is not 'Present'.")
        elif (datetime.strptime(employment_present["From"], '%m/%d/%Y')-datetime.now()).total_seconds()>0:
            p = self.doc.add_paragraph ("\t")
            p.add_run (employment_present["From"] + " To Present: ").bold = True
            p.add_run ("From date is not past date.")
        if len(self.spellCheck(employment_present["Company"])) > 0:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Company: ").bold = True
            p.add_run ("Spelling incorrect. Incorrect spelling word - " + ", ".join(self.spellCheck(employment_present["Company"])))
        full_addr = employment_present['Address'] + "," + employment_present['City'] + "," + employment_present['State'] + " " + employment_present['Zip']
        if not self.isValidAddress(full_addr):
            p = self.doc.add_paragraph ("\t")
            p.add_run (full_addr + ": ").bold = True
            p.add_run ("is invalid.")
        if employment_present["Position"] == "":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Position: ").bold = True
            p.add_run ("Should not blank. Identify seat position. For example, 'First Officer' or 'Captain'.")
        if employment_present["Position"] == "Pilot":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Position: ").bold = True
            p.add_run ("Please identify Seat position by first officer or captain.")
        p = self.doc.add_paragraph ("\t")
        p.add_run ("Duties: ").bold = True
        if len(employment_present["Duties"]) < 400:
            p.add_run ("Less than 400 characters. Recommend using all 500 characters.\n\t")
        query = 'Please check the following sentence. The sentence is "' + employment_present["Duties"] + '". Please check only spelling and grammar.  If correct, print just only string "Correct in spelling and grammar." simply. If not correct, print just only reason.'
        answer_string = self.getAnswerFromAI(query)
        p.add_run (answer_string + "\n\t")
        if employment_present["Duties"].strip()[-1] != ".":
            p.add_run ("Duties should be cut off.")
        if employment_present["AC Flown"] == "":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("AC Flown: ").bold = True
            p.add_run ("Should not be blank, Please fill with N/A.")
        if employment_present["Hours per Month"] in ["","0"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Hours per Month: ").bold = True
            p.add_run ("Invalid, if it's not specified, fill with N/A.")
        if employment_present["Supervisor"]=="":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Supervisor: ").bold = True
            p.add_run ("Should not be blank, Please fill with N/A.")
        if(self.checkFormat(self.patternPhoneNumber, employment_present['Phone']) == False):
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Phone: ").bold = True
            p.add_run ("Phone number format is not correct, Please input again.")
        if len (self.spellCheck(employment_present["Reason for Leaving"])) > 0:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Reason for Leaving: ").bold = True
            p.add_run ("Spelling incorrect. Incorrect spelling word - " + ', '.join(self.spellCheck(employment_present["Reason for Leaving"])))
        if employment_present["Reason for Leaving"] == "":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Reason for Leaving: ").bold = True
            p.add_run ("Should not be blank, Please fill with N/A.")

    def validateEmploymentHistory(self):
        self.doc.add_heading ('Employment - History', level = 1)
        ehs = self.data ['Employment History']
        for eh in ehs:
            self.doc.add_heading (eh["From"] + " To " + eh["To"] + ": ", level = 4)
            if (datetime.strptime(eh["From"], '%m/%d/%Y')-datetime.strptime(eh["From"], '%m/%d/%Y')).total_seconds()>0:
                p = self.doc.add_paragraph ("\t")
                p.add_run (eh["From"] + " To " + eh["To"] + ": ").bold = True
                p.add_run ("From date is not past date.")
            if len(self.spellCheck(eh["Company"])) > 0:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Company: ").bold = True
                p.add_run ("Spelling incorrect. Incorrect spelling word - " + ", ".join(self.spellCheck(eh["Company"])))
            full_addr = eh['Address'] + "," + eh['City'] + "," + eh['State'] + " " + eh['Zip']
            if not self.isValidAddress(full_addr):
                p = self.doc.add_paragraph ("\t")
                p.add_run (full_addr + ": ").bold = True
                p.add_run ("is invalid.")
            if eh["Position"] == "":
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Position: ").bold = True
                p.add_run ("Should not blank. Identify seat position. For example, 'First Officer' or 'Captain'.")
            if eh["Position"] == "Pilot":
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Position: ").bold = True
                p.add_run ("Please identify Seat position by first officer or captain.")
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Duties: ").bold = True
            if len(eh["Duties"]) < 400:
                p.add_run ("Less than 400 characters. Recommend using all 500 characters.\n\t")
            query = 'Please check the following sentence. The sentence is "' + eh["Duties"] + '". Please check only spelling and grammar.  If correct, print just only string "Correct in spelling and grammar." simply. If not correct, print just only reason.'
            answer_string = self.getAnswerFromAI(query)
            p.add_run (answer_string + "\n\t")
            if eh["Duties"].strip()[-1] != ".":
                p.add_run ("Duties should be cut off.")
            if eh["ACFlown"] == "":
                p = self.doc.add_paragraph ("\t")
                p.add_run ("AC Flown: ").bold = True
                p.add_run ("Should not be blank, Please fill with N/A.")
            if eh["HoursPerMonth"] in ["","0"]:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Hours per Month: ").bold = True
                p.add_run ("Invalid, if it's not specified, fill with N/A.")
            if eh["Supervisor"]=="":
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Supervisor: ").bold = True
                p.add_run ("Should not be blank, Please fill with N/A.")
            if(self.checkFormat(self.patternPhoneNumber, eh['Phone']) == False):
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Phone: ").bold = True
                p.add_run ("Phone number format is not correct, Please input again.")
            if len (self.spellCheck(eh["ReasonForLeaving"])) > 0:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Reason for Leaving: ").bold = True
                p.add_run ("Spelling incorrect. Incorrect spelling word - " + ', '.join(self.spellCheck(eh["ReasonForLeaving"])))
            if eh["ReasonForLeaving"] == "":
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Reason for Leaving: ").bold = True
                p.add_run ("Should not be blank, Please fill with N/A.")


    def validateUnemploymentFurlough(self):
        self.doc.add_heading ('Unemployment / Furlough', level = 1)

        unemployment_furlough = self.data["UnemploymentFurlough"]
        emfu_history = unemployment_furlough["History"]
        emfu_details = unemployment_furlough["Details"]

        for item in emfu_history:
            n=1
            self.doc.add_heading ('Item ' + str(n), level = 4)
            d_from = item["From"]
            d_type = item["Type"]
            d_desc = item["Description"]
            if d_type == "":
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Type: ").bold = True
                p.add_run ("Blank, please insert correct type.")
            if len(self.spellCheck(d_desc)) > 0:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Description: ").bold = True
                p.add_run ("Spelling incorrect. Incorrect spelling word - " + ", ".join(self.spellCheck(d_desc)))
            n=n+1

        if emfu_details in ["N/A",""]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Unemployment Details: ").bold = True
            if len(emfu_history)!=0:
                p.add_run ("Unemployment Details can be N/A only if there is no unemployment.")
            else:
                p.add_run (f"Please remove {emfu_details} Unemployment Details and insert summary of all periods unemployments.")

    def validateEmploymentMisc(self):
        self.doc.add_heading ('Employment - Misc', level = 1)
        em = self.data['Employment Misc']
        query = "\"" + em['Professional Memberships'] + "\"" + """\nThis is professional memberships. Is this invalid or correct spell? 
                Write the answer. Not write unnecessary description. Your answer string must be  'correct' or reason.If considered as correct, 
                you must print only 'correct' here. But if not correct, write the reason."""
        query += """\n Next item""" + em['Achievements and Awards'] + """
                This is Achievements and Awards string. Is this invalid? or Correct spell? Write the answer. Your answer string must 
                be 'correct' or reason If correct, you must print only 'correct' here. If not correct, write the reason.
                """
        query += """\n Next item""" + em['Volunteer Charity Work'] + """
                This is Volunteer Charity Work string. Is this invalid? or Correct spell? Write the answer. Your answer string must 
                be 'correct' or reason If correct, you must print only 'correct' here. If not correct, write the reason.
                """
        answer_string = self.getAnswerFromAI(query)
        pams = answer_string.split ('\n')
        i = 0
        for pam in pams:
            if pam.strip() == "": continue
            if i == 0:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Professional Memberships: ").bold = True
                p.add_run (pam)
            if i == 1:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Achievements and Awards: ").bold = True
                p.add_run (pam)
            if i == 2:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Volunteer Charity Work: ").bold = True
                p.add_run (pam)
            i += 1

    def validatePilotExperienceGeneral(self):
        self.doc.add_heading ('Pilot Experience - General', level = 1)
        peg = self.data['Pilot Experience General']
        for key in peg.keys():
            if not peg[key] in ["Yes", "No"]:
                p = self.doc.add_paragraph ("\t")
                p.add_run (key + ": ").bold = True
                p.add_run ("Should be 'Yes' or 'No'")

    def validatePilotCertificateRatings(self):
        self.doc.add_heading ('Pilot & FE Certificates and Ratings', level = 1)
        pcr = self.data['Pilot Certificates Ratings']

        if "" in list(pcr.values()):
            p = self.doc.add_paragraph ("\t")
            p.add_run ("No blanks allowed for entire sections").bold = True

        if not pcr["Cert. Number"].isdigit():
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Cert. Number: ").bold = True
            p.add_run ("Should consist of only numbers")
        if (datetime.strptime(pcr["Issue Date"], '%m/%d/%Y')-datetime.now()).total_seconds()>0:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Issue Date: ").bold = True
            p.add_run ("Issue Date is not past.")
        if not pcr["Flight Engineer Cert. Number"] in ["","N/A"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Flight Engineer Cert. Number: ").bold = True
            p.add_run ("Should not be blank.")
        if pcr["Flight Engineer Issue Date"] != "":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Flight Engineer Issue Date: ").bold = True
            p.add_run ("Should be blank.")
        if pcr["Flight Engineer FE Turbojet"] == "":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Flight Engineer FE Turbojet: ").bold = True
            p.add_run ("Blank, Please insert No in that case.")
        if pcr["Flight Engineer FE Reciprocating"] == "":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Flight Engineer FE Reciprocating: ").bold = True
            p.add_run ("Blank, Please insert No in that case.")
        if pcr["Flight Engineer FE Turboprop"] == "":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Flight Engineer FE Turboprop: ").bold = True
            p.add_run ("Blank, Please insert No in that case.")
        if pcr["Airplane SEL"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Airplane SEL: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")
            if pcr["Airplane SEL"] == "No": p.add_run ("Please check if the information is accurate.").blod = True
        if pcr["Airplane MEL"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Airplane MEL: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")
            if pcr["Airplane MEL"] == "No": p.add_run ("Please check if the information is accurate.").blod = True
        if pcr["Airplane SES"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Airplane SES: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")
        if pcr["Airplane MES"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Airplane MES: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")

        if pcr["Rotor Helicopter"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Rotor Helicopter: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")
        if pcr["Rotor Gyroplane"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Rotor Gyroplane: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")
        if pcr["Airship"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Airship: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")
        if pcr["Balloon"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Balloon: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")
        if pcr["Powered Lift"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Powered Lift: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")
        if pcr["Glider"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Glider: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")

        if pcr["Turbojet Typed"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Turbojet Typed: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")

        if pcr["B-737 Typed"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("B-737 Typed: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")

        if pcr["Large Aircraft Typed"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Large Aircraft Typed: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")

        if pcr["Instrument Airplane"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Instrument Airplane: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")

        if pcr["Instrument Airplane"] == "No":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Instrument Airplane: ").bold = True
            p.add_run ("Caution! Please make sure Instrument Airplane answer is correct")

        if pcr["Instrument Helicopter"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Instrument Helicopter: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")

        if pcr["Instrument Powered Lift"].capitalize() not in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Instrument Powered Lift: ").bold = True
            p.add_run ("Should be 'Yes' or 'No'")
    
    def validateInstructorCertificatesRatings(self):
        self.doc.add_heading ('Instructor Certificates and Ratings', level = 1)
        icr = self.data['Instructor Certificates Ratings']
        
        if icr["Flight Instructor"]=="":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Flight Instructor: ").bold = True
            p.add_run ("Should be N/A if blank")
        elif not icr["Flight Instructor"]=="N/A" and not icr["Flight Instructor"].isdigit():
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Flight Instructor: ").bold = True
            p.add_run ("Should be only numbers")
        if icr["Flight Instructor Issue Date"]!="":
            if (datetime.strptime(icr["Flight Instructor Issue Date"], '%m/%d/%Y')-datetime.now()).total_seconds()>0:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Flight Instructor Issue Date: ").bold = True
                p.add_run ("This value is the future date")

        if icr["Ground School"]=="":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Ground School: ").bold = True
            p.add_run ("Should be N/A if blank")
        elif not icr["Ground School"]=="N/A" and not icr["Ground School"].isdigit():
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Ground School: ").bold = True
            p.add_run ("Should be only numbers")

        if icr["Ground School Issue Date"]!="":
            if (datetime.strptime(icr["Ground School Issue Date"], '%m/%d/%Y')-datetime.now()).total_seconds()>0:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Ground School Issue Date: ").bold = True
                p.add_run ("This value is the future date")
        for key in icr.keys():
            if "Flight Instructor" in key or "Ground School" in key:
                continue
            if icr[key] == "":
                p = self.doc.add_paragraph ("\t")
                p.add_run (key + ": ").bold = True
                p.add_run ("Blank, please insert the field.")
            elif icr[key] not in ["Yes", "No"]:
                p = self.doc.add_paragraph ("\t")
                p.add_run (key + ": ").bold = True
                p.add_run ("Should be 'Yes' or 'No'.")

    def validateFAAWrittenTests(self):
        self.doc.add_heading ('FAA Written Tests', level = 1)
        fwt = self.data['FAA Written Tests']
        if (fwt["ATPDate"].strip() != "" and datetime.strptime(fwt["ATPDate"], '%m/%d/%Y')-datetime.now()).total_seconds()>0:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("ATP Date: ").bold = True
            p.add_run ("This value is the future date.")
        if fwt["FETurbojetDate"].strip() != "" :
            if (datetime.strptime(fwt["FETurbojetDate"], '%m/%d/%Y')-datetime.now()).total_seconds()>0:
                p = self.doc.add_paragraph ("\t")
                p.add_run ("FE Turboprop Date: ").bold = True
                p.add_run ("This value is the future date.")

        for key in fwt.keys():
            # print(key)
            if "Current" in key:
                if not fwt[key] in ["Yes", "No"]:
                    p = self.doc.add_paragraph ("\t")
                    p.add_run ("Current: ").bold = True
                    p.add_run ("This value should be Yes or No, please insert correct answer.")

    def validateMiscCertificates(self):
        self.doc.add_heading ('Misc Certificates', level = 1)
        mc = self.data['Misc Certificates']
        if not mc["Dispatcher"] in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Dispatcher: ").bold = True
            p.add_run ("Dispatcher should be Yes or No.")
        if not mc["Airframe&Powerplant"] in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Airframe & Powerplant: ").bold = True
            p.add_run ("Should be Yes or No.")
        if not mc["FCCPermit"] in ["Yes", "No"]:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("FCC Permit: ").bold = True
            p.add_run ("Should be Yes or No.")

    def validateFAAMedicals(self):
        self.doc.add_heading ('FAA Medicals', level = 1)
        fm = self.data['FAA Medicals']

        if fm["Class"]=="":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Class: ").bold = True
            p.add_run ("Make sure Class information is accurate. If you do not hold a First Class Medical we highly recommend obtaining it before publishing your application.")
        total_days = timedelta(days = 0)
        total_days += datetime.strptime(fm["Issued"], '%m/%d/%Y')-datetime.now()
        if total_days.days < 0:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Issued: ").bold = True
            p.add_run ("Issued Date was expired.")
        elif total_days.days() < 90:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Issued: ").bold = True
            p.add_run ("Your medical will be expired within the next 90 days.")
        if fm["Restrictions"] == "":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Restrictions: ").bold = True
            p.add_run ("If you do not have any restriction plan insert None and if you have any restriction place restriction that are on your medical in Restrictions section.")

    def validateFAAActions(self):
        self.doc.add_heading ('FAA Actions', level = 1)
        fa = self.data["FAA Actions"]
        
        for key in fa.keys():
            if key != "Details":
                if not fa[key] in ["Yes", "No"]:
                    p = self.doc.add_paragraph ("\t")
                    p.add_run (key + ": ").bold = True
                    p.add_run (f"{key} section should be Yes or No.")
        if fa["Details"] == "" or fa["Details"] == "N/A":
            if "Yes" in fa.values():
                p = self.doc.add_paragraph ("\t")
                p.add_run ("Details: ").bold = True
                p.add_run ("Details should be provided.")
            else:
                if fa["Details"] == "":
                    p = self.doc.add_paragraph ("\t")
                    p.add_run ("Details: ").bold = True
                    p.add_run ("Details should be N/A if blank.")

    def validateAircraftFlown(self):
        self.doc.add_heading ('Aircraft Flown', level = 1)
        af = self.data["Aircraft Flown"]
        for item in af:
            model = item["Model"]
            lastflown = item["LastFlown"]
            total_days = timedelta(days = 0)
            total_days += datetime.now()-datetime.strptime(lastflown, '%m/%Y')
            if total_days.days > 90:
                p = self.doc.add_paragraph ("\t")
                p.add_run (model + ": ").bold = True
                p.add_run (f"Caution for {model}! Please make sure your last flown date is correct. The information you provided is indicating that you have not flown within the last 90days.")

    def validateFlightTimeByConditions(self):
        self.doc.add_heading ('Flight Time By Conditions', level = 1)
        addendum = self.data["Addendum"]
        for item in addendum:

            question = item["question"]
            answer = item["answer"]
            if "flight time have you logged in" in question or "If yes" in question:
                if answer=="":
                    p = self.doc.add_paragraph ("\t")
                    p.add_run (question + ": ").bold = True
                    p.add_run ("You have to insert explanation for this question.")
            else:
                if not answer in ["Yes", "No"]:
                    p = self.doc.add_paragraph ("\t")
                    p.add_run (question + ": ").bold = True
                    p.add_run ("Answer to this question should be Yes or No.")
                else:
                    if "convicted of any felony" in question and answer=="No":
                        p = self.doc.add_paragraph ("\t")
                        p.add_run (question + ": ").bold = True
                        p.add_run ("Please make sure answer for this question is correct. This goes beyond the previous ten years period. This only applies to convictions.")
                    if "ever failed ANY checkrides" in question and answer == "No":
                        p = self.doc.add_paragraph ("\t")
                        p.add_run (question + ": ").bold = True
                        p.add_run ("Please make sure answer for this question is accurate. This includes any Part 61/141 stage checks, UPT stage checks/check rides or failures to include Form 8 Checkrides.")
                    if "I acknowledge that I have" in question and answer == "No":
                        p = self.doc.add_paragraph ("\t")
                        p.add_run (question + ": ").bold = True
                        p.add_run ("Please make sure answer for this question is correct.")

    def validateGeneralReferences (self):
        self.doc.add_heading ('General References', level = 1)
        gr = self.data["GeneralReferences"]
        if len(gr)<3:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("You must include 3 entries at a minimum.").bold = True
        span_days = []
        span_10_years = False

        for item in gr:
            self.doc.add_heading (item["Name"], level = 4)
            from_date = item["From"]
            to_date = item["To"]
            phone_number = item["PhoneNumber"]
            email_address = item["EmailAddress"]
            total_days = timedelta(days = 0)
            if "Present" in to_date:
                total_days += datetime.now()-datetime.strptime(from_date, '%m/%d/%Y')
            else:
                total_days += datetime.strptime(to_date, '%m/%d/%Y')-datetime.strptime(from_date, '%m/%d/%Y')
            span_days.append(total_days.days)
            if total_days.days>3650:
                span_10_years = True
            if self.checkFormat(self.patternPhoneNumber, phone_number) == False:
                p = self.doc.add_paragraph ("\t")
                p.add_run (f"Phone Number: {phone_number} - ").bold = True
                p.add_run ("Incorrect format. Please insert write format e.g. 601-209-5505.")
            if self.checkFormat(self.patternEmail, email_address) == False:
                p = self.doc.add_paragraph ("\t")
                p.add_run (f"Email address: {email_address} - ").bold = True
                p.add_run ("Incorrect format. Please insert write format e.g. abc@test.com.")
        if span_10_years:
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Caution! You need at least one reference that has known you for at least 10 years.").bold = True

    def validateTSRD(self):
        self.doc.add_heading ('Transportation Security Regulation Disclosure', level = 1)
        da = self.data["Disclosure Answer"]
        if da == "Yes":
            p = self.doc.add_paragraph ("\t")
            p.add_run ("Disclosure Answer: ").bold = True
            p.add_run ("Caution! Please check that your 'Yes' answer is correct.")

    def outputFile (self):
        self.validatePersonalInformation()
        self.validateAddressHistory()
        self.validateEducationHistory()
        self.validateDriversRecord()
        self.validateCriminalRecord()
        self.validateEmploymentGeneral()
        self.validateEmploymentPresent()
        self.validateEmploymentHistory()
        self.validateUnemploymentFurlough()
        self.validateEmploymentMisc ()
        self.validatePilotExperienceGeneral()
        self.validatePilotCertificateRatings ()
        self.validateInstructorCertificatesRatings ()
        self.validateFAAWrittenTests ()
        self.validateFAAMedicals ()
        self.validateFAAActions ()
        self.validateMiscCertificates ()
        self.validateAircraftFlown ()
        self.validateFlightTimeByConditions ()
        self.validateGeneralReferences ()
        self.validateTSRD ()
        self.doc.save (self.fileName + ".docx")

if __name__ == "__main__":
    r = Review ('input.json') #orientation='P', unit='mm', format='B5'
    r.outputFile ()