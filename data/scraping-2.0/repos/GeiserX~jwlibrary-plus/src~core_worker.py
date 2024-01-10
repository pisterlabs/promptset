import uuid
import shutil
import pytz
import os
import zipfile
import logging
import requests
import json
from datetime import datetime, timedelta
import sqlite3
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import subprocess
import langchain
from langchain.chat_models import ChatOpenAI
from langchain.cache import SQLiteCache, InMemoryCache, GPTCache
from gptcache import Cache
from gptcache.manager.factory import manager_factory
from gptcache.processor.pre import get_prompt
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder, SystemMessagePromptTemplate, HumanMessagePromptTemplate
from langchain.callbacks import get_openai_callback # TODO

logging.basicConfig(format="%(asctime)s - %(name)s - %(levelname)s - %(message)s", level=logging.INFO)
logger = logging.getLogger(__name__)

def init_gptcache(cache_obj: Cache, llm: str):
    cache_obj.init(
        pre_embedding_func=get_prompt,
        data_manager=manager_factory(manager="sqlite,faiss,local", data_dir=f"dbs/map_cache_{llm}", vector_params={"dimension": "128"}, max_size=100000),
    )

#######################################
### HELPER: DESCRIBE JWLIBRARY FILE ###
#######################################

def describe_jwlibrary(telegram_user):
    logger.info("describe_jwlibrary - Telegram User: {0}".format(telegram_user))
    jwfile = "userBackups/{0}.jwlibrary".format(telegram_user)

    with zipfile.ZipFile(jwfile, 'r') as zip_ref:
            files = zip_ref.namelist()
            zip_ref.extractall("userBackups/{0}/".format(telegram_user))
    
    uploadedDb = "userBackups/{0}/{1}".format(telegram_user, [zipname for zipname in files if zipname.endswith(".db")][0])

    connection = sqlite3.connect(uploadedDb)
    cursor = connection.cursor()
    cursor.execute("SELECT Count(*) FROM Note")
    notesN = cursor.fetchall()[0][0]
    cursor.execute("SELECT Count(*) FROM InputField")
    inputN = cursor.fetchall()[0][0]
    cursor.execute("SELECT Count(*) FROM TagMap")
    tagMaptN = cursor.fetchall()[0][0]
    cursor.execute("SELECT Count(*) FROM Tag")
    tagN = cursor.fetchall()[0][0]
    cursor.execute("SELECT Count(*) FROM Bookmark")
    bookmarkN = cursor.fetchall()[0][0]
    cursor.execute("SELECT LastModified FROM LastModified")
    lastModified = cursor.fetchall()[0][0]
    cursor.execute("SELECT Count(*) FROM UserMark")
    userMarkN = cursor.fetchall()[0][0]
    connection.close()

    shutil.rmtree("userBackups/{0}/".format(telegram_user))

    return notesN, inputN, tagMaptN, tagN, bookmarkN, lastModified, userMarkN

#######################
### EXTRACTING HTML ###
#######################

def w_extract_html(url, get_all):
    logger.info("w_extract_html - URL: {0} - Full Run: {1}".format(url, get_all))

    html = requests.get(url).text
    soup = BeautifulSoup(html, features="html5lib")
    title = soup.find("h1").text
    classArticleId = soup.find("article", {"id" : "article"}).get("class")
    articleId = next(x for x in classArticleId if x.startswith("iss"))[4:] + "00"
    articleN = soup.find("p", {"id":"p1"}).text

    if get_all:
        base_text = soup.find("p", {"id":"p3"}).text
        song = soup.find("p",{"id":"p4"}).text
        summary = soup.find("div", {"id": "footnote1"}).find("p").text
        documentId = soup.find("input", {"name": "docid"}).get("value")
        p_elements = soup.find("div", {"class":"bodyTxt"})
        questions = p_elements.find_all("p", {"id": lambda x: x and x.startswith("q")})
        paragraphs = p_elements.find_all("p", {"id": lambda x: x and x.startswith("p")})

        # Example q_map = {0 : [q1, [p1]], 1 : [q2&3, [p2, p3]]}
        q_map = {}
        i = 0
        for q in questions:
            q_map[i] = [q]
            q_map[i].append([p for p in paragraphs if p.has_attr('data-rel-pid') if p.get('data-rel-pid').strip('[]') in q.get('data-pid')])
            i = i+1
        
        return title, base_text, song, summary, questions, documentId, articleId, q_map
    else:
        return title, articleId, articleN
    
def mwb_extract_html(url, get_all): # TODO
    logger.info("w_extract_html - URL: {0} - Full Run: {1}".format(url, get_all))

    html = requests.get(url).text
    soup = BeautifulSoup(html, features="html5lib")

    for i in soup.find_all("a"):
        em = i.find("em")
        if em:
            if em.get_text() == "lff":
                study = i.get("href")


    study_html = requests.get("https://wol.jw.org" + study).text
    soup_study = BeautifulSoup(study_html, features="html5lib")



    # title = soup.find("h1").text
    # classArticleId = soup.find("article", {"id" : "article"}).get("class")
    # articleId = next(x for x in classArticleId if x.startswith("iss"))[4:] + "00"
    # articleN = soup.find("p", {"id":"p1"}).text

    # if get_all:
    #     base_text = soup.find("p", {"id":"p3"}).text
    #     song = soup.find("p",{"id":"p4"}).text
    #     summary = soup.find("div", {"id": "footnote1"}).find("p").text
    #     documentId = soup.find("input", {"name": "docid"}).get("value")
    #     p_elements = soup.find("div", {"class":"bodyTxt"})
    #     questions = p_elements.find_all("p", {"id": lambda x: x and x.startswith("q")})
    #     paragraphs = p_elements.find_all("p", {"id": lambda x: x and x.startswith("p")})

    #     # Example q_map = {0 : [q1, [p1]], 1 : [q2&3, [p2, p3]]}
    #     q_map = {}
    #     i = 0
    #     for q in questions:
    #         q_map[i] = [q]
    #         q_map[i].append([p for p in paragraphs if p.has_attr('data-rel-pid') if p.get('data-rel-pid').strip('[]') in q.get('data-pid')])
    #         i = i+1
        
    #     return title, base_text, song, summary, questions, documentId, articleId, q_map
    # else:
    #     return title, articleId, articleN

####################
### QUERY OPENAI ###
####################

def w_query_openai(title, base_text, song, summary, q_map, qs_user):
    logger.info("w_query_openai  - Title: {0} - Base Text: {1} - Song: {2} - Summary: {3} - Questions User: {4}".format(title, base_text, song, summary, qs_user))
    #langchain.llm_cache = SQLiteCache(database_path="dbs/langchain.db") # Not working
    #langchain.llm_cache = InMemoryCache()
    langchain.llm_cache = GPTCache(init_gptcache)

    questions = [f"{i}. {question}" for i, question in enumerate(qs_user, start=1) if question]
    questions_text = "\n".join(questions)

    llm = ChatOpenAI(model_name="gpt-3.5-turbo")
    prompt = ChatPromptTemplate.from_messages([ # Make questions modifiable by user input
        SystemMessagePromptTemplate.from_template("""
Eres un asistente que únicamente usa jw.org y las publicaciones de los Testigos de Jehová para mejorar la preparación de reuniones.
Yo estoy preparándome la Atalaya, edición de estudio, de los Testigos de Jehová.
Proveerás información extra proveniente de la literatura disponible en cada uno de los párrafos que te voy a ir mandando en los sucesivos prompts.
La Atalaya de esta semana se titula {0}, se basa en el texto de {1}, cantaremos la '{2}', y el resumen es el siguiente: 
{3}
Para cada pregunta y párrafo o párrafos que te vaya enviando a partir de ahora, responderás en una lista lo siguiente:
{4}
No escribas estas preguntas de nuevo en la respuesta. Separa las respuestas con dos retornos de carro.""".format(title, base_text, song, summary, questions_text)),
        MessagesPlaceholder(variable_name="history"),
        HumanMessagePromptTemplate.from_template("{input}")
    ])

    notes = {}
    i=0
    for q in q_map.values():
        conversation = ConversationChain(llm=llm, verbose=False, memory=ConversationBufferMemory(return_messages=True), prompt=prompt)
        flattened_paragraph = ""
        for p in q[1]:
            flattened_paragraph = flattened_paragraph + p.text
        notes[i] = conversation.predict(input="Pregunta: {0} -- Párrafo(s): {1}".format(q[0].text, flattened_paragraph))
        logger.info("w_query_openai(Note) - Note: {0}".format(notes[i])) # TODO: Reduce logs in the future when everything works stable
        i=i+1
    
    return notes


############################
### WRITE JWLIBRARY FILE ###
############################

def write_jwlibrary(documentId, articleId, title, questions, notes, telegram_user):

    logger.info("write_jwlibrary - Document ID: {0} - Article ID: {1} - Title: {2} - Questions: {3} - Notes: {4} - Telegram User: {5}".format(documentId, articleId, title, questions, notes, telegram_user))
    uploadedJwLibrary = 'userBackups/{0}.jwlibrary'.format(telegram_user)

    os.makedirs("/app/userBackups/{0}".format(telegram_user), exist_ok=True)

    now = datetime.now(pytz.timezone('Europe/Madrid'))
    now_date = now.strftime("%Y-%m-%d")
    hour_minute_second = now.strftime("%H-%M-%S")
    now_iso = now.isoformat("T", "seconds")

    j = '{{"name":"jwlibrary-plus-backup_{0}","creationDate":"{1}","version":1,"type":0,"userDataBackup":{{"lastModifiedDate":"{2}","deviceName":"jwlibrary-plus","databaseName":"userData.db","schemaVersion":8}}}}'.format(now_date, now_date, now_iso)
    manifest = json.loads(j)

    if(os.path.isfile(uploadedJwLibrary)):
        logger.info("Archivo .jwlibrary encontrado")
        with zipfile.ZipFile(uploadedJwLibrary, 'r') as zip_ref:
            files = zip_ref.namelist()
            zip_ref.extractall("userBackups/{0}/".format(telegram_user))
        
        uploadedDb = "userBackups/{0}/{1}".format(telegram_user, [zipname for zipname in files if zipname.endswith(".db")][0])
        manifestUser = "userBackups/{0}/manifest.json".format(telegram_user)

        manifest_file = 'userBackups/{0}/manifest-{0}-{1}.json'.format(telegram_user, now_date)
        with open(manifest_file, 'w') as f:
            json.dump(manifest, f)

        connection = sqlite3.connect(uploadedDb)
        cursor = connection.cursor()
        cursor.execute("SELECT LocationId FROM Location WHERE DocumentId={0}".format(documentId))
        locationId = cursor.fetchall()
        if locationId:
            locationId = locationId[0][0]
        else:
            cursor.execute("SELECT max(LocationId) FROM Location")
            locationId = cursor.fetchall()[0][0] + 1
            cursor.execute("""INSERT INTO Location (LocationId, DocumentId, IssueTagNumber, KeySymbol, MepsLanguage, Type, Title)
            VALUES ({0}, {1}, {2}, "w", 1, 0, "{3}");""".format(locationId, documentId, articleId, title))
        
        cursor.execute("SELECT TagId FROM Tag WHERE Name = 'jwlibrary-plus'")
        tagId = cursor.fetchall()
        if not tagId:
            cursor.execute("SELECT max(TagId) FROM Tag") # There will be always some tag, even on a brand-new install
            tagId = cursor.fetchall()[0][0] + 1
            cursor.execute("INSERT INTO Tag ('TagId', 'Type', 'Name') VALUES ('{0}', '1', 'jwlibrary-plus')".format(tagId))
            tagId +=1
        else:
            tagId = tagId[0][0]
 

        cursor.execute("SELECT * FROM UserMark LIMIT 1")
        nonEmptyUserMark = cursor.fetchall()
        if nonEmptyUserMark:
            cursor.execute("SELECT max(UserMarkId) FROM UserMark")
            userMarkId = cursor.fetchall()[0][0] + 1
        else:
            userMarkId = 1

        cursor.execute("SELECT BlockRangeId FROM BlockRange LIMIT 1")
        nonEmptyBlockRangeId = cursor.fetchall()
        if nonEmptyBlockRangeId:
            cursor.execute("SELECT max(BlockRangeId) FROM BlockRange")
            blockRangeId = cursor.fetchall()[0][0] + 1
        else:
            blockRangeId = 1

        cursor.execute("SELECT * FROM Note LIMIT 1")
        nonEmptyNote = cursor.fetchall()
        if nonEmptyNote:
            cursor.execute("SELECT max(NoteId) FROM Note")
            noteId = cursor.fetchall()[0][0] + 1
        else:
            noteId = 1

        cursor.execute("SELECT * FROM TagMap LIMIT 1")
        nonEmptyTagMap = cursor.fetchall()
        if nonEmptyTagMap:
            cursor.execute("SELECT max(TagMapId) FROM TagMap")
            tagMapId = cursor.fetchall()[0][0] + 1
            cursor.execute("SELECT max(Position) FROM TagMap")
            Position = cursor.fetchall()[0][0] + 1
        else:
            tagMapId = 1
            Position = 0

        for i in notes:
            uuid_value = str(uuid.uuid4())
            uuid_value2 = str(uuid.uuid4())

            cursor.execute("""INSERT INTO UserMark ('UserMarkId', 'ColorIndex', 'LocationId', 'StyleIndex', 'UserMarkGuid', 'Version')
            VALUES ('{0}', '2', '{1}', '0', '{2}', '1');""".format(userMarkId, locationId, uuid_value))
            
            cursor.execute ("""INSERT INTO "BlockRange" ("BlockRangeId", "BlockType", "Identifier", "StartToken", "EndToken", "UserMarkId")
            VALUES ('{0}', '1', '{1}', '0', '{2}', '{3}');""".format(blockRangeId, questions[i].get("data-pid"), questions[i].text.find(".")-1, userMarkId))
            

            cursor.execute("""INSERT INTO Note ("NoteId", "Guid", "UserMarkId", "LocationId", "Title", "Content", "LastModified", "BlockType", "BlockIdentifier") 
            VALUES ('{0}', '{1}', '{2}', '{3}', '{4}', '{5}', '{6}', '1', '{7}');""".format(noteId, uuid_value2, userMarkId, locationId, questions[i].text, notes[i].replace("'", '"'), now_iso, questions[i].get("data-pid")))

            cursor.execute("INSERT INTO TagMap ('TagMapId', 'NoteId', 'TagId', 'Position') VALUES ('{0}', '{1}', '{2}', '{3}')".format(tagMapId, noteId, tagId, Position))
            userMarkId += 1
            blockRangeId += 1
            noteId +=1
            tagMapId += 1
            Position +=1


        cursor.execute("UPDATE LastModified SET LastModified = '{0}'".format(now_iso))

        connection.commit()
        connection.close()

        fileName = "userBackups/{0}/jwlibrary-plus-{1}-{2}.jwlibrary".format(telegram_user, documentId, now_date)
        zf = zipfile.ZipFile(fileName, "w")
        zf.write(uploadedDb, arcname= "userData.db") # TODO
        zf.write(manifest_file, arcname="manifest.json")
        zf.close()

        os.remove(uploadedDb) # Remove all data from the user except the newly generated .jwlibrary file, which will be deleted after being sent
        os.remove(manifest_file)
        os.remove(uploadedJwLibrary)
        os.remove(manifestUser)

    else:

        dbOriginal = "dbs/userData.db.original"
        dbFromUser = "userBackups/{0}/userData-{0}-{1}_{2}.db".format(telegram_user, now_date, hour_minute_second)
        shutil.copyfile(src=dbOriginal, dst=dbFromUser)

        manifest_file = 'userBackups/{0}/manifest-{0}-{1}.json'.format(telegram_user, now_date)
        with open(manifest_file, 'w') as f:
            json.dump(manifest, f)

        connection = sqlite3.connect(dbFromUser)
        cursor = connection.cursor()

        cursor.execute("""INSERT INTO Location (LocationId, DocumentId, IssueTagNumber, KeySymbol, MepsLanguage, Type, Title)
        VALUES (1, {0}, {1}, "w", 1, 0, "{2}");""".format(documentId, articleId, title))

        cursor.execute("INSERT INTO Tag ('TagId', 'Type', 'Name') VALUES ('2', '1', 'jwlibrary-plus')")

        for i in notes:
            uuid_value = str(uuid.uuid4())
            uuid_value2 = str(uuid.uuid4())

            cursor.execute("""INSERT INTO UserMark ('UserMarkId', 'ColorIndex', 'LocationId', 'StyleIndex', 'UserMarkGuid', 'Version')
            VALUES ('{0}', '2', '1', '0', '{1}', '1');""".format(i+1,uuid_value))

            cursor.execute ("""INSERT INTO "BlockRange" ("BlockRangeId", "BlockType", "Identifier", "StartToken", "EndToken", "UserMarkId")
            VALUES ('{0}', '1', '{1}', '0', '{2}', '{3}');""".format(i+1, questions[i].get("data-pid"), questions[i].text.find(".")-1, i+1))

            cursor.execute("""INSERT INTO Note ("NoteId", "Guid", "UserMarkId", "LocationId", "Title", "Content", "LastModified", "BlockType", "BlockIdentifier") 
            VALUES ('{0}', '{1}', '{2}', '1', '{3}', '{4}', '{5}', '1', '{6}');""".format(i+1, uuid_value2, i+1, questions[i].text, notes[i].replace("'", '"'), now_iso, questions[i].get("data-pid")))

            cursor.execute("INSERT INTO TagMap ('TagMapId', 'NoteId', 'TagId', 'Position') VALUES ('{0}', '{1}', '2', '{2}')".format(i+1,i+1,i))

        cursor.execute("UPDATE LastModified SET LastModified = '{0}'".format(now_iso))

        connection.commit()
        connection.close()

        fileName = "userBackups/{0}/jwlibrary-plus-{1}-{2}.jwlibrary".format(telegram_user, documentId, now_date)
        zf = zipfile.ZipFile(fileName, "w")
        zf.write(dbFromUser, arcname= "userData.db")
        zf.write(manifest_file, arcname="manifest.json")
        zf.close()

        os.remove(dbFromUser)    
        os.remove(manifest_file)

    return fileName


def write_docx_pdf(documentId, title, questions, notes, telegram_user):
    
    now_date = datetime.now(pytz.timezone('Europe/Madrid')).strftime("%Y-%m-%d")
    document = Document()

    bold_style = document.styles.add_style('Bold List Number', WD_STYLE_TYPE.PARAGRAPH)
    bold_style.font.bold = True

    document.add_heading(title, 0)
    document.add_paragraph('By JW Library Plus - https://github.com/GeiserX/jwlibrary-plus', style="Subtitle")

    for i in range(len(questions)):
        p = document.add_paragraph(style='Bold List Number')
        p.add_run(questions[i].text).font.size = Pt(12)
        document.add_paragraph(notes[i])
    fileNameDoc = "userBackups/{0}/jwlibrary-plus-{1}-{2}.docx".format(telegram_user, documentId, now_date)
    document.save(fileNameDoc)

    fileNamePDF = "userBackups/{0}/jwlibrary-plus-{1}-{2}.pdf".format(telegram_user, documentId, now_date)
    cmd_str = "abiword --to=pdf --to-name={0} {1}".format(fileNamePDF, fileNameDoc)
    subprocess.run(cmd_str, shell=True)
    return fileNameDoc, fileNamePDF

def main(url, telegram_user, qs_user) -> None:

    title, base_text, song, summary, questions, documentId, articleId, q_map = w_extract_html(url, get_all=True)
    notes = w_query_openai(title, base_text, song, summary, q_map, qs_user)
    filenamejw = write_jwlibrary(documentId, articleId, title, questions, notes, telegram_user)
    filenamedoc, filenamepdf = write_docx_pdf(documentId, title, questions, notes, telegram_user)
    return filenamejw, filenamedoc, filenamepdf

if __name__ == "__main__":
    main()

    