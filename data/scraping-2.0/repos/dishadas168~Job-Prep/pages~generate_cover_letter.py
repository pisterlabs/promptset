import streamlit as st
from database import Database
import config
from langchain.chains.llm import LLMChain
from langchain.prompts import PromptTemplate
from langchain.chains.combine_documents.stuff import StuffDocumentsChain
from langchain.chat_models import ChatOpenAI
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document
from langchain.chains.summarize import load_summarize_chain
import docx 
from docx.shared import Pt 
import re
import logging

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s - %(lineno)d")

db = Database()

def main():

    st.title("Generate Cover Letter")
    st.write("Paste a job description below along with optional details to get a cover letter fit for this position.")
    job_desc = st.text_area("")
    extra_info = st.text_area("Optional background information")
    button = st.button("Generate cover letter!")

    if button:
        resume_text = db.get_resume()
        if resume_text is None:
            st.write("Please upload your resume.")
        else:
            with st.spinner('Generating cover letter...'):

                logging.info("Generating cover letter...")
                try:
                    text_splitter = CharacterTextSplitter()
                    texts = text_splitter.split_text(job_desc)
                    docs = [Document(page_content=t) for t in texts]

                    # Define prompt
                    prompt_template = f"""
                    You are given my resume, my background information and a job description. 
                    Based on the details provided in the job description below, align my pertinent 
                    skills and experience with the specific language and requirements outlined in 
                    the job description to create a cover letter. This approach should enhance the 
                    likelihood of being considered a strong candidate for the position. 
                
                    RESUME:
                    {resume_text["content"]}

                    BACKGROUND INFORMATION:
                    {extra_info}

                    """ + "\n\nJOB DESCRIPTION: \n\n{job_desc}"

                    prompt = PromptTemplate.from_template(prompt_template)

                    # Define LLM chain
                    llm = ChatOpenAI(openai_api_key=config.openai_api_key, temperature=0, model_name="gpt-3.5-turbo-16k")
                    llm_chain = LLMChain(llm=llm, prompt=prompt)

                    # Define StuffDocumentsChain
                    stuff_chain = StuffDocumentsChain(
                        llm_chain=llm_chain,
                        document_variable_name="job_desc"
                    )

                    output = stuff_chain.run(docs)

                    st.write(output)
                except Exception as e:
                    logging.error("An error occurred while generating cover letter : %s", str(e))

                try:
                    # Create an instance of a word document 
                    logging.info("Converting to DOCX...")
                    doc = docx.Document() 
                    output = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+', '', output)
                    doc.add_paragraph(output) 
                    doc.save('cover_letter.docx')

                    with open('cover_letter.docx', 'rb') as f:
                        st.download_button('Download as Docx', f, file_name='cover_letter.docx')
                except Exception as e:
                    logging.error("An error occurred while generating DOCX file : %s", str(e))

    
if __name__ == '__main__':
    main()