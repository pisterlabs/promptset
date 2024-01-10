import anthropic
from docx import Document
import gradio as gr
import os
import platform
import tempfile

# Run in debugging mode on Mac OS
if platform.system() == "Darwin":
    DEBUG = True
else:
    DEBUG = False

# Turn on auth if PATENTGEN_USERNAME and PATENTGEN_PASSWORD are set
USERNAME = os.environ.get("PATENTGEN_USERNAME")
PASSWORD = os.environ.get("PATENTGEN_PASSWORD")

MAX_TOKENS = 1000

SYSTEM_MESSAGE = "You are the world's best patent attorney.  You are drafting a US patent application based on the attached transcript of an invention disclosure meeting."

CLAIMS_PROMPT = "Draft The Claims section with 10 claims.  Only return the Claims and nothing else."
TITLE_PROMPT = "Draft the title for this patent application.  Only return the Title and nothing else."
TECHFIELD_PROMPT = "Draft the Technical Field section.  Only return the Technical Field and nothing else."
BACKGROUND_PROMPT = "Draft the Background section with 3 paragraphs.  Only return the Background and nothing else."
EMBODIMENTS_PROMPT = "Draft the Summary of Example Embodiments section with 3 example embodiments.  Only return the embodiments and nothing else."
DRAWINGS_PROMPT = "Draft the Brief Description of the Drawings section.  Only return the Drawings and nothing else."

def generate(new_user_message, history=[], temperature=1):
    if 'ANTHROPIC_API_KEY' not in os.environ:
        raise Exception("This model will be run from www.anthropic.com - Please obtain an API key from https://console.anthropic.com/account/keys and then set the following environment variable before running this app:\n```\nexport ANTHROPIC_API_KEY=<your key>\n```")

    client = anthropic.Anthropic()
    prompt = SYSTEM_MESSAGE + "\n"

    for user_message, assistant_response in history:
        if user_message.strip() and assistant_response.strip():
            prompt += anthropic.HUMAN_PROMPT + user_message + "\n" + anthropic.AI_PROMPT + assistant_response + "\n"
        
    prompt += anthropic.HUMAN_PROMPT + new_user_message + anthropic.AI_PROMPT

    if DEBUG:
        print(prompt)
        print("----------------------------------")

    stream = client.completions.create(
        model="claude-2",
        prompt=prompt,
        temperature=temperature,
        max_tokens_to_sample=MAX_TOKENS,
        stream=True
    )

    return stream
    
def gen_section_fn(index):

    def gen_section(transcript_file, *args):
        prompt_list = list(args)

        new_user_message = prompt_list[index]
        prompt_list[index] = ""

        messages = []

        if not transcript_file:
            raise gr.Error("Please upload a transcript of the invention disclosure meeting first!")

        with open(transcript_file.name, 'r') as f:
            transcript = f.read()
            if transcript:
                messages.append((transcript, "Thank you, I will use this as background info when drafting the patent application."))

        for i in range(0, len(prompt_list), 2):
            messages.append((prompt_list[i], prompt_list[i+1]))

        response = ""
        stream = generate(new_user_message, history=messages)

        for chunk in stream:
            response += chunk.completion
            yield response

        return response

    return gen_section

def gen_word_doc(claims_gen, title_gen, techfield_gen, background_gen, embodiments_gen, drawings_gen):
    doc = Document()

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    temp_path = temp_file.name

    doc.add_heading('TITLE', level=1)
    doc.add_paragraph(title_gen)
    
    doc.add_heading('TECHNICAL FIELD', level=1)
    doc.add_paragraph(techfield_gen)
    
    doc.add_heading('BACKGROUND', level=1)
    doc.add_paragraph(background_gen)
    
    doc.add_heading('SUMMARY OF EXAMPLE EMBODIMENTS', level=1)
    doc.add_paragraph(embodiments_gen)
    
    doc.add_heading('BRIEF DESCRIPTION OF THE DRAWINGS', level=1)
    doc.add_paragraph(drawings_gen)

    doc.add_heading('CLAIMS', level=1)
    doc.add_paragraph(claims_gen)

    doc.save(temp_path)
    
    return temp_path


with gr.Blocks() as demo:

    transcript = gr.File(label="Transcript of invention disclosure meeting")

    claims_prompt = gr.Textbox(label="Prompt", value=CLAIMS_PROMPT, interactive=True)
    claims_gen = gr.Textbox(lines=5, label="Claims", interactive=True, show_copy_button=True)
    claims_button = gr.Button(value="Generate Claims")

    title_prompt = gr.Textbox(label="Prompt", value=TITLE_PROMPT, interactive=True)
    title_gen = gr.Textbox(lines=1, label="Title", interactive=True, show_copy_button=True)
    title_button = gr.Button(value="Generate Title")

    techfield_prompt = gr.Textbox(label="Prompt", value=TECHFIELD_PROMPT, interactive=True)
    techfield_gen = gr.Textbox(lines=5, label="Technical Field", interactive=True, show_copy_button=True)
    techfield_button = gr.Button(value="Generate Technical Field")

    background_prompt = gr.Textbox(label="Prompt", value=BACKGROUND_PROMPT, interactive=True)
    background_gen = gr.Textbox(lines=5, label="Background", interactive=True, show_copy_button=True)
    background_button = gr.Button(value="Generate Background")
    
    embodiments_prompt = gr.Textbox(label="Prompt", value=EMBODIMENTS_PROMPT, interactive=True)
    embodiments_gen = gr.Textbox(lines=5, label="Embodiments", interactive=True, show_copy_button=True)
    embodiments_button = gr.Button(value="Generate Embodiments")

    drawings_prompt = gr.Textbox(label="Prompt", value=DRAWINGS_PROMPT, interactive=True)
    drawings_gen = gr.Textbox(lines=5, label="Drawings", interactive=True, show_copy_button=True)
    drawings_button = gr.Button(value="Generate Drawings")

    word_doc = gr.File(label="Output Word Doc")
    combine_button = gr.Button(value="Combine All Sections Into Word Doc", variant="primary")

    inputs = [
        transcript,
        claims_prompt, claims_gen, # 0
        title_prompt, title_gen, # 2
        techfield_prompt, techfield_gen, # 4
        background_prompt, background_gen, # 6
        embodiments_prompt, embodiments_gen, # 8
        drawings_prompt, drawings_gen, # 10
    ]

    claims_button.click(gen_section_fn(0), inputs=inputs, outputs=claims_gen)
    title_button.click(gen_section_fn(2), inputs=inputs, outputs=title_gen)
    techfield_button.click(gen_section_fn(4), inputs=inputs, outputs=techfield_gen)
    background_button.click(gen_section_fn(6), inputs=inputs, outputs=background_gen)
    embodiments_button.click(gen_section_fn(8), inputs=inputs, outputs=embodiments_gen)
    drawings_button.click(gen_section_fn(10), inputs=inputs, outputs=drawings_gen)

    combine_button.click(
        gen_word_doc, 
        inputs=[
            claims_gen, 
            title_gen, 
            techfield_gen, 
            background_gen, 
            embodiments_gen, 
            drawings_gen
        ],
        outputs=word_doc
    )

if USERNAME and PASSWORD:
    demo.queue().launch(auth=(USERNAME, PASSWORD), share=False, debug=DEBUG)
else:
    demo.queue().launch(share=False, debug=DEBUG)
