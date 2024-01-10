import os

try:
    import openai
    import dotenv
    import requests
    from pytube import YouTube
    from docx import Document
except ImportError as e:
  print("Missing dependencies detected. Run pip install -r requirements.txt to install...")
  print(e)

MODEL = "gpt-3.5-turbo-16k-0613" # The best model for this application - large text, but stay cheap

dotenv.load_dotenv()

STUCKINVIM_KEY = os.getenv("STUCKINVIM_KEY")

TEMPERATURE = 0.35 # Tweaked manually
MAX_TOKENS = 1000

if not STUCKINVIM_KEY:
    print("Missing API key. Please set the STUCKINVIM_KEY environment variable.")
    exit(1)

API_KEY = None

try:
    result = requests.get("http://getkey.stuckinvim.com/api/data?api_key=" + STUCKINVIM_KEY)

    if result is not None:
        result = result.json()
    else:
        "Failed to get API key from StuckInVim API."

    if result.get("status", "400") == "200":
        API_KEY = result["key"]
    else:
        raise Exception("Failed to get API key from StuckInVim API.")

    openai.api_key = API_KEY
except Exception as e:
    print("Failed to get OpenAI API key. Script will continue to run with limited features")

def get_video_description(video_url):
    if video_url.startswith("C:") or video_url.startswith("\\") or video_url.startswith("/"):
        return None
    elif not video_url.startswith("http") and not video_url.startswith("www") and not video_url.startswith("youtube"):
        video_url = "https://www.youtube.com/watch?v=" + video_url
    print(video_url)
    try:
        yt = YouTube(video_url)
        return yt.initial_data["engagementPanels"][1]["engagementPanelSectionListRenderer"]["content"]["structuredDescriptionContentRenderer"]["items"][1]["expandableVideoDescriptionBodyRenderer"]["attributedDescriptionBodyText"]["content"]
    except Exception as e:
        print(e)
        return None

def chatbot(questions_path, transcript_path, save_path, youtube_url=None):
    questions = ""
    if questions_path:
        try:
            doc = Document(questions_path)
        except:
            print("Error")
        if questions_path.endswith(".docx") or questions_path.endswith(".doc"):
            for paragraph in doc.paragraphs:
                questions += paragraph.text
        elif questions_path.endswith(".txt"):
            with open(questions_path, encoding="utf-8") as txt:
                questions = txt.read()
    elif youtube_url:
        print("Getting from URL")
        questions = get_video_description(youtube_url)
    else:
        print("No questions provided.")
        exit(1)

    if questions is None or questions == "":
        print("Failed to get questions from description.")
        exit(1) 

    transcript = ""
    with open(transcript_path, encoding="utf-8") as txt:
       transcript = txt.read()

    system_message_path = os.path.join(os.path.dirname(__file__), "system_message.txt")
    system_message = ""
    if os.path.exists(system_message_path):
        with open(system_message_path, encoding="utf-8") as txt:    
            system_message = txt.read()

    print("System message: " + system_message)
    print("Transcript: " + transcript[0:100])
    print("Questions: " + questions[0:100])

    messages = [
        {"role": "system", "content": system_message},
    ]

    messages.append({"role": "user", "content": transcript})

    print(messages)

    print("Sending prompt with transcript...")

    response = openai.chat.completions.create(
        model=MODEL,
        messages=messages,
        temperature=TEMPERATURE,
        max_tokens=MAX_TOKENS
    )
    print(response)
    print("Sending prompt with questions...")

    messages.append({"role": "user", "content": questions})
    print(messages)
    response = openai.chat.completions.create(
        model=MODEL,
        messages=messages,
        temperature=TEMPERATURE,
        max_tokens=MAX_TOKENS
    )

    try:
        response_text = response.choices[0].message.content

        if save_path.endswith(".docx") or save_path.endswith(".doc"):
            doc = Document()
            doc.add_paragraph(response_text)
            doc.save(save_path)
        elif save_path.endswith(".txt"):
            with open(save_path, "w", encoding="utf-8") as txt:
                txt.write(response_text)
    except:
        print("Error while saving answers. Will print to console instead.")

        print(response)