import openai
import os
import shutil
from loguru import *
from docx import *
#from change_key import main as change_key
#from get_key import main as get_key
#from get_path import get_temp_path, get_backup_path
from docx import Document
from tools.change_key import main as change_key
from tools.change_key import change_key_by_rate_limit
from tools.get_key import main as get_key
from tools.get_path import get_temp_path, get_backup_path
from tools.get_time_now import get_time
from tools.write_log import main as write_log
from docx.shared import Pt
from docx.oxml.ns import qn
from tools.get_ac_name import get_name
from docx.enum.text import WD_LINE_SPACING,WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm

def save_docx(article_data, path, params,name):
    save_path = f'{path}/{article_data["title"]}.docx'

    # 获取标题
    title = article_data['title']

    # 获取文章内容
    content = article_data['content']

    # 为答案设置编码字符集
    content.strip().encode('utf-8')

    # 创建document对象
    doc = Document()

    # 设置西方字体格式
    doc.styles['Normal'].font.name = 'Times New Roman'

    # 设置中文字体格式
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 获取标题所在段落
    p1 = doc.add_paragraph()

    # 为标题段落赋值并设置字体大小和加粗设置
    text1 = p1.add_run(title)  # 这里写标题

    # 字体大小
    text1.font.size = Pt(18)

    text1.bold = True

    # 设置标题内容居中
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 获取标题所在段落
    p2 = doc.add_paragraph()

    # 为正文赋值并设置首行缩进，字体大小和行距
    text2 = p2.add_run(content)  # 写入若干段落 # 写读取到的数据

    # 首行缩进0.75cm
    p2.paragraph_format.first_line_indent = Cm(0.75)

    # 字体大小
    text2.font.size = Pt(14)

    # 1.5倍行距
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # 修改作者信息，不然默认作者是"python-docx"
    doc.core_properties.author = "Administrator"

    info = f'{name}: 文章正在保存'
    logger.info(info)
    write_log(name,get_time()+' '+info, params)
    doc.save(save_path)

    if params['backup_switch']:
        backup_path = get_backup_path()
        if not os.path.exists(backup_path):
            os.makedirs(backup_path)
        info = f'{name}: 文章正在备份'
        logger.info(info)
        write_log(name,get_time()+' '+info, params)
        shutil.copy(save_path, backup_path)


# def write_article(title, params, name):
#     dic,_ = get_name()
#     openai.api_key = get_key()
#     path = f'{get_temp_path()}/{name}'
#     if not os.path.exists(path):
#         path = f'{get_temp_path()}/{dic[name]}'
#     info = f'{name}: 开始写文章   `{title}`'
#     logger.info(info)
#     write_log(name,get_time()+' '+info, params)
#     prompt = f"{title}。 用中文生成内容长度700字左右;并且按照列表划分好。"
#     prompt.encode('utf-8')
#     response = openai.Completion.create(
#         model="text-davinci-003",
#         prompt=prompt,
#         temperature=1,
#         max_tokens=2000,
#         frequency_penalty=0,
#         presence_penalty=0
#     )
#     article = response.choices[0].text
#     article_data = {
#         'title': title,
#         'content': article
#     }
#     save_docx(article_data, path, params,name)

def write_article(title, params, name):
    dic, _ = get_name()
    openai.api_key = get_key()
    path = f'{get_temp_path()}/{name}'
    if not os.path.exists(path):
        path = f'{get_temp_path()}/{dic[name]}'

    write = True
    for files in os.listdir(path):
        if title in files:
            write = False
            break
    if not write:
        info = info = f'{name}: 文章 `{title}` 已经存在,跳过该标题'
        logger.info(info)
        write_log(name, get_time() + ' ' + info, params)
        pass
    else:
        info = f'{name}: 开始写文章  `{title}`'
        logger.info(info)
        write_log(name, get_time() + ' ' + info, params)
        os.environ["http_proxy"] = "http://127.0.0.1:7890"
        os.environ["https_proxy"] = "http://127.0.0.1:7890"
        prompt = f"{title}。 用中文生成内容长度700字左右;并且按照列表划分好。"
        prompt.encode('utf-8')
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Serve me as a writing and programming assistant."},
                {"role": "user", "content": prompt},

            ],
            max_tokens=2000,
            temperature=1,
            frequency_penalty=0,
            presence_penalty=0
        )
        article = response['choices'][0]['message']['content']
        article_data = {
            'title': title,
            'content': article
        }
        save_docx(article_data, path, params, name)


import time
def main(title, params, name):
    exam = 0
    error = 0
    while True:
        try:
            write_article(title, params, name)
        except Exception as e:
            if str(e) == 'The server is overloaded or not ready yet.':
                info = f'{name}: 服务器过载,将在5秒后重试'
                logger.info(info)
                write_log(name,get_time()+' '+info,params)
                time.sleep(5)
                continue
            elif str(e) == 'You exceeded your current quota, please check your plan and billing details.':
                try:
                    info = f'{name}: key用光正在尝试更换'
                    logger.info(info)
                    write_log(name,get_time()+' '+info,params)
                    change_key()
                    openai.api_key = get_key()
                    continue
                except:
                    info = f'{name}: keys用光了'
                    logger.warning(info)
                    write_log(name,get_time()+' '+info,params)
                    exam = 1
                    break
            elif 'Rate limit reached' in str(e):
                info = f'{name}: 速率被限制，等待100s后换key重试(此问题可通过将openai绑定信用卡解决)'
                logger.warning(info)
                write_log(name,get_time()+' '+info,params)
                time.sleep(100)
                try:
                    change_key_by_rate_limit()
                    openai.api_key = get_key()
                    continue
                except:
                    info = f'{name}: 当前只有一个key，无可用key进行更换'
                    logger.warning(info)
                    write_log(name,get_time()+' '+info,params)
                    continue  
            else:
                error += 1
                if error >= 20:
                    info = f"{name}: 异常过多,请检查"
                    logger.warning(info)
                    write_log(name,get_time()+' '+info,params)
                    write_log(name,str(e),params)
                    exam = 1
                    break
                info = f'{name}: 出现第{str(error)}次异常'
                logger.warning(info)
                write_log(name,get_time() + ' ' +info,params)
                print(str(e))
                continue
        else:
            info = f'{name}: 文章书写完成'
            logger.info(info)
            write_log(name,get_time() + ' '+ info,params)
            break

    return exam


if __name__ == '__main__':
    params = {

    # 最大文章数量 这个没啥用 还没写 现在只会跑满
    "max_article": 100,

    # 文章种类排序
    #"sort": ['行业资料', '互联网', '基础教育', '学前教育', '高校与高等教育', '推荐', '实用模板',
    #        '语言/资格考试', '建筑', '商品说明书', '政务民生', '法律'],
    "sort":['推荐','实用模板'],
            
    # 是否显示浏览器窗口 True为显示,False为不显示
    "display": True,

    #是否使用用户名作为消息提醒开启则为用户名,关闭则为账号提示
    "use_name":True,

    #标题信息显示
    "title_info":False,

    # 备份开关,true为开启备份,false为不开启
    "backup_switch": False,

    # 自动关机,true为跑完自动关机,false为不关机
    "shutdown": False,

    #是否上传pdf格式的文件
    "ues_pdf": False,

    # 多线程开关 True为开启多线程 False为关闭
    "Multithreading": True,

    # 线程数 同时跑多少个账号
    "num": 2,
}
    main('我是你爸爸',params,'马克思')
