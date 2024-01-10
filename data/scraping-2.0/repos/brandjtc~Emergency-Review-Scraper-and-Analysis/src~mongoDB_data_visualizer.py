import matplotlib.pyplot as plt
import numpy as np
import os
import pickle
from pymongo import MongoClient
import pandas as pd
import datetime as dt
from tzlocal import get_localzone
from docx import Document as doc
import random
import time
import openAI_main as Jarvis
from dotenv import load_dotenv, find_dotenv
_ = load_dotenv(find_dotenv())

#Constants
LONGREVSCORE=True
#Establishes connection to mongo DB
mongoConnectStr=os.getenv("MONGO_KEY")
client = MongoClient(mongoConnectStr)
#Selects EmergencyReviewDB as the database in Mongo to populate
dbString=os.getenv("DATABASE")
proj_db=client[dbString]

#Connects to tables in aforementioned DB
colString=os.getenv("COLLECTION")
review_collection = proj_db[colString]

#Variable instantiation for time functions
start= dt.datetime.now(tz=get_localzone())


def timeStart():
    print("-"*40)
    fmt = "%m/%d/%y - %T %p"  
    start= dt.datetime.now(tz=get_localzone())
    print("Began at:"+start.strftime(fmt))

def timeEnd():
    print('-'*40)
    fmtend = "%m/%d/%y - %T %p"  
    end= dt.datetime.now(tz=get_localzone())
    print("Ended at: "+end.strftime(fmtend))
    print(f"Time elapsed: {end - start}")
    print("-"*40)
    print("\n")

#Variable declaration for data gathering, used in report doc generation later
playStoreReviewCount=0
playStorePercent=0
appStoreReviewCount=0
appStorePercent=0
meanLength=0
standardDeviation=0
shortLenList=[]
shortCount=0
shortReviewPercent=0
middleCount=0
middleReviewPercent=0
largeCount=0
largeReviewPercent=0
sampledShortRev=[]

print("Converting input excel info to lists and calculating total apps and generating figure")
timeStart()

#Gathering Data from input excel
excel_filepath = os.getenv("EXCELFILEPATH")
excelDoc = pd.read_excel(excel_filepath,usecols='A,L,M')
apple_link_list = list(excelDoc['Link to App Store'])
google_link_list = list(excelDoc['Link to Google Play Store'])

#Gathering data from screening excel that filters out irrelevant apps.
screening_excel=pd.read_excel(os.getenv("EXCELFILEPATH"), sheet_name="recon",engine='openpyxl')
reviewed_app_ids = list(screening_excel['appId'])
scores = list(screening_excel['total_yes'])
app_name_list=list(excelDoc['App-Name'])
app_id_list=list()
for i in range(0,len(apple_link_list)):
    if (str(google_link_list[i])!="nan"):
        tempSplit=(google_link_list[i].split("?id="))[1]
        if "&" in tempSplit:
            tempSplit=tempSplit.split("&")[0]
        temp=tempSplit
    else:
        #The link is converted to a string for this conditional to make it easier to process NAs,
        #when there is no link. 
        if (str(apple_link_list[i])!="nan"):
            tempSplit = apple_link_list[i].split("/")
            idSplit=tempSplit[6].split("d")

            #If statement checks for lingering language tag in the ID extracted from the app link. 
            #If present, it is removed and the temp variable is adtempApp.reviews[j]usted accordingly to be added to the app ID list.
            if ("?" in idSplit[1]):
                idSplit=idSplit[1].split("?")
                temp=idSplit[0]
            else:
                temp = idSplit[1]
        #Appends ID to the list. If there is no ID in the Apple app store, NA is added as the ID instead.
        #Same system with country codes
    app_id_list.append(temp)

timeEnd()

print("Compiling apps that fit the emergency app criteria and setting up future data inputs")
timeStart()
# Create a list of apps to be excluded; found based on score
excluded_app_ids = []
excluded_app_names=[]
for app_id, score, app_name in zip(reviewed_app_ids, scores, app_name_list):
    if score == 0:
        excluded_app_ids.append(app_id)
        excluded_app_names.append(app_name)

# Set up MongoDB query

cursor = review_collection.find({})

# Prepare data for Review Log xlsx file
review_logs = []
#Setting up input for histogram using data from this query
histogramInput = []
#Setting up input for percentage data from this query
percentageInput=[]

#Lists used to gather info on 100 longest reviews
revList=list()
lenList=list()
idList=list()
nameList=list()

#Increments through the MongoDB call and fills all required lists and data metrics for later use in one go
for document in cursor:
    # Extract desired traits from the document and add them to the data list
    if document['content']!=None:
        if len(document['content'])>25:
            #Lists used to gather info on 100 longest reviews
            lenList.append(len(document["content"]))
            revList.append(document["content"])
            idList.append(document['_id'])
            nameList.append(document['app_name'])

            review_logs.append({
                'Document Id': document['_id'],
                'App Name': document['app_name'],
                'App Id': document['app_id'],
                'Username': document['userName'],
                'Review': document['content'],
                'Score': document['score'],
                'Translated Review': document.get('translated_content', ''),
                'Language': document.get('language',''),
                'Platform': document['platform']
            })
            histogramInput.append(len(document['content']))
        if document.get('translated_content',None)==None:
            percentageInput.append([document['content'],document['platform'],document['app_name'],document['app_id']])
        else:
            percentageInput.append([document['translated_content'],document['platform'],document['app_name'],document['app_id']])

# Save the data list object using pickle
pklFilePath=os.getenv("PKLFILEPATH")
with open(pklFilePath, 'wb') as pklFile:
    pickle.dump(review_logs, pklFile)

client.close()
timeEnd()

print("Creating charts/visual graphics")
timeStart()
#Variable declaration for graphic generation
playstoreAndIos=0
exclusivePlaystore=0
exclusiveAppstore=0
for i in range(0,len(apple_link_list)):
    if (app_id_list[i] not in excluded_app_ids):
        if (str(apple_link_list[i])!="nan")&(str(google_link_list[i])!="nan"):
            playstoreAndIos+=1
        elif (str(apple_link_list[i])!="nan")&(str(google_link_list[i])=="nan"):
            exclusiveAppstore+=1
        elif (str(google_link_list[i])!="nan")&(str(apple_link_list[i])=="nan"):
            exclusivePlaystore+=1


barXinp=np.array(['Exclusive to\nGoogle PlayStore','Exclusive to\nIOS App Store','IOS and\n Google Play'])
barYinp=np.array([exclusivePlaystore,exclusiveAppstore,playstoreAndIos])
fig, ax = plt.subplots()
bars = ax.bar(barXinp, barYinp)

ax.bar_label(bars)

plt.title("App distribution\n among App Stores")
plt.ylabel("Amount on specified platform")
plt.savefig('./Generated Files/storeCount.png', 
           transparent=False,  
           facecolor='white', 
           bbox_inches="tight")

timeEnd()


print("Gathering review percentage counts")
timeStart()
longReviewList=list()
shortReviewList=list()
for item in percentageInput:
    if(item[0]!=None):
        if(len(item[0])>25):
            if (item[1]=="Google Play Store"):
                playStoreReviewCount+=1
            else:
                appStoreReviewCount+=1
            if(len(item[0])<200):
                middleCount+=1
            else:
                largeCount+=1
                longReviewList.append([item[0],item[2],item[3]])
        else:
            shortCount+=1
            shortLenList.append(len(item[0]))
            shortReviewList.append(item[0])
    else:
        shortCount+=1

preRoundRevTotal=middleCount+shortCount+largeCount

middleReviewPercent=round((middleCount/preRoundRevTotal)*100,2)
largeReviewPercent=round((largeCount/preRoundRevTotal)*100,2)

reviewTotal=playStoreReviewCount+appStoreReviewCount
playStorePercent=round((playStoreReviewCount/reviewTotal)*100,2)
appStorePercent=round((appStoreReviewCount/reviewTotal)*100,2)
timeEnd()


print("Generating review length histogram")
timeStart()

    

#Gathers the percentage of short reviews. These will all be excluded from final dataset for the rest of these calculations.
shortReviewPercent=round(shortCount/preRoundRevTotal*100,2)
# Creation of histogram containing review lengths as dataset
fig, ax = plt.subplots()
values, bins, bars = plt.hist(np.log10(histogramInput),linewidth=0.5, edgecolor="white")
plt.bar_label(bars)
plt.margins(x=0.1,y=0.1)

plt.title("Review Length Frequency")
plt.ylabel("Amount")
plt.xlabel("Log (Base 10) of Review Length in Characters")
plt.savefig('./Generated Files/reviewLength_Histogram.png', 
           transparent=False,  
           facecolor='white', 
           bbox_inches="tight")
timeEnd()

print("Building report doc")
timeStart()
#Histogram input already has the short lengths removed meaning it can be
#used for calculating standard deviation and mean
standardDeviation=round(np.std(histogramInput),2)
meanLength=round(np.mean(histogramInput),2)
sampledShortRev=random.sample(shortReviewList,3)
reportDoc = doc()

reportDoc.add_heading('Scraper Report', 0)
reportDoc.add_heading('Summary',level=1)
summaryString=f'''The length of the reviews, in characters, varied (M={meanLength}, SD={standardDeviation}, not factoring in reviews less than 25 characters). Approximately {shortReviewPercent}% consisted of just a few words (<25 characters), Approximately {middleReviewPercent}% contained at most one sentence (<200 characters), while {largeReviewPercent}% contained multiple sentences (>201 characters). Since reviews shorter than a few words (<25 characters) likely contain trivial comments such as "{sampledShortRev[0]}", "{sampledShortRev[1]}", and "{sampledShortRev[2]}." Their content is unlikely to have the depth or breadth needed to significantly invoke any of the dimensions or sub-dimensions presented in our theoretical model. Therefore, we excluded them from the dataset. Thus, after removing {shortCount} reviews, our final dataset consisted of {reviewTotal} reviews, with approximately ({playStorePercent}% being from the Google Play Store and {appStorePercent}% from the Apple AppStore).'''
reportDoc.add_paragraph(summaryString)
reportDoc.add_heading('Figure 1', level=1)
reportDoc.add_picture('./Generated Files/storeCount.png')
reportDoc.add_heading('Figure 2', level=1)
reportDoc.add_picture('./Generated Files/reviewLength_Histogram.png')
reportDoc.save('./Generated Files/report.docx')
timeEnd()

print("Sampling and Scoring 15 reviews over 200 characters")
timeStart()
sampledLongRev=random.sample(longReviewList,15)
LongReviewScoreList=list()
reviewNum=0
for u in range(0,len(sampledLongRev)):
    reviewNum+=1
    timeStart()
    print(f"Input to OpenAI: {sampledLongRev[u][0]}")
    LongReviewScoreList.append(Jarvis.aiInput(sampledLongRev[u][0],sampledLongRev[u][1],sampledLongRev[u][2],u))
    timeEnd()
    time.sleep(19)
    

#Creating dataframe
longRevDF = pd.json_normalize(LongReviewScoreList)

# Saving the DataFrame to an Excel file
longRevDF.to_excel("./Generated Files/Top 15 Reviews Scored.xlsx", index=False)
timeEnd()

print("List of JSON objects converted to Excel successfully!")

print("Generating Excel with 100 Longest Reviews")
timeStart()

print("Grabbing index of 100 longest reviews")
#Grabs the index of the top 100 longest reviews.

index=np.argsort(lenList)[::-1][:100]
client.close()

dictList=list()
print("-"*40)
print("Compiling reviews into list of dictionaries.")
#Populates list of dictionaries to later be put in PD dataframe
for i in range(0,len(index)):
    dictList.append({"App Name":nameList[index[i]],"Review":revList[index[i]],"Revew Length\n in Characters":lenList[index[i]],"Mongo ID":idList[index[i]]})

#Inserts list of dicts into a pandas dataframe
dfDictList=pd.DataFrame(dictList)
print(dfDictList)
print("-"*40)
print("Excel file created")
dfDictList.to_excel("./Generated Files/Top 100 Longest Reviews.xlsx")

timeEnd()
