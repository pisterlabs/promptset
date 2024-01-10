# Contains OpenAI functions form Streamlit

from flask import Blueprint
import os
from datetime import datetime  # using datetime for title
import openai
from dotenv import load_dotenv

# For printing out a remote service report
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert

# new convert module
import converting
import threading

load_dotenv()  # Load environment variables from .env file

openai.api_key = os.getenv("OPENAI_API_KEY")

class WhisperFunctions:

  def whisper(self, filepath):
    
    audio_file= open(filepath, "rb")
    print('name of audio file is: ', audio_file)
    print('name of filepath is: ', filepath)
    speech = openai.Audio.transcribe("whisper-1", audio_file)
    print('Transcription completed')
    return str(speech)
  

  def diarise(self, speech):
    # prompt completion model with instructions on diarisation
    messages = [
    {
      "role":
      "user",
      "content": "Please split the following transcript into a script between 2 people and format it so each person's dialog starts on a new paragraph."
      },  
      {
      "role": 
      "user", 
      "content": speech
      }
    ]

    response = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages=messages)
    diarised_transcript = response["choices"][0]["message"]["content"]

    return str(diarised_transcript) # this needs to be saved to disk so the following functions can utilise it

  def summary(self, transcript):
    messages = [
    {"role": "user", "content": "Please summarise the following transcript between an SEW Hotline Technician and a customer."},
    {"role": "user", "content": transcript}
    ]

    summary = openai.ChatCompletion.create(
      model="gpt-3.5-turbo", 
      messages = messages
    )
    return summary['choices'][0]['message']['content']
  
  def subject(self, summary):
    messages = [
    {"role": "user", "content": "Please reduce this summary of a call between an SEW Hotline Technician and a customer to a one-sentence subject line."},
    {"role": "user", "content": summary}
    ]

    subject = openai.ChatCompletion.create(
      model="gpt-3.5-turbo", 
      messages = messages
    )
    return subject['choices'][0]['message']['content']
  
class AiQuery:
  
  def ask_gpt(self, question, diarised_transcript):

    prompt = question

    messages = [
    {"role": "user", "content": "Please read the following transcript summary between an SEW Hotline Technician and a customer, and then answer the question in the following prompt."},
    {"role": "user", "content": diarised_transcript},
    {"role": "user", "content": prompt}
    ]
    response = openai.ChatCompletion.create(
    model="gpt-3.5-turbo", 
    messages = messages
    )
    answer = response['choices'][0]['message']['content']

    return answer

# Where the PDF report is created
class Template:
  def createBoothPrintout(self, document, summary, details, resolution, length, nameTechnician, nameCustomer, subject, product, serialNumber, companyName):
    # Defines a new style with the given identifier. Font is set to Lato as default. Size needs to be specified.
    def font(identifier, size, font='Lato'):

      styles = document.styles
      style = styles.add_style(identifier, WD_STYLE_TYPE.PARAGRAPH)
      style.font.name = font
      style.font.size = Pt(size)

    font('summary', 12)
    font('details', 12)
    font('resolution', 12)
    font('length', 12)

    for paragraph in document.paragraphs:
      if 'NameTec' in paragraph.text:
        paragraph.text = paragraph.text.replace('NameTec', '')
        paragraph.add_run('Technician: ').bold = True
        paragraph.add_run(nameTechnician)
        paragraph.style = document.styles['summary']

      if 'NameCus' in paragraph.text:
        paragraph.text = paragraph.text.replace('NameCus', '') 
        paragraph.add_run('Customer: ').bold = True
        paragraph.add_run(nameCustomer)
        paragraph.add_run(', ' + companyName)
        paragraph.style = document.styles['summary'] 

      if 'TextSum' in paragraph.text:
        paragraph.text = paragraph.text.replace('TextSum', summary)
        paragraph.style = document.styles['summary']
                                
      if 'DetailSum' in paragraph.text:
        paragraph.text = paragraph.text.replace('DetailSum', details)
        paragraph.style = document.styles['details']

      if 'ResSum' in paragraph.text:
        paragraph.text = paragraph.text.replace('ResSum', resolution)
        paragraph.style = document.styles['resolution']

      if 'subjectLine' in paragraph.text:
        paragraph.text = paragraph.text.replace('subjectLine', '')
        paragraph.add_run('Subject: ').bold = True
        paragraph.add_run(subject)
        paragraph.style = document.styles['summary']

      if 'product' in paragraph.text:
        paragraph.text = paragraph.text.replace('productName', '')
        paragraph.add_run('Product: ').bold = True
        paragraph.add_run(product)
        paragraph.style = document.styles['summary']

      if 'serialNumber' in paragraph.text:
        paragraph.text = paragraph.text.replace('serialNumber', '')
        paragraph.add_run('Product ID: ').bold = True
        paragraph.add_run(serialNumber)
        paragraph.style = document.styles['summary']

    """ Building and saving the word document path and general filename for this case report"""

    # create a path for the word document
    now = datetime.now().strftime('%Y%m%d_%H%M') # get current date and time

    # get current working directory
    current_directory = os.getcwd()
    print("Current Directory:", current_directory) # Current Directory: C:\Users\<yourusername>\Projects\SEW-Frontend

    # create filename (used by both word and pdf documents)
    filename = f'{now}_booth_printout' 

    # create filepath for word document
    docx_path = os.path.join(f'backend/data/word_reports/{filename}.docx')
    print('filepath: ', docx_path) # print filepath

    # save word document
    document.save(docx_path) 
    print('Word document saved')

    # create a path for the PDF
    pdf_path = os.path.join(f'backend/data/reports/{filename}.pdf')

    print("Attempting to convert word document to PDF...")
    thread = threading.Thread(target=converting.convert_docx_to_pdf(docx_path, pdf_path))
    thread.start()

    return pdf_path
    
    # convertSuccess = converting.convert_docx_to_pdf(docx_path, pdf_path)

    # if convertSuccess:
    # print("Word document converted to PDF successfully? : ", convertSuccess)
      # return pdf_path

class ReportFunctions:
  # Clean up text
  def cleanText(text):
    while text.startswith('\n'):
      text = text[1:]
    return text

	# Get details and resolution function
  def getDetailsandResolution(diarised_transcript, technician, customer):

    messages = [
    {"role": "user", "content": f"Read the following transcript between {technician} (an SEW Hotline Technician) and {customer} (a customer from Dublin airport). Return the details of the call in bullet points."},
    {"role": "user", "content": diarised_transcript},
    ]

    response = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages = messages)
    details = response['choices'][0]['message']['content']

    details = ReportFunctions.cleanText(details)
    
    messages = [
    {"role": "user", "content": f"Read the following transcript between {technician} (an SEW Hotline Technician) and {customer} (a customer from Dublin airport). Return a short summary of the resolution of the call."},
    {"role": "user", "content": diarised_transcript},
    ]

    response = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages = messages)
    resolution = response['choices'][0]['message']['content']

    resolution = ReportFunctions.cleanText(resolution)

    return details, resolution