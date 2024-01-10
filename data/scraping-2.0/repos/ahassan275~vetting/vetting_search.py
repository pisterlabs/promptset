import streamlit as st
from langchain.agents import AgentType, initialize_agent, Tool
from langchain.document_loaders import PyPDFLoader
from langchain.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores.faiss import FAISS
from langchain_core.prompts import PromptTemplate
import openai
import os
import requests
from vetting_questions import extracted_questions
import uuid
from docx import Document
from langchain.schema import SystemMessage
import base64
from langchain.agents.agent_toolkits import create_retriever_tool
from langchain.agents.agent_toolkits import create_conversational_retrieval_agent

# Set OpenAI API key
openai.api_key = os.environ["OPENAI_API_KEY"]
# GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
# OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
# GOOGLE_CSE_ID = st.secrets["GOOGLE_CSE_ID"]
GOOGLE_API_KEY = os.environ["GOOGLE_API_KEY"]
GOOGLE_CSE_ID = os.environ["GOOGLE_CSE_ID"]


def get_file_content_as_string(file_path):
    with open(file_path, 'rb') as f:
        binary_file_data = f.read()
    return base64.b64encode(binary_file_data).decode('utf-8')


def create_download_link(file_path, file_name):
    file_content = get_file_content_as_string(file_path)
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{file_content}" download="{file_name}">Download the responses</a>'
    return href


@st.cache_resource
def process_document(file_path):
    loader = PyPDFLoader(file_path)
    pages = loader.load_and_split()
    splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
    docs = splitter.split_documents(pages)
    embeddings = OpenAIEmbeddings()
    retriever = FAISS.from_documents(docs, embeddings).as_retriever()
    return retriever


def google_search(query):
    try:
        endpoint = "https://www.googleapis.com/customsearch/v1"
        params = {
            "key": GOOGLE_API_KEY,
            "cx": GOOGLE_CSE_ID,
            "q": query
        }
        response = requests.get(endpoint, params=params)
        results = response.json().get("items", [])
        return [result["link"] for result in results]
    except Exception as e:
        st.error(f"Error during web search: {e}")
        return []


import tempfile


def handle_uploaded_file(uploaded_file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(uploaded_file.getvalue())
        return tmp.name


def vetting_assistant_page():
    st.title("Vetting Assistant Chatbot")

    if "uploaded_pdf_path" not in st.session_state or "retriever" not in st.session_state:
        uploaded_file = st.file_uploader("Upload a PDF containing the terms of service", type=["pdf"])

        if uploaded_file:
            file_path = handle_uploaded_file(uploaded_file)
            st.session_state.uploaded_pdf_path = file_path
            st.session_state.retriever = process_document(st.session_state.uploaded_pdf_path)
    else:
        st.write("Using previously uploaded PDF. If you want to use a different PDF, please refresh the page.")

    app_name = st.text_input("Enter the name of the app:")

    if "retriever" in st.session_state:
        llm = ChatOpenAI(temperature=0.5, model="gpt-3.5-turbo-16k")
        tools = [
            Tool(
                name="vetting_tool",
                description="Tool for retrieving infomration related to security and privacy",
                func=RetrievalQA.from_llm(llm=llm, retriever=st.session_state.retriever, return_source_documents=True)
            )
        ]
        # tool = create_retriever_tool(
        #     st.session_state.retriever,
        #     "search_terms_service",
        #     "Searches and returns an application's privacy and data policies and terms of use.",
        # )
        # tools = [tool]
        agent_kwargs = {
            "system_message": SystemMessage(content="You are an intelligent Vetting Assistant, "
                                                    "expertly designed to analyze and extract key "
                                                    "information from terms of service documents. "
                                                    "Your goal is to assist users in understanding "
                                                    "complex legal documents and provide clear, "
                                                    "concise answers to their queries.")
        }
        agent = initialize_agent(agent=AgentType.OPENAI_FUNCTIONS, tools=tools, llm=llm, agent_kwargs=agent_kwargs,
                                 verbose=True)
        # agent = create_conversational_retrieval_agent(llm, tools)

        st.write("Ask any question related to the vetting process:")
        query_option = st.selectbox("Choose a predefined query:", extracted_questions)
        user_input = st.text_input("Your Question:", value=query_option)

        if st.button('Start Vetting') and user_input:
            with st.spinner('Processing your question...'):
                try:
                    response = agent.run(user_input)
                    # response = agent({"input": "user_input"})
                    st.write(f"Answer: {response}")
                except Exception as e:
                    st.error(f"An error occurred: {e}")

        st.write("Note: The chatbot retrieves answers from the uploaded document.")

        if 'running_queries' not in st.session_state:
            st.session_state.running_queries = False

        placeholder_message = f"{app_name} is being vetted for compliance and its policies provided in context. Does {app_name} meet this criteria?"
        all_queries = [f"{question} {placeholder_message}" for question in extracted_questions]

        if st.button('Run All Queries'):
            with st.spinner('Processing all queries...'):
                st.session_state.running_queries = True
                doc = Document()
                doc.add_heading('Vetting Assistant Responses', 0)

                for question in all_queries:
                    if not st.session_state.running_queries:
                        break
                    try:
                        response = agent.run(question)
                        doc.add_heading('Q:', level=1)
                        doc.add_paragraph(question)
                        doc.add_heading('A:', level=1)
                        doc.add_paragraph(response)
                    except Exception as e:
                        doc.add_paragraph(f"Error for question '{question}': {e}")

                doc_path = "vetting_responses.docx"
                doc.save(doc_path)
                st.markdown(create_download_link(doc_path, "vetting_responses.docx"), unsafe_allow_html=True)

        if st.button('Stop Queries'):
            st.session_state.running_queries = False

        if st.button('Search Web'):
            with st.spinner('Searching the web...'):
                links = google_search(user_input)
                st.write("Top search results:")
                for link in links:
                    st.write(link)


def pdf_chatbot_page():
    st.title("PDF-based Chatbot")

    if "uploaded_pdf_path" not in st.session_state:
        uploaded_file = st.file_uploader("Upload a PDF", type=["pdf"])

        if uploaded_file:
            file_path = handle_uploaded_file(uploaded_file)
            st.session_state.uploaded_pdf_path = file_path
            st.session_state.retriever = process_document(st.session_state.uploaded_pdf_path)
    else:
        st.write("Using previously uploaded PDF. If you want to use a different PDF, please refresh the page.")

    if "retriever" in st.session_state:
        llm = ChatOpenAI(temperature=0.5, model="gpt-3.5-turbo-16k")
        tools = [
            Tool(
                name="pdf_tool",
                description="Tool for querying based on document content",
                func=RetrievalQA.from_chain_type(llm=llm, retriever=st.session_state.retriever)
            )
        ]
        agent_kwargs = {
            "system_message": SystemMessage(content="You are an intelligent Vetting Assistant, "
                                                    "expertly designed to analyze and extract key "
                                                    "information from terms of service documents. "
                                                    "Your goal is to assist users in understanding "
                                                    "complex legal documents and provide clear, "
                                                    "concise answers to their queries.")
        }
        agent = initialize_agent(agent=AgentType.OPENAI_FUNCTIONS, tools=tools, llm=llm, agent_kwargs=agent_kwargs,
                                 verbose=True)

        instructions_container = st.container()
        with instructions_container:
            st.header("Instructions")
            st.write("""
            - This chatbot provides answers based on the content of the uploaded PDF.
            - Type in your question in the chat input below.
            - Adjust the slider to control the specificity of the chatbot's responses.
            """)

        input_container = st.container()
        with input_container:
            temperature = st.slider("Adjust chatbot specificity:", min_value=0.0, max_value=1.0, value=0.5, step=0.05)
            llm.temperature = temperature

        chat_container = st.container()
        with chat_container:
            if "messages" not in st.session_state:
                st.session_state.messages = []
            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

            user_input = st.text_input("Ask a question about the uploaded PDF:")

            if st.button('Query PDF') and user_input:
                with st.spinner('Processing your question...'):
                    try:
                        response = agent.run(user_input)
                        st.session_state.messages.append({"role": "assistant", "content": response})
                        with st.chat_message("assistant"):
                            st.markdown(response)
                    except Exception as e:
                        st.error(f"An error occurred: {e}")

            if st.button('Search Web'):
                with st.spinner('Searching the web...'):
                    links = google_search(user_input)
                    st.write("Top search results:")
                    for link in links:
                        st.write(link)


# Streamlit UI Configuration
st.set_page_config(page_title="Vetting Assistant Chatbot", layout="wide", initial_sidebar_state="expanded")
page = st.sidebar.selectbox("Choose a Tool:", ["Vetting Assistant", "PDF Chatbot"])

if page == "Vetting Assistant":
    vetting_assistant_page()
elif page == "PDF Chatbot":
    pdf_chatbot_page()
