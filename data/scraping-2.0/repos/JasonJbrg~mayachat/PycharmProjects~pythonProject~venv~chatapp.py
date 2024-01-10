import openai
import streamlit as st
from streamlit_chat import message as msg
from translate import Translator
from pydub import AudioSegment
from pydub.playback import play
import os
from dotenv import load_dotenv
import json
import docx
import io
from docx import Document
from datetime import datetime


# Load environment variables from .env file
load_dotenv()

# Read the config.json file
with open("config.json") as file:
    config = json.load(file)

# Extract the values from the config dictionary
task_selection = config["task_selection"]
initial_context = {
    task: f"{config['initial_context'][task]} Please provide brief and concise responses."
    for task in task_selection
}

load_dotenv()  # take environment variables from .env.
openai.api_key = os.getenv('OPENAI_KEY')

# Set page configuration
st.set_page_config(
    page_title="Maya Lingo",
    page_icon="/Users/jasons/PycharmProjects/pythonProject/venv/static/jedburghlogo_webicon.png",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Apply styles

streamlit_style = """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono&display=swap');

    html, body, [class*="css"] {
        font-family: 'IBM Plex Mono', monospace;
        background: #4F5223;
    }
    .custom-title {
        overflow: hidden;
        color: white;
        font-size: 1.3em;
        animation: typewriter 4s steps(50) 1s both;
        white-space: nowrap;
        padding-bottom: 50px;
    }
    @keyframes typewriter {
        0% {
            width: 0;
        }
        100% {
            width: 100%;
        }
    }
    </style>
"""

st.markdown(streamlit_style, unsafe_allow_html=True)
st.markdown('<h1 class="custom-title">WELCOME, AIRMAN ALLIE</h1>', unsafe_allow_html=True)

# If 'custom_title' is not in st.session_state, assign a default title
if 'custom_title' not in st.session_state:
    st.session_state.custom_title = "WELCOME, AIRMAN ALLIE"
app_name = st.session_state.custom_title

# Create subheader
st.markdown("""
    <div style='border: 2px solid white; padding: 10px;'>
        <h2 style='margin: 0; font-size: 14px; padding: 1em; font-family: 'IBM Plex Mono', monospace;'>Hamza, Cafe Owner</h2>
    </div>
""", unsafe_allow_html=True)




# Set default task
if 'selected_task' not in st.session_state:
    st.session_state.selected_task = task_selection[0]
st.session_state.selected_task = st.sidebar.radio("Select Task", task_selection)

# Initialize the Translator
translator = Translator(to_lang="en", from_lang="ar")

# Initialize chat history in session state
if 'hst_chat' not in st.session_state:
    st.session_state.hst_chat = []
if 'hst_chat_time' not in st.session_state:
    st.session_state.hst_chat_time = []

# Get user input
user_prompt = st.text_input("Start your chat (in Arabic):")
btn_enter = st.button("Enter")

# When 'Enter' button is clicked
if btn_enter:
    # Get the current timestamp
    current_time = datetime.now()

    # Add user's message and timestamp to chat history
    st.session_state.hst_chat.append({"role": "user", "content": user_prompt})
    st.session_state.hst_chat_time.append(current_time)

    # Load specific words from tcv.txt file
    with open("/Users/jasons/PycharmProjects/pythonProject/venv/tcv.txt", "r", encoding="utf-8") as file:
        specific_words = [word.strip() for word in file.readlines()]

    # Check if user's input has any of the specific words
    # If yes, play ding sound
    user_input_words = user_prompt.split()
    matching_words = set(specific_words).intersection(user_input_words)
    if matching_words:
        ding_sound_path = "/Users/jasons/PycharmProjects/pythonProject/venv/audio/tcv_match.mp3"
        ding_sound = AudioSegment.from_file(ding_sound_path)
        play(ding_sound)

    MAX_TOKENS = 500
    MAX_TOKENS_PER_MESSAGE = 50
    # Prepare the conversation for the chat model
    conversation = [
        {"role": "assistant", "content": initial_context[st.session_state.selected_task]},
    ] + st.session_state.hst_chat

    # Calculate the total number of tokens in the conversation
    total_tokens = sum(len(message['content'].split()) for message in conversation)

    # Check if the total tokens exceed the maximum allowed limit
    if total_tokens > MAX_TOKENS:
        # Remove messages until the total tokens is below the limit
        excess_tokens = total_tokens - MAX_TOKENS
        removed_tokens = 0
        removed_messages = 0

        # Iterate through the conversation messages from the beginning
        for i in range(len(conversation) - 1, -1, -1):
            message_tokens = len(conversation[i]['content'].split())
            if removed_tokens + message_tokens <= excess_tokens:
                # Remove the entire message
                removed_tokens += message_tokens
                removed_messages += 1
            else:
                # Remove a portion of the message
                tokens_to_remove = excess_tokens - removed_tokens
                conversation[i]['content'] = ' '.join(conversation[i]['content'].split()[:-tokens_to_remove])
                break

        # Remove the excess messages from the conversation
        conversation = conversation[:-removed_messages]

        # Split messages into multiple parts if they exceed the maximum tokens per message
        split_conversation = []
        current_message = {"role": conversation[0]["role"], "content": ""}
        for message in conversation[1:]:
            tokens_in_message = len(message["content"].split())
            if len(current_message["content"].split()) + tokens_in_message > MAX_TOKENS_PER_MESSAGE:
                split_conversation.append(current_message)
                current_message = {"role": message["role"], "content": message["content"]}
            else:
                current_message["content"] += " " + message["content"]

        if current_message["content"]:
            split_conversation.append(current_message)

        # Use OpenAI API to get a response from the chat model
        responses = []
        for split_message in split_conversation:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[split_message],
                max_tokens=MAX_TOKENS_PER_MESSAGE,
                n=1
            )
            responses.append(response['choices'][0]['message']['content'])

        # Add assistant's response to the chat history
        for response in responses:
            assistant_response = response
            st.session_state.hst_chat.append({"role": "assistant", "content": assistant_response})
            st.session_state.hst_chat_time.append(datetime.now())
    else:
        # Use OpenAI API to get a response from the chat model
        return_openai = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=conversation,
            max_tokens=MAX_TOKENS,
            n=1
        )

        # Add assistant's response to the chat history
        assistant_response = return_openai['choices'][0]['message']['content']
        st.session_state.hst_chat.append({"role": "assistant", "content": assistant_response})
        st.session_state.hst_chat_time.append(datetime.now())




# Display chat history
if st.session_state.hst_chat:
    for i in range(len(st.session_state.hst_chat)):
        if i % 2 == 0:
            # msg("You: " + st.session_state.hst_chat[i]['content'], is_user=True)
            st.markdown(f"<div style='text-align: left; color: black; background-color: rgba(206, 187, 163, 0.5); '>You: {st.session_state.hst_chat[i]['content']}</div>", unsafe_allow_html=True)
        else:
            # msg(st.session_state.selected_task + ": " + st.session_state.hst_chat[i]['content'])
            st.markdown(f"<div style='text-align: left; color: black; background-color: rgba(206, 187, 163, 1.0);'>{st.session_state.selected_task}: {st.session_state.hst_chat[i]['content']}</div>", unsafe_allow_html=True)


    # Translation button for user input
        if i % 2 == 0:
            translation_expander = st.expander("Show User Translation")
            with translation_expander:
                translation_result = translator.translate(st.session_state.hst_chat[i]['content'])
                if isinstance(translation_result, str):
                    translation = translation_result
                else:
                    translation = translation_result.text
                st.write(translation)
        # Translation button for assistant responses
        else:
            translation_expander = st.expander("Show Assistant Translation")
            with translation_expander:
                translation_result = translator.translate(st.session_state.hst_chat[i]['content'])
                if isinstance(translation_result, str):
                    translation = translation_result
                else:
                    translation = translation_result.text
                st.write(translation)

    # If chat history exists, show the 'Save & Export' button
    btn_save = st.button("Save & Export")
    if btn_save:
        # Create a Word document with chat history
        doc = Document()

        # Add the custom title with date and time to the document
        custom_title = f"{st.session_state.custom_title} - {datetime.now().strftime('%m/%d/%Y %I:%M:%S %p')}"
        doc.add_paragraph(custom_title)
        doc.add_paragraph("")

        # Calculate the total duration
        total_duration = st.session_state.hst_chat_time[-1] - st.session_state.hst_chat_time[0]

        # Add the total duration to the document
        doc.add_paragraph(f"Total Duration: {total_duration}")

        # Add the custom title, task selection and initial context to the document
        doc.add_paragraph(f"Custom Title: {st.session_state.custom_title}")
        doc.add_paragraph(f"Task Selection: {st.session_state.selected_task}")
        doc.add_paragraph(f"Initial Context: {initial_context[st.session_state.selected_task]}")
        doc.add_paragraph("")

        # Add the chat history to the document
        for message in st.session_state.hst_chat:
            doc.add_paragraph(f"{message['role']}: {message['content']}")

        # Save the Document into memory
        f = io.BytesIO()
        doc.save(f)

        # Format current date and time
        now = datetime.now()
        date_time = now.strftime("%m/%d/%Y %I:%M:%S %p")

        # Append date and time to the file name
        f.name = st.session_state.custom_title + "_" + date_time + '.docx'
        f.seek(0)

        # Download button for chat history Word document
        st.download_button(
            label="Download chat history",
            data=f,
            file_name=f.name,
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )