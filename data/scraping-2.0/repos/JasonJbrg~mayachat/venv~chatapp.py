import openai
import streamlit as st
from streamlit_chat import message as msg
from translate import Translator
from pydub import AudioSegment
from pydub.playback import play
import os
from dotenv import load_dotenv
import json
import docx
import io
from docx import Document
from datetime import datetime
import random

# Load environment variables from .env file
dotenv_path = "PycharmProjects/.env"
load_dotenv(dotenv_path)

# Read the config.json file
with open("venv/config.json") as file:
    config = json.load(file)



# Extract the values from the config dictionary
task_selection = config["task_selection"]
initial_context = {
    task: f"{config['initial_context'][task]} Please provide brief and concise responses."
    for task in task_selection
}
greetings = config["greetings"]

load_dotenv('/Users/jasons/PycharmProjects/pythonProject/PycharmProjects/.env')

# Set OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")

# Set page configuration
st.set_page_config(
    page_title="Maya Lingo",
    layout="wide",    
)

hide_streamlit_style = """
            <style>
            @import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono&display=swap');
            
            .css-uf99v8 {
                font-family: 'IBM Plex Mono', monospace;
                background-color: #4F5223;
               
            }

            .stChatFloatingInputContainer.css-usj992.ehod42b2 {
                font-family: 'IBM Plex Mono', monospace;
                background-color: #4F5223;
            }
    
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            header {visibility: hidden;}
            </style>
            """
st.markdown(hide_streamlit_style, unsafe_allow_html=True) 

# Custom title
st.markdown("""
    <div style="
        background-color: black;
        padding: 5px;
        font-family: 'IBM Plex Mono', monospace;
        font-size: 35px;
        color: white;
        margin-top: -90px;
        width: 100%;
        font-weight: bold;
    ">
        <p style="margin: 0; text-align: left;">Welcome,</p>
        <p style="margin: 0; text-align: left;">Airman Allie</p>
    </div>
    """, unsafe_allow_html=True)




# Add a default option to the languages dictionary
languages = {'Select Language': ''}

languages.update({
    'Chinese': 'zh',
    'Russian': 'ru',
    'Arabic': 'ar',
    'Japanese': 'ja',
    'Farsi': 'fa',
    'Spanish': 'es',
    'German': 'de',
    'Levantine Arabic': 'apc'  # ISO 639-3 code for Levantine Arabic
})

# Define default values for selected_language and selected_task
selected_language = 'Select Language'
selected_task = 'Select Topic'


# Initialize new_message and return_openai to None
new_message = None
return_openai = None

# Get user input for language selection
selected_language = st.selectbox("", list(languages.keys()), key='language_selection')

if selected_language != 'Select Language':
    # Initialize the Translator with the selected language
    translator = Translator(to_lang="en", from_lang=languages[selected_language])

    # Initialize two Translator objects with appropriate language settings
    translator_to_en = Translator(from_lang=languages[selected_language], to_lang="en")
    translator_from_en = Translator(from_lang="en", to_lang=languages[selected_language])

    # Add a default option to the task_selection list
    task_selection = ['Select Topic'] + task_selection

    # Get user input for task selection
    selected_task = st.selectbox(" ", task_selection, key='task_selection')

    # Only update the selected task in session state if a task is selected
    if selected_task != 'Select Topic':
        st.session_state.selected_task = selected_task

        # Only proceed if a task is selected and the chat history is empty
        if not st.session_state.hst_chat:
            # Update the selected task in session state
            st.session_state.selected_task = selected_task
            # Choose a random greeting for the selected task
            greeting = random.choice(greetings[selected_task])
            # Translate the greeting to the target language using translator_from_en
            greeting_translated = translator_from_en.translate(greeting)
            st.session_state.hst_chat.append({"role": "assistant", "content": greeting_translated})
            st.session_state.hst_chat_time.append(datetime.now())
    
# Get user input
if 'selected_task' in st.session_state:
    prompt = st.chat_input("Say something")
    if prompt:
        new_message = {"role": "user", "content": prompt}
            
    # Initialize conversation in session state if not already present
    if 'conversation' not in st.session_state:
        st.session_state.conversation = []
    
    # Check if a new message was submitted
    if new_message is not None:
        # Add user's original response to the chat history
        st.session_state.hst_chat.append(new_message)
        st.session_state.hst_chat_time.append(datetime.now())
    
        # Add user's response to the conversation
        st.session_state.conversation.append(new_message)
    
        # Only generate a response if the last message was from the user
        if len(st.session_state.hst_chat) >= 2 and st.session_state.hst_chat[-2]["role"] == "user":
            # Use OpenAI API to get a response from the chat model
            return_openai = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=st.session_state.conversation,
                max_tokens=MAX_TOKENS,
                n=1
            )
    
            # Add assistant's response to the chat history
            if return_openai:
                assistant_response = return_openai['choices'][0]['message']['content']
                st.session_state.hst_chat.append({"role": "assistant", "content": assistant_response})
                st.session_state.hst_chat_time.append(datetime.now())






# Add a default option to the task_selection list
task_selection = ['Select Topic'] + task_selection

# Initialize chat history in session state if not already present
if 'hst_chat' not in st.session_state:
    st.session_state.hst_chat = []
if 'hst_chat_time' not in st.session_state:
    st.session_state.hst_chat_time = []

# Only proceed if a task is selected and the chat history is empty
if selected_task != 'Select Topic' and not st.session_state.hst_chat:
    # Update the selected task in session state
    st.session_state.selected_task = selected_task
    # Choose a random greeting for the selected task
    greeting = random.choice(greetings[selected_task])
    # Translate the greeting to the target language using translator_from_en
    greeting_translated = translator_from_en.translate(greeting)
    st.session_state.hst_chat.append({"role": "assistant", "content": greeting_translated})
    st.session_state.hst_chat_time.append(datetime.now())


# Update the selected task in session state
st.session_state.selected_task = selected_task






MAX_TOKENS = 500
MAX_TOKENS_PER_MESSAGE = 50

# Define a function to get the initial context
def get_initial_context(task):
    if task is not None and task in initial_context:
        return initial_context[task]
    else:
        return "Please select a task."

# Initialize conversation to an empty list
conversation = []
        
# Prepare the conversation for the chat model
if 'selected_task' in st.session_state and st.session_state.selected_task is not None and st.session_state.selected_task in initial_context:
    conversation = [
        {"role": "assistant", "content": initial_context[st.session_state.selected_task]},
    ] + st.session_state.hst_chat
else:
    # Handle case where st.session_state.selected_task is None or does not exist in initial_context
    conversation = [
        {"role": "assistant", "content": "Please select a valid task."},
    ] + st.session_state.hst_chat


# Only generate a response if the last message was from the user
if conversation and conversation[-1]["role"] == "user":
    # Use OpenAI API to get a response from the chat model
    return_openai = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=conversation,
        max_tokens=MAX_TOKENS,
        n=1
    )

    # Calculate the total number of tokens in the conversation
    total_tokens = sum(len(message['content'].split()) for message in conversation)

    # Check if the total tokens exceed the maximum allowed limit
    if total_tokens > MAX_TOKENS:
        # Remove messages until the total tokens is below the limit
        excess_tokens = total_tokens - MAX_TOKENS
        removed_tokens = 0
        removed_messages = 0

        # Iterate through the conversation messages from the beginning
        for i in range(len(conversation) - 1, -1, -1):
            message_tokens = len(conversation[i]['content'].split())
            if removed_tokens + message_tokens <= excess_tokens:
                # Remove the entire message
                removed_tokens += message_tokens
                removed_messages += 1
            else:
                # Remove a portion of the message
                tokens_to_remove = excess_tokens - removed_tokens
                conversation[i]['content'] = ' '.join(conversation[i]['content'].split()[:-tokens_to_remove])
                break

        # Remove the excess messages from the conversation
        conversation = conversation[:-removed_messages]

        # Split messages into multiple parts if they exceed the maximum tokens per message
        split_conversation = []
        current_message = {"role": conversation[0]["role"], "content": ""}
        for message in conversation[1:]:
            tokens_in_message = len(message["content"].split())
            if len(current_message["content"].split()) + tokens_in_message > MAX_TOKENS_PER_MESSAGE:
                split_conversation.append(current_message)
                current_message = {"role": message["role"], "content": message["content"]}
            else:
                current_message["content"] += " " + message["content"]

        if current_message["content"]:
            split_conversation.append(current_message)


        # Use OpenAI API to get a response from the chat model
        responses = []
        for split_message in split_conversation:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[split_message],
                max_tokens=MAX_TOKENS_PER_MESSAGE,
                n=1
            )
            responses.append(response['choices'][0]['message']['content'])

        # Add assistant's response to the chat history
        for response in responses:
            assistant_response = response
            st.session_state.hst_chat.append({"role": "assistant", "content": assistant_response})
            st.session_state.hst_chat_time.append(datetime.now())
    else:
        # Use OpenAI API to get a response from the chat model
        return_openai = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=conversation,
            max_tokens=MAX_TOKENS,
            n=1
        )

        # Add assistant's response to the chat history
        assistant_response = return_openai['choices'][0]['message']['content']
        st.session_state.hst_chat.append({"role": "assistant", "content": assistant_response})
        st.session_state.hst_chat_time.append(datetime.now())

    

# Display chat history
if st.session_state.hst_chat:
    for i in range(len(st.session_state.hst_chat)):
        if st.session_state.hst_chat[i]["role"] == "user":
            st.markdown(
                f"<div style='text-align: left; color: black; background-color: rgba(206, 187, 163, 0.5); '>You: {st.session_state.hst_chat[i]['content']}</div>",
                unsafe_allow_html=True)
        elif st.session_state.hst_chat[i]["role"] == "assistant":
            st.markdown(
                f"<div style='text-align: left; color: black; background-color: rgba(206, 187, 163, 1.0);'>{st.session_state.selected_task}: {st.session_state.hst_chat[i]['content']}</div>",
                unsafe_allow_html=True)

        # Translation expander for user input
        if i % 2 == 0:
            translation_expander = st.expander("Show User Translation", expanded=False)
            with translation_expander:
                # Use translator_to_en for user's messages
                translation_result = translator_to_en.translate(st.session_state.hst_chat[i]['content'])
                st.write(translation_result)

        # Translation expander for assistant responses
        else:
            translation_expander = st.expander("Show Assistant Translation")
            with translation_expander:
                # Use translator_to_en for assistant's responses
                # We are assuming that the assistant's responses are not in English. 
                # If they are in English, you do not need to translate them. 
                translation_result = translator_to_en.translate(st.session_state.hst_chat[i]['content'])
                st.write(translation_result)


# If chat history exists, show the 'Save & Export' button
if st.session_state.hst_chat:
    btn_save = st.button("Save & Export")
    if btn_save:
        # Create a Word document with chat history
        doc = Document()

        # Add the current date and time to the document
        doc.add_paragraph(datetime.now().strftime('%m/%d/%Y %I:%M:%S %p'))

        # Calculate the total duration
        total_duration = st.session_state.hst_chat_time[-1] - st.session_state.hst_chat_time[0]

        # Add the total duration to the document
        doc.add_paragraph(f"Total Duration: {total_duration}")

        # Add the chat history to the document
        for message in st.session_state.hst_chat:
            doc.add_paragraph(f"{message['role']}: {message['content']}")

        # Save the Document into memory
        f = io.BytesIO()
        doc.save(f)

        # Format current date and time
        now = datetime.now()
        date_time = now.strftime("%m%d%Y_%H%M%S")  # Changed format to remove slashes and colons

        # Append date and time to the file name
        f.name = "Chat_History_" + date_time + '.docx'  # Changed to a static string "Chat_History_"
        f.seek(0)

        # Download button for chat history Word document
        st.download_button(
            label="Download chat history",
            data=f,
            file_name=f.name,
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
