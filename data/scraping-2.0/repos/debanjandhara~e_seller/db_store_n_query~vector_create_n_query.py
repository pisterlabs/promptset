# ------------

from dotenv import load_dotenv
import os
import openai

import shutil

# -------------

import magic
import docx
import json
import PyPDF2


# -------------

load_dotenv()

openai.api_key = os.getenv("OPENAI_API_KEY")

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.llms import OpenAI
from langchain.chat_models import ChatOpenAI
from langchain.chains.question_answering import load_qa_chain
from langchain.callbacks import get_openai_callback
import os
import openai

def read_file_to_string(file_path):
    try:
        # Initialize an empty string to store the file contents
        file_contents = ""

        # Open the file and read its contents
        with open(file_path, 'r', encoding='utf-8') as file:
            file_contents = file.read()

        return file_contents
    except FileNotFoundError:
        return "File not found."


def create_vector(content, vector_folder_name):

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
        )
    chunks = text_splitter.split_text(text=content)
    
    embeddings = OpenAIEmbeddings()
    
    print(vector_folder_name)
    
    VectorStore = FAISS.from_texts(chunks, embedding=embeddings)
    
    # os.makedirs(os.path.dirname(f"{vector_name}.pkl"), exist_ok=True)
    
    # directory = os.path.dirname(vector_name)
    # if not os.path.exists(directory):
    #     os.makedirs(directory)
    
    if os.path.exists(vector_folder_name):
        try:
            shutil.rmtree(vector_folder_name)
            print(f"Folder '{vector_folder_name}' and its contents deleted successfully.")
        except Exception as e:
            print(f"Error deleting folder '{vector_folder_name}': {e}")
    else:
        print(f"Folder '{vector_folder_name}' does not exist.")
        
    VectorStore.save_local(vector_folder_name)
    
    # with open(f"{vector_name}.pkl", "wb") as f:
    #     print("Dumping in ",f"{vector_name}.pkl")
    #     pickle.dump(VectorStore, f)

    return vector_folder_name


def query_from_vector(query, user_id):

    vector_folder_name = f"data/{user_id}/merged_vector"
    
    if not os.path.exists(vector_folder_name):
        response_for_empty_folder = "Sorry ! You need to upload documents first to be able to chat with it..."
        return response_for_empty_folder


    # if os.path.exists(f"{vector_name}.pkl"):
    #     with open(f"{vector_name}.pkl", "rb") as f:
    #         VectorStore = pickle.load(f)
    
    embeddings = OpenAIEmbeddings()    
    VectorStore = FAISS.load_local(vector_folder_name, embeddings=embeddings)


    # query = "what colour is the sky"

    docs = VectorStore.similarity_search(query=query, k=3)

    # llm = OpenAI()
    llm = ChatOpenAI(model_name='gpt-3.5-turbo')
    chain = load_qa_chain(llm=llm, chain_type="stuff")
    with get_openai_callback() as cb:
        response = chain.run(input_documents=docs, question=query)
        print(cb)
    # print("\n\nResponse : ",response)

    return response


def read_document(file_path):
    
    # Get the file extension
    _, file_extension = os.path.splitext(file_path)

    # Convert the extension to lowercase for case-insensitive comparison
    file_extension = file_extension.lower()

    # Check file type based on extension
    if file_extension == '.txt':
        print(f"The file {file_path} is a Text file.")
        with open(file_path, 'r', encoding='utf-8') as txt_file:
            content = txt_file.read()
    elif file_extension == '.docx':
        print(f"The file {file_path} is a Word document.")
        doc = docx.Document(file_path)
        content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    elif file_extension == '.pdf':
        print(f"The file {file_path} is a PDF document.")
        # Handle PDF files here
        pdfFileObj = open(file_path, 'rb')
    
        # creating a pdf reader object
        pdfReader = PyPDF2.PdfReader(pdfFileObj)
    
        # printing number of pages in pdf file
        # print(len(pdfReader.pages))

        # creating a page object
        pageObj = pdfReader.pages[0]

        # extracting text from page
        content = pageObj.extract_text()

        # closing the pdf file object
        pdfFileObj.close()
    elif file_extension == '.json':
        print(f"The file {file_path} is a JSON file.")
        with open(file_path, 'r') as json_file:
            content = json.load(json_file)
    else:
        print(f"The file type of {file_path} is not recognized.")
    

    
    # try:
    #     file_type = magic.Magic()
    #     detected_type = file_type.from_file(file_path)

    #     content = ''

    #     # Read content based on the detected file type
    #     if detected_type.startswith('Microsoft Word'):
    #         if file_path.endswith('.docx'):
    #             # For .docx files
    #             doc = docx.Document(file_path)
    #             content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    #     elif detected_type.startswith('JSON'):
    #         # For .json files
    #         with open(file_path, 'r') as json_file:
    #             content = json.load(json_file)
    #     elif detected_type.startswith('ASCII') or detected_type.startswith('text/plain') or detected_type.startswith('UTF-8 Unicode text'):
    #         # For plain text (.txt) files
    #         with open(file_path, 'r', encoding='utf-8') as txt_file:
    #             content = txt_file.read()
    #     elif detected_type.startswith('PDF'):
    #         # Handle PDF files here
    #         pdfFileObj = open(file_path, 'rb')
        
    #         # creating a pdf reader object
    #         pdfReader = PyPDF2.PdfReader(pdfFileObj)
    
    #         # printing number of pages in pdf file
    #         # print(len(pdfReader.pages))

    #         # creating a page object
    #         pageObj = pdfReader.pages[0]

    #         # extracting text from page
    #         content = pageObj.extract_text()

    #         # closing the pdf file object
    #         pdfFileObj.close()

    #     else:
    #         print(f"Unsupported file type: {detected_type}")
    # except FileNotFoundError:
    #     print(f"File not found: {file_path}")
    # except Exception as e:
    #     print(f"An error occurred: {e}")

    return content

def merge_db(user_id):
    vector_base_folder = f"data/{user_id}/vectors"
    final_folder = f"data/{user_id}/merged_vector"
    print(f"----------\n\n{vector_base_folder}\n\n{final_folder}\n\n--------")
    embeddings = OpenAIEmbeddings()
    all_items  = os.listdir(vector_base_folder)
    folders = [item for item in all_items if os.path.isdir(os.path.join(vector_base_folder, item))]
    print(len(folders))
    if len(folders)==1:
        VectorStore1 = FAISS.load_local(f"{vector_base_folder}/{folders[0]}", embeddings=embeddings)
        VectorStore1.save_local(final_folder)
        # return "Merged - Single"
        return "success"
    VectorStore1 = FAISS.load_local(f"{vector_base_folder}/{folders[0]}", embeddings=embeddings)
    VectorStore2 = FAISS.load_local(f"{vector_base_folder}/{folders[1]}", embeddings=embeddings)
    VectorStore2.merge_from(VectorStore1)
    VectorStore2.save_local(final_folder)
    for i in range(1,len(folders)):
        VectorStore1 = FAISS.load_local(final_folder, embeddings=embeddings)
        VectorStore2 = FAISS.load_local(f"{vector_base_folder}/{folders[i]}", embeddings=embeddings)
        VectorStore2.merge_from(VectorStore1)
        VectorStore2.save_local(final_folder)

    # response = "Merged - Multiple"
    response = "success"

    return response


def delete_folder(folder_path):
    try:
        # Check if the folder exists
        if os.path.exists(folder_path):
            # Remove the folder and its contents
            shutil.rmtree(folder_path)
            print(f"The folder '{folder_path}' has been successfully deleted.")
        else:
            print(f"The folder '{folder_path}' does not exist.")
    except Exception as e:
        print(f"An error occurred: {e}")
    response = "success"

    return response




# print(query_from_vector("give me a summary about networking", r"data/uuid2502/merged"))

# vector_base_folder = "data/uuid2502/vectors"

# filename = "test"

# vector_folder_name = f"{vector_base_folder}/{filename}"

# print(create_vector(read_document(r"data/uuid2502/uploads/sample.txt"), vector_folder_name))
# print(query_from_vector("give me summary about operating system", vector_folder_name))

# print(read_document(r"data/uuid2502/uploads/docx_file.docx"))
# vector_base_folder = r"data/uuid2502/vectors"

# vector_folder_name = f"{vector_base_folder}/{filename}"

# final_folder = r"data/uuid2502/merged_vector"

# merge_db(vector_base_folder, final_folder)----------

# query_from_vector("summary about networks", final_folder)----------------

# print(read_document(r"data/uuid2502/apple.com.filtered.json"))------------

# create_vector(read_document(r"data/uuid2502/apple.com.filtered.json"), vector_base_folder)--------------

# print(merge_db("uuid2300"))
