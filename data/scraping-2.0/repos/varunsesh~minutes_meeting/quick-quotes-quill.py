from openai import OpenAI
from dotenv import load_dotenv
from record_audio import AudioRecorder, RecordingApp
from datetime import datetime
import os
from docx import Document



def process_segments(segments, extraction_function):
    """Process each segment and combine the results."""
    print("processing segments...")
    combined_result = ""
    for segment in segments:
        result = extraction_function(segment)
        combined_result += result + "\n"
    return combined_result.strip()


def meeting_minutes(transcription):
    segments = split_transcript(transcription)
    segment_summaries = [summarize_segment(segment) for segment in segments]
    final_meeting_summary = final_summary(segment_summaries)
    return final_meeting_summary#truncate_to_word_limit(final_meeting_summary, 5000)


def truncate_to_word_limit(content, word_limit):
    print("truncating to limit...")
    words = content.split()
    if len(words) > word_limit:
        return ' '.join(words[:word_limit])
    return content



def save_as_docx(summary, filename, meeting_date):
    new_doc = Document()

    # Add the new title, date, and summary
    new_doc.add_paragraph("Meeting Minutes", style='Heading 1')
    new_doc.add_paragraph(meeting_date, style='Normal')
    new_doc.add_paragraph(summary, style='Normal')
    new_doc.add_paragraph()  # Line break

    # If the document exists, append its content to the new document
    if os.path.exists(filename):
        old_doc = Document(filename)
        for element in old_doc.element.body:
            new_doc.element.body.append(element)

    # Save the updated content in the filename
    new_doc.save(filename)



def split_transcript(transcription, max_length=4000):
    words = transcription.split()
    segments = []
    current_segment = []

    for word in words:
        if len(' '.join(current_segment + [word])) > max_length:
            segments.append(' '.join(current_segment))
            current_segment = [word]
        else:
            current_segment.append(word)

    segments.append(' '.join(current_segment))
    return segments

#using gpt4
def summarize_segment(segment):
    print("Summarizing segment...")
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0.7,
        messages=[
            {
                "role": "system",
                "content": "You are an AI that summarizes text. Please provide a comprehensive summary of the following segment, include separate sections for summary, action items and key points:"
            },
            {
                "role": "user",
                "content": segment
            }
        ]
    )
    print(response.choices[0].message.content)
    return response.choices[0].message.content

def final_summary(segments):
    i = 1
    combined_summary = ""
    # for segment in segments:
    #     title = "\n" + "\n" + "Summary Section: " + str(i) + "\n"
    #     combined_summary += title
    #     combined_summary += segment
    #     i+=1
    combined_summary = summarize_segment(' '.join(segments))
    return combined_summary



if __name__=="__main__":
    load_dotenv()
    
    client = OpenAI(
        # defaults to os.environ.get("OPENAI_API_KEY")
        # api_key="My API Key",
    )

    # Get the current date and time
    now = datetime.now()

# Format the date and time as "Month Day, Year"
    formatted_date = now.strftime("%B %d, %Y")
    recorder = AudioRecorder()
    app = RecordingApp(recorder)
    app.mainloop()
    if(app.get_permission()):
        transcription = app.get_transcript()
        summary = meeting_minutes(transcription)
        doc_name = 'minutes_' + now.strftime("%B_%d_%Y") + '.docx'
        save_as_docx(summary, doc_name, formatted_date)
    
