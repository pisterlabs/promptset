#Framework supporting MLOps Apps
import streamlit as st 
#Large Language Model Library
from langchain.llms import OpenAI
import docx
import datetime
import time
import base64


#Front-End
with st.sidebar:
    st.markdown("<h1 style='text-align:center;font-family:Georgia'>💸 MoneyMentor </h1>",unsafe_allow_html=True)
    st.markdown("This app is designed to assist Entrepreneurs and Startup founders in navigating the financial and legal aspects of \n\
                developing their businesses. It provides valuable information and tools related to startup financing, addressing \n\
                various Legal considerations such as Locale and Type of Funding Sought using the power of Large Language Models by Openai and Langchain.")
    st.markdown("<h2 style='text-align:center;font-family:Georgia'>Features</h1>",unsafe_allow_html=True)
    st.markdown(" - 🤑 MoneyMentor FinanceGPT - This Bot is ready to answer your business needs")
    st.markdown(" - 🧾 Financial Report Generator - This tool allows you to generate Financial Reports for your Start Up ")
    st.markdown("-------")
    openai_api_key = st.text_input('Enter OpenAI API Key', type='password')
    st.markdown("-------")
    st.markdown("<h1 style='text-align:center;font-family:Georgia'>🧾 Financial Report Generator</h1>",unsafe_allow_html=True)
    start_up_name=st.text_input("What is the name of your Start Up")
    start_up_description=st.text_input("Please describe what your start up is about and how you intend to generate revenue")
    country = st.selectbox(
    'Where is your Start Up is based?',
    ("Afghanistan", "Albania", "Algeria", "Andorra", "Angola", "Antigua and Barbuda", "Argentina", "Armenia", "Australia",
    "Austria", "Azerbaijan", "Bahamas", "Bahrain", "Bangladesh", "Barbados", "Belarus", "Belgium", "Belize", "Benin",
    "Bhutan", "Bolivia", "Bosnia and Herzegovina", "Botswana", "Brazil", "Brunei", "Bulgaria", "Burkina Faso", "Burundi",
    "Cabo Verde", "Cambodia", "Cameroon", "Canada", "Central African Republic", "Chad", "Chile", "China", "Colombia",
    "Comoros", "Congo (Congo-Brazzaville)", "Costa Rica", "Croatia", "Cuba", "Cyprus", "Czechia (Czech Republic)",
    "Democratic Republic of the Congo (Congo-Kinshasa)", "Denmark", "Djibouti", "Dominica", "Dominican Republic",
    "Ecuador", "Egypt", "El Salvador", "Equatorial Guinea", "Eritrea", "Estonia", "Eswatini (fmr. Swaziland)", "Ethiopia",
    "Fiji", "Finland", "France", "Gabon", "Gambia", "Georgia", "Germany", "Ghana", "Greece", "Grenada", "Guatemala",
    "Guinea", "Guinea-Bissau", "Guyana", "Haiti", "Holy See", "Honduras", "Hungary", "Iceland", "India", "Indonesia",
    "Iran", "Iraq", "Ireland", "Israel", "Italy", "Ivory Coast", "Jamaica", "Japan", "Jordan", "Kazakhstan", "Kenya",
    "Kiribati", "Kuwait", "Kyrgyzstan", "Laos", "Latvia", "Lebanon", "Lesotho", "Liberia", "Libya", "Liechtenstein",
    "Lithuania", "Luxembourg", "Madagascar", "Malawi", "Malaysia", "Maldives", "Mali", "Malta", "Marshall Islands",
    "Mauritania", "Mauritius", "Mexico", "Micronesia", "Moldova", "Monaco", "Mongolia", "Montenegro", "Morocco", "Mozambique",
    "Myanmar (formerly Burma)", "Namibia", "Nauru", "Nepal", "Netherlands", "New Zealand", "Nicaragua", "Niger", "Nigeria",
    "North Korea", "North Macedonia (formerly Macedonia)", "Norway", "Oman", "Pakistan", "Palau", "Palestine State",
    "Panama", "Papua New Guinea", "Paraguay", "Peru", "Philippines", "Poland", "Portugal", "Qatar", "Romania", "Russia",
    "Rwanda", "Saint Kitts and Nevis", "Saint Lucia", "Saint Vincent and the Grenadines", "Samoa", "San Marino",
    "Sao Tome and Principe", "Saudi Arabia", "Senegal", "Serbia", "Seychelles", "Sierra Leone", "Singapore", "Slovakia",
    "Slovenia", "Solomon Islands", "Somalia", "South Africa", "South Korea", "South Sudan", "Spain", "Sri Lanka", "Sudan",
    "Suriname", "Sweden", "Switzerland", "Syria", "Tajikistan", "Tanzania", "Thailand", "Timor-Leste", "Togo", "Tonga",
    "Trinidad and Tobago", "Tunisia", "Turkey", "Turkmenistan", "Tuvalu", "Uganda", "Ukraine", "United Arab Emirates",
    "United Kingdom", "United States of America", "Uruguay", "Uzbekistan", "Vanuatu", "Venezuela", "Vietnam", "Yemen",
    "Zambia", "Zimbabwe"))
    sector = st.multiselect('What is your Start Up about', ["Technology and Software", "Healthcare and Biotechnology", "Agriculture and AgriTech", "E-commerce and Retail", "Fintech (Financial Technology)", "Food and Beverage", "Manufacturing and Industry", "Clean Energy and Sustainability", "Education and EdTech", "Transportation and Mobility", "Real Estate and Property Tech (PropTech)", "Entertainment and Media", "Travel and Tourism", "Social Impact and Nonprofits", "Space and Aerospace", "Fashion and Apparel", "Artificial Intelligence (AI) and Machine Learning (ML)", "LegalTech", "Blockchain and Cryptocurrency", "Sports and Fitness", "Gaming and Esports", "Cybersecurity", "AI in Healthcare (HealthTech)", "Supply Chain and Logistics", "Emerging Technologies"])
    funding = st.multiselect('What sort of Funding are you looking for ?', ["Angel Investment", "Venture Capital", "Seed Funding", "Series A Funding", "Series B Funding", "Series C Funding", "Crowdfunding", "Debt Financing", "Corporate Investment", "Government Grants", "Accelerator Programs", "Strategic Partnerships", "Initial Coin Offerings (ICOs)", "Initial Public Offerings (IPOs)", "Private Equity", "Convertible Notes", "Revenue-based Financing", "Equity Crowdfunding", "Strategic Investments"])
    st.markdown("-------")

    generatebutt=st.button("Generate Financial Report")

    

def generate_report(Company_name,country, report_date, Funding_text_summary=None,Legal_text_summary=None):
    
    doc = docx.Document()

    # Add Title Page followed by section summary
    doc.add_heading("Financial Report", 0)
    doc.add_paragraph('Authored By: MoneyMentor FinGPT LLM')
    doc.add_paragraph(f'Created On: {str(report_date)}')
    doc.add_paragraph(f'Created For: {Company_name}')
    doc.add_paragraph(f'Country based: {country}')
    doc.add_heading(f'Navigating the Intersection: A Comprehensive Guide to {Company_name} Financial and Legal Strategies')
    # Funding Strategies
    doc.add_heading('Funding Strategies')
    doc.add_paragraph(Funding_text_summary)
    # Legal Strategies
    doc.add_heading('Legal Strategies')
    doc.add_paragraph(Legal_text_summary)
    doc.save('Financial Report.docx')
    data = open('Financial Report.docx', "rb").read()
    encoded = base64.b64encode(data)
    decoded = base64.b64decode(encoded)

    st.download_button('Download Here', decoded, "Financial Report.docx")

def generate_response(input_text):
    llm = OpenAI(temperature=0.3, openai_api_key=openai_api_key)
    output=llm(input_text)
    return output




# App framework
st.markdown("<h1 style='text-align:justified;font-family:Georgia'>🤑Chatbot</h1>",unsafe_allow_html=True)
# User input field
# Initialize chat history


if "messages" not in st.session_state:
    st.session_state.messages = []
    st.session_state.messages.append({"role": "🤑", "content": "Hey there, I'm your MoneyMentor , here to advise you on how to kickstart \n\
                                      your Business. Ask literally anything but before that, enter your API Key on the side bar"})
# Display chat messages from history on app rerun
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        with st.spinner('Starting Bot ...'):
            st.markdown(message["content"])

if generatebutt:
    if openai_api_key.startswith('sk-'):
        date_today=datetime.date.today()
        Funding_generated_summary=generate_response(f"I'm currently in the process of exploring funding options for my startup named {start_up_name}, and I'd like to gather as much information as possible.The description of my start up is {start_up_description}.\n\
            I'm particularly interested in understanding the various funding sources available to early-stage startups like mine and any specific tips or considerations.I've selected the following {funding} options. \n\
                                    To begin,I'd appreciate an overview of the different types of funding sources that are accessible to my startups in {country} and related to {sector}.Moreover, I'd like to understand the eligibility requirements and criteria that startups typically need to meet for each funding source I've selected. This information will be invaluable as I evaluate which options align with \n\
                                    my startup's current stage and objectives.Preparing strong applications or pitches is crucial when seeking funding. Therefore, I would welcome any advice or tips on how to present a compelling case to potential investors or funding organizations.\n\
                                        Understanding what investors look for can significantly enhance my chances of securing the necessary funds. Networking is often a vital aspect of the fundraising process.\n\
                                        If you could provide strategies for building connections with potential investors or organizations that provide funding, I would greatly appreciate it. Insights into effective networking can be a game-changer.\n\
                                        Additionally, I'd like to be aware of common challenges or pitfalls that startups frequently encounter during the fundraising process.Learning from these experiences can help me avoid potential setbacks and navigate the process more effectively.\n\
                                        Lastly, timing and planning are critical considerations. Insights into when it's the right time to seek funding and how to plan for a successful fundraising campaign would be highly valuable. \n\
                                        If you could also share any relevant resources, articles, or additional advice on this topic, it would be greatly appreciated. Your assistance in this matter is of utmost importance to me as I embark on this funding journey for my startup. Put it in point form and complete each point and up to date specified information.")
        
        Legal_generated_summary=generate_response(f"I am seeking your legal expertise to guide me through the process of launching my startup called {start_up_name} in a specific {country}. I would appreciate comprehensive advice that covers all relevant legal requirements, regulations, and considerations unique to this jurisdiction and related to {sector}.\n\
                                    Please provide as much information as possible to ensure a successful and compliant startup launch in this country. Your insights are invaluable in navigating the legal landscape effectively. Put it in point form and complete each point and up to date specified information.")
        
        generate_report(start_up_name, country, date_today, Funding_text_summary=Funding_generated_summary,Legal_text_summary=Legal_generated_summary)
    else:
        st.warning('Please enter your OpenAI API key!', icon='⚠')
      
if (prompt := st.chat_input("What is up?")): 
    if  openai_api_key.startswith('sk-'):
        # Display user message in chat message container
        #user=🙂
        with st.chat_message("🙂"):
            st.markdown(prompt)
        # Add user message to chat history
        st.session_state.messages.append({"role": "🙂", "content": prompt})

        # Display assistant response in chat message container
        #assistant=🤑
    # Display assistant response in chat message container
        with st.chat_message("🤑"):
            message_placeholder = st.empty()
            full_response = ""
            with st.spinner('Wait for it...'):
                assistant_response = generate_response(prompt)

        for chunk in assistant_response.split():
            full_response += chunk + " "
            time.sleep(0.05)
            # Add a blinking cursor to simulate typing
            message_placeholder.markdown(full_response + "▌")
        message_placeholder.markdown(full_response)

        # Add assistant response to chat history
        st.session_state.messages.append({"role": "🤑", "content": full_response})
    
    else:
        st.warning('Please enter your OpenAI API key!', icon='⚠')



