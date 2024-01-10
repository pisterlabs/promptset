from sentence_transformers import SentenceTransformer, util
import docx as dx
from sentence_transformers import SentenceTransformer
from langchain.chains.llm import LLMChain
from langchain.document_loaders import Docx2txtLoader

from langchain.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from langchain.document_loaders import PyPDFLoader
from langchain.document_loaders import TextLoader
from langchain.schema.document import Document


import os
# from langchain.llms import AzureOpenAI
from langchain.chat_models.azure_openai import AzureChatOpenAI
from langchain.memory import ConversationBufferMemory
from langchain.text_splitter import RecursiveCharacterTextSplitter

# os.environ["OPENAI_API_TYPE"] = "azure"
# os.environ["OPENAI_API_BASE"] = "https://utterancesresource.openai.azure.com/"
# os.environ["OPENAI_API_KEY"] = "5ea3e8e59b8a418e9cc3c066f853b0c0"
# os.environ["OPENAI_API_VERSION"] = "2023-07-01-preview"

os.environ["OPENAI_API_KEY"]= 'e63ed695495543d58595fab4e27e4ff1'
os.environ['OPENAI_API_VERSION'] = '2023-07-01-preview'
os.environ['OPENAI_API_BASE'] = 'https://tv-llm-applications.openai.azure.com/'
os.environ['OPENAI_API_TYPE'] = 'azure'

def read_docx(file_path):
    doc = dx.Document(file_path)
    full_text = []

    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    return '\n'.join(full_text)


def pdf_query_updated(query, text_splitter, llm, query_options, memory):  

    documents_query = []
    for file in query_options:
        if file.endswith('.pdf'):
            pdf_path = './documents/' + file
            loader = PyPDFLoader(pdf_path)
            documents_query.extend(loader.load())
        elif file.endswith('.docx') or file.endswith('.doc'):
            doc_path = './documents/' + file
            print(doc_path,'*'*117)
            loader = Docx2txtLoader(file)
            documents_query.extend(loader.load())
        elif file.endswith('.txt'):
            text_path = './documents/' + file
            loader = TextLoader(text_path)
            documents_query.extend(loader.load())


    # docs=text_splitter.split_documents(documents_query)
    # # embeddings=HuggingFaceEmbeddings(model_name = 'sentence-transformers/all-MiniLM-L6-v2')
    # embeddings = OpenAIEmbeddings(deployment='ada-embed',
    #                               openai_api_key='e63ed695495543d58595fab4e27e4ff1',
    #                               openai_api_base= 'https://tv-llm-applications.openai.azure.com/',
    #                               openai_api_type="azure",
    #                               openai_api_version='2023-07-01-preview',
    #                               chunk_size=16)
    bi_encoder = SentenceTransformer('multi-qa-MiniLM-L6-cos-v1')
    bi_encoder.max_seq_length = 256     #Truncate long passages to 256 tokens
    top_k = 32                          #Number of passages we want to retrieve with the bi-encoder


    
    # document_search = FAISS.from_texts([t.page_content for t in docs], embeddings)


    template = """You are an AI having a conversation with a human.

    Given the following extracted parts of a long document and a question, create a final answer.
    And if you can't find the answer, strictly mention "Currently I am unable to answer this questions based on my knowledge. Please reach out to us - support@guardsman.com"

    {context}

    {chat_history}
    Human: {human_input}
    AI:"""

    prompt = PromptTemplate(input_variables=["chat_history", "human_input", "context"], template=template)
    chain = LLMChain(llm = llm, prompt = prompt,memory = memory)
    
    doc_path = "Guardsman Group FAQ.docx"
    text_from_docx = read_docx(doc_path)
    # model = SentenceTransformer('multi-qa-MiniLM-L6-cos-v1')
    # document = model.encode(text_from_docx)
    local_doc = Document(page_content=text_from_docx,metadata={})


    result, context_docs = llm_query(query, local_doc, chain)

    return result, context_docs


def llm_query(query, document_search, chain):
    
    # docs = document_search.similarity_search(query)
    chain_res = chain.predict(human_input = query, context = document_search).split('Human:')[0]
    return chain_res + '\n\n', document_search


# file_path = r"QandA.docx"
# memory = ConversationBufferMemory()
# query_options = [file_path]
# human_query = 'what leave policies do I have'
# text_splitter=RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
# llm = AzureOpenAI(deployment_name='gpt-0301', temperature = 0)
# memory = ConversationBufferMemory(memory_key="chat_history", input_key = 'human_input')

# response, context_docs = pdf_query_updated(query = human_query, text_splitter = text_splitter, llm = llm, query_options = query_options, memory = memory)