#!/usr/bin/env python3
import os                                                                           
import openai                                                                       
import PyPDF2
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tkinter as tk
from tkinter import filedialog

script_dir = os.path.dirname(os.path.abspath(__file__))
os.chdir(script_dir)

def select_file():
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
    return file_path                    

resume_file = select_file()

with open(resume_file, 'rb') as file:
    reader = PyPDF2.PdfReader(file)
    resume_contents = ""

    for page in range(len(reader.pages)):
        resume_contents += reader.pages[page].extract_text()

linkedin_url = input("Please enter the LinkedIn URL: ")                                                                        
response = requests.get(linkedin_url)                                                                                                                                                         
soup = BeautifulSoup(response.content, 'html.parser')                                                                                                                                                    
job_description_element = soup.find('div', class_='description__text')
#print(job_description_element)
job_description = job_description_element.get_text(strip=True)                                    
                                                                                       

openai.api_key = "your OpenAI API key here"                
                                                                                       
response = openai.ChatCompletion.create(                                                
    model="gpt-3.5-turbo",                
    messages=[
        {"role": "system", "content": "You are a job seeker writing a cover letter tailored to a specific a job position"},
        {"role": "user", "content": "Here's your resume" + resume_contents + "\n\nHere's the job description" + job_description + "\n\nPlease write the contents of a cover letter starting with 'Dear Hiring Manager'"},
    ],                                                         
    max_tokens=500,  # Adjust the max tokens as needed                              
    temperature=0.5,  # Adjust the temperature as needed                                  
)                                                                                   
                                                                                       
cover_letter_draft = response['choices'][0]['message']['content'].strip()
name = input("Name: ")
location = input("Your location ('City, State'): ")
email = input("Email: ")
phone = input("Phone number: ")
personal_url = input("Your personal LinkedIn URL: ")

doc = Document()
paragraph = doc.add_paragraph(f'{name}\n{location}\n{email} - {phone}\n{Personal_URL}\n\n')
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph(cover_letter_draft)
doc.save('Cover_Letter.docx')




                          
                                                                                       
