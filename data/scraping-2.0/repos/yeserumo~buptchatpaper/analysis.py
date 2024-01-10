import logging
import requests
from typing import List, Dict, Any, Optional, Mapping, Union
from utils.logger import MyLogger
from utils.util import extract_method, extract_conclusion, extract_introduction
import langchain
import pickle
from texfiles import TexFiles
from texfiles import read_file_without_comments
from texfiles import delete_image_table, delete_others
from chatvicuna import ChatVicuna
from langchain.text_splitter import CharacterTextSplitter
from langchain.cache import InMemoryCache  
from langchain.chains.summarize import load_summarize_chain
from langchain.chains import AnalyzeDocumentChain
from text import DownloadTools
from text import ArxivTools
from tenacity import retry, wait_fixed, stop_after_attempt
import arxiv
from utils.util import extract_all_section, get_section_name, merge_tex_files
from utils.backend import output_information, get_timestamp, construct_prompt, make_dir_if_not_exist
from paperInfo import paperInfo
from load_config import load_config
import time
import warnings
import os
import re
from tqdm import tqdm
from text_filter import TextFilter
import argparse
import openai
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from output_parser import OutputParser
from langchain.prompts import PromptTemplate
from utils.util import chat_with_llm_api
from utils.prompt import construct_query_for_generate_abstract, review_from_abstract, review_from_method, review_from_conclusion, construct_query_for_generate_introduction_part1, construct_query_for_generate_introduction_part2, construct_query_for_generate_introduction_part3
logging.basicConfig(level=logging.INFO)
warnings.filterwarnings("ignore")
logger = MyLogger().get_logger()
langchain.llm_cache = InMemoryCache()

def llm_chat_with_one_paper(llm, title: str, timestamp = '20230722102536') -> None:
    '''
    接收一篇论文的Introduction，返回改论文的概述并保存为本地txt文件
    '''
    # 1. 获取论文的introduction
    getTexFiles = TexFiles(title, timestamp)
    texFiles = getTexFiles()
    if len(texFiles) == 1:
        text = read_file_without_comments(texFiles[0])
        introduction = extract_introduction(text)
    else:
        image_pattern = r'\\begin\{figure\}(\[.*?\])?.*?\\end\{figure\}'
        match_file = [file for file in texFiles if any(file.split('/')[-1].startswith(prefix) for prefix in ['intro', 'Intro', 'int'])]
        introduction = read_file_without_comments(match_file[0])
        introduction = re.sub(image_pattern, '', introduction)
    text_spliter = CharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=20,
            length_function=len,
        )
    CHAIN_TYPE = 'map_reduce'
    summary_chain = load_summarize_chain(llm, chain_type=CHAIN_TYPE)
    summarize_document_chain = AnalyzeDocumentChain(   
        combine_docs_chain=summary_chain,
        text_spliter=text_spliter,
    )
    start_time = time.time()
    logger.info(f"Start summarizing the document: {title}")
    res = summarize_document_chain.run(introduction)
    end_time = time.time()
    logger.info(f"Summarizing the document: {title} takes {end_time - start_time} seconds.")
        
    # 将模型输出的introduction保存到本地
    output_information(res, os.path.join(os.path.dirname(os.path.abspath(__file__)), 'introduction'), title)
    return introduction, res

# Testing Finished!
def get_introduction(title: str, timestamp):
    getTexFiles = TexFiles(title, timestamp)
    texFiles = getTexFiles()
    filter = TextFilter(level=3)
    if len(texFiles) == 1:
        text = read_file_without_comments(texFiles[0])
        introduction = extract_introduction(text)
        if not introduction:
            section_list = extract_all_section(text)
            introduction_name = get_section_name(section_list)['Introduction']
            introduction = extract_introduction(text, introduction_name)
        filter_introduction = filter._call(introduction)
    else:
        merge_tex_files(texFiles) # 将所有tex文件暂时合并为一个，并保存到tmp.tex中
        text = read_file_without_comments(config['utils']['tmp_file'])
        introduction = extract_introduction(text)
        if not introduction:
            section_list = extract_all_section(text)
            introduction_name = get_section_name(section_list)['Introduction']
            introduction = extract_introduction(text, introduction_name)
        if len(introduction) < 100: # 需要转向其他tex文件，该文件的全部内容均为introduction
            match_file = [file for file in texFiles if any(file.split('/')[-1].startswith(prefix) for prefix in ['intro', 'Intro', 'int'])]
            introduction = read_file_without_comments(match_file[0])
            filter_introduction = filter._call(introduction)
        else: # 和只有一个tex文件处理相同，已经获取到完整的Introduction
            filter_introduction = filter._call(introduction)
    return filter_introduction

def get_method(title: str, timestamp):
    getTexFiles = TexFiles(title, timestamp)  
    texFiles = getTexFiles()
    filter = TextFilter(level=3)
    if len(texFiles) == 1:
        text = read_file_without_comments(texFiles[0])
        method = extract_method(text)
        if not method:
            section_list = extract_all_section(text)
            method_name = get_section_name(section_list)['Method']
            method = extract_method(text, method_name)
        filter_method = filter(method)
        # method = delete_others(method)
    else:
        merge_tex_files(texFiles) # 将所有tex文件暂时合并为一个，并保存到tmp.tex中
        text = read_file_without_comments(config['utils']['tmp_file'])
        method = extract_method(text)
        if not method: # 如果method命名不是Method，需要重新获取
            section_list = extract_all_section(text) 
            method_name = get_section_name(section_list)['Method']
            method = extract_method(text, method_name)
        if len(method) < 100: # method内容可能是链接到其他文件的链接代码
            # 匹配单独method文件
            match_file = [file for file in texFiles if any(file.split('/')[-1].startswith(prefix) for prefix in ['met', 'Met', 'the', 'The', 'THE'])]
            method = read_file_without_comments(match_file[0])
            filter_method = filter(method)
        else:
            filter_method = filter(method)
    return filter_method
def get_conclusion(title: str, timestamp):
    getTexFiles = TexFiles(title, timestamp)  
    texFiles = getTexFiles()
    if len(texFiles) == 1:
        text = read_file_without_comments(texFiles[0])
        text = delete_image_table(text)  
        conclusion = extract_conclusion(text)
        conclusion = delete_others(conclusion)
    else:
        match_file = [file for file in texFiles if any(file.split('/')[-1].startswith(prefix) for prefix in ['conclusion', 'Conclusion', 'con'])]
        conclusion = read_file_without_comments(match_file[0])
        conclusion = delete_image_table(conclusion)
        conclusion = delete_others(conclusion)
    return conclusion
def analysis_introdution(llm, introduction: str) -> str:
    '''
    从论文的Introduction部分分析出以下三部分内容：
    1. 论文的研究背景
    2. 所要解决的问题
    3. 解决该问题的过去方案
    4. 论文的动机或创新点
    '''
    text_spliter = CharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=20,
            length_function=len,
        )
    CHAIN_TYPE = 'map_reduce'
    summary_chain = load_summarize_chain(llm, chain_type=CHAIN_TYPE)
    summarize_document_chain = AnalyzeDocumentChain(   
        combine_docs_chain=summary_chain,
        text_spliter=text_spliter,
    )
    start_time = time.time()
    logger.info(f"Start summarizing the document")
    res = summarize_document_chain.run(introduction)
    end_time = time.time()
    logger.info(f"Summarizing the document takes {end_time - start_time} seconds.")
    # 将模型输出的introduction保存到本地
    output_information(res, os.path.join(os.path.dirname(os.path.abspath(__file__)), 'introduction'), 'test')
    return res


def analysis_method(llm, method: str) -> str:
    '''
    Describe in detail the methodological idea of this article:
    - (1):...
    - (2):...
    - (3):...
    - .......
    '''
    text_spliter = CharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=20,
            length_function=len,
        )
    CHAIN_TYPE = 'map_reduce'
    summary_chain = load_summarize_chain(llm, chain_type=CHAIN_TYPE)
    summarize_document_chain = AnalyzeDocumentChain(   
        combine_docs_chain=summary_chain,
        text_spliter=text_spliter,
    )
    start_time = time.time()
    logger.info(f"Start summarizing the document")
    res = summarize_document_chain.run(method)
    end_time = time.time()
    logger.info(f"Summarizing the document takes {end_time - start_time} seconds.")
    # 将模型输出的introduction保存到本地
    output_information(res, os.path.join(os.path.dirname(os.path.abspath(__file__)), 'introduction'), 'test')
    return res

def analysis_conclusion(llm, conclusion: str) -> str:
    '''
    Describe the main conclusions of this article:
    - (1):...
    - (2):...
    - (3):...
    - .......
    '''
    text_spliter = CharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=20,
            length_function=len,
        )
    CHAIN_TYPE = 'map_reduce'
    summary_chain = load_summarize_chain(llm, chain_type=CHAIN_TYPE)
    summarize_document_chain = AnalyzeDocumentChain(   
        combine_docs_chain=summary_chain,
        text_spliter=text_spliter,
    )
    start_time = time.time()
    logger.info(f"Start summarizing the document")
    res = summarize_document_chain.run(conclusion)
    end_time = time.time()
    logger.info(f"Summarizing the document takes {end_time - start_time} seconds.")
    # 将模型输出的introduction保存到本地
    output_information(res, os.path.join(os.path.dirname(os.path.abspath(__file__)), 'introduction'), 'test')
    return res

def get_one_paper_introduction_abstract(title: str, timestamp: str):
    '''
    提取给定title的论文的introduction经过模型处理后的内容（从本地txt中提取），需要将所有论文的abstract部分一起输入LLM得到最终的summary
    '''

    # title结尾可能有一个'，需要特别注意
    abstract_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'introduction', title+'.txt')
    try:
        with open(abstract_path, 'r') as f:
            this_abstract = f.read()
    except FileNotFoundError:
        logger.error(f"File '{abstract_path}' not found.")
    return this_abstract

def get_all_paper_summary(llm, abstracts: List[str]):
    '''
    将所有论文的abstract输入LLM得到最终的summary
    abstracts: List[str]，每个元素是一篇论文的abstract的本地绝对地址
    '''
    summary_input = construct_prompt(abstracts, 2000)
    text_spliter = CharacterTextSplitter(
        chunk_size=1500,
        chunk_overlap=20,
        length_function=len,
    )
    pure_prompt = "{text}"
    PROMPT = PromptTemplate(
        template = pure_prompt,
        input_variable = ["text"],
    )
    CHAIN_TYPE = 'map_reduce'
    summary_chain = load_summarize_chain(llm, chain_type=CHAIN_TYPE)
    summarize_document_chain = AnalyzeDocumentChain(
        combine_docs_chain=summary_chain,
        text_spliter=text_spliter,
        prompt=PROMPT,
    )
    start_time = time.time()
    logger.info(f"Start Generate Summary for All Papers")
    res = summarize_document_chain.run(summary_input)
    end_time = time.time()
    logger.info(f"Generate Summary for All Papers takes {end_time - start_time} seconds.")
    return res

def summary_with_llm(llm, input, summary_type: str):
    text_spliter = CharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=20,
            length_function=len,
        )
    CHAIN_TYPE = 'map_reduce'
    summary_chain = load_summarize_chain(llm, chain_type=CHAIN_TYPE)
    summarize_document_chain = AnalyzeDocumentChain(   
        combine_docs_chain=summary_chain,
        text_spliter=text_spliter,
    )
    start_time = time.time()
    logger.info(f"Start summarizing the document")
    # 标识符
    res = summarize_document_chain.run(summary_type + input)
    end_time = time.time()
    logger.info(f"Summarizing the document takes {end_time - start_time} seconds.")
    # 将模型输出保存到本地
    # output_information(res, os.path.join(os.path.dirname(os.path.abspath(__file__)), 'review' + timestamp), 'review.txt')
    return res
def summary_abstract(paperInfo_list: list) -> str:
    '''
    从所有论文的摘要部分总结出综述的摘要
    '''
    abstracts = [paper.original_abstract for paper in paperInfo_list]
    input_abstract = ''
    # 将所有论文的abstract拼接起来
    # 1. XXX
    # 2. XXX
    for i, abstract in enumerate(abstracts):
        input_abstract += str(i+1) + ': ' + abstract.strip() + '\n'
    summary_abstract = summary_with_llm(llm, input_abstract, timestamp, '4')
    return summary_abstract

def summary_introduction(paperInfo_list: list) -> str:
    '''
    概括
    A: Definition of the XX
    B: Importance of the XX
    C: purpose of this Survey
    D: Contribution of this Survey
    E: Overview of the survey structure
    '''
    introductions = [paper.introduction_summary for paper in paperInfo_list]
    input_introduction = ''
    # 将所有论文的introduction拼接起来
    for i, intro in enumerate(introductions):
        input_introduction += str(i+1) + ': ' + intro.strip() + '\n'
    summary_introduction = summary_with_llm(llm, input_introduction, timestamp, '5')
    return summary_introduction

def summary_method(paperInfo_list: list) -> str:
    '''
    将每篇论文的method总结成一段话，然后堆叠即可
    形式为\textbf{XXX} XXXXX
    '''
    for paper in paperInfo_list:
        # 输入论文标题、作者和method，总结成一段话
        paper_method = summary_with_llm(llm, paper.title + '|' + paper.autho + '|' + paper.method, '6')


@retry(stop=stop_after_attempt(3), wait=wait_fixed(3))
def get_review_abstract(key_words_information: str) -> str:
    # 输入为配置文件中的关键词信息
    query = construct_query_for_generate_abstract(key_words_information)
    try:
        abstract = chat_with_llm_api(query)
        return abstract
    except Exception as e:
        raise e

@retry(stop=stop_after_attempt(3), wait=wait_fixed(3))
def get_introduction_part1(keywords_information:str) -> str:
    # 输入为配置文件中的关键词信息
    query = construct_query_for_generate_introduction_part1(keywords_information)
    try:
        part_1 = chat_with_llm_api(query)
        return part_1
    except Exception as e:
        raise e
@retry(stop=stop_after_attempt(3), wait=wait_fixed(3))
def get_introduction_part2(keywords_information:str) -> str:
    # 输入为配置文件中的关键词信息
    query = construct_query_for_generate_introduction_part2(keywords_information)
    try:
        part_2 = chat_with_llm_api(query)
        return part_2
    except Exception as e:
        raise e
@retry(stop=stop_after_attempt(3), wait=wait_fixed(3))
def review_table_from_abstract(abstract: str) -> str:
    '''
    从论文的摘要中生成综述的表格内容
    '''
    parser = OutputParser("get_review_table_from_abstract")
    query = review_from_abstract(abstract)
    try:
        information = chat_with_llm_api(query)
        return parser(information) # 返回一个字典
    except Exception as e:
        raise e
@retry(stop=stop_after_attempt(3), wait=wait_fixed(3))
def review_table_from_method(method: str) -> str:
    '''
    从论文的方法中生成综述的表格内容
    '''
    parser = OutputParser("get_review_table_from_method")
    query = review_from_method(method)
    try:
        information = chat_with_llm_api(query)
        return parser(information) # 返回一个字典
    except Exception as e:
        raise e
@retry(stop=stop_after_attempt(3), wait=wait_fixed(3))
def review_table_from_conclusion(conclusion: str) -> str:
    '''
    从论文的结论中生成综述的表格内容
    '''
    parser = OutputParser("get_review_table_from_conclusion")
    query = review_from_conclusion(conclusion)
    try:
        information = chat_with_llm_api(query)
        return parser(information) # 返回一个字典
    except Exception as e:
        raise e
def insert_table_to_word(paperInfo_list: list, doc_path: str) -> None:
    headers = ['Name', 'Author', 'Year', 'Method', 'Main Result', 'Advantages and Disadvantages', 'Innovation']
    doc = Document(doc_path)
    table = doc.add_table(rows=len(paperInfo_list)+1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.columns[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    for i, row in enumerate(paperInfo_list):
        data = [row.column_1, row.author, row.published, row.review_table_information_from_method, row.review_table_information_from_conclusion, row.column_2, row.column_3]
        for j, cell_value in enumerate(data):
            table.cell(i+1, j).text = str(cell_value)
    doc.add_paragraph("Table" + "caption").bold = True
    doc.save(doc_path)

if __name__ == '__main__':
    config = load_config()
    
    # Todo: 
    llm = ChatVicuna()
    timestamp = get_timestamp()

    # 下载论文
    # Todo: Parameterize
    arxivtool = ArxivTools(config['ARXIV']['search_keywords'], int(config['ARXIV']['search_nums']), arxiv.SortCriterion.SubmittedDate)
    paper_results = arxivtool() # arxiv列表
    print('Check the proxy if meet the error: AttributeError: object has no attribute \'status\'')
    downloadtools = DownloadTools(paper_results, timestamp)
    downloadtools.download_latex_file(url=list(map(lambda x: x.entry_id, paper_results)), save_path=list(map(lambda x: x.title, paper_results)))
    title_list = list(map(lambda x: x.title[0], paper_results))
    paperInfoList = []
    # 解析每篇论文
    with tqdm(total=int(config['ARXIV']['search_nums']), desc="Analysis Each Paper", unit="item") as pbar:
        for arxiv_paper in paper_results:
            title = arxiv_paper.title[0]
            authors = [author.__str__() for author in arxiv_paper.author[0]]
            url = arxiv_paper.entry_id
            published = arxiv_paper.published
            original_abstract = arxiv_paper.summary
            summary_sentence = summary_with_llm(llm, original_abstract, '7')
            try:
                review_table_information_from_abstract = review_table_from_abstract(config['ARXIV']['keyword_information'], original_abstract)
                column_1 = review_table_information_from_abstract['1. Method Name'] 
                column_2 = review_table_information_from_abstract['2. Strengths and weaknesses']
                column_3 = review_table_information_from_abstract['3. Innovation point']

                introduction = get_introduction(title, timestamp)
                method = get_method(title, timestamp)
                conclusion = get_conclusion(title, timestamp)
                introduction = '0' + introduction
                introduction_summary = analysis_introdution(llm, introduction)

                input_method = "1" + "<summary>" + introduction_summary + "\n\n<Methods>:\n\n" + method
                method_summary = analysis_method(llm, input_method)
                review_table_information_from_method = review_table_from_method(method_summary)

                input_conclusion = "2" + "<summary>" + introduction_summary + "\n <Method summary>:\n" + method_summary + "\n\n<Conclusion>:\n\n" + conclusion
                conclusion_summary = analysis_conclusion(llm, input_conclusion)
                # review_table_information_from_conclusion = review_table_from_conclusion(conclusion_summary)
                review_table_information_from_conclusion = 'conclusion'
                paper = paperInfo(title, authors, url, original_abstract, introduction_summary, method_summary, conclusion_summary, column_1, column_2, column_3, published, review_table_information_from_method, review_table_information_from_conclusion)
                paperInfoList.append(paper) 
                # introduciton, extracted_abstract = llm_chat_with_one_paper(llm, title, timestamp)
                #paper_info = paperInfo(title, authors, url, original_abstract, introduction_summary)
                pbar.update(1)
            except Exception as e:
                logger.error(f"Error: {e}")
    # 保存论文初步提取的信息
    with open('paper_information.pickle', 'wb') as f:
        pickle.dump(paperInfoList, f)
    logger.info(f"All Paper Analysis Finish!")
    total_summary = ''
    abstract_summary = summary_abstract(paperInfoList)
    total_summary += abstract_summary + '\n\n'

    introduction_summary = summary_introduction(paperInfoList)
    # 综述 1. abstract
    review_abstract = get_review_abstract(config['ARXIV']['keyword_information'])
    review_keywords = 'Keywords: ' + config['ARXIV']['search_keywords'] 
    # 综述 2. introduction
    review_introduction_part_1 = get_introduction_part1(config['ARXIV']['keyword_information'])
    # 2.1 Definition of the Voice Conversation
    review_introduction_part_2 = get_introduction_part2(config['ARXIV']['keyword_information'])

    # 2.2 Importance of the Voice Conversation
    







    # ToDo: 生成所有论文的summary
    # logger.info(f"Start to generate the summary of all papers")
    # 读取所有论文的introduction
    # all_abstract = []
    # for paper in title_list:
    #     try:
    #         all_abstract.append(get_one_paper_introduction_abstract(paper, timestamp))
    #     except FileNotFoundError:
    #         logger.error(f"File '{paper}' not found.")
    # # 生成所有论文的summary
    # all_summary = get_all_paper_summary(llm, all_abstract)
    # make_dir_if_not_exist(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'summary', timestamp))



    

    
