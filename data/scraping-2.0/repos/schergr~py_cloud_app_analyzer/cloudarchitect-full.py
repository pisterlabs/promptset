import logging
import config
import os
import re
import sys
import glob
import time
import json
import argparse
import requests
import warnings
import pandas as pd
import openpyxl
import openai
from transformers import GPT2Tokenizer
from docx import Document
from pptx import Presentation
from pptx.util import Inches, Pt
from pptx.enum.text import PP_ALIGN
from markdownify import markdownify as md

spreadsheet_name = "BCICAppListandQuestionnaire.xlsx"
architecture_prompt_file = "architecture_prompt.txt"
client_principles_prompt_file = "client_considerations.txt"
migration_prompt_file = "migration_principles.txt"
csp_selection_prompt_file = "csp_principles.txt"
output_sheet_name = "processed_data.xlsx"

discovery_process_status_column = 'Discovery Status'  # Adjust the column name as per your spreadsheet
app_name_column = 'Application'  # Adjust the column name as per your spreadsheet

root_path = "G:\\My Drive\\MyCode\\projects\\py_cloud_app_analyzer\\"
root_data_path = root_path + "localdata\\"
training_data_path = root_path + "training\\"
root_docs_path = root_path + "output\\"
root_logs_path = root_path + "logs\\"
full_path = root_data_path + spreadsheet_name
application_summary_path = root_docs_path
output_file_path = root_docs_path + output_sheet_name

log_filename = root_logs_path + "log.txt"

font_name = "Arial"
font_size = Pt(9)

### prompts
CISO_prompt = "You are the CISO of this company. Your job is to ensure that your team strikes the best balance between security and business agility. Review the analaysis provided, and provide expert feedback on ways to improve the strategy from a CISO perspective."
COO_prompt = "You are the COO of this company. Your job is to ensure that your team strikes the best balance between business agility, and financial responsibility. Review the analaysis provided, and provide expert feedback on ways to improve the strategy from an COO perspective."
final_analysis_prompt = "Thats an excellent analysis. You did a great job comparing those two documents. Now, providing as many facts and details as possible from the original source info that I gave you,  I'd like to take the best of both documents, and combine them into a final comprehensive analysis."
summary_prompt_part1 = "Thats an excellent analysis. You did a great job. Now, providing as many facts and details as possible from the original data I gave you. Here it is again, remember?"
summary_prompt_part2 = " I'd like you to create a 3 paragraph executive overview of our strategy, approach, and sequencing. Please watch your grammar."

perplexityai_url = "https://api.perplexity.ai/chat/completions"
perplexity_headers = {
        "accept": "application/json",
        "content-type": "application/json",
        "authorization": config.perplexity_api_key
    }

def remove_openai_from_docx(directory):
    for filename in os.listdir(directory):
        if filename.endswith('-openai Analysis.docx'):
            # Process the file
            file_path = os.path.join(directory, filename)
            doc = Document(file_path)

            # Edit the body of the document
            for section in doc.sections:
                for header in section.header.paragraphs:
                    header.text = header.text.replace('-openai', '')

            # Save the edited document
            new_file_path = os.path.join(directory, filename.replace('-openai', ''))
            doc.save(new_file_path)

            # Optionally, remove the original file
            os.remove(file_path)
            logging.info(f"Processed and renamed file: {new_file_path}")
        elif filename.endswith('.docx'):
            file_path = os.path.join(directory, filename)
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                 if "openai" in paragraph.text:
                    # Replace the text
                    logging.info(f"Current header text {paragraph.text}")
                    paragraph.text = paragraph.text.replace("openai", '')
                    logging.info(f"New Header Text {paragraph.text}")
            # Save the edited document
            doc.save(file_path)
            logging.info(f"Processed file: {file_path}")
        else:
            logging.info(f"No Files matching criteria.")

def create_conversation_data(file_path):
    # Read the spreadsheet into a DataFrame
    df = pd.read_excel(file_path)

    # List to hold training data
    messages = [{"role": "system", "content": "You are a cloud architect, and these are facts to help you."}]

    # Iterate through each row and column
    for index, row in df.iterrows():
        application_value = row['Application'] if 'Application' in df.columns else "Unknown Application"
        for col in df.columns:
            question = f"{application_value} {col}"  # Append application value to question
            answer = row[col]  # Cell content as answer
            messages.append({"role": "user", "content": question})
            messages.append({"role": "assistant", "content": answer})
    # Wrap the messages in a dictionary
    conversation_data = {"messages": messages}
    return conversation_data

def docx_to_text(file_path):
    """Converts a docx file to plain text"""
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def docx_to_markdown(file_path):
    """Converts a docx file to markdown format"""
    text = docx_to_text(file_path)
    return md(text)

def convert_files_in_directory(directory_path, convert_to='text'):
    """Converts all docx files in the given directory to text or markdown"""
    results = []
    for filename in os.listdir(directory_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(directory_path, filename)
            if convert_to == 'postprocessing':
                # Read .docx file
                doc = Document(os.path.join(directory_path, filename))
                # Extract text
                text = ' '.join([p.text for p in doc.paragraphs])
                # Process text and get results
                #summary = process_text(text)  # Replace with your function
                #migration_strategy = get_migration_strategy(text)  # Replace with your function
                #cloud_recommendation = get_cloud_recommendation(text)  # Replace with your function
                # Store results
                #results.append([file, summary, migration_strategy, cloud_recommendation])
                #df = pd.DataFrame(results, columns=['File', 'Summary', 'Migration Strategy', 'Cloud Recommendation'])
                # Write DataFrame to Excel file
                #df.to_excel('results.xlsx', index=False)
                logging.info(f"Postprocessing of {filename} complete.")
            elif convert_to == 'text':
                converted_text = docx_to_text(file_path)
                output_filename = filename.replace('.docx', '.txt' if convert_to == 'text' else '.md')
                output_path = os.path.join(directory_path, output_filename)
                with open(output_path, 'w') as file:
                    file.write(converted_text)
                logging.info(f"Converted {filename} to {output_filename}")
            elif convert_to == 'markdown':
                converted_text = docx_to_markdown(file_path)
                output_filename = filename.replace('.docx', '.txt' if convert_to == 'text' else '.md')
                output_path = os.path.join(directory_path, output_filename)
                with open(output_path, 'w') as file:
                    file.write(converted_text)
                logging.info(f"Converted {filename} to {output_filename}")
            else:
                raise ValueError("Invalid conversion format specified")

def add_file_handler(logger, log_file):
    file_handler = logging.FileHandler(log_file, 'w', 'utf-8')
    file_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
    logger.addHandler(file_handler)

def read_md_files(base_data_path):
    all_md_content = ""
    file_count=0
    logging.info (f"Loading Training Data...")
    for root, dirs, files in os.walk(base_data_path):
        for file in files:
            if file.endswith(".md"):
                file_count=file_count+1
                logging.debug (f"Preparing to read file {file_count}: {file}")
                file_path = os.path.join(root, file)
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        all_md_content += file.read() + "\n\n"
                        logging.debug (f"Read File: {file}")
                except IOError as e:
                    logging.error(f"Error reading file {file_path}: {e}")
                except Exception as e:
                    logging.error(f"An unexpected error occurred while processing {file_path}: {e}")
    logging.info(f"Accurate token count: {count_tokens(all_md_content)}")
    logging.info(f"Training on {file_count} files.")
    input("Press Enter to continue...")
    chunks = split_into_chunks(all_md_content)
    return chunks

def count_tokens(text):
    tokenizer = GPT2Tokenizer.from_pretrained("gpt2")
    tokens = tokenizer.encode(text)
    return len(tokens)

def split_into_chunks(text, max_token_size=9999):
    try:
        tokenizer = GPT2Tokenizer.from_pretrained("gpt2")
    except Exception as e:
        logging.error(f"Error loading tokenizer: {e}")
        return []
    try:
        tokenized_text = tokenizer.encode(text)
    except MemoryError:
        logging.error("Memory Error: The input text is too large.")
        return []
    except Exception as e:
        logging.error(f"Error during tokenization: {e}")
        return []

    chunks = []
    current_chunk = []
    current_length = 0

    for token in tokenized_text:
        current_chunk.append(token)
        current_length += 1
        if current_length == max_token_size:
            try:
                chunks.append(tokenizer.decode(current_chunk))
            except Exception as e:
                logging.error(f"Error decoding tokens: {e}")
                # Optionally, continue to the next chunk instead of returning
                return []
            current_chunk = []
            current_length = 0

    if current_chunk:
        try:
            chunks.append(tokenizer.decode(current_chunk))
        except Exception as e:
            logging.error(f"Error decoding tokens: {e}")
            return []

    return chunks

def create_word_document_with_completion(summary_text, application_name2, file_path_name):
    """
    Creates a Word document with completion.

    Args:
        summary_text (str): The summary text for the document.
        raw_data (str): The raw data to be added to the document.
        source (str): The source of the analysis.
        application_name2 (str): The name of the application.
        file_path_name (str): The file path and name to save the document.

    Returns:
        None
    """
    logging.info(f"Creating document for {application_name2}")
    doc = Document()
    logging.info(f"Initialized document {doc}")
    doc.add_heading(application_name2 + ' Analysis', 0)
    logging.debug(f"Added Document Heading for {application_name2}")
    doc.add_paragraph(summary_text)
    logging.debug(f"Added paragraph.")
    logging.debug(f"Saving {file_path_name} {application_name2} Analysis.docx")
    doc.save(file_path_name + application_name2 + " Analysis.docx")
    logging.info(f"Document {file_path_name} {application_name2} Analysis.docx saved")

def format_textbox(textbox, text):
    textbox.text = text
    for paragraph in textbox.text_frame.paragraphs:
        for run in paragraph.runs:
            run.font.size = font_size
            run.font.name = font_name
            paragraph.alignment = PP_ALIGN.LEFT

def list_applications(full_path):
    """
	Lists the applications from a given Excel file.

	Parameters:
	- full_path (str): The full path of the Excel file.

	Returns:
	None

	Raises:
	None
	"""
    warnings.filterwarnings("ignore", message="Workbook contains no default style, apply openpyxl's default")
    workbook = openpyxl.load_workbook(full_path)
    sheet = workbook.active

    for row in range(1, sheet.max_row + 1):
        warnings.filterwarnings("ignore", message="Workbook contains no default style, apply openpyxl's default")    
        application_name = sheet.cell(row=row+1, column=2).value
        application_lifecycle_status = sheet.cell(row=row+1, column=5).value
        application_discovery_status = sheet.cell(row=row+1, column=4).value

        logging.info(f"{row}: {application_name} ({application_lifecycle_status}) ({application_discovery_status})")

def get_response_from_model(persona,system_prompt,user_prompt,ai_model):
    perplexityai_url = "https://api.perplexity.ai/chat/completions"
    perplexity_headers = {
        "accept": "application/json",
        "content-type": "application/json",
        "authorization": config.perplexity_api_key
    }
    if ai_model=="openai":
        """
        Generates a response using OpenAI's chat completion API.

        :param persona: The persona used for the analysis.
        :type persona: str
        :param system_prompt: The system prompt for the chat completion.
        :type system_prompt: str
        :param user_prompt: The user prompt for the chat completion.
        :type user_prompt: str
        :return: The generated response.
        :rtype: str
        """
        logging.info(f"Processing {persona} Analysis with OpenAI")
        payload = openai.chat.completions.create(
            model="gpt-4",
            messages=[
                {
                    "role": "system",
                    "content": system_prompt,
                },
                {
                    "role": "user",
                    "content": user_prompt,
                },

            ],
        )
    if ai_model=="perplexity":
        """
        Sends a POST request to the Perplexity AI API to get the response perplexity.
        
        Args:
            persona (str): The persona for the conversation.
            system_prompt (str): The system prompt for the conversation.
            user_prompt (str): The user prompt for the conversation.
            perplexityai_url (str): The URL of the Perplexity AI API.
            perplexity_headers (dict): The headers for the Perplexity AI API request.
            
        Returns:
            str: The content of the response message from the Perplexity AI API.
            None: If there was an error in the request (HTTP error, connection error, timeout, or general request exception).
        """
        logging.info(f"Processing {persona} Analysis with Perplexity.AI")
        payload = {
            "model": "mistral-7b-instruct",
            "messages": [
                {
                    "role": "system",
                    "content": system_prompt,
                },
                {
                    "role": "user",
                    "content": user_prompt,
                }
            ]
        }
        try:
            response = requests.post(perplexityai_url, json=payload, headers=perplexity_headers)
            response.raise_for_status()  # Check if the request was successful
        except requests.exceptions.HTTPError as errh:
            logging.error (f"HTTP Error: {errh}")
            return None
        except requests.exceptions.ConnectionError as errc:
            logging.error (f"Error Connecting: {errc}")
            return None
        except requests.exceptions.Timeout as errt:
            logging.error (f"Timeout Error: {errt}")
            return None
        except requests.exceptions.RequestException as err:
            logging.error (f"Something went wrong: {err}")
            return None
    try:
        response_content = response.json()
        return response_content["choices"][0]["message"]["content"]
    except KeyError as e:
        logging.error(f"Key error in parsing response: {e}")
        return None

class PromptReader:
    def __init__(self, root_data_path, training_data_path,architecture_prompt_file,client_principles_prompt_file,migration_prompt_file,csp_selection_prompt_file):
        self.arch_prompt_path = os.path.join(root_data_path, architecture_prompt_file)
        self.clientprinciples_prompt_path = os.path.join(training_data_path, client_principles_prompt_file)
        self.msazureprinciples_prompt_path = os.path.join(training_data_path, migration_prompt_file)
        self.cspprinciples_prompt_path = os.path.join(training_data_path, csp_selection_prompt_file)

    def read_file(self, file_path):
        try:
            with open(file_path, 'r') as file:
                return file.read()
        except IOError as e:
            logging.error(f"Error reading file {file_path}: {e}")
            return None

    def get_prompts(self):
        architect_prompt = self.read_file(self.arch_prompt_path)
        client_principles_prompt = self.read_file(self.clientprinciples_prompt_path)
        MSAzure_principles_prompt = self.read_file(self.msazureprinciples_prompt_path)
        CSP_principles_prompt = self.read_file(self.cspprinciples_prompt_path)
        return architect_prompt, client_principles_prompt, MSAzure_principles_prompt, CSP_principles_prompt

class OpenAIChatSession:
    def __init__(self, model="gpt-4", system_prompt=""):
        self.model = model
        self.messages = [{"role": "system", "content": system_prompt}]

    def add_message(self, role, content):
        self.messages.append({"role": role, "content": content})

    def get_response(self, user_prompt):
        self.add_message("user", user_prompt)
        try:
            payload = openai.chat.completions.create(
                model=self.model,
                messages=self.messages,
                temperature=.5,
                user ="Greg"
            )
            logging.debug(f"Payload: {payload}\n")
            response = payload.choices[0].message.content
            payload_num_tokens=count_tokens(response)
            logging.debug(f"This payload has {payload_num_tokens} tokens.")
            self.add_message("assistant", response)
            logging.debug(f"Payload: {response}\n")
            return response
        except openai.error.RateLimitError as e:
            logging.error(f"Rate Limit Exceeded: {e}")
            # Extract wait time from the error message
            match = re.search(r'Please try again in (\d+)m(\d+\.\d+)s', str(e))
            if match:
                #wait_time = match.group(1)
                minutes = int(match.group(1))
                seconds = float(match.group(2))
                total_seconds = (minutes * 60) + seconds
                logging.warning (f"Rate limit exceeded. Trying again in {minutes}m{seconds}s.")
                time.sleep(total_seconds)
            else:
                logging.error(f"Rate limit exceeded. Please try again later.")
                return "Rate limit exceeded. Please try again later."
        except openai.OpenAIError as e:
            logging.error(f"OpenAI API Error: {e}")
            return None
        except Exception as e:
            logging.error(f"Error in OpenAI API call: {e}")
            return None
    
    def get_history(self):
        history = ""
        for message in self.messages:
            history += (f"{message['role'].title()}: {message['content']}\n")
        return history

def get_response_openai_chat(persona,system_prompt,user_prompt,client_principles_prompt,MSAzure_principles_prompt,CSP_principles_prompt,training_data):

    use_conversation = 1
    logging.info(f"Processing {persona} Analysis with OpenAI")
    if use_conversation == 1:
        session = OpenAIChatSession(
            model="gpt-4",
            system_prompt=system_prompt
        )
        # Add context or principles as needed
        session.add_message("assistant", client_principles_prompt)
        session.add_message("assistant", MSAzure_principles_prompt)
        session.add_message("assistant", CSP_principles_prompt)
        #session.add_message("assistant", training_data)
        payload = session.get_response(user_prompt)
        conversation_history = session.get_history()
        conv_num_tokens=count_tokens(conversation_history)
        
        logging.debug(f"This conversation has {conv_num_tokens} tokens.")
        logging.debug(f"Conversation History: {conversation_history}")
    else:
        payload = openai.chat.completions.create(
            model="gpt-4-1106-preview",
            messages=[
                {
                    "role": "system",
                    "content": system_prompt,
                },
                {
                    "role": "user",
                    "content": user_prompt,
                },

            ],
        )
    return payload

def get_response_openai(persona,system_prompt,user_prompt):
    """
	Generates a response using OpenAI's chat completion API.

	:param persona: The persona used for the analysis.
	:type persona: str
	:param system_prompt: The system prompt for the chat completion.
	:type system_prompt: str
	:param user_prompt: The user prompt for the chat completion.
	:type user_prompt: str
	:return: The generated response.
	:rtype: str
	"""
    logging.info(f"Processing {persona} Analysis with OpenAI")
    payload = openai.chat.completions.create(
        model="gpt-4",
        messages=[
            {
                "role": "system",
                "content": system_prompt,
            },
            {
                "role": "user",
                "content": user_prompt,
            },

        ],
    )
    return payload.choices[0].message.content

def get_response_perplexity_ai (persona, system_prompt, user_prompt, perplexityai_url, perplexity_headers):
    """
    Sends a POST request to the Perplexity AI API to get the response perplexity.
    
    Args:
        persona (str): The persona for the conversation.
        system_prompt (str): The system prompt for the conversation.
        user_prompt (str): The user prompt for the conversation.
        perplexityai_url (str): The URL of the Perplexity AI API.
        perplexity_headers (dict): The headers for the Perplexity AI API request.
        
    Returns:
        str: The content of the response message from the Perplexity AI API.
        None: If there was an error in the request (HTTP error, connection error, timeout, or general request exception).
    """
    logging.info(f"Processing {persona} Analysis with Perplexity.AI")
    payload = {
        "model": "mistral-7b-instruct",
        "messages": [
            {
                "role": "system",
                "content": system_prompt,
            },
            {
                "role": "user",
                "content": user_prompt,
            }
        ]
    }
    try:
        response = requests.post(perplexityai_url, json=payload, headers=perplexity_headers)
        response.raise_for_status()  # Check if the request was successful
    except requests.exceptions.HTTPError as errh:
        logging.error (f"HTTP Error: {errh}")
        return None
    except requests.exceptions.ConnectionError as errc:
        logging.error (f"Error Connecting: {errc}")
        return None
    except requests.exceptions.Timeout as errt:
        logging.error (f"Timeout Error: {errt}")
        return None
    except requests.exceptions.RequestException as err:
        logging.error (f"Something went wrong: {err}")
        return None
    try:
        response_content = response.json()
        return response_content["choices"][0]["message"]["content"]
    except KeyError as e:
        logging.error(f"Key error in parsing response: {e}")
        return None

def process_row_with_perplexityai(allorone,data_row,headers):
    """
    Processes a row using PerplexityAI.

    Parameters:
        None

    Returns:
        str: The summary content generated by PerplexityAI.
    """
    arch_prompt_path = os.path.join(root_data_path, "architecture_prompt.txt")
    gaps_prompt_path = os.path.join(root_data_path, "gap_analysis_prompt.txt")

    with open(arch_prompt_path, 'r') as file:
        architect_prompt = file.read()
    
    with open(gaps_prompt_path, 'r') as file:
        gaps_prompt = file.read()

    if allorone == "--single":
        app_spreadsheet_full_path = os.path.join(root_data_path, spreadsheet_name)

        try:
            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                df = pd.read_excel(app_spreadsheet_full_path)
        except Exception as e:
            logging.error(f"Error reading the spreadsheet: {e}")
            return

        list_applications(app_spreadsheet_full_path)

        # Prompt user for row number
        row_number = int(input("Please enter the row number of the spreadsheet to process: "))
        if row_number < 0 or row_number >= len(df):
            logging.warning("Invalid row number.")
            return

        # Retrieve column headers and data from the specified row
        headers = df.columns.tolist()
        data_row = df.iloc[row_number-1].to_dict()
        user_prompt = "\n".join([f"{header}: {data_row[header]}" for header in headers])
    else:
        user_prompt = "\n".join([f"{header}: {data_row[header]}" for header in headers])
    
    logging.info (user_prompt)
    BA_response = get_response_perplexity_ai("BA", gaps_prompt, user_prompt, perplexityai_url, perplexity_headers)
    CloudArchitect_response = get_response_perplexity_ai("Cloud Architect", architect_prompt, user_prompt, perplexityai_url, perplexity_headers)
    CISO_response = get_response_perplexity_ai("CISO", CISO_prompt, user_prompt + "\n" + CloudArchitect_response, perplexityai_url, perplexity_headers)
    COO_response = get_response_perplexity_ai("COO", COO_prompt, user_prompt + "\n" + CloudArchitect_response, perplexityai_url, perplexity_headers)
    summary_content = CloudArchitect_response + "\n" + CISO_response + "\n" + COO_response + "\n" + BA_response
    
    return summary_content, data_row['Application']

def process_row_with_openai(allorone,data_row,headers,training_data):

    prompt_reader = PromptReader(root_data_path, training_data_path,architecture_prompt_file,client_principles_prompt_file,migration_prompt_file,csp_selection_prompt_file)
    architect_prompt, client_principles_prompt, MSAzure_principles_prompt,csp_selection_prompt = prompt_reader.get_prompts()
    
    if allorone == "--single":
        full_path = os.path.join(root_data_path, spreadsheet_name)
        try:
            warnings.filterwarnings("ignore", message="Workbook contains no default style, apply openpyxl's default")
            df = pd.read_excel(full_path)
        except Exception as e:
            logging.error(f"Error reading the spreadsheet: {e}")
            return

        list_applications(full_path)

        # Prompt user for row number
        row_number = int(input("Please enter the row number of the spreadsheet to process: "))
        if row_number < 0 or row_number >= len(df):
            logging.warning("Invalid row number.")
            return

        # Retrieve column headers and data from the specified row
        headers = df.columns.tolist()
        data_row = df.iloc[row_number - 1].to_dict()
        # Combine headers and row data into a single string
        user_prompt = "\n".join([f"{header}: {data_row[header]}" for header in headers])
    else:
        user_prompt = "\n".join([f"{header}: {data_row[header]}" for header in headers])            

    CloudArchitect_response = get_response_openai_chat("Cloud Architect", architect_prompt, user_prompt, client_principles_prompt, MSAzure_principles_prompt,csp_selection_prompt,"")
    logging.info(f"Cloud Architect Response:\n\n{CloudArchitect_response}\n")
    #CISO_response = get_response_openai_chat("CISO", CISO_prompt, user_prompt + "\n" + CloudArchitect_response,client_principles_prompt, MSAzure_principles_prompt,"","")
    #logging.debug(f"CISO Response:\n {CISO_response}\n")
    #COO_response = get_response_openai_chat("COO", COO_prompt, user_prompt + "\n" + CloudArchitect_response,client_principles_prompt, MSAzure_principles_prompt,"","")
    #logging.debug(f"COO Response:\n {COO_response}\n")
    #summary_content = CloudArchitect_response + "\n" + CISO_response + "\n" + COO_response + "\n"
    summary_content = CloudArchitect_response
    #summary_response = get_response_openai_chat("Executive Overview Creator", summary_prompt_part1 + "\n" + user_prompt  + "\n" + summary_prompt_part2, summary_content, client_principles_prompt, MSAzure_principles_prompt,"","")
    #logging.info(f"Executive Summary:\n\n{summary_response}\n")

    return summary_content, data_row['Application']

def show_processing_items (set,set_description):
    logging.info(f"{set_description}: {len(set)}")
    logging.info(f"-----------------------------------------------------------")
    for item in set:
        logging.info(f"* {item}")
    return len(set)

def main():
    #os.system('cls')
    
    parser = argparse.ArgumentParser(description="Cloud Analyzer application.")
    parser.add_argument('mode', type=str, choices=['cleanupdocs','createtrainingdata','convertdocx','preprocess-single', 'preprocess-all', 'all','single','postprocessing','postprocess'], help='Processing Mode')
    # Define optional arguments
    parser.add_argument('--destformat', type=str, choices=['text','markdown'], help='Convert Output Docs to [markdown|text]',default='markdown')
    parser.add_argument('--model', type=str, help='[OpenAI|PerplexityAI] ', default='openai')
    parser.add_argument('--live', type=str, choices=['yes', 'no'], default='no')
    parser.add_argument('--logging', type=str, choices=['debug', 'verbose', 'warning'], default='warning')
    parser.add_argument('--training', type=str, choices=['yes', 'no'], default='no')
    parser.add_argument('--env', type=str, choices=['showenv'], default='hideenv')  
    args = parser.parse_args()
    # Parse the arguments
    if args.env in ["showenv"]:
        print(f"Mode Parameter Value: {args.mode}")
        print(f"Model Parameter Value: {args.model}")
        print(f"Logging Parameter Value: {args.logging}")
        print(f"Training Parameter Value: {args.training}")
        print(f"Run Process Parameter Value: {args.live}")
    #CLI: python ./cloudarchitect-full.py --single --openai --verbose --train
    #CLI: python cloudarchitect-full.py --all --openai --verbose --train
    logger = logging.getLogger()
    if args.logging in ["verbose","VERBOSE", "v", "V"]:
        logger.setLevel(logging.INFO)
    elif args.logging in ["debug", "d", "D","DEBUG"]:
        logger.setLevel(logging.DEBUG)
    elif args.logging in ["warning", "WARNING","w", "W"]:
        logger.setLevel(logging.WARNING)
    else:
        logger.setLevel(logging.ERROR)
    console_handler = logging.StreamHandler()
    console_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
    logger.addHandler(console_handler)
    add_file_handler(logger, log_filename)
    if args.training in ["yes", "y", "Y", "YES", "Yes"]:
        md_training_data = read_md_files(training_data_path)
    else: 
        md_training_data=""
    if args.mode in ["preprocess-all","all"]:
        app_spreadsheet_full_path = os.path.join(root_data_path, spreadsheet_name)
        logging.info (f"Source Analysis File: {app_spreadsheet_full_path}")
        try:
            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                df_scope = pd.read_excel(app_spreadsheet_full_path)
                scope_set = set(df_scope['Application'])
                logging.info("Populating full scope data frame with all rows from Source Analysis File.")
                processed_rows = []
        except Exception as e:
            logging.error(f"Error reading the spreadsheet: {e}")
            sys.exit()
        headers = df_scope.columns.tolist()
        #processed_rows.append(headers)
        try:
            logging.info("Attempting to pick up where we left off.")
            if os.path.exists(output_file_path):
                logging.info(f"The file {output_file_path} already exists.")
                existing_df = pd.read_excel(output_file_path)
                if 'Application' in existing_df.columns:
                    existing_set = set(existing_df['Application'])
                else:
                    logging.warning("Column 'Application' not found in the existing file.")
                    existing_set = set()
            else:
                logging.info(f"The file {output_file_path} does not exist. Starting from the beginning.")
                existing_df = pd.DataFrame(columns=headers)
                existing_df.to_excel(output_file_path, index=False)
                existing_set = set()
        except pd.errors.EmptyDataError:
            logging.error(f"The existing file is empty.")
            # Handle empty file scenario
            existing_set = set()
        except Exception as e:
            logging.error(f"No Previous runs identified: {e}")            
            logging.info(f"The file {output_file_path} does not exist. Starting from the beginning.")
        unprocessed_items = list(scope_set - existing_set)
        num_existing_apps = show_processing_items (existing_set,"Already Processed")
        num_unprocessed_apps = show_processing_items (unprocessed_items,"Unprocessesed Items")
        num_total_apps = show_processing_items (scope_set,"Complete List")
        logging.info (f"Preparing to process {num_unprocessed_apps} out of {num_total_apps}. {num_existing_apps} completed in prior runs.")
        num_items_processed = 0
        num_items_skipped = 0
        num_items_excluded = 0
        num_items_pending = 0
        num_items_in_progress = 0
        num_items_already_complete = 0
        it_count = 0
        for index, row in df_scope.iterrows():
            row_dict = row.to_dict()
            it_count += 1
            if row_dict[discovery_process_status_column] == 'Complete' and row_dict[app_name_column] in unprocessed_items:
                if args.model in ["--openai", "-openai", "openai","OpenAI"]:
                    logging.info(f"Processing {row[app_name_column]} with OpenAI")
                    num_items_processed+=1
                    if args.live in ["yes","Yes","YES","Y","y"]:
                        processed_data, application_name = process_row_with_openai(args.mode, row, headers,md_training_data)
                        logging.info (f"Processed Data {processed_data}\n")
                        logging.info (f"Row {row}\n")
                        processed_rows.append(row_dict)
                elif args.model in ["--perplexityai", "-perplexityai", "perplexityai","--perplexity", "-perplexity", "perplexity"]:
                    logging.info(f"Processing {row[app_name_column]} with Perplexity.AI")
                    if args.live in ["yes","Yes","YES","Y","y"]:
                        processed_data, application_name = process_row_with_perplexityai(args.mode, row, headers)
                        logging.info (f"Processed Data {processed_data}")
                        logging.info (f"Row {row}")               
                        processed_rows.append(row_dict)
                else:
                    logging.warning("Missing Command Line Argument for Model type.")
                    sys.exit()
                output_df = pd.DataFrame(processed_rows)
                # Write the DataFrame to a new Excel file
                logging.info(f"Completed Processing {row[app_name_column]}")
                if args.live in ["yes","Yes","YES","Y","y"]:
                    logging.info(f"Processed Data: {processed_data}")
                    #Read existing data
                    #existing_df = pd.read_excel(output_file_path)
                    # Concatenate with the new data
                    combined_df = pd.concat([existing_df, output_df], ignore_index=True)
                    # Write the combined DataFrame back to Excel
                    combined_df.to_excel(output_file_path, index=False)
                    logging.info(f"Processed data written to {output_file_path}")  
                    create_word_document_with_completion(processed_data, application_name,application_summary_path)
            elif row_dict[discovery_process_status_column] == 'Complete':
                num_items_already_complete+=1
            elif row_dict[discovery_process_status_column] == 'Exclude':
                num_items_excluded+=1
            elif row_dict[discovery_process_status_column] == 'Pending':
                num_items_pending+=1
            elif row_dict[discovery_process_status_column] == 'In Progress':
                num_items_in_progress+=1
            else:            
                num_items_skipped+=1
        else:
            logging.info(f"{num_items_processed} Completed Processing")
            logging.info(f"{num_items_already_complete} Already Processed")
            logging.info(f"{num_items_pending} Skipped - Pending Discovery")
            logging.info(f"{num_items_in_progress} Skipped - Discovery In Progress")
            logging.info(f"{num_items_excluded} Excluded") 
            logging.info(f"{num_items_skipped} Skipped")
            sys.exit()
    if args.mode in ["single","preprocess-single"]:
        logging.debug("Processing a single row.\n")
        row=""
        headers=""
        if args.model in ["--openai", "-openai", "openai"]:
            processed_data, application_name = process_row_with_openai(args.mode, row, headers,md_training_data)
            logging.info (f"{processed_data}\n")
            create_word_document_with_completion(processed_data, application_name, application_summary_path)
        elif args.model in ["--perplexityai", "-perplexityai", "perplexityai","--perplexity", "-perplexity", "perplexity"]:
            processed_data, application_name = process_row_with_perplexityai(args.mode, row, headers)
            logging.info (f"{processed_data}\n")
            create_word_document_with_completion(processed_data, application_name, application_summary_path)
        else:
            logging.warning("Missing Command Line Argument for AI Model")
            sys.exit()

    if args.mode in ["postprocessing"]:        
        # Example usage
        file_path = root_docs_path
        convert_files_in_directory(file_path, convert_to=args.mode)
        sys.exit()

    if args.mode in ["createtrainingdata"]:        
        # Example usage
        file_path = output_file_path
        conversation_json = create_conversation_data(file_path)
        # Print the formatted conversation data
        logging.info("%s", conversation_json)
        sys.exit()

    if args.mode in ["cleanupdocs"]:        
        remove_openai_from_docx(root_docs_path)
        sys.exit()

    if args.mode in ["convertdocx"]:
        logging.debug(f"Begining .docx to {args.destformat} Conversion.")
        directory_path = root_docs_path  # Replace with your directory path
        if args.destformat in ["text"]:
            convert_files_in_directory(directory_path, convert_to=args.destformat)  # or 'markdown'
        else:
            convert_files_in_directory(directory_path, convert_to=args.destformat)  # or 'markdown'
        sys.exit()

    else:
        logging.warning("Missing Command Line Argument for Processing Mode (--all, --single)")
        sys.exit()
if __name__ == "__main__":
    main()