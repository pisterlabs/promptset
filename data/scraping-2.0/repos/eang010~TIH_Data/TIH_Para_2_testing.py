import requests
import json
import openai
from docx import Document

api_url = "https://api.stb.gov.sg/content/common/v2/search"
api_key = ""

query_params = {
    "dataset": "attractions",
    "limit": 2,
    "offset": 0,
}

headers = {
    "X-API-KEY": api_key
}

# Initialize a list to store selected dataset data
selected_data = []


    # Make the API request with the specified query parameters and headers
response = requests.get(api_url, params=query_params, headers=headers)

if response.status_code == 200:
    data = response.json()
    for item in data['data']:
                # Select the fields you want to include in the JSON
                selected_item = {
                    "name": item['name'],
                    "category": item['categoryDescription'],
                    "type": item['type'],
                    "description": item['description'],
                    "rating": item['rating'],
                    "nearestMrtStation": item['nearestMrtStation'],
                    "officialWebsite": item['officialWebsite'],
                    "admissionInfo": item['admissionInfo'],
                    "pricing": item.get('pricing', {}).get('others'),
                    "amenities": item['amenities'],
                    "address": " ".join(part for part in [item['address']['block'], item['address']['streetName'], item['address']['postalCode']] if part),
                    "businessHour": [
                        {
                            "day": business_hour['day'],
                            "openTime": business_hour['openTime'],
                            "closeTime": business_hour['closeTime'],
                            "description": business_hour['description']
                        }
                        for business_hour in item.get("businessHour", [])
                    ],
                    "businessHourNotes": item.get('businessHourNotes', {}).get('notes')
                }
                selected_data.append(selected_item)

            
else:
    print(f"Error: {response.status_code} - {response.text}")


# Sort selected_data by rating in descending order
selected_data = sorted(selected_data, key=lambda x: x['rating'], reverse=True)

# Create a dictionary to store the data for each item
items_data = {}

for item in selected_data:
    name = item["name"]
    items_data[name] = item

''' For testing
# Now you can access the data for each item by name
for item_name, item_data in items_data.items():
    item_string = f"Item Name: {item_name}\n{json.dumps(item_data, indent=4)}"
    print(item_string)
'''


# # # # # # # # # # # # Into Paragraphs # # # # # # # # # # # #

# Initialize OpenAI with your API key

openai.organization = ""
openai.api_key = ""

# Initialize a Word document
doc = Document()

# Messages for the initial conversation
messages = [
    {"role": "system", "content": "You are a writer."},
    {"role": "user", "content": "Based on the content provided, please make it into a paragraph."}
]

# Initialize an empty list to store individual responses
responses = []

# Split the JSON data into sections for input to GPT-3.5
import json
import time

for item_name, item_data in items_data.items():
    # Extract the category from item_data
    category = item_data.get("category", "Unknown Category")

    item_string = f"Item Name: {item_name}\n{json.dumps(item_data, indent=4)}\n"
    message = {
        "role": "assistant",
        "content": item_string
    }
    print(item_string)

    # Create a new list for messages for each item
    item_messages = list(messages)  # Create a copy of the initial messages

    item_messages.append(message)  # Append the message specific to this item

    # Create a completion with OpenAI GPT-3.5 Turbo using the item-specific messages
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=item_messages
    )

    response_content = response['choices'][0]['message']['content']
    responses.append((category, item_name, response_content))  # Append category, item name, and response

    # Add a delay of a few seconds (adjust as needed) between requests
    time.sleep(1)  # Adjust the delay as needed

# Create a separate paragraph in the Word document for each response
for category, item_name, response_content in responses:
    category_header = doc.add_heading(f"{category}", level=1)
    item_name_paragraph = doc.add_paragraph(f"Name: {item_name}")
    response_paragraph = doc.add_paragraph(response_content)
    doc.add_paragraph("\n")  # Add an empty line for separation

# Save the Word document
doc.save(f'{category}_Para.docx')