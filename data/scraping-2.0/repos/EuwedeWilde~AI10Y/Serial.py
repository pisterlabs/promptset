import serial
from PIL import Image, ImageTk, ImageDraw, ImageFont
import cv2
import openai
import torch
from diffusers import StableDiffusionImg2ImgPipeline
from deepface import DeepFace
import tkinter as tk
from tkinter import Label
import docx
from docx.shared import Inches, Pt, Mm
# from postermaker import create_a3_template, add_text, add_image, add_squares, add_texts_below_squares, add_image_process, add_text_process, add_logo, save_template

# Empty Cude Cache
torch.cuda.empty_cache()
print("cjh")
# Connect to Arduino
ser = serial.Serial('COM5', 115200, timeout=0.1)
print("cjh")


# Open camera and lie out foundation for face recognition ###########################################################
# Open the camera - 0 = built-in webcam or first camera, 1 is external webcam
cap = cv2.VideoCapture(0)
face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
radius = 50


# Check if the camera opened successfully
if not cap.isOpened():
    print("Error: Could not open camera.")
    exit()


# Generation settings ###############################################################################################
steps = 20 # amount of inference steps
GS = 10 # guidance scale, higher means more closely linked to the prompt (min 1)
NP = "human, people, nudity, naked" #negative prompt
stren = 0.5 #how much the image will be transformed (0-1)


# Connect AI's ######################################################################################################
# Set your OpenAI API key
openai.api_key = ""

#start diffusion model
model_id = "runwayml/stable-diffusion-v1-5" #model_id must be same as the one used in the folder the User.py file is in
print("c")
# start pipeline and send to GPU
device = "cuda" if torch.cuda.is_available() else "cpu"
print("c")
pipeline = StableDiffusionImg2ImgPipeline.from_pretrained(model_id, torch_dtype=torch.float16, safety_checker = None, requires_safety_checker = False)
print("c")
pipeline.to(device)

# Values to use in case assignment of customized values fails: ######################################################
emotions_for_display = "Main emotion: Happy, Secondary emotin: Surprise"
prompt = "Beautiful paining with inspiring brushstrokes. "
start_img = 'start1.png'
transformation_slider = 0.750
abstractness_slider = 0.5
main_emotions = "Mysteriously happy"
display_text = f"Your main facial expression seems to be Happy, and your secondary facial expression seems to be Neutral."
m_emo = "Neutral"
s_emo = "Happy"

print("loading check 1")

# Create a window for display: ######################################################################################
image_width = 320
image_height = 240
window_width = 600
window_height = 1024

window = tk.Tk()
window.configure(bg = "black")
window.title("Camera Feed and Prompt")
window.geometry(f"{str(window_width)}x{str(window_height)}")  # Adjust the dimensions as needed

# Create a label for displaying the camera feed
camera_label = Label(window)
camera_label.pack(pady=10)

# Create a label for displaying the prompt
prompt_label = Label(window, text="", font=("Helvetica", 14), wraplength = window_width, fg = "white", bg = "black")
prompt_label.pack()

# Create a label for displaying the 'webcam_capture.jpg' image
captured_image_label = Label(window)
captured_image_label.pack(pady=10)

print("initial code loade")




# Define required functions #########################################################################################
# Take a picture
def pic():
    print("start picure taking code")

    # Check if the frame was captured successfully
    if not ret:
        print("Error: Could not read frame.")
        exit()

    # Save the captured frame as an image
    cv2.imwrite("webcam_capture.jpg", frame)

    print("picture taken")

# Analyze emotions of picture taken
def analyze_emotions(input_pic):
    print("start emotino recognition code")
    global emotions_for_display

    frame = cv2.imread(input_pic)

    # Detect faces in the frame
    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    faces = face_cascade.detectMultiScale(gray, 1.3, 5)
    
    if len(faces) == 0:
        print("No face detected")
        return "No face detected", "Neutral", "Neutral"

    # Take the first detected face for analysis
    x, y, w, h = faces[0]
    
    # Crop the face from the frame
    cropped_face = frame[y:(y-radius)+(h+2*radius), x:(x-radius)+(w+2*radius)]
    
    # Save the cropped face as an image
    cropped_img_path = "cropped_face.jpg"
    cv2.imwrite(cropped_img_path, cropped_face)
    
    try:
        # Analyze the cropped face using DeepFace
        face_analysis = DeepFace.analyze(img_path=cropped_img_path, actions=['emotion'])
        
        # If face_analysis is a list, get the first item
        if isinstance(face_analysis, list):
            
            emotions = face_analysis[0]['emotion']
            print("### ANALYZE FACE completed")
        else:
            emotions = face_analysis['emotion']
        
        # Extracting the main and secondary emotions
        emotion_scores = face_analysis[0]['emotion']
        sorted_emotions = sorted(emotion_scores.items(), key=lambda x: x[1], reverse=True)
        main_emotion = sorted_emotions[0][0]
        secondary_emotion = sorted_emotions[1][0]

        print("main emotion is: " + main_emotion)
        print("secondary emotion is: " + secondary_emotion)
        print(f"{emotions}, {main_emotion}, {secondary_emotion}")
        return emotions, main_emotion, secondary_emotion #it is likely that also main and secondary emotion need to be returned
    
    except Exception as e:
        print("Error analyzing face:", e)
        return "No face detected", "Neutral", "Neutral"

# Generate prompt based on image taken
def generate_prompt(first, second):

    # Generate prompt
    print("start prompt generation code")
    response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a painter. You create image descriptions of maximum 65 words based on emotions provided."
                    },
                    {
                        "role": "user",
                        "content": f"Generate a very very (!!!) short image description based mainly on emotion: {first}, but also on emotion: {second}."
                    }
                ],
            )
    image_prompt = response.choices[0].message['content']
    image_prompt = image_prompt
    print(f"Generated Image Prompt: {image_prompt}")
    return image_prompt



# Generate new image based on prompt
def generate_image(generation_prompt, strenn, abstractness, img):
    print("start image generation code")
    
    # Handle abstractness
    abstractness = classify_art_movement(abstractness)
    # Update prompt
    generation_prompt = abstractness + generation_prompt 
    print("The final prompt used is: " + generation_prompt)

    # Use correct settings
    stren = strenn
    init_image = Image.open(img)
    init_image = init_image.convert('RGB')


    # Generate image
    print("Generating image...")
    image = pipeline(prompt = generation_prompt, num_inference_steps = steps, guidance_scale = GS, negative_prompt = NP, strength = stren, image = init_image).images[0]
    new_filename = "show_image.png"

    # Save image to display
    image.save(new_filename)

    #Use previous image as new image in case generation button is pressed more often
    init_image = Image.open(new_filename)

    print("generation completed: " + new_filename)

# Assigning art movements to different values of abstractness
def classify_art_movement(abstractness):
    print("Assigning art movements to different values of abstractness")
    if abstractness < 0.1:
        return "Suprematism, "
    elif abstractness < 0.2:
        return "Dadaism, "
    elif abstractness < 0.3:
        return "Surrealism, "
    elif abstractness < 0.4:
        return "Art Nouveau, "
    elif abstractness < 0.5:
        return "Post-Impressionism, "
    elif abstractness < 0.6:
        return "Neoclassicism, "
    elif abstractness < 0.7:
        return "Baroque, "
    elif abstractness < 0.8:
        return "Hudson River School, "
    elif abstractness < 0.9:
        return "Hyperrealism, "
    else:
        return "Digital High Realism, "
    

def create_a3_template():
    # A3 size in pixels (300 DPI)
    width, height = int(11.7 * 300), int(16.5 * 300)

    # Create a blank image with a white background
    template = Image.new("RGB", (width, height), (255, 255, 255))

    # Draw a square outline in the middle (250x250 mm)
    square_size = (int(250 / 25.4 * 300), int(250 / 25.4 * 300))
    square_position = ((width - square_size[0]) // 2, (height - square_size[1]) // 2)
    
    draw = ImageDraw.Draw(template)
    draw.rectangle([square_position, (square_position[0] + square_size[0], square_position[1] + square_size[1])], outline=(0, 0, 0), width=2)

    return template

def add_text(template, text, position, font_size=20, text_color=(0, 0, 0), font_path=None):
    # Add text to the template at the specified position
    draw = ImageDraw.Draw(template)
    font = ImageFont.truetype(font_path, font_size)
    draw.text(position, text, font=font, fill=text_color)

def add_image(template, image_path, square_size):
    # Open the image
    image = Image.open(image_path)

    # Calculate the scaling factors for width and height
    scale_width = square_size[0] / image.width
    scale_height = square_size[1] / image.height

    # Choose the minimum scaling factor to ensure the entire image fits within the square
    scale_factor = min(scale_width, scale_height)
    max_dimension = max(image.width, image.height)

    # Resize the image while maintaining its aspect ratio
    new_width = int(max_dimension * scale_factor)
    new_height = int(max_dimension * scale_factor)
    resized_image = image.resize((new_width, new_height))

    # Create a white background image with the specified square size
    background = Image.new("RGB", square_size, (255, 255, 255))

    # Calculate the position to center the resized image in the white background
    position = ((square_size[0] - new_width) // 2, (square_size[1] - new_height) // 2)

    # Paste the resized image onto the white background
    background.paste(resized_image, position)

    # Calculate the position to paste the white background onto the template
    template_position = ((template.width - square_size[0]) // 2, (template.height - square_size[1]) // 2)

    # Paste the white background with the resized image onto the template
    template.paste(background, template_position)

def add_squares(template, num_squares, square_size, spacing):
    # Calculate the total width of all squares and spacing
    total_width = num_squares * square_size[0] + (num_squares - 1) * spacing

    # Calculate the starting position to center the squares
    start_position = (template.width - total_width) // 2

    # Draw the squares and lines on the template and return square positions
    square_positions = []
    draw = ImageDraw.Draw(template)
    for i in range(num_squares - 1):  # Iterate up to num_squares - 1
        square_position = (start_position + i * (square_size[0] + spacing), template.height - square_size[1] - 220)
        draw.rectangle([square_position, (square_position[0] + square_size[0], square_position[1] + square_size[1])], outline=(0, 0, 0), width=2)
        square_positions.append(square_position)

        # Draw lines between squares
        line_start = (square_position[0] + square_size[0], square_position[1] + square_size[1] // 2)
        line_end = (square_positions[i][0] + square_size[0] + spacing, square_positions[i][1] + square_size[1] // 2)
        draw.line([line_start, line_end], fill=(0, 0, 0), width=2)

    # Add the last square without drawing a line
    last_square_position = (start_position + (num_squares - 1) * (square_size[0] + spacing), template.height - square_size[1] - 220)
    draw.rectangle([last_square_position, (last_square_position[0] + square_size[0], last_square_position[1] + square_size[1])], outline=(0, 0, 0), width=5)
    square_positions.append(last_square_position)

    return square_positions

def add_texts_below_squares(template, texts, square_positions, text_pos, font_size=20, text_color=(0, 0, 0), font_path=None):
    # Add text below each square based on the provided positions
    draw = ImageDraw.Draw(template)
    font = ImageFont.truetype(font_path, font_size)

    for i, text in enumerate(texts):
        # Calculate the bounding box for the text
        text_bbox = draw.textbbox(square_positions[i], text, font=font)
        
        # Calculate the center position for the text
        text_width = text_bbox[2] - text_bbox[0]
        text_position = (
            square_positions[i][0] + (square_size[0] - text_width) // 2,
            square_positions[i][1] + square_size[1] + text_pos
        )

        add_text(template, text, text_position, font_size=font_size, text_color=text_color, font_path=font_path)

def add_image_process(template, image_path, square_size, square_position):
    # Open the image
    image = Image.open(image_path)

    # Calculate the scaling factors for width and height
    scale_width = square_size[0] / image.width
    scale_height = square_size[1] / image.height

    # Choose the minimum scaling factor to ensure the entire image fits within the square
    scale_factor = min(scale_width, scale_height)
    max_dimension = max(image.width, image.height)

    # Resize the image while maintaining its aspect ratio
    new_width = int(max_dimension * scale_factor)
    new_height = int(max_dimension * scale_factor)
    resized_image = image.resize((new_width, new_height))

    # Create a white background image with the specified square size
    background = Image.new("RGB", square_size, (255, 255, 255))

    # Calculate the position to center the resized image in the white background
    position = ((square_size[0] - new_width) // 2, (square_size[1] - new_height) // 2)

    # Paste the resized image onto the white background
    background.paste(resized_image, position)

    # Calculate the position to paste the white background onto the template
    template_position = (
        square_position[0] + (square_size[0] - new_width) // 2,
        square_position[1] + (square_size[1] - new_height) // 2
    )

    # Paste the white background with the resized image onto the template
    template.paste(background, template_position)

def add_text_process(template, text, square_size, square_position, font_size=20, text_color=(0, 0, 0), font_path=None):
    # Add text to the template inside the specified square
    draw = ImageDraw.Draw(template)
    font = ImageFont.truetype(font_path, font_size)

    # Calculate the available width and height for the text
    available_width = square_size[0] - 10
    available_height = square_size[1] - 10

    # Break the text into lines to fit within the available width
    lines = []
    current_line = ""
    for word in text.split():
        # Check the width of the line with the new word
        test_line = current_line + word + " "
        text_bbox = draw.textbbox(square_position, test_line, font=font)
        text_width = text_bbox[2] - text_bbox[0]

        # If the line exceeds the available width, start a new line
        if text_width > available_width:
            lines.append(current_line.strip())
            current_line = word + " "
        else:
            current_line = test_line

    # Add the last line
    lines.append(current_line.strip())

    # Calculate the total height of the text
    total_text_height = sum(draw.textbbox(square_position, line, font=font)[3] - draw.textbbox(square_position, line, font=font)[1] for line in lines)

    # Calculate the remaining space above the text after subtracting the total text height
    remaining_space = available_height - total_text_height

    # Calculate the starting position for the text to be vertically centered
    start_position = (
        square_position[0] + (available_width - text_width) // 2,
        square_position[1] + remaining_space // 2
    )

    # Add each line of text to the template
    for line in lines:
        # Calculate the width of each line for horizontal centering
        line_width = draw.textbbox(start_position, line, font=font)[2] - draw.textbbox(start_position, line, font=font)[0]
        horizontal_position = square_position[0] + spacing // 10 + (available_width - line_width) // 2

        # Draw the line
        draw.text((horizontal_position, start_position[1]), line, font=font, fill=text_color)

        # Move the starting position to the next line
        start_position = (start_position[0], start_position[1] + (draw.textbbox(square_position, line, font=font)[3] - draw.textbbox(square_position, line, font=font)[1]))

def add_logo(template, logo_path, position, logo_size):
    # Open the logo image
    logo = Image.open(logo_path)

    # Resize the logo to the specified size
    logo = logo.resize(logo_size)

    # Paste the logo onto the template at the specified position
    template.paste(logo, position)

def save_template(template, output_path):
    # Save the full-size A3 template
    template.save(output_path)
    
def poster():
    global square_size
    global spacing
    # Create A3 poster:
    template = create_a3_template()

    # Variables from the Serial file

    #start_image_path = "start3.png" # start image
    #emo_detection = "Main emotion: Happy Secondary emotion: Neutral" # emotion detection
    #abstr_transf = "Abstractness: Dadaism Transformation: 14"  # abstractness and transformation
    #image_prompt = "A beautiful landscape bla bla placeholder text or something something lorem ipsum" # image prompt
    start_image_path = start_img # start image
    emo_detection = f"Main emotion: {m_emo} Secondary emotion: {s_emo}" # emotion detection
    abstr_transf = f"Abstractness: {abstractness_slider} Transformation: {transformation_slider}" # abstractness and transformation
    facial_expr_path = "cropped_face.jpg" # facial expression
    image_prompt = prompt # image prompt
    art_path = "show_image.png" # main image

    # Fonts
    fira_bold = "FiraCode-Bold.ttf"
    fira_semibold = "FiraCode-SemiBold.ttf"
    fira_medium = "FiraCode-Medium.ttf"
    bahnschrift = "BAHNSCHRIFT 1.TTF"

    # Add the image to the template
    add_image(template, art_path, (int(250 / 25.4 * 300), int(250 / 25.4 * 300)))

    # Add a title right above the image
    title_text = "AI10Y"
    add_text(template, title_text, (300, 300), font_size=500, font_path=fira_bold)

    # Add a subtitle right below the title
    subtitle_text = "AI GENERATED ART BASED ON YOU"
    add_text(template, subtitle_text, (300, 800), font_size=125, font_path=fira_semibold)

    # Add a logo
    logo_path = "logov2.jpg"
    logo_size = (500, 500)
    logo_position = ((template.width - logo_size[0] - 280), (410))
    add_logo(template, logo_path, logo_position, logo_size)

    # Process subheader
    process_text = "PROCESS"
    add_text(template, process_text, (300, template.height - 840), font_size=100, font_path=fira_semibold)

    # Process visual (squares)
    square_size = (450, 450)
    spacing = 170
    square_positions = add_squares(template, num_squares=5, square_size=square_size, spacing=spacing)

    # Specify texts for each step
    square_texts_1 = ["Chosen starting image", "User input detection", "Facial expression", "Image prompt", "Final artwork"]
    square_texts_2 = [" ", " ", "recognition", "generation", "generation"]
    add_texts_below_squares(template, square_texts_1, square_positions, 20, font_size=50, text_color=(0, 0, 0), font_path=bahnschrift)
    add_texts_below_squares(template, square_texts_2, square_positions, 70, font_size=50, text_color=(0, 0, 0), font_path=bahnschrift)

    # Add the input into squares
    add_image_process(template, start_image_path, square_size, square_positions[0]) # start image
    add_text_process(template, abstr_transf, square_size, square_positions[1], font_size=40, font_path=fira_medium) # abstractness & transformation
    add_image_process(template, facial_expr_path, (square_size[0] - 150, square_size[1] - 150), square_positions[2]) # facial expression recognition
    add_text_process(template, emo_detection, (square_size[0], square_size[1] + 300), square_positions[2], font_size=30, font_path=fira_medium) # facial expression recognition text
    add_text_process(template, image_prompt, square_size, square_positions[3], font_size=30, font_path=fira_medium) # image prompt
    add_image_process(template, art_path, square_size, square_positions[4]) # final artwork generation

    save_template(template, "AI10Y Poster.png")

t1 = "Introduction"
p1 = "This page hopes to answer some of the technical questions that might arise. To do so, we will go over the different steps that are taken to generate a personalized artwork based on your facial expressions. To find your own personalized artwork, please refer to the other side of this poster. "
t2 = "The different steps"
p2 = """Step 1: initial images
In the first part of the process, users can select a starting image. This image serves as a starting point in the image generation process. When Stable Diffusion, the image generation AI used, creates an image, it always uses another image as starting point, from which multiple steps are taken to change the image in such a way that with every step it looks more like the image description (prompt) provided. The more steps are taken, the less the initial image is visible in the final result. 

Step 2: settings 
In the settings section, transformation and abstractness can be influenced. Transformation determines the number of steps the AI takes to generate an image. When it is set lower, more of the starting image can be seen. Abstractness determines the art style that the final artwork will be made in. When set towards the abstract mark, the art style could be suprematism or dadaism, while a setting more towards the realistic mark will result in hyperrealism or the Hudson River school. These art styles are used and interpreted by Stable Diffusion. 

Step 3: facial expression recognition 
After taking a picture that the visitor is content with, they can push the accept button to let the AI interpret their emotions. Here, DeepFace compares the picture taken with its database, which results in a list of emotions with percentages of how strong the AI considers these emotions to be present in your picture. This is transformed to only show the main and secondary emotion detected. 

Step 4: generating image description
To generate a creative image description to be used to generate a new artwork, the list of emotions is sent to ChatGPT, which is asked to write a creative image description based on this information. ChatGPT's answer is saved as the image prompt to be used in step 5.

Step 5: generating the artwork 
To generate the final result, the following information is sent to Stable Diffusion: the image description generated by ChatGPT is used in combination with the starting image, the transformation setting to determine the number of steps that should be taken, the selected art style that results from the abstractness setting, as well as a number of settings that the used cannot influence such as the guidance scale (which determines how much the AI listens to what it is being asked to do), and a negative prompt (that forbids Stable Diffusion to generate human people). All of this information is considered, and after a few seconds a new artwork appears. 

Step 6: building the poster 
When the artwork has been generated, all settings, pictures and results are gathered and put into a poster for visitors to take home. This way, they can refer back to all steps, think more about their experience and learn more about AI.  
"""
t3 = "The technical process"
p3 = """In order to create this experience, various different techniques and products had to be combined. The physical aspects of the AI10Y dashboard (such as buttons and sliders) are connected to an Arduino Mega. Arduino runs a code that sends information of which buttons are pressed to Python, where the main code is hosted. Python refers to a number of downloaded libraries (large chunks of code) and files that enable us to generate images, make pictures, and analyze emotions. In order to connect our system to ChatGPT, we use the OpenAI API to communicate over the internet with ChatGPT using Python code.  
"""
def create_custom_style(doc, style_name, font_size):
    style = doc.styles.add_style(style_name, 1)  # Add a new paragraph style
    style.font.size = Pt(font_size)  # Set the font size
    return style

def save_as_docx():
    # Create an instance of a word document in A4 size
    doc = docx.Document()
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7) 
    doc.add_picture("AI10Y Poster.png", width = Mm((210-(2*25.4))), height = Mm((297-(2*25.4)-(2*12.7)))) # use image path, not image using PIL
    
    doc.add_page_break()

    custom_style = create_custom_style(doc, 'small_text', 9)

    doc.add_heading(t1, 3)
    doc.add_paragraph(p1, style = "small_text")
    doc.add_heading(t2, 3)
    doc.add_paragraph(p2, style = "small_text")
    doc.add_heading(t3, 3)
    doc.add_paragraph(p3, style = "small_text")

    
    
    # Now save the document to a location
    doc.save('gfg.docx')


def rotate_image(image, angle):
    (h, w) = image.shape[:2]
    center = (w / 2, h / 2)
    M = cv2.getRotationMatrix2D(center, angle, 1.0)
    rotated = cv2.warpAffine(image, M, (w, h))
    return rotated

# Main loop #########################################################################################################
while True:
    print("loop running")
    # Keep camera on all the time to reduce picture taking time and improve quality
    ret, frame = cap.read()
    frame = rotate_image(frame, -90)


 #.decode().strip().split(",")
    # Read serial port and split into an array
    data = ser.readline().decode('utf-8').rstrip().split(',')
    print(data)

    # Select starting image
    if data[0] == '0':
        if data[1] == '1':
            start_img = "start1.png"
        elif data[1] == '2':
            start_img = "start2.png"
        elif data[1] == '3':
            start_img = "start4.png"
        else: # data[1] == 3
            start_img = "start3.png"
        data[0] = "5" # reset data array

    # Update sliders
    elif data[0] == '1':
        
        abstractness_slider = float(data[1]) #using this float function is likely not required in the serial version of this code
        
        transformation_slider = float(data[2])
        transformation_slider = 0.5 + (0.5 * transformation_slider)
        data[0] = "5" # reset data array

    # Pictures and prompt generation
    elif data[0] == '2':
        if data[1] == 'c':
            pic()
        elif data[1] == 'r':
            pic()
        elif data[1] == 'a':
            main_emotions, m_emo, s_emo = analyze_emotions("webcam_capture.jpg")
            if m_emo == s_emo:
                display_text = "Your emotions could not be detected, please try again."
            else: 
                display_text = f"Your main facial expression seems to be {m_emo}, and your secondary facial expression seems to be {s_emo}"
        else: #data[1] == 'l'
            prompt = generate_prompt(m_emo, s_emo)
        data[0] = "5" # reset data array
            

    # Generate image
    elif data[0] == '3':
        display_text = "Your image is being generated..."
        prompt_label.config(text=display_text)
        window.update_idletasks()
        window.update()
        generate_image(prompt, transformation_slider, abstractness_slider, start_img)
        poster()
        save_as_docx()
        display_text = "Image generation complete!"
        data[0] = "5" # reset data array
        
    
    # Tkinter updates ########################################################################################
    # Must be at end of the while-loop to contain the most relevant contents
    # Convert the frame to RGB format for tkinter
    frame = cv2.resize(frame, (image_width,image_height)) # Resize the frame to fit within the window
    frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
    img = Image.fromarray(frame_rgb)
    imgtk = ImageTk.PhotoImage(image=img)

    # Update the camera feed label
    camera_label.imgtk = imgtk
    camera_label.config(image=imgtk)

    # Display the 'webcam_capture.jpg' image
    captured_image = Image.open('webcam_capture.jpg')
    captured_image = captured_image.resize((image_width,image_height)) 
    captured_imgtk = ImageTk.PhotoImage(image=captured_image)

    # Update the captured image label
    captured_image_label.imgtk = captured_imgtk
    captured_image_label.config(image=captured_imgtk)

    # Update the prompt label
    prompt_label.config(text=display_text)
    # Update the GUI
    window.update_idletasks()
    window.update()
    #################


