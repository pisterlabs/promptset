import os
import openai
import json
from dotenv import load_dotenv
from dev_tools import read as rd
from docx import Document


# Template for the manual_formatter
def create_recipe(title, portions, kcal, main_ingredients, spices, extras,
                  preparation_time, preparation_steps, cooking_steps, serving_steps
                  ) -> Document():
    # Create a new document
    document = Document()

    # Title
    document.add_heading(title, level=1)

    # Ingredients
    document.add_heading('Zutaten', level=2)
    document.add_paragraph(f'Für {portions} Portionen ca. {kcal} kcal')
    document.add_heading('Hauptzutaten', level=3)
    for ingredient in main_ingredients:
        document.add_paragraph(ingredient, style='List Bullet')
    document.add_heading('Gewürze', level=3)
    for spice in spices:
        document.add_paragraph(spice, style='List Bullet')
    document.add_heading('Weiteres', level=3)
    for extra in extras:
        document.add_paragraph(extra, style='List Bullet')

    # Preparation
    document.add_heading('Zubereitung', level=2)
    document.add_paragraph(f'Gesamtzeit ca. {preparation_time} Minuten')
    document.add_heading('Vorbereitung', level=3)
    for step in preparation_steps:
        document.add_paragraph(step, style='List Bullet')
    document.add_heading('Zubereitung', level=3)
    for step in cooking_steps:
        document.add_paragraph(step, style='List Bullet')
    document.add_heading('Anrichten', level=3)
    for step in serving_steps:
        document.add_paragraph(step, style='List Bullet')

    # Return the document .save(f'{title}.docx')
    return document


def function_call(content, description):
    api_response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": description},
                  {"role": "user", "content": content}],
        temperature=1,
        max_tokens=1024,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0)
    return api_response


def save_recipe(recipe):
    # Define the file path
    file_path = f'.//recipes//{recipe["title"]}.docx'

    # Create the recipe
    document = create_recipe(
        title=recipe['title'],
        portions=recipe['portions'],
        kcal=recipe['kcal'],
        main_ingredients=recipe['main_ingredients'],
        spices=recipe['spices'],
        extras=recipe['extras'],
        preparation_time=recipe['preparation_time'],
        preparation_steps=recipe['preparation_steps'],
        cooking_steps=recipe['cooking_steps'],
        serving_steps=recipe['serving_steps']
    )

    # Check if the directory exists, if not, create it
    directory = os.path.dirname(file_path)
    if not os.path.exists(directory):
        os.makedirs(directory)

    # Save the document
    document.save(file_path)


# Load the API key from the .env file.
load_dotenv()
openai.api_key = os.getenv('OPENAI_KEY')

# Make the API call
response = function_call(rd.text_data('../recipe.txt'), rd.text_data('recipe_description.txt'))

# Extract the details from the API response
recipe_details = json.loads(response['choices'][0]['message']['content'])

# Save the recipe
save_recipe(recipe_details)
