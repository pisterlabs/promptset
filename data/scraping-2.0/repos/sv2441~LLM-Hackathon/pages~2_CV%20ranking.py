import streamlit as st
from langchain.chat_models import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from dotenv import load_dotenv
import docx
import os
import pdfplumber

load_dotenv()

# os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]

os.environ["OPENAI_API_KEY"] = os.getenv('OPENAI_API_KEY')
chat_llm = ChatOpenAI(temperature=0.0, request_timeout=120)

def extract_text_from_pdf(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text()
    return text

def summary(text):
    title_template = """you are a HR Recruiter bot.you are given a text from resume. 
                Summarize the "{topic}" into 50 to 60 words including key skills and technology.
            """           
    prompt = ChatPromptTemplate.from_template(template=title_template)
    messages = prompt.format_messages(topic=text)
    response = chat_llm(messages)
    
    return response.content

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        return str(e)
      
def cv_rank(topic):
    file_path = "enhanced_jd.docx"
    jd = read_docx(file_path)
    title_template = """you are a HR Recruiter bot.
    "{topic}" is Resume summary . Score the Summary based on "{jd}". Give the Rate out of 10.
            """           
    prompt = ChatPromptTemplate.from_template(template=title_template)
    messages = prompt.format_messages(topic=topic,jd=jd)
    response = chat_llm(messages)
    
    return response.content

def save_as_docx(text, filename):
    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(filename)

def main():
    st.title("CV Ranking")

    uploaded_files = st.file_uploader("Upload PDF Resumes", type=["pdf"], accept_multiple_files=True)

    if uploaded_files:
        st.write("Uploaded Resumes:")
        for resume in uploaded_files:
            st.write(resume.name)
            text = extract_text_from_pdf(resume)
            response=summary(text)
            st.text(response) 
            # Display extracted text on the app
        if st.button("Rank"):
            rank=cv_rank(response)
            st.text(rank) 
        if st.button("Save"):
            save_as_docx(response, "summary.docx")
            
                
if __name__ == "__main__":
    main()

