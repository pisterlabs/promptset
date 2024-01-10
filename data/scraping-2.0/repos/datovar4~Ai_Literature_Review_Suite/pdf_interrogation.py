# -*- coding: utf-8 -*-
"""
Created on Sat May 13 19:00:28 2023

@author: David Tovar
"""

import fitz
import re
import numpy as np
import tensorflow_hub as hub
import openai
from sklearn.neighbors import NearestNeighbors
import tkinter as tk
import requests
import time
import tkinter.filedialog as filedialog
from docx import Document
from docx.shared import RGBColor
import os
import tkinter.simpledialog as simpledialog
from api_key_checker import check_api_key

# Set your OpenAI API Key here
check_api_key("OpenAI_API.txt", "OpenAI")
with open("OpenAI_API.txt", "r") as f:
    openai_api_key = f.read().strip()
    
openai.api_key = openai_api_key
    

   # Additional methods
def preprocess(text):
     text = text.replace('\n', ' ')
     text = re.sub('\s+', ' ', text)
     return text

def pdf_to_text(path, start_page=1, end_page=None):
    doc = fitz.open(path)
    total_pages = doc.page_count

    if end_page is None:
        end_page = total_pages

    text_list = []

    for i in range(start_page-1, end_page):
        text = doc.load_page(i).get_text("text")
        text = preprocess(text)
        text_list.append(text)

    doc.close()
    return text_list

def text_to_chunks(texts, word_length=250, start_page=1):
    text_toks = [t.split(' ') for t in texts]
    page_nums = []
    chunks = []

    for idx, words in enumerate(text_toks):
        for i in range(0, len(words), word_length):
            chunk = words[i:i+word_length]
            if (i+word_length) > len(words) and (len(chunk) < word_length) and (
                    len(text_toks) != (idx+1)):
                text_toks[idx+1] = chunk + text_toks[idx+1]
                continue
            chunk = ' '.join(chunk).strip()
            chunk = f'[Page no. {idx+start_page}]' + \
                ' ' + '"' + chunk + '"'
            chunks.append(chunk)
    return chunks

class SemanticSearch:

    def __init__(self):
        self.use = hub.load(
            'https://tfhub.dev/google/universal-sentence-encoder/4')
        self.fitted = False

    def fit(self, data, batch=1000, n_neighbors=5):
        self.data = data
        self.embeddings = self.get_text_embedding(data, batch=batch)
        n_neighbors = min(n_neighbors, len(self.embeddings))
        self.nn = NearestNeighbors(n_neighbors=n_neighbors)
        self.nn.fit(self.embeddings)
        self.fitted = True

    def __call__(self, text, return_data=True):
        inp_emb = self.use([text])
        neighbors = self.nn.kneighbors(inp_emb, return_distance=False)[0]

        if return_data:
            return [self.data[i] for i in neighbors]
        else:
            return neighbors

    def get_text_embedding(self, texts, batch=1000):
        embeddings = []
        for i in range(0, len(texts), batch):
            text_batch = texts[i:(i+batch)]
            emb_batch = self.use(text_batch)
            embeddings.append(emb_batch)
        embeddings = np.vstack(embeddings)
        return embeddings

# New function to load the recommender
def load_recommender(path, start_page=1):
    global recommender
    texts = pdf_to_text(path, start_page=start_page)
    chunks = text_to_chunks(texts, start_page=start_page)
    recommender.fit(chunks)
    return 'Corpus Loaded.'

# GUI
class QuestionAnsweringGUI(tk.Tk):
    def __init__(self, recommender):
        tk.Tk.__init__(self)
        self.current_pdf_directory = None
        self.doc = Document()  # initialize a Word document
        self.title("ChatGPT-3.5 Question Answering")
        self.recommender = recommender
        self.geometry("800x600")
        self.all_conversation = "\n"  # initialize with a line break

        self.question_var = tk.StringVar()

        self.entry = tk.Entry(
            self, textvariable=self.question_var, width=70)
        self.display = tk.Text(self, width=70, height=30, wrap=tk.WORD)
        self.display.configure(state='disabled')  # Make it read-only

        self.load_button = tk.Button(
            self, text="Load PDF", command=self.load_pdf)
        self.submit_button = tk.Button(
            self, text="PDF Specific Question", command=lambda: self.submit_question("PDF Specific"))
        self.submit_general_button = tk.Button(
            self, text="General Question", command=lambda: self.submit_question("General"))
        self.end_button = tk.Button(
            self, text="End", command=self.end_program)

        self.entry.pack()
        self.display.pack()
        self.load_button.pack()
        self.submit_button.pack()
        self.submit_general_button.pack()
        self.end_button.pack()

    def load_pdf(self):
        filepath = filedialog.askopenfilename(
            filetypes=[("PDF files", "*.pdf")])
        if filepath:
            load_recommender(filepath)
            self.display.configure(state='normal')
            filename = os.path.basename(filepath)
            self.current_pdf_directory = os.path.dirname(filepath)
            self.display.insert(tk.END, "\n")
            self.display.insert(tk.END, "\n")
            message = "Article: " + filename + "\n"
            self.display.insert(tk.END, message)
            self.all_conversation += "\n" + message  # add this line
            self.display.configure(state='disabled')

            # Log the loaded PDF name in the Word document
            run = self.doc.add_paragraph().add_run("\nArticle: " + filename + "\n")
            run.font.color.rgb = RGBColor(
                255, 0, 0)  # Red color for the text
            self.doc.add_paragraph()  # Add space after the loaded PDF name

    def submit_question(self, question_type):
        question = self.question_var.get()
        if question:
            # Generate an answer depending on the type of the question
            answer = self.generate_answer(
                question) if question_type == 'PDF Specific' else self.generate_text(question)

            # Add a couple of spaces from the previous question/answer
            self.display.configure(state='normal')
            self.display.insert(tk.END, "\n")
            self.display.configure(state='disabled')

            # Add the question type to the GUI
            self.display.configure(state='normal')
            self.display.insert(tk.END, f"\n({question_type}) ")
            self.display.configure(state='disabled')

            # Insert question and answer into the Text widget with a couple of spaces in between
            # Temporarily make it editable
            self.display.configure(state='normal')
            user_message = "\nUser: " + question
            gpt_message = "\nGPT-3.5: " + answer
            self.display.insert(tk.END, user_message + "\n" + gpt_message)
            self.all_conversation += "\nQuestion Type: " + question_type + \
                user_message + "\n" + gpt_message + "\n"  # add this line
            # Make it read-only again
            self.display.configure(state='disabled')

            # Log the question and answer in the Word document
            self.doc.add_paragraph()
            run = self.doc.add_paragraph().add_run(
                "\nQuestion Type: " + question_type + "\n")
            run.font.color.rgb = RGBColor(
                0, 0, 0)  # Black color for the text
            # Add question in a new paragraph
            self.doc.add_paragraph(user_message + "\n")
            # Add answer in a new paragraph
            self.doc.add_paragraph(gpt_message + "\n")
            # Add an empty paragraph for space
            self.doc.add_paragraph("\n")

            # Autoscroll to the end of the Text widget
            self.display.see(tk.END)

            # Clear the Entry widget
            self.question_var.set("")

    def generate_text(self, prompt, engine="gpt-3.5-turbo"):
        max_attempts = 20
        sleep_time = 30  # Number of seconds to wait between attempts
        for attempt in range(max_attempts):
            try:
                if engine == "gpt-3.5-turbo":
                    completions = openai.ChatCompletion.create(
                        model=engine,
                        messages=[
                            {"role": "system", "content": "You are an expert scientific writer who writes review articles for Nature Neuroscience Reviews. Provide answes in an single paragraph"},
                            {"role": "user", "content": prompt}
                        ],

                        max_tokens=300,
                        temperature=0.7
                    )

                else:
                    completions = openai.Completion.create(
                        engine=engine,
                        prompt=prompt,
                        max_tokens=700,
                        n=1,
                        stop=None,
                        temperature=0.7,
                    )
                message = completions.choices[0].text.strip(
                ) if engine != "gpt-3.5-turbo" else completions.choices[0].message['content'].strip()
                return message

            except requests.exceptions.RequestException as e:
                if 'rate limit' in str(e):
                    if attempt + 1 == max_attempts:
                        raise e  # If this was the last attempt, raise the exception
                    else:
                        # Wait for a while before trying again
                        time.sleep(sleep_time)
                else:
                    raise e  # If this was not a rate limit error, raise the exception

    def generate_answer(self, question):
        topn_chunks = self.recommender(question)

        prompt = ""
        prompt += 'search results:\n\n'
        for c in topn_chunks:
            prompt += c + '\n\n'

        prompt += f"For the above text, {question}. Compose a comprehensive reply to the question using the search results given. "\
                  "Cite each reference using [ Page Number] notation at the end of each sentence (every result has this number at the beginning). "\
                  "Citation should be done at the end of each sentence. If the search results mention multiple subjects "\
                  "with the same name, create separate answers for each. Only include information found in the results and "\
                  "don't add any additional information. Make sure the answer is correct and don't and don't output false content. "\
                  "If the text does not relate to the query, simply state 'Text Not Found in PDF'. Ignore outlier "\
                  "search results which has nothing to do with the question. Only answer what is asked. The "\
                  "answer should be short and concise. Answer step-by-step."

        print("Calling GPT3 API")
        answer = self.generate_text(prompt)
        return answer

    # Add this method
    def end_program(self):
        title = simpledialog.askstring(
            "Title", "Enter the conversation title:")
        if title is not None:  # If the user didn't click "Cancel"
            filename = os.path.join(
                self.current_pdf_directory, title) if self.current_pdf_directory else title
            document = Document(
                filename + ".docx") if os.path.exists(filename + ".docx") else Document()

            # Save conversation to the Word document in black
            document.add_paragraph().add_run(self.all_conversation)

            document.save(filename + ".docx")
        self.destroy()

if __name__ == "__main__":
    recommender = SemanticSearch()
    gui = QuestionAnsweringGUI(recommender)
    gui.mainloop()
