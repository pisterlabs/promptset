import os
import sys

import pandas as pd
import docx.oxml.ns as ns
from docx import Document
import re
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx.oxml as oxml
import openai
from datetime import datetime

from pandas import NaT

API_KEY_FILE = 'apikey.txt'

# 1. read csv file
def read_csv_file(file_path):
    try:
        # Read the CSV file into a DataFrame
        df = pd.read_csv(file_path)
        # Keep only the specified fields (columns)
        fields = ["Title", "Funder", "Deadline", "Amount", "Eligibility", "Abstract", "More Information"]
        filtered_df = df[fields]
        # remove html tag in each cell
        filtered_df = filtered_df.map(remove_html_tag)
        # Return the DataFrame
        return filtered_df
    except FileNotFoundError:
        print("File not found!")
        return None
    except KeyError:
        print("Some specified fields do not exist in the CSV file!")
        return None
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def remove_html_tag(value):
    if not isinstance(value, str):
        value = str(value)
    value = re.sub(r"<[^>]*>", "", value)
    return value

# 2. convert csv file to word file and format word file
def format_word_file(data_frame, head_title):
    get_api_key()

    # Create a new Word document for the formatted content
    formatted_doc = Document()
    
    # Add head_title as the first line, centered and bold
    if head_title:
        title = formatted_doc.add_paragraph(head_title)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(14)

    data_frame['ClosestFutureDate'] = data_frame['Deadline'].apply(extract_closest_future_date)
    # Sort the Deadline
    data_frame = data_frame.sort_values(by='ClosestFutureDate')
    
    # Iterate over each row in the DataFrame, skipping the first row
    for _, row in data_frame.iterrows():
        p = formatted_doc.add_paragraph()

        # title
        title_text = f"{row['Funder']} | {row['Title']}"
        if row['More Information']:
            hyperlink_run = add_hyperlink(p, row['More Information'], title_text)
            hyperlink_run.font.color.rgb = RGBColor(0, 0, 255)
            hyperlink_run.bold = True
            p.add_run("\n")
        else:
            title_run = p.add_run(title_text)
            title_run.bold = True
            p.add_run("\n")

        # Deadline
        if row['Deadline']:
            deadline_txt = row['Deadline']
            bold_run = p.add_run(f"Due Date: ")
            bold_run.bold = True
            closest_future_date = row['ClosestFutureDate']

            # Find the line in the text containing the closest future
            closest_date_line = [line for line in deadline_txt.split('\n')
                                 if closest_future_date is not NaT and
                                 closest_future_date.strftime("%d %b %Y") in line]
            closest_date = closest_date_line[0] if closest_date_line else "No upcoming date found"
            p.add_run(f"{closest_date}\n")

        # Amount
        if row['Amount']:
            amount = row['Amount']
            print(f'Summarizing amount: {amount}')
            try:
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",  # Specify the chat model
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant who's good at summarization."},
                        {"role": "user", "content": f"Summarize the following text by extracting award amount in"
                                                f"USD. Is amount upper exists, just use that number is enough."
                                                f"Do not include any notes or explanations:\n\n{amount}"}
                    ],
                    max_tokens=150,  # Set the maximum length for the summary
                    temperature=0.7  # Adjusts randomness in the response. Lower is more deterministic.
                )
                # The response format is different for chat completions
                summary = response['choices'][0]['message']['content'].strip()
                print(f'Summarized amount: {summary}')
            except Exception as e:
                print(f"API call failed for amount: {e}")
                summary = amount
            bold_run = p.add_run("Award Amount: ")
            bold_run.bold = True
            p.add_run(f"{summary}\n")

        # Eligibility
        if row['Eligibility']:
            eligibility = row['Eligibility']
            print(f'Summarizing eligibility: {eligibility}')
            try:
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",  # Specify the chat model
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant who's good at summarization."},
                        {"role": "user", "content": f"Summarize the following text by extracting which level "
                                                f"of faculty is eligible. If the level is not mentioned, "
                                                f"simply return Any level faculty"
                                                f"Also include information if this requires MD or PhD if the "
                                                f"information is available. Include the tenure information "
                                                    f"as well."
                                                    f" Don’t add any prefix like ‘Eligible faculty level:’."
                                                    f"Do not include any notes or explanations. And do not say things like:"
                                                    f"'MD or PhD is not mentioned. Tenure information is not provided.'"
                                                    f"Please be concise: "
                                                
                                                f" \n\n{eligibility}"}
                    ],
                    max_tokens=150,  # Set the maximum length for the summary
                    temperature=0.7  # Adjusts randomness in the response. Lower is more deterministic.
                )
                summary = response['choices'][0]['message']['content'].strip()
                print(f'Summarized eligibility: {summary}')
            except Exception as e:
                print(f"API call failed for Eligibility: {e}")
                summary = eligibility
            # The response format is different for chat completions
            bold_run = p.add_run("Eligibility: ")
            bold_run.bold = True
            p.add_run(f"{summary}\n")

        # Abstract
        if row['Abstract']:
            abstract = row['Abstract']
            print(f'Summarizing abstract: {abstract}')
            try:
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",  # Specify the chat model
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant who's good at summarization."},
                        {"role": "user", "content": f"Summarize the following text in a concise way."
                                                    f"Don’t do the explanation on the foundation itself. Must include "
                                                    f"the specific type of research. Don’t repeat the eligibility nor "
                                                    f"include the budget:\n\n{abstract}"}
                    ],
                    max_tokens=150,  # Set the maximum length for the summary
                    temperature=0.7  # Adjusts randomness in the response. Lower is more deterministic.
                )
                # The response format is different for chat completions
                summary = response['choices'][0]['message']['content'].strip()
                print(f'Summarized abstract: {summary}')
            except Exception as e:
                print(f"API call failed for Abstract: {e}")
                summary = abstract
            bold_run = p.add_run("Program Goal: ")
            bold_run.bold = True
            p.add_run(f"{summary}\n")

    # Save the formatted Word document
    # formatted_doc.save(output_file_path)
    return formatted_doc

def save_file(formatted_doc, output_file_path):
    # Save the formatted Word document
    formatted_doc.save(output_file_path)

# Helper functions:
# Add hyperlink
def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = oxml.OxmlElement('w:hyperlink')
    hyperlink.set(oxml.ns.qn('r:id'), r_id)

    new_run = oxml.OxmlElement('w:r')
    rPr = oxml.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    r = paragraph.add_run()
    r._r.append(hyperlink)

    return r

def extract_closest_future_date(deadline_txt):
    if deadline_txt:
        current_date = datetime.now()
        date_pattern = r"\d{2} \w{3} \d{4}"
        dates = re.findall(date_pattern, deadline_txt)
        future_dates = [datetime.strptime(date, "%d %b %Y") for date in dates if datetime.strptime(date, "%d %b %Y") > current_date]

        if future_dates:
            return min(future_dates, key=lambda x: (x - current_date))
    return None

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_path, relative_path)

def get_api_key():
    with open(resource_path(API_KEY_FILE), 'r') as file:
        openai.api_key = file.read().strip()

def unify_line_endings(file_path):
    with open(file_path, 'r', encoding='utf-8', newline=None) as file:
        content = file.read().replace('\r\n', '\n').replace('\r', '\n')

    with open(file_path, 'w', encoding='utf-8', newline='\n') as file:
        file.write(content)
        
def file_process(file_path, head_title):
    unify_line_endings(file_path)
    # 1. read csv file
    data_frame = read_csv_file(file_path)
    # 2. convert csv file to word file and format word file
    if data_frame is not None:
        formatted_doc = format_word_file(data_frame, head_title)
        return formatted_doc
    else:
        raise ValueError("No data found in the file")

if __name__ == "__main__":
    file_path = "sample_data/opps_export.csv"
    formatted_word_file_path = "output_word/formattedOutput.docx"
    head_title = "test"
    formatted_doc = file_process(file_path, head_title)
    save_file(formatted_doc, formatted_word_file_path)
