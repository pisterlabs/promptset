import discord
from discord import app_commands
import discord.ext.commands
import discord.ext.tasks

from utils.log_utils import colors
from utils.redis_utils import save_pickle_to_redis, load_pickle_from_redis
from utils.postgres_utils import fetch_key, fetch_keys_table, upsert_key, delete_key
from utils.tool_utils import dummy_sync_function
from tools import (
    get_image_from_search,
    get_organic_results,
    get_shopping_results,
    question_answer_webpage,
    summarize_webpage,
    get_full_blip,
    view_webpage_window,
)

import sys
import asyncio
from io import StringIO
from typing import Dict, Optional
from pydantic import BaseModel, Field

import time
import asyncio
import os
import openai
import datetime
from transformers import GPT2TokenizerFast
import re
import requests
import itertools
import pydot
import PyPDF2
import io
import textwrap
import chardet
import aiohttp
import logging
import subprocess
import pandas as pd
from docx import Document

from langchain.chat_models import ChatOpenAI
from langchain.llms import OpenAI
from langchain.chains.llm import LLMChain
from langchain.callbacks import get_openai_callback
from langchain.chains.conversation.memory import ConversationSummaryBufferMemory
from langchain.memory.buffer_window import ConversationBufferWindowMemory
from langchain.memory.token_buffer import ConversationTokenBufferMemory
from langchain.agents import Tool
from langchain.agents.conversational.base import ConversationalAgent
from langchain.agents.agent import AgentExecutor
from output_parser import ConvoOutputParser
from langchain.text_splitter import TokenTextSplitter
from langchain.schema import (
    AIMessage,
    HumanMessage,
    SystemMessage
)
from langchain.prompts.chat import (
    ChatPromptTemplate,
    SystemMessagePromptTemplate,
    AIMessagePromptTemplate,
    HumanMessagePromptTemplate,
)

from constants import (
    ORGANIC_RESULTS_ASK_TOOL_DESCRIPTION,
    QA_WEBPAGE_ASK_TOOL_DESCRIPTION,
    WEBPAGE_WINDOW_ASK_TOOL_DESCRIPTION,
    IMAGE_SEARCH_ASK_TOOL_DESCRIPTION,
    RECOGNIZE_IMAGE_ASK_TOOL_DESCRIPTION,
    SUMMARIZE_WEBPAGE_ASK_TOOL_DESCRIPTION,
    PYTHON_REPL_ASK_TOOL_DESCRIPTION,
    
    QA_WEBPAGE_CHAT_TOOL_DESCRIPTION,
    IMAGE_SEARCH_CHAT_TOOL_DESCRIPTION,
    ORGANIC_RESULTS_CHAT_TOOL_DESCRIPTION,
    RECOGNIZE_IMAGE_CHAT_TOOL_DESCRIPTION,
    SUMMARIZE_WEBPAGE_CHAT_TOOL_DESCRIPTION,
    
    get_ask_prefix,
    get_ask_custom_format_instructions,
    get_ask_suffix,
    
    get_chat_prefix,
    get_chat_custom_format_instructions,
    get_chat_suffix,
    
    get_thread_namer_prompt,
    
    FEATURES,
)

GOOGLE_API_KEY = os.getenv("GOOGLE_API_TOKEN")
GOOGLE_CSE_ID = os.getenv("GOOGLE_CSE_ID")
DISCORD_TOKEN = os.getenv("DISCORD_TOKEN") # load discord app token
GUILD_ID = os.getenv("GUILD_ID") # load dev guild
SERPAPI_API_KEY = os.getenv("SERPAPI_API_KEY")
NEWS_API_KEY = os.getenv("NEWS_API_KEY")
WOLFRAM_ALPHA_APPID = os.getenv("WOLFRAM_ALPHA_APPID")
DATABASE_URL = os.getenv("DATABASE_URL")

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

tokenizer = GPT2TokenizerFast.from_pretrained("gpt2") # initialize tokenizer

intents = discord.Intents.default() # declare intents
intents.message_content = True
intents.presences = False
intents.members = False

client = discord.Client(intents=intents)
tree = app_commands.CommandTree(client)

active_users = {} # dict of lists
active_names = {} # dict of strings

@client.event
async def on_ready():
    
    await fetch_keys_table()
    await tree.sync()

    logging.info(f"Registered {len(client.guilds)} guilds.")
    logging.info(f"Logged in as @{client.user.name}.")
    
@client.event
async def on_guild_join(guild):
    
    logging.info(f"Client added to guild {guild}.")
    
    await tree.sync(guild=guild)

@client.event
async def on_message(message):

    if message.author == client.user:
        return
    
    agent_mention = client.user.mention

    if "<@&1053339601449779383>" in message.content or agent_mention in message.content:
        
        global active_users
        global active_names
        
        active_users = await load_pickle_from_redis('active_users')
        chat_mems = await load_pickle_from_redis('chat_mems')
        
        # Get the current timestamp
        timestamp = datetime.datetime.now()
        time = timestamp.strftime(r"%Y-%m-%d %I:%M:%S")
        itis = timestamp.strftime(r"%B %d, %Y")
        clock = timestamp.strftime(r"%I:%M %p")

        channel_id = message.channel.id
            
        if channel_id not in chat_mems:
            chat_mems[channel_id] = None
        if channel_id not in active_users:
            active_users[channel_id] = []
            
        guild_name = message.guild
        if guild_name == None:
            guild_name = "DM"
        bot = client.user.display_name
        user_name = message.author.name
        id = message.author.id
        user_mention = message.author.mention
        prompt = message.content
        attachments = message.attachments
        openai_key = ""
        
        prompt = prompt.replace("<@1050437164367880202>", "")
        prompt = prompt.strip()
        
        attachment_text = ''
        file_placeholder = ''
        blip_text = ''
            
        async with message.channel.typing():
            
            result = await fetch_key(id)
            user_settings = await load_pickle_from_redis('user_settings')
            
            chat_model = user_settings.get(id, {}).get('model', 'gpt-3.5-turbo')
            #temperature = user_settings.get(id, {}).get('temperature', 0.5)
            temperature = 0.5
            
            if chat_model == "gpt-3.5-turbo-0613":
                chat_model = "gpt-3.5-turbo"
            
            max_tokens = 4096
            
            if chat_model == "gpt-4":
                max_tokens = 8192
            
            if result != None:
                openai.api_key=result
                openai_key=result
                
            else:
                embed = discord.Embed(description=f'<:ivanotify:1051918381844025434> {user_mention} Use `/setup` to register API key first or `/help` for more info. You can find your API key at [openai.com](https://beta.openai.com).', color=discord.Color.dark_theme())
                await message.channel.send(embed=embed)
                return
            
            text_splitter = TokenTextSplitter()
            logical_llm = ChatOpenAI(
                openai_api_key=openai_key,
                temperature=0,
                verbose=True,
                #callback_manager=manager,
                request_timeout=600,
                )
            
            async def parse_organic_results_input(url_comma_question):
                #a, b = url_comma_question.split(",", maxsplit=1)
                #answer = await get_organic_results(query=a, recency_days=int(b), llm=logical_llm)
                answer = await get_organic_results(query=url_comma_question, recency_days=None, llm=logical_llm)
                return f"{answer}"
            
            async def parse_qa_webpage_input(url_comma_question):
                a, b = url_comma_question.split(",", maxsplit=1)
                answer = await question_answer_webpage(url=a, question=b, llm=logical_llm)
                return f"{answer}\n"
            
            async def parse_summary_webpage_input(url):
                summary = await summarize_webpage(url, llm=logical_llm)
                return summary
            
            async def parse_blip_recognition(url_comma_question):
                a, b = url_comma_question.split(",", maxsplit=1)
                output = await get_full_blip(image_url=a, question=b)
                return output
            
            async def parse_view_webpage_input(url_comma_page_index):
                a, b = url_comma_page_index.split(",", maxsplit=1)
                output = await view_webpage_window(url=a, span_index=int(b))
                return output
            
            # STRINGIFY ACTIVE USERS
                
            if f"{user_name} ({user_mention})" not in active_users[channel_id]:
                active_users[channel_id].append(f"{user_name} ({user_mention})")
            
            active_names[channel_id] = ", ".join(active_users[channel_id])
            
            try:
                
                files = []

                if chat_model != "text-davinci-003":
                    chat_llm = ChatOpenAI(
                        temperature=temperature,
                        model_name=chat_model,
                        openai_api_key=openai_key,
                        request_timeout=600,
                        verbose=True,
                        )
                else:
                    chat_llm = OpenAI(
                        temperature=temperature,
                        model_name=chat_model,
                        openai_api_key=openai_key,
                        request_timeout=600,
                        verbose=True,
                    )

                tools = []
                
                def dummy_sync_function(tool_input: str) -> str:
                    raise NotImplementedError("This tool only supports async")
                
                tools.append(Tool(
                    name = "Search",
                    func=dummy_sync_function,
                    coroutine=parse_organic_results_input,
                    description=ORGANIC_RESULTS_CHAT_TOOL_DESCRIPTION,
                ))
                """
                tools.append(Tool(
                    name = "Summarize Webpage",
                    func=dummy_sync_function,
                    coroutine=parse_summary_webpage_input,
                    description=SUMMARIZE_WEBPAGE_CHAT_TOOL_DESCRIPTION,
                ))
                
                tools.append(Tool(
                    name = "Query Webpage",
                    func=dummy_sync_function,
                    coroutine=parse_qa_webpage_input,
                    description=QA_WEBPAGE_CHAT_TOOL_DESCRIPTION,
                ))
                """
                tools.append(Tool(
                    name = "Webpage",
                    func=dummy_sync_function,
                    coroutine=parse_view_webpage_input,
                    description=WEBPAGE_WINDOW_ASK_TOOL_DESCRIPTION,
                ))
                """
                tools.append(Tool(
                    name = "Vision",
                    func=dummy_sync_function,
                    coroutine=parse_blip_recognition,
                    description=RECOGNIZE_IMAGE_CHAT_TOOL_DESCRIPTION,
                ))
                
                tools.append(Tool(
                    name = "Images",
                    func=dummy_sync_function,
                    coroutine=get_image_from_search,
                    description=IMAGE_SEARCH_CHAT_TOOL_DESCRIPTION,
                ))
                """
                tool_names = [tool.name for tool in tools]
                
                prefix = await get_chat_prefix(active_names=active_names.get(channel_id, ''), itis=itis)
                
                custom_format_instructions = await get_chat_custom_format_instructions(tool_names=tool_names, user_name=user_name)
                
                suffix = await get_chat_suffix()
                
                if attachments != []:
                    for file in attachments:
                        
                        file_type = file.content_type
                        attachment_bytes = await file.read()
                        file_name = file.filename
                        
                        with open(f'{file_name}', 'wb') as f:
                            f.write(attachment_bytes)
                        
                        if file_type in ('image/jpeg', 'image/jpg', 'image/png'):
                            blip_text += f"\n\n{file_name} attached and saved to working directory: {file.url}"
                            file_placeholder += f"\n\n:frame_photo: **{file_name}**"
                        
                        elif "text/plain" in file_type: #txt
                            # Detect encoding
                            detected = chardet.detect(attachment_bytes)
                            encoding = detected['encoding']
                            # Decode using the detected encoding
                            raw_text = attachment_bytes.decode(encoding)
                            
                            file_tokens = len(tokenizer(prefix + custom_format_instructions + suffix + raw_text, truncation=True, max_length=12000)['input_ids'])

                            if file_tokens >= max_tokens:
                                
                                attachment_text += f"\n\n{file_name} is too large for you to view, but it has still been saved to the directory if you'd like to use Python REPL to interact with it. Here is a preview of the file:\n--- {file_name} ---\n\n{raw_text[:100]} [...]"
                                
                            else:
                                attachment_text += f"\n\n{file_name} has been saved to the working directory\n--- {file_name} ---\n\n{attachment_bytes.decode(encoding)}"
                                
                            file_placeholder += f"\n\n:page_facing_up: **{file_name}**"
                            
                        elif "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in file_type: #docx
                            
                            file_like_object = io.BytesIO(attachment_bytes)
                            
                            doc = Document(file_like_object)
                            full_text = []
                            for para in doc.paragraphs:
                                full_text.append(para.text)
                            raw_text = "\n".join(full_text)
                            
                            file_tokens = len(tokenizer(prefix + custom_format_instructions + suffix + raw_text, truncation=True, max_length=12000)['input_ids'])

                            if file_tokens >= max_tokens:
                                
                                attachment_text += f"\n\n{file_name} is too large for you to view, but it has still been saved to the directory if you'd like to use Python REPL to interact with it. Here is a preview of the file:\n--- {file_name} ---\n\n{raw_text[:100]} [...]"
                                
                            else:
                                attachment_text += f"\n\n{file_name} has been saved to the working directory\n--- {file_name} ---\n\n{raw_text}"

                            file_placeholder += f"\n\n:page_facing_up: **{file_name}**"
                        
                        elif "application/pdf" in file_type: #pdf

                            pdf_file = io.BytesIO(attachment_bytes)
                            pdf_reader = PyPDF2.PdfReader(pdf_file)
                            pdf_content = ""
                            for page in range(len(pdf_reader.pages)):
                                page_text = pdf_reader.pages[page].extract_text()
                                # Replace multiple newlines with a single space
                                page_text = re.sub(r'\n+', ' ', page_text)
                                pdf_content += page_text
                                
                            file_tokens = len(tokenizer(prefix + custom_format_instructions + suffix + pdf_content, truncation=True, max_length=12000)['input_ids'])

                            if file_tokens >= max_tokens:
                                
                                attachment_text += f"\n\n{file_name} is too large for you to view, but it has still been saved to the directory if you'd like to use Python REPL to interact with it. Here is a preview of the file:\n--- {file_name} ---\n\n{pdf_content[:100]} [...]"
                                
                            else:
                                attachment_text += f"\n\n{file_name} has been saved to the working directory\n--- {file_name} ---\n\n{pdf_content}"
                                
                            file_placeholder += f"\n\n:page_facing_up: **{file_name}**"
                            
                        elif "text/csv" in file_type: #csv
                            
                            try:
                                # Detect encoding
                                detected = chardet.detect(attachment_bytes)
                                encoding = detected['encoding']
                                # Decode using the detected encoding
                                raw_text = attachment_bytes.decode(encoding)
                                    
                                data = pd.read_csv(file_name)
                                
                                attachment_text += f"\n\n{file_name} has been saved to the working directory. Here is a preview of the file head:\n--- {file_name} ---\n\n{data.head()}"
                                    
                                file_placeholder += f"\n\n:page_facing_up: **{file_name}**"
                            except:
                                # Detect encoding
                                detected = chardet.detect(attachment_bytes)
                                encoding = detected['encoding']
                                # Decode using the detected encoding
                                raw_text = attachment_bytes.decode(encoding)
                                
                                attachment_text += f"\n\n{file_name} is too large for you to view, but it has still been saved to the directory if you'd like to use Python REPL to interact with it. Here is a preview of the file:\n--- {file_name} ---\n\n{raw_text[:100]} [...]"
                                
                                file_placeholder += f"\n\n:page_facing_up: **{file_name}**"
                            
                        else:
                            try:
                                # Detect encoding
                                detected = chardet.detect(attachment_bytes)
                                encoding = detected['encoding']
                                # Decode using the detected encoding
                                raw_text = attachment_bytes.decode(encoding)
                                
                                file_tokens = len(tokenizer(prefix + custom_format_instructions + suffix + raw_text, truncation=True, max_length=12000)['input_ids'])

                                if file_tokens >= max_tokens:
                                    
                                    attachment_text += f"\n\n{file_name} is too large for you to view, but it has still been saved to the directory if you'd like to use Python REPL to interact with it. Here is a preview of the file:\n--- {file_name} ---\n\n{raw_text[:100]} [...]"
                                    
                                else:
                                    attachment_text += f"\n\n{file_name} has been saved to the working directory\n--- {file_name} ---\n\n{attachment_bytes.decode(encoding)}"
                                    
                                file_placeholder += f"\n\n:page_facing_up: **{file_name}**"
                                
                            except:
                                embed = discord.Embed(description=f'<:ivanotify:1051918381844025434> {user_mention} the attachment\'s file type is unknown. consider converting it to plain text such as `.txt`.', color=discord.Color.dark_theme())
                                await message.channel.send(embed=embed)
                                return
                
                guild_prompt = ConversationalAgent.create_prompt(
                    tools=tools,
                    prefix=textwrap.dedent(prefix).strip(),
                    suffix=textwrap.dedent(suffix).strip(),
                    format_instructions=textwrap.dedent(custom_format_instructions).strip(),
                    input_variables=["input", "chat_history", "agent_scratchpad"],
                    ai_prefix = f"Iva",
                    human_prefix = f"",
                )
                
                if chat_mems[channel_id] != None:
                    
                    guild_memory = chat_mems[channel_id]
                    guild_memory.max_token_limit = 2000
                    guild_memory.ai_prefix = f"Iva"
                    guild_memory.human_prefix = f""
                    
                else:
                    
                    guild_memory = ConversationSummaryBufferMemory(
                        llm=chat_llm,
                        max_token_limit=2000,
                        memory_key="chat_history",
                        input_key="input",
                        ai_prefix = f"Iva",
                        human_prefix = f"",
                    )
                
                llm_chain = LLMChain(
                    llm=chat_llm,
                    verbose=True,
                    prompt=guild_prompt,
                )
                
                output_parser = ConvoOutputParser(
                    ai_prefix="Iva",
                )
                
                agent = ConversationalAgent(
                    llm_chain=llm_chain,
                    tools=tools,
                    verbose=True,
                    ai_prefix=f"Iva",
                    llm_prefix=f"Iva",
                    output_parser=output_parser,
                )
                
                agent_chain = AgentExecutor.from_agent_and_tools(
                    agent=agent,
                    tools=tools,
                    verbose=True,
                    memory=guild_memory,
                    ai_prefix=f"Iva",
                    llm_prefix=f"Iva",
                    max_execution_time=600,
                    #max_iterations=3,
                    #early_stopping_method="generate",
                    #return_intermediate_steps=False,
                )
                
                try:

                    reply = await agent_chain.arun(input=f"{user_name} ({user_mention}): {prompt}{attachment_text}")

                except Exception as e:
                    if str(e).startswith("Could not parse LLM output:"):
                        reply = str(e).replace("Could not parse LLM output: `", "")
                        reply = reply.replace("Thought: Do I need to use a tool? No", "")
                        reply = reply.strip("`")
                        mem_list = guild_memory.chat_memory.messages
                        extend_mems_list = [
                            HumanMessage(
                                content=prompt,
                                additional_kwargs={},
                            ),
                            AIMessage(
                                content=reply,
                                additional_kwargs={},
                            )]
                        mem_list.extend(extend_mems_list)
                        
                    else:
                        logging.error(e)
                        embed = discord.Embed(description=f'<:ivanotify:1051918381844025434> {user_mention} `{type(e).__name__}` {e}\n\nuse `/help` or seek https://discord.com/channels/1053335631159377950/1053336180692897943 if the issue persists.')
                        await message.channel.send(embed=embed)
                        return

            except Exception as e:
                logging.error(e)
                embed = discord.Embed(description=f'error', color=discord.Color.dark_theme())
                await message.channel.send(embed=embed)
                return
        try:
            
            reply = reply.replace("Iva: ", "")
            reply = reply.replace("Do I need to use a tool? No", "")
            
            if len(reply) > 2000:
                embed = discord.Embed(description=reply, color=discord.Color.dark_theme())
                await message.channel.send(embed=embed)
                return
            else:
                await message.channel.send(content=f"{reply}", files=files)
            
            chat_mems[channel_id] = guild_memory
            
            await save_pickle_to_redis('active_users', active_users)
            await save_pickle_to_redis('chat_mems', chat_mems)
        
        except Exception as e:
            logging.error(e)
            embed = discord.Embed(description=f'error', color=discord.Color.dark_theme())
            await message.channel.send(embed=embed)
            return
                

    return

class Opt(discord.ui.View):
    def __init__(self, key: str):
        super().__init__(timeout=60 * 60 * 24 * 365)
        self.value = None
        self.key = key
        
    async def on_timeout(self) -> None:

        for item in self.children:
            item.disabled = True

        await self.message.edit(view=self)
    
    @discord.ui.button(label="Agree and Continue", emoji="<:ivaup:1101609056604524594>", style=discord.ButtonStyle.grey)
    async def agree(self, interaction: discord.Interaction, button: discord.ui.Button):
        user_id = interaction.user.id
        mention = interaction.user.mention
        
        await upsert_key(str(user_id), self.key)
        
        embed = discord.Embed(description=f"<:ivathumbsup:1051918474299056189> **Key registered for {mention}. Welcome to Iva!**", color=discord.Color.dark_theme())
        await interaction.response.edit_message(embed=embed, delete_after=10, view=None)
        return

    @discord.ui.button(label="Disagree", emoji="<:ivadown:1101609054729666610>", style=discord.ButtonStyle.grey)
    async def disagree(self, interaction: discord.Interaction, button: discord.ui.Button):
        user_id = interaction.user.id
        mention = interaction.user.mention
        
        await delete_key(user_id)
        
        embed = discord.Embed(description=f"<:ivathumbsup:1051918474299056189> **You have opted out.**", color=discord.Color.dark_theme())
        await interaction.response.edit_message(embed=embed, delete_after=10, view=None)
        return
        
class Menu(discord.ui.View):
    def __init__(self):
        super().__init__(timeout=60 * 60 * 24 * 365)
        self.value = None
        
    async def on_timeout(self) -> None:
        # Step 2
        for item in self.children:
            item.disabled = True

        # Step 3
        await self.message.edit(view=self)
    
    @discord.ui.button(emoji="<:ivadelete:1095559772754952232>", style=discord.ButtonStyle.grey)
    async def delete(self, interaction: discord.Interaction, button: discord.ui.Button):
        
        guild_id = interaction.guild_id
        user_id = interaction.user.id
        channel_id = interaction.channel.id
        mention = interaction.user.mention
        
        ask_mems = await load_pickle_from_redis('ask_mems')
        
        if channel_id in ask_mems and user_id in ask_mems[channel_id] and ask_mems[channel_id][user_id]["user_id"] is not None:
            original_user_id = ask_mems[channel_id][user_id]["user_id"]
        else:
            embed = discord.Embed(description=f'<:ivanotify:1051918381844025434> {mention} You do not own this context line', color=discord.Color.dark_theme())
            await interaction.response.send_message(embed=embed, ephemeral=True, delete_after=10)
            return
        
        if original_user_id != user_id:
            embed = discord.Embed(description=f'<:ivanotify:1051918381844025434> {mention} You do not own this context line', color=discord.Color.dark_theme())
            await interaction.response.send_message(embed=embed, ephemeral=True, delete_after=10)
            return
        else:
            try:
                if channel_id in ask_mems and user_id in ask_mems[channel_id] and ask_mems[channel_id][user_id]["memory"] is not None:
                    
                    memory = ask_mems[channel_id][user_id]["memory"]
                    memory.chat_memory.messages = memory.chat_memory.messages[:-2]
                    await save_pickle_to_redis('ask_mems', ask_mems)
                    
            except Exception as e:
                embed = discord.Embed(description=f'<:ivanotify:1051918381844025434> {mention} `{type(e).__name__}` {e}\n\nuse `/help` or seek https://discord.com/channels/1053335631159377950/1053336180692897943 if the issue persists.')
                await interaction.channel.send(content=None, embed=embed)
        
        embed = discord.Embed(description=f'<:ivadelete:1095559772754952232>', color=discord.Color.dark_theme())
        await interaction.message.edit(content=None, embed=embed, view=None, delete_after=5)
        return
    
    @discord.ui.button(emoji="<:ivareset:1051691297443950612>", style=discord.ButtonStyle.grey)
    async def reset(self, interaction: discord.Interaction, button: discord.ui.Button):
        
        guild_id = interaction.guild_id
        channel_id = interaction.channel.id
        user_id = interaction.user.id
        mention = interaction.user.mention
        
        ask_mems = await load_pickle_from_redis('ask_mems')
        
        if channel_id in ask_mems and user_id in ask_mems[channel_id] and ask_mems[channel_id][user_id]["user_id"] is not None:
            original_user_id = ask_mems[channel_id][user_id]["user_id"]

        else:
            embed = discord.Embed(description=f'<:ivanotify:1051918381844025434> {mention} You do not own this context line', color=discord.Color.dark_theme())
            await interaction.response.send_message(embed=embed, ephemeral=True, delete_after=10)
            return

        if original_user_id != user_id:
            embed = discord.Embed(description=f'<:ivanotify:1051918381844025434> {mention} You do not own this context line', color=discord.Color.dark_theme())
            await interaction.response.send_message(embed=embed, ephemeral=True, delete_after=10)
            return
        else:
            
            if channel_id in ask_mems and user_id in ask_mems[channel_id] and ask_mems[channel_id][user_id]["memory"] is not None:
                
                ask_mems[channel_id][user_id]["memory"] = None
                
            if channel_id in ask_mems and user_id in ask_mems[channel_id] and ask_mems[channel_id][user_id]["last_message_id"] is not None:
                
                ask_mems[channel_id][user_id]["last_message_id"] = None
                
            await save_pickle_to_redis('ask_mems', ask_mems)
        
        embed = discord.Embed(description="<:ivareset:1051691297443950612>", color=discord.Color.dark_theme())
        button.disabled = True
        embeds = interaction.message.embeds
        attachments = interaction.message.attachments
        embeds.append(embed)
        await interaction.message.edit(view=None, embeds=embeds, attachments=attachments)
        #await interaction.channel.send(embed=embed)

@tree.command(name = "iva", description="write a prompt")
@app_commands.describe(prompt = "prompt", file_one = "file one", file_two = "file two", file_three = "file three")
async def iva(interaction: discord.Interaction, prompt: str, file_one: discord.Attachment = None, file_two: discord.Attachment = None, file_three: discord.Attachment = None):
    
    start_time = time.monotonic()
    
    guild_id = interaction.guild_id
    guild_name = interaction.guild
    user_id = interaction.user.id
    user = interaction.user
    channel_id = interaction.channel.id
    mention = interaction.user.mention
    bot = client.user.display_name
    user_name = interaction.user.name
    channel = interaction.channel
    
    try:
        
        await interaction.response.defer()
            
        # fetch the row with the given id
        result = await fetch_key(user_id)
        openai_key = ""
        
        if result != None:
            openai.api_key=result
            openai_key=result
        else:
            embed = discord.Embed(description=f'<:ivanotify:1051918381844025434> {mention} Use `/setup` to register API key first or `/help` for more info. You can find your API key at [openai.com](https://beta.openai.com).', color=discord.Color.dark_theme())
            await interaction.followup.send(embed=embed, ephemeral=True)
            return
        
        if isinstance(interaction.channel, discord.TextChannel):
            
            followup_message = await interaction.followup.send(content=channel.jump_url)
            await followup_message.delete()
            
            try:
                thread_namer = ChatOpenAI(temperature=1.0, openai_api_key=openai_key)
                template = await get_thread_namer_prompt(user_name)
                system_message_prompt = SystemMessagePromptTemplate.from_template(template)
                human_template = f"\"{{text}}\""
                human_message_prompt = HumanMessagePromptTemplate.from_template(human_template)
                chat_prompt = ChatPromptTemplate.from_messages([system_message_prompt, human_message_prompt])
                thread_namer_chain = LLMChain(llm=thread_namer, prompt=chat_prompt)
                
                thread_name = await thread_namer_chain.arun(f"{prompt}")
                thread_name = thread_name.strip("'").replace('.', '').replace('"', '').replace("Title: ", "")
                thread_name = thread_name[:100] #s lice if larger than 100 chars
                
                channel = await interaction.channel.create_thread(
                    type=discord.ChannelType.public_thread,
                    name=thread_name,
                )
                await channel.add_user(user)
                channel_id = channel.id
                
                thinking_message = await channel.send(content="<a:ivaloading:1102305649246867561>  iva is thinking...")
                
            except Exception as e:
                logging.error(e)
                embed = discord.Embed(description=f'<:ivanotify:1051918381844025434> {mention} `{type(e).__name__}` {e}\n\nuse `/help` or seek https://discord.com/channels/1053335631159377950/1053336180692897943 if the issue persists.')
                await interaction.followup.send(embed=embed, ephemeral=True)
                return
        
        default_user_data = {
            "last_message_id": None,
            "user_id": None,
            "memory": None,
        }

        user_settings = await load_pickle_from_redis('user_settings')
        ask_mems = await load_pickle_from_redis('ask_mems')
        
        ask_mems.setdefault(channel_id, {}).setdefault(user_id, default_user_data)
        
        chat_model = user_settings.get(user_id, {}).get('model', 'gpt-3.5-turbo')
        temperature = user_settings.get(user_id, {}).get('temperature', 0.5)
        
        if chat_model == "gpt-3.5-turbo-0613":
            chat_model = "gpt-3.5-turbo"
        
        if chat_model == "gpt-4":
            max_tokens = 8192
        elif chat_model == "gpt-3.5-turbo":
            max_tokens = 4096
        elif chat_model == "gpt-3.5-turbo-16k":
            max_tokens = 16384
        else:
            max_tokens = 4096
        
        # Get the current timestamp
        timestamp = datetime.datetime.now()
        itis = timestamp.strftime(r"%B %d, %Y")

        view = Menu()
        
        logical_llm = ChatOpenAI(
            openai_api_key=openai_key,
            temperature=0,
            verbose=True,
            #model_name=chat_model,
            #callback_manager=manager,
            request_timeout=600,
            )
        
        async def parse_organic_results_input(url_comma_question):
            #a, b = url_comma_question.split(",", maxsplit=1)
            #answer = await get_organic_results(query=a, recency_days=int(b), llm=logical_llm)
            answer = await get_organic_results(query=url_comma_question, recency_days=None, llm=logical_llm)
            return f"{answer}"
        
        async def parse_qa_webpage_input(url_comma_question):
            a, b = url_comma_question.split(",", maxsplit=1)
            answer = await question_answer_webpage(url=a, question=b, llm=logical_llm)
            return f"{answer}\n"
        
        async def parse_summary_webpage_input(url):
            summary = await summarize_webpage(url, llm=logical_llm)
            return summary
        
        async def parse_blip_recognition(url_comma_question):
            a, b = url_comma_question.split(",", maxsplit=1)
            output = await get_full_blip(image_url=a, question=b)
            return output
        
        async def parse_view_webpage_input(url_comma_page_index):
            url_comma_page_index = url_comma_page_index.strip("[").strip("]")
            url_comma_page_index = url_comma_page_index.strip()
            a, b = url_comma_page_index.split(",", maxsplit=1)
            output = await view_webpage_window(url=a, span_index=int(b))
            return output
        
        tools = []
        embeds = []
        files = []
        embeds_overflow = []
        files_overflow = []
        file_count=0
        
        class PythonREPL(BaseModel):
            """Simulates a standalone Python REPL."""
            globals: Optional[Dict] = Field(default_factory=dict, alias="_globals")
            locals: Optional[Dict] = Field(default_factory=dict, alias="_locals")

            def run(self, command: str) -> str:
                
                logging.info("using sync run repl")
                
                #command = autopep8.fix_code(command, options={"aggressive": 2})
                
                """Run command with own globals/locals and returns anything printed."""
                old_stdout = sys.stdout
                sys.stdout = mystdout = StringIO()
                try:
                    exec(command, self.globals, self.locals)
                    sys.stdout = old_stdout
                    output = mystdout.getvalue()
                except Exception as e:
                    sys.stdout = old_stdout
                    output = str(e)
                return output

            async def arun(self, command: str) -> str:
                
                logging.info("using async run repl")
                
                #command = autopep8.fix_code(command, options={"aggressive": 2})
                
                """Run command (sync or async) with own globals/locals and returns anything printed."""
                old_stdout = sys.stdout
                sys.stdout = mystdout = StringIO()

                async def run_sync_code():
                    loop = asyncio.get_running_loop()
                    return await loop.run_in_executor(None, self.run, command)

                async def run_async_code():
                    exec(
                        f"async def __arun_inner(scope):\n"
                        f"    async with scope:\n"
                        f"        {command}\n",
                        self.globals,
                        self.locals,
                    )
                    coroutine = self.locals["__arun_inner"](asyncio.get_event_loop())
                    await coroutine

                try:
                    if "async " in command:
                        logging.info("detected async code")
                        await run_async_code()
                    else:
                        logging.info("detected sync code")
                        await run_sync_code()
                    sys.stdout = old_stdout
                    output = mystdout.getvalue()
                except Exception as e:
                    logging.error(e)
                    sys.stdout = old_stdout
                    output = str(e)

                return output
        
        async def python_repl(command):
            
            command = command.strip().replace("```python", "").replace("```py", "").strip("```").replace(".show()", ".savefig('output.png')")
            
            if "!pip" in command:
                pip_install_lines = re.findall(r'^!pip install .*\n?', command, re.MULTILINE)
                command = re.sub(r'^!pip install .*\n?', '', command, flags=re.MULTILINE)

                for pip in pip_install_lines:
                    logging.info(f"PIP INSTALL COMMAND: {pip}")
                    # Handle pip install
                    package = pip.strip().split(' ')[-1]
                    result = subprocess.check_call([sys.executable, "-m", "pip", "install", package])
                    if result == 0:  # if the pip install command was successful
                        await python_repl(command)
                        return
            '''
            command = f"""try:
                {command}
            except Exception as e:
                print(str(e))"""
            '''
            logging.info(f"SANITIZED COMMAND: {command}")
            
            repl = PythonREPL()
            
            # Get the list of files before running the command
            before_files = set(os.listdir())
            
            try:
                output = await repl.arun(command)
            except Exception as e:
                logging.error(e)
                
            # Get the list of files after running the command
            after_files = set(os.listdir())
            logging.info(after_files)
                
            # Get the list of created files
            created_files = list(after_files - before_files)
            
            for file in created_files:
                if file.startswith("."):
                    continue
                else:
                    try:
                        output += f"{file} attached. "
                        logging.info(f"FILE ATTACHED {file}")
                        files.append(discord.File(fp=file))
                        os.remove(file)
                    except IsADirectoryError as e:
                        continue
            return output
            
        tools.append(Tool(
            name = "Search",
            func=dummy_sync_function,
            coroutine=parse_organic_results_input,
            description=ORGANIC_RESULTS_ASK_TOOL_DESCRIPTION,
        ))
        """
        tools.append(Tool(
            name = "Summarize Webpage",
            func=dummy_sync_function,
            coroutine=parse_summary_webpage_input,
            description=SUMMARIZE_WEBPAGE_ASK_TOOL_DESCRIPTION,
        ))
        """
        """
        tools.append(Tool(
            name = "Query Webpage",
            func=dummy_sync_function,
            coroutine=parse_qa_webpage_input,
            description=QA_WEBPAGE_ASK_TOOL_DESCRIPTION,
        ))
        """
        tools.append(Tool(
            name = "Webpage",
            func=dummy_sync_function,
            coroutine=parse_view_webpage_input,
            description=WEBPAGE_WINDOW_ASK_TOOL_DESCRIPTION,
        ))
        """
        tools.append(Tool(
            name = "Python REPL",
            func=dummy_sync_function,
            coroutine=python_repl,
            description=PYTHON_REPL_ASK_TOOL_DESCRIPTION,
        ))
        
        tools.append(Tool(
            name = "Vision",
            func=dummy_sync_function,
            coroutine=parse_blip_recognition,
            description=RECOGNIZE_IMAGE_ASK_TOOL_DESCRIPTION,
        ))
        
        tools.append(Tool(
            name = "Images",
            func=dummy_sync_function,
            coroutine=get_image_from_search,
            description=IMAGE_SEARCH_ASK_TOOL_DESCRIPTION,
        ))
        """
        tool_names = [tool.name for tool in tools]

        prefix = await get_ask_prefix(itis=itis)
        
        custom_format_instructions = await get_ask_custom_format_instructions(tool_names=tool_names)
        
        suffix = await get_ask_suffix()
        
        blip_text = ""
        
        attached_files = [file_one, file_two, file_three]
        
        attachment_text = ""
        file_placeholder = ""
        
        for file in attached_files:
        
            if file != None:
                
                attachment_bytes = await file.read()
                file_type = file.content_type
                file_name = file.filename
                
                with open(f'{file_name}', 'wb') as f:
                    f.write(attachment_bytes)
                    
                files.append(discord.File(f"{file_name}"))

                file_count += 1
                
                if file_type in ('image/jpeg', 'image/jpg', 'image/png'):
                    blip_text += f"\n\n{file_name} attached and saved to working directory: {file.url}"
                    file_placeholder += f"\n\n:frame_photo: **{file_name}**"
                
                elif "text/plain" in file_type: #txt
                    # Detect encoding
                    detected = chardet.detect(attachment_bytes)
                    encoding = detected['encoding']
                    # Decode using the detected encoding
                    raw_text = attachment_bytes.decode(encoding)
                    
                    file_tokens = len(tokenizer(prefix + custom_format_instructions + suffix + raw_text, truncation=True, max_length=12000)['input_ids'])

                    if file_tokens >= max_tokens:
                        
                        attachment_text += f"\n\n{file_name} is too large for you to view, but it has still been saved to the directory if you'd like to use Python REPL to interact with it. Here is a preview of the file:\n--- {file_name} ---\n\n{raw_text[:100]} [...]"
                        
                    else:
                        attachment_text += f"\n\n{file_name} has been saved to the working directory\n--- {file_name} ---\n\n{attachment_bytes.decode(encoding)}"
                        
                    file_placeholder += f"\n\n:page_facing_up: **{file_name}**"
                    
                elif "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in file_type: #docx
                    
                    file_like_object = io.BytesIO(attachment_bytes)
                    
                    doc = Document(file_like_object)
                    full_text = []
                    for para in doc.paragraphs:
                        full_text.append(para.text)
                    raw_text = "\n".join(full_text)
                    
                    file_tokens = len(tokenizer(prefix + custom_format_instructions + suffix + raw_text, truncation=True, max_length=12000)['input_ids'])

                    if file_tokens >= max_tokens:
                        
                        attachment_text += f"\n\n{file_name} is too large for you to view, but it has still been saved to the directory if you'd like to use Python REPL to interact with it. Here is a preview of the file:\n--- {file_name} ---\n\n{raw_text[:100]} [...]"
                        
                    else:
                        attachment_text += f"\n\n{file_name} has been saved to the working directory\n--- {file_name} ---\n\n{raw_text}"

                    file_placeholder += f"\n\n:page_facing_up: **{file_name}**"
                
                elif "application/pdf" in file_type: #pdf

                    pdf_file = io.BytesIO(attachment_bytes)
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    pdf_content = ""
                    for page in range(len(pdf_reader.pages)):
                        page_text = pdf_reader.pages[page].extract_text()
                        # Replace multiple newlines with a single space
                        page_text = re.sub(r'\n+', ' ', page_text)
                        pdf_content += page_text
                        
                    file_tokens = len(tokenizer(prefix + custom_format_instructions + suffix + pdf_content, truncation=True, max_length=12000)['input_ids'])

                    if file_tokens >= max_tokens:
                        
                        attachment_text += f"\n\n{file_name} is too large for you to view, but it has still been saved to the directory if you'd like to use Python REPL to interact with it. Here is a preview of the file:\n--- {file_name} ---\n\n{pdf_content[:100]} [...]"
                        
                    else:
                        attachment_text += f"\n\n{file_name} has been saved to the working directory\n--- {file_name} ---\n\n{pdf_content}"
                        
                    file_placeholder += f"\n\n:page_facing_up: **{file_name}**"
                    
                elif "text/csv" in file_type: #csv
                    
                    try:
                        # Detect encoding
                        detected = chardet.detect(attachment_bytes)
                        encoding = detected['encoding']
                        # Decode using the detected encoding
                        raw_text = attachment_bytes.decode(encoding)
                            
                        data = pd.read_csv(file_name)
                        
                        attachment_text += f"\n\n{file_name} has been saved to the working directory. Here is a preview of the file head:\n--- {file_name} ---\n\n{data.head()}"
                            
                        file_placeholder += f"\n\n:page_facing_up: **{file_name}**"
                    except:
                        # Detect encoding
                        detected = chardet.detect(attachment_bytes)
                        encoding = detected['encoding']
                        # Decode using the detected encoding
                        raw_text = attachment_bytes.decode(encoding)
                        
                        attachment_text += f"\n\n{file_name} is too large for you to view, but it has still been saved to the directory if you'd like to use Python REPL to interact with it. Here is a preview of the file:\n--- {file_name} ---\n\n{raw_text[:100]} [...]"
                        
                        file_placeholder += f"\n\n:page_facing_up: **{file_name}**"
                    
                else:
                    try:
                        # Detect encoding
                        detected = chardet.detect(attachment_bytes)
                        encoding = detected['encoding']
                        # Decode using the detected encoding
                        raw_text = attachment_bytes.decode(encoding)
                        
                        file_tokens = len(tokenizer(prefix + custom_format_instructions + suffix + raw_text, truncation=True, max_length=12000)['input_ids'])

                        if file_tokens >= max_tokens:
                            
                            attachment_text += f"\n\n{file_name} is too large for you to view, but it has still been saved to the directory if you'd like to use Python REPL to interact with it. Here is a preview of the file:\n--- {file_name} ---\n\n{raw_text[:100]} [...]"
                            
                        else:
                            attachment_text += f"\n\n{file_name} has been saved to the working directory\n--- {file_name} ---\n\n{attachment_bytes.decode(encoding)}"
                            
                        file_placeholder += f"\n\n:page_facing_up: **{file_name}**"
                        
                    except:
                        embed = discord.Embed(description=f'<:ivanotify:1051918381844025434> {mention} the attachment\'s file type is unknown. consider converting it to plain text such as `.txt`.', color=discord.Color.dark_theme())
                        if isinstance(interaction.channel, discord.TextChannel):
                            await thinking_message.edit(content=None, embed=embed)
                        else:
                            await interaction.followup.send(embed=embed, ephemeral=True)
                        return
            
        try:
            if channel_id in ask_mems and user_id in ask_mems[channel_id] and ask_mems[channel_id][user_id]["last_message_id"] is not None:
                
                original_message = await interaction.channel.fetch_message(ask_mems[channel_id][user_id]["last_message_id"])
                await original_message.edit(content="⠀", view=None)

        except discord.errors.HTTPException as e:
            logging.error(e)
        
        if chat_model != "text-davinci-003":
            ask_llm = ChatOpenAI(
                temperature=temperature,
                model_name=chat_model,
                openai_api_key=openai_key,
                request_timeout=600,
                verbose=True,
                #callback_manager=manager,
                #max_tokens=max_tokens,
                )
        else:
            ask_llm = OpenAI(
                temperature=temperature,
                model_name=chat_model,
                openai_api_key=openai_key,
                request_timeout=600,
                verbose=True,
            )
        
        k_limit = 3
        total_cost = None
        
        if channel_id in ask_mems and user_id in ask_mems[channel_id] and ask_mems[channel_id][user_id]["memory"] is not None:
            
            memory = ask_mems[channel_id][user_id]["memory"]
            
        else:
            
            memory = ConversationTokenBufferMemory(
                return_messages=False,
                human_prefix="User",
                ai_prefix="Iva",
                llm=ask_llm,
                memory_key="chat_history",
                max_token_limit=2000,
            )
            
            ask_mems[channel_id][user_id]["memory"] = None
        
        guild_prompt = ConversationalAgent.create_prompt(
            tools=tools,
            prefix=textwrap.dedent(prefix).strip(),
            suffix=textwrap.dedent(suffix).strip(),
            format_instructions=textwrap.dedent(custom_format_instructions).strip(),
            input_variables=["input", "chat_history", "agent_scratchpad"],
            ai_prefix = f"Iva",
            human_prefix = f"User",
        )
            
        llm_chain = LLMChain(
            llm=ask_llm,
            prompt=guild_prompt,
            verbose=True
        )
        
        output_parser = ConvoOutputParser(
            ai_prefix="Iva",
        )
            
        agent = ConversationalAgent(
            llm_chain=llm_chain,
            allowed_tools=tool_names,
            ai_prefix=f"Iva",
            output_parser=output_parser,
        )
        
        agent_chain = AgentExecutor.from_agent_and_tools(
            memory=memory,
            agent=agent,
            tools=tools,
            verbose=True,
            ai_prefix=f"Iva",
            max_execution_time=600,
            #max_iterations=3,
            #early_stopping_method="generate",
            #return_intermediate_steps=True,
        )
        
        url_pattern = re.compile(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+')
        links = url_pattern.findall(prompt)
        link_guidance = ""
        #if links:
        #    link_guidance = " (open the link with a tool)"
        
        try:
            
            with get_openai_callback() as cb:
                reply = await agent_chain.arun(input=f"{prompt}{link_guidance}{blip_text}{attachment_text}")
                total_cost = cb.total_cost
                
        except Exception as e:
            if str(e).startswith("Could not parse LLM output:"):
                logging.error(e)
                reply = str(e).replace("Could not parse LLM output: `", "")
                reply = reply.replace("Thought: Do I need to use a tool? No", "")
                reply = reply.strip("`")
                mem_list = memory.chat_memory.messages
                extend_mems_list = [
                    HumanMessage(
                        content=prompt,
                        additional_kwargs={},
                    ),
                    AIMessage(
                        content=reply,
                        additional_kwargs={},
                    )]
                mem_list.extend(extend_mems_list)
            else:
                logging.error(e)
                embed = discord.Embed(description=f'<:ivanotify:1051918381844025434> {mention} `{type(e).__name__}` {e}\n\nuse `/help` or seek https://discord.com/channels/1053335631159377950/1053336180692897943 if the issue persists.')
                if isinstance(interaction.channel, discord.TextChannel):
                    await thinking_message.edit(content=None, embed=embed)
                else:
                    await interaction.followup.send(embed=embed, ephemeral=True)
                return
            
        reply = reply.replace("Iva: ", "")
        reply = reply.replace("Do I need to use a tool? No", "")
        reply = reply.replace("```C#", "```cs")
        # Regex pattern for Markdown inline images
        pattern = r'\n?!\[.*?\]\(.*?\)'
        # Substituting the pattern with an empty string
        reply = re.sub(pattern, '', reply)
        
        dash_count = ""
        interaction_count = (len(memory.buffer)//2)-1
        
        for i in range(interaction_count):
            dash_count += "-"

        embed = discord.Embed(description=reply, color=discord.Color.dark_theme())
        
        file_count += 1
        
        if '$$' in reply or '```dot' in reply:

            # Use the findall() method of the re module to find all occurrences of content between $$
            dpi = "{200}"
            color = "{white}"
            
            tex_pattern = re.compile(r"\$\$(.*?)\$\$", re.DOTALL)
            dot_pattern = re.compile(r'```dot\s*([\s\S]*?)\s*```', re.DOTALL)
            
            tex_matches = tex_pattern.findall(reply)
            dot_matches = dot_pattern.finditer(reply)
            dot_matches = [match.group(1).strip() for match in dot_matches]
            
            non_matches = re.sub(r"```dot\s*[\s\S]*?\s*```|(\$\$|\%\%|\@\@).*?(\@\@|\%\%|\$\$)", "~~", reply, flags=re.DOTALL)

            non_matches = non_matches.split("~~")
            
            try:
                
                for (tex_match, dot_match, non_match) in itertools.zip_longest(tex_matches, dot_matches, non_matches):
                    
                    if non_match != None and non_match != "" and non_match != "\n" and non_match != "." and non_match != "\n\n" and non_match != " " and non_match != "\n> " and non_match.isspace() != True and non_match.startswith("![") != True:
                        
                        non_match = non_match.replace("$", "`")
                        non_match_embed = discord.Embed(description=non_match, color=discord.Color.dark_theme())
                        
                        if len(embeds) >= 9:
                            embeds_overflow.append(non_match_embed)
                        else:
                            embeds.append(non_match_embed)
                        
                    if tex_match != None and tex_match != "" and tex_match != "\n" and tex_match != " " and tex_match.isspace() != True:
                        
                        tex_match = tex_match.strip()
                        tex_match = tex_match.replace("\n", "")
                        tex_match = tex_match.strip("$")
                        tex_match = tex_match.split()
                        tex_match = "%20".join(tex_match)
                        match_embed = discord.Embed(color=discord.Color.dark_theme())

                        image_url = f"https://latex.codecogs.com/png.image?\dpi{dpi}\color{color}{tex_match}"

                        img_data = requests.get(image_url, verify=False).content
                        subfolder = 'tex'
                        if not os.path.exists(subfolder):
                            os.makedirs(subfolder)
                        with open(f'{subfolder}/latex{file_count}.png', 'wb') as handler:
                            handler.write(img_data)
                        tex_file = discord.File(f'{subfolder}/latex{file_count}.png')
                        match_embed.set_image(url=f"attachment://latex{file_count}.png")

                        file_count += 1
                        
                        if len(embeds) >= 9:
                            embeds_overflow.append(match_embed)
                            files_overflow.append(tex_file)
                        else:
                            embeds.append(match_embed)
                            files.append(tex_file)
                            
                    if dot_match != None and dot_match != "" and dot_match != "\n" and dot_match.isspace() != True:
                        
                        pattern = r'((di)?graph\s+[^{]*\{)'
                        replacement = r'\1\nbgcolor="#36393f";\nnode [fontcolor=white, color=white];\nedge [fontcolor=white, color=white];\n'
                        dot_match = re.sub(pattern, replacement, dot_match)
                        
                        graphs = pydot.graph_from_dot_data(dot_match)
                        
                        graph = graphs[0]
                        subfolder = 'graphviz'

                        if not os.path.exists(subfolder):
                            os.makedirs(subfolder)

                        graph.write_png(f'{subfolder}/graphviz{file_count}.png')
                        
                        dot_file = discord.File(f'{subfolder}/graphviz{file_count}.png')
                        match_embed = discord.Embed(color=discord.Color.dark_theme())
                        match_embed.set_image(url=f"attachment://graphviz{file_count}.png")
                        
                        file_count += 1

                        if len(embeds) >= 9:
                            embeds_overflow.append(match_embed)
                            files_overflow.append(dot_file)
                        else:
                            embeds.append(match_embed)
                            files.append(dot_file)
                    
            except Exception as e:
                logging.error(e)
        else:
            if len(reply) > 4096:
                try:
                    embeds = []
                    substrings = []
                    for i in range(0, len(reply), 4096):
                        substring = reply[i:i+4096]
                        substrings.append(substring)
                    for string in substrings:
                        embed_string = discord.Embed(description=string, color=discord.Color.dark_theme())
                        embeds.append(embed_string)
                except Exception as e:
                    logging.error(e)
                    embed = discord.Embed(description=f'<:ivaerror:1051918443840020531> **{mention} 4096 character response limit reached. Response contains {len(reply)} characters. Use `/reset`.**', color=discord.Color.dark_theme())
                    if isinstance(interaction.channel, discord.TextChannel):
                        await thinking_message.edit(content=None, embed=embed)
                    else:
                        await interaction.followup.send(embed=embed, ephemeral=True)
            else:
                embeds.append(embed)
            
        try:
            
            end_time = time.monotonic()
            
            word_count = len(reply.split())
            elapsed_time = end_time - start_time
            minutes, seconds = divmod(elapsed_time, 60)
            elapsed_time_format = f"{int(minutes):02}:{int(seconds):02}"
            
            if total_cost is not None:
                prompt_embed = discord.Embed(description=f"{dash_count}→ {prompt}{file_placeholder}\n\n`{chat_model}`  `{temperature}`  `{round(total_cost, 3)}`  `{elapsed_time_format}`  `{word_count}`")
            else:
                prompt_embed = discord.Embed(description=f"{dash_count}→ {prompt}{file_placeholder}\n\n`{chat_model}`  `{temperature}`  `{elapsed_time_format}`  `{word_count}`")
                
            embeds.insert(0, prompt_embed)
            
            if isinstance(interaction.channel, discord.TextChannel):
                await thinking_message.delete()
                initial_message = await channel.send(files=files, embeds=embeds, view=view)
                message_id = initial_message.id

            else:
                followup_message = await interaction.followup.send(files=files, embeds=embeds, view=view)
                message_id = followup_message.id
            
            if len(embeds_overflow) > 0:
                await channel.send(files = files_overflow, embeds=embeds_overflow)
            
            url_pattern = re.compile(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+')
            links = url_pattern.findall(reply)
            stripped_links = [link.rstrip(',.:)]') for link in links]
            
            if len(stripped_links) > 0:
                stripped_links = list(set(stripped_links))
                formatted_links = "\n".join(stripped_links)
                await channel.send(content=formatted_links)
            
            ask_mems[channel_id][user_id]["last_message_id"] = message_id
            ask_mems[channel_id][user_id]["user_id"] = user_id
            ask_mems[channel_id][user_id]["memory"] = memory
            await save_pickle_to_redis('ask_mems', ask_mems)
                
            return
        except Exception as e:
            logging.error(e, stack_info=True)
    except discord.errors.NotFound as e:
        logging.error(e, stack_info=True)

@tree.command(name = "reset", description="start a new conversation")
async def reset(interaction):
    
    channel_id = interaction.channel_id
    guild_id = interaction.guild_id
    user_id = interaction.user.id
    
    active_users = await load_pickle_from_redis('active_users')
    ask_mems = await load_pickle_from_redis('ask_mems')
    chat_mems = await load_pickle_from_redis('chat_mems')
    
    try:
        if channel_id in ask_mems and user_id in ask_mems[channel_id] and ask_mems[channel_id][user_id]["last_message_id"] is not None:
                
                original_message = await interaction.channel.fetch_message(ask_mems[channel_id][user_id]["last_message_id"])
                await original_message.edit(content="⠀", view=None)
            
    except discord.errors.HTTPException as e:
        logging.error(e)

    if channel_id in ask_mems and user_id in ask_mems[channel_id] and ask_mems[channel_id][user_id]["last_message_id"] is not None:
        
        ask_mems[channel_id][user_id]["last_message_id"] = None
        
    if channel_id in ask_mems and user_id in ask_mems[channel_id] and ask_mems[channel_id][user_id]["memory"] is not None:
        
        ask_mems[channel_id][user_id]["memory"] = None
        
    chat_mems[channel_id] = None
    active_users[channel_id] = []
    
    await save_pickle_to_redis('ask_mems', ask_mems)
    await save_pickle_to_redis('active_users', active_users)
    await save_pickle_to_redis('chat_mems', chat_mems)
    
    embed = discord.Embed(description="<:ivareset:1051691297443950612>", color=discord.Color.dark_theme())
    await interaction.response.send_message(embed=embed, ephemeral=False)

    
@tree.command(name = "help", description="get started")
async def help(interaction):
    
    mention = interaction.user.mention

    embed = discord.Embed(title=f"Welcome. Let's **Get Started**.\n\n", color=discord.Color.dark_theme())
    embed.set_thumbnail(url=client.user.avatar.url)
    embed.add_field(name="Step One", value="Iva uses **[OpenAI](https://beta.openai.com)** to generate responses. Create an account with them to start.")
    embed.add_field(name="Step Two", value="Visit your **[API Keys](https://beta.openai.com/account/api-keys)** page and click **`+ Create new secret key`**.")
    embed.add_field(name="Step Three", value=f"Copy and paste that secret key (`sk-...`) when you run `/setup` with {client.user.mention}")
    
    embed1 = discord.Embed(title="Step One", color=discord.Color.dark_theme())
    embed2 = discord.Embed(title="Step Two", color=discord.Color.dark_theme())
    embed3 = discord.Embed(title="Step Three", color=discord.Color.dark_theme())
    
    embed1.set_image(url="https://media.discordapp.net/attachments/1053423931979218944/1055535479140929606/Screenshot_2022-12-21_233858.png?width=960&height=546")
    embed2.set_image(url="https://media.discordapp.net/attachments/1053423931979218944/1055535478817947668/Screenshot_2022-12-21_234629.png?width=960&height=606")
    embed3.set_image(url="https://media.discordapp.net/attachments/1053423931979218944/1055535478507585578/Screenshot_2022-12-21_234900.png")
    
    await interaction.response.send_message(embeds=[embed, embed1, embed2, embed3], ephemeral=True)

@tree.command(name = "tutorial", description="how to talk with iva")
async def tutorial(interaction):
    
    mention = interaction.user.mention

    embed_main = discord.Embed(title="Introduction to Iva", description="there are two *separate* ways to talk to iva, both with their own conversation history: `@iva` and `/iva`. let's go over their differences, in addition to a other helpful tools.", color=discord.Color.dark_theme())
    embed_main.set_thumbnail(url=client.user.avatar.url)
    
    embed_chat = discord.Embed(title="`@iva`", description="provides **chat** and **conversation** oriented answers. has personality, asks questions back, is more creative.", color=discord.Color.dark_theme())

    embed_ask = discord.Embed(title="`/iva`", description="provides **academic** and **work** oriented answers. has less personality, is more focused on consistency and reliability.", color=discord.Color.dark_theme())
    #embed_ask.add_field(inline=True, name="<:ivacontinue1:1051714712242491392> `Continue`", value="say more, extend the last prompt's response")
    #embed_ask.add_field(inline=True, name="<:ivaregenerate:1051697145713000580> `Regenerate`", value="replace the last prompt's response with a different one")
    embed_ask.add_field(inline=True, name="<:ivadelete:1095559772754952232> `Delete`", value="delete the last interaction with iva in the conversation.")
    embed_ask.add_field(inline=True, name="<:ivareset:1051691297443950612> `Reset`", value="reset conversation history, clear iva's memory with you in the channel.")
    
    embed_other = discord.Embed(title="Other", color=discord.Color.dark_theme())
    embed_other.add_field(inline=True, name="`/reset`", value="reset `@iva` and `/iva` conversation history.")
    embed_other.add_field(inline=True, name="`/model`", value="switch between `gpt-4` and `gpt-3.5` models.")
    embed_other.add_field(inline=True, name="`/temperature`", value="change the temperature.")
    embed_other.add_field(inline=True, name="`/help`", value="show instructions for setup.")
    embed_other.add_field(inline=True, name="`/setup`", value="enter your key. `/help` for more info.")
    
    await interaction.response.send_message(embeds=[embed_main, embed_chat, embed_ask, embed_other], ephemeral=True)
    
@tree.command(name = "features", description="learn all the features iva has to offer")
async def tutorial(interaction):
    
    features_string = FEATURES
    
    features_intro = discord.Embed(title="Features", description="Becoming familiar with all Iva has to offer will allow you to maximize your workflow. This list is constantly being updated, so be on the look out!", color=discord.Color.dark_theme())
    features_intro.set_thumbnail(url=client.user.avatar.url)
    
    feature_list = discord.Embed(description=textwrap.dedent(features_string).strip(), color=discord.Color.dark_theme())

    embeds = [
        features_intro,
        feature_list,
    ]
    
    await interaction.response.send_message(embeds=embeds, ephemeral=True)
    
@tree.command(name = "setup", description="register your key")
@app_commands.describe(key = "key")
async def setup(interaction, key: str = None):

    id = interaction.user.id
    mention = interaction.user.mention
    
    if key is None:
        
        await delete_key(id)
        
        embed = discord.Embed(description=f"<:ivathumbsup:1051918474299056189> **Key deleted for {mention}.**", color=discord.Color.dark_theme())
        await interaction.response.send_message(embed=embed, ephemeral=True, delete_after=10)
        return

    # Use the `SELECT` statement to fetch the row with the given id
    result = await fetch_key(id)

    if result != None:

        # Access the values of the columns in the row
        if key != result[0]:
            
            await upsert_key(str(id), key)
            
            embed = discord.Embed(description=f"<:ivathumbsup:1051918474299056189> **Key updated for {mention}.**", color=discord.Color.dark_theme())
            await interaction.response.send_message(embed=embed, ephemeral=True, delete_after=10)
            return
            
        elif key == result[0]:
            
            embed = discord.Embed(description=f"<:ivaerror:1051918443840020531> **Key already registered for {mention}.**", color=discord.Color.dark_theme())
            await interaction.response.send_message(embed=embed, ephemeral=True, delete_after=10)
            return

    else:
        
        view = Opt(key=key)
        
        embed = discord.Embed(description=f"<:ivanotify:1051918381844025434> **{mention} In order to use Iva, you must agree to our [Privacy Policy](https://iva.gg/privacy) and [Terms of Service](https://iva.gg/terms)**.\n\nPlease take a few minutes to read and understand them both.", color=discord.Color.dark_theme())
        
        await interaction.response.send_message(embed=embed, ephemeral=True, view=view)
        
@tree.command(name = "model", description="choose a completion model")
@app_commands.choices(choices=[
        app_commands.Choice(name="gpt-3.5-turbo-4k ($0.002 / 1k tokens)", value="gpt-3.5-turbo"),
        app_commands.Choice(name="gpt-3.5-turbo-16k ($0.004 / 1k tokens)", value="gpt-3.5-turbo-16k"),
        app_commands.Choice(name="gpt-4-8k ($0.06 / 1k tokens)", value="gpt-4"),
    ])
async def model(interaction, choices: app_commands.Choice[str] = None):
    
    id = interaction.user.id
    mention = interaction.user.mention
    
    user_settings = await load_pickle_from_redis('user_settings')
    
    if choices is not None:
    
        user_settings.setdefault(id, {})['model'] = choices.value
        
        await save_pickle_to_redis('user_settings', user_settings)
        
        embed = discord.Embed(description=f"<:ivamodel:1096498759040520223> **set model to `{choices.value}` for {mention}.**", color=discord.Color.dark_theme())
        
    else:
        
        try:
            current_model = user_settings.get(id)["model"]
        except KeyError:
            current_model = "gpt-3.5-turbo"
        except TypeError:
            current_model = "gpt-3.5-turbo"
            
        embed = discord.Embed(description=f"<:ivamodel:1096498759040520223> **Current Model:** `{current_model}`", color=discord.Color.dark_theme())
    
    await interaction.response.send_message(embed=embed, ephemeral=True, delete_after=30)
    
    return
    
@tree.command(name = "temperature", description="set a default temperature to use with iva.")
@app_commands.describe(temperature = "temperature")
async def temperature(interaction, temperature: float = None):
    
    id = interaction.user.id
    mention = interaction.user.mention
    
    user_settings = await load_pickle_from_redis('user_settings')
    
    if temperature is not None:
    
        if not (temperature >= 0.0 and temperature <= 2.0):
            
            embed = discord.Embed(description=f"<:ivaerror:1051918443840020531> **{mention} `temperature` must be a float value from 0.0-2.0.**", color=discord.Color.dark_theme())
            
            await interaction.response.send_message(embed=embed, ephemeral=True, delete_after=30)
            
            return
        
        user_settings.setdefault(id, {})['temperature'] = temperature
        
        await save_pickle_to_redis('user_settings', user_settings)
        
        embed = discord.Embed(description=f"<:ivatemp:1097754157747818546>**set temperature to `{temperature}` for {mention}.**", color=discord.Color.dark_theme())
        
    else:
        
        try:
            temperature = user_settings.get(id)["temperature"]
        except KeyError:
            temperature = "0.5"
        except TypeError:
            temperature = "0.5"
        
        embed = discord.Embed(description=f"<:ivatemp:1097754157747818546>**Current Temperature:** `{temperature}`", color=discord.Color.dark_theme())
    
    await interaction.response.send_message(embed=embed, ephemeral=True, delete_after=30)
    
    return
    
client.run(DISCORD_TOKEN)