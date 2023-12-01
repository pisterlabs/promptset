# -*- coding: utf-8 -*-
"""Streamlit article generator.ipynb

Automatically generated by Colaboratory.

Original file is located at
    https://colab.research.google.com/drive/1nBLQHUOhjNvRm01MJ7RyIiaMLX4d5Ija
"""

import subprocess

"""# Article Generation"""

# Import necessary libraries
import openai
import pandas as pd
import requests
from docx import Document
from docx.shared import Inches
from PIL import Image
import cv2
import numpy as np
from io import BytesIO
import subprocess
import requests
import tempfile
from docx import Document
from docx.shared import Inches
from PIL import Image
import cv2
import numpy as np
from io import BytesIO
import subprocess
import io
import zipfile
import requests
import os

def get_terms(string):
    terms = []
    start_index = 0
    end_index = string.find('\n', start_index)
    while end_index != -1:
        term = string[start_index:end_index]
        if term.strip() != "":
            terms.append(term)
        start_index = end_index + 1
        end_index = string.find('\n', start_index)
    term = string[start_index:]
    if term.strip() != "":
        terms.append(term)
    return terms



def create_content(input):
  # Generate the article content using the GPT-3 API
  image_prompt = openai.Completion.create(model="text-davinci-003",
  prompt="write 4 image generation prompt starting the prompt with 'HQ, 4k,fine details cinematic intricate scenery, artistic, real photography of' of 15 words each to generate images using generative AI for an article about "+input,
  max_tokens=4000,
  temperature=0.9,top_p=1,n=1,)
  image_content = image_prompt.choices[0].text
  image_prompts=get_terms(image_content)
  image_urls = []
  # Get the URL of the generated image
  for i in range(0,len(image_prompts)):
    # Generate an image using the DALL-E API
    response = openai.Image.create(prompt=image_prompts[i],n=1,size="1024x1024")
    image_urls.append(response.data[0].url)
  # Generate an image using the DALL-E API
  response = openai.Image.create(prompt='realistic photography of '+input,n=4,size="1024x1024")
  image_urls.append(response.data[0].url)
  image_urls.append(response.data[1].url)
  image_urls.append(response.data[2].url)
  image_urls.append(response.data[3].url)
  pr="write with an adequate semantical structure and information structure to make it SEO compliant an article about "+str(input)+" with 4 different parts containing 400 words each and dealing with different aspects about "+str(input)
  # Generate the article content using the GPT-3 API
  completions = openai.Completion.create(model="text-davinci-003",prompt=pr,
  max_tokens=4000,temperature=0.9,top_p=1,n=1,)
  # Get the generated article content
  article_content = completions.choices[0].text
  # Create a DataFrame with the generated data
  data = {'image_url': image_urls, 'article_content': article_content,'image prompt':image_content}
  # Print the DataFrame
  return data


def create_docx(text, images_url, watermark_image_url):
    # Create a new Word document
    document = Document()

    # Add the text to the document
    document.add_paragraph(text)

    # Add the images to the document
    for url in images_url:
        response = requests.get(url)
        img = cv2.imdecode(np.frombuffer(response.content, np.uint8), -1)
        _, img_encoded = cv2.imencode('.jpg', img)
        # Create a file-like object from the encoded image data
        img = BytesIO(img_encoded)
        document.add_picture(img, width=Inches(6))

    # Add the watermark image to the top right corner of each page
    response = requests.get(watermark_image_url)
    img = cv2.imdecode(np.frombuffer(response.content, np.uint8), -1)
    _, img_encoded = cv2.imencode('.jpg', img)
    # Create a file-like object from the encoded image data
    img = BytesIO(img_encoded)
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    logo_run = paragraph.add_run()
    logo_run.add_picture(img, width=Inches(1))
    text_run = paragraph.add_run()
    text_run.text = '\t' + "This was made by creAIte" # For center align of text
    text_run.style = "Heading 2 Char"

    # Save the document
    document.save(input1+'.docx')
    subprocess.run(['pandoc' ,'-s' ,input1+'.docx','-o',input1+'.pdf'])

def create_docx(text, images_url, watermark_image_url):
    # Create a new Word document
    document = Document()

    # Add the text to the document
    document.add_paragraph(text)

    # Add the images to the document
    for url in images_url:
        response = requests.get(url)
        img = cv2.imdecode(np.frombuffer(response.content, np.uint8), -1)
        _, img_encoded = cv2.imencode('.jpg', img)
        # Create a file-like object from the encoded image data
        img = BytesIO(img_encoded)
        document.add_picture(img, width=Inches(6))

    # Add the watermark image to the top right corner of each page
    response = requests.get(watermark_image_url)
    img = cv2.imdecode(np.frombuffer(response.content, np.uint8), -1)
    _, img_encoded = cv2.imencode('.jpg', img)
    # Create a file-like object from the encoded image data
    img = BytesIO(img_encoded)
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    logo_run = paragraph.add_run()
    logo_run.add_picture(img, width=Inches(1))
    text_run = paragraph.add_run()
    text_run.text = '\t' + "This was made by creAIte" # For center align of text
    text_run.style = "Heading 2 Char"

    # Save the document to a temporary file
    with tempfile.TemporaryDirectory() as tempdir:
        temp_file = tempdir + '/temp.docx'
        document.save(temp_file)
        subprocess.run(['pandoc' ,'-s' ,temp_file,'-o',input1+'.pdf'])

def download_images(urls, zip_name):
  # create a zip file object
  zip_buffer = io.BytesIO()
  zip_file = zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED)

  # download each image and add it to the zip file
  for url in urls:
    response = requests.get(url)
    image_data = response.content

    # create a unique file name for the image
    file_name = url.split("/")[-1] + ".png"

    # add the image to the zip file
    zip_file.writestr(file_name, image_data)

  # close the zip file
  zip_file.close()

  # write the zip file to disk
  with open(zip_name, "wb") as f:
    f.write(zip_buffer.getvalue())

def imgshow(url,imgnum):
  response = requests.get(url)
  img = Image.open(BytesIO(response.content))
  imgnum=st.image(img, width=200)


def file_downloader(file_path, file_name):
    with open(file_path, "rb") as f:
        file_data = f.read()
    with open(file_name, "wb") as f:
        f.write(file_data)
    return file_name

import streamlit as st
import requests
from PIL import Image
from io import BytesIO

# Display some text
header=st.header('Article Generator')
st.write('This streamlit app was developped by CreAIte, the goal is to prototype what are the possibilities for content creation that generative AI bring to the table. Please feel free to use it and give us some feedback ! Please have fun creating with us !')


result=st.button('Generate your article')

input1=st.text_input('What is the subject of the article ? ')
input2=st.text_input('What is the name of the zip file you want for download ')
input3=st.text_input('What is your OPEN AI API Key ?')



if result==True:
  # Set the API key
  openai.api_key = input3
  data=create_content(input1)
  urls= data['image_url']
  create_docx(data['article_content'],data['image_url'],data['image_url'][0])
  zip_name = str(input2)+".zip"
  download_images(data["image_url"], zip_name)
  response1 = requests.get(urls[0])
  response2 = requests.get(urls[1])
  response3 = requests.get(urls[2])
  response4 = requests.get(urls[3])
  img1 = Image.open(BytesIO(response1.content))  
  img2 = Image.open(BytesIO(response2.content))  
  img3 = Image.open(BytesIO(response3.content))  
  img4 = Image.open(BytesIO(response4.content))
  imgshow1=st.image(img1, width=200)
  imgshow2=st.image(img2, width=200)
  imgshow3=st.image(img3, width=200)
  imgshow4=st.image(img4, width=200)
  
  st.write("List of all files in the '/content/' directory:")
  for root, dirs, files in os.walk("/content/"):
    for filename in files:
        st.write(os.path.join(root, filename))



  with open(zip_name, "rb") as file:
    btn = st.download_button(label="Download image zip file",data=file,file_name=input2+'.zip',mime="application/octet-stream")
  with open(input1+".docx", "rb") as file:
    btn = st.download_button(label="Download docx file",data=file,file_name=input1+'.docx',mime="application/octet-stream")
  with open(input1+".pdf", "rb") as file:
    btn = st.download_button(label="Download pdf file",data=file,file_name=input1+'.pdf',mime="application/octet-stream")
  for url in urls:
    response = requests.get(url)
    img = Image.open(BytesIO(response.content))
    st.image(img, width=200)

cwd=os.getcwd()

#if st.button("Download File.docx"):
#    file_path = cwd+ "/"+input1+".docx"
#    file_name = file_downloader(file_path, input1+".docx")
 #   st.success(f"{file_name} downloaded successfully")

#if st.button("Download File.pdf"):
#    file_path = cwd+ "/"+input1+".pdf"
 #   file_name = file_downloader(file_path, input1+".pdf")
  #  st.success(f"{file_name} downloaded successfully")