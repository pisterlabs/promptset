import os
import csv
import sys
sys.path.append("..")
sys.path.append("../library")
sys.path.append("../library/dicionarios")
import mdlDicLivros
import pandas as pd
import requests
import re
from bs4 import BeautifulSoup
from docx import Document
import json
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from openpyxl import Workbook
import openpyxl
from openpyxl.styles import Alignment

from docx import Document
from docx.shared import Inches

from docx.shared import RGBColor
import openai
#Deuteronomy.16.1
# criando o documento word
document = Document()



def formata(str, paragrafo,rgb,inches):
    # Remove all tags except "i", "br" and "n"

    text = str

    text = re.sub(r'<a[^>]*>([^<]+)</a>', '\\1', text)
    text = re.sub(r'<sup[^>]*>([^<]+)</sup>', '\\1', text)


   # text = "S HANGED IS A קללת אלהים — i.e., a degradation of the Divine King"
    #r = romanize3.__dict__['heb']
    #transliterated_text = r.convert(text)
   # print(transliterated_text)


   # texto = "ola <br/> <n> estou aqui </n> vc nao sabe"
    segmentado = re.split(r'(<[^>]*>)', text)
    print(segmentado)



    paragrafo.paragraph_format.left_indent = Inches(inches)
    tagi = False
    tagbr = False
    tagn = False
    for i in segmentado:
       # print(rashi_paragraph)
        if "<" in i:
            run = paragrafo.add_run("")
        else:
            run = paragrafo.add_run(i)
        run.font.color.rgb = rgb
        if "<i" in i:
            tagi = True

        if "<br/>"  in i:

            run.add_break()
        if "<br>" in i:
            run.add_break()

        if "<b>" in i:
            tagn = True


        if "</i>" in i:
            tagi = False

        if "</b>" in i:
            tagn = False

        if tagi:
            run.italic = True
        else:
            run.italic = False
        if tagn:
            run.bold = True
        else:
            run.bold = False

    return paragrafo;
        #run.font.color.rgb = RGBColor(0, 0, 255)
# iterando pelas linhas do dataframe
def extrair_passagemnvi(livro, cap, ver):
    # URL da passagem a ser extraída
    url = f'https://www.bibliaonline.com.br/nvi/{livro}/{cap}/{ver}'

    # Fazendo a requisição
    page = requests.get(url)
    soup = BeautifulSoup(page.content, 'html.parser')

    # Extrair o texto da passagem
    passagem = soup.find('p', class_='jss35')
    if passagem is None:
        print(f"Não foi possível encontrar a passagem {livro}/{cap}/{ver}")
        return None
    return passagem.get_text()

def extrair_passagem( cap, ver,livroreduzido,url):

    # URL da passagem a ser extraída
    url = f'https://www.biblegateway.com{url}'
    print(url)
    # Fazendo a requisição
    page = requests.get(url)
    soup = BeautifulSoup(page.content, 'html.parser')

    # Extrair o texto da passagem
    print(f'Extraindo passagem {livroreduzido}.{cap}.{ver}...')
    passagem2 = soup.find('span', class_=f'{livroreduzido}-{cap}-{ver}')
    if passagem2 is None:
        print(f"Não foi possível encontrar a passagem {livroreduzido}.{cap}.{ver}")
        return None
    return passagem2.get_text().replace(';', '.,')
# cria o dataframe


dfversoes = pd.DataFrame(columns=['textoSefaria', 'textoPortugues', 'textoOJB', 'bookAnc', 'chapterAnc', 'verseAnc'])
def pesquisar(livro, capitulo, versiculo,arquivo,linkObj):
    listtext = []
    cntMusar =0
    cntMidrash = 0
    cntTalmude = 0
    print('0.1')
    Versiculo = versiculo
    url = 'https://www.sefaria.org/api/related/' + livro+'.'+str(capitulo) + '.' + str(Versiculo) + '?with_sheet_links=1'
    print(url)
    response = requests.get(url)
    #provisorio, ]e preciso entender mais a fundo
    try:
      data = response.json()
    except:
        return
    df = pd.DataFrame(data["links"])
    df = df[['anchorRef', 'category', 'index_title', 'type', 'anchorVerse', 'sourceHasEn',
             'ref']]
    for index, row in df.iterrows():
        try:
            ref = row['ref'].split()
            book = ref[0]
            chapter = ref[1].split(':')[0]
            verse = ref[1].split(':')[1]
        except:

            book = ''
            chapter = ''
            verse = ''
        df.at[index, "book"] = book
        df.at[index, "chapter"] = chapter
        df.at[index, "verse"] = verse

    # seleciona as colunas desejadas

    print('0.2')

    for index, row in df.iterrows():
        rejeita = False
        try:
            if len(row['ref'].split(":"))> 2: #se tiver mais de 2 pontos na referencia rejeita
                if( row['ref'].split(":")[2]!='1'):
                    rejeita=True
        except:
            rejeita=True
        if (row['category'] == 'Talmud' or row['category'] == 'Targum' or row['category'] == 'Midrash'
            or row['category'] == 'Commentary') or (row['category']=='Musar') and rejeita == False:
            try:


                if row['category'] == 'Midrash' or row['category'] == "Talmud":
                    capComentario = row['ref'] #.split(':')[0].strip()

                else:
                    capComentario = row['ref']

                url = 'https://www.sefaria.org/api/texts/' + capComentario + '?commentary=0&context=1&pad=0&wrapLinks=1&wrapNamedEntities=1&multiple=0&stripItags=0&transLangPref=&firstAvailableRef=1&fallbackOnDefaultVersion=1'
                print(url)


                response = requests.get(url)
                data1 = response.json()
                print(row['category'])
                print(len(data1['text']))
                #se tiver ingles

                if(row['category'] == 'Targum'):
                    print(row['category'])
                    print (row['ref'])
                    text = data1['text'][int(row['ref'].split(':')[1])-1]

                    print(row['category'] +" sucesso")
                    print(text)
                elif row['category'] == 'Midrash' or row['category'] == "Talmud":
                    #print('FASE 1')
                    for it in range(0, len(data1['text'])):
                        texto = data1['text'][it]
                       # print('FASE 2'+ texto)
                        if (row['anchorRef'] in texto):
                            print(row['category'] + " sucesso")
                            #print('***********FASE 2' )
                            text = texto



                else:

                     text =  data1['text'][0]
                if (row['category'] == "Talmud"):
                    cntTalmude = cntTalmude + 1
                    if (cntTalmude > 3):
                        text = ''
                if (row['category'] == "Midrash"):
                    cntMidrash = cntMidrash + 1
                    if (cntMidrash > 3):
                        text = ''
                if (row['category'] == "Musar"):
                    cntMusar = cntMusar + 1
                    if (cntMusar > 3):
                        text = ''


            except:
              #  print('erro: '+ row['category'] + ' '+ row['ref'])
                textref = ''
                text = ''
        else:
            textref = ''
            text = ''


       # df.at[index, "textref"] = textref
        if text not in listtext:
            df.at[index, "text"] = text
            listtext.append(text)
        else:
            df.at[index, "text"] = ''


    print('0.3')



    df = df[df["text"].str.len() != 0]

    for index, row in df.iterrows():

            try:
                ref = row['anchorRef'].split()

                if len(ref)>2:
                    if(ref[0].strip()=='I'):
                        ref[0]='1'
                    if(ref[0].strip()=='II'):
                        ref[0]='2'

                   # print(ref[0].strip() + ' ' + ref[1].strip())
                   # print(mdlDicLivros.tradSefariaNvi[ref[0].strip() + ' ' + ref[1].strip()])
                    book = mdlDicLivros.tradSefariaNvi[ref[0].strip() + ' ' + ref[1].strip()]
                    chapter = ref[2].split(':')[0]
                    verse = ref[2].split(':')[1]

                else:
                    #print(ref[0] )

                    book = mdlDicLivros.tradSefariaNvi[ref[0].strip()]
                    chapter = ref[1].split(':')[0]
                    verse = ref[1].split(':')[1]
            except:
                print( 'nao achou: ' , ref[0])
                book = ''
                chapter = ''
                verse = ''
            df.at[index, "bookAnc"] = book
            df.at[index, "chapterAnc"] = chapter
            df.at[index, "verseAnc"] = verse
           # df.at[index, "textoptbr"] = extrair_passagemnvi(book, chapter, verse)
    print('0.4')

    #number of rows dataframe
    #provisorio
    if df.shape[0] == 0:
       return



    print(df)
    #livro, capitulo, versiculoro, capitulo, versiculo
    try:
        url = 'https://www.sefaria.org/api/texts/' + df.at[
            0, 'anchorRef'].split(":")[0] + ''
        print('traducao' + url)
        response = requests.get(url)
        data1 = response.json()

        textoSefaria = data1['text'][int(versiculo)-1]
    except:
        return
    #textoPortugues = extrair_passagemnvi(df.at[index1, "bookAnc"], df.at[index1, "chapterAnc"], df.at[index1, "verseAnc"])
    textoPortugues = extrair_passagemnvi( mdlDicLivros.tradSefariaNvi[livro], capitulo, versiculo)
   # textoOJB=extrair_passagem(df.at[index1, 'chapterAnc'], df.at[index1, 'verseAnc'],df.at[index1, "bookAnc"] ,linkObj)
    textoOJB = extrair_passagem(capitulo, versiculo, livro,
                                linkObj)


    dfversoes = pd.DataFrame(
        columns=['textoSefaria', 'textoPortugues', 'textoOJB', 'bookAnc', 'chapterAnc', 'verseAnc'])
    dfversoes=  dfversoes.append(
        {'textoSefaria': textoSefaria, 'textoPortugues': textoPortugues, 'textoOJB': textoOJB, 'bookAnc': mdlDicLivros.tradSefariaNvi[livro] ,
         'chapterAnc': capitulo, 'verseAnc': versiculo}, ignore_index=True)
#    #versoes.append({'textoPortugues': textoPortugues, 'textoSefaria': textoSefaria, 'textoOJB': textoOJB, 'chapterAnc': df.at[index1, 'chapterAnc'], 'verseAnc': df.at[index1, 'verseAnc'], 'bookAnc': df.at[index1, 'bookAnc']},ignore_index=True)
    dfversoes.to_csv(arquivo + 'b.csv', index=False,header=False, mode='a')

    print("pos ***** "+str(dfversoes.shape[0]))


    df = df[['anchorRef', 'category', 'index_title', 'type', 'anchorVerse', 'sourceHasEn','book','chapter','verse','ref',"bookAnc","chapterAnc","verseAnc","text" ]]

    print(df)
    df_talmud = df.loc[df['category'] == 'Talmud']
    print(df_talmud)

    df["text"].replace("\n", " ", inplace=True)
#anchorRef,category,index_title,type,anchorVerse,sourceHasEn,book,chapter,verse,ref,bookAnc,chapterAnc,verseAnc,text
    print('0.6')

    df.to_csv(arquivo+'.csv', index=False,header=False,mode='a')
    #exit()




def referencia(nomelivro, rangein,nomearquivo1):
    with open(nomearquivo1 + '.csv', 'w', newline='') as csvfile:
        writer = csv.writer(csvfile)
        # Escrever uma linha vazia
        writer.writerow(
            ['anchorRef', 'category', 'index_title', 'type', 'anchorVerse', 'sourceHasEn', 'book', 'chapter', 'verse',
             'ref', 'bookAnc', 'chapterAnc', 'verseAnc', 'text'])
    with open(nomearquivo1 + 'b.csv', 'w', newline='') as csvfile:
        writer = csv.writer(csvfile)
        # Escrever uma linha vazia
        writer.writerow(['textoSefaria', 'textoPortugues', 'textoOJB', 'bookAnc', 'chapterAnc', 'verseAnc'])

    dfnorder = pd.DataFrame(
        {'texto1': [], 'nvi': [], 'passagem': [], 'url': [], 'livroreduzido': [], 'Livro': [], 'Capitulo': [],
         'Versiculo': [], 'siglaBr': []})
    dfcorder = pd.DataFrame(
        {'texto1': [], 'nvi': [], 'passagem': [], 'url': [], 'livroreduzido': [], 'Livro': [], 'Capitulo': [],
         'Versiculo': [], 'siglaBr': []})

    arquivofinal = nomelivro
    arquivoordenado = nomelivro + 'od'
    arquivofinalordenado = nomelivro + 'odF'
    extensao = '.txt'
    arquivointermediario = arquivofinal + 'bf'

    #  arquivointermediario=arquivointermediario+extensao
    # Dicionário com as traduções das siglas em inglês para português



    # Cria uma lista vazia para armazenar as referências
    referencias = []
    referenciasfinal = []
    # Itera sobre os capítulos especificados
    nomelivropath = nomelivro.replace(' ', '%20')
    strlinks = []
    siglas = []
    a_tags = []
    for capitulo in rangein:

        url = f'https://www.biblegateway.com/passage/?search={nomelivropath}%20{capitulo}&version=OJB'
        print('url ' + url)
        page = requests.get(url)
        soup = BeautifulSoup(page.content, 'html.parser')

        # Seleciona todas as tags <a> com a classe "bibleref" e a propriedade "data-bibleref"https://www.biblegateway.com/passage/?search=Shemot%2024%3A7&version=OJB

        a_tags.extend(soup.select('a.bibleref[data-bibleref]'))
        versiculosRef=[] ;
        for tag in a_tags:
            #rint('tag '+tag.parent)
            print('tag '+tag.parent.text)
            if tag.parent.text not in versiculosRef:
                texto = tag.parent.text
                versiculosRef.append(texto)





        links = soup.find_all("a", href=lambda value: value and value.startswith("/passage/?search="))

        # Iterar sobre cada link e imprimir o conteúdo href
        for link in links:
            link.parent
            stringln = link["href"]
            array = stringln.split(".")
            sigla1 = array[0]
            # from mdlDicLivros import tradOjbSefaria, tradOjbNvi, tradNviOjb, tradSefariaOjb, tradNviSefaria,   tradSefariaNvi,tradOjbUrlToOjbLink,tradOjbLinkToOjbUrl
            sigla1 = sigla1.replace("/passage/?search=", "")
            siglas.append(sigla1)
            if sigla1 in mdlDicLivros.tradOjbUrlToOjbLink:
                strlinks.append(link["href"])
                # print(link["href"])
    print('atags ',a_tags[0])
        # Armazena o valor da propriedade "data-bibleref" na lista
    for a_tag in a_tags:
        referencias.append(a_tag['data-bibleref'])
    print(' ref '+referencias[0])

    nome_arquivo = arquivointermediario

    if True:

        # Itera sobre as referências e substitui os pontos por \
        referencias_expandidas = []
        for referencia in referencias:

            referencia = referencia.split(";")[0]
            if "-" in referencia:
                referencia1, referencia2 = referencia.split("-")
                verso1 = referencia1.split(".")[2]
                verso2 = referencia2.split(".")[2]
                cap = referencia2.split(".")[1]
                livro = referencia2.split(".")[0]

                inicio, fim = verso1, verso2
                for i in range(int(inicio), int(fim) + 1):
                    referencia_expandida = f"{livro}.{cap}.{i}"

                    referencias_expandidas.append(referencia_expandida)
            else:
                referencias_expandidas.append(referencia)
        # Remove duplicatas
        new_list = []
        for item in referencias_expandidas:
            if item not in new_list:
                new_list.append(item)
        referencias_expandidas = new_list
        print('ref_expandidas ',referencias_expandidas[0])
        versiculos = referencias_expandidas
        #referencias_expandidasfinal = []
        versiculosfinal = []
        for v in referencias_expandidas:
            try:
               # from mdlDicLivros import tradOjbSefaria, tradOjbNvi, tradNviOjb, tradSefariaOjb, tradNviSefaria, \
                #    tradSefariaNvi, tradOjbUrlToOjbLink, tradOjbLinkToOjbUrl
                print(v)
                versiculosfinal.append(v)

                                                #1,2,3
                dfnorder.loc[len(dfnorder)] = ['', '', v,
                                                #4
                                               f"/passage/?search={v.split('.')[0] + '.' + v.split('.')[1] + '.' + v.split('.')[2]}&version=OJB",
                                                  #5,6,7
                                               v.split('.')[0], mdlDicLivros.tradOjbSefaria[v.split('.')[0]], v.split('.')[1],
                                               v.split('.')[2], mdlDicLivros.tradOjbNvi[v.split('.')[0]]]
                # dfnorder   #{'texto1', 'passagem', 'url', 'livroreduzido', 'Livro', 'Capitulo', 'Versiculo'})
                #               dfnorder.append({'passagem': v, 'url': f"/passage/?search={v.split('.')[0] + '.' + v.split('.')[1] + '.' + v.split('.')[2]}&version=OJB" ,
                #                                    'livroreduzido':v.split('.')[0] , 'Livro': traducoes2[v.split('.')[0]], 'Capitulo': v.split('.')[1], 'Versiculo': v.split('.')[2]}, ignore_index=True)
         #       referencias_expandidasfinal.append(
         #           v + ";" + f"/passage/?search={v.split('.')[0] + '.' + v.split('.')[1] + '.' + v.split('.')[2]}&version=OJB" + ";" +
         #           v.split('.')[0] + ";" + mdlDicLivros.tradOjbSefaria[v.split('.')[0]] + ";" + v.split('.')[
          #              1] + ';' + v.split('.')[2])
            except:
                print('nao achou')
        print(dfnorder)  # print(referencias_expandidasfinal)

        # print(referencias_expandidasfinal)

        # Cria uma lista com a ordem dos livros na Bíblia
        ordem_livros = ["Gn", "Ex", "Lv", "Nm", "Dt", "Js", "Jz", "Rt", "1Sm", "2Sm", "1Rs", "2Rs", "1Cr", "2Cr", "Ed",
                        "Ne", "Et", "Jó", "Sl", "Pv", "Ec", "Ct", "Is", "Jr", "Lm", "Ez", "Dn", "Os", "Jl", "Am", "Ob",
                        "Mq", "Na", "Hc", "Sf", "Ag", "Zc", "Ml"]

        # Ordena a lista de versículos usando a ordem dos livros

        versiculos_ordenados = []
        for v in versiculosfinal:
            print(livro)
            v1 = v
            v = v.split(";")[0]

            parts = v.split(".")
            livro = parts[0]
            if len(parts) != 3:
                print("Versiculo invalido: ", v)
                continue

            try:
                if mdlDicLivros.tradOjbNvi[livro] not in ordem_livros:
                    print("Livro invalido: ", livro)
                    continue
            except:
                print("Livro invalido: ", livro)
                continue
            try:
                cap = int(parts[1])
                verso = int(parts[2])
            except ValueError:
                print("Capitulo ou versiculo invalido: ", v)
                continue
            versiculos_ordenados.append(f"{parts[0]}.{parts[1]}.{parts[2]}")
            # versiculos_ordenados.append(v1)
        # print(versiculos_ordenados)


        # versiculos_ordenados = sorted(versiculos_ordenados, key=lambda x: (ordem_livros.index(x[0]), x[1], x[2]))
        # print(versiculos_ordenados)
        versiculos_ordenados = sorted(versiculos_ordenados, key=lambda x: (
        ordem_livros.index(mdlDicLivros.tradOjbNvi[x.split(".")[0]]), int(x.split(".")[1]), int(x.split(".")[2].split(";")[0])))
        versiculos_ordenadosfinal = []
        for v in versiculos_ordenados:
            #   , dfcorder
            dfcorder.loc[len(dfcorder)] = ['', '', v,
                                           f"/passage/?search={v.split('.')[0] + '.' + v.split('.')[1] + '.' + v.split('.')[2]}&version=OJB",
                                           v.split('.')[0], mdlDicLivros.tradOjbSefaria[v.split('.')[0]], v.split('.')[1],
                                           v.split('.')[2], mdlDicLivros.tradOjbNvi[v.split('.')[0]]]
            # {'texto1', 'passagem', 'url', 'livroreduzido', 'Livro', 'Capitulo', 'Versiculo'})
            #versiculos_ordenadosfinal.append(*
            #    v + ";" + f"/passage/?search={v.split('.')[0] + '.' + v.split('.')[1] + '.' + v.split('.')[2]}&version=OJB" + ";" +
            #    v.split('.')[0] + ";" + mdlDicLivros.tradOjbSefaria[v.split('.')[0]] + ";" + v.split('.')[1] + ';' + v.split('.')[2])

        # print(versiculos_ordenadosfinal)
        # Imprime a lista de versículos ordenada
      #  print(dfcorder)
    # limpa csv

    for index, row in dfnorder.iterrows():

       pesquisar(row["Livro"] , row["Capitulo"] ,row["Versiculo"], nomearquivo1,row['url'])
def representaCategoriasNumericamente(categoria):
    if categoria == 'Targum':
        return 1
    elif categoria == 'Commentary':
        return 2
    elif categoria == 'Midrash':
        return 3
    elif categoria == 'Talmud':
        return 4
    elif categoria == 'Musar':
        return 5
    return 0

def assinatura(doc, nomearquivo,capitulon=''):
    # Adiciona uma nova página


    # Adiciona o título da capa
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if capitulon != '':
        run = title.add_run(nomearquivo + " \ncapitulo: " + capitulon)
    else:
        run = title.add_run(nomearquivo )
    run.bold = True
    run.font.size = Pt(24)
    document.add_page_break()

    # Adiciona o nome
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run("Levantamento de Renato Nati")
    run.font.size = Pt(12)
    doc.add_page_break()



def save_word_file(nomearquivo,capitulon=''):

    file_path=nomearquivo + ' cap ' + capitulon+'.docx'
    if os.path.isfile(file_path):
        os.remove(file_path)
        print(f"{file_path} has been removed.")
    else:
        print(f"{file_path} not found.")
    dfversoes = pd.read_csv(nomearquivo+'b.csv')
    df = pd.read_csv(nomearquivo+'.csv')
    print(df)
    print("@@" )
    print(dfversoes)
    document = Document()
    assinatura(document, nomearquivo,capitulon)
    #document.add_heading(nomearquivo, 1)
    for index1  , row1 in dfversoes.iterrows():

        print('124')
        # criando o documento word

        # adicionando a primeira linha com o livro, capítulo e versículo
        document.add_heading(f'{row1["bookAnc"]} {row1["chapterAnc"]}:{row1["verseAnc"]}',1)

        # adicionando a segunda linha com o conteúdo da coluna nvi

        document.add_paragraph('PTBR: '+ row1["textoPortugues"])
        # adicionando a terceira linha com o conteúdo da coluna texto1 em itálico
        #texto1_paragraph = document.add_paragraph("textoOjb"+row1["textoOJB"])
        texto1_paragraph2 = document.add_paragraph('Sefaria: ' + row1["textoSefaria"])
       # texto1_paragraph.italic = True
        texto1_paragraph2.italic = True

        df['categoria_numerica'] = df['category'].apply(representaCategoriasNumericamente)
        df.sort_values(by='categoria_numerica', inplace=True)


        for index, row in df.iterrows():

            #try:

            bookeq= row['bookAnc'].strip() == row1["bookAnc"].strip()
            chapeq= str(row['chapterAnc']).strip() == str(row1["chapterAnc"]).strip()
            verseq= str(row['verseAnc']).strip() == str(row1["verseAnc"]).strip()
            if bookeq and chapeq and verseq:
                response_text2 = ""
                if not pd.isna(row["text"]) and len(row["text"]) >100 and len(row)<800:

                    paragraph = document.add_paragraph(row["category"] + ": " + row["ref"]+ "\n")

                    #paragraph.add_run(row["category"] + ": " + row["ref"])

                    try:
                        prompt = "Explique este texto para uma pessoa mediana entender de forma resumida em poucas palavras capturando os pontos principais, e após isso demonstre como eu poderia usar isso para falar do evangelho, citando alguns versiculos que mostre correlação, e também forneça as palavras chaves que funcionem como marcadores do texto.  e em português: " + \
                             row["text"]
                        response = openai.Completion.create(
                            engine="text-davinci-003",
                            prompt=prompt,
                            max_tokens=len(row["text"]) +200,
                            n=1,
                            stop=None,
                            temperature=0.3,
                        )
                        print(response)
                        response_text = response['choices'][0]['text']
                        response_text = '\n__________________________________\nAnalise através de IA quanto a este Comentario acima: \n' + response_text + '\n_________________________________';
                        response_text2 = " "
                    except:
                        response_text2=""
                        print("erro na chamada a openai")


                    if index % 2 == 0:



                        try:

                            paragraph = formata(row["text"], paragraph, RGBColor(107, 94, 155),0.2)

                            if response_text2 != "" :

                                paragraph = formata(response_text, paragraph, RGBColor(123, 61, 0),0.3)

                            #document.add_heading(row["category"] + ": " + row["ref"], 3)

                        except:
                            paragraph = formata(row["text"], paragraph, RGBColor(30, 106, 57),0.2)
                    else:
                       # document.add_heading(row["category"] + ": " + row["ref"], 3)




                       try:


                           paragraph = formata(row["text"], paragraph, RGBColor(30, 106, 57),0.2)
                           #paragraph.add_run(row["text"])
                           if response_text2 != "":

                               paragraph = formata(response_text, paragraph, RGBColor(123, 61, 0), 0.3)
                           #paragraph.add_run(response_text)
                       except:
                            paragraph = formata(row["text"], paragraph, RGBColor(30, 106, 57),0.2)

                    # paragraph.runs[0].font.color.rgb = RGBColor(107, 94, 155)

            #except:
            #    print("erro")
    capitulon= ' '+capitulon
    document.save(nomearquivo + ' cap ' + capitulon+'.docx')
     #extrair_e_salvar(arquivofinal, dfnorder)
# extrair_e_salvar(arquivoordenado,dfcorder)
#referencia('Kefa I',[1,2,3,4,5], 'Kefa I')
for i in range(2, 5):
     referencia('Yaakov',[i], 'Yaakov')
     save_word_file('Yaakov',str(i))
#referencia('Yaakov',[1], 'Yaakov')
exit(0)
capComentario = "Sotah.9b.4"
category="Midrash"
url = 'https://www.sefaria.org/api/texts/' + capComentario + '?commentary=0&context=1&pad=0&wrapLinks=1&wrapNamedEntities=1&multiple=0&stripItags=0&transLangPref=&firstAvailableRef=1&fallbackOnDefaultVersion=1'
print(url)

response = requests.get(url)
data1 = response.json()
print(len(data1['text']))
# se tiver ingles
if (category == 'Targum'):
    text = data1['text'][int("Numbers 5:21".split(':')[1]) - 1]

elif category == 'Midrash' or category == "Talmud":
    for pos in range(0, len(data1['text'])):

        texto=data1['text'][pos]
        if ("Numbers 5:21" in texto):
            print(pos)
            text = texto

else:
    text = data1['text'][0]
print(text)

