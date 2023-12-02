import tkinter as tk
from tkinter import ttk

from Company_Info_Web_Scraper import *
from GetArticles import *
from Stock_Analyzer import *
import config
from docx import Document
from datetime import datetime
import yfinance as yf
import pandas as pd
import openai


# set openai api key
openai.api_key = config.API_KEY


def get_stock_data(symbol, start_date, end_date):
    ticker = yf.download(symbol, start_date, end_date)

    # save stock data to csv
    file = ticker.to_csv("Price_data\\"+symbol + '_Price_Data.csv')

    return file


def main():

    # Set default start and end dates
    default_start_date = "2020-01-01"
    default_end_date = datetime.today().strftime('%Y-%m-%d')

    def fetch_data():
        # Get user input
        ticker_symbol = ticker_entry.get()
        start_date = start_date_entry.get()
        end_date = end_date_entry.get()

        # download and save stock data
        info = get_stock_data(ticker_symbol, start_date, end_date)

        # Create a new document
        doc = Document()
        # doc.save(f"{ticker_symbol}_report.docx")

        # Add a heading of level 0 (largest heading)
        doc.add_heading(f'Stock Report for {ticker_symbol}', 0)

        # Get company info and stock data
        info = get_company_info(ticker_symbol)

        # Add company info to document
        doc.add_heading('Company Information:', 1)
        doc.add_paragraph(f"Name: {info['name']}")
        doc.add_paragraph(f"Sector: {info['sector']}")
        doc.add_paragraph(f"Industry: {info['industry']}")
        doc.add_paragraph(f"Summary: {info['summary']}")

        # read stock price data from csv
        filename = "Price_Data\\" + ticker_symbol + '_Price_Data.csv'
        df = pd.read_csv(filename)

        # get wiki info
        wiki_url = get_wiki_info(info['name'])

        # Add wiki url to document
        doc.add_heading('Wikipedia Information', 1)
        doc.add_paragraph(wiki_url)

        # Add analysis to document
        doc.add_heading('Stock Analysis', 1)
        # doc.add_paragraph(get_stock_indicators(
        #     ticker_symbol, start_date, end_date))
        doc.add_paragraph(analyze_stock(
            ticker_symbol, start_date, end_date))

        # plot stock data
        fig = plot_stock_with_moving_averages_from_csv(filename)
        fig.savefig(f'{ticker_symbol}_stock_plot.png')
        doc.add_picture(f'{ticker_symbol}_stock_plot.png')

        # get articles
        articles = get_MW_Articles(ticker_symbol, 1)

        # Add articles to document
        doc.add_heading('Articles', 1)
        for article in articles:
            doc.add_paragraph(article['title'])
            doc.add_paragraph(article['url'])
            doc.add_paragraph(summarize_article(article))

        # Save the document
        doc.save(
            f'C:\\Users\\Aidan\\Desktop\\USB\\Projects\\Python\\MoneyBots\\{ticker_symbol}_report.docx')

        # close the root window
        root.destroy()

    # Create root window
    root = tk.Tk()

    # Create frames
    input_frame = tk.Frame(root)
    input_frame.pack(side="top", fill="x")

    output_frame = tk.Frame(root)
    output_frame.pack(side="bottom", fill="both", expand=True)

    # User Input
    ticker_label = tk.Label(input_frame, text="Enter Ticker Symbol:")
    ticker_entry = tk.Entry(input_frame)
    start_date_label = tk.Label(input_frame, text="Start date:")
    start_date_entry = tk.Entry(input_frame)
    end_date_label = tk.Label(input_frame, text="End date:")
    end_date_entry = tk.Entry(input_frame)

    # Set default values for start and end date entries
    start_date_entry.insert(tk.END, default_start_date)
    end_date_entry.insert(tk.END, default_end_date)

    # Position widgets using grid
    ticker_label.grid(row=0, column=0)
    ticker_entry.grid(row=0, column=1)
    start_date_label.grid(row=1, column=0)
    start_date_entry.grid(row=1, column=1)
    end_date_label.grid(row=2, column=0)
    end_date_entry.grid(row=2, column=1)

    fetch_button = ttk.Button(
        input_frame, text="Get Stock Data", command=fetch_data)
    fetch_button.grid(row=3, column=0, columnspan=2)

    root.mainloop()


if __name__ == "__main__":

    try:
        main()

    except Exception as e:
        print(e)
