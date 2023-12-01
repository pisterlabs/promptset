import docx
import sys
import os
from openai import OpenAI

import tiktoken
import json
from docx.oxml import OxmlElement
import re
from docassemble.base.util import get_config

os.environ["OPENAI_API_KEY"] = get_config("openai api key")

from typing import List, Tuple, Optional, Union

__all__ = [
    "get_labeled_docx_runs",
    "update_docx",
    "modify_docx_with_openai_guesses",
]


def add_paragraph_after(paragraph, text):
    p = OxmlElement("w:p")
    p.text = text
    paragraph._element.addnext(p)


def add_paragraph_before(paragraph, text):
    p = OxmlElement("w:p")
    p.text = text
    paragraph._element.addprevious(p)


def update_docx(
    document: Union[docx.Document, str], modified_runs: List[Tuple[int, int, str, int]]
) -> docx.Document:
    """Update the document with the modified runs.

    Args:
        document: the docx.Document object, or the path to the DOCX file
        modified_runs: a tuple of paragraph number, run number, the modified text, a question (not used), and whether a new paragraph should be inserted (for conditional text)

    Returns:
        The modified document.
    """
    ## Sort modified_runs in reverse order so inserted paragraphs are in the correct order
    # modified_runs = sorted(modified_runs, key=lambda x: x[0], reverse=True)
    #
    ## also sort each run in the modified_runs so that the runs are in the correct order
    # modified_runs = sorted(modified_runs, key=lambda x: x[1], reverse=True)

    if isinstance(document, str):
        document = docx.Document(document)

    for paragraph_number, run_number, modified_text, new_paragraph in modified_runs:
        paragraph = document.paragraphs[paragraph_number]
        run = paragraph.runs[run_number]
        # if new_paragraph == 1:
        #    add_paragraph_after(paragraph, modified_text)
        # elif new_paragraph == -1:
        #    add_paragraph_before(paragraph, modified_text)
        # else:
        run.text = modified_text
    return document


def get_labeled_docx_runs(
    docx_path: str,
    custom_people_names: Optional[Tuple[str, str]] = None,
    openai_client: Optional[OpenAI] = None,
) -> List[Tuple[int, int, str, int]]:
    """Scan the DOCX and return a list of modified text with Jinja2 variable names inserted.

    Args:
        docx_path: path to the DOCX file
        custom_people_names: a tuple of custom names and descriptions to use in addition to the default ones. Like: ("clients", "the person benefiting from the form")

    Returns:
        A list of tuples, each containing a paragraph number, run number, and the modified text of the run.
    """
    if not openai_client:
        openai_client = OpenAI()

    role_description = """
    You will process a DOCX document and return a JSON structure that turns the DOCX file into a template 
    based on the following guidelines and examples. The DOCX will be provided as an annotated series of
    paragraphs and runs.

    Steps:
    1. Analyze the document. Identify placeholder text and repeated _____ that should be replaced with a variable name.
    2. Insert jinja2 tags around a new variable name that represents the placeholder text.
    3. Mark optional paragraphs with conditional Jinja2 tags.
    4. Text intended for verbatim output in the final document will remain unchanged.
    5. The result will be a JSON structure that indicates which paragraphs and runs in the DOCX require modifications,
    the new text of the modified run with Jinja2 inserted, and a draft question to provide a definition of the variable.

    Example input, with paragraph and run numbers indicated:
    [
        [0, 1, "Dear John Smith:"],
        [1, 0, "This sentence can stay as is in the output and will not be in the reply."],
        [2, 0, "[Optional: if you are a tenant, include this paragraph]"],
    ]

    Example reply, indicating paragraph, run, the new text, and a number indicating if this changes the 
    current paragraph, adds one before, or adds one after (-1, 0, 1):

    {
        "results": [
            [0, 1, "Dear {{ other_parties[0] }}:", 0],
            [2, 0, "{%p if is_tenant %}", -1],
            [3, 0, "{%p endif %}", "", 1],
        ]
    }

    The reply ONLY contains the runs that have modified text.
    """

    custom_name_text = ""
    if custom_people_names:
        assert isinstance(custom_people_names, list)
        for name, description in custom_people_names:
            custom_name_text += f"    {name} ({description}), \n"

    rules = f"""
    Rules for variable names:
        1. Variables usually refer to people or their attributes.
        2. People are stored in lists.
        3. We use Docassemble objects and conventions.
        4. Use variable names and patterns from the list below. Invent new variable names when it is appropriate.

    List names for people:
        {custom_people_names}
        users (for the person benefiting from the form, especially when for a pro se filer)
        other_parties (the opposing party in a lawsuit or transactional party)
        plaintiffs
        defendants
        petitioners
        respondents
        children
        spouses
        parents
        caregivers
        attorneys
        translators
        debt_collectors
        creditors
        witnesses
        guardians_ad_litem
        guardians
        decedents
        interested_parties

        Name Forms:
            users (full name of all users)
            users[0] (full name of first user)
            users[0].name.full() (Alternate full name of first user)
            users[0].name.first (First name only)
            users[0].name.middle (Middle name only)
            users[0].name.middle_initial() (First letter of middle name)
            users[0].name.last (Last name only)
            users[0].name.suffix (Suffix of user's name only)

    Attribute names (replace `users` with the appropriate list name):
        Demographic Data:
            users[0].birthdate (Birthdate)
            users[0].age_in_years() (Calculated age based on birthdate)
            users[0].gender (Gender)
            users[0].gender_female (User is female, for checkbox field)
            users[0].gender_male (User is male, for checkbox field)
            users[0].gender_other (User is not male or female, for checkbox field)
            users[0].gender_nonbinary (User identifies as nonbinary, for checkbox field)
            users[0].gender_undisclosed (User chose not to disclose gender, for checkbox field)
            users[0].gender_self_described (User chose to self-describe gender, for checkbox field)
            user_needs_interpreter (User needs an interpreter, for checkbox field)
            user_preferred_language (User's preferred language)

        Addresses:
            users[0].address.block() (Full address, on multiple lines)
            users[0].address.on_one_line() (Full address on one line)
            users[0].address.line_one() (Line one of the address, including unit or apartment number)
            users[0].address.line_two() (Line two of the address, usually city, state, and Zip/postal code)
            users[0].address.address (Street address)
            users[0].address.unit (Apartment, unit, or suite)
            users[0].address.city (City or town)
            users[0].address.state (State, province, or sub-locality)
            users[0].address.zip (Zip or postal code)
            users[0].address.county (County or parish)
            users[0].address.country (Country)

        Other Contact Information:
            users[0].phone_number (Phone number)
            users[0].mobile_number (A phone number explicitly labeled as the "mobile" number)
            users[0].phone_numbers() (A list of both mobile and other phone numbers)
            users[0].email (Email)

        Signatures:
            users[0].signature (Signature)
            signature_date (Date the form is completed)

        Information about Court and Court Processes:
            trial_court (Court's full name)
            trial_court.address.county (County where court is located)
            trial_court.division (Division of court)
            trial_court.department (Department of court)
            docket_number (Case or docket number)
            docket_numbers (A comma-separated list of docket numbers)
            
    When No Existing Variable Name Exists:
        1. Craft short, readable variable names in python snake_case.
        2. Represent people with lists, even if only one person.
        3. Use valid Python variable names within complete Jinja2 tags, like: {{ new_variable_name }}.

        Special endings:
            Suffix _date for date values.
            Suffix _value or _amount for currency values.

        Examples: 
        "(State the reason for eviction)" transforms into `{{ eviction_reason }}`.
    """
    encoding = tiktoken.encoding_for_model("gpt-4")

    doc = docx.Document(docx_path)

    items = []
    for pnum, para in enumerate(doc.paragraphs):
        for rnum, run in enumerate(para.runs):
            items.append([pnum, rnum, run.text])

    encoding = tiktoken.encoding_for_model("gpt-4")
    token_count = len(encoding.encode(role_description + rules + repr(items)))
    if token_count > 128000:
        raise Exception(
            f"Input to OpenAI is too long ({token_count} tokens). Maximum is 128000 tokens."
        )

    response = openai_client.chat.completions.create(
        model="gpt-4-1106-preview",
        messages=[
            {"role": "system", "content": role_description + rules},
            {"role": "user", "content": repr(items)},
        ],
        response_format={"type": "json_object"},
        temperature=0.5,
        max_tokens=4096,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0,
    )

    assert isinstance(response.choices[0].message.content, str)
    guesses = json.loads(response.choices[0].message.content)["results"]
    return guesses


def modify_docx_with_openai_guesses(docx_path: str) -> docx.Document:
    """Uses OpenAI to guess the variable names for a document and then modifies the document with the guesses.

    Args:
        docx_path (str): Path to the DOCX file to modify.

    Returns:
        docx.Document: The modified document, ready to be saved to the same or a new path
    """
    guesses = get_labeled_docx_runs(docx_path)

    return update_docx(docx.Document(docx_path), guesses)


if __name__ == "__main__":
    new_doc = modify_docx_with_openai_guesses(sys.argv[1])
    new_doc.save(sys.argv[1] + ".output.docx")