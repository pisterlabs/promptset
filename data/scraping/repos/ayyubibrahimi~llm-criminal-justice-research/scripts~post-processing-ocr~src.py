import aiohttp
import asyncio
import json
from docx import Document as PyDocx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document
import openai
import os

class Reformat:
    def __init__(self, api_key):
        self.api_key = api_key
        openai.api_key = api_key
        self.completion = openai.ChatCompletion()

    def extract_text_from_docx(self, docx_path):
        doc = Document(docx_path)
        return "\n".join([p.text for p in doc.paragraphs])

    def split_text_into_chunks(self, text, chunk_size=4090):
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=50,
            length_function=len,
        )
        return text_splitter.split_text(text)

    async def ask_gpt(self, session, question, text_chunk):
        chat_log = [
            {'role': 'system', 'content': 'You are an expert in court transcripts.'},
            {'role': 'user', 'content': question},
            {'role': 'system', 'content': text_chunk},
        ]

        data = {
            "model": "gpt-3.5-turbo",
            "messages": chat_log,
        }

        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
        }

        async with session.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            data=json.dumps(data),
        ) as response:
            response.raise_for_status()
            result = await response.json()
            return result["choices"][0]["message"]["content"]

    
    async def generate_and_save_response(self, docx_path, question):
        docx_text = self.extract_text_from_docx(docx_path)
        text_chunks = self.split_text_into_chunks(docx_text)

        file_name = os.path.splitext(os.path.basename(docx_path))[0]
        output_docx_path = f"pdf/{file_name}_output.docx"

        async with aiohttp.ClientSession() as session:
            tasks = [self.ask_gpt(session, question, text_chunk) for text_chunk in text_chunks]
            answers = await asyncio.gather(*tasks)

        final_answer = " ".join(answers)

        new_doc = PyDocx()
        new_doc.add_paragraph(final_answer)
        new_doc.save(output_docx_path)
