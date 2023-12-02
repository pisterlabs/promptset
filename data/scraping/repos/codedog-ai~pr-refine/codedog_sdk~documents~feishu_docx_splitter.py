import copy
import io
from typing import Optional

from docx import Document as load_document
from docx.document import Document as WordDocument
from docx.text.paragraph import Paragraph
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter, TextSplitter

from codedog_sdk.documents.base import Chapter, DocumentIndex

FONT_SIZE_TITLE = 330200
FONT_SIZE_H1 = 228600
FONT_SIZE_H2 = 203200
FONT_SIZE_H3 = 190500
FONT_SIZE_H4 = 177800
FONT_SIZE_H5 = 152400
FONT_SIZE_NORMAL = 139700


class FeishuDocxSplitter:
    def __init__(
        self,
        chapter_splitter: Optional[TextSplitter] = None,
    ):
        self.chapter_splitter = chapter_splitter or RecursiveCharacterTextSplitter()

    def generate_doci(
        self,
        doc_path_or_obj: Optional[str | io.BytesIO] = None,
        doc: Optional[WordDocument] = None,
    ):
        raw_chapters, documents = self.split_documents(doc_path_or_obj, doc)

        doci = DocumentIndex()
        last_level = 100
        last_chapter = None
        i = 0
        for raw_chapter, document in zip(raw_chapters, documents):
            if not raw_chapter:
                if not document.page_content:
                    continue
                else:
                    raw_chapter = [f"root {i}"]
            curr_chapter = Chapter(raw_chapter)
            document.metadata["chapter_key"] = curr_chapter.key

            if curr_chapter.level > last_level:
                parent_chapter = last_chapter
            else:
                tmp_parent = last_chapter
                tmp_counter = last_level - curr_chapter.level + 1
                while tmp_counter > 0 and tmp_parent is not None:
                    tmp_parent = tmp_parent.parent
                    tmp_counter -= 1
                parent_chapter = tmp_parent

            if parent_chapter:
                parent_chapter.add_child(curr_chapter)
                curr_chapter.parent = parent_chapter
                document.metadata["parent_chapter_key"] = parent_chapter.key

            last_chapter = curr_chapter
            last_level = curr_chapter.level
            doci.add_chapter(curr_chapter, doc=document, root=parent_chapter is None)
        return doci

    def split_documents(
        self,
        doc_path_or_obj: Optional[str | io.BytesIO] = None,
        doc: Optional[WordDocument] = None,
    ) -> tuple[list, list[Document]]:
        if doc_path_or_obj:
            doc = load_document(doc_path_or_obj)

        assert doc
        raw_chapters, documents = self._split_documents(doc)
        # chapters = self._merge_chapters(raw_chapters)
        return raw_chapters, documents

    def _split_documents(
        self, word_doc: WordDocument
    ) -> tuple[list[list[str]], list[Document]]:
        documents = []
        metadata = {}
        curr_header = []
        curr_texts = []
        chapters = []

        for paragraph in word_doc.paragraphs:
            if not paragraph.text:
                continue

            paragraph_type = self._get_paragraph_type(paragraph)

            match paragraph_type:
                case "title":
                    metadata["title"] = paragraph.text
                case "header1":
                    if curr_texts:
                        tmp_doc = Document(
                            page_content=self._build_content(curr_texts, curr_header)
                        )
                    else:
                        tmp_doc = Document(page_content="")
                    self._add_header_meta(tmp_doc, curr_header)
                    documents.append(tmp_doc)
                    curr_texts.clear()
                    chapters.append(curr_header)
                    curr_header = self._update_header(curr_header, paragraph.text, 1)
                case "header2":
                    if curr_texts:
                        tmp_doc = Document(
                            page_content=self._build_content(curr_texts, curr_header)
                        )
                    else:
                        tmp_doc = Document(page_content="")
                    self._add_header_meta(tmp_doc, curr_header)
                    documents.append(tmp_doc)
                    curr_texts.clear()
                    chapters.append(curr_header)
                    curr_header = self._update_header(curr_header, paragraph.text, 2)
                case "header3":
                    if curr_texts:
                        tmp_doc = Document(
                            page_content=self._build_content(curr_texts, curr_header)
                        )
                    else:
                        tmp_doc = Document(page_content="")
                    self._add_header_meta(tmp_doc, curr_header)
                    documents.append(tmp_doc)
                    curr_texts.clear()
                    chapters.append(curr_header)
                    curr_header = self._update_header(curr_header, paragraph.text, 3)
                case "header4":
                    if curr_texts:
                        tmp_doc = Document(
                            page_content=self._build_content(curr_texts, curr_header)
                        )
                    else:
                        tmp_doc = Document(page_content="")
                    self._add_header_meta(tmp_doc, curr_header)
                    documents.append(tmp_doc)
                    curr_texts.clear()
                    chapters.append(curr_header)
                    curr_header = self._update_header(curr_header, paragraph.text, 4)
                case "header5":
                    if curr_texts:
                        tmp_doc = Document(
                            page_content=self._build_content(curr_texts, curr_header)
                        )
                    else:
                        tmp_doc = Document(page_content="")

                    self._add_header_meta(tmp_doc, curr_header)
                    documents.append(tmp_doc)
                    curr_texts.clear()
                    chapters.append(curr_header)
                    curr_header = self._update_header(curr_header, paragraph.text, 5)
                case "text":
                    curr_texts.append(paragraph.text)
                case "table":
                    tbl_docs: list[Document] = self._parse_table_docs(paragraph)
                    for tbl_doc in tbl_docs:
                        self._add_header_meta(tbl_doc, curr_header)
                        tbl_doc.metadata["special"] = "table"
                    documents.append(tbl_docs)
                case "image":
                    continue
                case _:
                    continue

        if curr_texts:
            tmp_doc = Document(page_content="\n\n".join(curr_texts))
            chapters.append(curr_header)
            self._add_header_meta(tmp_doc, curr_header)
        if metadata:
            for doc in documents:
                doc.metadata.update(copy.copy(metadata))

        return chapters, documents

    def _get_paragraph_type(self, paragraph: Paragraph) -> str:
        if not paragraph.text:
            return ""
        fontsize = paragraph.runs[0].font.size
        fontbold = paragraph.runs[0].font.bold

        if not fontbold:
            if fontsize == FONT_SIZE_NORMAL:
                return "text"
            else:
                return ""
        if fontsize == FONT_SIZE_H1:
            return "header1"
        elif fontsize == FONT_SIZE_H2:
            return "header2"
        elif fontsize == FONT_SIZE_H3:
            return "header3"
        elif fontsize == FONT_SIZE_H4:
            return "header4"
        elif fontsize == FONT_SIZE_H5:
            return "header5"
        elif fontsize == FONT_SIZE_TITLE:
            return "title"
        return ""

    def _build_content(self, curr_texts: list[str], curr_header: list[str]) -> str:
        return (
            " - ".join(h for h in curr_header if h) + "\n\n" + "\n\n".join(curr_texts)
        )

    def _update_header(
        self, curr_header: list[str], header: str, level: int
    ) -> list[str]:
        curr_header = copy.copy(curr_header)
        if len(curr_header) < level - 1:
            curr_header += [""] * (level - 1 - len(curr_header))
        if len(curr_header) > level - 1:
            curr_header = curr_header[: level - 1]
        curr_header.append(header)
        return curr_header

    def _parse_table_docs(self, paragraph: Paragraph) -> list[Document]:
        raise NotImplementedError

    def _add_header_meta(self, doc: Document, curr_header: list[str]):
        for i, header in enumerate(curr_header):
            doc.metadata[f"header{i}"] = header


if __name__ == "__main__":
    doci = FeishuDocxSplitter().generate_doci("tmp/1.docx")

    # for chapter in doci.chapters:
    for chapter in doci.root_chapters:
        print("----")
        print(
            f"[chapter:{chapter.name}][parent:{chapter.parent.name if chapter.parent else 'root'}]{'/'.join(chapter.path)}"
        )
        print(f"children:{','.join(c.name for c in chapter.children)}")
        print("----")
