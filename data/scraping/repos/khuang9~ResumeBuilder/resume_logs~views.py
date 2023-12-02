from django.shortcuts import render, redirect

from django.contrib.auth.decorators import login_required
from django.http import Http404

from .models import Topic, QuestionHistory
from .forms import EntryForm, QuestionForm

from fpdf import FPDF


RESULT_FILE_PDF = 'resume.pdf'
RESULT_FILE_TXT = 'resume.txt'
RESULT_FILE_DOC = 'resume.docx'
KNOWLEDGE_FILE = 'ApplicationKnowledge.csv'
ai_suggested_text = ""
ai_suggested_texts = [""]

questions = [
    ["Is there a school you attended that you would like to add to your resume?", "In which years did you attend this school?", "What was your GPA?", "Any notable events or achievements at this school?", "Great! Is there another school you would like to add?"],
    ["What is, in your opinion, the most notable test/competition you have taken or participated in?", "What score did you get on this test?", "Can you think of any areas you were able to improve in through this experience?", "Thank you for the information! Any other tests/competitions you would like to talk about?"],
    ["Out of all the courses you have taken so far, which one do you think deserves to be discussed first?", "What mark did you get in this course?", "What do you think were the most important things that you learned from this course?", "How much did you enjoy taking this course?", "Good to know! Any other courses you'd like to mention?"],
    ["What is the first extracurricular activity you have participated in that comes to mind?", "Can you provide some extra details on the things you did while partaking in this extracurricular activity?", "What kinds of technical and transferable skills were you able to hone through this extracurricular?", "Any notable achievements related to this extracurricular?", "Thank you! Any more extracurricular activities worth mentioning?"],
  ]

questions_keyword = [
    ["School", "Year", "GPA", "achievements", "AnotherSchool"],
    ["test/competition name", "the score on this test", "areas to improve through this experience", "other tests/competitions"],
    ["first course", "mark", "important things learnt", "How much did you enjoy", "other courses to mention"],
    ["first extracurricular activity", "extra details", "technical and transferable skills able to hone", "notable achievements", "more extracurricular activities"],
  ]

prompts = [
    ["school name", "the graduation year", "GPA", "achievements", "school name"],
    ["test/competition name", "the score on this test", "areas to improve through this experience", "other tests/competitions"],
    ["first course", "mark", "important things learnt", "How much did you enjoy", "other courses to mention"],
    ["first extracurricular activity", "extra details", "technical and transferable skills able to hone", "notable achievements", "more extracurricular activities"],
  ]
  

from django.db import models
import openai


resume_personal_info = "Kevin Huang\nCell: 226-123-4567\n375 Hagey Blvd, Waterloo, ON N2L 6R5"
resume_summary = "Hard-working, motivated, and reliable student who excels in science and technology."


import os

os.environ['OPENAI_API_KEY'] = "to be replaced"

from langchain.indexes import VectorstoreIndexCreator

from langchain.document_loaders.csv_loader import CSVLoader
csvloader = CSVLoader(file_path="C:\Anaconda3\envs\WebDev\kprjs\media\measurements.txt", source_column="measurement type")


from .constants import OPENAIKEY
# Replace 'YOUR_API_KEY' with your actual API key or token
openai.api_key = OPENAIKEY

GPT_countdown = 10
def chatgpt_completions(prompt, text):
  global GPT_countdown
  GPT_countdown -= 1
  if GPT_countdown < 0:
    print("You've used all lives :))))))))))")
    return text
  prompt=f'{prompt} "{text}"'
  response = openai.Completion.create(
    engine="text-davinci-002",  # Use the GPT-3.5 engine
    prompt=prompt,
    max_tokens=100,  # Control the length of the generated text
    temperature=0.7  # Controls the randomness of the generated text
  )
  reply = response.choices[0].text.strip()
  
  print(f'========== prompt = {prompt}\nreply = {reply}')
  return reply  

def profile_agent_get_fieldinfo(input_info, output_info, delimiter = ':'):
  if 'but no ' not in output_info and 'is no ' not in output_info and 'are no ' not in output_info  and ((len(output_info) < 10) or (input_info != output_info)):
    return output_info.lstrip('-').rstrip()
  else:
    return "Unknown"


def profile_agent_process_answer_simple(question_key, prompt, text):
  information = chatgpt_completions(f'Please use the resume style to extract "{prompt}" from the text ', text)
  information = profile_agent_get_fieldinfo(text, information)
  text = f'{question_key}: {information}'
  
  return text


profile_chat_history = []

# Create views here.
def index(request):
  return render(request, 'resume_logs/index.html')

@login_required
def topics(request):
  topics = Topic.objects.filter(owner=request.user).order_by('id')
  context = {'topics' : topics}
  return render(request, 'resume_logs/topics.html', context)

def get_question(request, topic_id):
  topic = Topic.objects.get(id=topic_id)
  
  topics = Topic.objects.filter(owner=request.user).order_by('date_added')
  local_topic_id = topic_id - topics[0].id
      
  return questions[local_topic_id][topic.question_index]

def get_question_total(request, topic_id):
  topics = Topic.objects.filter(owner=request.user).order_by('date_added')
  local_topic_id = topic_id - topics[0].id
  
  return len(questions[local_topic_id])


def get_question_keyword(request, topic_id):
  topic = Topic.objects.get(id=topic_id)
  
  topics = Topic.objects.filter(owner=request.user).order_by('date_added')
  local_topic_id = topic_id - topics[0].id
  
  return questions_keyword[local_topic_id][topic.question_index]

def get_prompt(request, topic_id):
  topic = Topic.objects.get(id=topic_id)
  
  topics = Topic.objects.filter(owner=request.user).order_by('date_added')
  local_topic_id = topic_id - topics[0].id
  
  return prompts[local_topic_id][topic.question_index]

@login_required
def topic(request, topic_id):
  """Show a single topic and its chat history."""
  topic = Topic.objects.get(id=topic_id)
  
  if topic.owner != request.user:
    raise Http404

  entries = topic.entry_set.order_by('-id')
  
  incomplete_entries = get_question_total(request, topic_id)
  answered_questions = 0
  for entry in entries:
    print(f"entry.gpt={entry.gpt}")
    if entry.gpt != "" and "Unknown" not in entry.gpt:
      incomplete_entries -= 1
    if entry.text != "":
      answered_questions += 1
  
  prefix = ""
  if answered_questions == 0:
    prefix = f'Hello! Welcome to the "{topic.text}" section. let\'s get started. '
  print(f'================={prefix}')
  
  context = {
    'topic': topic,
	'form': EntryForm(),
	'entries': entries,
	'profile_chat_history' : profile_chat_history,
	'question': prefix + get_question(request, topic_id),
    'incomplete_entries' : incomplete_entries,
    'temp_info' : models.TextField(),
  }
  
  return render(request, 'resume_logs/topic.html', context)
  
@login_required
def generate_resume(request):
  context = {'items' : [], 'personal_info' : resume_personal_info, 'summary' : resume_summary}
  return render(request, 'resume_logs/generate_resume.html', context)

@login_required
def extract_entries(request):
  topics = Topic.objects.filter(owner=request.user).order_by('id')
  
  items = []
  for topic in topics:
    topic_text = topic.text + ':\n'
    entries = topic.entry_set.order_by('id')
    profile_text = ""
    for entry_number in range(len(entries)):
      print(f'================{entries[entry_number].gpt}')
      gpt_info = entries[entry_number].gpt.split(':')
      if len(gpt_info) < 2:
        continue
      elif len(gpt_info[1].strip()) == 0:
        continue
      elif 'N/A' in gpt_info[1]:
        continue
      elif 'nothing' in gpt_info[1]:
        continue
        
      profile_text += entries[entry_number].gpt + '\n'
      
    if profile_text != "":
      resume_text = chatgpt_completions(f"Please use the resume style to generate the {topic_text} section of the resume based on ", profile_text)
      topic.resume_text = resume_text
      print(f'{profile_text}')
      print(f'{resume_text}')
      items.append({'topic_text': topic_text, 'profile_text' : profile_text, 'resume_text' : resume_text})
    else:
      topic.resume_text = ""
  
      topic.save()        
      
  context = {'items' : items, 'personal_info' : resume_personal_info, 'summary' : resume_summary}
  return render(request, 'resume_logs/generate_resume.html', context)


@login_required
def new_answer(request, topic_id):
  """Add a new answer to a question for a particular topic."""
  global profile_chat_history
  initial_info = ""
  topic = Topic.objects.get(id=topic_id)
  
  topics = Topic.objects.filter(owner=request.user).order_by('date_added')
  local_topic_id = topic_id - topics[0].id
  
  if request.method != 'POST':
    # No data submitted; create a blank form.
    form = EntryForm(data=None)
  else:
    # POST data submitted; process data.
    action = request.POST['action']
    local_topic_id = topic_id - topics[0].id
    
    if action == 'Backward':
      topic.question_index = (topic.question_index + 1) % (len(questions[local_topic_id]))
      topic.save()
      return redirect('resume_logs:topic', topic_id=topic_id)
    elif action == 'Forward':
      topic.question_index = (topic.question_index - 1) % (len(questions[local_topic_id]))
      topic.save()
      return redirect('resume_logs:topic', topic_id=topic_id)
    
    
    topic = Topic.objects.get(id=topic_id)
    if topic.owner != request.user:
      raise Http404

    form = EntryForm(data=request.POST)
    if form.is_valid():
      topics = Topic.objects.filter(owner=request.user).order_by('date_added')
      local_topic_id = topic_id - topics[0].id
      
  
      entries = topic.entry_set.order_by('-id')
      new_entry = None
      keyword = get_question_keyword(request, topic_id)
      question = get_question(request, topic_id)
      
      for entry in entries:            
        if entry.keyword == keyword:
          new_entry = entry
          new_entry.text = entry.text
          break

      if new_entry == None:
        new_entry = form.save(commit=False)
        new_entry.topic = topic
        new_entry.section = question
        new_entry.keyword = keyword

      new_entry.text = form.cleaned_data["text"]
      print(f'-----------{new_entry.text}')
      new_entry.gpt = profile_agent_process_answer_simple(keyword, get_prompt(request, topic_id), new_entry.text)
      new_entry.save()
      
      topic.question_index = (topic.question_index + 1) % (len(questions[local_topic_id]))
      topic.save()
      
      profile_chat_history.append({'question' : question, 'answer' : new_entry.text})
      
      return redirect('resume_logs:topic', topic_id=topic_id)

  # Display a blank or invalid form.
  context = {'topic': topic, 'form': form, 'initial_info' : initial_info}
  return render(request, 'resume_logs/topic.html', context)
  
	
from django.http import HttpResponse
from django.conf import settings
import os


@login_required
def generate_txt(request):
  lines = []
  lines += resume_personal_info.split('\n')  
  lines += resume_summary.split('\n')  

  topics = Topic.objects.filter(owner=request.user).order_by('id')

  for topic in topics:
    lines.append(topic.text)
    lines += topic.resume_text.split('\n')  

  file_path = os.path.join(settings.MEDIA_ROOT, RESULT_FILE_TXT)
  with open(file_path, "w") as file:
    for line in lines:
      file.write(line + "\n")

  # Open the file for reading
  with open(file_path, 'rb') as file:
    response = HttpResponse(file.read(), content_type='application/octet-stream')
    response['Content-Disposition'] = 'attachment; filename=' + RESULT_FILE_TXT
    return response
    
def generate_pdf(request):
        
  pdf = FPDF()
  
  top_margin = 10   # Adjust as needed
  left_margin = 10  # Adjust as needed
  right_margin = 50 # Adjust as needed
  pdf.set_margins(left = 10.0, top = 10.9, right = 10.0)
  # Set auto page break mode and margins
  pdf.set_auto_page_break(auto=True, margin=top_margin)
  pdf.set_left_margin(left_margin)
  pdf.set_right_margin(right_margin)

  pdf.add_page()
  firstline = True
  pdf.set_font("Arial", style='B', size=12)
    
  # Personal information
  for line in resume_personal_info.split('\n'):
    pdf.cell(200,5, txt = line, ln=True, align='C')
    if firstline:
      firstline = False
      pdf.set_font("Arial", size=10)

  # Summary
  pdf.set_font("Arial", size=10)
  for line in resume_summary.split('\n'):
    pdf.multi_cell(200,10, txt = line, align='L')

  topics = Topic.objects.filter(owner=request.user).order_by('id')

  for topic in topics:
    pdf.set_font("Arial", style='B', size=12)
    pdf.multi_cell(200, 10, txt = topic.text, align='L')
    
    pdf.set_font("Arial", size=11)
    for line in topic.resume_text.split('\n'):
      pdf.multi_cell(200, 5, txt = line, align='L')

  file_path = os.path.join(settings.MEDIA_ROOT, RESULT_FILE_PDF)
  pdf.output(file_path)

  # Open the file for reading
  with open(file_path, 'rb') as file:
    response = HttpResponse(file.read(), content_type='application/octet-stream')
    response['Content-Disposition'] = 'attachment; filename=' + RESULT_FILE_PDF
    return response
    

def generate_docx(request):

  from docx import Document
  from docx.shared import Pt
  doc = Document()

  firstline = True
    
  # Personal information
  for line in resume_personal_info.split('\n'):
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = 1
    run = paragraph.runs[0]
    if firstline:
      firstline = False
      run.bold = True
      run.font.size = Pt(16)
    else:
      run.bold = False
      run.font.size = Pt(12)

  # Summary
  for line in resume_summary.split('\n'):
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = 0
    paragraph.runs[0].bold = False
          
  topics = Topic.objects.filter(owner=request.user).order_by('id')

  for topic in topics:
    paragraph = doc.add_paragraph(topic.text)
    paragraph.runs[0].font.size = Pt(14)
    
    for line in topic.resume_text.split('\n'):
      paragraph = doc.add_paragraph(line)
  
  # Save the document
  file_path = os.path.join(settings.MEDIA_ROOT, RESULT_FILE_DOC)
  doc.save(file_path)

  from docx2pdf import convert
  convert(file_path)
  print(f'{file_path}')
  # Open the file for reading
  with open(file_path, 'rb') as file:
    response = HttpResponse(file.read(), content_type='application/octet-stream')
    response['Content-Disposition'] = 'attachment; filename=' + RESULT_FILE_DOC
    return response
    
os.environ['OPENAI_API_KEY'] = OPENAIKEY


from langchain.document_loaders.csv_loader import CSVLoader
csvloader = CSVLoader(file_path=os.path.join(settings.MEDIA_ROOT, KNOWLEDGE_FILE), source_column="Program Name")


def qa_agent_answer(query):

  index = VectorstoreIndexCreator().from_loaders([csvloader])
  answer = index.query(query).strip()
  
  return answer
  
  
@login_required
def q_and_a(request):
  history = QuestionHistory.objects.filter(owner=request.user).order_by('id')
  reverse_history = QuestionHistory.objects.filter(owner=request.user).order_by('-id')
  
  context = {'history': history, 'reverse_history': reverse_history, 'form': QuestionForm()}
  return render(request, 'resume_logs/q_and_a.html', context)


@login_required
def send_question(request):
  history = QuestionHistory.objects.filter(owner=request.user).order_by('id')
  reverse_history = QuestionHistory.objects.filter(owner=request.user).order_by('-id')
  
  if request.method != 'POST':
    form = QuestionForm()
  else:
    form = QuestionForm(data=request.POST)
    if form.is_valid():
      
      new_question = form.save(commit=False)
      new_question.owner = request.user
      print(f"======== Q and A question: {new_question.question}")
      new_question.answer = qa_agent_answer(new_question.question)
      print(f"======== Q and A answer: {new_question.answer}")
      new_question.save()
      
      return redirect('resume_logs:q_and_a')
    
  context = {'history': history, 'reverse_history': reverse_history, 'form': form}
  return render(request, 'resume_logs/q_and_a.html', context)