import openai
from docx import Document
import os
import soundfile as sf
import numpy as np
import streamlit as st
import time
from pydub import AudioSegment
from moviepy.editor import *
import math
from openai import OpenAI
client = OpenAI()


openai.api_key = st.secrets["OPENAI_API_KEY"]

# Streamlit Configuration
st.set_page_config(
    page_title="Meeting Minutes",
    page_icon=":microphone:"
)

def get_state_variable(var_name, default_value):
    if 'st_state' not in st.session_state:
        st.session_state['st_state'] = {}
    if var_name not in st.session_state['st_state']:
        st.session_state['st_state'][var_name] = default_value
    return st.session_state['st_state'][var_name]

# Initialize session state for authentication
if 'is_authenticated' not in st.session_state:
    st.session_state.is_authenticated = False


def display_page():

    # Function to split audio
    def split_audio_sf(audio_path, target_chunk_size, max_file_size=25 * 1024 * 1024):
        # Read the audio data
        data, samplerate = sf.read(audio_path)
        total_samples = len(data)

        while True:
            chunk_samples = int(target_chunk_size * samplerate)

            # Check if the chunk size produces files larger than max_file_size
            num_chunks = math.ceil(total_samples / chunk_samples)
            print(f"Splitting audio into {num_chunks} chunks of {chunk_samples} samples each")
            estimated_size_per_chunk = os.path.getsize(audio_path) / num_chunks
            print(f"Estimated size per chunk: {estimated_size_per_chunk} bytes")

            if estimated_size_per_chunk > max_file_size:
                # If chunks are too large, reduce the target chunk size and recalculate
                target_chunk_size -= 10  # Reduce by 10 seconds and retry
                if target_chunk_size <= 0:
                    raise ValueError("Cannot split audio into chunks smaller than the max file size.")
                continue  # If still too large, loop back and try with reduced chunk size

            # If the estimated size per chunk is within the limit, proceed to create chunks
            chunks = [(data[start:start + chunk_samples], samplerate) for start in range(0, total_samples, chunk_samples)]
            break  # If chunk sizes are okay, break out of the loop and return chunks

        return chunks

    def transcribe_chunks_sf(chunks):
        transcriptions = []
        
        for i, (data, samplerate) in enumerate(chunks):
            temp_path = f'temp_chunk_{i}.wav'
            sf.write(temp_path, data, samplerate)

            # Check file size
            file_size = os.path.getsize(temp_path)
            if file_size > 26214400:  # 25 MB in bytes
                print(f"Skipping chunk {i} due to large size: {file_size} bytes")
                os.remove(temp_path)
                continue

            transcription = transcribe_audio(temp_path)
            transcriptions.append(transcription)
            os.remove(temp_path)
            
        return ' '.join(transcriptions)


    def transcribe_audio(audio_file_path):
        with open(audio_file_path, 'rb') as audio_file:
            transcription = openai.audio.transcriptions.create(
                model="whisper-1", 
                file=audio_file,
                language="en"  # Set the language to English
            )
        return transcription.text  # Here, we use .text instead of ['text']



    def meeting_minutes(transcription, additional_context):
        word_count = len(transcription.split())
        # Calculate sleep duration: 4 seconds for every 1000 words
        sleep_duration = (word_count / 1000) * .5

        # Ensure minimum sleep time (e.g., 2 seconds) if transcription is very short
        sleep_duration = max(sleep_duration, 1)

        # Sleep dynamically based on the number of words in the transcription
        time.sleep(sleep_duration)

        abstract_summary = abstract_summary_extraction(transcription, additional_context)
        time.sleep(sleep_duration)  # Repeat sleep after each processing step

        key_points = key_points_extraction(transcription, additional_context)
        time.sleep(sleep_duration)

        action_items = action_item_extraction(transcription, additional_context)
        time.sleep(sleep_duration)

        sentiment = sentiment_analysis(transcription, additional_context)
        # no need to sleep after the last step since the function is about to return

        return {
            'abstract_summary': abstract_summary,
            'key_points': key_points,
            'action_items': action_items,
            'sentiment': sentiment
        }
        
    def abstract_summary_extraction(transcription, additional_context):
        prompt = f"{additional_context}\nTranscription:\n\"{transcription}\""
        response = openai.chat.completions.create(
            model="gpt-4-1106-preview",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_tokens=2000,
        )
        # Access the content directly from the response object's attributes
        summary_content = response.choices[0].message.content
        #st.write(prompt)
        return summary_content


    def key_points_extraction(transcription, additional_context):
        prompt = f"{additional_context}\nTranscription:\n\"{transcription}\""
        response = openai.chat.completions.create(
            model="gpt-4-1106-preview",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_tokens=2000,
        )
        # Access the content directly from the response object's attributes
        summary_content = response.choices[0].message.content
        return summary_content


    def action_item_extraction(transcription, additional_context):
        prompt = f"{additional_context}\nTranscription:\n\"{transcription}\""
        response = openai.chat.completions.create(
            model="gpt-4-1106-preview",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_tokens=2000,
        )
        # Access the content directly from the response object's attributes
        summary_content = response.choices[0].message.content
        return summary_content

    def sentiment_analysis(transcription, additional_context):
        prompt = f"{additional_context}\nTranscription:\n\"{transcription}\""
        response = openai.chat.completions.create(
            model="gpt-4-1106-preview",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_tokens=2000,
        )
        # Access the content directly from the response object's attributes
        summary_content = response.choices[0].message.content
        #st.write(prompt)
        return summary_content
    
    def save_as_docx(minutes, filename):
        doc = Document()
        for key, value in minutes.items():
            # Replace underscores with spaces and capitalize each word for the heading
            heading = ' '.join(word.capitalize() for word in key.split('_'))
            doc.add_heading(heading, level=1)
            doc.add_paragraph(value)
            # Add a line break between sections
            doc.add_paragraph()
        doc.save(filename)

    # Parameters
    chunk_length = 2 * 60  # 3 minutes in seconds

    def main():
        # Initialize Streamlit interface
        st.title('Meeting Transcription and Minutes Generation')

        # File uploader allows user to add their own audio
        uploaded_audio = st.file_uploader("Upload audio", type=["wav", "mp3", "m4a"])

        if uploaded_audio is not None:
            # Button to trigger transcription and processing
            # Text area for additional context or notes
            if 'additional_context' not in st.session_state:
                st.session_state.additional_context = ""

            st.session_state.additional_context = st.text_area("Provide additional context or notes for the AI:", value=st.session_state.additional_context)
            if st.button('Transcribe and Analyze'):
                # Process the audio
                with st.spinner('Processing audio...'):
                    audio_path = save_uploaded_file(uploaded_audio)

                    # Split audio
                    with st.spinner('Splitting audio into chunks...'):
                        audio_chunks = split_audio_sf(audio_path, chunk_length)

                    # Transcribe each chunk and concatenate
                    with st.spinner('Transcribing audio... this might take a while'):
                        full_transcription = transcribe_chunks_sf(audio_chunks)

                    # Delete temp file
                    if os.path.exists(audio_path):
                        os.remove(audio_path)

                    # Display transcription
                    st.subheader("Transcription")
                    st.text_area("Full Transcription", full_transcription, height=300)

                    # Generate meeting minutes
                    with st.spinner('Generating meeting minutes... this might take a while'):
                        minutes = meeting_minutes(full_transcription, st.session_state.additional_context)

                        # Display meeting minutes details
                        st.subheader("Meeting Minutes")
                        for key, value in minutes.items():
                            st.markdown(f"#### {key.replace('_', ' ').title()}")
                            st.write(value)

                        # Save as DOCX
                        docx_file_path = 'meeting_minutes.docx'
                        save_as_docx(minutes, docx_file_path)
                        # Provide download link
                        with open(docx_file_path, "rb") as file:
                            btn = st.download_button(
                                label="Download Meeting Minutes as DOCX",
                                data=file,
                                file_name=docx_file_path,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )

    def convert_to_mp3(audio_path, target_path, bitrate="128k"):
        try:
            audio_clip = AudioFileClip(audio_path)
            audio_clip.write_audiofile(target_path, bitrate=bitrate)
            audio_clip.close()  # It's good practice to close the clip when done
            return True
        except Exception as e:
            print(f"Failed to convert {audio_path} to MP3: {e}")
            return False

    def save_uploaded_file(uploaded_file):
        try:
            file_extension = os.path.splitext(uploaded_file.name.lower())[1]
            base_name = os.path.splitext(uploaded_file.name)[0]
            save_folder = 'temp_files'

            if not os.path.isdir(save_folder):
                os.mkdir(save_folder)

            temp_path = os.path.join(save_folder, uploaded_file.name)
            target_path = os.path.join(save_folder, f"{base_name}.mp3")

            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
                print(f"File saved temporarily at {temp_path}")

            # If the uploaded file is .m4a, convert it to .mp3
            if file_extension == '.m4a':
                print(f"Converting {temp_path} to {target_path}")
                # Verify tempfile exists before attempting conversion
                if os.path.exists(temp_path):
                    convert_to_mp3(temp_path, target_path)
                    # Verify conversion was successful
                    if os.path.exists(target_path):
                        print(f"Conversion successful, file saved at {target_path}")
                        os.remove(temp_path)  # Remove the .m4a file after conversion
                        return target_path
                    else:
                        print(f"Conversion failed, file at {target_path} not found")
                        return None
                else:
                    print(f"Temp file at {temp_path} does not exist, cannot convert")
                    return None
            elif file_extension in ['.mp3', '.wav']:
                return temp_path

        except Exception as e:
            st.error(f"Error handling file: {e}")
            return None

    # start the Streamlit app
    if __name__ == "__main__":
        main()







# Logic for password checking
def check_password():
    if not st.session_state.is_authenticated:
        password = st.text_input("Enter Password:", type="password")


            
        
        if password == st.secrets["db_password"]:
            st.session_state.is_authenticated = True
            st.rerun()
        elif password:
            st.write("Please enter the correct password to proceed.")
            
        blank, col_img, col_title = st.columns([2, 1, 3])

        # Upload the image to the left-most column
        with col_img:
            st.image("https://s3-eu-west-1.amazonaws.com/tpd/logos/5a95521f54e2c70001f926b8/0x0.png")


        # Determine the page selection using the selectbox in the right column
        with col_title:
            #st.title("Created By Halo")
            st.write("")
            st.markdown('<div style="text-align: left; font-size: 40px; font-weight: normal;">Created By Halo*</div>', unsafe_allow_html=True)
            
        blank2, col_img2, col_title2 = st.columns([2, 1, 3])

        # Upload the image to the left-most column
        with col_img2:
            st.image("https://th.bing.com/th/id/OIP.42nU_MRx_INTLq_ejdHxBQHaCe?pid=ImgDet&rs=1")


        # Determine the page selection using the selectbox in the right column
        with col_title2:
            
            #st.title("Powered By IRS")
            st.markdown('<div style="text-align: left; font-size: 30px; font-weight: normal;">Powered By IRS</div>', unsafe_allow_html=True)
        # Fill up space to push the text to the bottom
        for _ in range(20):  # Adjust the range as needed
            st.write("")

        # Write your text at the bottom left corner
        st.markdown('<div style="text-align: right; font-size: 10px; font-weight: normal;">* Trenton Dambrowitz, Special Projects Manager, is "Halo" in this case.</div>', unsafe_allow_html=True)



    else:
        print("Access granted, welcome to the app.")
        display_page()


check_password()
