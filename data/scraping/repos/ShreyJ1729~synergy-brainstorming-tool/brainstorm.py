import openai
from typing import Dict, List, Tuple
import os
from docx import Document

# Set OpenAI API key
openai.api_key = os.environ.get('OPENAI_API_KEY')

def guess_curr_thread(prev_transcription, curr_transcription, conversations):
    """
    Guesses the current thread of conversation
    """
    # Prepare the list of conversation threads
    thread_list = list(conversations.keys())

    # Prepare the system message
    thread_list = '\n'.join(thread_list)
    system_message = f"You are a proficient AI with a specialty in understanding and following conversation threads.\nThe following are the threads that you have identified in the previous conversation:\n{thread_list}\n\nThe partial transcription of what conversation partner last said was '{prev_transcription}'.\n\nYour conversation partner just said the following:\n'{curr_transcription}'"

    # Use OpenAI API to predict the current thread
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": system_message
            },
            {
                "role": "user",
                "content": "We need what thread you think we're in now. Please enter a number from the list above."
            }
        ]
    )

    # Return the predicted thread
    return response['choices'][0]['message']['content']

def initialize_conversation_structure(thread_list: List[str]) -> Dict[str, Tuple[List[str], str]]:
    """
    Initialize a dictionary where the keys are the values in the list that you made and the value is a tuple where the first element is a list of strings and the second element is a string.
    """
    return {thread: ([], '') for thread in thread_list}


def transcribe_audio(audio_file_path):
    """
    Transcribes audio file using OpenAI API
    """
    with open(audio_file_path, 'rb') as audio_file:
        transcription = openai.Audio.transcribe("whisper-1", audio_file)

    return transcription['text']


def brainstorm(transcription):
    """
    Extracts meeting minutes from transcription
    """
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_items_extraction(transcription)
    sentiment = sentiment_analysis(transcription)

    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }


def abstract_summary_extraction(transcription):
    """
    Extracts abstract summary from transcription
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )

    return response['choices'][0]['message']['content']


def key_points_extraction(transcription):
    """
    Extracts key points from transcription
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )

    return response['choices'][0]['message']['content']


def action_items_extraction(transcription):
    """
    Extracts action items from transcription
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )

    return response['choices'][0]['message']['content']


def sentiment_analysis(transcription):
    """
    Extracts sentiment from transcription
    """
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting sentiment. Please review the text and identify the overall sentiment of the conversation. This could be positive, negative, or neutral. Please provide a brief explanation of why you chose this sentiment."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def save_as_markdown(minutes, filename):
    with open(filename, 'w') as f:
        for key, value in minutes.items():
            # Replace underscores with spaces and capitalize each word for the heading
            heading = ' '.join(word.capitalize() for word in key.split('_'))
            # Write the heading and the paragraph to the file
            f.write(f'# {heading}\n\n{value}\n\n')


def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)


def thread_splitter(input_text):
    """
    Splits the input text into threads for brainstorming
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in organizing thoughts into distinct conversation threads. Based on the following response from the conversation, identify and list the main threads that could be discussed or brought up later. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. This list is made so that another AI will organize the coming conversation into the categories that you've defined for them. Your goal is to make sure you have the most important threads listed so that the other AI can place certain pieces of the coming conversation in one thread over another. The following is an example input and output. Aftwerwards, you will be given a new input and you will be asked to produce a similar output.\n\nInput:\nSo I saw a Twitter demo of a cool realtime transcription software where the LLM talks back right after. I also saw a startup using this type of thing for interactive forms, but I thought it was wack. I didn't like it because I just wanted to finish the form, and it was making me take longer than I'd like. It'd be cool to do this in edtech for young kids. I also learned about how these researchers in Spain at around 2003 found the CRISPR mechanism in bacteria in salt mines. I wanted to think through the lightbulb moment and think through that in a discovery fiction sort of way. So those are the things I want to talk to you about today.\n\nExpected output:\n1. Startup ideas for realtime transcription\n2. CRISPR discovery fiction"
            },
            {
                "role": "user",
                "content": input_text
            }
        ]
    )

    return response['choices'][0]['message']['content']

def update_conversations(curr_thread, curr_transcription, conversations):
    for key in conversations.keys():
        if curr_thread in key:
            conversations[key][0].append(curr_transcription)
            return conversations

def check_switching(curr_transcription):
    """
    Checks if the user wants to switch threads
    """
    # Use OpenAI API to predict if the user wants to switch threads
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": f"You are a proficient AI with a specialty in understanding and following conversation threads. Your conversation partner just said the following:\n'{curr_transcription}'"
            },
            {
                "role": "user",
                "content": "We need to know if you think we should switch threads based on what was just said. Please return 1 if we should switch threads and 0 if we should not."
            }
        ]
    )

    # Return the predicted decision
    return int(response['choices'][0]['message']['content'])

# Initial prompt
print('What do you want to brainstorm about today?')

input1 = "So I saw a Twitter demo of a cool realtime transcription software where the LLM talks back right after. I also saw a startup using this type of thing for interactive forms, but I thought it was wack. I didn't like it because I just wanted to finish the form, and it was making me take longer than I'd like. It'd be cool to do this in edtech for young kids. I also learned about how these researchers in Spain at around 2003 found the CRISPR mechanism in bacteria in salt mines. I wanted to think through the lightbulb moment and think through that in a discovery fiction sort of way. So those are the things I want to talk to you about today."

print("These are the threads that I'm splitting this into for this session")

ans1 = thread_splitter(input1)
thread_list = ans1.split('\n')  # Split the string into a list
conversations = initialize_conversation_structure(thread_list)  # Initialize the conversation structure

print(ans1)
print('Okay nice! What do you want to talk about first?')

input2 = "I want to talk about the transcription one first"

prev_thread = None
curr_thread = None
prev_transcription = input1

while True:
    # Base case where you just started the conversation
    if prev_thread is None:
        curr_transcription = input2
        curr_thread = guess_curr_thread(prev_transcription, curr_transcription,  conversations)
        # Update the conversations dictionary
        conversations = update_conversations(curr_thread, curr_transcription, conversations)
        prev_thread = curr_thread
        continue
        
    # Get the transcription
    curr_transcription = "I thought the application of the realtime transcription to the startup was wack"
    curr_thread = guess_curr_thread(prev_transcription, curr_transcription,  conversations)
    if curr_thread != prev_thread:
        print(f"I thought that you were trying to talk about {prev_thread} but now it seems we're talking about {curr_thread}, should we switch to {curr_thread}?")
        curr_transcription = "Yes please go to that thread instead"
        print(curr_transcription)
        should_switch = check_switching(curr_transcription)
        if should_switch == 1:
            conversations = update_conversations(curr_thread, curr_transcription, conversations)
            prev_thread = curr_thread
        else:
            conversations = update_conversations(prev_thread, curr_transcription, conversations)
        
# Test the brainstorm function.
# minutes = brainstorm(input1)

# save_as_docx(minutes, 'brainstorm.docx')
# save_as_markdown(minutes, 'brainstorm.md')

